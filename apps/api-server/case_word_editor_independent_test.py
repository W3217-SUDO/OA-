"""Independent runtime checks for the case Word editor's atomicity guarantees."""

from __future__ import annotations

import asyncio
import importlib
import tempfile
import unittest
import zipfile
from io import BytesIO
from pathlib import Path

import httpx
from docx import Document
from docx.oxml import OxmlElement
from sqlalchemy import event
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

from app.config import settings
from app.database import Base, get_db
from app.main import app
from app.models import BusinessRecord, FileAttachment, User
from app.security import current_identity


main_module = importlib.import_module("app.main")
IDENTITY = {"username": "CODEX-WORD-independent", "display_name": "独立验证人", "role": "admin", "role_ids": ["admin"], "department": "验证部"}


class IndependentWordEditorTest(unittest.IsolatedAsyncioTestCase):
    async def asyncSetUp(self):
        self.temp = tempfile.TemporaryDirectory(prefix="codex-word-independent-")
        self.old_root = main_module.UPLOAD_ROOT
        main_module.UPLOAD_ROOT = Path(self.temp.name)
        database = Path(self.temp.name) / "word-editor.db"
        self.engine = create_async_engine(f"sqlite+aiosqlite:///{database}")
        self.sessions = async_sessionmaker(self.engine, expire_on_commit=False, class_=AsyncSession)
        async with self.engine.begin() as connection:
            await connection.run_sync(Base.metadata.create_all)
        document = Document()
        document.add_paragraph("可编辑正文")
        protected = document.add_paragraph("受保护 XML")
        field = OxmlElement("w:fldSimple")
        field.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr", "DATE")
        protected._p.append(field)
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "合并单元格"
        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(1, 0).text = "普通表格"
        table.cell(1, 1).text = "右侧"
        content = BytesIO(); document.save(content)
        self.original = Path(self.temp.name) / "original.docx"
        self.original.write_bytes(content.getvalue())
        async with self.sessions() as db:
            record = BusinessRecord(module="case", serial_no="CODEX-WORD-INDEPENDENT", title="独立 Word 验证", customer="", status="文书准备", owner=IDENTITY["username"], department=IDENTITY["department"], data={})
            db.add_all([record, User(username=IDENTITY["username"], display_name=IDENTITY["display_name"], department=IDENTITY["department"], role="admin", role_ids=["admin"], password_hash="test")])
            await db.flush()
            attachment = FileAttachment(record_id=record.id, category="法院诉讼文书", original_name="independent.docx", stored_name=self.original.name, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", size=self.original.stat().st_size, path=str(self.original), uploader=IDENTITY["username"], remark="")
            db.add(attachment); await db.commit()
            self.case_id, self.attachment_id = record.id, attachment.id

        async def override_db():
            async with self.sessions() as db:
                yield db
        self.saved_overrides = dict(app.dependency_overrides)
        app.dependency_overrides[get_db] = override_db
        app.dependency_overrides[current_identity] = lambda: IDENTITY
        self.client = httpx.AsyncClient(transport=httpx.ASGITransport(app=app), base_url="http://independent.test")

    async def asyncTearDown(self):
        await self.client.aclose()
        app.dependency_overrides.clear(); app.dependency_overrides.update(self.saved_overrides)
        main_module.UPLOAD_ROOT = self.old_root
        await self.engine.dispose(); self.temp.cleanup()

    def url(self, suffix: str) -> str:
        return f"{settings.api_prefix}/cases/{self.case_id}/attachments/{self.attachment_id}/word-editor{suffix}"

    async def open_editor(self) -> dict:
        response = await self.client.get(self.url("/content"))
        self.assertEqual(response.status_code, 200, response.text)
        return response.json()

    async def test_complex_xml_and_merged_cells_are_preserved(self):
        opened = await self.open_editor()
        merged = [block for block in opened["blocks"] if block["text"] == "合并单元格"]
        self.assertEqual(len(merged), 1)
        protected = next(block for block in opened["blocks"] if block["text"] == "受保护 XML")
        self.assertFalse(protected["editable"])
        blocks = opened["blocks"]
        next(block for block in blocks if block["text"] == "可编辑正文")["text"] = "保存后的正文"
        saved = await self.client.put(self.url("/content"), json={"lock_token": opened["lock_token"], "version": opened["version"], "blocks": blocks})
        self.assertEqual(saved.status_code, 200, saved.text)
        async with self.sessions() as db:
            item = await db.get(FileAttachment, self.attachment_id)
            final_path = Path(item.path)
        self.assertEqual(Document(final_path).paragraphs[1].text, "受保护 XML")
        with zipfile.ZipFile(final_path) as archive:
            self.assertIn(b"fldSimple", archive.read("word/document.xml"))

    async def test_simultaneous_saves_allow_exactly_one_winner(self):
        opened = await self.open_editor()
        blocks = opened["blocks"]
        next(block for block in blocks if block["text"] == "可编辑正文")["text"] = "并发保存"
        body = {"lock_token": opened["lock_token"], "version": opened["version"], "blocks": blocks}
        first, second = await asyncio.gather(
            self.client.put(self.url("/content"), json=body),
            self.client.put(self.url("/content"), json=body),
        )
        self.assertEqual(sorted([first.status_code, second.status_code]), [200, 409], (first.text, second.text))

    async def test_commit_failure_keeps_original_pointer_and_file(self):
        opened = await self.open_editor()
        blocks = opened["blocks"]
        next(block for block in blocks if block["text"] == "可编辑正文")["text"] = "不应提交"
        fail_commit = {"armed": True}

        def reject_save_commit(_connection):
            if fail_commit["armed"]:
                fail_commit["armed"] = False
                raise RuntimeError("forced commit failure")

        event.listen(self.engine.sync_engine, "commit", reject_save_commit)
        try:
            # ASGITransport deliberately re-raises unhandled server errors;
            # that is stronger evidence than a synthesized HTTP 500 response.
            with self.assertRaisesRegex(RuntimeError, "forced commit failure"):
                await self.client.put(self.url("/content"), json={"lock_token": opened["lock_token"], "version": opened["version"], "blocks": blocks})
        finally:
            event.remove(self.engine.sync_engine, "commit", reject_save_commit)
        async with self.sessions() as db:
            item = await db.get(FileAttachment, self.attachment_id)
            self.assertEqual(Path(item.path), self.original)
        self.assertTrue(self.original.is_file())
        self.assertEqual(Document(self.original).paragraphs[0].text, "可编辑正文")
        self.assertEqual(list(Path(self.temp.name).glob("*.docx")), [self.original])


if __name__ == "__main__":
    unittest.main()
