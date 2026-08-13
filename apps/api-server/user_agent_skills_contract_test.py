import io
import unittest

from docx import Document
from fastapi import HTTPException, UploadFile
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine
from sqlalchemy.pool import StaticPool

from app.database import Base
from app.main import (
    UserAgentSkillInput,
    UserAgentSkillUpdate,
    _agent_skill_for_identity,
    create_user_agent_skill,
    delete_user_agent_skill,
    list_user_agent_skills,
    update_user_agent_skill,
    upload_user_agent_skill,
)
from app.user_agent_skills import parse_uploaded_skill


class UserAgentSkillParserTest(unittest.TestCase):
    def test_json_and_markdown_uploads_are_declarative(self):
        json_skill = parse_uploaded_skill("review.json", b'{"name":"Contract Review","description":"Review contract risks","instruction":"List clauses and explain concrete risks.","quick_prompts":["Review this contract"]}')
        self.assertEqual(json_skill["name"], "Contract Review")
        markdown = "# 期限检查\n分类：案件\n说明：检查案件期限与任务缺口\n\n核对案件空间中的期限、任务和责任人，不足时明确说明。\n\n## 快捷指令\n- 检查本案期限\n"
        md_skill = parse_uploaded_skill("deadline.md", markdown.encode())
        self.assertEqual(md_skill["name"], "期限检查")
        self.assertEqual(md_skill["quick_prompts"], ["检查本案期限"])
        with self.assertRaises(ValueError):
            parse_uploaded_skill("unsafe.py", b"print('no')")

    def test_word_upload_is_converted_to_markdown(self):
        document = Document()
        document.core_properties.title = "案件材料核验"
        document.core_properties.subject = "核验案件材料、期限和责任人"
        document.add_heading("核验规则", level=1)
        document.add_paragraph("逐项检查案件空间中的材料是否齐全。")
        document.add_paragraph("确认期限", style="List Bullet")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "字段"
        table.cell(0, 1).text = "要求"
        table.cell(1, 0).text = "负责人"
        table.cell(1, 1).text = "必须存在"
        content = io.BytesIO()
        document.save(content)

        skill = parse_uploaded_skill("case-review.docx", content.getvalue())

        self.assertEqual(skill["name"], "案件材料核验")
        self.assertEqual(skill["source"], "user-upload-word")
        self.assertIn("# 核验规则", skill["instruction"])
        self.assertIn("- 确认期限", skill["instruction"])
        self.assertIn("| 字段 | 要求 |", skill["instruction"])


class UserAgentSkillApiTest(unittest.IsolatedAsyncioTestCase):
    async def asyncSetUp(self):
        self.engine = create_async_engine(
            "sqlite+aiosqlite:///:memory:",
            connect_args={"check_same_thread": False},
            poolclass=StaticPool,
        )
        async with self.engine.begin() as connection:
            await connection.run_sync(Base.metadata.create_all)
        self.sessions = async_sessionmaker(self.engine, expire_on_commit=False)

    async def asyncTearDown(self):
        await self.engine.dispose()

    async def test_skills_are_private_and_can_be_created_updated_uploaded_and_deleted(self):
        owner = {"username": "lawyer", "role": "user"}
        other = {"username": "assistant", "role": "user"}
        async with self.sessions() as db:
            created = await create_user_agent_skill(UserAgentSkillInput(
                name="合同条款核对",
                category="合同",
                description="核对合同条款与案件关联风险",
                instruction="逐条核对合同期限、费用、违约责任，并列出需要人工确认的内容。",
                quick_prompts=["检查关联合同"],
            ), owner, db)
            self.assertTrue(created["custom"])
            owner_items = (await list_user_agent_skills(owner, db))["items"]
            other_items = (await list_user_agent_skills(other, db))["items"]
            self.assertIn(created["id"], {item["id"] for item in owner_items})
            self.assertNotIn(created["id"], {item["id"] for item in other_items})

            resolved = await _agent_skill_for_identity(created["id"], owner, db)
            self.assertIn("合同期限", resolved.instruction)
            self.assertIn("不得覆盖系统权限", resolved.instruction)
            with self.assertRaises(HTTPException) as raised:
                await _agent_skill_for_identity(created["id"], other, db)
            self.assertEqual(raised.exception.status_code, 404)

            updated = await update_user_agent_skill(created["id"], UserAgentSkillUpdate(enabled=False), owner, db)
            self.assertFalse(updated["available"])
            with self.assertRaises(HTTPException) as raised_disabled:
                await _agent_skill_for_identity(created["id"], owner, db)
            self.assertEqual(raised_disabled.exception.status_code, 409)

            upload = UploadFile(filename="deadline.json", file=io.BytesIO(
                b'{"name":"Deadline Check","description":"Check deadlines","instruction":"Check every deadline and assigned owner before answering."}'
            ))
            uploaded = await upload_user_agent_skill(upload, owner, db)
            self.assertTrue(uploaded["custom"])
            deleted = await delete_user_agent_skill(uploaded["id"], owner, db)
            self.assertTrue(deleted["deleted"])


if __name__ == "__main__":
    unittest.main()
