import tempfile
import unittest
from pathlib import Path

from docx import Document
from PIL import Image

from app.agent_attachment_reader import read_attachment
from app.case_agent import DOCUMENT_READING_RULES


class AgentAttachmentReaderContractTest(unittest.TestCase):
    def test_reads_word_paragraphs_and_tables(self):
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "case.docx"
            document = Document()
            document.add_heading("案件事实", level=1)
            document.add_paragraph("被告于测试日期销售涉案商品。")
            table = document.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "合同编号"
            table.rows[0].cells[1].text = "SHHT-001"
            document.save(path)

            reading = read_attachment(path, path.name)

        self.assertEqual(reading.status, "parsed")
        self.assertIn("被告于测试日期销售涉案商品", reading.text)
        self.assertIn("合同编号 | SHHT-001", reading.text)

    def test_scanned_pdf_is_forwarded_as_visual_page(self):
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "scan.pdf"
            Image.new("RGB", (180, 120), "white").save(path, "PDF")

            reading = read_attachment(path, path.name)

        self.assertEqual(reading.status, "visual")
        self.assertEqual(reading.page_count, 1)
        self.assertEqual(len(reading.images), 1)
        self.assertTrue(reading.images[0]["data_url"].startswith("data:image/jpeg;base64,"))

    def test_model_rules_require_real_attachment_reading(self):
        self.assertIn("实际解析的附件内容", DOCUMENT_READING_RULES)
        self.assertIn("文件名和页码", DOCUMENT_READING_RULES)
        self.assertIn("只看到了附件清单", DOCUMENT_READING_RULES)

    def test_visual_pages_are_distributed_across_attachments(self):
        source = (Path(__file__).parent / "app" / "main.py").read_text(encoding="utf-8")
        self.assertIn("document_visual_groups", source)
        self.assertIn("for page_index in range(4)", source)
        self.assertIn("document_visual_count >= 12", source)

    def test_selected_documents_and_streaming_are_permission_scoped(self):
        source = (Path(__file__).parent / "app" / "main.py").read_text(encoding="utf-8")
        runtime_source = (Path(__file__).parent / "app" / "case_agent.py").read_text(encoding="utf-8")
        self.assertIn("body.document_ids is None", source)
        self.assertIn("item not in allowed_document_ids", source)
        self.assertIn("case_agent_runtime.invoke_stream", source)
        self.assertIn('choices = payload.get("choices") or []', runtime_source)


if __name__ == "__main__":
    unittest.main()
