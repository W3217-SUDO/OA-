import importlib
import tempfile
import unittest
from datetime import datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path

import httpx
from docx import Document
from docx.oxml import OxmlElement
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine
from sqlalchemy.pool import StaticPool

from app.config import settings
from app.database import Base, get_db
from app.main import app
from app.models import BusinessRecord, FileAttachment, User
from app.security import current_identity


main_module = importlib.import_module("app.main")
OWNER = {"username": "CODEX-WORD-owner", "display_name": "Word 编辑人", "role": "user", "department": "测试部"}
OTHER = {"username": "CODEX-WORD-other", "display_name": "另一编辑人", "role": "admin", "role_ids": ["admin"], "department": "测试部"}
API = settings.api_prefix


class CaseWordEditorContractTest(unittest.IsolatedAsyncioTestCase):
    async def asyncSetUp(self):
        self.temp_dir = tempfile.TemporaryDirectory(prefix="codex-word-editor-")
        self.previous_upload_root = main_module.UPLOAD_ROOT
        main_module.UPLOAD_ROOT = Path(self.temp_dir.name)
        self.engine = create_async_engine("sqlite+aiosqlite:///:memory:", connect_args={"check_same_thread": False}, poolclass=StaticPool)
        self.sessions = async_sessionmaker(self.engine, expire_on_commit=False, class_=AsyncSession)
        async with self.engine.begin() as connection:
            await connection.run_sync(Base.metadata.create_all)
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("原正文").bold = True
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "原表格"
        output = BytesIO(); document.save(output)
        self.docx_path = Path(self.temp_dir.name) / "word-editor.docx"
        self.docx_path.write_bytes(output.getvalue())
        self.doc_path = Path(self.temp_dir.name) / "legacy.doc"
        self.doc_path.write_bytes(b"legacy binary")
        async with self.sessions() as db:
            case = BusinessRecord(module="case", serial_no="CODEX-WORD-CASE", title="Word 在线编辑", customer="测试客户", status="文书准备", owner=OWNER["username"], department=OWNER["department"], data={})
            db.add_all([
                case,
                User(username=OWNER["username"], display_name=OWNER["display_name"], department=OWNER["department"], role="user", role_ids=["user"], password_hash="test-only"),
                User(username=OTHER["username"], display_name=OTHER["display_name"], department=OTHER["department"], role="admin", role_ids=["admin"], password_hash="test-only"),
            ])
            await db.flush(); self.case_id = case.id
            attachment = FileAttachment(record_id=case.id, category="法院诉讼文书", original_name="word-editor.docx", stored_name=self.docx_path.name, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", size=self.docx_path.stat().st_size, path=str(self.docx_path), uploader=OWNER["username"], remark="")
            legacy = FileAttachment(record_id=case.id, category="法院诉讼文书", original_name="legacy.doc", stored_name=self.doc_path.name, content_type="application/msword", size=self.doc_path.stat().st_size, path=str(self.doc_path), uploader=OWNER["username"], remark="")
            db.add_all([attachment, legacy]); await db.commit()
            self.attachment_id, self.legacy_id = attachment.id, legacy.id
        self.identity = OWNER

        async def override_db():
            async with self.sessions() as db:
                yield db
        self.previous_overrides = dict(app.dependency_overrides)
        app.dependency_overrides[get_db] = override_db
        app.dependency_overrides[current_identity] = lambda: self.identity
        self.client = httpx.AsyncClient(transport=httpx.ASGITransport(app=app), base_url="http://word-editor.test")

    async def asyncTearDown(self):
        await self.client.aclose()
        app.dependency_overrides.clear(); app.dependency_overrides.update(self.previous_overrides)
        main_module.UPLOAD_ROOT = self.previous_upload_root
        await self.engine.dispose(); self.temp_dir.cleanup()

    def _url(self, suffix: str) -> str:
        return f"{API}/cases/{self.case_id}/attachments/{self.attachment_id}/word-editor{suffix}"

    async def test_edits_text_blocks_preserves_table_and_requires_current_lease_and_version(self):
        opened = await self.client.get(self._url("/content"))
        self.assertEqual(opened.status_code, 200, opened.text)
        payload = opened.json()
        self.assertEqual(len(payload["version"]), 64)
        self.assertTrue(payload["lock_token"])
        self.assertIn("原表格", payload["content"])

        self.identity = OTHER
        blocked = await self.client.post(self._url("/lock"))
        self.assertEqual(blocked.status_code, 409, blocked.text)
        self.assertEqual(blocked.json()["detail"]["lock_holder"], OWNER["username"])
        self.identity = OWNER
        async with self.sessions() as db:
            locked_row = await db.get(FileAttachment, self.attachment_id)
            self.assertEqual(locked_row.word_editor_locked_by, OWNER["username"])
            self.assertTrue(locked_row.word_editor_lock_token)
            now = main_module._word_editor_now()
            self.assertGreater(locked_row.word_editor_lock_expires_at, now)
        denied_rename = await self.client.put(f"{API}/cases/attachments/{self.attachment_id}/rename", json={"original_name": "renamed.docx"})
        self.assertEqual(denied_rename.status_code, 409, denied_rename.text)

        blocks = payload["blocks"]
        body_block = next(block for block in blocks if block["text"] == "原正文")
        body_block["text"] = "已修改正文"
        saved = await self.client.put(self._url("/content"), json={"lock_token": payload["lock_token"], "version": payload["version"], "blocks": blocks})
        self.assertEqual(saved.status_code, 200, saved.text)
        self.assertNotEqual(saved.json()["version"], payload["version"])
        self.assertNotEqual(saved.json()["lock_token"], payload["lock_token"])
        async with self.sessions() as db:
            saved_attachment = await db.get(FileAttachment, self.attachment_id)
            saved_path = Path(saved_attachment.path)
        self.assertTrue(saved_path.is_file())
        self.assertFalse(self.docx_path.exists())
        result = Document(saved_path)
        changed_paragraph = next(paragraph for paragraph in result.paragraphs if paragraph.text == "已修改正文")
        self.assertTrue(changed_paragraph.runs[0].bold)
        self.assertEqual(result.tables[0].cell(0, 0).text, "原表格")

        stale = await self.client.put(self._url("/content"), json={"lock_token": payload["lock_token"], "version": payload["version"], "blocks": blocks})
        self.assertEqual(stale.status_code, 409, stale.text)
        released = await self.client.request("DELETE", self._url("/lock"), json={"lock_token": saved.json()["lock_token"]})
        self.assertEqual(released.status_code, 200, released.text)

    async def test_rejects_legacy_doc_and_expired_lock(self):
        legacy = await self.client.get(f"{API}/cases/{self.case_id}/attachments/{self.legacy_id}/word-editor/content")
        self.assertEqual(legacy.status_code, 422, legacy.text)
        opened = await self.client.get(self._url("/content")); self.assertEqual(opened.status_code, 200, opened.text)
        async with self.sessions() as db:
            row = await db.get(FileAttachment, self.attachment_id)
            row.word_editor_lock_expires_at = main_module._word_editor_now() - timedelta(seconds=1)
            await db.commit()
        expired = await self.client.post(self._url("/lock/renew"), json={"lock_token": opened.json()["lock_token"]})
        self.assertEqual(expired.status_code, 409, expired.text)

    async def test_complex_word_block_is_read_only_and_cannot_be_overwritten(self):
        document = Document()
        paragraph = document.add_paragraph("受保护段落")
        field = OxmlElement("w:fldSimple")
        field.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr", "DATE")
        paragraph._p.append(field)
        output = BytesIO(); document.save(output)
        self.docx_path.write_bytes(output.getvalue())
        opened = await self.client.get(self._url("/content")); self.assertEqual(opened.status_code, 200, opened.text)
        payload = opened.json()
        protected = next(block for block in payload["blocks"] if block["text"] == "受保护段落")
        self.assertFalse(protected["editable"])
        self.assertTrue(protected["read_only_reason"])
        protected["text"] = "不应覆盖"
        rejected = await self.client.put(self._url("/content"), json={"lock_token": payload["lock_token"], "version": payload["version"], "blocks": payload["blocks"]})
        self.assertEqual(rejected.status_code, 422, rejected.text)
        self.assertIn("受保护段落", Document(self.docx_path).paragraphs[0].text)


if __name__ == "__main__":
    unittest.main()
