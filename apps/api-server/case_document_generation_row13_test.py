"""Row 13: legacy-equivalent case document generation and permission contract."""

from __future__ import annotations

import importlib
import tempfile
import unittest
from pathlib import Path

import httpx
from docx import Document
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine
from sqlalchemy.pool import StaticPool

from app.config import settings
from app.database import Base, get_db
from app.main import app
from app.models import BusinessRecord, FileAttachment, JobRole, SystemConfig, User, WorkflowEvent
from app.security import current_identity


main_module = importlib.import_module("app.main")
API = settings.api_prefix
EDITOR = {"username": "row13-lawyer", "role": "user", "display_name": "第十三行律师", "department": "诉讼部"}
VIEWER = {"username": "row13-viewer", "role": "user", "display_name": "第十三行查看员", "department": "诉讼部"}

DOCUMENTS = {
    "identification_letter": ("法院诉讼文书", "鉴定函"),
    "archive-cover": ("庭审及庭后文件", "归档封面"),
    "authorization-letter": ("主体及委托资料", "授权委托书"),
    "first-instance-appellant-lawyer-letter": ("法院诉讼文书", "一审（我方原告）律所函"),
    "first-instance-appellee-lawyer-letter": ("法院诉讼文书", "一审（我方被告）律所函"),
    "second-instance-appellant-lawyer-letter": ("法院诉讼文书", "二审（我方上诉）律所函"),
    "second-instance-appellee-lawyer-letter": ("法院诉讼文书", "二审（对方上诉）律所函"),
    "execution-lawyer-letter": ("法院诉讼文书", "执行律所函"),
    "identity-certificate": ("主体及委托资料", "法定代表人身份证明"),
    "settlement-list": ("庭审及庭后文件", "结算提成表"),
    "compensation-payment-application": ("庭审及庭后文件", "代收代付赔偿款申请单"),
}


class CaseDocumentGenerationRow13Test(unittest.IsolatedAsyncioTestCase):
    async def asyncSetUp(self) -> None:
        self.identity = EDITOR
        self.temp_dir = tempfile.TemporaryDirectory(prefix="codex-827-13-docs-")
        self.previous_upload_root = main_module.UPLOAD_ROOT
        main_module.UPLOAD_ROOT = Path(self.temp_dir.name)
        self.engine = create_async_engine(
            "sqlite+aiosqlite:///:memory:", connect_args={"check_same_thread": False}, poolclass=StaticPool,
        )
        self.sessions = async_sessionmaker(self.engine, expire_on_commit=False, class_=AsyncSession)
        async with self.engine.begin() as connection:
            await connection.run_sync(Base.metadata.create_all)
        async with self.sessions() as db:
            customer = BusinessRecord(
                module="customer", serial_no="CODEX-827-13-CUSTOMER", title="第十三行测试公司", customer="第十三行测试公司",
                status="正常", owner=EDITOR["username"], department="诉讼部",
                data={"customer_no": "CODEX-827-13-CUSTOMER", "legal_representative": "测试法人", "legal_agent_id_no": "310101199001011313", "legal_agent_title": "执行董事"},
            )
            db.add(customer)
            db.add_all([
                User(username=EDITOR["username"], display_name=EDITOR["display_name"], department=EDITOR["department"], role="user", password_hash="x", is_active=True, profile={"permission_role_code": "ROW13-EDITOR", "mobile": "13800138130"}),
                User(username=VIEWER["username"], display_name=VIEWER["display_name"], department=VIEWER["department"], role="user", password_hash="x", is_active=True, profile={"permission_role_code": "ROW13-VIEWER"}),
                User(username="row13-assistant", display_name="第十三行助理", department="诉讼部", role="user", password_hash="x", is_active=True, profile={"mobile": "13900139130"}),
                JobRole(code="ROW13-EDITOR", name="案件文书办理岗", permissions=["case-mine", "案件承办"], data_scope="全所数据", is_active=True),
                JobRole(code="ROW13-VIEWER", name="案件文书查看岗", permissions=["case-mine"], data_scope="全所数据", is_active=True),
                SystemConfig(key="company_profile", label="公司设置", group="机构配置", value={"name": "上海申浩律师事务所", "address": "上海市徐汇区华山路1954号浩然高科技大厦", "phone": "021-64484005"}, description="test", updated_by="system"),
            ])
            await db.flush()
            complete = BusinessRecord(
                module="case", serial_no="CODEX-827-13-COMPLETE", title="第十三行完整文书案件", customer=customer.title,
                status="一审准备开庭", owner="another-owner", department="诉讼部",
                data={
                    "customer_id": customer.id, "customer_no": customer.serial_no, "case_type": "民事案件",
                    "case_creation_step": "completed", "case_creation_approval_status": "已通过",
                    "handling_lawyers": [EDITOR["display_name"]], "handling_lawyer_usernames": [EDITOR["username"]],
                    "hearing_lawyer": EDITOR["display_name"], "hearing_lawyer_username": EDITOR["username"],
                    "assistant": "第十三行助理", "assistant_username": "row13-assistant",
                    "plaintiff": customer.title, "opponent": "第十三行被告", "cause_or_charge": "合同纠纷",
                    "first_court_name": "上海市第一中级人民法院", "second_court_name": "上海市高级人民法院",
                    "execution_court_name": "上海市徐汇区人民法院", "case_team_usernames": [EDITOR["username"], "row13-assistant"],
                },
            )
            incomplete = BusinessRecord(
                module="case", serial_no="CODEX-827-13-INCOMPLETE", title="第十三行字段缺失案件", customer=customer.title,
                status="一审准备开庭", owner="another-owner", department="诉讼部",
                data={
                    "customer_id": customer.id, "case_type": "民事案件", "case_creation_step": "basic",
                    "case_creation_approval_status": "已通过", "handling_lawyers": [EDITOR["display_name"]],
                    "handling_lawyer_usernames": [EDITOR["username"]], "case_team_usernames": [EDITOR["username"]],
                },
            )
            db.add_all([complete, incomplete])
            await db.flush()
            self.complete_id = complete.id
            self.incomplete_id = incomplete.id
            await db.commit()
        self.previous_overrides = dict(app.dependency_overrides)
        app.dependency_overrides[get_db] = self.override_db
        app.dependency_overrides[current_identity] = lambda: self.identity
        self.client = httpx.AsyncClient(transport=httpx.ASGITransport(app=app), base_url="http://row13.test")

    async def asyncTearDown(self) -> None:
        await self.client.aclose()
        app.dependency_overrides.clear()
        app.dependency_overrides.update(self.previous_overrides)
        main_module.UPLOAD_ROOT = self.previous_upload_root
        await self.engine.dispose()
        self.temp_dir.cleanup()

    async def override_db(self):
        async with self.sessions() as db:
            yield db

    async def test_all_document_actions_generate_distinct_downloadable_docx_files(self) -> None:
        generated_ids = []
        for document_type, (category, expected_title) in DOCUMENTS.items():
            response = await self.client.post(f"{API}/cases/{self.complete_id}/documents/{document_type}")
            self.assertEqual(response.status_code, 201, response.text)
            body = response.json()
            generated_ids.append(body["id"])
            self.assertEqual(body["category"], category)
            self.assertIn(expected_title, body["original_name"])
            async with self.sessions() as db:
                persisted = await db.get(FileAttachment, body["id"])
                path = Path(persisted.path)
            self.assertTrue(path.is_file())
            self.assertGreater(path.stat().st_size, 1000)
            download = await self.client.get(f"{API}/attachments/{body['id']}/download")
            self.assertEqual(download.status_code, 200, download.text)
            self.assertGreater(len(download.content), 1000)
            visible_text = "\n".join(paragraph.text for paragraph in Document(path).paragraphs)
            if document_type == "identification_letter":
                self.assertEqual(Document(path).core_properties.title, "鉴定函")
                for expected in ["鉴定函", "CODEX-827-13-COMPLETE", "第十三行完整文书案件", "第十三行测试公司", "第十三行律师", "13800138130", "鉴定事项及要求：________________", "送鉴材料：________________"]:
                    self.assertIn(expected, visible_text)
            elif document_type == "authorization-letter":
                self.assertIn("代理权限为特别授权", visible_text)
                self.assertIn("代为申请执行并收取执行款项", visible_text)
            elif document_type == "identity-certificate":
                self.assertIn("测试法人", visible_text)
                self.assertIn("执行董事", visible_text)
            elif document_type == "archive-cover":
                self.assertIn("案件卷宗", visible_text)
                self.assertIn("CODEX-827-13-COMPLETE", visible_text)
            elif document_type == "settlement-list":
                self.assertIn("结算提成表", visible_text)
                self.assertIn("诉讼标的金额", visible_text)
            elif document_type == "compensation-payment-application":
                self.assertIn("代收代付赔偿款申请单", visible_text)
                self.assertIn("第十三行被告", visible_text)
            else:
                self.assertIn("上海申浩律师事务所 函", visible_text)
                self.assertIn("合同纠纷", visible_text)
                self.assertIn("第十三行助理", visible_text)
        self.assertEqual(len(set(generated_ids)), len(DOCUMENTS))
        async with self.sessions() as db:
            self.assertEqual(await db.scalar(select(func.count()).select_from(FileAttachment).where(FileAttachment.record_id == self.complete_id)), len(DOCUMENTS))
            self.assertEqual(await db.scalar(select(func.count()).select_from(WorkflowEvent).where(WorkflowEvent.record_id == self.complete_id, WorkflowEvent.action == "生成案件文书")), len(DOCUMENTS))

    async def test_missing_legacy_fields_fail_atomically_without_file_or_database_residue(self) -> None:
        before = set(Path(self.temp_dir.name).glob("*"))
        response = await self.client.post(f"{API}/cases/{self.incomplete_id}/documents/first-instance-appellant-lawyer-letter")
        self.assertEqual(response.status_code, 422, response.text)
        for label in ["一审法院", "开庭律师", "律师助理", "原告", "被告", "案由"]:
            self.assertIn(label, response.text)
        self.assertEqual(set(Path(self.temp_dir.name).glob("*")), before)
        async with self.sessions() as db:
            self.assertEqual(await db.scalar(select(func.count()).select_from(FileAttachment).where(FileAttachment.record_id == self.incomplete_id)), 0)
            self.assertEqual(await db.scalar(select(func.count()).select_from(WorkflowEvent).where(WorkflowEvent.record_id == self.incomplete_id)), 0)

    async def test_identification_letter_rejects_locked_cases_without_files_or_events(self) -> None:
        for locked_status in ["已归档", "已合并"]:
            async with self.sessions() as db:
                record = await db.get(BusinessRecord, self.complete_id)
                record.status = locked_status
                await db.commit()
            response = await self.client.post(f"{API}/cases/{self.complete_id}/documents/identification_letter")
            self.assertEqual(response.status_code, 409, response.text)
            self.assertEqual(list(Path(self.temp_dir.name).iterdir()), [])
            async with self.sessions() as db:
                self.assertEqual(await db.scalar(select(func.count()).select_from(FileAttachment)), 0)
                self.assertEqual(await db.scalar(select(func.count()).select_from(WorkflowEvent)), 0)

    async def test_frontend_capability_and_backend_permission_use_the_same_gate(self) -> None:
        capability = await self.client.get(f"{API}/cases/{self.complete_id}/action-capabilities")
        self.assertEqual(capability.status_code, 200, capability.text)
        self.assertTrue(capability.json()["can_write"])
        self.assertTrue(capability.json()["can_generate_document"])
        incomplete_capability = await self.client.get(f"{API}/cases/{self.incomplete_id}/action-capabilities")
        self.assertEqual(incomplete_capability.status_code, 200, incomplete_capability.text)
        # Row 14 made basic-information editing independent from the three-step
        # creation wizard, while document generation keeps its own permission gate.
        self.assertTrue(incomplete_capability.json()["can_write"])
        self.assertTrue(incomplete_capability.json()["can_generate_document"])
        self.identity = VIEWER
        denied_capability = await self.client.get(f"{API}/cases/{self.complete_id}/action-capabilities")
        self.assertEqual(denied_capability.status_code, 200, denied_capability.text)
        self.assertTrue(denied_capability.json()["can_write"])
        self.assertTrue(denied_capability.json()["can_generate_document"])
        allowed = await self.client.post(f"{API}/cases/{self.complete_id}/documents/authorization-letter")
        self.assertEqual(allowed.status_code, 201, allowed.text)


if __name__ == "__main__":
    unittest.main()
