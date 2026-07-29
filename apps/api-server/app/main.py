import asyncio
from contextlib import asynccontextmanager, suppress
import csv
import hashlib
from datetime import date, datetime, timedelta, timezone
from decimal import Decimal, ROUND_UP
import io
import json
import logging
from pathlib import Path
import re
import unicodedata
import zipfile
from urllib.parse import quote
from uuid import uuid4
from xml.sax.saxutils import escape as xml_escape

import httpx
from fastapi import Depends, FastAPI, File, Form, HTTPException, Query, Response, UploadFile, status
from docx import Document
from fastapi.middleware.cors import CORSMiddleware
from fastapi.exceptions import RequestValidationError
from fastapi.security import OAuth2PasswordRequestForm
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy import and_, delete, false, func, inspect, or_, select, text, update
from sqlalchemy.exc import IntegrityError
from sqlalchemy.ext.asyncio import AsyncSession

from .config import settings
from .database import Base, SessionLocal, engine, get_db
from .models import AgentDocument, BusinessRecord, CommunicationLog, ContractApprovalStep, ContractEvent, Department, DocumentTemplate, FileAttachment, FinanceTransaction, HearingSchedule, HrSubrecord, IncomingPayment, JobRole, Notification, ReceivablePlan, ReconciliationBatch, RolePermission, SealAsset, SecurityPolicy, SystemConfig, SystemMenu, SystemParameter, User, WorkflowEvent
from .security import create_token, current_identity, hash_password, verify_password


logger = logging.getLogger(__name__)


CASE_CREATE_PERMISSION_BY_TYPE = {
    "民事案件": "case-new-civil",
    "刑事案件": "case-new-criminal",
    "行政案件及国家赔偿": "case-new-administrative",
    "法律顾问": "case-new-counsel",
    "仲裁": "case-new-arbitration",
}
CASE_CREATABLE_TYPES = {"民事案件", "刑事案件", "行政案件及国家赔偿", "法律顾问", "仲裁"}
NORMAL_CASE_BASIC_TYPES = {"民事案件", "刑事案件", "行政案件及国家赔偿"}
CASE_BASIC_EDITABLE_PHASES = {"等待公证书", "等待审核公证书", "待立案审批", "新案待分配", "文书准备", "一审立案受理", "一审准备开庭", "待上诉", "二审", "执行"}
CASE_SOURCE_CONTRACT_STATUSES = {"已通过", "履行中", "已完成"}
REQUIRED_SEAL_ASSETS = (
    ("YZ-HT-001", "合同章", "行政部保险柜 A01"),
    ("YZ-GZ-001", "公章", "行政部保险柜 A02"),
    ("YZ-SH-001", "所函专用章", "行政部保险柜 A03"),
    ("YZ-FR-001", "法人章", "行政部保险柜 A04"),
    ("YZ-FP-001", "发票章", "财务部保险柜 F01"),
    ("YZ-CW-001", "财务专用章", "财务部保险柜 F02"),
    ("YZ-CS-001", "财务三排章", "财务部保险柜 F03"),
)
REQUIRED_SEAL_TYPES = {seal_type for _, seal_type, _ in REQUIRED_SEAL_ASSETS}
ADMINISTRATIVE_CLIENT_POSITIONS = {"原告/申请人", "被告/被申请人", "第三人"}
CASE_CREATE_PERMISSION_KEYS = list(CASE_CREATE_PERMISSION_BY_TYPE.values())
CASE_CLIENT_POSITIONS_BY_TYPE = {
    "民事案件": {"原告/申请人", "被告/被申请人", "第三人"},
    "刑事案件": {"被告人/犯罪嫌疑人", "被害人"},
    "行政案件及国家赔偿": {"原告/申请人", "被告/被申请人", "第三人"},
    "仲裁": {"原告/申请人", "被告/被申请人", "第三人"},
}
CRIMINAL_JUDICIAL_PREFIXES = ("public_security_", "first_procuratorate_", "second_procuratorate_", "retrial_procuratorate_")
COURT_JUDICIAL_KEYS = {
    "court", "court_case_no", "courtroom", "judge", "clerk", "judge_phone", "filing_date", "hearing_date", "hearing_time",
    "first_court_name", "first_court_case_no", "first_court_courtroom", "first_court_judge", "first_court_clerk", "first_court_filing_date", "first_court_hearing_date",
    "second_court_name", "second_court_case_no", "second_court_courtroom", "second_court_judge", "second_court_clerk", "second_court_filing_date", "second_court_hearing_date",
    "retrial_court_name", "retrial_court_case_no", "retrial_court_courtroom", "retrial_court_judge", "retrial_court_clerk", "retrial_court_filing_date", "retrial_court_hearing_date",
}
MENU_KEYS = ["task", "seal", "customer", "customer-conflict", "contract", "case", *CASE_CREATE_PERMISSION_KEYS, "investigation", "documents", "finance", "platform-finance", "user-center", "hr", "system", "warehouse", "reports"]
DEFAULT_SYSTEM_MENUS = [
    ("dashboard", "", "控制台", "dashboard", 0),
    ("seal", "", "用印中心", "file-text", 10), ("seal-my", "seal", "我的用印申请", "", 11), ("seal-audit", "seal", "用印审核", "", 12), ("seal-admin", "seal", "行政用印", "", 13),
    ("task", "", "事务中心", "file-text", 20), ("task-my", "task", "我的任务", "", 21), ("task-dept", "task", "部门任务", "", 22), ("task-company", "task", "公司任务", "", 23),
    ("customer", "", "客户管理", "team", 30), ("customer-new", "customer", "新建客户", "", 31), ("customer-mine", "customer", "我的客户", "", 32), ("customer-recycle", "customer", "个人回收站", "", 33), ("customer-dept", "customer", "部门客户", "", 34), ("customer-dept-recycle", "customer", "部门回收站", "", 35), ("customer-company", "customer", "公司客户", "", 36), ("customer-public", "customer", "公海客户", "", 37), ("customer-shared", "customer", "我的共享客户", "", 38), ("customer-recent-contact", "customer", "最近联系的客户", "", 39), ("customer-recent-update", "customer", "最近更新的客户", "", 40), ("customer-company-recycle", "customer", "公司回收站", "", 41), ("customer-conflict", "customer", "客户利益检索", "", 42),
    ("contract", "", "合同中心", "file-text", 50), ("contract-new", "contract", "合同新建", "", 51), ("contract-mine", "contract", "我的合同", "", 52), ("contract-audit", "contract", "合同审批", "", 53), ("contract-dept", "contract", "部门合同", "", 54), ("contract-company", "contract", "公司合同", "", 55), ("contract-receivable", "contract", "应收款", "", 56),
    ("case", "", "案件中心", "file-text", 60), ("case-new", "case", "新建案件", "", 61), ("case-mine", "case", "我的案件", "", 62), ("case-dept", "case", "部门案件", "", 63), ("case-company", "case", "公司案件", "", 64), ("case-files", "case", "案件文件", "", 65), ("case-archive", "case", "案件归档审核", "", 66),
    ("investigation", "", "调查大厅", "search", 70), ("clue", "investigation", "我的调查线索", "", 71), ("notary", "investigation", "公证管理", "", 72), ("evidence", "investigation", "证据管理", "", 73),
    ("documents", "", "收发文台", "file-text", 80), ("documents-official", "documents", "官文收文", "", 1), ("documents-my", "documents", "我的收文", "", 2), ("documents-company", "documents", "公司收文", "", 3), ("documents-register", "documents", "收发文登记", "", 84), ("documents-files", "documents", "文件附件", "", 85), ("documents-template", "documents", "文书模板", "", 86), ("documents-agent", "documents", "AI 智能文档", "", 87), ("documents-archive", "documents", "归档材料", "", 88),
    ("finance", "", "财务中心", "file-text", 90),
    ("platform-finance", "", "平台财务中心", "bank", 100),
    ("user-center", "", "用户中心", "user", 110), ("user-messages", "user-center", "消息通知", "", 111), ("user-communications", "user-center", "沟通日志", "", 112), ("user-account", "user-center", "账户管理", "", 113),
    ("hr", "", "人事中心", "team", 120), ("hr-new", "hr", "新建员工", "", 1), ("hr-all", "hr", "员工管理", "", 2), ("hr-departments", "hr", "部门管理", "", 3), ("hr-roles", "hr", "角色管理", "", 4),
    ("system", "", "系统中心", "user", 130), ("system-parameters", "system", "系统参数", "", 1), ("system-management", "system", "系统管理", "", 2),
    ("warehouse", "", "仓库管理", "file-text", 140), ("warehouse-list", "warehouse", "仓库一览表", "", 1), ("reports", "", "报表中心", "dashboard", 150),
]
DEFAULT_SYSTEM_MENUS += [
    ("seal-my-pending", "seal-my", "待审批", "", 1), ("seal-my-stamping", "seal-my", "待用印", "", 2), ("seal-my-used", "seal-my", "已用印", "", 3), ("seal-my-refused", "seal-my", "已拒绝", "", 4), ("seal-my-withdrawn", "seal-my", "已撤回", "", 5),
    ("seal-audit-pending", "seal-audit", "待审批用印", "", 1), ("seal-audit-stamping", "seal-audit", "已审待用印", "", 2), ("seal-audit-refused", "seal-audit", "已拒绝用印", "", 3),
    ("seal-admin-pending", "seal-admin", "待用印", "", 1), ("seal-admin-used", "seal-admin", "已用印", "", 2), ("seal-admin-query", "seal-admin", "用印查询", "", 3),
    ("task-my-created", "task-my", "我发起的任务", "", 1), ("task-my-accepted", "task-my", "我接受的任务", "", 2), ("task-my-collaborating", "task-my", "我协作的任务", "", 3), ("task-my-unread", "task-my", "未读新消息的任务", "", 4),
    ("task-dept-created", "task-dept", "部门发起的任务", "", 1), ("task-dept-accepted", "task-dept", "部门接受的任务", "", 2), ("task-dept-collaborating", "task-dept", "部门协作的任务", "", 3),
    ("task-company-created", "task-company", "公司发起的任务", "", 1), ("task-company-accepted", "task-company", "公司接受的任务", "", 2), ("task-company-collaborating", "task-company", "公司协作的任务", "", 3),
    ("contract-audit-pending", "contract-audit", "待审批合同", "", 1), ("contract-audit-refused", "contract-audit", "已驳回合同", "", 2), ("contract-audit-approved", "contract-audit", "已审批合同", "", 3),
    ("contract-receivable-mine", "contract-receivable", "我的应收款", "", 1), ("contract-receivable-dept", "contract-receivable", "部门应收款", "", 2), ("contract-receivable-company", "contract-receivable", "公司应收款", "", 3), ("contract-receivable-detail", "contract-receivable", "应收款明细", "", 4),
    ("case-new-civil", "case-new", "民事争议", "", 1), ("case-new-criminal", "case-new", "刑事案件", "", 2), ("case-new-administrative", "case-new", "行政案件及国家赔偿", "", 3), ("case-new-counsel", "case-new", "法律顾问", "", 4), ("case-new-arbitration", "case-new", "仲裁", "", 5),
    ("case-mine-civil", "case-mine", "民事争议", "", 1), ("case-mine-criminal", "case-mine", "刑事案件", "", 2), ("case-mine-administrative", "case-mine", "行政案件及国家赔偿", "", 3), ("case-mine-counsel", "case-mine", "法律顾问", "", 4), ("case-mine-arbitration", "case-mine", "仲裁", "", 5), ("case-mine-schedule", "case-mine", "我的案件开庭排期", "", 6), ("case-mine-execution", "case-mine", "执行案件", "", 7), ("case-mine-unclaimed", "case-mine", "未申请提成的到账案件", "", 8),
    ("case-dept-civil", "case-dept", "民事争议", "", 1), ("case-dept-criminal", "case-dept", "刑事案件", "", 2), ("case-dept-administrative", "case-dept", "行政案件及国家赔偿", "", 3), ("case-dept-counsel", "case-dept", "法律顾问", "", 4), ("case-dept-arbitration", "case-dept", "仲裁", "", 5), ("case-dept-schedule", "case-dept", "部门案件开庭排期", "", 6),
    ("case-company-civil", "case-company", "民事争议", "", 1), ("case-company-criminal", "case-company", "刑事案件", "", 2), ("case-company-administrative", "case-company", "行政案件及国家赔偿", "", 3), ("case-company-counsel", "case-company", "法律顾问", "", 4), ("case-company-arbitration", "case-company", "仲裁", "", 5), ("case-company-schedule", "case-company", "公司案件开庭排期", "", 6), ("case-company-execution", "case-company", "执行案件", "", 7), ("case-company-unclaimed", "case-company", "未申请提成的到账案件", "", 8), ("case-company-stage", "case-company", "案件阶段统计", "", 9), ("case-company-no-refund", "case-company", "不再办理退费案件", "", 10),
    ("case-files-receipt", "case-files", "案件票据文件", "", 1), ("case-files-invoice", "case-files", "案件发票文件", "", 2),
    ("case-archive-pending", "case-archive", "待审核列表", "", 1), ("case-archive-done", "case-archive", "已归档列表", "", 2), ("case-archive-refused", "case-archive", "已拒绝列表", "", 3), ("case-archive-loss-internal", "case-archive", "亏损内审列表", "", 4), ("case-archive-loss-audit", "case-archive", "亏损审核列表", "", 5), ("case-archive-loss-done", "case-archive", "亏损归档列表", "", 6), ("case-archive-loss-refused", "case-archive", "亏损拒绝列表", "", 7),
]
DEFAULT_SYSTEM_MENUS += [
    ("investigation-task-published", "investigation", "我发布的调查任务", "", 1), ("investigation-task-mine", "investigation", "我的调查任务", "", 2), ("investigation-task-overdue", "investigation", "过期调查任务", "", 3), ("investigation-task-unassigned", "investigation", "待我分配的调查任务", "", 4), ("investigation-task-sub-published", "investigation", "我发布的调查子任务", "", 5), ("investigation-task-sub-mine", "investigation", "我的调查任务", "", 6),
    ("clue-my-draft", "clue", "待提交线索", "", 1), ("clue-my-pending", "clue", "待审核线索", "", 2), ("clue-my-customer", "clue", "待客户审核", "", 3), ("clue-my-collect", "clue", "待取证线索", "", 4), ("clue-my-collected", "clue", "已取证线索", "", 5), ("clue-my-refused", "clue", "已拒绝线索", "", 6), ("clue-my-no-fee", "clue", "未申请费用线索", "", 7), ("clue-my-fee", "clue", "已申请费用线索", "", 8),
    ("clue-audit", "investigation", "调查线索审核", "", 72), ("clue-audit-pending", "clue-audit", "待审批线索", "", 1), ("clue-audit-customer", "clue-audit", "待客户审核", "", 2), ("clue-audit-refused", "clue-audit", "已拒绝线索", "", 3), ("clue-audit-collect", "clue-audit", "待取证线索", "", 4), ("clue-audit-collected", "clue-audit", "已取证线索", "", 5),
    ("clue-company", "investigation", "公司调查线索", "", 73), ("clue-company-draft", "clue-company", "待提交线索", "", 1), ("clue-company-pending", "clue-company", "待审核线索", "", 2), ("clue-company-collect", "clue-company", "待取证线索", "", 3), ("clue-company-collected", "clue-company", "已取证线索", "", 4), ("clue-company-refused", "clue-company", "已拒绝线索", "", 5), ("clue-company-no-fee", "clue-company", "未申请费用线索", "", 6), ("clue-company-fee", "clue-company", "已申请费用线索", "", 7),
    ("notary-import-storage", "notary", "公证书号仓库信息导入", "", 1), ("notary-import-files", "notary", "公证书文件导入", "", 2), ("notary-import-invoices", "notary", "发票文件导入", "", 3), ("notary-query-files", "notary", "公证文件查询", "", 4),
    ("finance-receipts", "finance", "回款管理", "", 1), ("finance-receipts-icbc", "finance-receipts", "回款(工行)", "", 1), ("finance-receipts-citic", "finance-receipts", "回款(中信)", "", 2), ("finance-receipts-boc", "finance-receipts", "回款(中行)", "", 3), ("finance-receipts-new", "finance-receipts", "新增回款", "", 4), ("finance-receipts-manage", "finance-receipts", "回款管理", "", 5), ("finance-receipts-claim", "finance-receipts", "回款领取", "", 6), ("finance-receipts-pending", "finance-receipts", "待分配回款", "", 7), ("finance-receipts-allocated", "finance-receipts", "已分配回款", "", 8), ("finance-receipts-query", "finance-receipts", "到账查询", "", 9),
    ("finance-payment", "finance", "付款管理", "", 2), ("finance-payment-mine", "finance-payment", "我的请款单", "", 1), ("finance-payment-audit", "finance-payment", "请款单审批", "", 2), ("finance-payment-waiting", "finance-payment", "待付款列表", "", 3), ("finance-payment-print", "finance-payment", "付款单打印", "", 4), ("finance-payment-writeoff", "finance-payment", "待核销列表", "", 5), ("finance-payment-query", "finance-payment", "付款单查询", "", 6),
    ("finance-internal", "finance", "内部费用", "", 3), ("finance-internal-mine", "finance-internal", "我的请款单", "", 1), ("finance-internal-settle", "finance-internal", "内部提成-待结算", "", 2), ("finance-internal-archive", "finance-internal", "内部提成-待归档", "", 3), ("finance-internal-audit", "finance-internal", "内部提成-待审核", "", 4), ("finance-internal-fee-audit", "finance-internal", "内部费用-待审核", "", 5), ("finance-internal-refused", "finance-internal", "内部提成-已拒绝", "", 6), ("finance-internal-void", "finance-internal", "内部提成-已作废", "", 7), ("finance-internal-refund-audit", "finance-internal", "内部提成(退费)-待审核", "", 8), ("finance-internal-payment", "finance-internal", "待付款列表", "", 9), ("finance-internal-writeoff", "finance-internal", "待核销列表", "", 10), ("finance-internal-query", "finance-internal", "付款单查询", "", 11), ("finance-internal-done", "finance-internal", "已核销列表", "", 12), ("finance-internal-detail", "finance-internal", "内部费用明细", "", 13), ("finance-internal-company", "finance-internal", "内部费用明细(公司)", "", 14),
    ("finance-invoice", "finance", "开票管理", "", 4), ("finance-invoice-mine", "finance-invoice", "我的开票", "", 1), ("finance-invoice-pending", "finance-invoice", "待处理开票", "", 2), ("finance-invoice-company", "finance-invoice", "公司开票", "", 3), ("finance-invoice-unissued", "finance-invoice", "未开票", "", 4), ("finance-invoice-company-unissued", "finance-invoice", "公司未开票", "", 5),
    ("finance-settlement", "finance", "结算管理", "", 5), ("finance-settlement-pending", "finance-settlement", "待结算", "", 1), ("finance-settlement-audit", "finance-settlement", "待审核", "", 2), ("finance-settlement-payment", "finance-settlement", "待付款", "", 3), ("finance-settlement-paid", "finance-settlement", "已付款", "", 4), ("finance-settlement-refused", "finance-settlement", "已拒绝", "", 5),
    ("finance-archive-fee", "finance", "归档费结算", "", 6), ("finance-archive-fee-pending", "finance-archive-fee", "待归档", "", 1), ("finance-archive-fee-payment", "finance-archive-fee", "待支付", "", 2), ("finance-archive-fee-paid", "finance-archive-fee", "已支付", "", 3), ("finance-archive-fee-refused", "finance-archive-fee", "已拒绝", "", 4),
    ("finance-fee-query", "finance", "费用查询", "", 7),
    ("platform-finance-overview", "platform-finance", "回款管理", "", 1), ("platform-finance-overview-icbc", "platform-finance-overview", "回款(工行)", "", 1), ("platform-finance-overview-citic", "platform-finance-overview", "回款(中信)", "", 2), ("platform-finance-overview-boc", "platform-finance-overview", "回款(中行)", "", 3), ("platform-finance-overview-new", "platform-finance-overview", "新增回款", "", 4), ("platform-finance-overview-manage", "platform-finance-overview", "回款管理", "", 5), ("platform-finance-overview-claim", "platform-finance-overview", "回款领取", "", 6), ("platform-finance-overview-pending", "platform-finance-overview", "待分配回款", "", 7), ("platform-finance-overview-allocated", "platform-finance-overview", "已分配回款", "", 8), ("platform-finance-overview-query", "platform-finance-overview", "到账查询", "", 9),
    ("platform-finance-invoice", "platform-finance", "开票管理", "", 3), ("platform-finance-invoice-mine", "platform-finance-invoice", "我的开票", "", 1), ("platform-finance-invoice-pending", "platform-finance-invoice", "待处理开票", "", 2), ("platform-finance-invoice-company", "platform-finance-invoice", "公司开票", "", 3),
    ("reports-brand", "reports", "品牌资金运营情况统计", "", 1), ("reports-lawyer", "reports", "律师资金运营情况统计", "", 2), ("reports-refund", "reports", "退费进度案件统计", "", 3), ("reports-execution-1", "reports", "执行进度案件统计1", "", 4), ("reports-execution-2", "reports", "执行进度案件统计2", "", 5), ("reports-execution-3", "reports", "执行进度案件统计3", "", 6),
]
DEFAULT_SYSTEM_MENUS += [
    ("platform-finance-payment", "platform-finance", "付款管理", "", 2), ("platform-finance-payment-mine", "platform-finance-payment", "我的请款单", "", 1), ("platform-finance-payment-audit", "platform-finance-payment", "请款单审批", "", 2), ("platform-finance-payment-waiting", "platform-finance-payment", "待付款列表", "", 3), ("platform-finance-payment-print", "platform-finance-payment", "付款单打印", "", 4), ("platform-finance-payment-writeoff", "platform-finance-payment", "待核销列表", "", 5), ("platform-finance-payment-query", "platform-finance-payment", "付款单查询", "", 6),
    ("platform-finance-settlement", "platform-finance", "结算管理", "", 4), ("platform-finance-settlement-pending", "platform-finance-settlement", "待结算", "", 1), ("platform-finance-settlement-audit", "platform-finance-settlement", "待审核", "", 2), ("platform-finance-settlement-payment", "platform-finance-settlement", "待付款", "", 3), ("platform-finance-settlement-paid", "platform-finance-settlement", "已付款", "", 4), ("platform-finance-settlement-refused", "platform-finance-settlement", "已拒绝", "", 5),
    ("platform-finance-archive-fee", "platform-finance", "归档费结算", "", 5), ("platform-finance-archive-fee-pending", "platform-finance-archive-fee", "待归档", "", 1), ("platform-finance-archive-fee-payment", "platform-finance-archive-fee", "待支付", "", 2), ("platform-finance-archive-fee-paid", "platform-finance-archive-fee", "已支付", "", 3), ("platform-finance-archive-fee-refused", "platform-finance-archive-fee", "已拒绝", "", 4),
    ("platform-finance-fee-query", "platform-finance", "费用查询", "", 6),
]
DEFAULT_SYSTEM_MENUS += [
    ("system-parameters-case-type", "system-parameters", "案件类型", "", 1),
    ("system-parameters-fee-type", "system-parameters", "费用类型", "", 2),
    ("system-parameters-case-phase", "system-parameters", "案件阶段", "", 3),
    ("system-parameters-court", "system-parameters", "法院设置", "", 4),
    ("system-parameters-notary", "system-parameters", "公证处设置", "", 5),
    ("system-parameters-cause", "system-parameters", "案由设置", "", 6),
    ("system-parameters-payment", "system-parameters", "付款类型", "", 7),
    ("system-parameters-company", "system-parameters", "公司设置", "", 8),
    ("system-management-cache", "system-management", "缓存管理", "", 1),
    ("system-management-menu", "system-management", "菜单管理", "", 2),
    ("system-management-config", "system-management", "系统配置", "", 3),
]
SYSTEM_MENU_ROUTE_KEYS = {key for key, *_ in DEFAULT_SYSTEM_MENUS}
LEGACY_FINANCE_MENU_KEYS = {
    "finance-fees", "finance-audit", "finance-refund", "finance-transactions", "finance-reconcile",
    "platform-finance-reconcile",
}
LEGACY_ADMIN_MENU_KEYS = {
    "hr-active", "hr-probation", "hr-offboard",
    "system-users", "system-roles", "system-audit", "system-security",
}
ORIGINAL_FINANCE_MENU_KEYS = {
    key for key, _, _, _, _ in DEFAULT_SYSTEM_MENUS
    if key == "finance" or key.startswith("finance-") or key == "platform-finance" or key.startswith("platform-finance-")
}
ORIGINAL_ADMIN_MENU_KEYS = {
    key for key, _, _, _, _ in DEFAULT_SYSTEM_MENUS
    if key == "hr" or key.startswith("hr-") or key == "system" or key.startswith("system-") or key == "warehouse" or key.startswith("warehouse-")
}
ORIGINAL_INVESTIGATION_MENU_KEYS = {
    key for key, _, _, _, _ in DEFAULT_SYSTEM_MENUS
    if key == "investigation" or key in {"clue", "notary", "evidence"} or key.startswith(("investigation-", "clue-", "notary-"))
}
LEGACY_INVESTIGATION_MENU_KEYS = {"notary-import-info"}
LEGACY_TASK_MENU_KEYS = {"task-reminders"}
FIELD_KEYS = ["customer.billing", "customer.bank", "customer.legal", "contract.amount", "finance.amount", "hr.identity"]
DEFAULT_ROLE_PERMISSIONS = {
    "admin": {"display_name": "系统管理员", "data_scope": "全所数据", "menu_keys": MENU_KEYS, "field_keys": FIELD_KEYS},
    "manager": {"display_name": "部门负责人", "data_scope": "本部门数据", "menu_keys": ["task", "seal", "customer", "customer-conflict", "contract", "case", *CASE_CREATE_PERMISSION_KEYS, "investigation", "documents", "finance", "user-center", "hr", "warehouse", "reports"], "field_keys": FIELD_KEYS},
    "auditor": {"display_name": "审批人员", "data_scope": "授权审批数据", "menu_keys": ["task", "seal", "contract", "case", "investigation", "finance", "platform-finance", "user-center", "reports"], "field_keys": ["contract.amount", "finance.amount"]},
    "user": {"display_name": "普通用户", "data_scope": "本人及共享数据", "menu_keys": ["task", "customer", "customer-conflict", "contract", "case", *CASE_CREATE_PERMISSION_KEYS, "investigation", "documents", "finance", "user-center"], "field_keys": ["customer.legal", "contract.amount"]},
}
SYSTEM_PARAMETER_CATEGORIES = {
    "case_type": "案件类型",
    "fee_type": "费用类型",
    "case_phase": "案件阶段",
    "court": "法院设置",
    "notary_office": "公证处设置",
    "cause": "案由设置",
    "payment_type": "付款类型",
}
DEFAULT_SYSTEM_PARAMETERS = [
    ("case_type", "110", "民事争议", {"letter_code": "MS"}),
    ("case_type", "120", "刑事案件", {"letter_code": "XS"}),
    ("case_type", "130", "行政案件及国家赔偿", {"letter_code": "XZ"}),
    ("case_type", "140", "法律顾问", {"letter_code": "GW"}),
    ("case_type", "150", "仲裁", {"letter_code": "ZC"}),
    ("fee_type", "OFFICIAL", "官方费用", {"group": "案件费用"}),
    ("fee_type", "INTERNAL", "内部费用", {"group": "内部费用"}),
    ("fee_type", "SETTLEMENT", "结算费用", {"group": "结算费用"}),
    ("fee_type", "ARCHIVE", "归档费用", {"group": "归档费用"}),
    ("case_phase", "NEW", "新案待分配", {"case_type": "民事争议", "parent_code": "", "sort_order": 10}),
    ("case_phase", "DOCUMENT", "文书准备", {"case_type": "民事争议", "parent_code": "NEW", "sort_order": 20}),
    ("case_phase", "FIRST", "一审", {"case_type": "民事争议", "parent_code": "DOCUMENT", "sort_order": 30}),
    ("case_phase", "SECOND", "二审", {"case_type": "民事争议", "parent_code": "FIRST", "sort_order": 40}),
    ("case_phase", "EXECUTION", "执行", {"case_type": "民事争议", "parent_code": "SECOND", "sort_order": 50}),
    ("court", "SHBS", "上海市宝山区人民法院", {}),
    ("court", "SHXH", "上海市徐汇区人民法院", {}),
    ("court", "SHIP", "上海知识产权法院", {}),
    ("notary_office", "SHDF", "上海市东方公证处", {"number_template": "（{year}）沪东证经字第{serial}号"}),
    ("notary_office", "SHXH", "上海市徐汇公证处", {"number_template": "（{year}）沪徐证经字第{serial}号"}),
    ("cause", "IP-TRADEMARK", "侵害商标权纠纷", {"parent_code": "知识产权纠纷"}),
    ("cause", "IP-COPYRIGHT", "侵害著作权纠纷", {"parent_code": "知识产权纠纷"}),
    ("cause", "IP-PATENT", "侵害专利权纠纷", {"parent_code": "知识产权纠纷"}),
    ("payment_type", "COURT", "法院费用", {"nature": "对公", "payee": "人民法院", "account": ""}),
    ("payment_type", "NOTARY", "公证费用", {"nature": "对公", "payee": "公证处", "account": ""}),
]
DEFAULT_SYSTEM_CONFIGS = {
    "customer_share_policy": {
        "label": "客户共享时间设置", "group": "业务配置", "description": "客户在不同等级或状态下进入共享范围的天数。",
        "value": {"all_days": 100, "filed_days": 360, "premium_days": 30, "standard_days": 40, "basic_days": 60, "shared_days": 1000},
    },
    "company_profile": {
        "label": "公司设置", "group": "机构配置", "description": "系统抬头、开票和联系信息。",
        "value": {"name": "上海申浩律师事务所", "code": "31280", "short_code": "SH", "address": "上海市徐汇区华山路1954号浩然高科技大厦", "phone": "021-64484005", "fax": "", "email": "", "postal_code": "", "bank_name": "", "bank_account": "", "bank_address": ""},
    },
    "application_settings": {
        "label": "系统配置", "group": "运行配置", "description": "网页端显示与业务运行参数。",
        "value": {"system_name": "申浩律师协作平台", "default_department": "上海分所", "page_size": 20, "attachment_limit_mb": 20, "maintenance_mode": False},
    },
}
SYSTEM_PARAMETER_CACHE: dict[str, list[dict]] = {}
SYSTEM_CACHE_META = {"system-parameters": {"last_cleared_at": None, "last_cleared_by": ""}}
DEFAULT_DEPARTMENTS = [
    ("SHONE", "诉讼一部"), ("SHTWO", "诉讼二部"), ("SH-THREE", "诉讼三部"), ("SHFOUR", "诉讼四部"),
    ("IP", "知识产权中心"), ("DCQZ", "调查取证部"), ("CW", "财务部"), ("SH", "审核部"),
    ("DA", "档案部"), ("LC", "行政流程部"), ("HZ", "合作律师部"), ("HH", "合伙人律师部"),
    ("SHNJ", "申浩南京办公室"), ("GENERAL", "综合管理部"), ("SH-OFFICE", "上海分所"),
    ("BJ-OFFICE", "北京分所"), ("HZ-OFFICE", "杭州分所"), ("SZ-OFFICE", "深圳分所"),
]
DEFAULT_JOB_ROLES = [
    ("SYSTEM-ADMIN", "系统管理员", ["系统配置", "用户管理", "权限分配", "全部模块管理"]),
    ("BUSINESS-SPECIALIST", "业务专员", ["客户新建", "客户编辑", "联系人管理"]),
    ("CUSTOMER-SUPERVISOR", "客户主管", ["客户审批", "客户分级调整", "客户服务端开通"]),
    ("CONTRACT-ADMIN", "合同管理员", ["合同创建", "合同变更", "外部合同号管理"]),
    ("INVESTIGATION-SUPERVISOR", "调查主管", ["调查任务发布", "线索审批", "取证安排"]),
    ("INVESTIGATOR-ROLE", "调查员", ["线索提交", "取证执行", "扫描上传"]),
    ("NOTARY", "公证员", ["公证书号码登记", "公证审核"]),
    ("CASE-SUPERVISOR", "案件主管", ["案件审批", "任务分配", "阶段流转"]),
    ("HANDLING-LAWYER", "承办律师", ["案件办理", "费用申请", "文档上传"]),
    ("FINANCE-SPECIALIST", "财务专员", ["付款录入", "回款录入", "开票录入", "退费录入"]),
    ("FINANCE-SUPERVISOR", "财务主管", ["财务审批", "财务对账", "归档结算"]),
    ("ARCHIVIST", "档案管理员", ["原件接收", "档案移交", "归档登记"]),
    ("GENERAL-USER", "普通用户", ["控制台查看", "待办处理"]),
    ("PARTNER", "合伙人律师", ["案件承办", "合同审批", "部门管理"]),
    ("LEAD-LAWYER", "主办律师", ["案件承办", "开庭办理", "任务派发"]),
    ("ASSISTANT", "律师助理", ["文书准备", "材料归档", "任务办理"]),
    ("INVESTIGATOR", "调查专员", ["调查取证", "证据整理"]),
    ("FINANCE", "财务人员", ["费用审核", "收付款", "对账"]),
    ("ADMINISTRATION", "行政人员", ["用印管理", "收发文", "仓库管理"]),
]


def _upgrade_schema(connection) -> None:
    """为已有本地数据库补充 create_all 不会自动增加的兼容字段。"""
    columns = {item["name"] for item in inspect(connection).get_columns("file_attachments")}
    if "finance_transaction_id" not in columns:
        connection.execute(text("ALTER TABLE file_attachments ADD COLUMN finance_transaction_id INTEGER"))
        connection.execute(text("CREATE INDEX IF NOT EXISTS ix_file_attachments_finance_transaction_id ON file_attachments (finance_transaction_id)"))
    user_columns = {item["name"] for item in inspect(connection).get_columns("users")}
    if "department" not in user_columns:
        connection.execute(text("ALTER TABLE users ADD COLUMN department VARCHAR(64) NOT NULL DEFAULT '上海分所'"))
        connection.execute(text("CREATE INDEX IF NOT EXISTS ix_users_department ON users (department)"))
    for column, definition in {
        "profile": "JSON NOT NULL DEFAULT '{}'",
        "failed_login_attempts": "INTEGER NOT NULL DEFAULT 0",
        "locked_until": "DATETIME",
        "last_login_at": "DATETIME",
        "password_changed_at": "DATETIME",
        "must_change_password": "BOOLEAN NOT NULL DEFAULT FALSE",
    }.items():
        if column not in user_columns: connection.execute(text(f"ALTER TABLE users ADD COLUMN {column} {definition}"))
    role_columns = {item["name"] for item in inspect(connection).get_columns("role_permissions")}
    if "field_keys" not in role_columns:
        default_fields = json.dumps(FIELD_KEYS, ensure_ascii=False)
        connection.execute(text(f"ALTER TABLE role_permissions ADD COLUMN field_keys JSON NOT NULL DEFAULT '{default_fields}'"))
        for role, config in DEFAULT_ROLE_PERMISSIONS.items():
            fields = json.dumps(config["field_keys"], ensure_ascii=False).replace("'", "''")
            connection.execute(text(f"UPDATE role_permissions SET field_keys = '{fields}' WHERE role = '{role}'"))
    connection.execute(text("CREATE TABLE IF NOT EXISTS schema_migrations (key VARCHAR(128) PRIMARY KEY)"))
    conflict_capability_migrated = connection.execute(text(
        "SELECT key FROM schema_migrations WHERE key = 'customer_conflict_leaf_v1'"
    )).first()
    if not conflict_capability_migrated:
        role_rows = connection.execute(text("SELECT role, menu_keys FROM role_permissions")).mappings().all()
        for role_row in role_rows:
            raw_keys = role_row["menu_keys"]
            keys = list(raw_keys if isinstance(raw_keys, list) else json.loads(raw_keys or "[]"))
            if role_row["role"] != "auditor" and "customer" in keys and "customer-conflict" not in keys:
                encoded_keys = json.dumps([*keys, "customer-conflict"], ensure_ascii=False).replace("'", "''")
                role = str(role_row["role"]).replace("'", "''")
                connection.execute(text(f"UPDATE role_permissions SET menu_keys = '{encoded_keys}' WHERE role = '{role}'"))
        connection.execute(text("INSERT INTO schema_migrations (key) VALUES ('customer_conflict_leaf_v1')"))
    case_create_capabilities_migrated = connection.execute(text(
        "SELECT key FROM schema_migrations WHERE key = 'case_create_leaf_capabilities_v1'"
    )).first()
    if not case_create_capabilities_migrated:
        role_rows = connection.execute(text("SELECT role, menu_keys FROM role_permissions")).mappings().all()
        for role_row in role_rows:
            raw_keys = role_row["menu_keys"]
            keys = list(raw_keys if isinstance(raw_keys, list) else json.loads(raw_keys or "[]"))
            if role_row["role"] != "auditor" and "case" in keys:
                migrated_keys = [*keys, *(key for key in CASE_CREATE_PERMISSION_KEYS if key not in keys)]
                encoded_keys = json.dumps(migrated_keys, ensure_ascii=False).replace("'", "''")
                role = str(role_row["role"]).replace("'", "''")
                connection.execute(text(f"UPDATE role_permissions SET menu_keys = '{encoded_keys}' WHERE role = '{role}'"))
        connection.execute(text("INSERT INTO schema_migrations (key) VALUES ('case_create_leaf_capabilities_v1')"))
    # Remove the short-lived internal marker used by an earlier development
    # build; internal migrations must never appear in editable system config.
    connection.execute(text("DELETE FROM system_configs WHERE key = 'permission_capability_migrations'"))
    notification_columns = {item["name"] for item in inspect(connection).get_columns("notifications")}
    for column, definition in {
        "sender": "VARCHAR(64) NOT NULL DEFAULT 'system'",
        "notification_type": "VARCHAR(32) NOT NULL DEFAULT '系统通知'",
        "recipient_deleted": "BOOLEAN NOT NULL DEFAULT 0",
        "sender_deleted": "BOOLEAN NOT NULL DEFAULT 0",
    }.items():
        if column not in notification_columns: connection.execute(text(f"ALTER TABLE notifications ADD COLUMN {column} {definition}"))
    connection.execute(text("CREATE INDEX IF NOT EXISTS ix_notifications_sender ON notifications (sender)"))
    connection.execute(text("CREATE INDEX IF NOT EXISTS ix_notifications_notification_type ON notifications (notification_type)"))
    agent_document_columns = {item["name"] for item in inspect(connection).get_columns("agent_documents")}
    timestamp_type = "TIMESTAMP WITH TIME ZONE" if connection.dialect.name == "postgresql" else "DATETIME"
    for column, definition in {
        "content_version": "INTEGER NOT NULL DEFAULT 1",
        "confirmed_by": "VARCHAR(64) NOT NULL DEFAULT ''",
        "confirmed_at": timestamp_type,
        "confirmed_content_hash": "VARCHAR(64) NOT NULL DEFAULT ''",
    }.items():
        if column not in agent_document_columns:
            connection.execute(text(f"ALTER TABLE agent_documents ADD COLUMN {column} {definition}"))


@asynccontextmanager
async def lifespan(_: FastAPI):
    if settings.app_env.strip().lower() == "production":
        unsafe_secret = (
            len(settings.secret_key) < 64
            or "CHANGE_ME" in settings.secret_key.upper()
            or settings.secret_key == "replace-this-before-production"
        )
        unsafe_admin_password = (
            len(settings.initial_admin_password) < 12
            or "CHANGE_ME" in settings.initial_admin_password.upper()
            or settings.initial_admin_password == "20230616601"
            or settings.initial_admin_password.lower() in {"admin", "password", "12345678"}
        )
        if unsafe_secret:
            raise RuntimeError("生产环境 SECRET_KEY 不安全，必须使用至少 64 位随机值")
        if unsafe_admin_password:
            raise RuntimeError("生产环境 INITIAL_ADMIN_PASSWORD 不安全，必须使用至少 12 位强随机一次性密码")
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)
        await connection.run_sync(_upgrade_schema)
    async with SessionLocal() as db:
        existing = await db.scalar(select(User).where(User.username == settings.initial_admin_username))
        if not existing:
            if not settings.initial_admin_password:
                raise RuntimeError("首次初始化必须通过 INITIAL_ADMIN_PASSWORD 配置一次性管理员密码")
            db.add(User(
                username=settings.initial_admin_username,
                display_name=settings.initial_admin_display_name,
                department=settings.initial_admin_department,
                role="admin",
                password_hash=hash_password(settings.initial_admin_password),
                must_change_password=True,
            ))
        if not await db.get(SecurityPolicy, 1):
            db.add(SecurityPolicy(id=1, min_password_length=8, max_failed_attempts=5, lock_minutes=30, token_minutes=settings.access_token_minutes, updated_by="system"))
        existing_roles = set((await db.scalars(select(RolePermission.role))).all())
        for role, config in DEFAULT_ROLE_PERMISSIONS.items():
            if role not in existing_roles:
                db.add(RolePermission(role=role, **config))
        admin_permission = await db.scalar(select(RolePermission).where(RolePermission.role == "admin"))
        if admin_permission:
            admin_config = DEFAULT_ROLE_PERMISSIONS["admin"]
            admin_permission.display_name = admin_config["display_name"]
            admin_permission.data_scope = admin_config["data_scope"]
            admin_permission.menu_keys = list(MENU_KEYS)
            admin_permission.field_keys = list(FIELD_KEYS)
        if not await db.scalar(select(func.count()).select_from(SystemParameter)):
            db.add_all([
                SystemParameter(category=category, code=code, name=name, extra=extra, sort_order=index, created_by="system", updated_by="system")
                for index, (category, code, name, extra) in enumerate(DEFAULT_SYSTEM_PARAMETERS, start=1)
            ])
        existing_configs = {item.key: item for item in (await db.scalars(select(SystemConfig))).all()}
        for key, config in DEFAULT_SYSTEM_CONFIGS.items():
            if key not in existing_configs:
                db.add(SystemConfig(key=key, **config, updated_by="system"))
            else:
                current_value = existing_configs[key].value or {}
                missing_defaults = {name: value for name, value in config["value"].items() if name not in current_value}
                if missing_defaults:
                    existing_configs[key].value = {**current_value, **missing_defaults}
        existing_menus = {item.key: item for item in (await db.scalars(select(SystemMenu))).all()}
        for key, parent_key, label, icon, sort_order in DEFAULT_SYSTEM_MENUS:
            if key not in existing_menus:
                db.add(SystemMenu(key=key, parent_key=parent_key, label=label, icon=icon, sort_order=sort_order, updated_by="system"))
            elif key in ORIGINAL_FINANCE_MENU_KEYS or key in ORIGINAL_ADMIN_MENU_KEYS or key in ORIGINAL_INVESTIGATION_MENU_KEYS or (
                key.startswith("customer-") and existing_menus[key].updated_by == "system"
            ) or (
                key in {"documents-official", "documents-my", "documents-company"}
                and existing_menus[key].updated_by == "system"
            ):
                existing_menus[key].parent_key = parent_key
                existing_menus[key].label = label
                existing_menus[key].icon = icon
                existing_menus[key].sort_order = sort_order
                existing_menus[key].is_visible = True
                existing_menus[key].is_active = True
        for key in LEGACY_FINANCE_MENU_KEYS:
            if key in existing_menus:
                existing_menus[key].is_visible = False
                existing_menus[key].is_active = False
        for key in LEGACY_ADMIN_MENU_KEYS:
            if key in existing_menus:
                existing_menus[key].is_visible = False
                existing_menus[key].is_active = False
        for key in LEGACY_INVESTIGATION_MENU_KEYS:
            if key in existing_menus:
                existing_menus[key].is_visible = False
                existing_menus[key].is_active = False
        for key in LEGACY_TASK_MENU_KEYS:
            if key in existing_menus:
                existing_menus[key].is_visible = False
                existing_menus[key].is_active = False
        existing_department_codes = set((await db.scalars(select(Department.code))).all())
        for index, (code, name) in enumerate(DEFAULT_DEPARTMENTS, start=1):
            if code not in existing_department_codes: db.add(Department(code=code, name=name, sort_order=index, created_by="system", updated_by="system"))
        existing_job_role_codes = set((await db.scalars(select(JobRole.code))).all())
        for index, (code, name, permissions) in enumerate(DEFAULT_JOB_ROLES, start=1):
            if code not in existing_job_role_codes: db.add(JobRole(code=code, name=name, permissions=permissions, sort_order=index, created_by="system", updated_by="system"))
        # 七类印章是合同用印流程所需的基础资料，不属于演示数据。这里只补缺，
        # 不覆盖管理员已经维护的保管人、位置、状态、用印次数等真实台账字段。
        existing_seal_assets = (await db.scalars(select(SealAsset))).all()
        seal_assets_by_type = {item.seal_type: item for item in existing_seal_assets}
        seal_assets_by_code = {item.code: item for item in existing_seal_assets}
        legacy_default_types = {"合同专用章": "合同章", "律师事务所专用章": "所函专用章"}
        for code, seal_type, location in REQUIRED_SEAL_ASSETS:
            if seal_type in seal_assets_by_type:
                continue
            existing_asset = seal_assets_by_code.get(code)
            if (
                existing_asset
                and existing_asset.name.startswith("申浩律师事务所")
                and legacy_default_types.get(existing_asset.seal_type) == seal_type
            ):
                existing_asset.name = f"申浩律师事务所{seal_type}"
                existing_asset.seal_type = seal_type
                seal_assets_by_type[seal_type] = existing_asset
                continue
            if existing_asset:
                # 默认编号被真实台账占用时不覆盖用户数据，换一个系统补缺编号。
                code = f"{code.rsplit('-', 1)[0]}-SYS-001"
            asset = SealAsset(
                code=code,
                name=f"申浩律师事务所{seal_type}",
                seal_type=seal_type,
                custodian="admin",
                location=location,
                remark="系统基础印章资料；管理员可在用印中心维护保管信息",
            )
            db.add(asset)
            seal_assets_by_type[seal_type] = asset
        record_count = await db.scalar(select(func.count()).select_from(BusinessRecord))
        if settings.seed_demo_data and not record_count:
            db.add_all(_seed_business_records())
            await db.flush()
        # SQLite 默认不强制外键，主动清理孤立流程记录，并为初始化数据补一条留痕。
        await db.execute(delete(WorkflowEvent).where(WorkflowEvent.record_id.not_in(select(BusinessRecord.id))))
        records = (await db.scalars(select(BusinessRecord))).all()
        original_customer = next(
            (record for record in records if record.module == "customer" and record.serial_no == "SHKH1810649"),
            None,
        )
        if settings.seed_demo_data and original_customer is None:
            original_customer = BusinessRecord(
                module="customer", serial_no="SHKH1810649", title="test", customer="test",
                status="正常", owner="admin", department="上海分所",
                data={
                    "source_person": "管理者", "customer_managers": ["管理者"],
                    "customer_type": "客户", "invoice_address": "test",
                    "customer_source": "管理者", "is_shared": "否",
                    "level": "立案客户", "is_assisted": "否",
                    "file_date": "2018-07-29", "last_contact_at": "2018-07-29",
                    "last_modified_date": "2018-07-29", "contact_count": 0,
                    "contract_count": 3, "civil_case_count": 2,
                    "agency_fee_due": 0, "official_fee_unreceived": -4000,
                },
            )
            db.add(original_customer)
            await db.flush()
            records.append(original_customer)
        elif settings.seed_demo_data and original_customer.title == "test" and (original_customer.data or {}).get("source_person") == "管理者":
            # Keep the historical read-only reference fixture complete across upgrades.
            # This branch is restricted to the app-owned demo row and never overwrites
            # non-empty values, so user-created customer data remains untouched.
            fixture_defaults = {
                "customer_type": "客户", "invoice_address": "test",
                "customer_source": "管理者", "is_shared": "否",
                "level": "立案客户", "is_assisted": "否",
            }
            fixture_data = dict(original_customer.data or {})
            for key, value in fixture_defaults.items():
                if not fixture_data.get(key):
                    fixture_data[key] = value
            original_customer.data = fixture_data
        original_contracts = [
            ("SHHT2610035", "test_合同", "审批中", {"contract_body": "律所"}),
            ("SHHT2510026", "test_合同", "审批中", {
                "contract_body": "律所", "official_paid": 0, "official_received": 4000,
                "official_unreceived": -4000, "official_loss": 0, "agency_total": 6000,
                "agency_received": 6000, "agency_due": 0, "other_total": 0,
                "other_paid": 0, "other_due": 0, "invoice_opened": 0,
                "invoice_should": 6000, "invoice_excess": 0,
            }),
            ("SHHT1810328", "test_合同", "已归档", {"contract_body": "律所"}),
        ]
        if not settings.seed_demo_data:
            original_contracts = []
        existing_contract_nos = {record.serial_no for record in records if record.module == "contract"}
        for serial_no, title, contract_status, data in original_contracts:
            if serial_no in existing_contract_nos:
                continue
            original_contract = BusinessRecord(
                module="contract", serial_no=serial_no, title=title, customer="test",
                status=contract_status, owner="admin", department="上海分所",
                data={
                    "type": "争议解决合同", "fee_type": "固定收费", "signed_at": "",
                    "source_person": "管理者", "amount": data.get("agency_total", 0),
                    "official_paid": 0, "official_received": 0, "official_unreceived": 0,
                    "official_loss": 0, "agency_total": 0, "agency_received": 0,
                    "agency_due": 0, "other_total": 0, "other_paid": 0,
                    "other_due": 0, "invoice_opened": 0, "invoice_should": 0,
                    "invoice_excess": 0, **data,
                },
            )
            db.add(original_contract)
            await db.flush()
            records.append(original_contract)
        original_cases = [
            ("SHMS2300502", "一审待客户回款", "上海台享餐饮管理有限公司", "长寿区娅娅小吃店", "重庆市自由贸易试验区人民法院", "（2023）渝0192民初10300号", "外部合作律师", "外部合作律师", "2023-12-29", 928, "结算规档任务", "本案SHMS2300502已到账超过30日,请尽快提交结算并归档.", "外部合作律师", "2025-09-26"),
            ("SHMS2400031", "一审判决结案", "中饮巴比食品股份有限公司", "高新区芭比特包包子铺", "成都高新技术产业开发区人民法院", "(2024)川0191民初18219号", "System", "刘波", "2026-02-12", 152, "结算规档任务", "本案SHMS2400031已到账超过30日,请尽快提交结算并归档.", "刘波", "2026-03-25"),
            ("SHMS2400065", "一审判决结案", "中饮巴比食品股份有限公司", "璧山区段世华面馆", "重庆市自由贸易试验区人民法院", "(2024)渝0192民初10299号", "System", "刘波", "2026-02-12", 152, "结算规档任务", "本案SHMS2400065已到账超过30日,请尽快提交结算并归档.", "刘波", "2026-03-25"),
            ("SHMS2500709A", "已归档", "上海天路人造草坪有限公司", "常州莱因人造草坪科技有限公司", "江苏省苏州市中级人民法院", "（2025）苏05民初1478号", "陶勇刚", "陶勇刚", "2026-07-08", 6, "结算归档一审和解结案", "结算归档", "陶亮", "2026-05-23"),
            ("SHMS2400317", "等待公证书", "珠海双喜电器股份有限公司", "义乌市热康日用品厂", "", "", "System", "", "2024-05-19", 786, "案件审核", "品管回复停止取证", "System", "2026-07-15"),
            ("SH171000067", "一审待客户回款", "珠海格力电器股份有限公司", "常州市天宁区天宁正和电子经营部", "常州市天宁区人民法院", "（2018）苏0402民初4642号", "崔铧尹", "李晓岩,朱莹", "2023-03-15", 1217, "案件跟进回款", "这几个格力案件，现在什么情况？", "陶国南", "2026-07-16"),
            ("SH171000093", "一审待客户回款", "珠海格力电器股份有限公司", "常州市钱达电器经营部", "常州市天宁区人民法院", "（2018）苏0402民初4643号", "崔铧尹", "李晓岩,朱莹", "2023-02-13", 1247, "案件跟进回款", "这几个格力案件，现在什么情况？", "陶国南", "2026-07-16"),
            ("SHMS2500647", "文书准备", "九牧王股份有限公司", "亳州市谯城区衣家园服装批发店（个体工商户）", "利辛县人民法院", "", "李佳妮", "张美莹", "2026-01-07", 188, "案件审核", "", "李佳妮", "2026-07-16"),
            ("SH191000297", "执行终本", "中粮集团有限公司", "上海联华快客便利有限公司习勤店,蓬莱华夏葡园酒业有限公司,上海联华快客便利有限公司", "上海市徐汇区人民法院", "（2024）沪0104执7123号、（2025）沪0104执异495号", "陶勇刚", "李佳妮", "2025-11-24", 232, "终本案件，先到账的先结算发提成，后面还要继续追讨", "", "审核管理（赵媛）", "2026-07-16"),
            ("SHMS2500149", "文书准备", "广东三雄极光照明股份有限公司", "王勇,上海寻梦信息技术有限公司", "上海市长宁区人民法院", "", "王晓英", "郝蕴", "2025-08-13", 335, "文书审核", "已修改上传系统，是否可以盖章", "郝蕴", "2026-07-16"),
        ]
        if not settings.seed_demo_data:
            original_cases = []
        existing_case_nos = {record.serial_no for record in records if record.module == "case"}
        for case_item in original_cases:
            serial_no, case_status, plaintiff, defendant, court, court_case_no, lawyer, assistant, changed_at, days, task_name, task_content, task_handler, task_time = case_item
            if serial_no in existing_case_nos:
                continue
            original_case = BusinessRecord(
                module="case", serial_no=serial_no, title=f"{plaintiff}诉{defendant}", customer=plaintiff,
                status=case_status, owner="admin", department="上海分所",
                data={"case_type": "民事案件", "plaintiff": plaintiff, "opponent": defendant,
                      "court": court, "court_case_no": court_case_no, "hearing_lawyer": lawyer,
                      "handling_lawyers": [lawyer] if lawyer else [], "assistant": assistant,
                      "phase_changed_at": changed_at, "phase_days": days, "task_name": task_name,
                      "task_content": task_content, "task_handler": task_handler, "task_time": task_time},
            )
            db.add(original_case); await db.flush(); records.append(original_case)
        event_record_ids = set((await db.scalars(select(WorkflowEvent.record_id).distinct())).all())
        for record in records:
            if settings.seed_demo_data and record.id not in event_record_ids:
                db.add(WorkflowEvent(record_id=record.id, action="系统初始化", to_status=record.status, operator="system", comment="初始化示例业务数据"))
        if settings.seed_demo_data and not await db.scalar(select(func.count()).select_from(ReceivablePlan)):
            contracts = {record.serial_no: record for record in records if record.module == "contract"}
            if contracts.get("HT2026070018"):
                db.add_all([
                    ReceivablePlan(contract_record_id=contracts["HT2026070018"].id, phase="合同签订首付款", due_date=date(2026, 7, 20), amount=140000, received_amount=0, status="待收款", payer=contracts["HT2026070018"].customer),
                    ReceivablePlan(contract_record_id=contracts["HT2026070018"].id, phase="项目办结尾款", due_date=date(2026, 12, 20), amount=140000, received_amount=0, status="待收款", payer=contracts["HT2026070018"].customer),
                ])
            if contracts.get("HT2026060097"):
                db.add(ReceivablePlan(contract_record_id=contracts["HT2026060097"].id, phase="年度顾问费", due_date=date(2026, 6, 30), amount=120000, received_amount=80000, status="部分收款", payer=contracts["HT2026060097"].customer))
        if settings.seed_demo_data and not await db.scalar(select(func.count()).select_from(HearingSchedule)):
            cases = {record.serial_no: record for record in records if record.module == "case"}
            if cases.get("SH191000382B"):
                db.add(HearingSchedule(case_record_id=cases["SH191000382B"].id, hearing_date=date(2026, 7, 15), hearing_time="09:00", court="上海市宝山区人民法院", courtroom="第六法庭", hearing_type="一审开庭", hearing_lawyer="陈名涛"))
            if cases.get("SHMS2600387"):
                db.add(HearingSchedule(case_record_id=cases["SHMS2600387"].id, hearing_date=date(2026, 7, 20), hearing_time="14:00", court="杭州市余杭区人民法院", courtroom="第二法庭", hearing_type="证据交换", hearing_lawyer="陶勇刚"))
        if settings.seed_demo_data and not await db.scalar(select(func.count()).select_from(DocumentTemplate)):
            db.add_all([
                DocumentTemplate(name="民事起诉状", category="诉讼文书", version="2026.1", description="知识产权民事案件起诉状标准模板", fields=["原告", "被告", "诉讼请求", "事实与理由"]),
                DocumentTemplate(name="律师函", category="非诉文书", version="2026.1", description="侵权告知及停止侵权律师函", fields=["委托人", "收函人", "事实", "法律意见"]),
                DocumentTemplate(name="案件归档目录", category="归档文书", version="2026.1", description="案件归档材料目录标准模板", fields=["案号", "客户", "材料清单", "归档日期"]),
            ])
        await db.flush()
        assets_by_type = {item.seal_type: item for item in (await db.scalars(select(SealAsset))).all()}
        for record in records:
            if record.module == "seal" and not (record.data or {}).get("seal_asset_id"):
                asset = assets_by_type.get((record.data or {}).get("seal_type")) or assets_by_type.get("公章")
                if asset:
                    record.data = {**(record.data or {}), "seal_asset_id": asset.id, "seal_name": asset.name}
        approval_contracts = [record for record in records if settings.seed_demo_data and record.module == "contract" and record.status == "审批中"]
        for contract in approval_contracts:
            if not await db.scalar(select(func.count()).select_from(ContractApprovalStep).where(ContractApprovalStep.contract_record_id == contract.id)):
                db.add(ContractApprovalStep(contract_record_id=contract.id, step_order=1, approver="admin", status="待审批", comment="历史合同补充默认审批节点"))
        await db.commit()
    rule_task = asyncio.create_task(_business_rule_loop())
    try:
        yield
    finally:
        rule_task.cancel()
        with suppress(asyncio.CancelledError):
            await rule_task


app = FastAPI(title=settings.app_name, version="0.1.0", lifespan=lifespan)


@app.exception_handler(RequestValidationError)
async def request_validation_error_handler(_, exc: RequestValidationError):
    details = []
    for error in exc.errors():
        location = ".".join(str(part) for part in error.get("loc", []) if part != "body")
        details.append(f"{location}：{error.get('msg', '参数格式错误')}" if location else error.get("msg", "参数格式错误"))
    logger.warning("Request validation failed: %s", "；".join(details))
    return JSONResponse(status_code=422, content={"detail": "；".join(details)})
app.add_middleware(CORSMiddleware, allow_origins=["http://localhost", "http://127.0.0.1"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])
UPLOAD_ROOT = Path(__file__).resolve().parent.parent / "uploads"
UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)


class DifyRequest(BaseModel):
    query: str
    conversation_id: str | None = None


class AgentDocumentInput(BaseModel):
    template_id: int
    record_id: int | None = None
    title: str
    instruction: str = ""


class AgentDocumentUpdate(BaseModel):
    title: str | None = None
    content: str | None = None


class AgentDocumentConfirmInput(BaseModel):
    comment: str = Field(default="", max_length=1000)


class RecordInput(BaseModel):
    module: str
    serial_no: str
    title: str
    customer: str = ""
    status: str = "草稿"
    owner: str = "管理者"
    department: str = "上海分所"
    description: str = ""
    data: dict = Field(default_factory=dict)


class RecordUpdate(BaseModel):
    title: str | None = None
    customer: str | None = None
    status: str | None = None
    owner: str | None = None
    department: str | None = None
    description: str | None = None
    data: dict | None = None


class TransitionInput(BaseModel):
    to_status: str
    comment: str = ""


class DocumentTransitionInput(BaseModel):
    to_status: str
    action_date: date = Field(default_factory=date.today)
    handler: str = ""
    archive_no: str = ""
    archive_location: str = ""
    comment: str = ""


class OfficialDocumentProcessInput(BaseModel):
    record_ids: list[int] = Field(min_length=1, max_length=100)
    processed: bool
    comment: str = Field(default="", max_length=1000)


class OfficialDocumentReceiptDateInput(BaseModel):
    """Dedicated command for the legacy official-receipt date correction."""
    record_ids: list[int] = Field(min_length=1, max_length=100)
    document_date: date
    comment: str = Field(default="", max_length=1000)


class HrTransitionInput(BaseModel):
    to_status: str
    effective_date: date = Field(default_factory=date.today)
    reason: str = ""
    handover_to: str = ""
    comment: str = ""


class HrSubrecordInput(BaseModel):
    kind: str = Field(pattern="^(leave|matter|commission)$")
    data: dict = Field(default_factory=dict)


class HrSubrecordUpdate(BaseModel):
    data: dict


class DepartmentInput(BaseModel):
    code: str = Field(min_length=1, max_length=64)
    name: str = Field(min_length=1, max_length=128)
    manager: str = Field(default="", max_length=64)
    sort_order: int = Field(default=0, ge=0, le=99999)
    is_active: bool = True


class DepartmentUpdate(BaseModel):
    code: str | None = Field(default=None, min_length=1, max_length=64)
    name: str | None = Field(default=None, min_length=1, max_length=128)
    manager: str | None = Field(default=None, max_length=64)
    sort_order: int | None = Field(default=None, ge=0, le=99999)
    is_active: bool | None = None


class JobRoleInput(BaseModel):
    code: str = Field(min_length=1, max_length=64)
    name: str = Field(min_length=1, max_length=128)
    permissions: list[str] = Field(default_factory=list, max_length=50)
    description: str = Field(default="", max_length=1000)
    sort_order: int = Field(default=0, ge=0, le=99999)
    is_active: bool = True


class JobRoleUpdate(BaseModel):
    code: str | None = Field(default=None, min_length=1, max_length=64)
    name: str | None = Field(default=None, min_length=1, max_length=128)
    permissions: list[str] | None = Field(default=None, max_length=50)
    description: str | None = Field(default=None, max_length=1000)
    sort_order: int | None = Field(default=None, ge=0, le=99999)
    is_active: bool | None = None


class WarehouseBorrowInput(BaseModel):
    borrower: str = Field(min_length=1, max_length=64)
    due_date: date
    purpose: str = ""
    comment: str = ""


class WarehouseReturnInput(BaseModel):
    comment: str = ""


class WarehouseReturnConfirmInput(BaseModel):
    condition: str = Field(default="完好", min_length=1, max_length=64)
    location: str = ""
    comment: str = ""


class WarehouseScrapInput(BaseModel):
    reason: str = Field(min_length=2, max_length=1000)


class WarehouseEvidenceInput(BaseModel):
    serial_no: str = Field(min_length=1, max_length=128)
    warehouse: str = Field(min_length=1, max_length=128)
    location: str = Field(min_length=1, max_length=128)
    notary_no: str = Field(default="", max_length=128)
    case_no: str = Field(default="", max_length=128)
    shop_name: str = Field(min_length=1, max_length=255)
    investigator: str = Field(min_length=1, max_length=128)
    notary_office: str = Field(default="", max_length=255)
    rights_holder: str = Field(min_length=1, max_length=255)
    evidence_date: date
    description: str = Field(default="", max_length=1000)


class WarehouseEvidenceCheckInInput(BaseModel):
    warehouse: str = Field(min_length=1, max_length=128)
    location: str = Field(min_length=1, max_length=128)
    comment: str = Field(default="", max_length=1000)


class WarehouseEvidenceCheckOutInput(BaseModel):
    recipient: str = Field(min_length=1, max_length=128)
    purpose: str = Field(min_length=1, max_length=500)
    comment: str = Field(default="", max_length=1000)


class WarehouseEvidenceRecheckInInput(BaseModel):
    warehouse: str = Field(min_length=1, max_length=128)
    location: str = Field(min_length=1, max_length=128)
    condition: str = Field(default="完好", min_length=1, max_length=128)
    comment: str = Field(default="", max_length=1000)


class WarehouseEvidenceDestroyInput(BaseModel):
    reason: str = Field(min_length=2, max_length=1000)


class ReceivableInput(BaseModel):
    contract_record_id: int
    phase: str
    due_date: date
    amount: float = Field(gt=0)
    payer: str = ""
    remark: str = ""


class ReceivePaymentInput(BaseModel):
    amount: float = Field(gt=0)
    comment: str = ""


class NotaryReviewInput(BaseModel):
    approved: bool
    comment: str = ""
    case_type: str = "民事案件"
    court: str = ""


class ClueReviewInput(BaseModel):
    approved: bool
    comment: str = Field(min_length=2, max_length=1000)


class ClueCollectionInput(BaseModel):
    collected_at: date
    notary_institution: str = Field(min_length=2, max_length=255)
    comment: str = ""


class EvidenceCreateInput(BaseModel):
    title: str
    owner: str
    source: str = "调查取证"
    description: str = ""


class NotaryCertificateInput(BaseModel):
    certificate_no: str = Field(min_length=2, max_length=128)
    issued_date: date
    storage_location: str = Field(min_length=2, max_length=255)
    physical_received: bool = False
    comment: str = ""


class InvestigationTaskInput(BaseModel):
    title: str
    owner: str
    deadline: date
    priority: str = "普通"
    parent_task_id: int | None = None
    description: str = ""


class BatchClueCaseInput(BaseModel):
    clue_ids: list[int] = Field(min_length=1, max_length=100)
    contract_record_id: int
    case_type: str = "民事案件"
    court: str = ""


class InvestigationAssignmentInput(BaseModel):
    investigator: str = Field(min_length=1, max_length=128)
    comment: str = ""


class InvestigationBatchDeleteInput(BaseModel):
    record_ids: list[int] = Field(min_length=1, max_length=100)
    comment: str = ""


class InvestigationFeeInput(BaseModel):
    amount: float = Field(gt=0, le=100000000)
    fee_type: str = Field(min_length=1, max_length=64)
    description: str = ""


INVESTIGATION_RECORD_MODULES = {"investigation", "clue", "notary", "evidence"}
INVESTIGATION_CREATE_STATUS_BY_MODULE = {
    "investigation": "待分配",
    "clue": "草稿",
    "notary": "待审核",
    "evidence": "待整理",
}
INVESTIGATION_EDIT_DATA_FIELDS = {
    "region", "address", "right_type", "deadline", "priority", "platform",
    "product", "source", "infringement_method",
}


class CaseAssignmentInput(BaseModel):
    customer_manager: str = ""
    hearing_lawyer: str
    handling_lawyers: list[str] = Field(default_factory=list)
    assistant: str = ""
    comment: str = ""


class CaseBatchUpdateInput(BaseModel):
    case_ids: list[int] = Field(default_factory=list, max_length=100)
    case_nos: list[str] = Field(default_factory=list, max_length=100)
    hearing_lawyer: str | None = Field(default=None, max_length=128)
    handling_lawyers: list[str] | None = Field(default=None, max_length=20)
    assistant: str | None = Field(default=None, max_length=128)
    case_stage: str | None = Field(default=None, max_length=128)
    comment: str = Field(default="", max_length=500)


class CaseReminderInput(BaseModel):
    reminder_date: date
    deadline: date
    content: str = Field(min_length=1, max_length=1000)


class CaseLogInput(BaseModel):
    content: str = Field(min_length=1, max_length=1000)


class CaseBatchFeeInput(BaseModel):
    case_ids: list[int] = Field(min_length=1, max_length=100)
    amount: float = Field(gt=0, le=100000000)
    expense_scope: str = Field(pattern="^(律所|平台|内部)$")
    expense_subtype: str = Field(pattern="^(官费|第三方费用|代理费|其他费用|内部费用)$")
    handler: str = Field(default="", max_length=128)
    description: str = Field(default="", max_length=1000)


class AttachmentBatchInput(BaseModel):
    attachment_ids: list[int] = Field(min_length=1, max_length=100)


class CaseAttachmentRenameInput(BaseModel):
    original_name: str = Field(min_length=1, max_length=255)


class CaseProgressInput(BaseModel):
    first_instance_court: str = ""
    first_instance_case_no: str = ""
    courtroom: str = ""
    judge: str = ""
    clerk: str = ""
    judgment_date: date | None = None
    judgment_document_no: str = ""
    second_instance_court: str = ""
    second_instance_case_no: str = ""
    comment: str = ""


class CaseCreateInput(BaseModel):
    contract_record_id: int = Field(gt=0)
    serial_no: str = Field(default="", max_length=64)
    title: str = Field(min_length=1, max_length=256)
    status: str = "新案待分配"
    owner: str = Field(default="", max_length=128)
    case_type: str = Field(default="刑事案件", min_length=1, max_length=64)
    opponent: str = ""
    court: str = ""
    client_position: str = ""
    cause_or_charge: str = Field(default="", max_length=256)
    right_type: str = ""
    counsel_type: str = ""
    counsel_start: date | None = None
    counsel_end: date | None = None
    handling_lawyers: list[str] = Field(default_factory=list, max_length=20)
    assistant: str = ""
    investigator: str = ""
    investigation_clue: str = ""
    party_type: str = ""
    party_id_type: str = ""
    party_id_no: str = ""
    party_contact: str = ""
    party_phone: str = ""
    party_address: str = ""
    legal_representative: str = ""
    party_remark: str = ""
    court_case_no: str = ""
    judge: str = ""
    judge_phone: str = ""
    filing_date: date | None = None
    hearing_date: date | None = None
    hearing_time: str = ""
    courtroom: str = ""
    judicial_remark: str = ""
    description: str = ""


class CaseLitigantsInput(BaseModel):
    plaintiffs: list[str] = Field(default_factory=list, max_length=50)
    plaintiff_agents: list[str] = Field(default_factory=list, max_length=50)
    defendants: list[str] = Field(default_factory=list, max_length=50)
    defendant_agents: list[str] = Field(default_factory=list, max_length=50)
    third_parties: list[str] = Field(default_factory=list, max_length=50)
    third_party_agents: list[str] = Field(default_factory=list, max_length=50)
    comment: str = Field(default="", max_length=500)


class CaseCreationCompleteInput(BaseModel):
    comment: str = Field(default="", max_length=500)


class CaseCounselBasicInput(BaseModel):
    title: str = Field(min_length=1, max_length=256)
    counsel_type: str = Field(min_length=1, max_length=128)
    counsel_start: date
    counsel_end: date
    handling_lawyers: list[str] = Field(min_length=1, max_length=20)
    assistant: str = Field(default="", max_length=128)
    comment: str = Field(default="", max_length=500)


class CaseNormalBasicInput(BaseModel):
    """Old-system type-specific basic-information editor for ordinary cases.

    This deliberately does not reuse the legal-counsel endpoint: ordinary cases
    have a case phase, cause/charge and clue/investigator fields instead.
    """
    customer_record_id: int = Field(gt=0)
    title: str = Field(min_length=1, max_length=256)
    case_phase: str = Field(min_length=1, max_length=64)
    cause_or_charge: str = Field(min_length=1, max_length=256)
    handling_lawyers: list[str] = Field(min_length=1, max_length=20)
    assistant: str = Field(default="", max_length=128)
    business_owner: str = Field(default="", max_length=128)
    investigator: str = Field(default="", max_length=128)
    investigation_clue_ids: list[int] = Field(default_factory=list, max_length=50)
    right_type: str = Field(default="", max_length=128)
    comment: str = Field(default="", max_length=500)


class CaseArbitrationBasicInput(BaseModel):
    """Dedicated legacy arbitration basic-information branch (not normal/counsel)."""
    customer_record_id: int = Field(gt=0)
    title: str = Field(min_length=1, max_length=256)
    case_phase: str = Field(min_length=1, max_length=64)
    cause_or_charge: str = Field(min_length=1, max_length=256)
    handling_lawyers: list[str] = Field(min_length=1, max_length=20)
    assistant: str = Field(default="", max_length=128)
    investigator: str = Field(default="", max_length=128)
    investigation_clue_ids: list[int] = Field(default_factory=list, max_length=50)
    comment: str = Field(default="", max_length=500)


class CriminalPublicSecurityMaintenanceInput(BaseModel):
    public_security_name: str = Field(default="", max_length=256)
    public_security_case_no: str = Field(default="", max_length=128)
    public_security_address: str = Field(default="", max_length=500)
    public_security_operator: str = Field(default="", max_length=128)
    comment: str = Field(default="", max_length=500)


class CriminalProcuratorateMaintenanceInput(BaseModel):
    first_procuratorate_name: str = Field(default="", max_length=256); first_procuratorate_case_no: str = Field(default="", max_length=128); first_procuratorate_operator: str = Field(default="", max_length=128)
    second_procuratorate_name: str = Field(default="", max_length=256); second_procuratorate_case_no: str = Field(default="", max_length=128); second_procuratorate_operator: str = Field(default="", max_length=128)
    retrial_procuratorate_name: str = Field(default="", max_length=256); retrial_procuratorate_case_no: str = Field(default="", max_length=128); retrial_procuratorate_operator: str = Field(default="", max_length=128)
    comment: str = Field(default="", max_length=500)


class CriminalCourtMaintenanceInput(BaseModel):
    first_court_name: str = Field(default="", max_length=256); first_court_case_no: str = Field(default="", max_length=128); first_court_courtroom: str = Field(default="", max_length=128); first_court_judge: str = Field(default="", max_length=128); first_court_clerk: str = Field(default="", max_length=128)
    second_court_name: str = Field(default="", max_length=256); second_court_case_no: str = Field(default="", max_length=128); second_court_courtroom: str = Field(default="", max_length=128); second_court_judge: str = Field(default="", max_length=128); second_court_clerk: str = Field(default="", max_length=128)
    retrial_court_name: str = Field(default="", max_length=256); retrial_court_case_no: str = Field(default="", max_length=128); retrial_court_courtroom: str = Field(default="", max_length=128); retrial_court_judge: str = Field(default="", max_length=128); retrial_court_clerk: str = Field(default="", max_length=128)
    comment: str = Field(default="", max_length=500)


class CounselCaseSearchInput(BaseModel):
    scope: str = "company"
    customer: str = Field(default="", max_length=256)
    serial_no: str = Field(default="", max_length=128)
    keyword: str = Field(default="", max_length=256)
    counsel_start: date | None = None
    counsel_end: date | None = None
    counsel_type: str = Field(default="", max_length=128)
    case_status: str = Field(default="", max_length=64)
    handling_lawyer: str = Field(default="", max_length=128)
    assistant: str = Field(default="", max_length=128)
    document_name: str = Field(default="", max_length=255)
    sort_order: str = "updated_desc"
    page: int = Field(default=1, ge=1)
    page_size: int = Field(default=10, ge=1, le=200)
    selected_ids: list[int] = Field(default_factory=list, max_length=200)
    selected_only: bool = False


class CaseJudicialInput(BaseModel):
    # 当前网页端使用的兼容法院字段。
    court: str = Field(default="", max_length=256)
    court_case_no: str = Field(default="", max_length=128)
    courtroom: str = Field(default="", max_length=128)
    judge: str = Field(default="", max_length=128)
    clerk: str = Field(default="", max_length=128)
    judge_phone: str = Field(default="", max_length=64)
    filing_date: date | None = None
    hearing_date: date | None = None
    hearing_time: str = Field(default="", max_length=8)
    judicial_remark: str = Field(default="", max_length=1000)
    description: str = Field(default="", max_length=4000)

    # 刑事案件原站脚本中出现的公安机关和三级检察院字段。
    public_security_name: str = Field(default="", max_length=256)
    public_security_case_no: str = Field(default="", max_length=128)
    public_security_address: str = Field(default="", max_length=500)
    public_security_phone: str = Field(default="", max_length=64)
    public_security_operator: str = Field(default="", max_length=128)
    first_procuratorate_name: str = Field(default="", max_length=256)
    first_procuratorate_case_no: str = Field(default="", max_length=128)
    first_procuratorate_address: str = Field(default="", max_length=500)
    first_procuratorate_phone: str = Field(default="", max_length=64)
    first_procuratorate_operator: str = Field(default="", max_length=128)
    second_procuratorate_name: str = Field(default="", max_length=256)
    second_procuratorate_case_no: str = Field(default="", max_length=128)
    second_procuratorate_address: str = Field(default="", max_length=500)
    second_procuratorate_phone: str = Field(default="", max_length=64)
    second_procuratorate_operator: str = Field(default="", max_length=128)
    retrial_procuratorate_name: str = Field(default="", max_length=256)
    retrial_procuratorate_case_no: str = Field(default="", max_length=128)
    retrial_procuratorate_address: str = Field(default="", max_length=500)
    retrial_procuratorate_phone: str = Field(default="", max_length=64)
    retrial_procuratorate_operator: str = Field(default="", max_length=128)

    # 原站法院页可勾选一审、二审和再审；字段分别保存，避免后续互相覆盖。
    first_court_enabled: bool = False
    first_court_name: str = Field(default="", max_length=256)
    first_court_case_no: str = Field(default="", max_length=128)
    first_court_courtroom: str = Field(default="", max_length=128)
    first_court_judge: str = Field(default="", max_length=128)
    first_court_clerk: str = Field(default="", max_length=128)
    first_court_filing_date: date | None = None
    first_court_hearing_date: date | None = None
    second_court_enabled: bool = False
    second_court_name: str = Field(default="", max_length=256)
    second_court_case_no: str = Field(default="", max_length=128)
    second_court_courtroom: str = Field(default="", max_length=128)
    second_court_judge: str = Field(default="", max_length=128)
    second_court_clerk: str = Field(default="", max_length=128)
    second_court_filing_date: date | None = None
    second_court_hearing_date: date | None = None
    retrial_court_enabled: bool = False
    retrial_court_name: str = Field(default="", max_length=256)
    retrial_court_case_no: str = Field(default="", max_length=128)
    retrial_court_courtroom: str = Field(default="", max_length=128)
    retrial_court_judge: str = Field(default="", max_length=128)
    retrial_court_clerk: str = Field(default="", max_length=128)
    retrial_court_filing_date: date | None = None
    retrial_court_hearing_date: date | None = None


class HearingInput(BaseModel):
    case_record_id: int
    hearing_date: date
    hearing_time: str
    court: str
    courtroom: str = ""
    hearing_type: str = "开庭"
    hearing_lawyer: str
    remark: str = ""


class ArchiveCheckInput(BaseModel):
    case_closed: bool = False
    fees_settled: bool = False
    documents_complete: bool = False
    finance_complete: bool = False
    archive_no: str = ""
    paper_archive_location: str = ""
    paper_volume_count: int = Field(default=1, ge=1, le=999)
    comment: str = ""
    submit: bool = False


class ArchiveReviewInput(BaseModel):
    approved: bool
    comment: str = Field(min_length=2, max_length=1000)


class TaskInput(BaseModel):
    title: str
    customer: str = ""
    owner: str
    deadline: date
    priority: str = "普通"
    source: str = "日常任务"
    collaborators: list[str] = Field(default_factory=list, max_length=20)
    case_no: str = ""
    description: str = ""


class TaskHandoffInput(BaseModel):
    recipient: str
    comment: str = ""


class TaskActionInput(BaseModel):
    comment: str = ""


class TaskExceptionRequestInput(BaseModel):
    action: str = Field(pattern="^(挂起|取消)$")
    reason: str = Field(min_length=2, max_length=1000)


class TaskExceptionReviewInput(BaseModel):
    approved: bool
    comment: str = Field(default="", max_length=1000)


class TaskBatchUpdateInput(BaseModel):
    task_ids: list[int] = Field(min_length=1, max_length=100)
    owner: str | None = Field(default=None, max_length=128)
    deadline: date | None = None
    priority: str | None = None
    comment: str = Field(default="", max_length=1000)


class TaskBatchLifecycleInput(BaseModel):
    task_ids: list[int] = Field(min_length=1, max_length=100)
    action: str = Field(pattern="^(accept|complete|handoff|withdraw)$")
    recipient: str = Field(default="", max_length=128)
    comment: str = Field(default="", max_length=1000)


class TaskBatchReadInput(BaseModel):
    """Selected task rows on the personal unread-message page."""
    task_ids: list[int] = Field(min_length=1, max_length=100)


class TemplateInput(BaseModel):
    name: str
    category: str
    version: str = "1.0"
    description: str = ""
    fields: list[str] = Field(default_factory=list)


class TemplateUpdate(BaseModel):
    name: str | None = None
    category: str | None = None
    version: str | None = None
    description: str | None = None
    fields: list[str] | None = None
    is_active: bool | None = None


class FinanceFeeInput(BaseModel):
    title: str
    customer: str = ""
    amount: float
    fee_type: str
    expense_scope: str | None = Field(default=None, pattern="^(律所|平台|内部)$")
    expense_subtype: str | None = Field(default=None, pattern="^(官费|第三方费用|代理费|其他费用|内部费用)$")
    case_no: str = ""
    handler: str
    court: str = ""
    document_no: str = ""
    payee: str = ""
    description: str = ""
    contract_record_id: int | None = None
    case_record_id: int | None = None


class FinanceActionInput(BaseModel):
    comment: str = ""


class FinanceSettlementMarkInput(BaseModel):
    fee_ids: list[int] = Field(min_length=1, max_length=100)
    comment: str = Field(default="", max_length=500)


class FinanceFeeReviewInput(BaseModel):
    approved: bool
    comment: str = Field(default="", max_length=1000)


class FinanceFeeBatchReviewInput(FinanceFeeReviewInput):
    fee_ids: list[int] = Field(min_length=1, max_length=100)


class FinancePaymentPackagePreviewInput(BaseModel):
    fee_ids: list[int] = Field(min_length=1, max_length=100)


class FinancePaymentPackageCreateInput(FinancePaymentPackagePreviewInput):
    package_no: str = Field(pattern=r"^P\d{6}-\d{8}$")
    comment: str = Field(default="", max_length=500)


class FinancePaymentPackageWriteoffInput(BaseModel):
    amount: float
    paid_date: date
    payment_method: str
    invoice_no: str = Field(min_length=1, max_length=128)
    remark: str = Field(default="", max_length=500)


class InvoiceApplicationInput(BaseModel):
    customer: str
    case_no: str = ""
    amount: float = Field(gt=0)
    invoice_title: str
    taxpayer_id: str
    invoice_phone: str = ""
    bank_account: str = ""
    bank_name: str = ""
    invoice_address: str = ""
    extra_amount: float = Field(default=0, ge=0)
    invoice_type: str = "增值税普通发票"
    invoice_content: str = "法律服务费"
    delivery_method: str = "电子发票"
    recipient: str = ""
    recipient_phone: str = ""
    email: str = ""
    delivery_address: str = ""
    remark: str = ""
    contract_record_id: int | None = None
    case_record_id: int | None = None
    case_fee_ids: list[int] = Field(default_factory=list, max_length=100)


class InvoiceIssueInput(BaseModel):
    invoice_no: str = Field(min_length=3, max_length=128)
    invoice_date: date
    invoice_holder: str = Field(default="", max_length=128)
    extra_amount: float = Field(default=0, ge=0)
    comment: str = ""


class InvoiceVoidInput(BaseModel):
    reason: str = Field(min_length=2, max_length=1000)


class InvoiceNumberChangeInput(BaseModel):
    invoice_no: str = Field(min_length=1, max_length=128)


class InvoiceDateChangeInput(BaseModel):
    application_date: date
    invoice_date: date


class FinanceReviewInput(BaseModel):
    approved: bool
    comment: str = Field(min_length=2, max_length=1000)


class LitigationRefundInput(BaseModel):
    fee_record_id: int | None = Field(default=None, ge=1)
    customer: str
    case_no: str
    court: str
    original_payment_no: str
    amount: float = Field(gt=0)
    applicant: str
    refund_account_name: str
    refund_bank: str
    refund_account: str
    expected_date: date | None = None
    reason: str = "诉讼费退费"
    remark: str = ""


class RefundCompleteInput(BaseModel):
    actual_date: date
    voucher_no: str = Field(min_length=2, max_length=128)
    comment: str = ""


class FinanceTransactionInput(BaseModel):
    finance_record_id: int | None = None
    transaction_type: str
    amount: float = Field(gt=0)
    transaction_date: date
    voucher_no: str = ""
    counterparty: str = ""
    remark: str = ""


class FinanceWriteoffInput(BaseModel):
    voucher_no: str = Field(min_length=2, max_length=128)
    comment: str = ""


class IncomingPaymentInput(BaseModel):
    received_date: date
    amount: float = Field(gt=0)
    payer_name: str = Field(min_length=2, max_length=255)
    bank_reference: str = Field(min_length=2, max_length=128)
    remark: str = ""


class IncomingPaymentClaimInput(BaseModel):
    customer: str = Field(min_length=2, max_length=255)
    comment: str = ""


class IncomingPaymentSettlementItem(BaseModel):
    fee_record_id: int | None = None
    fee_type: str = Field(min_length=1, max_length=64)
    amount: float = Field(gt=0)
    settlement_amount: float = Field(ge=0)
    archive_fee: float = Field(ge=0)


class IncomingPaymentAllocationItem(BaseModel):
    receivable_plan_id: int
    amount: float = Field(gt=0)
    case_no: str = ""
    payment_method: str = Field(default="", max_length=64)
    settlement_items: list[IncomingPaymentSettlementItem] = Field(default_factory=list, max_length=50)


class IncomingPaymentAllocateInput(BaseModel):
    allocations: list[IncomingPaymentAllocationItem] = Field(min_length=1, max_length=50)
    comment: str = ""


class FinanceSettlementApplyInput(BaseModel):
    receipt_ids: list[int] = Field(min_length=1, max_length=100)
    comment: str = ""


class FinanceSettlementReviewInput(BaseModel):
    application_ids: list[int] = Field(min_length=1, max_length=100)
    approved: bool
    comment: str = Field(default="", max_length=2000)


class FinanceSettlementPaymentInput(BaseModel):
    application_ids: list[int] = Field(min_length=1, max_length=100)
    action: str = Field(pattern="^(paid|rollback)$")
    comment: str = Field(default="", max_length=2000)


class FinanceSettlementReapplyInput(BaseModel):
    application_ids: list[int] = Field(min_length=1, max_length=100)
    comment: str = Field(min_length=1, max_length=2000)


class ArchiveSettlementPaymentReviewInput(BaseModel):
    settlement_ids: list[str] = Field(min_length=1, max_length=100)
    approved: bool
    comment: str = Field(default="", max_length=2000)


class ArchiveSettlementRollbackInput(BaseModel):
    record_ids: list[int] = Field(min_length=1, max_length=100)
    comment: str = Field(min_length=1, max_length=2000)


class ArchiveSettlementRejectedActionInput(BaseModel):
    record_ids: list[int] = Field(min_length=1, max_length=100)
    comment: str = Field(default="", max_length=2000)


class ReconciliationInput(BaseModel):
    period_type: str
    date_from: date
    date_to: date
    discrepancy_amount: float = 0
    remark: str = ""


class SystemUserInput(BaseModel):
    username: str = Field(min_length=3, max_length=64)
    display_name: str = Field(min_length=1, max_length=64)
    department: str = Field(default="上海分所", min_length=1, max_length=64)
    password: str = Field(min_length=8, max_length=128)
    role: str = "user"
    is_active: bool = True
    must_change_password: bool = False
    profile: dict = Field(default_factory=dict)


class SystemUserUpdate(BaseModel):
    username: str | None = Field(default=None, min_length=3, max_length=64)
    display_name: str | None = Field(default=None, min_length=1, max_length=64)
    department: str | None = Field(default=None, min_length=1, max_length=64)
    role: str | None = None
    is_active: bool | None = None
    password: str | None = Field(default=None, min_length=8, max_length=128)
    profile: dict | None = None


class SystemUserPasswordResetInput(BaseModel):
    """Administrator-issued one-time password for an existing account."""

    new_password: str = Field(min_length=8, max_length=128)


class HrEmployeeUpdateInput(BaseModel):
    username: str = Field(min_length=3, max_length=64)
    display_name: str = Field(min_length=1, max_length=64)
    department: str = Field(min_length=1, max_length=64)
    role: str
    position: str = Field(min_length=1, max_length=128)
    is_active: bool = True
    email: str = Field(default="", max_length=128)
    mobile: str = Field(default="", max_length=32)
    office_phone: str = Field(default="", max_length=32)
    joined_at: date
    left_at: date | None = None
    data: dict = Field(default_factory=dict)


class HrEmployeeCreateInput(BaseModel):
    # Only an "employee account" has a system-login counterpart.  Keeping this
    # optional lets HR retain customer/external personnel files without creating
    # a privileged or orphaned system account by accident.
    username: str = Field(default="", max_length=64)
    display_name: str = Field(min_length=1, max_length=64)
    employee_no: str = Field(min_length=1, max_length=64)
    company: str = Field(min_length=1, max_length=255)
    department: str = Field(min_length=1, max_length=64)
    password: str = Field(default="", max_length=128)
    role: str = "user"
    position: str = Field(min_length=1, max_length=128)
    is_active: bool = True
    account_type: str = Field(default="员工账号", max_length=32)
    data: dict = Field(default_factory=dict)


class CurrentUserUpdate(BaseModel):
    display_name: str | None = Field(default=None, min_length=1, max_length=64)
    email: str | None = Field(default=None, max_length=128)
    office_phone: str | None = Field(default=None, max_length=32)
    mobile: str | None = Field(default=None, max_length=32)
    menu_auto_collapse: str | None = Field(default=None, pattern="^(yes|no)$")
    current_password: str | None = Field(default=None, min_length=1, max_length=128)
    new_password: str | None = Field(default=None, min_length=8, max_length=128)


class UserMessageInput(BaseModel):
    recipients: list[str] = Field(min_length=1, max_length=50)
    title: str = Field(min_length=1, max_length=255)
    content: str = Field(min_length=1, max_length=4000)


class CommunicationLogInput(BaseModel):
    customer_record_id: int
    contact: str = Field(default="", max_length=128)
    phone: str = Field(default="", max_length=64)
    content: str = Field(min_length=1, max_length=4000)
    occurred_at: datetime


class CommunicationLogUpdate(BaseModel):
    contact: str | None = Field(default=None, max_length=128)
    phone: str | None = Field(default=None, max_length=64)
    content: str | None = Field(default=None, min_length=1, max_length=4000)
    occurred_at: datetime | None = None


class RolePermissionUpdate(BaseModel):
    data_scope: str = Field(min_length=1, max_length=64)
    menu_keys: list[str]
    field_keys: list[str]


class SecurityPolicyUpdate(BaseModel):
    min_password_length: int = Field(ge=8, le=32)
    max_failed_attempts: int = Field(ge=3, le=10)
    lock_minutes: int = Field(ge=1, le=1440)
    token_minutes: int = Field(ge=15, le=1440)


class SystemParameterInput(BaseModel):
    category: str = Field(min_length=2, max_length=32)
    code: str = Field(min_length=1, max_length=64)
    name: str = Field(min_length=1, max_length=255)
    extra: dict = Field(default_factory=dict)
    sort_order: int = Field(default=0, ge=0, le=99999)
    is_active: bool = True


class SystemParameterUpdate(BaseModel):
    code: str | None = Field(default=None, min_length=1, max_length=64)
    name: str | None = Field(default=None, min_length=1, max_length=255)
    extra: dict | None = None
    sort_order: int | None = Field(default=None, ge=0, le=99999)
    is_active: bool | None = None


class SystemConfigUpdate(BaseModel):
    value: dict


class SystemMenuUpdate(BaseModel):
    label: str | None = Field(default=None, min_length=1, max_length=128)
    icon: str | None = Field(default=None, max_length=64)
    sort_order: int | None = Field(default=None, ge=0, le=99999)
    is_visible: bool | None = None
    is_active: bool | None = None


class SystemMenuInput(BaseModel):
    key: str = Field(min_length=1, max_length=128, pattern=r"^[a-z0-9][a-z0-9-]*$")
    parent_key: str = Field(default="", max_length=128)
    label: str = Field(min_length=1, max_length=128)
    icon: str = Field(default="", max_length=64)
    sort_order: int = Field(default=0, ge=0, le=99999)
    is_visible: bool = True
    is_active: bool = True


class ReportInput(BaseModel):
    title: str
    report_type: str
    period: str
    format: str = "CSV"
    description: str = ""


class CustomerShareInput(BaseModel):
    recipients: list[str] = Field(min_length=1, max_length=200)
    comment: str = ""


class CustomerActionInput(BaseModel):
    comment: str = ""


class CustomerContactInput(BaseModel):
    name: str
    project_role: str = ""
    phone: str = ""
    office_phone: str = ""
    im_account: str = ""
    email: str = ""
    position: str = ""
    contact_status: str = "正常联系"
    is_valid: bool = True
    is_primary: bool = False
    remark: str = ""


class CustomerManagersInput(BaseModel):
    managers: list[str] = Field(min_length=1, max_length=20)
    comment: str = ""


class CustomerNoteInput(BaseModel):
    content: str = Field(min_length=1, max_length=4000)
    note_type: str = Field(default="跟进记录", max_length=32)


class CustomerLevelChangeInput(BaseModel):
    level: str = Field(min_length=2, max_length=32)
    comment: str = Field(default="", max_length=1000)


class CustomerLevelReviewInput(BaseModel):
    approved: bool
    comment: str = Field(default="", max_length=1000)


class CustomerKeyChangeInput(BaseModel):
    title: str = Field(min_length=1, max_length=255)
    credit_code: str = Field(default="", max_length=64)
    comment: str = Field(min_length=2, max_length=1000)


class CustomerKeyChangeReviewInput(BaseModel):
    approved: bool
    comment: str = Field(default="", max_length=1000)


class CustomerPortalActionInput(BaseModel):
    comment: str = Field(default="", max_length=1000)


class CustomerPortalLoginInput(BaseModel):
    account: str = Field(min_length=3, max_length=128)
    activation_code: str = Field(min_length=16, max_length=128)


class CustomerPortalDemandInput(CustomerPortalLoginInput):
    title: str = Field(min_length=2, max_length=200)
    content: str = Field(min_length=2, max_length=2000)
    case_no: str = Field(default="", max_length=128)


class CustomerCreateInput(BaseModel):
    serial_no: str = ""
    title: str = ""
    status: str = ""
    owner: str = ""
    department: str = ""
    description: str = ""
    customer_managers: list[str] = Field(default_factory=list, max_length=20)
    customer_type: str | None = None
    level: str | None = None
    is_shared: str | bool | None = None
    is_assisted: str | bool | None = None
    fee_reduction: str | bool | None = None
    contact: str | None = None
    phone: str | None = None
    credit_code: str | None = None
    legal_representative: str | None = None
    registered_address: str | None = None
    invoice_title: str | None = None
    taxpayer_id: str | None = None
    invoice_address: str | None = None
    invoice_phone: str | None = None
    bank_name: str | None = None
    bank_account: str | None = None
    short_name: str | None = None
    fax: str | None = None
    legal_agent_id_no: str | None = None
    legal_agent_title: str | None = None
    customer_source: str | None = None
    file_date: str | None = None
    province: str | None = None
    postal_code: str | None = None
    patent_customer_type: str | None = None
    industry: str | None = None
    output_value: str | None = None
    cooperation_status: str | None = None
    gb_classification: str | None = None
    website: str | None = None
    organization_nature: str | None = None
    organization_code: str | None = None
    registration_region: str | None = None
    registration_postal_code: str | None = None
    registered_capital: str | None = None
    registration_year: str | None = None
    data: dict = Field(default_factory=dict)


class ContractSubmitInput(BaseModel):
    approvers: list[str] = Field(min_length=1, max_length=10)
    comment: str = ""


class ContractDraftInput(BaseModel):
    serial_no: str = Field(min_length=1, max_length=128)
    title: str = Field(min_length=1, max_length=255)
    customer: str = Field(min_length=1, max_length=255)
    owner: str = Field(min_length=1, max_length=64)
    department: str = Field(min_length=1, max_length=64)
    description: str = Field(default="", max_length=2000)
    data: dict = Field(default_factory=dict)


class ContractApprovalInput(BaseModel):
    approved: bool
    comment: str = ""


class CaseCreationReviewInput(BaseModel):
    approved: bool
    comment: str = Field(default="", max_length=1000)


class CaseUnarchiveRequestInput(BaseModel):
    reason: str = Field(min_length=2, max_length=1000)


class CaseUnarchiveReviewInput(BaseModel):
    approved: bool
    comment: str = Field(default="", max_length=1000)


class ContractSealApplicationInput(BaseModel):
    seal_asset_id: int
    copies: int = Field(ge=1, le=999)
    purpose: str = Field(min_length=1, max_length=500)
    use_date: date
    delivery_method: str = "现场用印"
    document_names: str = ""
    description: str = ""
    submit: bool = False


class ContractInvestigationInput(BaseModel):
    title: str = Field(min_length=2, max_length=200)
    owner: str = Field(default="", max_length=100)
    authorized_from: date
    authorized_to: date
    region: str = Field(default="", max_length=300)
    right_type: str = Field(default="商标", max_length=50)
    customer_review: bool = False
    description: str = Field(default="", max_length=2000)


class ContractChangeInput(BaseModel):
    change_type: str
    reason: str = Field(min_length=2, max_length=1000)
    title: str | None = None
    amount: float | None = Field(default=None, ge=0)
    external_contract_no: str | None = None
    external_contract_numbers: list[str] | None = Field(default=None, max_length=50)
    end_date: date | None = None


class ContractChangeReviewInput(BaseModel):
    approved: bool
    comment: str = Field(default="", max_length=1000)


class ContractEventInput(BaseModel):
    content: str = Field(min_length=1, max_length=1000)


class SealApplicationInput(BaseModel):
    title: str
    customer: str = ""
    case_no: str = ""
    contract_no: str = ""
    use_type: str = ""
    seal_asset_id: int
    copies: int = Field(ge=1, le=999)
    purpose: str
    use_date: date
    delivery_method: str = "现场用印"
    is_electronic_seal: bool = False
    is_offline_print: bool = False
    document_names: str = ""
    description: str = ""


class SealPackageDownloadInput(BaseModel):
    application_ids: list[int] = Field(min_length=1, max_length=100)


class SealApprovalInput(BaseModel):
    approved: bool
    comment: str = ""


class SealStampInput(BaseModel):
    actual_copies: int = Field(ge=1, le=999)
    operator: str = ""
    archive_no: str = ""
    comment: str = ""


class SealAssetInput(BaseModel):
    code: str = Field(min_length=2, max_length=64)
    name: str = Field(min_length=2, max_length=128)
    seal_type: str
    custodian: str
    location: str = ""
    remark: str = ""


class SealAssetUpdate(BaseModel):
    name: str | None = None
    seal_type: str | None = None
    custodian: str | None = None
    location: str | None = None
    status: str | None = None
    remark: str | None = None


WORKFLOW_TRANSITIONS: dict[str, dict[str, list[str]]] = {
    "customer": {"跟进中": ["正常", "公海"], "正常": ["待共享", "公海"], "待共享": ["正常"], "公海": ["跟进中", "已回收"]},
    "contract": {"草稿": ["审批中"], "审批中": ["已通过", "已拒绝"], "已通过": ["履行中"], "履行中": ["已完成"]},
    "case": {"等待公证书": ["等待审核公证书"], "等待审核公证书": ["新案待分配"], "新案待分配": ["文书准备"], "文书准备": ["一审立案受理"], "一审立案受理": ["一审准备开庭"], "一审准备开庭": ["待上诉", "二审", "执行", "已归档"], "待上诉": ["二审", "执行", "已归档"], "二审": ["执行", "已归档"], "执行": ["已归档"]},
    "task": {"待处理": ["处理中", "已撤回"], "处理中": ["已完成", "已逾期", "已撤回"], "已逾期": ["处理中", "已完成"]},
    "clue": {"草稿": ["待审批"], "待审批": ["待取证", "已驳回"], "待取证": ["已取证"], "已取证": ["待公证", "已转案件"], "待公证": ["已转案件", "已驳回"]},
    "notary": {"待审核": ["审核通过", "审核驳回"], "审核驳回": ["待审核"]},
    "evidence": {"待整理": ["已整理"], "已整理": ["已入卷"]},
    "seal": {"草稿": ["待审批", "已撤回"], "待审批": ["待用印", "已拒绝", "已撤回"], "待用印": ["已用印", "已撤回"], "已用印": ["已归档"]},
    "finance": {"草稿": ["待审批"], "待审批": ["已审批", "已退回"], "已审批": ["已付款"], "已付款": ["已对账"]},
    "document": {"待登记": ["待签收"], "待签收": ["已签收"], "已签收": ["已归档"]},
    "hr": {"试用": ["在职", "离职"], "在职": ["离职", "停用"]},
    "warehouse": {"在库": ["借出", "报废"], "借出": ["归还中"], "归还中": ["在库"]},
    "report": {"生成中": ["已生成"], "已生成": ["已发布"]},
    "system": {"启用": ["停用"], "停用": ["启用"]},
}

# 通用记录接口只服务于没有专用生命周期的简单资料。受保护模块的
# 状态、审批、附件和删除必须由各模块的 command API 完成。
GENERIC_RECORD_EDITABLE_MODULES = {"customer", "report", "system"}
GENERIC_RECORD_TRANSITION_MODULES = {"report", "system"}
GENERIC_RECORD_DELETABLE_MODULES = {"report"}


FIELD_PERMISSION_DATA_KEYS = {
    "customer.billing": {"invoice_title", "taxpayer_id", "invoice_address", "invoice_phone"},
    "customer.bank": {"bank_name", "bank_account"},
    "customer.legal": {"credit_code", "legal_representative", "registered_address"},
    "contract.amount": {"amount"},
    "finance.amount": {"amount", "paid_amount", "invoice_amount", "refund_amount"},
    "hr.identity": {"id_no"},
}


def _record_dict(record: BusinessRecord, allowed_fields: set[str] | None = None) -> dict:
    data = dict(record.data or {})
    if allowed_fields is not None:
        for permission, keys in FIELD_PERMISSION_DATA_KEYS.items():
            if permission not in allowed_fields:
                for key in keys: data.pop(key, None)
    return {
        "id": record.id, "module": record.module, "serial_no": record.serial_no,
        "title": record.title, "customer": record.customer, "status": record.status,
        "owner": record.owner, "department": record.department,
        "description": record.description, "data": data,
        "created_at": record.created_at, "updated_at": record.updated_at,
    }


def _receivable_dict(plan: ReceivablePlan, contract: BusinessRecord) -> dict:
    remaining = max(plan.amount - plan.received_amount, 0)
    effective_status = plan.status
    if remaining > 0 and plan.due_date < date.today() and plan.status != "已收款":
        effective_status = "已逾期"
    return {
        "id": plan.id, "contract_record_id": plan.contract_record_id,
        "contract_no": contract.serial_no, "contract_title": contract.title,
        "customer": contract.customer, "phase": plan.phase, "due_date": plan.due_date,
        "amount": plan.amount, "received_amount": plan.received_amount,
        "remaining_amount": remaining, "status": effective_status,
        "payer": plan.payer, "owner": contract.owner, "remark": plan.remark,
        "updated_at": plan.updated_at,
    }


def _hearing_dict(item: HearingSchedule, case_record: BusinessRecord) -> dict:
    return {
        "id": item.id, "case_record_id": item.case_record_id,
        "case_no": case_record.serial_no, "case_title": case_record.title,
        "customer": case_record.customer, "hearing_date": item.hearing_date,
        "weekday": "星期" + "一二三四五六日"[item.hearing_date.weekday()],
        "hearing_time": item.hearing_time, "court": item.court, "courtroom": item.courtroom,
        "hearing_type": item.hearing_type, "hearing_lawyer": item.hearing_lawyer,
        "status": item.status, "remark": item.remark,
    }


def _task_dict(record: BusinessRecord) -> dict:
    data = record.data or {}
    try:
        deadline = date.fromisoformat(str(data.get("deadline", "")))
        days_remaining = (deadline - date.today()).days
    except ValueError:
        deadline = None
        days_remaining = None
    workflow_status = record.status
    # 兼容早期本地数据：旧“待确认”等同于原系统“进行中-已完成”。
    effective_status = "已完成" if record.status == "待确认" else record.status
    if days_remaining is not None and days_remaining < 0 and record.status in {"待接收", "待处理", "处理中"}:
        effective_status = "已逾期"
    reminder_due = days_remaining in {1, 3} or (days_remaining is not None and days_remaining < 0 and abs(days_remaining) % 3 == 0)
    reminder_text = ""
    if days_remaining in {1, 3}:
        reminder_text = f"{days_remaining} 天后到期"
    elif days_remaining is not None and days_remaining < 0:
        reminder_text = f"已逾期 {abs(days_remaining)} 天" + ("，今日提醒" if reminder_due else "")
    elif days_remaining == 0:
        reminder_text = "今日到期"
    return {
        **_record_dict(record), "status": effective_status, "workflow_status": workflow_status,
        "deadline": deadline, "days_remaining": days_remaining,
        "priority": data.get("priority", "普通"), "source": data.get("source", "日常任务"),
        "initiator": data.get("initiator", ""), "collaborators": data.get("collaborators", []),
        "case_no": data.get("case_no", ""), "rejected_reason": data.get("rejected_reason", ""),
        "plaintiff": data.get("plaintiff", ""), "defendant": data.get("defendant", ""),
        "case_stage": data.get("case_stage", ""),
        "accepted_at": data.get("accepted_at", ""),
        "verified_at": data.get("confirmed_at") or data.get("auto_confirmed_at", ""),
        "completion_auto_confirm_at": data.get("completion_auto_confirm_at", ""),
        "reminder_due": reminder_due, "reminder_text": reminder_text,
        "handoff_recipient": data.get("handoff_recipient", ""),
        "handoff_auto_complete_at": data.get("handoff_auto_complete_at", ""),
        "handoff_restarted": bool(data.get("handoff_restarted")),
        "auto_completed": bool(data.get("auto_completed")),
        "performance_impact": data.get("performance_impact", {}),
        "exception_request": data.get("exception_request", {}),
        "parent_task_id": data.get("parent_task_id"),
        "parent_task_no": data.get("parent_task_no", ""),
        "investigation_record_id": data.get("investigation_record_id"),
        "investigation_no": data.get("investigation_no", ""),
        "investigation_module": data.get("investigation_module", ""),
    }


async def _add_task_message_notifications(
    task: BusinessRecord,
    event: WorkflowEvent,
    db: AsyncSession,
    *,
    content: str,
) -> None:
    """Persist one idempotent unread lifecycle/communication message per participant."""
    db.add(event)
    await db.flush()
    data = task.data or {}
    recipients = {
        task.owner,
        str(data.get("initiator") or ""),
        *(str(value) for value in data.get("collaborators", []) if value),
    } - {""}
    if event.operator != "system":
        recipients.discard(event.operator)
    active_recipients = set((await db.scalars(select(User.username).where(
        User.username.in_(recipients), User.is_active.is_(True),
    ))).all()) if recipients else set()
    existing_keys = set((await db.scalars(select(Notification.source_key).where(
        Notification.source_key.in_([
            f"task-message-{task.id}-{event.id}-{recipient}" for recipient in active_recipients
        ])
    ))).all()) if active_recipients else set()
    for recipient in active_recipients:
        source_key = f"task-message-{task.id}-{event.id}-{recipient}"
        if source_key in existing_keys:
            continue
        db.add(Notification(
            source_key=source_key, source_type="task", source_id=task.id,
            sender=event.operator or "system", recipient=recipient,
            notification_type="系统通知", title=f"任务新消息：{task.serial_no}",
            content=content.strip() or event.action, level="info", is_read=False,
        ))


async def _delete_task_notifications(task_id: int, db: AsyncSession) -> None:
    """Delete every notification owned by a task before deleting that task record."""
    await db.execute(delete(Notification).where(
        Notification.source_type == "task", Notification.source_id == task_id,
    ))


def _attachment_dict(item: FileAttachment, record: BusinessRecord | None = None) -> dict:
    return {
        "id": item.id, "record_id": item.record_id, "finance_transaction_id": item.finance_transaction_id,
        "record_no": record.serial_no if record else "",
        "record_title": record.title if record else "", "category": item.category,
        "original_name": item.original_name, "content_type": item.content_type,
        "size": item.size, "uploader": item.uploader, "remark": item.remark,
        "created_at": item.created_at, "download_url": f"{settings.api_prefix}/attachments/{item.id}/download",
    }


def _template_dict(item: DocumentTemplate) -> dict:
    return {"id": item.id, "name": item.name, "category": item.category, "version": item.version, "description": item.description, "fields": item.fields or [], "is_active": item.is_active, "created_at": item.created_at, "updated_at": item.updated_at}


ARCHIVE_REQUIRED_CATEGORIES = {"委托材料", "证据材料", "诉讼文书", "裁判文书"}


async def _sync_case_document_readiness(case_record: BusinessRecord, db: AsyncSession) -> bool:
    categories = set((await db.scalars(select(FileAttachment.category).where(FileAttachment.record_id == case_record.id))).all())
    complete = ARCHIVE_REQUIRED_CATEGORIES.issubset(categories)
    case_record.data = {**(case_record.data or {}), "documents_complete": complete, "archive_material_categories": sorted(categories)}
    return complete


def _record_links_to_case(record: BusinessRecord, case_record: BusinessRecord) -> bool:
    """Prefer the persisted case id; use the case number only for legacy rows."""
    record_data = record.data or {}
    linked_case_id = int(record_data.get("case_id") or record_data.get("case_record_id") or 0)
    if linked_case_id:
        return linked_case_id == case_record.id
    return str(record_data.get("case_no") or "") == case_record.serial_no


async def _case_archive_checks(case_record: BusinessRecord, db: AsyncSession) -> dict[str, bool]:
    """Calculate archive readiness from persisted business facts, never client checkboxes."""
    data = dict(case_record.data or {})
    documents_complete = await _sync_case_document_readiness(case_record, db)
    related_rows = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module.in_({"finance", "invoice", "refund"})))).all()

    related = [item for item in related_rows if _record_links_to_case(item, case_record)]
    fees = [item for item in related if item.module == "finance"]
    invoices = [item for item in related if item.module == "invoice"]
    refunds = [item for item in related if item.module == "refund"]
    fee_terminal = {"已付款", "已核销", "已对账", "已作废", "已撤销"}
    invoice_terminal = {"已开票", "已作废", "已撤回"}
    refund_terminal = {"已退款", "已作废", "已撤回"}
    fees_settled = all(item.status in fee_terminal for item in fees)
    contract_id = int(data.get("contract_id") or 0)
    receivables = (await db.scalars(select(ReceivablePlan).where(ReceivablePlan.contract_record_id == contract_id))).all() if contract_id else []
    receivables_complete = all(float(item.amount) - float(item.received_amount or 0) <= 0.001 for item in receivables)
    finance_complete = fees_settled and all(item.status in invoice_terminal for item in invoices) and all(item.status in refund_terminal for item in refunds) and receivables_complete
    checks = {
        "case_closed": bool(data.get("case_closed_at") and data.get("case_closed_by")),
        "fees_settled": fees_settled,
        "documents_complete": documents_complete,
        "finance_complete": finance_complete,
    }
    case_record.data = {**(case_record.data or {}), **checks}
    return checks


def _finance_transaction_dict(item: FinanceTransaction, record: BusinessRecord | None = None, attachments: list[FileAttachment] | None = None, *, show_amount: bool = True) -> dict:
    vouchers = attachments or []
    return {
        "id": item.id, "finance_record_id": item.finance_record_id,
        "finance_no": record.serial_no if record else "", "finance_title": record.title if record else "",
        "transaction_type": item.transaction_type, "amount": item.amount if show_amount else None, "transaction_date": item.transaction_date,
        "voucher_no": item.voucher_no, "counterparty": item.counterparty, "operator": item.operator,
        "remark": item.remark, "created_at": item.created_at,
        "voucher_count": len(vouchers), "voucher_categories": sorted({x.category for x in vouchers}),
        "vouchers": [_attachment_dict(x, record) for x in vouchers],
    }


def _incoming_payment_dict(item: IncomingPayment, *, show_amount: bool = True) -> dict:
    amount = float(item.amount); allocated = float(item.allocated_amount or 0)
    return {"id": item.id, "receipt_no": item.receipt_no, "received_date": item.received_date, "amount": amount if show_amount else None, "payer_name": item.payer_name, "bank_reference": item.bank_reference, "status": item.status, "claimed_customer": item.claimed_customer, "claimant": item.claimant, "allocated_amount": allocated if show_amount else None, "remaining_amount": max(amount - allocated, 0) if show_amount else None, "allocations": item.allocations or [], "operator": item.operator, "remark": item.remark, "created_at": item.created_at, "updated_at": item.updated_at}


def _reconciliation_dict(item: ReconciliationBatch, *, show_amount: bool = True) -> dict:
    return {"id": item.id, "period_type": item.period_type, "date_from": item.date_from, "date_to": item.date_to, "transaction_count": item.transaction_count, "total_amount": item.total_amount if show_amount else None, "discrepancy_amount": item.discrepancy_amount if show_amount else None, "status": item.status, "operator": item.operator, "remark": item.remark, "created_at": item.created_at}


def _seed_business_records() -> list[BusinessRecord]:
    rows = [
        ("customer", "KH20260714001", "光明乳业股份有限公司", "光明乳业股份有限公司", "正常", "朱菁芸", {"contact": "法务部", "phone": "021-12345678", "level": "重点客户"}),
        ("customer", "KH20260714002", "萨普托乳业（中国）有限公司", "萨普托乳业（中国）有限公司", "跟进中", "朱淑旖", {"contact": "品牌保护部", "phone": "021-87654321", "level": "重点客户"}),
        ("contract", "HT2026070018", "知识产权维权专项法律服务合同", "迈大食品（上海）有限公司", "审批中", "陈名涛", {"amount": "280000.00", "signed_at": "2026-07-08", "type": "专项服务"}),
        ("contract", "HT2026060097", "常年法律顾问合同", "上海天路人造草坪有限公司", "履行中", "陶勇刚", {"amount": "120000.00", "signed_at": "2026-06-20", "type": "法律顾问"}),
        ("case", "SH191000382B", "光明乳业商标侵权纠纷", "光明乳业股份有限公司", "文书准备", "陈名涛", {"court": "上海市宝山区人民法院", "case_type": "民事案件", "opponent": "安徽鑫牛食品有限公司"}),
        ("case", "SHMS2600387", "龙角散商标侵权纠纷", "株式会社龙角散", "一审立案受理", "陶勇刚", {"court": "杭州市余杭区人民法院", "case_type": "民事案件", "opponent": "杭州取道贸易有限公司"}),
        ("task", "RW20260714001", "准备开庭代理词及证据目录", "上海天路人造草坪有限公司", "处理中", "陶勇刚", {"deadline": "2026-07-15", "priority": "紧急", "source": "案件任务"}),
        ("task", "RW20260714002", "审核合同付款节点", "迈大食品（上海）有限公司", "待处理", "朱淑旖", {"deadline": "2026-07-17", "priority": "普通", "source": "合同任务"}),
        ("clue", "XS2026070015", "线上店铺销售疑似侵权产品", "北京汇源食品饮料有限公司", "待审批", "卢愿", {"platform": "淘宝", "product": "果汁饮料", "notary": "待申请"}),
        ("seal", "YY2026070042", "民事起诉状用印申请", "株式会社龙角散", "待审批", "陶勇刚", {"seal_type": "公章", "copies": 3, "purpose": "法院立案"}),
        ("finance", "FY2026070093", "上海市宝山区人民法院诉讼费", "光明乳业股份有限公司", "待审批", "陈名涛", {"amount": "3500.00", "fee_type": "官方费用", "case_no": "SH191000382B"}),
        ("document", "SW2026070031", "上海市徐汇区人民法院开庭传票", "上海益民食品一厂有限公司", "已签收", "江彤", {"direction": "收文", "received_at": "2026-07-14", "case_no": "SHMS2200026"}),
    ]
    return [BusinessRecord(module=m, serial_no=no, title=title, customer=customer, status=st, owner=owner, data=data) for m, no, title, customer, st, owner, data in rows]


@app.get("/health")
async def health():
    return {"status": "ok"}


async def _permission_payload(role: str, db: AsyncSession) -> dict:
    permission = await db.scalar(select(RolePermission).where(RolePermission.role == role))
    config = DEFAULT_ROLE_PERMISSIONS.get(role, DEFAULT_ROLE_PERMISSIONS["user"])
    if role == "admin":
        return {
            "menu_keys": list(MENU_KEYS),
            "data_scope": config["data_scope"],
            "field_keys": list(FIELD_KEYS),
        }
    return {
        "menu_keys": list(permission.menu_keys if permission else config["menu_keys"]),
        "data_scope": permission.data_scope if permission else config["data_scope"],
        "field_keys": list(permission.field_keys if permission else config["field_keys"]),
    }


async def _user_permission_payload(user: User, db: AsyncSession) -> dict:
    """Expose the contract approval workbench to users assigned that job permission."""
    permission = await _permission_payload(user.role, db)
    can_approve_contract = await _user_has_job_permission(user, "合同审批", db)
    menu_keys = list(permission["menu_keys"])
    if can_approve_contract and "contract" not in menu_keys:
        menu_keys.append("contract")
    return {**permission, "menu_keys": menu_keys, "can_approve_contract": can_approve_contract}


async def _allowed_field_keys(identity: dict, db: AsyncSession) -> set[str]:
    if identity.get("role") == "admin":
        return set(FIELD_KEYS)
    permission = await db.scalar(select(RolePermission).where(RolePermission.role == identity.get("role")))
    config = DEFAULT_ROLE_PERMISSIONS.get(identity.get("role"), DEFAULT_ROLE_PERMISSIONS["user"])
    return set(permission.field_keys if permission else config["field_keys"])


async def _record_dict_for_identity(record: BusinessRecord, identity: dict, db: AsyncSession) -> dict:
    return _record_dict(record, await _allowed_field_keys(identity, db))


async def _resolve_active_customer_managers(values: list[object], db: AsyncSession) -> list[str]:
    """Resolve customer managers to stable usernames, accepting unique display names for legacy UI clients."""
    tokens = list(dict.fromkeys(str(value or "").strip() for value in values if str(value or "").strip()))
    if not tokens:
        raise HTTPException(status_code=422, detail="至少保留一名客户管理人")
    users = list((await db.scalars(select(User).where(User.is_active.is_(True), or_(User.username.in_(tokens), User.display_name.in_(tokens))))).all())
    by_username = {user.username: user.username for user in users}
    by_display: dict[str, list[str]] = {}
    for user in users:
        by_display.setdefault(user.display_name, []).append(user.username)
    resolved: list[str] = []
    invalid: list[str] = []
    for token in tokens:
        username = by_username.get(token)
        if not username:
            matches = by_display.get(token, [])
            if len(matches) == 1:
                username = matches[0]
        if not username:
            invalid.append(token)
        elif username not in resolved:
            resolved.append(username)
    if invalid:
        raise HTTPException(status_code=422, detail=f"客户管理人不存在、已停用或姓名不唯一：{'、'.join(invalid)}")
    return resolved


async def _resolve_active_case_people(values: list[object], db: AsyncSession, *, field_name: str) -> tuple[list[str], list[str]]:
    """Resolve case-team inputs once and persist usernames alongside display values.

    The UI remains compatible with the old name-based forms, but access control
    never depends on a mutable/duplicate display name after a case is saved.
    """
    labels = list(dict.fromkeys(str(value or "").strip() for value in values if str(value or "").strip()))
    if not labels:
        return [], []
    users = list((await db.scalars(select(User).where(
        User.is_active.is_(True),
        or_(User.username.in_(labels), User.display_name.in_(labels)),
    ))).all())
    by_username = {user.username: user for user in users}
    by_display: dict[str, list[User]] = {}
    for user in users:
        by_display.setdefault(user.display_name, []).append(user)
    resolved_labels: list[str] = []
    resolved_usernames: list[str] = []
    invalid: list[str] = []
    for label in labels:
        user = by_username.get(label)
        if not user:
            matches = by_display.get(label, [])
            user = matches[0] if len(matches) == 1 else None
        if not user:
            invalid.append(label)
            continue
        # Preserve the UI's submitted label for backward-compatible display;
        # the parallel username list below is the authoritative ACL identity.
        resolved_labels.append(label)
        if user.username not in resolved_usernames:
            resolved_usernames.append(user.username)
    if invalid:
        raise HTTPException(status_code=422, detail=f"{field_name}不存在、已停用或姓名不唯一：{'、'.join(invalid)}")
    return resolved_labels, resolved_usernames


def _case_team_payload(
    case_data: dict, handling_lawyers: list[str], handling_usernames: list[str], assistant: str, assistant_username: str,
) -> dict:
    """Keep legacy display fields and the stable access-control projection in sync."""
    return {
        **case_data,
        "handling_lawyers": handling_lawyers,
        "assistant": assistant,
        "handling_lawyer_usernames": handling_usernames,
        "assistant_username": assistant_username,
        "case_team_usernames": list(dict.fromkeys([*handling_usernames, *([assistant_username] if assistant_username else [])])),
    }


async def _case_team_role(case_record: BusinessRecord, identity: dict, db: AsyncSession) -> str:
    """Return manager/handling_lawyer/assistant/none for a visible case.

    New writes use stable usernames.  The narrowly-scoped legacy fallback keeps
    old rows usable only where the current user's display name is unique.
    """
    if identity.get("role") == "admin" or case_record.owner == identity["username"]:
        return "manager"
    user = await db.scalar(select(User).where(User.username == identity["username"]))
    if identity.get("role") == "manager" and user and case_record.department == user.department:
        return "manager"
    data = case_record.data or {}
    username = identity["username"]
    handling_usernames = {str(value or "").strip() for value in data.get("handling_lawyer_usernames", [])}
    assistant_username = str(data.get("assistant_username") or "").strip()
    if username in handling_usernames:
        return "handling_lawyer"
    if username and username == assistant_username:
        return "assistant"
    if data.get("case_team_usernames"):
        return "none"
    # Legacy-only fallback.  Display names may be used only when unique.
    display_name = str(user.display_name or "").strip() if user else ""
    display_is_unique = bool(display_name) and (await db.scalar(select(func.count(User.id)).where(User.is_active.is_(True), User.display_name == display_name)) == 1)
    handler_labels = {str(value or "").strip() for value in data.get("handling_lawyers", [])}
    assistant_label = str(data.get("assistant") or "").strip()
    if username in handler_labels or (display_is_unique and display_name in handler_labels):
        return "handling_lawyer"
    if username == assistant_label or (display_is_unique and display_name == assistant_label):
        return "assistant"
    return "none"


async def _record_scope_conditions(identity: dict, db: AsyncSession) -> list:
    if identity.get("role") == "admin":
        return []
    user = await db.scalar(select(User).where(User.username == identity["username"]))
    if not user:
        raise HTTPException(status_code=401, detail="当前用户不存在")
    permission = await db.scalar(select(RolePermission).where(RolePermission.role == user.role))
    scope = permission.data_scope if permission else DEFAULT_ROLE_PERMISSIONS.get(user.role, DEFAULT_ROLE_PERMISSIONS["user"])["data_scope"]
    if scope == "全所数据":
        return []
    public_customer = and_(BusinessRecord.module == "customer", BusinessRecord.status == "公海")
    # JSON arrays are serialized with quoted string members.  Matching the
    # quoted username keeps ``ann`` from gaining access to rows assigned or
    # shared to ``joann`` while remaining portable across SQLite and
    # PostgreSQL.
    exact_username_token = f'"{user.username}"'
    managed_customer = and_(BusinessRecord.module == "customer", BusinessRecord.data["customer_managers"].as_string().contains(exact_username_token))
    shared_to_text = BusinessRecord.data["shared_to"].as_string()
    shared_with_text = BusinessRecord.data["shared_with"].as_string()
    shared_customer = and_(BusinessRecord.module == "customer", shared_with_text.contains(exact_username_token))
    exact_shared_to = shared_to_text.contains(exact_username_token)
    # A case team member must be able to find the assigned case even if their
    # ordinary role has only "own data" scope.  This is deliberately limited to
    # case records and stable username projections; it does not widen financial,
    # archive or team-management authority.
    case_team = and_(BusinessRecord.module == "case", BusinessRecord.data["case_team_usernames"].as_string().contains(exact_username_token))
    if scope == "本部门数据":
        return [or_(BusinessRecord.department == user.department, public_customer, managed_customer, shared_customer, exact_shared_to, case_team)]
    if scope == "授权审批数据":
        return [or_(BusinessRecord.owner == user.username, public_customer, managed_customer, shared_customer, exact_shared_to, case_team, and_(BusinessRecord.module.in_(["contract", "finance", "invoice", "refund", "seal", "clue", "notary"]), BusinessRecord.status.in_(["待审批", "审批中", "待审核"])))]
    return [or_(BusinessRecord.owner == user.username, public_customer, managed_customer, shared_customer, exact_shared_to, case_team)]


async def _ensure_record_visible(record_id: int, identity: dict, db: AsyncSession) -> BusinessRecord:
    conditions = [BusinessRecord.id == record_id, *(await _record_scope_conditions(identity, db))]
    record = await db.scalar(select(BusinessRecord).where(*conditions))
    if not record:
        raise HTTPException(status_code=404, detail="记录不存在或无权访问")
    return record


async def _ensure_attachment_record_visible(record_id: int, identity: dict, db: AsyncSession) -> BusinessRecord:
    """Resolve attachment parent visibility without losing task-participant access.

    Task collaborators and initiators deliberately have a narrower record scope than
    ordinary owners, but they are still first-class participants for task feedback.
    Attachments must therefore use the task participation rule before the generic
    business-record scope rule.
    """
    record = await db.get(BusinessRecord, record_id)
    if record and record.module == "task":
        if not _is_task_participant(record, identity):
            raise HTTPException(status_code=403, detail="只有任务参与人可以访问任务反馈附件")
        return record
    return await _ensure_record_visible(record_id, identity, db)


async def _ensure_record_module(record_id: int, module: str, identity: dict, db: AsyncSession) -> BusinessRecord:
    record = await _ensure_record_visible(record_id, identity, db)
    if record.module != module:
        raise HTTPException(status_code=404, detail="业务记录不存在")
    return record


async def _require_record_owner_or_manager(record: BusinessRecord, identity: dict, db: AsyncSession) -> None:
    if record.module == "customer" and record.status == "公海" and identity.get("role") != "admin":
        raise HTTPException(status_code=403, detail="公海客户必须先领取后才能修改")
    if identity.get("role") == "admin" or record.owner == identity["username"] or (
        record.module == "customer" and identity["username"] in (record.data or {}).get("customer_managers", [])
    ):
        return
    user = await db.scalar(select(User).where(User.username == identity["username"]))
    if identity.get("role") == "manager" and user and record.department == user.department:
        return
    raise HTTPException(status_code=403, detail="只有负责人、部门负责人或系统管理员可以执行此操作")


def _require_task_owner_or_initiator(task: BusinessRecord, identity: dict, *, action: str) -> None:
    """Protect task writes from department-manager privilege escalation.

    Department scope grants read access to task lists, not authority to impersonate
    the current owner.  A manager who is neither owner nor initiator must not first
    replace the owner through a generic/batch endpoint and then perform lifecycle
    actions as that new owner.  System administrators deliberately retain the
    documented all-firm override.
    """
    username = identity["username"]
    data = task.data or {}
    if identity.get("role") == "admin" or task.owner == username or data.get("initiator") == username:
        return
    raise HTTPException(status_code=403, detail=f"只有任务负责人、发起人或系统管理员可以{action}")


async def _visible_record_ids(identity: dict, db: AsyncSession) -> set[int]:
    conditions = await _record_scope_conditions(identity, db)
    return set((await db.scalars(select(BusinessRecord.id).where(*conditions))).all())


async def _filter_visible_attachments(items: list[FileAttachment], identity: dict, db: AsyncSession) -> list[FileAttachment]:
    if identity.get("role") == "admin":
        return items
    record_ids = await _visible_record_ids(identity, db)
    return [item for item in items if (item.record_id and item.record_id in record_ids) or (not item.record_id and item.uploader == identity["username"])]


async def _security_policy(db: AsyncSession) -> SecurityPolicy:
    policy = await db.get(SecurityPolicy, 1)
    if not policy:
        policy = SecurityPolicy(id=1, min_password_length=8, max_failed_attempts=5, lock_minutes=30, token_minutes=settings.access_token_minutes, updated_by="system")
        db.add(policy); await db.flush()
    return policy


def _security_policy_dict(policy: SecurityPolicy) -> dict:
    return {"min_password_length": policy.min_password_length, "max_failed_attempts": policy.max_failed_attempts, "lock_minutes": policy.lock_minutes, "token_minutes": policy.token_minutes, "updated_by": policy.updated_by, "updated_at": policy.updated_at}


@app.post(f"{settings.api_prefix}/auth/login")
async def login(form: OAuth2PasswordRequestForm = Depends(), db: AsyncSession = Depends(get_db)):
    user = await db.scalar(select(User).where(User.username == form.username))
    if not user or not user.is_active:
        raise HTTPException(status_code=401, detail="账号或密码错误")
    policy = await _security_policy(db)
    # PostgreSQL returns TIMESTAMP WITH TIME ZONE values as aware datetimes,
    # while SQLite returns the same model field as a naive datetime.
    # Match the persisted value's timezone before comparing so both local
    # development and Docker/PostgreSQL enforce login locks identically.
    now = datetime.now(user.locked_until.tzinfo) if user.locked_until and user.locked_until.tzinfo else datetime.now()
    if user.locked_until and user.locked_until > now:
        raise HTTPException(status_code=423, detail=f"登录失败次数过多，账号锁定至 {user.locked_until.strftime('%Y-%m-%d %H:%M:%S')}")
    if not verify_password(form.password, user.password_hash):
        user.failed_login_attempts = int(user.failed_login_attempts or 0) + 1
        if user.failed_login_attempts >= policy.max_failed_attempts:
            user.locked_until = now + timedelta(minutes=policy.lock_minutes)
        await db.commit()
        if user.locked_until: raise HTTPException(status_code=423, detail=f"登录失败次数过多，账号已锁定 {policy.lock_minutes} 分钟")
        raise HTTPException(status_code=401, detail=f"账号或密码错误，还可尝试 {policy.max_failed_attempts - user.failed_login_attempts} 次")
    user.failed_login_attempts = 0; user.locked_until = None; user.last_login_at = now
    await db.commit()
    permission = await _user_permission_payload(user, db)
    return {"access_token": create_token(user.username, user.role, policy.token_minutes), "token_type": "bearer", "expires_in": policy.token_minutes * 60, "must_change_password": user.must_change_password, "user": {"username": user.username, "display_name": user.display_name, "department": user.department, "role": user.role, "must_change_password": user.must_change_password, **permission}}


@app.get(f"{settings.api_prefix}/auth/me")
async def current_user_profile(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    user = await db.scalar(select(User).where(User.username == identity["username"]))
    if not user or not user.is_active:
        raise HTTPException(status_code=404, detail="当前用户不存在")
    return {**_system_user_dict(user), **(await _user_permission_payload(user, db))}


@app.patch(f"{settings.api_prefix}/auth/me")
async def update_current_user_profile(body: CurrentUserUpdate, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    user = await db.scalar(select(User).where(User.username == identity["username"]))
    if not user or not user.is_active:
        raise HTTPException(status_code=404, detail="当前用户不存在")
    if body.display_name is not None:
        user.display_name = body.display_name.strip()
    profile_changes = {
        key: value.strip() if isinstance(value, str) else value
        for key, value in {
            "email": body.email,
            "office_phone": body.office_phone,
            "mobile": body.mobile,
            "menu_auto_collapse": body.menu_auto_collapse,
        }.items()
        if value is not None
    }
    if profile_changes:
        user.profile = {**(user.profile or {}), **profile_changes}
    if body.new_password is not None:
        if not body.current_password or not verify_password(body.current_password, user.password_hash):
            raise HTTPException(status_code=400, detail="当前密码不正确")
        if body.new_password == body.current_password:
            raise HTTPException(status_code=400, detail="新密码不能与当前密码相同")
        policy = await _security_policy(db)
        if len(body.new_password) < policy.min_password_length: raise HTTPException(status_code=422, detail=f"新密码至少需要 {policy.min_password_length} 位")
        user.password_hash = hash_password(body.new_password)
        user.password_changed_at = datetime.now(); user.failed_login_attempts = 0; user.locked_until = None; user.must_change_password = False
    await db.commit()
    await db.refresh(user)
    return {**_system_user_dict(user), **(await _user_permission_payload(user, db))}


def _require_admin(identity: dict) -> None:
    if identity.get("role") != "admin":
        raise HTTPException(status_code=403, detail="仅系统管理员可以执行此操作")


def _system_user_dict(user: User) -> dict:
    profile = user.profile or {}
    return {"id": user.id, "username": user.username, "display_name": user.display_name, "department": user.department, "role": user.role, "is_active": user.is_active, "must_change_password": user.must_change_password, "profile": profile, "email": profile.get("email", ""), "office_phone": profile.get("office_phone", ""), "mobile": profile.get("mobile", ""), "menu_auto_collapse": profile.get("menu_auto_collapse", "no"), "failed_login_attempts": user.failed_login_attempts or 0, "locked_until": user.locked_until, "last_login_at": user.last_login_at, "password_changed_at": user.password_changed_at, "created_at": user.created_at}


@app.get(f"{settings.api_prefix}/system/users")
async def list_system_users(keyword: str = "", identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    statement = select(User)
    if keyword.strip():
        like = f"%{keyword.strip()}%"
        statement = statement.where(or_(User.username.ilike(like), User.display_name.ilike(like)))
    users = (await db.scalars(statement.order_by(User.id))).all()
    return {"items": [_system_user_dict(user) for user in users], "total": len(users)}


@app.post(f"{settings.api_prefix}/system/users", status_code=status.HTTP_201_CREATED)
async def create_system_user(body: SystemUserInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    if body.role not in {"admin", "manager", "auditor", "user"}:
        raise HTTPException(status_code=422, detail="角色值无效")
    username = body.username.strip().lower()
    if await db.scalar(select(User).where(User.username == username)):
        raise HTTPException(status_code=409, detail="登录账号已存在")
    policy = await _security_policy(db)
    if len(body.password) < policy.min_password_length: raise HTTPException(status_code=422, detail=f"密码至少需要 {policy.min_password_length} 位")
    user = User(username=username, display_name=body.display_name.strip(), department=body.department.strip(), role=body.role, profile=body.profile, password_hash=hash_password(body.password), is_active=body.is_active, password_changed_at=None if body.must_change_password else datetime.now(), must_change_password=body.must_change_password)
    db.add(user)
    await db.commit(); await db.refresh(user)
    return _system_user_dict(user)


def _replace_username_value(value: object, old_username: str, new_username: str) -> object:
    """Replace exact username tokens in JSON data without touching free text."""
    if isinstance(value, str):
        return new_username if value == old_username else value
    if isinstance(value, list):
        return [_replace_username_value(item, old_username, new_username) for item in value]
    if isinstance(value, dict):
        return {key: _replace_username_value(item, old_username, new_username) for key, item in value.items()}
    return value


async def _rename_system_username(user: User, requested_username: str, identity: dict, db: AsyncSession) -> str:
    new_username = requested_username.strip().lower()
    old_username = user.username
    if not re.fullmatch(r"[a-z0-9._-]+", new_username):
        raise HTTPException(status_code=422, detail="登录账号只能包含小写字母、数字、点、下划线或短横线")
    if new_username == old_username:
        return old_username
    if old_username == identity["username"]:
        raise HTTPException(status_code=409, detail="不能在当前登录会话中修改自己的登录账号，请由其他管理员操作")
    if await db.scalar(select(User.id).where(User.username == new_username, User.id != user.id)):
        raise HTTPException(status_code=409, detail="登录账号已存在")

    scalar_references = (
        (Department, Department.manager), (Department, Department.created_by), (Department, Department.updated_by),
        (JobRole, JobRole.created_by), (JobRole, JobRole.updated_by),
        (SecurityPolicy, SecurityPolicy.updated_by), (SystemParameter, SystemParameter.created_by),
        (SystemParameter, SystemParameter.updated_by), (SystemConfig, SystemConfig.updated_by),
        (SystemMenu, SystemMenu.updated_by), (BusinessRecord, BusinessRecord.owner),
        (WorkflowEvent, WorkflowEvent.operator), (HrSubrecord, HrSubrecord.created_by),
        (HrSubrecord, HrSubrecord.updated_by), (HearingSchedule, HearingSchedule.hearing_lawyer),
        (FileAttachment, FileAttachment.uploader), (FinanceTransaction, FinanceTransaction.operator),
        (ReconciliationBatch, ReconciliationBatch.operator), (IncomingPayment, IncomingPayment.claimant),
        (IncomingPayment, IncomingPayment.operator), (ContractApprovalStep, ContractApprovalStep.approver),
        (Notification, Notification.sender), (Notification, Notification.recipient),
        (CommunicationLog, CommunicationLog.operator), (AgentDocument, AgentDocument.creator),
        (AgentDocument, AgentDocument.confirmed_by), (SealAsset, SealAsset.custodian),
    )
    for model, field in scalar_references:
        await db.execute(update(model).where(field == old_username).values({field.key: new_username}))

    for record in (await db.scalars(select(BusinessRecord))).all():
        replaced = _replace_username_value(record.data or {}, old_username, new_username)
        if replaced != (record.data or {}):
            record.data = replaced
    for payment in (await db.scalars(select(IncomingPayment))).all():
        replaced = _replace_username_value(payment.allocations or [], old_username, new_username)
        if replaced != (payment.allocations or []):
            payment.allocations = replaced
    for config in (await db.scalars(select(SystemConfig))).all():
        replaced = _replace_username_value(config.value or {}, old_username, new_username)
        if replaced != (config.value or {}):
            config.value = replaced
    user.profile = _replace_username_value(user.profile or {}, old_username, new_username)
    user.username = new_username
    return new_username


@app.patch(f"{settings.api_prefix}/system/users/{{user_id}}")
async def update_system_user(user_id: int, body: SystemUserUpdate, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    user = await db.get(User, user_id)
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在")
    if body.username is not None:
        await _rename_system_username(user, body.username, identity, db)
    if body.role is not None:
        if body.role not in {"admin", "manager", "auditor", "user"}:
            raise HTTPException(status_code=422, detail="角色值无效")
        if user.username == identity["username"] and body.role != "admin":
            raise HTTPException(status_code=409, detail="不能取消当前登录账号的管理员角色")
        user.role = body.role
    if body.is_active is not None:
        if user.username == identity["username"] and not body.is_active:
            raise HTTPException(status_code=409, detail="不能停用当前登录账号")
        user.is_active = body.is_active
    if body.display_name is not None:
        user.display_name = body.display_name.strip()
    if body.department is not None:
        user.department = body.department.strip()
    if body.password is not None:
        policy = await _security_policy(db)
        if len(body.password) < policy.min_password_length: raise HTTPException(status_code=422, detail=f"密码至少需要 {policy.min_password_length} 位")
        user.password_hash = hash_password(body.password)
        user.password_changed_at = datetime.now(); user.failed_login_attempts = 0; user.locked_until = None
    if body.profile is not None:
        user.profile = {**(user.profile or {}), **body.profile}
    await db.commit(); await db.refresh(user)
    return _system_user_dict(user)


@app.delete(f"{settings.api_prefix}/system/users/{{user_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_system_user(user_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    user = await db.get(User, user_id)
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在")
    if user.username == identity["username"]:
        raise HTTPException(status_code=409, detail="不能删除当前登录账号")
    await db.delete(user); await db.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.post(f"{settings.api_prefix}/system/users/{{user_id}}/reset-password")
async def reset_system_user_password(user_id: int, body: SystemUserPasswordResetInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Reset a password as a distinct security action, not a profile edit.

    The recipient must change this administrator-issued password before using
    business APIs.  A reset also removes any stale login lock, which is the
    practical equivalent of the old system's separate "reset password" action.
    """
    _require_admin(identity)
    user = await db.get(User, user_id)
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在")
    policy = await _security_policy(db)
    if len(body.new_password) < policy.min_password_length:
        raise HTTPException(status_code=422, detail=f"新密码至少需要 {policy.min_password_length} 位")
    if verify_password(body.new_password, user.password_hash):
        raise HTTPException(status_code=400, detail="新密码不能与当前密码相同")
    user.password_hash = hash_password(body.new_password)
    user.password_changed_at = None
    user.must_change_password = True
    user.failed_login_attempts = 0
    user.locked_until = None
    await db.commit(); await db.refresh(user)
    return _system_user_dict(user)


@app.post(f"{settings.api_prefix}/system/users/{{user_id}}/unlock")
async def unlock_system_user(user_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    user = await db.get(User, user_id)
    if not user: raise HTTPException(status_code=404, detail="用户不存在")
    user.failed_login_attempts = 0; user.locked_until = None
    await db.commit(); await db.refresh(user); return _system_user_dict(user)


@app.get(f"{settings.api_prefix}/system/security-policy")
async def get_security_policy(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity); return _security_policy_dict(await _security_policy(db))


@app.patch(f"{settings.api_prefix}/system/security-policy")
async def update_security_policy(body: SecurityPolicyUpdate, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity); policy = await _security_policy(db)
    for key, value in body.model_dump().items(): setattr(policy, key, value)
    policy.updated_by = identity["username"]
    await db.commit(); await db.refresh(policy); return _security_policy_dict(policy)


def _system_parameter_dict(item: SystemParameter) -> dict:
    return {
        "id": item.id, "category": item.category, "category_name": SYSTEM_PARAMETER_CATEGORIES.get(item.category, item.category),
        "code": item.code, "name": item.name, "extra": item.extra or {}, "sort_order": item.sort_order,
        "is_active": item.is_active, "created_by": item.created_by, "updated_by": item.updated_by,
        "created_at": item.created_at, "updated_at": item.updated_at,
    }


def _clear_parameter_cache(category: str | None = None, operator: str = "system") -> None:
    if category:
        SYSTEM_PARAMETER_CACHE.pop(category, None)
        SYSTEM_PARAMETER_CACHE.pop("__all__", None)
    else:
        SYSTEM_PARAMETER_CACHE.clear()
    SYSTEM_CACHE_META["system-parameters"] = {"last_cleared_at": datetime.now().isoformat(), "last_cleared_by": operator}


@app.get(f"{settings.api_prefix}/system/parameter-categories")
async def list_system_parameter_categories(identity: dict = Depends(current_identity)):
    _require_admin(identity)
    return {"items": [{"key": key, "name": name} for key, name in SYSTEM_PARAMETER_CATEGORIES.items()]}


@app.get(f"{settings.api_prefix}/system/parameters")
async def list_system_parameters(category: str = "", keyword: str = "", identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    if category and category not in SYSTEM_PARAMETER_CATEGORIES: raise HTTPException(status_code=422, detail="参数分类无效")
    cache_key = category or "__all__"
    if not keyword.strip() and cache_key in SYSTEM_PARAMETER_CACHE:
        return {"items": SYSTEM_PARAMETER_CACHE[cache_key], "categories": SYSTEM_PARAMETER_CATEGORIES, "cached": True}
    statement = select(SystemParameter)
    if category: statement = statement.where(SystemParameter.category == category)
    if keyword.strip():
        term = f"%{keyword.strip()}%"
        statement = statement.where(or_(SystemParameter.code.ilike(term), SystemParameter.name.ilike(term)))
    items = (await db.scalars(statement.order_by(SystemParameter.sort_order, SystemParameter.id))).all()
    result = [_system_parameter_dict(item) for item in items]
    if not keyword.strip(): SYSTEM_PARAMETER_CACHE[cache_key] = result
    return {"items": result, "categories": SYSTEM_PARAMETER_CATEGORIES, "cached": False}


@app.post(f"{settings.api_prefix}/system/parameters", status_code=status.HTTP_201_CREATED)
async def create_system_parameter(body: SystemParameterInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    if body.category not in SYSTEM_PARAMETER_CATEGORIES: raise HTTPException(status_code=422, detail="参数分类无效")
    code, name = body.code.strip(), body.name.strip()
    duplicate = await db.scalar(select(SystemParameter).where(SystemParameter.category == body.category, or_(SystemParameter.code == code, SystemParameter.name == name)))
    if duplicate: raise HTTPException(status_code=409, detail="同一分类下参数代码或名称已存在")
    item = SystemParameter(**body.model_dump(exclude={"code", "name"}), code=code, name=name, created_by=identity["username"], updated_by=identity["username"])
    db.add(item); await db.commit(); await db.refresh(item); _clear_parameter_cache(body.category, identity["username"])
    return _system_parameter_dict(item)


@app.patch(f"{settings.api_prefix}/system/parameters/{{parameter_id}}")
async def update_system_parameter(parameter_id: int, body: SystemParameterUpdate, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    item = await db.get(SystemParameter, parameter_id)
    if not item: raise HTTPException(status_code=404, detail="系统参数不存在")
    code = body.code.strip() if body.code is not None else item.code
    name = body.name.strip() if body.name is not None else item.name
    duplicate = await db.scalar(select(SystemParameter).where(SystemParameter.category == item.category, SystemParameter.id != item.id, or_(SystemParameter.code == code, SystemParameter.name == name)))
    if duplicate: raise HTTPException(status_code=409, detail="同一分类下参数代码或名称已存在")
    for key, value in body.model_dump(exclude_unset=True).items(): setattr(item, key, value.strip() if key in {"code", "name"} else value)
    item.updated_by = identity["username"]
    await db.commit(); await db.refresh(item); _clear_parameter_cache(item.category, identity["username"])
    return _system_parameter_dict(item)


@app.delete(f"{settings.api_prefix}/system/parameters/{{parameter_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_system_parameter(parameter_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    item = await db.get(SystemParameter, parameter_id)
    if not item: raise HTTPException(status_code=404, detail="系统参数不存在")
    category = item.category
    await db.delete(item); await db.commit(); _clear_parameter_cache(category, identity["username"])
    return Response(status_code=status.HTTP_204_NO_CONTENT)


def _validate_system_config(key: str, value: dict) -> dict:
    if key == "customer_share_policy":
        required = {"all_days", "filed_days", "premium_days", "standard_days", "basic_days", "shared_days"}
        if set(value) != required or any(isinstance(value[name], bool) or not isinstance(value[name], int) or not 1 <= value[name] <= 3650 for name in required):
            raise HTTPException(status_code=422, detail="客户共享天数必须完整填写，范围为 1 至 3650 天")
    elif key == "company_profile":
        required = {"name", "code", "short_code", "address", "phone", "fax", "email", "postal_code", "bank_name", "bank_account", "bank_address"}
        if set(value) != required or any(not isinstance(value[name], str) for name in required) or not all(str(value[name]).strip() for name in {"name", "code", "short_code"}):
            raise HTTPException(status_code=422, detail="公司名称、代码和字母短写代码必填，且公司资料字段必须完整")
        value = {name: str(value[name]).strip() for name in required}
    elif key == "application_settings":
        required = {"system_name", "default_department", "page_size", "attachment_limit_mb", "maintenance_mode"}
        if set(value) != required or not isinstance(value["maintenance_mode"], bool) or not isinstance(value["page_size"], int) or not 10 <= value["page_size"] <= 200 or not isinstance(value["attachment_limit_mb"], int) or not 1 <= value["attachment_limit_mb"] <= 100:
            raise HTTPException(status_code=422, detail="系统配置字段或数值范围无效")
        if not str(value["system_name"]).strip() or not str(value["default_department"]).strip(): raise HTTPException(status_code=422, detail="系统名称和默认部门必填")
    else:
        raise HTTPException(status_code=404, detail="系统配置不存在")
    return value


@app.get(f"{settings.api_prefix}/system/configs")
async def list_system_configs(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    items = (await db.scalars(select(SystemConfig).order_by(SystemConfig.id))).all()
    return {"items": [{"key": item.key, "label": item.label, "group": item.group, "value": item.value or {}, "description": item.description, "updated_by": item.updated_by, "updated_at": item.updated_at} for item in items]}


@app.patch(f"{settings.api_prefix}/system/configs/{{config_key}}")
async def update_system_config(config_key: str, body: SystemConfigUpdate, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    item = await db.scalar(select(SystemConfig).where(SystemConfig.key == config_key))
    if not item: raise HTTPException(status_code=404, detail="系统配置不存在")
    item.value = _validate_system_config(config_key, body.value); item.updated_by = identity["username"]
    await db.commit(); await db.refresh(item)
    return {"key": item.key, "label": item.label, "group": item.group, "value": item.value, "description": item.description, "updated_by": item.updated_by, "updated_at": item.updated_at}


@app.get(f"{settings.api_prefix}/system/caches")
async def list_system_caches(identity: dict = Depends(current_identity)):
    _require_admin(identity)
    meta = SYSTEM_CACHE_META["system-parameters"]
    return {"items": [{"key": "system-parameters", "name": "系统参数字典缓存", "entry_count": sum(len(items) for items in SYSTEM_PARAMETER_CACHE.values()), "bucket_count": len(SYSTEM_PARAMETER_CACHE), **meta}]}


@app.post(f"{settings.api_prefix}/system/caches/{{cache_key}}/clear")
async def clear_system_cache(cache_key: str, identity: dict = Depends(current_identity)):
    _require_admin(identity)
    if cache_key != "system-parameters": raise HTTPException(status_code=404, detail="缓存不存在")
    _clear_parameter_cache(operator=identity["username"])
    return {"key": cache_key, "cleared": True, **SYSTEM_CACHE_META[cache_key]}


def _system_menu_dict(item: SystemMenu) -> dict:
    return {
        "id": item.id, "key": item.key, "parent_key": item.parent_key,
        "label": item.label, "icon": item.icon, "sort_order": item.sort_order,
        "is_visible": item.is_visible, "is_active": item.is_active,
        "is_system": item.key in SYSTEM_MENU_ROUTE_KEYS,
        "updated_by": item.updated_by, "updated_at": item.updated_at,
    }


@app.get(f"{settings.api_prefix}/system/menus/navigation")
async def navigation_menus(_: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    items = [
        item for item in (await db.scalars(
            select(SystemMenu).where(SystemMenu.is_active.is_(True), SystemMenu.is_visible.is_(True)).order_by(SystemMenu.sort_order, SystemMenu.id)
        )).all()
        if item.key in SYSTEM_MENU_ROUTE_KEYS
    ]
    visible_keys = {item.key for item in items if not item.parent_key}
    changed = True
    while changed:
        changed = False
        for item in items:
            if item.key not in visible_keys and item.parent_key in visible_keys:
                visible_keys.add(item.key); changed = True
    return {"items": [_system_menu_dict(item) for item in items if item.key in visible_keys]}


@app.get(f"{settings.api_prefix}/system/menus")
async def list_system_menus(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    items = (await db.scalars(select(SystemMenu).order_by(SystemMenu.sort_order, SystemMenu.id))).all()
    return {"items": [_system_menu_dict(item) for item in items], "total": len(items)}


@app.post(f"{settings.api_prefix}/system/menus", status_code=status.HTTP_201_CREATED)
async def create_system_menu(body: SystemMenuInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    menu_key = body.key.strip()
    parent_key = body.parent_key.strip()
    if menu_key not in SYSTEM_MENU_ROUTE_KEYS:
        raise HTTPException(status_code=422, detail="菜单标识不是已实现的系统路由，不能创建菜单入口")
    if await db.scalar(select(SystemMenu.id).where(SystemMenu.key == menu_key)):
        raise HTTPException(status_code=409, detail="菜单标识已经存在")
    if parent_key and not await db.scalar(select(SystemMenu.id).where(SystemMenu.key == parent_key)):
        raise HTTPException(status_code=422, detail="父级菜单不存在")
    item = SystemMenu(
        key=menu_key,
        parent_key=parent_key,
        label=body.label.strip(),
        icon=body.icon.strip(),
        sort_order=body.sort_order,
        is_visible=body.is_visible,
        is_active=body.is_active,
        updated_by=identity["username"],
    )
    db.add(item)
    await db.commit()
    await db.refresh(item)
    return _system_menu_dict(item)


@app.patch(f"{settings.api_prefix}/system/menus/{{menu_id}}")
async def update_system_menu(menu_id: int, body: SystemMenuUpdate, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    item = await db.get(SystemMenu, menu_id)
    if not item: raise HTTPException(status_code=404, detail="菜单不存在")
    changes = body.model_dump(exclude_none=True)
    if item.key in {"dashboard", "system", "system-management"} and (changes.get("is_active") is False or changes.get("is_visible") is False):
        raise HTTPException(status_code=422, detail="控制台和系统管理入口不能隐藏或停用")
    for key, value in changes.items():
        setattr(item, key, value.strip() if isinstance(value, str) else value)
    item.updated_by = identity["username"]
    await db.commit(); await db.refresh(item)
    return _system_menu_dict(item)


@app.delete(f"{settings.api_prefix}/system/menus/{{menu_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_system_menu(menu_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    item = await db.get(SystemMenu, menu_id)
    if not item:
        raise HTTPException(status_code=404, detail="菜单不存在")
    if item.key in {key for key, *_ in DEFAULT_SYSTEM_MENUS}:
        raise HTTPException(status_code=422, detail="系统预置菜单不能删除")
    if await db.scalar(select(SystemMenu.id).where(SystemMenu.parent_key == item.key)):
        raise HTTPException(status_code=409, detail="请先删除该菜单的子菜单")
    await db.delete(item)
    await db.commit()


@app.post(f"{settings.api_prefix}/system/menus/reset")
async def reset_system_menus(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    defaults = {key: (parent_key, label, icon, sort_order) for key, parent_key, label, icon, sort_order in DEFAULT_SYSTEM_MENUS}
    items = (await db.scalars(select(SystemMenu))).all()
    by_key = {item.key: item for item in items}
    for item in items:
        if item.key not in defaults:
            item.is_visible = False; item.is_active = False; item.updated_by = identity["username"]
    for key, (parent_key, label, icon, sort_order) in defaults.items():
        item = by_key.get(key)
        if not item:
            db.add(SystemMenu(key=key, parent_key=parent_key, label=label, icon=icon, sort_order=sort_order, updated_by=identity["username"]))
            continue
        item.parent_key = parent_key; item.label = label; item.icon = icon; item.sort_order = sort_order
        item.is_visible = True; item.is_active = True; item.updated_by = identity["username"]
    await db.commit()
    refreshed = (await db.scalars(select(SystemMenu).order_by(SystemMenu.sort_order, SystemMenu.id))).all()
    return {"items": [_system_menu_dict(item) for item in refreshed], "total": len(refreshed)}


def _role_permission_dict(item: RolePermission) -> dict:
    if item.role == "admin":
        config = DEFAULT_ROLE_PERMISSIONS["admin"]
        return {"role": item.role, "display_name": config["display_name"], "data_scope": config["data_scope"], "menu_keys": list(MENU_KEYS), "field_keys": list(FIELD_KEYS), "updated_at": item.updated_at}
    return {"role": item.role, "display_name": item.display_name, "data_scope": item.data_scope, "menu_keys": item.menu_keys, "field_keys": item.field_keys, "updated_at": item.updated_at}


@app.get(f"{settings.api_prefix}/system/role-permissions")
async def list_role_permissions(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    items = (await db.scalars(select(RolePermission).order_by(RolePermission.id))).all()
    return {"items": [_role_permission_dict(item) for item in items], "available_menu_keys": MENU_KEYS, "available_field_keys": FIELD_KEYS}


@app.patch(f"{settings.api_prefix}/system/role-permissions/{{role}}")
async def update_role_permission(role: str, body: RolePermissionUpdate, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    if role not in DEFAULT_ROLE_PERMISSIONS:
        raise HTTPException(status_code=404, detail="角色不存在")
    invalid = sorted(set(body.menu_keys) - set(MENU_KEYS))
    if invalid:
        raise HTTPException(status_code=422, detail=f"无效菜单权限：{', '.join(invalid)}")
    menu_keys = list(dict.fromkeys(body.menu_keys))
    if "user-center" not in menu_keys:
        raise HTTPException(status_code=422, detail="用户中心为基础权限，不能移除")
    if role == "admin" and set(menu_keys) != set(MENU_KEYS):
        raise HTTPException(status_code=422, detail="系统管理员必须保留全部菜单权限")
    if role == "admin" and body.data_scope.strip() != DEFAULT_ROLE_PERMISSIONS["admin"]["data_scope"]:
        raise HTTPException(status_code=422, detail="系统管理员必须保留全所数据权限")
    invalid_fields = sorted(set(body.field_keys) - set(FIELD_KEYS))
    if invalid_fields: raise HTTPException(status_code=422, detail=f"无效字段权限：{', '.join(invalid_fields)}")
    field_keys = list(dict.fromkeys(body.field_keys))
    if role == "admin" and set(field_keys) != set(FIELD_KEYS): raise HTTPException(status_code=422, detail="系统管理员必须保留全部字段权限")
    item = await db.scalar(select(RolePermission).where(RolePermission.role == role))
    if not item:
        config = DEFAULT_ROLE_PERMISSIONS[role]
        item = RolePermission(role=role, display_name=config["display_name"], data_scope=body.data_scope, menu_keys=menu_keys, field_keys=field_keys)
        db.add(item)
    else:
        item.data_scope = body.data_scope.strip()
        item.menu_keys = menu_keys
        item.field_keys = field_keys
    await db.commit(); await db.refresh(item)
    return _role_permission_dict(item)


@app.get(f"{settings.api_prefix}/dashboard")
async def dashboard(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    scope = await _record_scope_conditions(identity, db)
    records = (await db.scalars(select(BusinessRecord).where(*scope))).all()
    by_module = {module: [item for item in records if item.module == module] for module in {"case", "task", "finance", "refund", "contract", "clue", "seal"}}
    cases, tasks = by_module["case"], by_module["task"]
    finances = by_module["finance"] + by_module["refund"]
    username = identity["username"]
    pending_statuses = {"待审批", "审批中", "待审核"}
    def count(rows, statuses, owner_only=False, predicate=None):
        return sum(item.status in statuses and (not owner_only or item.owner == username) and (predicate is None or predicate(item.data or {})) for item in rows)
    def fee_match(word): return lambda data: word in str(data.get("fee_type") or data.get("expense_scope") or "")
    official = lambda data: any(word in str(data.get("fee_type") or data.get("expense_scope") or "") for word in ("官方", "律所"))
    unpaid_official = [item for item in finances if official(item.data or {}) and item.status not in {"已付款", "已完成", "已驳回", "已拒绝"}]
    unpaid_amount = sum(max(float((item.data or {}).get("amount") or 0) - float((item.data or {}).get("paid_amount") or 0), 0) for item in unpaid_official)
    metrics = [
        {"label": "待缴官费", "value": f"{len(unpaid_official)}件", "tone": "amber"},
        {"label": "待退费", "value": f"{sum(item.status not in {'已完成','已驳回','已拒绝'} for item in by_module['refund'])}件", "tone": "cyan"},
        {"label": "补充证据", "value": f"{sum('补充证据' in item.status for item in cases)}件", "tone": "green"},
        {"label": "补充意见", "value": f"{sum('补充意见' in item.status for item in cases)}件", "tone": "blue"},
        {"label": "待上诉", "value": f"{sum('待上诉' in item.status for item in cases)}件", "tone": "red"},
        {"label": "待执行", "value": f"{sum('待执行' in item.status for item in cases)}件", "tone": "purple"},
        {"label": "紧急案件", "value": f"{sum(bool((item.data or {}).get('urgent')) or '紧急' in item.status for item in cases)}件", "tone": "orange"},
        {"label": "未到官费金额", "value": f"{unpaid_amount:.2f}元", "tone": "navy"},
    ]
    todo_specs = [
        ("待处理任务", tasks, {"待接收", "待处理", "处理中"}, None, "待审批官方费用", finances, pending_statuses, official),
        ("待审批线索", by_module["clue"], pending_statuses, None, "待审批内部费用", finances, pending_statuses, fee_match("内部")),
        ("待审批合同", by_module["contract"], pending_statuses, None, "待审批结算费用", finances, pending_statuses, fee_match("结算")),
        ("待审批用印", by_module["seal"], pending_statuses, None, "待审批归档费用", finances, pending_statuses, fee_match("归档")),
        ("待审核归档", cases, {"待归档审核"}, None, "待审核预损费用", finances, pending_statuses, fee_match("预损")),
    ]
    todos = [[left, count(left_rows, left_states, True, left_pred), count(left_rows, left_states, False, left_pred), right, count(right_rows, right_states, True, right_pred), count(right_rows, right_states, False, right_pred)] for left, left_rows, left_states, left_pred, right, right_rows, right_states, right_pred in todo_specs]
    current_month = date.today().replace(day=1); month_keys = []
    for offset in range(9, -1, -1):
        year, month = current_month.year, current_month.month - offset
        while month <= 0: year -= 1; month += 12
        month_keys.append(f"{year:04d}-{month:02d}")
    case_trend = [{"date": key, "value": sum(item.created_at.strftime("%Y-%m") == key for item in cases)} for key in month_keys]
    stage_groups = [("立案待分配", lambda s: s in {"新案待分配", "立案待分配"}, "#f7474c"), ("文书准备", lambda s: "文书准备" in s, "#46b8b8"), ("一审", lambda s: "一审" in s, "#ffb45a"), ("二审", lambda s: "二审" in s, "#7f70b3"), ("再审", lambda s: "再审" in s, "#98a5b7"), ("执行", lambda s: "执行" in s, "#303030")]
    stage_counts = {label: 0 for label, _, _ in stage_groups}; other_count = 0
    for item in cases:
        matched = next((label for label, match, _ in stage_groups if match(item.status)), None)
        if matched: stage_counts[matched] += 1
        else: other_count += 1
    civil_distribution = [{"label": label, "value": stage_counts[label], "color": color} for label, _, color in stage_groups]
    if other_count: civil_distribution.append({"label": "其他", "value": other_count, "color": "#c5cbd3"})
    case_map = {item.id: item for item in cases}; visible_case_ids = set(case_map)
    hearing_rows = (await db.scalars(select(HearingSchedule).where(HearingSchedule.case_record_id.in_(visible_case_ids), HearingSchedule.hearing_date >= date.today()).order_by(HearingSchedule.hearing_date, HearingSchedule.hearing_time).limit(15))).all() if visible_case_ids else []
    weekdays = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]
    hearings = []
    for item in hearing_rows:
        case = case_map[item.case_record_id]; data = case.data or {}
        hearings.append({"weekday": weekdays[item.hearing_date.weekday()], "date": str(item.hearing_date), "time": item.hearing_time, "court": item.court, "case_no": case.serial_no, "client": case.customer, "lawyer": item.hearing_lawyer or data.get("hearing_lawyer", ""), "agent": ",".join(data.get("handling_lawyers", [])), "assistant": data.get("assistant", "")})
    latest_cases = []
    for item in sorted(cases, key=lambda value: (value.created_at, value.id), reverse=True)[:15]:
        data = item.data or {}
        latest_cases.append({"case_no": item.serial_no, "stage": item.status, "plaintiff": data.get("plaintiff") or item.customer, "defendant": data.get("opponent", ""), "date": str(item.created_at.date()), "manager": data.get("customer_manager", ""), "lawyer": data.get("hearing_lawyer", ""), "agent": ",".join(data.get("handling_lawyers", [])), "assistant": data.get("assistant", "")})
    return {"metrics": metrics, "todos": todos, "case_trend": case_trend, "civil_distribution": civil_distribution, "hearings": hearings, "latest_cases": latest_cases, "source": "realtime"}


@app.get(f"{settings.api_prefix}/search")
async def global_search(q: str = Query(min_length=2, max_length=100), identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    keyword = q.strip(); like = f"%{keyword}%"
    records = (await db.scalars(select(BusinessRecord).where(or_(BusinessRecord.serial_no.ilike(like), BusinessRecord.title.ilike(like), BusinessRecord.customer.ilike(like), BusinessRecord.owner.ilike(like), BusinessRecord.description.ilike(like)), *(await _record_scope_conditions(identity, db))).order_by(BusinessRecord.updated_at.desc()).limit(50))).all()
    attachments = (await db.scalars(select(FileAttachment).where(or_(FileAttachment.original_name.ilike(like), FileAttachment.remark.ilike(like))).order_by(FileAttachment.created_at.desc()).limit(20))).all()
    attachments = await _filter_visible_attachments(attachments, identity, db)
    templates = (await db.scalars(select(DocumentTemplate).where(or_(DocumentTemplate.name.ilike(like), DocumentTemplate.description.ilike(like))).order_by(DocumentTemplate.updated_at.desc()).limit(20))).all()
    route_map = {"customer": "customer-company", "contract": "contract-mine", "case": "case-company", "task": "task-company", "clue": "clue", "notary": "notary", "evidence": "evidence", "seal": "seal-my", "finance": "finance-fee-query", "document": "documents-register", "hr": "hr-all", "warehouse": "warehouse", "report": "reports"}
    items = [{"type": "record", "id": x.id, "module": x.module, "route": route_map.get(x.module, "dashboard"), "serial_no": x.serial_no, "title": x.title, "subtitle": x.customer or x.description, "status": x.status, "updated_at": x.updated_at} for x in records]
    items.extend({"type": "attachment", "id": x.id, "module": "attachment", "route": "documents-files", "serial_no": "附件", "title": x.original_name, "subtitle": f"{x.category}｜{x.remark}", "status": "", "updated_at": x.created_at} for x in attachments)
    items.extend({"type": "template", "id": x.id, "module": "template", "route": "documents-template", "serial_no": "模板", "title": x.name, "subtitle": f"{x.category}｜{x.description}", "status": "启用" if x.is_active else "停用", "updated_at": x.updated_at} for x in templates)
    return {"query": keyword, "items": items, "total": len(items)}


def _notification_dict(item: Notification) -> dict:
    return {"id": item.id, "source_type": item.source_type, "source_id": item.source_id, "sender": item.sender, "recipient": item.recipient, "notification_type": item.notification_type, "title": item.title, "content": item.content, "level": item.level, "is_read": item.is_read, "read_at": item.read_at, "created_at": item.created_at}


async def _sync_notifications(identity: dict, db: AsyncSession) -> None:
    username = identity["username"]; today = date.today(); candidates: list[dict] = []
    task_terminal_statuses = ["已完成", "待确认", "已验收", "已拒绝", "已撤回", "已停止", "已取消"]
    all_tasks = list((await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "task"))).all())
    tasks = [task for task in all_tasks if task.status not in task_terminal_statuses]
    terminal_task_ids = {task.id for task in all_tasks if task.status in task_terminal_statuses}
    current_user = await db.scalar(select(User).where(User.username == username))
    stale = (await db.scalars(select(Notification).where(Notification.recipient == username, Notification.source_type.in_(["task", "finance", "contract", "case"])))).all()
    existing_record_ids = set((await db.scalars(select(BusinessRecord.id))).all())
    if identity.get("role") != "admin":
        visible_tasks = [task for task in all_tasks if _is_task_participant(task, identity) or (identity.get("role") == "manager" and current_user and task.department == current_user.department)]
        tasks = [task for task in tasks if _is_task_participant(task, identity) or (identity.get("role") == "manager" and current_user and task.department == current_user.department)]
        visible_ids = await _visible_record_ids(identity, db)
        visible_task_ids = {task.id for task in visible_tasks}
    else:
        visible_ids = existing_record_ids
        visible_task_ids = {task.id for task in all_tasks}
    for notice in stale:
        is_auto_reminder = (
            (notice.source_type == "task" and notice.source_key.startswith("task-") and not notice.source_key.startswith(("task-history-", "task-message-")))
            or notice.source_key.startswith(("finance-approval-", "contract-approval-", "hearing-"))
        )
        # 旧版本的自动提醒键是全局唯一键，同一业务只能被第一个访问提醒页的
        # 用户取得。迁移为收件人维度，确保管理员与每个参与人都有独立提醒。
        if is_auto_reminder and not notice.source_key.endswith(f"-{username}"):
            recipient_key = f"{notice.source_key}-{username}"
            duplicate = await db.scalar(select(Notification).where(Notification.source_key == recipient_key))
            if duplicate:
                await db.delete(notice)
                continue
            notice.source_key = recipient_key
        allowed_ids = visible_task_ids if notice.source_type == "task" else visible_ids
        is_terminal_task_reminder = (
            notice.source_type == "task" and notice.source_id in terminal_task_ids
            and notice.source_key.startswith("task-")
            and not notice.source_key.startswith(("task-history-", "task-message-"))
        )
        if notice.source_id not in existing_record_ids or notice.source_id not in allowed_ids or is_terminal_task_reminder:
            await db.delete(notice)
    for task in tasks:
        info = _task_dict(task)
        if info["reminder_due"]:
            candidates.append({"source_key": f"task-{task.id}-{today}-{username}", "source_type": "task", "source_id": task.id, "title": info["reminder_text"], "content": f"{task.serial_no}｜{task.title}｜负责人：{task.owner}", "level": "error" if info["status"] == "已逾期" else "warning"})
    fees = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "finance", BusinessRecord.status == "待审批", *(await _record_scope_conditions(identity, db))))).all()
    for fee in fees:
        candidates.append({"source_key": f"finance-approval-{fee.id}-{username}", "source_type": "finance", "source_id": fee.id, "title": "费用待审批", "content": f"{fee.serial_no}｜{fee.title}｜{(fee.data or {}).get('amount', 0)} 元", "level": "warning"})
    current_steps = (await db.execute(select(ContractApprovalStep, BusinessRecord).join(BusinessRecord, BusinessRecord.id == ContractApprovalStep.contract_record_id).where(ContractApprovalStep.status == "待审批", BusinessRecord.status == "审批中"))).all()
    for step, contract in current_steps:
        if identity.get("role") == "admin" or step.approver == username:
            candidates.append({"source_key": f"contract-approval-{contract.id}-{step.id}-{username}", "source_type": "contract", "source_id": contract.id, "title": f"合同第 {step.step_order} 级待审批", "content": f"{contract.serial_no}｜{contract.title}｜审批人：{step.approver}", "level": "warning"})
    hearings = (await db.execute(select(HearingSchedule, BusinessRecord).join(BusinessRecord, BusinessRecord.id == HearingSchedule.case_record_id).where(HearingSchedule.hearing_date == today + timedelta(days=1), HearingSchedule.status == "已排期", *(await _record_scope_conditions(identity, db))))).all()
    for hearing, case_record in hearings:
        candidates.append({"source_key": f"hearing-{hearing.id}-{hearing.hearing_date}-{username}", "source_type": "case", "source_id": case_record.id, "title": "明日开庭提醒", "content": f"{case_record.serial_no}｜{hearing.hearing_time}｜{hearing.court}｜{hearing.hearing_lawyer}", "level": "info"})
    existing_keys = set((await db.scalars(select(Notification.source_key).where(Notification.source_key.in_([x["source_key"] for x in candidates])))).all()) if candidates else set()
    for item in candidates:
        if item["source_key"] not in existing_keys: db.add(Notification(**item, recipient=username))
    await db.commit()


@app.get(f"{settings.api_prefix}/notifications")
async def list_notifications(
    unread_only: bool = False, sent_only: bool = False, read_status: str = "", sender: str = "",
    keyword: str = "", notification_type: str = "", source_type: str = "", level: str = "",
    reminder_only: bool = False, date_from: date | None = None, date_to: date | None = None,
    page: int = Query(1, ge=1), page_size: int = Query(100, ge=1, le=200),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    await _sync_notifications(identity, db)
    if date_from and date_to and date_from > date_to: raise HTTPException(status_code=422, detail="开始日期不能晚于结束日期")
    conditions = [Notification.sender == identity["username"], Notification.sender_deleted.is_(False)] if sent_only else [Notification.recipient == identity["username"], Notification.recipient_deleted.is_(False)]
    if unread_only or read_status == "未读": conditions.append(Notification.is_read.is_(False))
    elif read_status == "已读": conditions.append(Notification.is_read.is_(True))
    elif read_status not in {"", "全部"}: raise HTTPException(status_code=422, detail="消息状态无效")
    if sender.strip(): conditions.append(Notification.sender.ilike(f"%{sender.strip()}%"))
    if keyword.strip():
        term = f"%{keyword.strip()}%"; conditions.append(or_(Notification.title.ilike(term), Notification.content.ilike(term)))
    if notification_type:
        if notification_type not in {"系统通知", "用户通知"}: raise HTTPException(status_code=422, detail="消息类型无效")
        conditions.append(Notification.notification_type == notification_type)
    if source_type:
        if source_type not in {"task", "finance", "contract", "case", "message"}: raise HTTPException(status_code=422, detail="消息来源无效")
        conditions.append(Notification.source_type == source_type)
    if level:
        if level not in {"info", "warning", "error"}: raise HTTPException(status_code=422, detail="提醒级别无效")
        conditions.append(Notification.level == level)
    if reminder_only:
        conditions.extend([
            Notification.source_type == "task",
            Notification.source_key.like("task-%"),
            ~Notification.source_key.like("task-message-%"),
            ~Notification.source_key.like("task-history-%"),
        ])
    if date_from: conditions.append(Notification.created_at >= datetime.combine(date_from, datetime.min.time()))
    if date_to: conditions.append(Notification.created_at <= datetime.combine(date_to, datetime.max.time()))
    total = int(await db.scalar(select(func.count()).select_from(Notification).where(*conditions)) or 0)
    items = (await db.scalars(select(Notification).where(*conditions).order_by(Notification.created_at.desc(), Notification.id.desc()).offset((page - 1) * page_size).limit(page_size))).all()
    unread = int(await db.scalar(select(func.count()).select_from(Notification).where(Notification.recipient == identity["username"], Notification.recipient_deleted.is_(False), Notification.is_read.is_(False))) or 0)
    return {"items": [_notification_dict(x) for x in items], "unread": unread, "total": total, "page": page, "page_size": page_size}


@app.get(f"{settings.api_prefix}/users/directory")
async def user_directory(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    items = (await db.scalars(select(User).where(User.is_active.is_(True)).order_by(User.display_name, User.username))).all()
    job_roles = (await db.scalars(select(JobRole).where(JobRole.is_active.is_(True)))).all()
    roles_by_name = {item.name: item for item in job_roles}
    payload = []
    for item in items:
        position = str((item.profile or {}).get("position") or (item.profile or {}).get("staff_role") or "")
        job_role = roles_by_name.get(position)
        job_permissions = list(job_role.permissions or []) if job_role else []
        payload.append({
            "username": item.username,
            "display_name": item.display_name,
            "department": item.department,
            "is_active": item.is_active,
            "role": item.role,
            "position": position,
            "staff_role": str((item.profile or {}).get("staff_role") or ""),
            "job_permissions": job_permissions,
            "can_approve_contract": item.role == "admin" or "合同审批" in set(job_permissions),
        })
    return {"items": payload}


@app.post(f"{settings.api_prefix}/notifications/send", status_code=status.HTTP_201_CREATED)
async def send_user_message(body: UserMessageInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    recipients = list(dict.fromkeys(value.strip() for value in body.recipients if value.strip()))
    active_users = set((await db.scalars(select(User.username).where(User.username.in_(recipients), User.is_active.is_(True)))).all())
    missing = sorted(set(recipients) - active_users)
    if missing: raise HTTPException(status_code=422, detail=f"接收人不存在或已停用：{', '.join(missing)}")
    batch = uuid4().hex
    items = []
    for recipient in recipients:
        item = Notification(source_key=f"user-message-{batch}-{recipient}", source_type="message", sender=identity["username"], recipient=recipient, notification_type="用户通知", title=body.title.strip(), content=body.content.strip(), level="info")
        db.add(item); items.append(item)
    await db.commit()
    for item in items: await db.refresh(item)
    return {"items": [_notification_dict(item) for item in items], "sent": len(items)}


@app.post(f"{settings.api_prefix}/notifications/{{notification_id}}/read")
async def read_notification(notification_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await db.get(Notification, notification_id)
    if not item or item.recipient != identity["username"] or item.recipient_deleted: raise HTTPException(status_code=404, detail="消息不存在")
    item.is_read = True; item.read_at = datetime.now(); await db.commit(); await db.refresh(item)
    return _notification_dict(item)


@app.post(f"{settings.api_prefix}/notifications/read-all")
async def read_all_notifications(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    items = (await db.scalars(select(Notification).where(Notification.recipient == identity["username"], Notification.recipient_deleted.is_(False), Notification.is_read.is_(False)))).all()
    now = datetime.now()
    for item in items: item.is_read = True; item.read_at = now
    await db.commit(); return {"updated": len(items)}


@app.delete(f"{settings.api_prefix}/notifications/{{notification_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_notification(notification_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await db.get(Notification, notification_id)
    if not item: raise HTTPException(status_code=404, detail="消息不存在")
    changed = False
    if item.recipient == identity["username"] and not item.recipient_deleted: item.recipient_deleted = True; changed = True
    if item.sender == identity["username"] and not item.sender_deleted: item.sender_deleted = True; changed = True
    if not changed: raise HTTPException(status_code=404, detail="消息不存在")
    if item.recipient_deleted and item.sender_deleted: await db.delete(item)
    await db.commit(); return Response(status_code=status.HTTP_204_NO_CONTENT)


def _communication_dict(item: CommunicationLog) -> dict:
    return {"id": item.id, "customer_record_id": item.customer_record_id, "customer_name": item.customer_name, "contact": item.contact, "phone": item.phone, "content": item.content, "occurred_at": item.occurred_at, "operator": item.operator, "created_at": item.created_at, "updated_at": item.updated_at}


def _parse_customer_contact_at(value: object) -> datetime | None:
    raw_value = str(value or "").strip()
    if not raw_value:
        return None
    try:
        parsed = datetime.fromisoformat(raw_value.replace("Z", "+00:00"))
    except ValueError:
        return None
    # Normalize explicit offsets and Z to one UTC-naive timeline.  Historical
    # naive values keep their wall-clock meaning for backward compatibility.
    if parsed.tzinfo is not None:
        parsed = parsed.astimezone(timezone.utc).replace(tzinfo=None)
    return parsed


def _sync_customer_contact_metrics(customer: BusinessRecord) -> None:
    """Recompute the denormalized contact date/count from real contact notes.

    Communication-log rows are mirrored into ``data.notes`` with the same
    stable note id.  Keeping one source of truth here also covers ordinary
    customer follow-up notes and makes edits/deletes recalculate the maximum
    event time instead of blindly overwriting it with the last API request.
    Customer-directory changes (for example adding a contact person) are not
    communication events and must never call this helper.
    """
    data = dict(customer.data or {})
    notes = [dict(note) for note in list(data.get("notes", [])) if isinstance(note, dict)]
    occurred_values: list[tuple[datetime, str]] = []
    for note in notes:
        raw_value = str(note.get("created_at") or "").strip()
        parsed = _parse_customer_contact_at(raw_value)
        if parsed is not None:
            occurred_values.append((parsed, raw_value))
    latest = max(occurred_values, key=lambda value: value[0])[1] if occurred_values else ""
    customer.data = {**data, "notes": notes, "last_contact_at": latest, "contact_count": len(occurred_values)}


CUSTOMER_MODIFICATION_ACTIONS = {
    "创建客户", "批量导入", "编辑", "编辑变更",
    "领取客户", "释放公海", "共享客户", "移入回收站", "恢复客户",
    "新增联系人", "删除联系人", "更新客户管理人",
    "新增客户跟进", "删除客户跟进",
    "新增沟通日志", "修改沟通日志", "删除沟通日志",
}


def _mark_customer_modified(customer: BusinessRecord, identity: dict) -> None:
    """Maintain the actor used by the original ``最近更新的客户`` projection.

    ``updated_at`` is the authoritative modification time maintained by the
    database.  The actor is kept in customer JSON because the legacy schema has
    no dedicated column.  A historical ``last_modified_date`` value must not
    override the authoritative database timestamp after a real local change.
    """
    data = dict(customer.data or {})
    data.pop("last_modified_date", None)
    data["last_modified_by"] = identity["username"]
    customer.data = data


@app.get(f"{settings.api_prefix}/communications")
async def list_communications(keyword: str = "", date_from: date | None = None, date_to: date | None = None, mine_only: bool = True, page: int = Query(1, ge=1), page_size: int = Query(50, ge=1, le=200), identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    conditions = []
    if identity.get("role") != "admin" or mine_only: conditions.append(CommunicationLog.operator == identity["username"])
    if identity.get("role") != "admin": conditions.append(CommunicationLog.customer_record_id.in_(await _visible_record_ids(identity, db)))
    if keyword.strip():
        term = f"%{keyword.strip()}%"; conditions.append(or_(CommunicationLog.customer_name.ilike(term), CommunicationLog.contact.ilike(term), CommunicationLog.phone.ilike(term), CommunicationLog.content.ilike(term)))
    if date_from: conditions.append(CommunicationLog.occurred_at >= datetime.combine(date_from, datetime.min.time()))
    if date_to: conditions.append(CommunicationLog.occurred_at <= datetime.combine(date_to, datetime.max.time()))
    total = int(await db.scalar(select(func.count()).select_from(CommunicationLog).where(*conditions)) or 0)
    items = (await db.scalars(select(CommunicationLog).where(*conditions).order_by(CommunicationLog.occurred_at.desc(), CommunicationLog.id.desc()).offset((page - 1) * page_size).limit(page_size))).all()
    return {"items": [_communication_dict(item) for item in items], "total": total, "page": page, "page_size": page_size}


@app.post(f"{settings.api_prefix}/communications", status_code=status.HTTP_201_CREATED)
async def create_communication(body: CommunicationLogInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    customer = await _customer_or_404(body.customer_record_id, identity, db)
    await _require_record_owner_or_manager(customer, identity, db)
    occurred_at = _parse_customer_contact_at(body.occurred_at)
    if occurred_at is None: raise HTTPException(status_code=422, detail="沟通时间格式无效")
    if occurred_at > datetime.now() + timedelta(minutes=5): raise HTTPException(status_code=422, detail="沟通时间不能晚于当前时间")
    note_id = uuid4().hex; contact, phone, content = body.contact.strip(), body.phone.strip(), body.content.strip()
    note = {"id": note_id, "type": "沟通日志", "content": content, "operator": identity["username"], "contact": contact, "phone": phone, "created_at": body.occurred_at.isoformat(timespec="seconds")}
    customer.data = {**(customer.data or {}), "notes": [note, *list((customer.data or {}).get("notes", []))]}
    _sync_customer_contact_metrics(customer)
    item = CommunicationLog(customer_record_id=customer.id, customer_name=customer.title, contact=contact, phone=phone, content=content, occurred_at=body.occurred_at, operator=identity["username"], note_id=note_id)
    db.add(item); db.add(_customer_event(customer, "新增沟通日志", identity, f"{contact or '客户联系人'}：{content[:120]}"))
    await db.commit(); await db.refresh(item); return _communication_dict(item)


@app.patch(f"{settings.api_prefix}/communications/{{communication_id}}")
async def update_communication(communication_id: int, body: CommunicationLogUpdate, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await db.get(CommunicationLog, communication_id)
    if not item or (identity.get("role") != "admin" and item.operator != identity["username"]): raise HTTPException(status_code=404, detail="沟通记录不存在")
    customer = await _customer_or_404(item.customer_record_id, identity, db)
    await _require_record_owner_or_manager(customer, identity, db)
    if body.occurred_at:
        occurred_at = _parse_customer_contact_at(body.occurred_at)
        if occurred_at is None: raise HTTPException(status_code=422, detail="沟通时间格式无效")
        if occurred_at > datetime.now() + timedelta(minutes=5): raise HTTPException(status_code=422, detail="沟通时间不能晚于当前时间")
    changes = body.model_dump(exclude_unset=True, exclude_none=True)
    for key, value in changes.items(): setattr(item, key, value.strip() if key in {"contact", "phone", "content"} else value)
    data = dict(customer.data or {}); note_updates = {"content": item.content, "contact": item.contact, "phone": item.phone, "created_at": item.occurred_at.isoformat(timespec="seconds")}
    notes = [{**dict(note), **note_updates} if note.get("id") == item.note_id else dict(note) for note in list(data.get("notes", []))]
    customer.data = {**data, "notes": notes}
    _sync_customer_contact_metrics(customer)
    db.add(_customer_event(customer, "修改沟通日志", identity, item.content[:120]))
    await db.commit(); await db.refresh(item); return _communication_dict(item)


@app.delete(f"{settings.api_prefix}/communications/{{communication_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_communication(communication_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await db.get(CommunicationLog, communication_id)
    if not item or (identity.get("role") != "admin" and item.operator != identity["username"]): raise HTTPException(status_code=404, detail="沟通记录不存在")
    customer = await db.get(BusinessRecord, item.customer_record_id)
    if customer:
        customer = await _customer_or_404(customer.id, identity, db)
        await _require_record_owner_or_manager(customer, identity, db)
        data = customer.data or {}; customer.data = {**data, "notes": [note for note in list(data.get("notes", [])) if note.get("id") != item.note_id]}
        _sync_customer_contact_metrics(customer)
        db.add(_customer_event(customer, "删除沟通日志", identity, item.content[:120]))
    await db.delete(item); await db.commit(); return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.get(f"{settings.api_prefix}/reports/summary")
async def report_summary(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    modules = ["customer", "contract", "case", "task", "finance", "hr", "warehouse"]
    scope_conditions = await _record_scope_conditions(identity, db)
    counts = {module: int(await db.scalar(select(func.count()).select_from(BusinessRecord).where(BusinessRecord.module == module, *scope_conditions)) or 0) for module in modules}
    reports = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "report", *scope_conditions))).all()
    return {"business_counts": counts, "total_reports": len(reports), "generated": sum(x.status == "已生成" for x in reports), "published": sum(x.status == "已发布" for x in reports)}


async def _report_analytics(view: str, identity: dict, db: AsyncSession, customer: str = "", court_lawyer: str = "", handling_lawyer: str = "", assistant: str = "", investigator: str = "", court: str = "", source_from: date | None = None, source_to: date | None = None, hearing_from: date | None = None, hearing_to: date | None = None, group_mode: str = "") -> dict:
    if view not in {"brand", "lawyer", "refund", "execution-1", "execution-2", "execution-3"}: raise HTTPException(status_code=422, detail="不支持的统计视图")
    if source_from and source_to and source_from > source_to: raise HTTPException(status_code=422, detail="案源开始日期不能晚于结束日期")
    if hearing_from and hearing_to and hearing_from > hearing_to: raise HTTPException(status_code=422, detail="开庭开始日期不能晚于结束日期")
    if group_mode and group_mode not in {"按律师分组统计", "按文书分组统计"}: raise HTTPException(status_code=422, detail="不支持的分组模式")
    scope = await _record_scope_conditions(identity, db)
    cases = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "case", *scope))).all()
    customers_selected = {value.strip() for value in customer.split(",") if value.strip()}
    court_lawyers_selected = {value.strip() for value in court_lawyer.split(",") if value.strip()}
    def matches(item: BusinessRecord) -> bool:
        data = item.data or {}; handlers = data.get("handling_lawyers", [])
        return (not customers_selected or any(value in item.customer for value in customers_selected)) and (not court_lawyers_selected or data.get("hearing_lawyer") in court_lawyers_selected) and (not handling_lawyer or handling_lawyer in handlers) and (not assistant or assistant in str(data.get("assistant", ""))) and (not investigator or investigator in str(data.get("investigator", "")))
    cases = [item for item in cases if matches(item) and (not source_from or item.created_at.date() >= source_from) and (not source_to or item.created_at.date() <= source_to)]
    if court or hearing_from or hearing_to:
        case_ids = {item.id for item in cases}
        conditions = [HearingSchedule.case_record_id.in_(case_ids)]
        if court: conditions.append(HearingSchedule.court.ilike(f"%{court}%"))
        if hearing_from: conditions.append(HearingSchedule.hearing_date >= hearing_from)
        if hearing_to: conditions.append(HearingSchedule.hearing_date <= hearing_to)
        heard_case_ids = set((await db.scalars(select(HearingSchedule.case_record_id).where(*conditions))).all()) if case_ids else set()
        cases = [item for item in cases if item.id in heard_case_ids]
    case_by_no = {item.serial_no: item for item in cases}
    customers = sorted({item.customer for item in cases if item.customer}); lawyers = sorted({str((item.data or {}).get("hearing_lawyer")) for item in cases if (item.data or {}).get("hearing_lawyer")})
    if view in {"refund", "execution-1", "execution-2", "execution-3"}:
        if view == "refund":
            refunds = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "refund", *scope))).all()
            refunds = [item for item in refunds if not customers_selected or any(value in item.customer for value in customers_selected)]
            specs = [("准备资料进度案件数量", {"草稿", "准备资料", "待提交"}), ("客户盖章进度案件数量", {"客户盖章"}), ("提交法院进度案件数量", {"提交法院", "已提交法院"}), ("等待客户回款进度案件数量", {"等待客户回款", "待回款"})]
            charts = [{"title": title, "unit": "个/案", "items": [{"name": "案件数", "value": sum(item.status in statuses for item in refunds)}]} for title, statuses in specs]
        else:
            execution_statuses = ["一审待执行", "二审待执行", "准备材料", "提交法院", "执行受理", "执行中止", "执行结案", "执行终本", "执行终结", "执行中止"]
            charts = [{"title": f"{status}案件数量", "unit": "个/案", "items": [{"name": "案件数", "value": sum(status in item.status for item in cases)}]} for status in execution_statuses]
        return {"view": view, "charts": charts, "filter_options": {"customers": customers, "lawyers": lawyers}, "source": "realtime"}
    can_view_amount = "finance.amount" in await _allowed_field_keys(identity, db)
    finances = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "finance", *scope))).all() if can_view_amount else []
    grouped: dict[str, dict[str, float]] = {}
    for case in cases:
        data = case.data or {}
        group = case.customer if view == "brand" else str((data.get("document_lawyer") or data.get("assistant") or "未分配") if group_mode == "按文书分组统计" else (data.get("hearing_lawyer") or "未分配"))
        grouped.setdefault(group, {"expense": 0, "received": 0, "loss": 0, "cycle_total": 0, "cycle_count": 0})
    for fee in finances:
        data = fee.data or {}; case = case_by_no.get(str(data.get("case_no") or ""))
        if not case: continue
        case_data = case.data or {}
        group = case.customer if view == "brand" else str((case_data.get("document_lawyer") or case_data.get("assistant") or "未分配") if group_mode == "按文书分组统计" else (case_data.get("hearing_lawyer") or "未分配")); bucket = grouped[group]
        amount = float(data.get("amount") or 0); received = float(data.get("paid_amount") or data.get("received_amount") or 0); loss = float(data.get("loss_amount") or 0)
        bucket["expense"] += amount; bucket["received"] += received; bucket["loss"] += loss
        recovery_days = data.get("recovery_days")
        if recovery_days not in {None, ""}: bucket["cycle_total"] += float(recovery_days); bucket["cycle_count"] += 1
    def series(metric):
        items = []
        for name, value in sorted(grouped.items(), key=lambda pair: pair[0]):
            if metric == "cycle": result = value["cycle_total"] / value["cycle_count"] if value["cycle_count"] else 0
            elif metric == "loss": result = value["loss"]
            elif metric == "return": result = value["received"] / value["expense"] * 100 if value["expense"] else 0
            else: result = value["loss"] / value["expense"] * 100 if value["expense"] else 0
            items.append({"name": name, "value": round(result, 2)})
        return items
    charts = [{"title": "资金回款周期统计", "unit": "天/案", "items": series("cycle")}, {"title": "资金亏损金额统计", "unit": "元", "items": series("loss")}, {"title": "资金回报率统计", "unit": "百分比", "items": series("return")}, {"title": "资金亏损率统计", "unit": "百分比", "items": series("loss_rate")}]
    return {"view": view, "charts": charts, "filter_options": {"customers": customers, "lawyers": lawyers}, "source": "realtime"}


@app.get(f"{settings.api_prefix}/reports/analytics")
async def report_analytics(view: str, customer: str = "", court_lawyer: str = "", handling_lawyer: str = "", assistant: str = "", investigator: str = "", court: str = "", source_from: date | None = None, source_to: date | None = None, hearing_from: date | None = None, hearing_to: date | None = None, group_mode: str = "", identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    return await _report_analytics(view, identity, db, customer, court_lawyer, handling_lawyer, assistant, investigator, court, source_from, source_to, hearing_from, hearing_to, group_mode)


@app.get(f"{settings.api_prefix}/reports/analytics/export")
async def export_report_analytics(view: str, customer: str = "", court_lawyer: str = "", handling_lawyer: str = "", assistant: str = "", investigator: str = "", court: str = "", source_from: date | None = None, source_to: date | None = None, hearing_from: date | None = None, hearing_to: date | None = None, group_mode: str = "", identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    data = await _report_analytics(view, identity, db, customer, court_lawyer, handling_lawyer, assistant, investigator, court, source_from, source_to, hearing_from, hearing_to, group_mode)
    output = io.StringIO(); writer = csv.writer(output); writer.writerow(["统计图", "分组", "数值", "单位"])
    for chart in data["charts"]:
        for item in chart["items"]: writer.writerow([chart["title"], item["name"], item["value"], chart["unit"]])
    return Response(content=("\ufeff" + output.getvalue()).encode("utf-8"), media_type="text/csv; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="report-{view}-{date.today()}.csv"'})


@app.post(f"{settings.api_prefix}/reports/generate", status_code=status.HTTP_201_CREATED)
async def generate_report(body: ReportInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    supported = {"综合经营报表", "案件统计报表", "客户统计报表", "财务统计报表", "人事统计报表", "仓库统计报表"}
    if body.report_type not in supported:
        raise HTTPException(status_code=422, detail="不支持的报表类型")
    module_labels = {"customer": "客户数量", "contract": "合同数量", "case": "案件数量", "task": "任务数量", "finance": "费用数量", "hr": "员工数量", "warehouse": "库存物品种类"}
    type_modules = {"案件统计报表": ["case"], "客户统计报表": ["customer"], "财务统计报表": ["finance"], "人事统计报表": ["hr"], "仓库统计报表": ["warehouse"]}
    modules = type_modules.get(body.report_type, list(module_labels))
    scope_conditions = await _record_scope_conditions(identity, db)
    metrics: dict[str, int | float] = {}
    for module in modules:
        metrics[module_labels[module]] = int(await db.scalar(select(func.count()).select_from(BusinessRecord).where(BusinessRecord.module == module, *scope_conditions)) or 0)
    if "finance" in modules and "finance.amount" in await _allowed_field_keys(identity, db):
        fees = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "finance", *scope_conditions))).all()
        metrics["费用总额"] = round(sum(float((x.data or {}).get("amount", 0) or 0) for x in fees), 2)
        metrics["待审批费用"] = sum(x.status == "待审批" for x in fees)
    serial = f"BB{datetime.now().strftime('%Y%m%d%H%M%S')}{uuid4().hex[:4].upper()}"
    item = BusinessRecord(module="report", serial_no=serial, title=body.title, status="已生成", owner=identity["username"], department="上海分所", description=body.description, data={"report_type": body.report_type, "period": body.period, "format": body.format, "metrics": metrics, "generated_at": datetime.now().isoformat(timespec="seconds")})
    db.add(item); await db.flush()
    db.add(WorkflowEvent(record_id=item.id, action="生成报表", to_status="已生成", operator=identity["username"], comment=f"{body.report_type}｜{body.period}"))
    await db.commit(); await db.refresh(item)
    return _record_dict(item)


@app.get(f"{settings.api_prefix}/reports/{{report_id}}/download")
async def download_report(report_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    report = await _ensure_record_module(report_id, "report", identity, db)
    data = report.data or {}; metrics = data.get("metrics", {})
    lines = ["指标,数值"]
    for label, value in metrics.items():
        safe_label = str(label).replace('"', '""'); safe_value = str(value).replace('"', '""')
        lines.append(f'"{safe_label}","{safe_value}"')
    content = "\ufeff" + "\r\n".join(lines)
    return Response(content=content.encode("utf-8"), media_type="text/csv; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="{report.serial_no}.csv"'})


@app.get(f"{settings.api_prefix}/audit/events")
async def list_audit_events(module: str = "", keyword: str = "", page: int = Query(1, ge=1), page_size: int = Query(50, ge=1, le=200), identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") != "admin": raise HTTPException(status_code=403, detail="仅管理员可以查看全所操作日志")
    conditions = []
    if module: conditions.append(BusinessRecord.module == module)
    if keyword.strip():
        like = f"%{keyword.strip()}%"
        conditions.append(or_(BusinessRecord.serial_no.ilike(like), BusinessRecord.title.ilike(like), WorkflowEvent.action.ilike(like), WorkflowEvent.operator.ilike(like), WorkflowEvent.comment.ilike(like)))
    base = select(WorkflowEvent, BusinessRecord).join(BusinessRecord, BusinessRecord.id == WorkflowEvent.record_id).where(*conditions)
    total = int(await db.scalar(select(func.count()).select_from(WorkflowEvent).join(BusinessRecord, BusinessRecord.id == WorkflowEvent.record_id).where(*conditions)) or 0)
    result = (await db.execute(base.order_by(WorkflowEvent.created_at.desc()).offset((page - 1) * page_size).limit(page_size))).all()
    return {"items": [{"id": event.id, "record_id": record.id, "module": record.module, "serial_no": record.serial_no, "title": record.title, "action": event.action, "from_status": event.from_status, "to_status": event.to_status, "operator": event.operator, "comment": event.comment, "created_at": event.created_at} for event, record in result], "total": total, "page": page, "page_size": page_size}


@app.get(f"{settings.api_prefix}/records/export")
async def export_records(module: str = Query(min_length=1, max_length=32), identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    conditions = [BusinessRecord.module == module, *(await _record_scope_conditions(identity, db))]
    records = (await db.scalars(select(BusinessRecord).where(*conditions).order_by(BusinessRecord.created_at))).all()
    allowed_fields = await _allowed_field_keys(identity, db)
    output = io.StringIO(); writer = csv.writer(output); writer.writerow(["业务编号", "标题", "客户/主体", "状态", "负责人", "部门", "说明", "扩展数据", "创建时间", "更新时间"])
    for item in records:
        visible = _record_dict(item, allowed_fields)
        writer.writerow([visible["serial_no"], visible["title"], visible["customer"], visible["status"], visible["owner"], visible["department"], visible["description"], json.dumps(visible.get("data") or {}, ensure_ascii=False), visible["created_at"], visible["updated_at"]])
    content = "\ufeff" + output.getvalue()
    return Response(content=content.encode("utf-8"), media_type="text/csv; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="{module}-{date.today()}.csv"'})


def _export_ids(value: str) -> list[int]:
    if not value.strip():
        return []
    try:
        return list(dict.fromkeys(int(item) for item in value.split(",") if item.strip()))
    except ValueError as exc:
        raise HTTPException(status_code=422, detail="导出记录编号格式错误") from exc


async def _scoped_export_records(module: str, ids: str, identity: dict, db: AsyncSession) -> list[BusinessRecord]:
    conditions = [BusinessRecord.module == module, *(await _record_scope_conditions(identity, db))]
    selected_ids = _export_ids(ids)
    if selected_ids:
        conditions.append(BusinessRecord.id.in_(selected_ids))
    return list((await db.scalars(select(BusinessRecord).where(*conditions).order_by(BusinessRecord.created_at, BusinessRecord.id))).all())


def _csv_response(filename: str, headers: list[str], rows: list[list[object]]) -> Response:
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(headers)
    writer.writerows(rows)
    content = "\ufeff" + output.getvalue()
    disposition = f"attachment; filename=export.csv; filename*=UTF-8''{quote(filename)}"
    return Response(content=content.encode("utf-8"), media_type="text/csv; charset=utf-8", headers={"Content-Disposition": disposition})


@app.get(f"{settings.api_prefix}/investigations/clues/export")
async def export_investigation_clues(ids: str = "", identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    records = await _scoped_export_records("clue", ids, identity, db)
    rows = []
    for item in records:
        data = item.data or {}
        rows.append([item.serial_no, item.title, item.customer, item.status, item.owner, data.get("platform", ""), data.get("product", ""), data.get("region") or data.get("address", ""), data.get("collected_at", ""), data.get("notary_institution", ""), data.get("converted_case_no", ""), item.description])
    return _csv_response(f"调查线索-{date.today()}.csv", ["线索编号", "店铺/事项", "权利人", "状态", "调查员", "调查平台", "侵权产品", "调查区域/地址", "取证日期", "公证机构", "关联案号", "说明"], rows)


@app.get(f"{settings.api_prefix}/investigations/clues/handover-export")
async def export_investigation_handover(ids: str = "", identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    records = await _scoped_export_records("clue", ids, identity, db)
    rows = []
    for item in records:
        data = item.data or {}
        rows.append([item.serial_no, data.get("converted_case_no", ""), item.customer, item.title, data.get("product", ""), item.owner, data.get("collected_at", ""), data.get("certificate_no", ""), data.get("warehouse", ""), data.get("handoff_recipient", ""), item.status, item.description])
    return _csv_response(f"调查线索交接清单-{date.today()}.csv", ["线索编号", "案件编号", "权利人", "店铺/事项", "侵权产品", "调查员", "取证日期", "公证书号", "仓库位置", "交接接收人", "当前状态", "交接说明"], rows)


@app.get(f"{settings.api_prefix}/documents/official/export")
async def export_official_documents(ids: str = "", identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    records = await _scoped_export_records("document", ids, identity, db)
    records = [item for item in records if (item.data or {}).get("direction", "收文") == "收文"]
    rows = []
    for item in records:
        data = item.data or {}
        rows.append([data.get("case_no") or item.serial_no, data.get("plaintiff") or item.customer, data.get("defendant") or data.get("sender", ""), item.title, data.get("document_date") or data.get("received_at", ""), data.get("uploaded_at") or data.get("registered_at", ""), data.get("uploader") or item.owner, data.get("import_status", "已导入"), data.get("business_process_status", "未处理"), item.status])
    return _csv_response(f"官文收文-{date.today()}.csv", ["案号", "原告", "被告", "文件名称", "文件日期", "上传日期", "上传人", "导入状态", "业务处理状态", "办理状态"], rows)


@app.get(f"{settings.api_prefix}/tasks/print-export")
async def export_task_print_table(ids: str = "", identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    records = await _scoped_export_records("task", ids, identity, db)
    rows = []
    for item in records:
        data = item.data or {}
        rows.append([item.serial_no, data.get("case_no", ""), item.title, item.description, item.customer, data.get("plaintiff", ""), data.get("defendant", ""), item.status, data.get("priority", ""), data.get("initiator", ""), item.owner, data.get("deadline", ""), data.get("source", "")])
    return _csv_response(f"案件任务打印表-{date.today()}.csv", ["任务编号", "案件编号", "任务标题", "任务内容", "客户", "原告", "被告", "状态", "优先级", "发起人", "负责人", "截止日期", "任务来源"], rows)


@app.post(f"{settings.api_prefix}/customers/import")
async def import_customers(file: UploadFile = File(...), identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if not (file.filename or "").lower().endswith(".csv"): raise HTTPException(status_code=422, detail="仅支持 UTF-8 CSV 文件")
    raw = await file.read()
    if len(raw) > 5 * 1024 * 1024: raise HTTPException(status_code=413, detail="导入文件不能超过 5MB")
    try: text = raw.decode("utf-8-sig")
    except UnicodeDecodeError as exc: raise HTTPException(status_code=422, detail="CSV 必须使用 UTF-8 编码") from exc
    reader = csv.DictReader(io.StringIO(text)); existing = set((await db.scalars(select(BusinessRecord.title).where(BusinessRecord.module == "customer"))).all()); seen: set[str] = set(); created_items: list[dict] = []; errors: list[dict] = []
    current_user = await db.scalar(select(User).where(User.username == identity["username"], User.is_active.is_(True)))
    if not current_user: raise HTTPException(status_code=401, detail="当前用户不存在或已停用")
    for row_no, row in enumerate(reader, 2):
        title = (row.get("客户名称") or row.get("title") or "").strip()
        if not title: errors.append({"row": row_no, "error": "客户名称为空"}); continue
        if title in existing or title in seen: errors.append({"row": row_no, "error": "客户名称已存在", "value": title}); continue
        seen.add(title); serial = f"KH{datetime.now().strftime('%Y%m%d%H%M%S')}{uuid4().hex[:4].upper()}"
        requested_owner = (row.get("负责人") or row.get("owner") or identity["username"]).strip()
        department = (row.get("部门") or row.get("department") or current_user.department).strip()
        try:
            if identity.get("role") == "user":
                owner = current_user.username; department = current_user.department
            else:
                owner = (await _resolve_active_customer_managers([requested_owner], db))[0]
                owner_user = await db.scalar(select(User).where(User.username == owner))
                if identity.get("role") != "admin":
                    department = current_user.department
                    if not owner_user or owner_user.department != current_user.department:
                        raise HTTPException(status_code=422, detail="部门负责人只能导入本部门负责人名下的客户")
            item = BusinessRecord(module="customer", serial_no=serial, title=title, customer=title, status="跟进中", owner=owner, department=department, data={"contact": (row.get("联系人") or row.get("contact") or "").strip(), "phone": (row.get("电话") or row.get("phone") or "").strip(), "level": (row.get("客户等级") or row.get("level") or "普通客户").strip(), "customer_managers": [owner], "imported_at": datetime.now().isoformat(timespec="seconds"), "last_modified_by": identity["username"]})
            db.add(item); await db.flush(); db.add(WorkflowEvent(record_id=item.id, action="批量导入", to_status="跟进中", operator=identity["username"], comment=f"CSV 第 {row_no} 行")); created_items.append({"id": item.id, "serial_no": item.serial_no, "title": item.title, "owner": owner, "department": department})
        except HTTPException as exc:
            errors.append({"row": row_no, "error": str(exc.detail), "value": requested_owner})
    await db.commit()
    return {"created": len(created_items), "failed": len(errors), "items": created_items, "errors": errors}


RECORD_IMPORT_COLUMNS = {
    "contract": ["业务编号", "合同名称", "客户/主体", "合同类型", "合同金额", "签订日期", "外部合同号", "负责人", "部门", "说明"],
    "case": ["业务编号", "案件名称", "关联合同号", "案件类型", "对方当事人", "法院", "负责人", "说明"],
    "task": ["业务编号", "任务内容", "客户/主体", "截止日期", "优先级", "来源", "关联案号", "负责人", "协作人", "部门", "说明"],
    "document": ["业务编号", "文件名称", "客户/主体", "收发类型", "文件日期", "关联案号", "来文/送达单位", "负责人", "部门", "说明"],
    "finance": ["业务编号", "费用名称", "客户/主体", "费用类型", "金额", "关联案号", "经办人", "部门", "说明"],
    "hr": ["员工编号", "姓名", "状态", "部门", "岗位", "联系电话", "入职日期", "用工类型", "证件号码", "邮箱", "说明"],
    "warehouse": ["物品编号", "物品名称", "物品类别", "数量", "单位", "存放位置", "部门", "供应商", "说明"],
    "seal": ["申请编号", "申请标题", "客户/主体", "印章编号", "份数", "用途", "计划日期", "办理方式", "文件名称", "负责人", "部门", "说明"],
}
RECORD_IMPORT_SAMPLES = {
    "contract": ["HT2026079999", "知识产权专项服务合同", "示例客户有限公司", "专项服务", "100000.00", str(date.today()), "EXT-001", "admin", "上海分所", "示例数据请删除"],
    "case": ["SHMS2699999", "示例商标侵权纠纷", "HT2026060097", "民事案件", "示例对方有限公司", "上海知识产权法院", "admin", "合同必须已经审批通过"],
    "task": ["RW2026079999", "准备案件证据目录", "示例客户有限公司", str(date.today() + timedelta(days=7)), "普通", "日常任务", "", "admin", "协作人甲、协作人乙", "上海分所", "示例数据请删除"],
    "document": ["SW2026079999", "法院开庭传票", "示例客户有限公司", "收文", str(date.today()), "SH191000382B", "上海市示例人民法院", "admin", "上海分所", "示例数据请删除"],
    "finance": ["FY2026079999", "示例诉讼费", "示例客户有限公司", "官方费用", "3500.00", "SH191000382B", "admin", "上海分所", "示例数据请删除"],
    "hr": ["YG2026079999", "示例员工", "试用", "上海分所", "律师助理", "13800000000", str(date.today()), "全职", "", "demo@example.com", "示例数据请删除"],
    "warehouse": ["WP2026079999", "示例笔记本电脑", "电子设备", "1", "台", "行政仓 A-01", "上海分所", "示例供应商", "示例数据请删除"],
    "seal": ["YY2026079999", "示例用印申请", "示例客户有限公司", "YZ-GZ-001", "2", "法院立案", str(date.today() + timedelta(days=1)), "现场用印", "民事起诉状", "admin", "上海分所", "示例数据请删除"],
}


@app.get(f"{settings.api_prefix}/records/import-template")
async def records_import_template(module: str = Query(min_length=1, max_length=32), _: dict = Depends(current_identity)):
    if module == "case": raise HTTPException(status_code=409, detail="案件必须使用分阶段专用入口创建，不能使用通用导入模板")
    columns = RECORD_IMPORT_COLUMNS.get(module)
    if not columns: raise HTTPException(status_code=422, detail="该业务模块不支持批量导入")
    output = io.StringIO(); writer = csv.writer(output); writer.writerow(columns); writer.writerow(RECORD_IMPORT_SAMPLES[module])
    return Response(content=("\ufeff" + output.getvalue()).encode("utf-8"), media_type="text/csv; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="{module}-import-template.csv"'})


def _csv_value(row: dict, *names: str, default: str = "") -> str:
    for name in names:
        value = row.get(name)
        if value is not None and str(value).strip(): return str(value).strip()
    return default


def _csv_date(value: str, label: str, *, required: bool = True) -> str:
    if not value and not required: return ""
    try: return str(date.fromisoformat(value))
    except ValueError as exc: raise ValueError(f"{label}必须为 YYYY-MM-DD") from exc


@app.post(f"{settings.api_prefix}/records/import")
async def import_business_records(module: str = Query(min_length=1, max_length=32), file: UploadFile = File(...), identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if module not in RECORD_IMPORT_COLUMNS: raise HTTPException(status_code=422, detail="该业务模块不支持批量导入")
    if module == "case": raise HTTPException(status_code=409, detail="案件必须使用分阶段专用入口创建，不能通过通用导入绕过")
    if module in {"hr", "warehouse"} and identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="当前角色不能批量导入该模块")
    if not (file.filename or "").lower().endswith(".csv"): raise HTTPException(status_code=422, detail="仅支持 UTF-8 CSV 文件")
    raw = await file.read()
    if len(raw) > 5 * 1024 * 1024: raise HTTPException(status_code=413, detail="导入文件不能超过 5MB")
    try: csv_text = raw.decode("utf-8-sig")
    except UnicodeDecodeError as exc: raise HTTPException(status_code=422, detail="CSV 必须使用 UTF-8 编码") from exc
    reader = csv.DictReader(io.StringIO(csv_text))
    if not reader.fieldnames: raise HTTPException(status_code=422, detail="CSV 缺少表头")
    existing = set((await db.scalars(select(BusinessRecord.serial_no))).all()); seen: set[str] = set(); errors: list[dict] = []; created_items: list[dict] = []
    scope = await _record_scope_conditions(identity, db)
    contracts = {item.serial_no: item for item in (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "contract", *scope))).all()}
    case_numbers = set((await db.scalars(select(BusinessRecord.serial_no).where(BusinessRecord.module == "case", *scope))).all())
    seal_assets = {item.code: item for item in (await db.scalars(select(SealAsset).where(SealAsset.status == "可用"))).all()}
    user = await db.scalar(select(User).where(User.username == identity["username"]))
    for row_no, row in enumerate(reader, 2):
        try:
            serial = _csv_value(row, "业务编号", "员工编号", "物品编号", "申请编号", "serial_no")
            title = _csv_value(row, "标题", "合同名称", "案件名称", "任务内容", "文件名称", "费用名称", "姓名", "物品名称", "申请标题", "title")
            if not serial or not title: raise ValueError("业务编号和名称不能为空")
            if serial in existing or serial in seen: raise ValueError(f"业务编号已存在：{serial}")
            owner = _csv_value(row, "负责人", "经办人", "owner", default=identity["username"])
            department = _csv_value(row, "部门", "department", default=(user.department if user else "上海分所"))
            if identity.get("role") != "admin":
                department = user.department if user else department
                if identity.get("role") == "user": owner = identity["username"]
            customer = _csv_value(row, "客户/主体", "customer")
            description = _csv_value(row, "说明", "description")
            data: dict = {"imported_at": datetime.now().isoformat(timespec="seconds"), "import_row": row_no}
            status_value = "草稿"
            if module == "contract":
                amount = float(_csv_value(row, "合同金额", "amount")); signed_at = _csv_date(_csv_value(row, "签订日期", "signed_at"), "签订日期")
                if amount < 0 or not customer: raise ValueError("客户不能为空，合同金额不能为负数")
                data.update({"type": _csv_value(row, "合同类型", "type", default="专项服务"), "amount": f"{amount:.2f}", "signed_at": signed_at, "external_contract_no": _csv_value(row, "外部合同号", "external_contract_no")})
            elif module == "case":
                contract_no = _csv_value(row, "关联合同号", "contract_no"); contract = contracts.get(contract_no)
                if not contract: raise ValueError("关联合同不存在或无权访问")
                if contract.status in {"草稿", "审批中", "已拒绝", "已撤回", "已作废"}: raise ValueError("关联合同尚未审批通过")
                customer = contract.customer; department = contract.department; status_value = "新案待分配"; contract_data = contract.data or {}
                data.update({"contract_id": contract.id, "contract_no": contract.serial_no, "external_contract_no": contract_data.get("external_contract_no", ""), "contract_title": contract.title, "case_type": _csv_value(row, "案件类型", "case_type", default="民事案件"), "opponent": _csv_value(row, "对方当事人", "opponent"), "court": _csv_value(row, "法院", "court")})
            elif module == "task":
                deadline = _csv_date(_csv_value(row, "截止日期", "deadline"), "截止日期")
                priority = _csv_value(row, "优先级", "priority", default="普通")
                if priority not in {"普通", "重要", "紧急"}: raise ValueError("任务优先级无效")
                try:
                    _validate_task_deadline(date.fromisoformat(deadline))
                    owner = await _active_task_username(owner, db, field_name="负责人")
                    collaborators = []
                    for value in [name.strip() for name in re.split(r"[、,，;；]", _csv_value(row, "协作人", "collaborators")) if name.strip()]:
                        collaborator = await _active_task_username(value, db, field_name="协作人")
                        if collaborator != owner and collaborator not in collaborators:
                            collaborators.append(collaborator)
                except HTTPException as exc:
                    raise ValueError(str(exc.detail)) from exc
                status_value = "待接收"; data.update({"deadline": deadline, "priority": priority, "source": _csv_value(row, "来源", "source", default="日常任务"), "case_no": _csv_value(row, "关联案号", "case_no"), "initiator": identity["username"], "collaborators": collaborators})
            elif module == "document":
                direction = _csv_value(row, "收发类型", "direction"); case_no = _csv_value(row, "关联案号", "case_no")
                if direction not in {"收文", "发文"}: raise ValueError("收发类型必须为收文或发文")
                if case_no and case_no not in case_numbers: raise ValueError("关联案件不存在或无权访问")
                status_value = "待登记"; data.update({"direction": direction, "document_date": _csv_date(_csv_value(row, "文件日期", "document_date"), "文件日期"), "case_no": case_no, "sender": _csv_value(row, "来文/送达单位", "sender")})
            elif module == "finance":
                amount = float(_csv_value(row, "金额", "amount")); case_no = _csv_value(row, "关联案号", "case_no")
                if amount <= 0: raise ValueError("费用金额必须大于 0")
                if case_no and case_no not in case_numbers: raise ValueError("关联案件不存在或无权访问")
                data.update({"fee_type": _csv_value(row, "费用类型", "fee_type", default="官方费用"), "amount": f"{amount:.2f}", "case_no": case_no, "handler": _csv_value(row, "经办人", "handler", default=owner)})
            elif module == "hr":
                position = _csv_value(row, "岗位", "position"); joined_at = _csv_date(_csv_value(row, "入职日期", "joined_at"), "入职日期")
                if not position: raise ValueError("岗位不能为空")
                requested_status = _csv_value(row, "状态", "status", default="试用"); status_value = "在职" if requested_status == "在职" else "试用"
                data.update({"position": position, "phone": _csv_value(row, "联系电话", "phone"), "joined_at": joined_at, "employment_type": _csv_value(row, "用工类型", "employment_type", default="全职"), "id_no": _csv_value(row, "证件号码", "id_no"), "email": _csv_value(row, "邮箱", "email")})
            elif module == "warehouse":
                quantity = int(_csv_value(row, "数量", "quantity")); category = _csv_value(row, "物品类别", "category"); location = _csv_value(row, "存放位置", "location")
                if quantity < 1 or not category or not location: raise ValueError("物品类别、数量和存放位置不能为空")
                customer = _csv_value(row, "供应商", "vendor"); status_value = "在库"; data.update({"category": category, "quantity": quantity, "unit": _csv_value(row, "单位", "unit", default="件"), "location": location, "vendor": customer, "borrower": "", "due_date": "", "borrow_purpose": ""})
            elif module == "seal":
                asset_code = _csv_value(row, "印章编号", "seal_code"); asset = seal_assets.get(asset_code); copies = int(_csv_value(row, "份数", "copies"))
                if not asset: raise ValueError("印章编号不存在或印章不可用")
                if copies < 1: raise ValueError("用印份数必须大于 0")
                data.update({"seal_asset_id": asset.id, "seal_name": asset.name, "seal_type": asset.seal_type, "copies": copies, "purpose": _csv_value(row, "用途", "purpose"), "use_date": _csv_date(_csv_value(row, "计划日期", "use_date"), "计划日期"), "delivery_method": _csv_value(row, "办理方式", "delivery_method", default="现场用印"), "document_names": _csv_value(row, "文件名称", "document_names")})
            item = BusinessRecord(module=module, serial_no=serial, title=title, customer=customer, status=status_value, owner=owner, department=department, description=description, data=data)
            db.add(item); await db.flush(); db.add(WorkflowEvent(record_id=item.id, action="批量导入", to_status=item.status, operator=identity["username"], comment=f"CSV 第 {row_no} 行"))
            seen.add(serial); created_items.append({"id": item.id, "serial_no": serial, "title": title})
        except (ValueError, TypeError) as exc:
            errors.append({"row": row_no, "error": str(exc) or "字段格式错误", "value": _csv_value(row, "业务编号", "员工编号", "物品编号", "申请编号", "serial_no")})
    await db.commit()
    return {"module": module, "created": len(created_items), "failed": len(errors), "items": created_items, "errors": errors}


@app.get(f"{settings.api_prefix}/records")
async def list_records(
    module: str = Query(min_length=1, max_length=32),
    keyword: str = "", record_status: str = "", page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    if module in {"notary", "case"}:
        await _apply_notary_auto_conversion(db)
    conditions = [BusinessRecord.module == module]
    conditions.extend(await _record_scope_conditions(identity, db))
    if keyword:
        like = f"%{keyword}%"
        conditions.append(or_(BusinessRecord.serial_no.ilike(like), BusinessRecord.title.ilike(like), BusinessRecord.customer.ilike(like), BusinessRecord.owner.ilike(like)))
    if record_status:
        conditions.append(BusinessRecord.status == record_status)
    total = await db.scalar(select(func.count()).select_from(BusinessRecord).where(*conditions))
    result = await db.scalars(select(BusinessRecord).where(*conditions).order_by(BusinessRecord.updated_at.desc()).offset((page - 1) * page_size).limit(page_size))
    allowed_fields = await _allowed_field_keys(identity, db)
    return {"items": [_record_dict(item, allowed_fields) for item in result], "total": total or 0, "page": page, "page_size": page_size}


CUSTOMER_CREATE_DATA_FIELDS = {
    "contact", "phone", "credit_code", "legal_representative", "registered_address",
    "invoice_title", "taxpayer_id", "invoice_address", "invoice_phone", "bank_name", "bank_account",
    "customer_type", "short_name", "fax", "legal_agent_id_no", "legal_agent_title", "customer_source",
    "is_shared", "is_assisted", "file_date", "level", "province", "postal_code", "patent_customer_type",
    "fee_reduction", "industry", "output_value", "cooperation_status", "gb_classification", "website",
    "organization_nature", "organization_code", "registration_region", "registration_postal_code",
    "registered_capital", "registration_year",
}
CUSTOMER_SYSTEM_DATA_FIELDS = {
    "notes", "last_contact_at", "contact_count", "last_modified_by", "last_modified_date",
    "status_before_recycle", "recycled_at", "recycled_by", "restored_at", "restored_by",
    "released_at", "released_by", "claimed_at", "claimed_by", "shared_with", "is_shared", "shared_at",
    "level_change", "key_change", "portal_access",
}
CUSTOMER_CREATE_STATUSES = {"潜在", "目标", "立项", "关怀", "签约", "谈判", "价值"}
CUSTOMER_LEVELS = {"潜在客户", "目标客户", "签约客户", "立案客户", "高级客户", "中级客户", "低级客户"}


def _normalize_customer_yes_no(value: object, field_name: str) -> str:
    if isinstance(value, bool):
        return "是" if value else "否"
    normalized = str(value or "").strip().casefold()
    if normalized in {"是", "yes", "true", "1"}: return "是"
    if normalized in {"否", "no", "false", "0", ""}: return "否"
    raise HTTPException(status_code=422, detail=f"{field_name}只能填写是或否")


def _normalize_customer_name(value: object) -> str:
    return unicodedata.normalize("NFKC", str(value or "").strip()).casefold()


async def _ensure_unique_customer_name(title: str, db: AsyncSession, *, exclude_id: int | None = None) -> None:
    normalized = _normalize_customer_name(title)
    if not normalized:
        raise HTTPException(status_code=422, detail="客户名称不能为空")
    rows = (await db.execute(
        select(BusinessRecord.id, BusinessRecord.title).where(BusinessRecord.module == "customer")
    )).all()
    if any(item_id != exclude_id and _normalize_customer_name(existing_title) == normalized for item_id, existing_title in rows):
        raise HTTPException(status_code=409, detail="客户名称已存在，不能创建或改为同名客户")


@app.post(f"{settings.api_prefix}/customers", status_code=status.HTTP_201_CREATED)
async def create_customer(body: CustomerCreateInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    title = body.title.strip()
    await _ensure_unique_customer_name(title, db)
    customer_status = body.status.strip()
    if customer_status and customer_status not in CUSTOMER_CREATE_STATUSES:
        raise HTTPException(status_code=422, detail="客户状态无效")
    data = dict(body.data or {})
    for field_name in CUSTOMER_CREATE_DATA_FIELDS:
        if field_name in body.model_fields_set:
            data[field_name] = getattr(body, field_name)
    # New customers have no lifecycle or sharing history.  Drop every
    # client-supplied system field after merging both ``data`` and top-level
    # aliases so neither input shape can forge audit state or list membership.
    for protected_customer_field in CUSTOMER_SYSTEM_DATA_FIELDS:
        data.pop(protected_customer_field, None)
    data["shared_with"] = []
    data["is_shared"] = "否"
    customer_type = str(data.get("customer_type") or "客户").strip()
    if customer_type not in {"客户", "当事人"}: raise HTTPException(status_code=422, detail="客户类型只能为客户或当事人")
    level = str(data.get("level") or "立案客户").strip()
    if level not in CUSTOMER_LEVELS: raise HTTPException(status_code=422, detail="客户等级无效")
    data["customer_type"] = customer_type; data["level"] = level
    for key, label in {"is_shared": "是否共享", "is_assisted": "上海市资助信息", "fee_reduction": "是否费减"}.items():
        data[key] = _normalize_customer_yes_no(data.get(key), label)
    credit_code = str(data.get("credit_code") or "").strip()
    if credit_code and any(character.isspace() for character in str(data.get("credit_code") or "")):
        raise HTTPException(status_code=422, detail="统一社会信用代码不允许包含空格")
    if credit_code:
        existing_customers = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "customer"))).all()
        if any(str((item.data or {}).get("credit_code") or "").strip().casefold() == credit_code.casefold() for item in existing_customers):
            raise HTTPException(status_code=409, detail="统一社会信用代码已存在")
    data["credit_code"] = credit_code
    current_user = await db.scalar(select(User).where(User.username == identity["username"], User.is_active.is_(True)))
    if not current_user: raise HTTPException(status_code=401, detail="当前用户不存在或已停用")
    requested_owner = body.owner.strip() or current_user.username
    department = body.department.strip() or current_user.department
    if identity.get("role") != "admin":
        department = current_user.department
        if identity.get("role") == "user": requested_owner = current_user.username
    owner = (await _resolve_active_customer_managers([requested_owner], db))[0]
    raw_managers = body.customer_managers if "customer_managers" in body.model_fields_set else data.get("customer_managers", [])
    managers = await _resolve_active_customer_managers(list(raw_managers or [owner]), db)
    data["customer_managers"] = [owner, *[manager for manager in managers if manager != owner]]
    serial_no = body.serial_no.strip()
    if not serial_no:
        serial_no = f"KH{datetime.now():%Y%m%d%H%M%S%f}{uuid4().hex[:4].upper()}"
    if await db.scalar(select(BusinessRecord.id).where(BusinessRecord.serial_no == serial_no)):
        raise HTTPException(status_code=409, detail="业务编号已存在")
    record = BusinessRecord(
        module="customer", serial_no=serial_no, title=title, customer=title, status=customer_status,
        owner=owner, department=department, description=body.description.strip(), data=data,
    )
    _mark_customer_modified(record, identity)
    db.add(record); await db.flush()
    db.add(WorkflowEvent(record_id=record.id, action="创建客户", to_status=record.status, operator=identity["username"], comment="通过新建客户专用入口创建"))
    await db.commit(); await db.refresh(record)
    return await _record_dict_for_identity(record, identity, db)


CASE_PLAINTIFF_FIELDS = (
    "plaintiff", "plaintiffs", "plaintiff_name", "plaintiff_names",
    "appellant", "appellants", "appellant_name", "appellant_names",
)
CASE_DEFENDANT_FIELDS = (
    "defendant", "defendants", "defendant_name", "defendant_names", "opponent",
    "appellee", "appellees", "appellee_name", "appellee_names",
)
CASE_THIRD_PARTY_FIELDS = (
    "third_party", "third_parties", "third_party_name", "third_party_names",
)
# Legacy Chinese multi-party strings use adjacent commas.  A comma followed by
# whitespace is retained as part of an English legal name (for example
# ``Foo, Inc.``) rather than treated as a party boundary.
CASE_PARTY_SEPARATOR = re.compile(r",(?!\s)|[，、;；\r\n]+")


def _normalize_conflict_entity(value: object) -> str:
    """Normalize typography without weakening whole-entity equality."""
    return " ".join(unicodedata.normalize("NFKC", str(value or "")).split()).casefold()


def _conflict_entity_tokens(value: object) -> list[str]:
    """Expand stored party arrays and legacy comma-delimited party fields."""
    if isinstance(value, (list, tuple, set)):
        # An explicit array already defines entity boundaries.  Do not split an
        # English legal name such as ``Foo, Inc.`` inside one array element.
        return list(dict.fromkeys(str(item or "").strip() for item in value if str(item or "").strip()))
    tokens: list[str] = []
    for token in CASE_PARTY_SEPARATOR.split(str(value or "")):
        token = token.strip()
        if token and token not in tokens:
            tokens.append(token)
    return tokens


def _case_party_values(data: dict, fields: tuple[str, ...]) -> list[str]:
    """Use the first populated historical field family, preserving stored order."""
    for field in fields:
        tokens = _conflict_entity_tokens(data.get(field))
        if tokens:
            return tokens
    return []


def _case_party_match_values(data: dict, fields: tuple[str, ...]) -> list[str]:
    """Match only the same highest-priority field that the result will show."""
    for field in fields:
        raw_value = data.get(field)
        displayed = _conflict_entity_tokens(raw_value)
        if not displayed:
            continue
        if isinstance(raw_value, (list, tuple, set)):
            return displayed
        full_value = str(raw_value or "").strip()
        # Keep the complete scalar as an exact candidate so an English legal
        # name containing a comma remains searchable, while retaining legacy
        # comma-delimited entities as additional exact candidates.
        return list(dict.fromkeys([full_value, *displayed]))
    return []


def _case_conflict_entities(record: BusinessRecord) -> list[str]:
    data = record.data or {}
    values = [record.customer.strip()] if record.customer.strip() else []
    for fields in (CASE_PLAINTIFF_FIELDS, CASE_DEFENDANT_FIELDS, CASE_THIRD_PARTY_FIELDS):
        for value in _case_party_match_values(data, fields):
            if value not in values:
                values.append(value)
    return values


def _case_filing_date(record: BusinessRecord) -> date | None:
    raw_value = str((record.data or {}).get("filing_date") or "").strip()
    if not raw_value:
        return None
    try:
        return date.fromisoformat(raw_value[:10])
    except ValueError:
        return None


async def _require_customer_conflict_permission(identity: dict, db: AsyncSession) -> None:
    if identity.get("role") == "admin":
        return
    permission = await _permission_payload(identity.get("role", "user"), db)
    if "customer-conflict" not in set(permission.get("menu_keys", [])):
        raise HTTPException(status_code=403, detail="当前角色没有客户利益检索权限")


async def _conflict_customer_managers(entity_name: str, db: AsyncSession) -> list[str]:
    normalized_name = _normalize_conflict_entity(entity_name)
    customers = list((await db.scalars(
        select(BusinessRecord).where(BusinessRecord.module == "customer")
    )).all())
    matching_customers = [
        customer for customer in customers
        if normalized_name in {
            _normalize_conflict_entity(customer.title),
            _normalize_conflict_entity(customer.customer),
        }
    ]
    if not matching_customers:
        return []
    # A current customer record wins over a recycled historical copy; within
    # the same lifecycle class the most recently updated row, then id, wins.
    matching_customers.sort(
        key=lambda customer: (
            customer.status != "已回收",
            customer.updated_at or customer.created_at,
            customer.id,
        ),
        reverse=True,
    )
    customer = matching_customers[0]
    manager_tokens = _conflict_entity_tokens((customer.data or {}).get("customer_managers"))
    if not manager_tokens and customer.owner:
        manager_tokens = [customer.owner]
    if not manager_tokens:
        return []
    users = list((await db.scalars(select(User).where(User.is_active.is_(True), User.username.in_(manager_tokens)))).all())
    display_names = {user.username: user.display_name or user.username for user in users}
    return [display_names.get(manager, manager) for manager in manager_tokens]


def _empty_customer_conflict_result(query: str) -> dict:
    return {
        "found": False, "query": query, "enterprise_name": "",
        "latest_case_no": "", "latest_case_date": "",
        "plaintiffs": [], "defendants": [], "third_parties": [],
        "our_customer": "", "customer_managers": [],
    }


@app.get(f"{settings.api_prefix}/customers/conflicts")
async def customer_conflicts(
    name: str = Query(min_length=1, max_length=100),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    """Return the latest case containing one exact, normalized enterprise party."""
    await _require_customer_conflict_permission(identity, db)
    query = name.strip()
    if not query:
        raise HTTPException(status_code=422, detail="企业名称不能为空")
    needle = _normalize_conflict_entity(query)
    if not needle:
        raise HTTPException(status_code=422, detail="企业名称不能为空")

    # Conflict checking deliberately spans the whole firm.  Only the minimum
    # original-page disclosure is returned and no source record id is exposed.
    cases = list((await db.scalars(
        select(BusinessRecord).where(BusinessRecord.module == "case")
    )).all())
    matching_cases = [
        case_record for case_record in cases
        if any(_normalize_conflict_entity(entity) == needle for entity in _case_conflict_entities(case_record))
    ]
    if not matching_cases:
        return _empty_customer_conflict_result(query)
    latest_case = max(
        matching_cases,
        key=lambda case_record: (_case_filing_date(case_record) or date.min, case_record.id),
    )
    data = latest_case.data or {}
    filing_date = _case_filing_date(latest_case)
    plaintiffs = _case_party_values(data, CASE_PLAINTIFF_FIELDS) or _conflict_entity_tokens(latest_case.customer)
    return {
        "found": True,
        "query": query,
        "enterprise_name": query,
        "latest_case_no": latest_case.serial_no,
        "latest_case_date": filing_date.isoformat() if filing_date else "",
        "plaintiffs": plaintiffs,
        "defendants": _case_party_values(data, CASE_DEFENDANT_FIELDS),
        "third_parties": _case_party_values(data, CASE_THIRD_PARTY_FIELDS),
        "our_customer": latest_case.customer,
        "customer_managers": await _conflict_customer_managers(query, db),
    }


async def _customer_or_404(customer_id: int, identity: dict, db: AsyncSession) -> BusinessRecord:
    try:
        return await _ensure_record_module(customer_id, "customer", identity, db)
    except HTTPException as exc:
        if exc.status_code == 404:
            raise HTTPException(status_code=404, detail="客户不存在或无权访问") from exc
        raise


async def _locked_customer_or_404(customer_id: int, identity: dict, db: AsyncSession) -> BusinessRecord:
    """Lock a visible customer and refresh it after waiting for a concurrent writer."""
    conditions = [
        BusinessRecord.id == customer_id,
        BusinessRecord.module == "customer",
        *(await _record_scope_conditions(identity, db)),
    ]
    customer = await db.scalar(
        select(BusinessRecord)
        .where(*conditions)
        .with_for_update()
        .execution_options(populate_existing=True)
    )
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在或无权访问")
    return customer


def _customer_event(customer: BusinessRecord, action: str, identity: dict, comment: str, from_status: str | None = None) -> WorkflowEvent:
    _mark_customer_modified(customer, identity)
    return WorkflowEvent(record_id=customer.id, action=action, from_status=from_status or customer.status, to_status=customer.status, operator=identity["username"], comment=comment)


@app.post(f"{settings.api_prefix}/customers/{{customer_id}}/claim")
async def claim_customer(customer_id: int, body: CustomerActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager", "user"}:
        raise HTTPException(status_code=403, detail="当前角色不能领取公海客户")
    customer = await _locked_customer_or_404(customer_id, identity, db)
    if customer.status != "公海": raise HTTPException(status_code=409, detail="只有公海客户可以领取")
    current_user = await db.scalar(select(User).where(User.username == identity["username"], User.is_active.is_(True)))
    if not current_user: raise HTTPException(status_code=401, detail="当前用户不存在或已停用")
    old = customer.status; customer.status = "潜在"; customer.owner = current_user.username; customer.department = current_user.department
    customer.data = {
        **(customer.data or {}),
        "customer_managers": [identity["username"]],
        "shared_with": [],
        "is_shared": "否",
        "claimed_at": datetime.now().isoformat(timespec="seconds"),
        "claimed_by": identity["username"],
    }
    db.add(_customer_event(customer, "领取客户", identity, body.comment or "从公海领取客户", old)); await db.commit(); await db.refresh(customer)
    return await _record_dict_for_identity(customer, identity, db)


@app.post(f"{settings.api_prefix}/customers/{{customer_id}}/release")
async def release_customer(customer_id: int, body: CustomerActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    customer = await _locked_customer_or_404(customer_id, identity, db)
    await _require_record_owner_or_manager(customer, identity, db)
    if customer.status not in {"正常", "跟进中", "已回收", *CUSTOMER_CREATE_STATUSES}: raise HTTPException(status_code=409, detail="当前客户状态不能释放到公海")
    old = customer.status; customer.status = "公海"; customer.owner = "公海"
    customer.data = {
        **(customer.data or {}),
        "shared_with": [],
        "is_shared": "否",
        "released_at": datetime.now().isoformat(timespec="seconds"),
        "released_by": identity["username"],
    }
    db.add(_customer_event(customer, "释放公海", identity, body.comment or "客户释放到公海", old)); await db.commit(); await db.refresh(customer)
    return await _record_dict_for_identity(customer, identity, db)


@app.post(f"{settings.api_prefix}/customers/{{customer_id}}/share")
async def share_customer(customer_id: int, body: CustomerShareInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    customer = await _locked_customer_or_404(customer_id, identity, db)
    await _require_record_owner_or_manager(customer, identity, db)
    if customer.status in {"公海", "已回收"}: raise HTTPException(status_code=409, detail="公海或回收站客户不能共享")
    recipients = await _resolve_active_customer_managers(body.recipients, db)
    customer_managers = {
        str(value).strip()
        for value in (customer.data or {}).get("customer_managers", [])
        if str(value).strip()
    }
    redundant = sorted(set(recipients) & ({str(customer.owner or "").strip()} | customer_managers))
    if redundant:
        raise HTTPException(status_code=422, detail=f"客户负责人或管理人无需重复共享：{'、'.join(redundant)}")
    existing = set((customer.data or {}).get("shared_with", [])); existing.update(recipients)
    customer.data = {**(customer.data or {}), "shared_with": sorted(existing), "is_shared": "是", "shared_at": datetime.now().isoformat(timespec="seconds")}
    db.add(_customer_event(customer, "共享客户", identity, body.comment or f"共享给：{'、'.join(recipients)}")); await db.commit(); await db.refresh(customer)
    return await _record_dict_for_identity(customer, identity, db)


@app.post(f"{settings.api_prefix}/customers/{{customer_id}}/recycle")
async def recycle_customer(customer_id: int, body: CustomerActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    customer = await _locked_customer_or_404(customer_id, identity, db)
    await _require_record_owner_or_manager(customer, identity, db)
    if customer.status == "公海": raise HTTPException(status_code=409, detail="公海客户必须先领取，不能直接移入回收站")
    if customer.status == "已回收": raise HTTPException(status_code=409, detail="客户已在回收站")
    old = customer.status
    customer.data = {
        **(customer.data or {}),
        "shared_with": [],
        "is_shared": "否",
        "status_before_recycle": old,
        "recycled_at": datetime.now().isoformat(timespec="seconds"),
        "recycled_by": identity["username"],
    }
    customer.status = "已回收"; db.add(_customer_event(customer, "移入回收站", identity, body.comment or "客户移入回收站", old)); await db.commit(); await db.refresh(customer)
    return await _record_dict_for_identity(customer, identity, db)


@app.post(f"{settings.api_prefix}/customers/{{customer_id}}/restore")
async def restore_customer(customer_id: int, body: CustomerActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    customer = await _locked_customer_or_404(customer_id, identity, db)
    await _require_record_owner_or_manager(customer, identity, db)
    if customer.status != "已回收": raise HTTPException(status_code=409, detail="只有回收站客户可以恢复")
    old = customer.status; previous = str((customer.data or {}).get("status_before_recycle", "跟进中"))
    if previous not in {"正常", "跟进中", "公海", *CUSTOMER_CREATE_STATUSES}: previous = "潜在"
    customer.status = previous; customer.data = {**(customer.data or {}), "restored_at": datetime.now().isoformat(timespec="seconds"), "restored_by": identity["username"]}
    db.add(_customer_event(customer, "恢复客户", identity, body.comment or f"恢复为{previous}", old)); await db.commit(); await db.refresh(customer)
    return await _record_dict_for_identity(customer, identity, db)


@app.post(f"{settings.api_prefix}/customers/{{customer_id}}/contacts", status_code=status.HTTP_201_CREATED)
async def add_customer_contact(customer_id: int, body: CustomerContactInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    customer = await _customer_or_404(customer_id, identity, db); await _require_record_owner_or_manager(customer, identity, db); data = customer.data or {}; contacts = list(data.get("contacts", []))
    if any(x.get("name") == body.name.strip() and x.get("phone", "") == body.phone.strip() for x in contacts): raise HTTPException(status_code=409, detail="相同联系人已存在")
    if body.is_primary: contacts = [{**x, "is_primary": False} for x in contacts]
    contact = {
        "id": uuid4().hex,
        "name": body.name.strip(),
        "project_role": body.project_role.strip(),
        "position": body.position.strip(),
        "phone": body.phone.strip(),
        "office_phone": body.office_phone.strip(),
        "im_account": body.im_account.strip(),
        "email": body.email.strip(),
        "contact_status": body.contact_status.strip() or "正常联系",
        "is_valid": body.is_valid,
        "is_primary": body.is_primary,
        "remark": body.remark.strip(),
    }
    contacts.append(contact); customer.data = {**data, "contacts": contacts}
    db.add(_customer_event(customer, "新增联系人", identity, f"联系人：{contact['name']}")); await db.commit()
    return contact


@app.put(f"{settings.api_prefix}/customers/{{customer_id}}/contacts/{{contact_id}}")
async def update_customer_contact(customer_id: int, contact_id: str, body: CustomerContactInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Update an existing contact in place; customer scope and edit authority are never bypassed."""
    customer = await _customer_or_404(customer_id, identity, db)
    await _require_record_owner_or_manager(customer, identity, db)
    data = customer.data or {}
    contacts = list(data.get("contacts", []))
    index = next((i for i, item in enumerate(contacts) if item.get("id") == contact_id), None)
    if index is None:
        raise HTTPException(status_code=404, detail="联系人不存在")
    name = body.name.strip()
    phone = body.phone.strip()
    if not name:
        raise HTTPException(status_code=422, detail="请输入联系人姓名")
    if any(item.get("id") != contact_id and item.get("name") == name and item.get("phone", "") == phone for item in contacts):
        raise HTTPException(status_code=409, detail="相同联系人已存在")
    previous = contacts[index]
    updated = {
        **previous,
        "name": name,
        "project_role": body.project_role.strip(),
        "position": body.position.strip(),
        "phone": phone,
        "office_phone": body.office_phone.strip(),
        "im_account": body.im_account.strip(),
        "email": body.email.strip(),
        "contact_status": body.contact_status.strip() or "正常联系",
        "is_valid": body.is_valid,
        "is_primary": body.is_primary,
        "remark": body.remark.strip(),
    }
    if body.is_primary:
        contacts = [{**item, "is_primary": False} if item.get("id") != contact_id else updated for item in contacts]
    else:
        contacts[index] = updated
    changed_labels = {
        "name": "姓名", "position": "职务", "project_role": "项目角色", "phone": "移动电话",
        "office_phone": "办公电话", "im_account": "IM", "email": "邮箱",
        "contact_status": "联系状态", "is_valid": "有效状态", "is_primary": "主要联系人", "remark": "备注",
    }
    changed = [label for key, label in changed_labels.items() if previous.get(key) != updated.get(key)]
    customer.data = {**data, "contacts": contacts}
    if changed:
        db.add(_customer_event(customer, "修改联系人", identity, f"联系人：{previous.get('name', '')} → {updated['name']}；修改字段：{'、'.join(changed)}"))
    await db.commit()
    return updated


@app.delete(f"{settings.api_prefix}/customers/{{customer_id}}/contacts/{{contact_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_customer_contact(customer_id: int, contact_id: str, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    customer = await _customer_or_404(customer_id, identity, db); await _require_record_owner_or_manager(customer, identity, db); data = customer.data or {}; contacts = list(data.get("contacts", [])); remaining = [x for x in contacts if x.get("id") != contact_id]
    if len(remaining) == len(contacts): raise HTTPException(status_code=404, detail="联系人不存在")
    customer.data = {**data, "contacts": remaining}; db.add(_customer_event(customer, "删除联系人", identity, "删除客户联系人")); await db.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.put(f"{settings.api_prefix}/customers/{{customer_id}}/managers")
async def update_customer_managers(customer_id: int, body: CustomerManagersInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    customer = await _customer_or_404(customer_id, identity, db)
    if customer.status == "公海": raise HTTPException(status_code=409, detail="公海客户必须先领取后才能分配管理人")
    await _require_record_owner_or_manager(customer, identity, db)
    managers = await _resolve_active_customer_managers(body.managers, db)
    customer.owner = managers[0]
    customer.data = {**(customer.data or {}), "customer_managers": managers}
    db.add(_customer_event(customer, "更新客户管理人", identity, body.comment or f"客户管理人：{'、'.join(managers)}"))
    await db.commit(); await db.refresh(customer)
    return await _record_dict_for_identity(customer, identity, db)


@app.post(f"{settings.api_prefix}/customers/{{customer_id}}/notes", status_code=status.HTTP_201_CREATED)
async def add_customer_note(customer_id: int, body: CustomerNoteInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    customer = await _customer_or_404(customer_id, identity, db)
    await _require_record_owner_or_manager(customer, identity, db)
    data = customer.data or {}
    notes = list(data.get("notes", []))
    note = {"id": uuid4().hex, "type": body.note_type.strip() or "跟进记录", "content": body.content.strip(), "operator": identity["username"], "created_at": datetime.now().isoformat(timespec="seconds")}
    notes.insert(0, note)
    customer.data = {**data, "notes": notes}
    _sync_customer_contact_metrics(customer)
    db.add(_customer_event(customer, "新增客户跟进", identity, f"{note['type']}：{note['content'][:120]}"))
    await db.commit(); await db.refresh(customer)
    return note


@app.delete(f"{settings.api_prefix}/customers/{{customer_id}}/notes/{{note_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_customer_note(customer_id: int, note_id: str, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    customer = await _customer_or_404(customer_id, identity, db)
    await _require_record_owner_or_manager(customer, identity, db)
    data = customer.data or {}; notes = list(data.get("notes", [])); remaining = [note for note in notes if note.get("id") != note_id]
    if len(remaining) == len(notes):
        raise HTTPException(status_code=404, detail="跟进记录不存在")
    customer.data = {**data, "notes": remaining}
    _sync_customer_contact_metrics(customer)
    db.add(_customer_event(customer, "删除客户跟进", identity, "删除一条客户跟进记录"))
    await db.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.post(f"{settings.api_prefix}/customers/{{customer_id}}/level-change")
async def submit_customer_level_change(customer_id: int, body: CustomerLevelChangeInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    customer = await _customer_or_404(customer_id, identity, db)
    await _require_record_owner_or_manager(customer, identity, db)
    requested_level = body.level.strip()
    if requested_level not in CUSTOMER_LEVELS:
        raise HTTPException(status_code=422, detail="客户等级无效")
    data = customer.data or {}
    if requested_level == str(data.get("level") or ""):
        raise HTTPException(status_code=409, detail="客户当前已经是该等级")
    pending = data.get("level_change") or {}
    if pending.get("status") == "待审批":
        raise HTTPException(status_code=409, detail="已有客户分级调整正在审批")
    customer.data = {
        **data,
        "level_change": {
            "status": "待审批", "from_level": data.get("level", ""), "to_level": requested_level,
            "requested_by": identity["username"], "requested_at": datetime.now().isoformat(timespec="seconds"),
            "comment": body.comment.strip(),
        },
    }
    db.add(_customer_event(customer, "提交客户分级调整", identity, body.comment or f"{data.get('level', '')} → {requested_level}"))
    await db.commit(); await db.refresh(customer)
    return await _record_dict_for_identity(customer, identity, db)


@app.post(f"{settings.api_prefix}/customers/{{customer_id}}/level-change/review")
async def review_customer_level_change(customer_id: int, body: CustomerLevelReviewInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或客户主管可以审批客户分级")
    customer = await _customer_or_404(customer_id, identity, db)
    data = customer.data or {}; pending = data.get("level_change") or {}
    if pending.get("status") != "待审批":
        raise HTTPException(status_code=409, detail="该客户没有待审批的分级调整")
    if not body.approved and not body.comment.strip():
        raise HTTPException(status_code=422, detail="驳回时必须填写原因")
    previous_level = str(data.get("level") or "")
    requested_level = str(pending.get("to_level") or "")
    if requested_level not in CUSTOMER_LEVELS:
        raise HTTPException(status_code=409, detail="待审批的客户等级已失效")
    reviewed = {
        **pending, "status": "已通过" if body.approved else "已驳回",
        "reviewed_by": identity["username"], "reviewed_at": datetime.now().isoformat(timespec="seconds"),
        "review_comment": body.comment.strip(),
    }
    customer.data = {**data, "level": requested_level if body.approved else previous_level, "level_change": reviewed}
    _mark_customer_modified(customer, identity)
    action = "客户分级审批通过" if body.approved else "客户分级审批驳回"
    db.add(_customer_event(customer, action, identity, body.comment or f"{previous_level} → {requested_level}"))
    await db.commit(); await db.refresh(customer)
    return await _record_dict_for_identity(customer, identity, db)


@app.post(f"{settings.api_prefix}/customers/{{customer_id}}/key-change")
async def submit_customer_key_change(customer_id: int, body: CustomerKeyChangeInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    customer = await _customer_or_404(customer_id, identity, db)
    await _require_record_owner_or_manager(customer, identity, db)
    data = customer.data or {}; pending = data.get("key_change") or {}
    if pending.get("status") == "待审批": raise HTTPException(status_code=409, detail="已有客户关键字段变更正在审批")
    title = body.title.strip(); credit_code = body.credit_code.strip()
    await _ensure_unique_customer_name(title, db, exclude_id=customer.id)
    if title == customer.title and credit_code == str(data.get("credit_code") or ""):
        raise HTTPException(status_code=422, detail="客户名称和统一社会信用代码均未变化")
    if credit_code and any(character.isspace() for character in credit_code): raise HTTPException(status_code=422, detail="统一社会信用代码不允许包含空格")
    request_data = {"status": "待审批", "before": {"title": customer.title, "credit_code": data.get("credit_code", "")}, "after": {"title": title, "credit_code": credit_code}, "requested_by": identity["username"], "requested_at": datetime.now().isoformat(timespec="seconds"), "comment": body.comment.strip()}
    customer.data = {**data, "key_change": request_data}
    db.add(_customer_event(customer, "提交客户关键字段变更", identity, body.comment))
    await db.commit(); await db.refresh(customer)
    return await _record_dict_for_identity(customer, identity, db)


@app.post(f"{settings.api_prefix}/customers/{{customer_id}}/key-change/review")
async def review_customer_key_change(customer_id: int, body: CustomerKeyChangeReviewInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="只有管理员或客户主管可以审批客户关键字段变更")
    customer = await _customer_or_404(customer_id, identity, db)
    data = customer.data or {}; pending = data.get("key_change") or {}
    if pending.get("status") != "待审批": raise HTTPException(status_code=409, detail="该客户没有待审批的关键字段变更")
    if not body.approved and not body.comment.strip(): raise HTTPException(status_code=422, detail="驳回时必须填写原因")
    after = pending.get("after") or {}; new_title = str(after.get("title") or "").strip(); new_credit_code = str(after.get("credit_code") or "").strip()
    if body.approved:
        await _ensure_unique_customer_name(new_title, db, exclude_id=customer.id)
    if body.approved and new_credit_code:
        customers = list((await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "customer", BusinessRecord.id != customer.id))).all())
        if any(str((item.data or {}).get("credit_code") or "").strip().casefold() == new_credit_code.casefold() for item in customers):
            raise HTTPException(status_code=409, detail="统一社会信用代码已被其他客户使用")
    reviewed = {**pending, "status": "已通过" if body.approved else "已驳回", "reviewed_by": identity["username"], "reviewed_at": datetime.now().isoformat(timespec="seconds"), "review_comment": body.comment.strip()}
    if body.approved:
        old_title = customer.title; customer.title = new_title; customer.customer = new_title
        customer.data = {**data, "credit_code": new_credit_code, "key_change": reviewed}
        linked = list((await db.scalars(select(BusinessRecord).where(BusinessRecord.customer == old_title, BusinessRecord.module.in_(["contract", "case", "task", "finance"])))).all())
        for item in linked: item.customer = new_title
    else:
        customer.data = {**data, "key_change": reviewed}
    _mark_customer_modified(customer, identity)
    db.add(_customer_event(customer, "客户关键字段审批通过" if body.approved else "客户关键字段审批驳回", identity, body.comment))
    await db.commit(); await db.refresh(customer)
    return await _record_dict_for_identity(customer, identity, db)


def _portal_code_hash(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


async def _portal_customer(body: CustomerPortalLoginInput, db: AsyncSession) -> BusinessRecord:
    customers = list((await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "customer"))).all())
    customer = next((item for item in customers if str(((item.data or {}).get("portal_access") or {}).get("account") or "").casefold() == body.account.strip().casefold()), None)
    portal = (customer.data or {}).get("portal_access") if customer else None
    if not customer or not portal or not portal.get("enabled") or portal.get("activation_code_hash") != _portal_code_hash(body.activation_code.strip()):
        raise HTTPException(status_code=401, detail="客户服务账号或激活码无效")
    return customer


@app.post(f"{settings.api_prefix}/customers/{{customer_id}}/portal/open")
async def open_customer_portal(customer_id: int, body: CustomerPortalActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或客户主管可以开通客户服务端")
    customer = await _customer_or_404(customer_id, identity, db)
    contracts = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "contract", BusinessRecord.customer == customer.title,
        BusinessRecord.status.in_(CASE_SOURCE_CONTRACT_STATUSES),
    ))).all())
    if (customer.data or {}).get("level") != "签约客户" and not contracts:
        raise HTTPException(status_code=409, detail="客户签约或存在审批通过的合同后才能开通客户服务端")
    old_portal = (customer.data or {}).get("portal_access") or {}
    account = str(old_portal.get("account") or f"vip-{customer.serial_no}").lower()
    activation_code = uuid4().hex
    portal = {
        "account": account, "enabled": True, "activation_code_hash": _portal_code_hash(activation_code),
        "opened_by": identity["username"], "opened_at": datetime.now().isoformat(timespec="seconds"),
        "comment": body.comment.strip(),
    }
    customer.data = {**(customer.data or {}), "portal_access": portal}
    db.add(_customer_event(customer, "开通客户服务端", identity, body.comment or f"服务账号：{account}"))
    await db.commit()
    return {"customer_id": customer.id, "account": account, "activation_code": activation_code, "notice": "激活码仅本次显示，请安全交付客户"}


@app.post(f"{settings.api_prefix}/customers/{{customer_id}}/portal/close")
async def close_customer_portal(customer_id: int, body: CustomerPortalActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或客户主管可以停用客户服务端")
    customer = await _customer_or_404(customer_id, identity, db)
    portal = (customer.data or {}).get("portal_access") or {}
    if not portal.get("enabled"):
        raise HTTPException(status_code=409, detail="客户服务端尚未开通或已经停用")
    customer.data = {**(customer.data or {}), "portal_access": {**portal, "enabled": False, "closed_by": identity["username"], "closed_at": datetime.now().isoformat(timespec="seconds"), "close_comment": body.comment.strip()}}
    db.add(_customer_event(customer, "停用客户服务端", identity, body.comment))
    await db.commit()
    return {"customer_id": customer.id, "account": portal.get("account", ""), "enabled": False}


@app.post(f"{settings.api_prefix}/customer-portal/overview")
async def customer_portal_overview(body: CustomerPortalLoginInput, db: AsyncSession = Depends(get_db)):
    customer = await _portal_customer(body, db)
    records = list((await db.scalars(select(BusinessRecord).where(BusinessRecord.customer == customer.title, BusinessRecord.module.in_(["contract", "case"])))).all())
    record_ids = [item.id for item in records]
    files = list((await db.scalars(select(FileAttachment).where(FileAttachment.record_id.in_(record_ids)))).all()) if record_ids else []
    return {
        "customer": {"id": customer.id, "serial_no": customer.serial_no, "name": customer.title, "level": (customer.data or {}).get("level", "")},
        "contracts": [_record_dict(item) for item in records if item.module == "contract"],
        "cases": [_record_dict(item) for item in records if item.module == "case"],
        "documents": [_attachment_dict(item, next((record for record in records if record.id == item.record_id), None)) for item in files],
    }


@app.post(f"{settings.api_prefix}/customer-portal/files/{{attachment_id}}/download")
async def customer_portal_download(attachment_id: int, body: CustomerPortalLoginInput, db: AsyncSession = Depends(get_db)):
    customer = await _portal_customer(body, db)
    attachment = await db.get(FileAttachment, attachment_id)
    record = await db.get(BusinessRecord, attachment.record_id) if attachment and attachment.record_id else None
    if not attachment or not record or record.customer != customer.title or record.module not in {"contract", "case"}:
        raise HTTPException(status_code=404, detail="客户文档不存在")
    path = Path(attachment.path)
    if not path.is_file():
        raise HTTPException(status_code=404, detail="客户文档文件不存在")
    return FileResponse(path, media_type=attachment.content_type, filename=Path(attachment.original_name).name)


@app.post(f"{settings.api_prefix}/customer-portal/demands", status_code=status.HTTP_201_CREATED)
async def customer_portal_create_demand(body: CustomerPortalDemandInput, db: AsyncSession = Depends(get_db)):
    customer = await _portal_customer(body, db)
    case_record = None
    if body.case_no.strip():
        case_record = await db.scalar(select(BusinessRecord).where(BusinessRecord.module == "case", BusinessRecord.customer == customer.title, BusinessRecord.serial_no == body.case_no.strip()))
        if not case_record:
            raise HTTPException(status_code=422, detail="所选案件不属于当前客户")
    owner = customer.owner
    if not await db.scalar(select(User.id).where(User.username == owner, User.is_active.is_(True))):
        raise HTTPException(status_code=409, detail="客户负责人不存在或已停用，暂不能提交需求")
    task = BusinessRecord(
        module="task", serial_no=f"RW{datetime.now():%Y%m%d%H%M%S%f}", title=f"客户需求—{body.title.strip()}",
        customer=customer.title, status="待接收", owner=owner, department=customer.department, description=body.content.strip(),
        data={"deadline": str(date.today() + timedelta(days=7)), "priority": "普通", "source": "客户任务", "task_type": "手动任务", "initiator": owner, "collaborators": [], "case_no": case_record.serial_no if case_record else "", "case_id": case_record.id if case_record else None, "portal_customer_id": customer.id, "portal_account": body.account.strip()},
    )
    db.add(task); await db.flush()
    await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action="客户服务端提交需求", to_status="待接收", operator=body.account.strip(), comment=body.content.strip()), db, content="客户提交了新的服务需求.")
    await db.commit(); await db.refresh(task)
    return {"id": task.id, "serial_no": task.serial_no, "status": task.status, "owner": task.owner}


def _approval_step_dict(step: ContractApprovalStep) -> dict:
    return {"id": step.id, "contract_record_id": step.contract_record_id, "step_order": step.step_order, "approver": step.approver, "status": step.status, "comment": step.comment, "acted_at": step.acted_at, "created_at": step.created_at}


def _normalize_external_contract_numbers(data: dict) -> dict:
    raw_values = data.get("external_contract_numbers")
    if raw_values is None:
        raw_values = [data.get("external_contract_no", "")]
    if not isinstance(raw_values, list):
        raise HTTPException(status_code=422, detail="外部合同号必须为数组")
    values = list(dict.fromkeys(str(value or "").strip() for value in raw_values if str(value or "").strip()))
    if len(values) > 50:
        raise HTTPException(status_code=422, detail="一个合同最多关联 50 个外部合同号")
    if any(len(value) > 128 for value in values):
        raise HTTPException(status_code=422, detail="外部合同号长度不能超过 128 个字符")
    return {**data, "external_contract_numbers": values, "external_contract_no": values[0] if values else ""}


async def _resolve_contract_customer(customer_name: str, data: dict, identity: dict, db: AsyncSession) -> BusinessRecord:
    """Resolve the selected customer and keep contract linkage authoritative."""
    name = customer_name.strip()
    customer_id = int(data.get("customer_id") or 0)
    conditions = [BusinessRecord.module == "customer", BusinessRecord.status != "已回收"]
    if customer_id:
        conditions.append(BusinessRecord.id == customer_id)
    else:
        conditions.append(BusinessRecord.title == name)
    customer = await db.scalar(
        select(BusinessRecord)
        .where(*conditions, *(await _record_scope_conditions(identity, db)))
        .order_by(BusinessRecord.id.desc())
    )
    if not customer or customer.title != name:
        raise HTTPException(status_code=422, detail="请选择当前账号可见的有效客户，不能手工录入未登记客户名称")
    return customer


async def _user_has_job_permission(user: User, permission_name: str, db: AsyncSession) -> bool:
    """Resolve business action permission from the employee's configured job role."""
    if user.role == "admin":
        return True
    profile = user.profile or {}
    job_role_name = str(profile.get("position") or profile.get("staff_role") or "").strip()
    if not job_role_name:
        return False
    job_role = await db.scalar(select(JobRole).where(JobRole.name == job_role_name, JobRole.is_active.is_(True)))
    permissions = set(job_role.permissions or []) if job_role else set()
    # 旧的人事岗位“调查专员”已经使用“调查取证”描述其工作范围；
    # 调查线索提交和现场材料扫描是该范围内不可拆分的专用动作。
    implied_permissions = {
        "调查取证": {"线索提交", "扫描上传"},
    }
    return permission_name in permissions or permission_name in set().union(*(implied_permissions.get(item, set()) for item in permissions))


@app.get(f"{settings.api_prefix}/investigations/action-capabilities")
async def investigation_action_capabilities(
    record_ids: str = Query(default="", max_length=1200),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    """Return the current user's real investigation actions for visible records.

    The investigation workbench must not infer approval authority from a display
    role.  In particular, clue review is a job-role permission, customer review
    belongs to the relevant customer manager, and certificate registration also
    requires authority over the notary record itself.
    """
    try:
        requested_ids = list(dict.fromkeys(int(value) for value in record_ids.split(",") if value.strip()))
    except ValueError:
        raise HTTPException(status_code=422, detail="记录编号格式无效")
    if not requested_ids:
        return {"items": {}}
    if len(requested_ids) > 100:
        raise HTTPException(status_code=422, detail="一次最多查询 100 条调查记录的操作权限")
    records = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.id.in_(requested_ids),
        BusinessRecord.module.in_({"clue", "notary"}),
        *(await _record_scope_conditions(identity, db)),
    ))).all())
    user = await db.scalar(select(User).where(User.username == identity["username"]))
    can_review_clue = bool(user and await _user_has_job_permission(user, "线索审批", db))
    can_review_notary = bool(user and await _user_has_job_permission(user, "公证审核", db))
    can_register_certificate = bool(user and await _user_has_job_permission(user, "公证书号码登记", db))
    customer_names = {record.customer.strip() for record in records if record.module == "clue" and record.customer.strip()}
    customer_managers: dict[str, set[str]] = {}
    if customer_names:
        customers = list((await db.scalars(select(BusinessRecord).where(
            BusinessRecord.module == "customer", BusinessRecord.title.in_(customer_names),
        ))).all())
        customer_managers = {
            customer.title.strip(): set((customer.data or {}).get("customer_managers") or [customer.owner])
            for customer in customers
        }
    items: dict[str, dict[str, bool]] = {}
    for record in records:
        can_manage_record = (
            identity.get("role") == "admin"
            or record.owner == identity["username"]
            or (identity.get("role") == "manager" and user and record.department == user.department)
        )
        items[str(record.id)] = {
            "review_clue": record.module == "clue" and can_review_clue,
            "review_customer_clue": record.module == "clue" and (
                identity.get("role") == "admin"
                or identity["username"] in customer_managers.get(record.customer.strip(), set())
            ),
            "review_notary": record.module == "notary" and can_review_notary,
            "register_notary_certificate": record.module == "notary" and can_register_certificate and can_manage_record,
        }
    return {"items": items}


async def _ensure_contract_approval_access(contract_id: int, identity: dict, db: AsyncSession) -> BusinessRecord:
    contract = await db.get(BusinessRecord, contract_id)
    if not contract or contract.module != "contract":
        raise HTTPException(status_code=404, detail="合同不存在")
    assigned = await db.scalar(select(ContractApprovalStep.id).where(
        ContractApprovalStep.contract_record_id == contract_id,
        ContractApprovalStep.approver == identity["username"],
    ))
    if assigned or identity.get("role") == "admin":
        return contract
    return await _ensure_record_module(contract_id, "contract", identity, db)


@app.post(f"{settings.api_prefix}/contracts", status_code=status.HTTP_201_CREATED)
async def create_contract_draft(body: ContractDraftInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if await db.scalar(select(BusinessRecord.id).where(BusinessRecord.serial_no == body.serial_no.strip())):
        raise HTTPException(status_code=409, detail="合同编号已存在")
    data = _normalize_external_contract_numbers(dict(body.data or {}))
    customer = await _resolve_contract_customer(body.customer, data, identity, db)
    customer_data = customer.data or {}
    data = {**data, "customer_id": customer.id, "customer_no": customer.serial_no, "customer_manager": "、".join(customer_data.get("customer_managers") or [customer.owner])}
    if float(data.get("amount") or 0) < 0:
        raise HTTPException(status_code=422, detail="合同金额不能小于零")
    owner = body.owner.strip()
    department = body.department.strip()
    if identity.get("role") != "admin":
        current_user = await db.scalar(select(User).where(User.username == identity["username"]));
        if not current_user: raise HTTPException(status_code=401, detail="当前用户不存在")
        department = current_user.department
        if identity.get("role") == "user": owner = identity["username"]
    item = BusinessRecord(module="contract", serial_no=body.serial_no.strip(), title=body.title.strip(), customer=customer.title, status="草稿", owner=owner, department=department, description=body.description.strip(), data=data)
    db.add(item); await db.flush()
    db.add(WorkflowEvent(record_id=item.id, action="创建合同草稿", to_status="草稿", operator=identity["username"], comment="通过合同专用入口创建"))
    await db.commit(); await db.refresh(item)
    return await _record_dict_for_identity(item, identity, db)


@app.patch(f"{settings.api_prefix}/contracts/{{contract_id}}")
async def update_contract_draft(contract_id: int, body: ContractDraftInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _ensure_record_module(contract_id, "contract", identity, db)
    await _require_record_owner_or_manager(item, identity, db)
    if item.status not in {"草稿", "已拒绝"}:
        raise HTTPException(status_code=409, detail="合同提交审批后不能直接编辑，请使用合同变更流程")
    duplicate = await db.scalar(select(BusinessRecord.id).where(BusinessRecord.serial_no == body.serial_no.strip(), BusinessRecord.id != item.id))
    if duplicate: raise HTTPException(status_code=409, detail="合同编号已存在")
    data = _normalize_external_contract_numbers(dict(body.data or {}))
    customer = await _resolve_contract_customer(body.customer, data, identity, db)
    customer_data = customer.data or {}
    data = {**data, "customer_id": customer.id, "customer_no": customer.serial_no, "customer_manager": "、".join(customer_data.get("customer_managers") or [customer.owner])}
    if float(data.get("amount") or 0) < 0: raise HTTPException(status_code=422, detail="合同金额不能小于零")
    owner = body.owner.strip(); department = body.department.strip()
    if identity.get("role") != "admin":
        current_user = await db.scalar(select(User).where(User.username == identity["username"]));
        if not current_user: raise HTTPException(status_code=401, detail="当前用户不存在")
        department = current_user.department
        if identity.get("role") == "user": owner = identity["username"]
    item.serial_no = body.serial_no.strip(); item.title = body.title.strip(); item.customer = customer.title
    item.owner = owner; item.department = department; item.description = body.description.strip(); item.data = data
    db.add(WorkflowEvent(record_id=item.id, action="修改合同草稿", from_status=item.status, to_status=item.status, operator=identity["username"], comment="通过合同专用入口修改"))
    await db.commit(); await db.refresh(item)
    return await _record_dict_for_identity(item, identity, db)


@app.get(f"{settings.api_prefix}/contracts/{{contract_id}}/approvals")
async def contract_approvals(contract_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    contract = await _ensure_contract_approval_access(contract_id, identity, db)
    steps = (await db.scalars(select(ContractApprovalStep).where(ContractApprovalStep.contract_record_id == contract_id).order_by(ContractApprovalStep.step_order))).all()
    return {"contract": await _record_dict_for_identity(contract, identity, db), "items": [_approval_step_dict(x) for x in steps], "current_step": next((_approval_step_dict(x) for x in steps if x.status == "待审批"), None)}


@app.post(f"{settings.api_prefix}/contracts/{{contract_id}}/submit")
async def submit_contract(contract_id: int, body: ContractSubmitInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    contract = await _ensure_record_module(contract_id, "contract", identity, db)
    await _require_record_owner_or_manager(contract, identity, db)
    if contract.status not in {"草稿", "已拒绝"}: raise HTTPException(status_code=409, detail="只有草稿或已拒绝合同可以重新提交")
    approvers = [x.strip() for x in body.approvers if x.strip()]
    if len(approvers) != 1: raise HTTPException(status_code=422, detail="合同审批只能选择一名具有合同审批权限的人员")
    if approvers[0] == identity["username"]: raise HTTPException(status_code=422, detail="合同发起人不能审批自己提交的合同")
    approver_user = await db.scalar(select(User).where(User.username == approvers[0], User.is_active.is_(True)))
    if not approver_user: raise HTTPException(status_code=422, detail="审批人不存在或已停用")
    if not await _user_has_job_permission(approver_user, "合同审批", db):
        raise HTTPException(status_code=422, detail="所选人员的岗位角色没有合同审批权限")
    await db.execute(delete(ContractApprovalStep).where(ContractApprovalStep.contract_record_id == contract_id))
    for index, approver in enumerate(approvers, 1):
        db.add(ContractApprovalStep(contract_record_id=contract_id, step_order=index, approver=approver, status="待审批" if index == 1 else "等待中"))
    old = contract.status
    contract.status = "审批中"
    contract.data = {
        **(contract.data or {}),
        "approval_count": len(approvers),
        "submitted_at": datetime.now().isoformat(timespec="seconds"),
        "submitted_by": identity["username"],
        "submit_comment": body.comment.strip(),
        "current_approver": approvers[0],
    }
    db.add(WorkflowEvent(record_id=contract.id, action="提交合同审批", from_status=old, to_status="审批中", operator=identity["username"], comment=body.comment or f"审批人：{approvers[0]}")); await db.commit(); await db.refresh(contract)
    return await _record_dict_for_identity(contract, identity, db)


@app.post(f"{settings.api_prefix}/contracts/{{contract_id}}/approve")
async def approve_contract(contract_id: int, body: ContractApprovalInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    contract = await _ensure_contract_approval_access(contract_id, identity, db)
    if contract.status != "审批中": raise HTTPException(status_code=409, detail="合同当前不在审批中")
    if not body.approved and not body.comment.strip(): raise HTTPException(status_code=422, detail="拒绝时必须填写审批意见")
    steps = (await db.scalars(select(ContractApprovalStep).where(ContractApprovalStep.contract_record_id == contract_id).order_by(ContractApprovalStep.step_order))).all()
    current = next((x for x in steps if x.status == "待审批"), None)
    if not current: raise HTTPException(status_code=409, detail="合同没有待处理的审批节点")
    if current.approver != identity["username"]: raise HTTPException(status_code=403, detail=f"当前节点应由 {current.approver} 审批，管理员也不能代替指定审批人操作")
    current.comment = body.comment.strip(); current.acted_at = datetime.now(); old = contract.status
    if not body.approved:
        current.status = "已拒绝"; contract.status = "已拒绝"
        contract.data = {**(contract.data or {}), "current_approver": ""}
        for step in steps:
            if step.status == "等待中": step.status = "已取消"
        action = "合同审批拒绝"
    else:
        current.status = "已通过"; next_step = next((x for x in steps if x.status == "等待中"), None)
        if next_step:
            next_step.status = "待审批"; action = "合同节点通过"
            contract.data = {**(contract.data or {}), "current_approver": next_step.approver}
        else:
            contract.status = "已通过"; action = "合同审批完成"
            contract.data = {**(contract.data or {}), "current_approver": ""}
            seal_application_id = int((contract.data or {}).get("seal_application_id") or 0)
            if seal_application_id and (contract.data or {}).get("sync_seal"):
                seal_application = await db.get(BusinessRecord, seal_application_id)
                if seal_application and seal_application.module == "seal" and seal_application.status == "草稿":
                    seal_file_count = int(await db.scalar(select(func.count()).select_from(FileAttachment).where(FileAttachment.record_id == seal_application.id, FileAttachment.category == "用印文件")) or 0)
                    if seal_file_count:
                        seal_application.status = "待审批"
                        contract.data = {**(contract.data or {}), "sync_seal_submitted_at": datetime.now().isoformat(timespec="seconds"), "sync_seal_file_required": False}
                        db.add(WorkflowEvent(record_id=seal_application.id, action="合同通过后自动提交同步用印", from_status="草稿", to_status="待审批", operator=identity["username"], comment=f"来源合同 {contract.serial_no} 已审批通过"))
                    else:
                        contract.data = {**(contract.data or {}), "sync_seal_file_required": True}
                        db.add(WorkflowEvent(record_id=seal_application.id, action="合同通过后等待用印文件", from_status="草稿", to_status="草稿", operator=identity["username"], comment=f"来源合同 {contract.serial_no} 已通过；请在用印中心上传真实用印文件后提交审批"))
    db.add(WorkflowEvent(record_id=contract.id, action=action, from_status=old, to_status=contract.status, operator=identity["username"], comment=f"第{current.step_order}级 {current.approver}：{body.comment}")); await db.commit(); await db.refresh(contract)
    return await _record_dict_for_identity(contract, identity, db)


@app.post(f"{settings.api_prefix}/contracts/{{contract_id}}/seal-application", status_code=status.HTTP_201_CREATED)
async def create_contract_seal_application(contract_id: int, body: ContractSealApplicationInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    contract = await _ensure_record_module(contract_id, "contract", identity, db)
    await _require_record_owner_or_manager(contract, identity, db)
    if contract.status == "审批中" and body.submit:
        raise HTTPException(status_code=409, detail="合同审批中只能保存同步用印资料，审批通过后系统会自动提交")
    if contract.status not in {"审批中", "已通过", "履行中", "已完成"}:
        raise HTTPException(status_code=409, detail="合同提交审批后才能配置同步用印")
    if body.submit:
        raise HTTPException(status_code=409, detail="请先保存合同用印草稿，在用印中心上传真实用印文件后再提交审批")
    existing_id = int((contract.data or {}).get("seal_application_id") or 0)
    if existing_id:
        existing = await db.get(BusinessRecord, existing_id)
        if existing:
            raise HTTPException(status_code=409, detail=f"合同已生成用印申请 {existing.serial_no}")
    asset = await db.get(SealAsset, body.seal_asset_id)
    if not asset:
        raise HTTPException(status_code=404, detail="印章不存在")
    if asset.status != "可用":
        raise HTTPException(status_code=409, detail=f"印章当前状态为“{asset.status}”，不能申请")
    serial = f"YY{datetime.now():%Y%m%d%H%M%S}{uuid4().hex[:3].upper()}"
    seal_status = "草稿"
    seal = BusinessRecord(
        module="seal",
        serial_no=serial,
        title=f"{contract.title}合同用印",
        customer=contract.customer,
        status=seal_status,
        owner=identity["username"],
        department=contract.department,
        description=body.description,
        data={
            "contract_record_id": contract.id,
            "contract_no": contract.serial_no,
            "use_type": "合同用印",
            "seal_asset_id": asset.id,
            "seal_type": asset.seal_type,
            "seal_name": asset.name,
            "copies": body.copies,
            "purpose": body.purpose,
            "use_date": str(body.use_date),
            "delivery_method": body.delivery_method,
            "document_names": body.document_names,
        },
    )
    db.add(seal)
    await db.flush()
    contract.data = {
        **(contract.data or {}),
        "seal_application_id": seal.id,
        "seal_application_no": seal.serial_no,
        "seal_requested_at": datetime.now().isoformat(timespec="seconds"),
        "sync_seal": contract.status == "审批中",
    }
    db.add_all([
        WorkflowEvent(record_id=seal.id, action="创建合同用印申请", to_status=seal_status, operator=identity["username"], comment=f"来源合同 {contract.serial_no}｜{asset.name}｜{body.copies}份"),
        WorkflowEvent(record_id=contract.id, action="配置同步用印" if contract.status == "审批中" else "发起合同用印", from_status=contract.status, to_status=contract.status, operator=identity["username"], comment=f"生成用印申请 {seal.serial_no}" + ("并提交审批" if body.submit else "，保存为草稿")),
    ])
    await db.commit()
    await db.refresh(seal)
    return await _seal_record_dict(seal, db)


@app.post(f"{settings.api_prefix}/contracts/{{contract_id}}/investigation", status_code=status.HTTP_201_CREATED)
async def create_contract_investigation(contract_id: int, body: ContractInvestigationInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    contract = await _ensure_record_module(contract_id, "contract", identity, db)
    await _require_record_owner_or_manager(contract, identity, db)
    if contract.status not in {"已通过", "履行中", "已完成"}:
        raise HTTPException(status_code=409, detail="合同审批通过后才能新建调查任务")
    if body.authorized_to < body.authorized_from:
        raise HTTPException(status_code=422, detail="授权结束日期不能早于开始日期")
    user = await db.scalar(select(User).where(User.username == identity["username"]))
    owner = body.owner.strip()
    if identity.get("role") == "user":
        owner = identity["username"]
    if owner:
        assignee = await db.scalar(select(User).where(or_(User.username == owner, User.display_name == owner), User.is_active.is_(True)))
        if not assignee:
            raise HTTPException(status_code=422, detail="调查任务负责人不存在或已停用")
        owner = assignee.username
    duplicate = await db.scalar(select(BusinessRecord).where(
        BusinessRecord.module == "investigation",
        BusinessRecord.data["contract_id"].as_integer() == contract.id,
        BusinessRecord.title == body.title.strip(),
        BusinessRecord.status.not_in({"已取消", "已完成"}),
    ))
    if duplicate:
        raise HTTPException(status_code=409, detail=f"该合同已有同名调查任务 {duplicate.serial_no}")
    serial = f"DC{datetime.now():%Y%m%d%H%M%S}{uuid4().hex[:4].upper()}"
    investigation = BusinessRecord(
        module="investigation",
        serial_no=serial,
        title=body.title.strip(),
        customer=contract.customer,
        status="进行中" if owner else "待分配",
        owner=owner,
        department=user.department if user else contract.department,
        description=body.description,
        data={
            "contract_id": contract.id,
            "contract_no": contract.serial_no,
            "authorized_from": str(body.authorized_from),
            "authorized_to": str(body.authorized_to),
            "region": body.region.strip(),
            "right_type": body.right_type.strip(),
            "customer_review": body.customer_review,
            "publisher": identity["username"],
            "assigner": identity["username"] if owner else "",
            "source_owner": (contract.data or {}).get("source_person") or contract.owner,
        },
    )
    db.add(investigation)
    await db.flush()
    linked_ids = list((contract.data or {}).get("investigation_ids", [])); linked_ids.append(investigation.id)
    contract.data = {**(contract.data or {}), "investigation_ids": list(dict.fromkeys(linked_ids)), "last_investigation_no": serial}
    db.add_all([
        WorkflowEvent(record_id=investigation.id, action="从合同创建调查任务", to_status=investigation.status, operator=identity["username"], comment=f"来源合同 {contract.serial_no}"),
        WorkflowEvent(record_id=contract.id, action="新建调查任务", from_status=contract.status, to_status=contract.status, operator=identity["username"], comment=f"生成调查任务 {serial}"),
    ])
    await db.commit(); await db.refresh(investigation)
    return _record_dict(investigation)


@app.post(f"{settings.api_prefix}/contracts/{{contract_id}}/archive")
async def archive_contract(contract_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    contract = await _ensure_record_module(contract_id, "contract", identity, db)
    await _require_record_owner_or_manager(contract, identity, db)
    if contract.status == "已归档":
        raise HTTPException(status_code=409, detail="合同已经归档")
    if contract.status not in {"已通过", "履行中", "已完成"}:
        raise HTTPException(status_code=409, detail="只有审批全部通过且处于履行或完成阶段的合同可以归档")
    pending_or_rejected = int(await db.scalar(select(func.count()).select_from(ContractApprovalStep).where(ContractApprovalStep.contract_record_id == contract.id, ContractApprovalStep.status != "已通过")) or 0)
    if pending_or_rejected:
        raise HTTPException(status_code=409, detail="合同仍有未通过的审批节点，不能归档")
    previous = contract.status
    contract.status = "已归档"
    contract.data = {**(contract.data or {}), "archived_at": datetime.now().isoformat(timespec="seconds")}
    db.add(WorkflowEvent(record_id=contract.id, action="合同归档", from_status=previous, to_status="已归档", operator=identity["username"], comment="合同列表批量操作归档"))
    await db.commit(); await db.refresh(contract)
    return await _record_dict_for_identity(contract, identity, db)


def _contract_event_dict(event: ContractEvent) -> dict:
    return {
        "id": event.id,
        "contract_record_id": event.contract_record_id,
        "content": event.content,
        "operator": event.operator,
        "created_at": event.created_at,
    }


@app.get(f"{settings.api_prefix}/contracts/{{contract_id}}/events")
async def list_contract_events(contract_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    contract = await _ensure_record_visible(contract_id, identity, db)
    if contract.module != "contract":
        raise HTTPException(status_code=404, detail="合同不存在")
    events = (await db.scalars(
        select(ContractEvent)
        .where(ContractEvent.contract_record_id == contract.id)
        .order_by(ContractEvent.created_at.desc(), ContractEvent.id.desc())
    )).all()
    return {"items": [_contract_event_dict(event) for event in events], "total": len(events)}


@app.post(f"{settings.api_prefix}/contracts/{{contract_id}}/events", status_code=status.HTTP_201_CREATED)
async def create_contract_event(contract_id: int, body: ContractEventInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    contract = await _ensure_record_module(contract_id, "contract", identity, db)
    await _require_record_owner_or_manager(contract, identity, db)
    if contract.status == "已归档":
        raise HTTPException(status_code=409, detail="已归档合同不可新增事项记录")
    content = body.content.strip()
    if not content:
        raise HTTPException(status_code=422, detail="事项内容不能为空")
    event = ContractEvent(contract_record_id=contract.id, content=content, operator=identity["username"])
    db.add(event)
    db.add(WorkflowEvent(
        record_id=contract.id,
        action="新增合同事项",
        from_status=contract.status,
        to_status=contract.status,
        operator=identity["username"],
        comment=content,
    ))
    await db.commit()
    await db.refresh(event)
    return _contract_event_dict(event)


@app.get(f"{settings.api_prefix}/contracts/{{contract_id}}/changes")
async def contract_changes(contract_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    contract = await _ensure_record_visible(contract_id, identity, db)
    if contract.module != "contract": raise HTTPException(status_code=404, detail="合同不存在")
    events = (await db.scalars(select(WorkflowEvent).where(WorkflowEvent.record_id == contract.id, WorkflowEvent.action.in_(["提交合同变更", "合同变更审批通过", "合同变更审批驳回"])).order_by(WorkflowEvent.created_at.desc(), WorkflowEvent.id.desc()))).all()
    items = []
    for event in events:
        try: detail = json.loads(event.comment or "{}")
        except json.JSONDecodeError: detail = {"reason": event.comment}
        items.append({"id": event.id, "operator": event.operator, "created_at": event.created_at, **detail})
    return {"items": items, "total": len(items)}


@app.post(f"{settings.api_prefix}/contracts/{{contract_id}}/changes", status_code=status.HTTP_201_CREATED)
async def change_contract(contract_id: int, body: ContractChangeInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    contract = await _ensure_record_visible(contract_id, identity, db)
    if contract.module != "contract": raise HTTPException(status_code=404, detail="合同不存在")
    await _require_record_owner_or_manager(contract, identity, db)
    if contract.status not in {"已通过", "履行中"}: raise HTTPException(status_code=409, detail="只有已通过或履行中的合同可以发起变更")
    data = dict(contract.data or {}); changes = []
    if (data.get("pending_change") or {}).get("status") == "待审批":
        raise HTTPException(status_code=409, detail="已有合同变更正在审批")
    requested_external_numbers = body.external_contract_numbers
    if requested_external_numbers is None and body.external_contract_no is not None:
        requested_external_numbers = [body.external_contract_no]
    normalized_external_numbers = None
    if requested_external_numbers is not None:
        normalized_external_numbers = _normalize_external_contract_numbers({"external_contract_numbers": requested_external_numbers})["external_contract_numbers"]
    candidates = {
        "title": (contract.title, body.title),
        "amount": (data.get("amount"), body.amount),
        "external_contract_numbers": (data.get("external_contract_numbers") or ([data.get("external_contract_no")] if data.get("external_contract_no") else []), normalized_external_numbers),
        "end_date": (data.get("end_date", ""), str(body.end_date) if body.end_date else None),
    }
    labels = {"title": "合同名称", "amount": "合同金额", "external_contract_numbers": "外部合同号", "end_date": "合同期限"}
    for key, (before, after) in candidates.items():
        if after is not None and after != before:
            changes.append({"field": key, "label": labels[key], "before": before, "after": after})
    if not changes: raise HTTPException(status_code=422, detail="至少填写一项发生变化的合同内容")
    detail = {"status": "待审批", "change_type": body.change_type.strip(), "reason": body.reason.strip(), "changes": changes, "requested_by": identity["username"], "requested_at": datetime.now().isoformat(timespec="seconds")}
    contract.data = {**data, "pending_change": detail}
    db.add(WorkflowEvent(record_id=contract.id, action="提交合同变更", from_status=contract.status, to_status=contract.status, operator=identity["username"], comment=json.dumps(detail, ensure_ascii=False)))
    await db.commit(); await db.refresh(contract)
    return {"contract": await _record_dict_for_identity(contract, identity, db), **detail}


@app.post(f"{settings.api_prefix}/contracts/{{contract_id}}/changes/review")
async def review_contract_change(contract_id: int, body: ContractChangeReviewInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager", "auditor"}:
        raise HTTPException(status_code=403, detail="只有合同管理员、财务主管或管理员可以审批合同变更")
    contract = await _ensure_record_visible(contract_id, identity, db)
    if contract.module != "contract": raise HTTPException(status_code=404, detail="合同不存在")
    data = dict(contract.data or {}); pending = data.get("pending_change") or {}
    if pending.get("status") != "待审批": raise HTTPException(status_code=409, detail="该合同没有待审批的变更")
    if not body.approved and not body.comment.strip(): raise HTTPException(status_code=422, detail="驳回时必须填写原因")
    if body.approved:
        for change in pending.get("changes", []):
            key = change.get("field"); value = change.get("after")
            if key == "title": contract.title = str(value)
            elif key == "external_contract_numbers": data = _normalize_external_contract_numbers({**data, "external_contract_numbers": value})
            elif key in {"amount", "end_date"}: data[key] = value
        data["last_changed_at"] = datetime.now().isoformat(timespec="seconds")
        data["change_count"] = int(data.get("change_count", 0)) + 1
    reviewed = {**pending, "status": "已通过" if body.approved else "已驳回", "reviewed_by": identity["username"], "reviewed_at": datetime.now().isoformat(timespec="seconds"), "review_comment": body.comment.strip()}
    contract.data = {**data, "pending_change": reviewed}
    action = "合同变更审批通过" if body.approved else "合同变更审批驳回"
    db.add(WorkflowEvent(record_id=contract.id, action=action, from_status=contract.status, to_status=contract.status, operator=identity["username"], comment=json.dumps(reviewed, ensure_ascii=False)))
    await db.commit(); await db.refresh(contract)
    return {"contract": await _record_dict_for_identity(contract, identity, db), **reviewed}


@app.get(f"{settings.api_prefix}/receivables")
async def list_receivables(
    keyword: str = "", receivable_status: str = "",
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    plans = (await db.scalars(select(ReceivablePlan).order_by(ReceivablePlan.due_date))).all()
    contract_ids = {plan.contract_record_id for plan in plans}
    contracts = {record.id: record for record in (await db.scalars(select(BusinessRecord).where(BusinessRecord.id.in_(contract_ids), *(await _record_scope_conditions(identity, db))))).all()} if contract_ids else {}
    items = [_receivable_dict(plan, contracts[plan.contract_record_id]) for plan in plans if plan.contract_record_id in contracts]
    if keyword:
        needle = keyword.casefold()
        items = [item for item in items if needle in " ".join([item["contract_no"], item["contract_title"], item["customer"], item["phase"], item["payer"]]).casefold()]
    if receivable_status:
        items = [item for item in items if item["status"] == receivable_status]
    all_amount = sum(item["amount"] for item in items)
    received = sum(item["received_amount"] for item in items)
    overdue = sum(item["remaining_amount"] for item in items if item["status"] == "已逾期")
    return {"items": items, "total": len(items), "summary": {"amount": all_amount, "received": received, "remaining": all_amount - received, "overdue": overdue}}


INVESTIGATION_MATERIAL_CATEGORIES = {
    "investigation": ["调查授权书", "权利证明", "调查资料", "客户指示"],
    "clue": ["现场照片", "网页截图", "购买记录", "调查报告"],
    "notary": ["公证书扫描件", "公证费发票", "公证实物登记"],
    "evidence": ["证据原件", "证据扫描件", "证据目录"],
}


async def _sync_investigation_materials(record: BusinessRecord, db: AsyncSession) -> list[str]:
    categories = sorted(set((await db.scalars(select(FileAttachment.category).where(FileAttachment.record_id == record.id))).all()))
    record.data = {**(record.data or {}), "material_categories": categories, "material_count": int(await db.scalar(select(func.count()).select_from(FileAttachment).where(FileAttachment.record_id == record.id)) or 0)}
    return categories


@app.get(f"{settings.api_prefix}/investigations/clues/import-template")
async def clue_import_template(_: dict = Depends(current_identity)):
    content = "\ufeff线索标题,客户,调查平台,侵权产品,负责人,对方主体,来源链接,说明\r\n线上店铺销售疑似侵权产品,示例客户,淘宝,示例商品,管理者,示例店铺,https://example.com,每一行仅填写一种侵权产品"
    return Response(content=content.encode("utf-8"), media_type="text/csv; charset=utf-8", headers={"Content-Disposition": 'attachment; filename="clue-import-template.csv"'})


@app.post(f"{settings.api_prefix}/investigations/records", status_code=status.HTTP_201_CREATED)
async def create_investigation_record(body: RecordInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if body.module not in INVESTIGATION_RECORD_MODULES:
        raise HTTPException(status_code=422, detail="调查中心记录类型无效")
    if body.module in {"notary", "evidence"}:
        raise HTTPException(status_code=422, detail="公证和证据必须从线索专用办理入口创建")
    if await db.scalar(select(BusinessRecord.id).where(BusinessRecord.serial_no == body.serial_no)):
        raise HTTPException(status_code=409, detail="业务编号已存在")
    payload = body.model_dump()
    payload["status"] = INVESTIGATION_CREATE_STATUS_BY_MODULE[body.module]
    user = await db.scalar(select(User).where(User.username == identity["username"]))
    if not user:
        raise HTTPException(status_code=401, detail="当前用户不存在")
    if body.module == "investigation" and identity.get("role") != "admin" and not await _user_has_job_permission(user, "调查任务发布", db):
        raise HTTPException(status_code=403, detail="当前岗位没有发布调查任务权限")
    if body.module == "clue":
        if identity.get("role") != "admin" and not await _user_has_job_permission(user, "线索提交", db):
            raise HTTPException(status_code=403, detail="当前岗位没有创建调查线索权限")
        source_task_id = int((payload.get("data") or {}).get("source_task_id") or 0)
        if not source_task_id:
            raise HTTPException(status_code=422, detail="创建线索必须关联已接收的调查任务")
        source_task = await _ensure_record_module(source_task_id, "task", identity, db)
        if not (source_task.data or {}).get("investigation_record_id"):
            raise HTTPException(status_code=422, detail="线索来源必须是调查中心任务")
        if identity.get("role") != "admin" and source_task.owner != identity["username"]:
            raise HTTPException(status_code=403, detail="只能在本人负责的调查任务下创建线索")
        payload["owner"] = source_task.owner
        payload["customer"] = source_task.customer
        payload["department"] = source_task.department
        payload["data"] = {**(payload.get("data") or {}), "source_task_id": source_task.id, "source_task_no": source_task.serial_no, "investigation_record_id": (source_task.data or {}).get("investigation_record_id"), "investigation_no": (source_task.data or {}).get("investigation_no"), "customer_review": bool((source_task.data or {}).get("customer_review"))}
    if identity.get("role") != "admin":
        payload["department"] = user.department
        if identity.get("role") == "user":
            payload["owner"] = user.username
    record = BusinessRecord(**payload)
    db.add(record)
    await db.flush()
    db.add(WorkflowEvent(record_id=record.id, action="创建调查中心记录", to_status=record.status, operator=identity["username"], comment=f"类型：{body.module}"))
    await db.commit()
    await db.refresh(record)
    return _record_dict(record)


@app.patch(f"{settings.api_prefix}/investigations/records/{{record_id}}")
async def update_investigation_record(record_id: int, body: RecordUpdate, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    record = await _ensure_record_visible(record_id, identity, db)
    if record.module not in INVESTIGATION_RECORD_MODULES:
        raise HTTPException(status_code=404, detail="调查中心记录不存在")
    await _require_record_owner_or_manager(record, identity, db)
    if record.module == "clue" and record.status not in {"草稿", "已驳回"}:
        raise HTTPException(status_code=409, detail="待审批、待客户审核及后续线索不可直接修改；请等待审核结果或走驳回后重提流程")
    changes = body.model_dump(exclude_unset=True)
    if changes.get("status") and changes["status"] != record.status:
        raise HTTPException(status_code=409, detail="调查中心状态必须通过专用审批或办理入口变更")
    if changes.get("owner") and changes["owner"] != record.owner:
        raise HTTPException(status_code=409, detail="调查员/负责人必须通过分配调查员入口修改")
    old_status = record.status
    for field in ("title", "customer", "description"):
        if field in changes and changes[field] is not None:
            setattr(record, field, changes[field])
    if "data" in changes and changes["data"] is not None:
        incoming_data = dict(changes["data"] or {})
        editable_data = {key: incoming_data[key] for key in INVESTIGATION_EDIT_DATA_FIELDS if key in incoming_data}
        record.data = {**(record.data or {}), **editable_data}
    db.add(WorkflowEvent(record_id=record.id, action="修改调查中心资料", from_status=old_status, to_status=record.status, operator=identity["username"], comment="通过调查中心专用入口修改基础资料"))
    await db.commit()
    await db.refresh(record)
    return _record_dict(record, await _allowed_field_keys(identity, db))


@app.post(f"{settings.api_prefix}/investigations/clues/import")
async def import_investigation_clues(file: UploadFile = File(...), identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if not (file.filename or "").lower().endswith(".csv"):
        raise HTTPException(status_code=422, detail="仅支持 UTF-8 CSV 文件")
    raw = await file.read()
    if len(raw) > 5 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="导入文件不能超过 5MB")
    try:
        content = raw.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise HTTPException(status_code=422, detail="CSV 必须使用 UTF-8 编码") from exc
    reader = csv.DictReader(io.StringIO(content))
    existing_rows = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "clue"))).all()
    existing = {(x.title.strip().casefold(), x.customer.strip().casefold(), str((x.data or {}).get("product", "")).strip().casefold()) for x in existing_rows}
    seen: set[tuple[str, str, str]] = set(); created = 0; errors: list[dict] = []
    for row_no, row in enumerate(reader, 2):
        title = (row.get("线索标题") or row.get("title") or "").strip()
        customer = (row.get("客户") or row.get("customer") or "").strip()
        platform = (row.get("调查平台") or row.get("platform") or "").strip()
        product = (row.get("侵权产品") or row.get("product") or "").strip()
        owner = (row.get("负责人") or row.get("owner") or identity["username"]).strip()
        if not title or not product:
            errors.append({"row": row_no, "error": "线索标题、侵权产品为必填项"}); continue
        if any(separator in product for separator in ["；", ";", "、", "\n"]):
            errors.append({"row": row_no, "error": "每行只能填写一种侵权产品，请拆分为多行"}); continue
        key = (title.casefold(), customer.casefold(), product.casefold())
        if key in existing or key in seen:
            errors.append({"row": row_no, "error": "相同客户、标题和侵权产品的线索重复"}); continue
        serial_no = f"XS{datetime.now():%Y%m%d%H%M%S}{row_no:04d}{uuid4().hex[:4].upper()}"
        item = BusinessRecord(
            module="clue", serial_no=serial_no, title=title, customer=customer,
            status="草稿", owner=owner or identity["username"], department="上海分所",
            description=(row.get("说明") or row.get("description") or "").strip(),
            data={"platform": platform, "product": product, "opponent": (row.get("对方主体") or row.get("opponent") or "").strip(), "source_url": (row.get("来源链接") or row.get("source_url") or "").strip(), "imported_at": datetime.now().isoformat(timespec="seconds")},
        )
        db.add(item); await db.flush()
        db.add(WorkflowEvent(record_id=item.id, action="批量导入线索", to_status="草稿", operator=identity["username"], comment=f"CSV 第 {row_no} 行"))
        seen.add(key); created += 1
    await db.commit()
    return {"created": created, "failed": len(errors), "errors": errors}


@app.get(f"{settings.api_prefix}/investigations/notaries/import-template")
async def notary_import_template(_: dict = Depends(current_identity)):
    content = "\ufeff来源线索编号,公证标题,负责人,审核截止日,公证书编号,签发日期,存放位置,实物已收,说明\r\nXS2026070015,线上侵权产品公证审核,管理者,2026-08-15,(2026)沪证字001号,2026-08-10,上海档案室A-01,是,来源线索必须已经存在"
    return Response(content=content.encode("utf-8"), media_type="text/csv; charset=utf-8", headers={"Content-Disposition": 'attachment; filename="notary-import-template.csv"'})


@app.post(f"{settings.api_prefix}/investigations/notaries/import", status_code=status.HTTP_201_CREATED)
async def import_investigation_notaries(file: UploadFile = File(...), identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if not (file.filename or "").lower().endswith(".csv"): raise HTTPException(status_code=422, detail="仅支持 UTF-8 CSV 文件")
    raw = await file.read()
    if len(raw) > 5 * 1024 * 1024: raise HTTPException(status_code=413, detail="导入文件不能超过 5MB")
    try: content = raw.decode("utf-8-sig")
    except UnicodeDecodeError as exc: raise HTTPException(status_code=422, detail="CSV 必须使用 UTF-8 编码") from exc
    clues = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "clue", *(await _record_scope_conditions(identity, db))))).all()
    clue_by_no = {item.serial_no.strip().casefold(): item for item in clues}
    existing_clue_ids = set((await db.scalars(select(BusinessRecord.data["clue_id"].as_integer()).where(BusinessRecord.module == "notary"))).all())
    existing_certificate_nos = {str(value).strip().casefold() for value in (await db.scalars(select(BusinessRecord.data["certificate_no"].as_string()).where(BusinessRecord.module == "notary"))).all() if value}
    seen: set[int] = set(); created_ids: list[int] = []; errors: list[dict] = []
    for row_no, row in enumerate(csv.DictReader(io.StringIO(content)), 2):
        clue_no = (row.get("来源线索编号") or row.get("clue_no") or "").strip()
        clue = clue_by_no.get(clue_no.casefold())
        if not clue: errors.append({"row": row_no, "error": "来源线索编号不存在"}); continue
        if clue.status not in {"已取证", "待公证", "已转案件"}: errors.append({"row": row_no, "error": "来源线索完成取证登记后才能导入公证信息"}); continue
        if clue.id in existing_clue_ids or clue.id in seen or (clue.data or {}).get("notary_record_id"):
            errors.append({"row": row_no, "error": "该线索已经存在公证记录"}); continue
        raw_due = (row.get("审核截止日") or row.get("review_due_date") or "").strip()
        try: due = date.fromisoformat(raw_due) if raw_due else date.today() + timedelta(days=30)
        except ValueError: errors.append({"row": row_no, "error": "审核截止日格式应为 YYYY-MM-DD"}); continue
        if due < date.today(): errors.append({"row": row_no, "error": "审核截止日不能早于今天"}); continue
        certificate_no = (row.get("公证书编号") or row.get("certificate_no") or "").strip(); raw_issued = (row.get("签发日期") or row.get("issued_date") or "").strip()
        if certificate_no and certificate_no.casefold() in existing_certificate_nos: errors.append({"row": row_no, "error": "公证书编号已经登记"}); continue
        try: issued_date = str(date.fromisoformat(raw_issued)) if raw_issued else ""
        except ValueError: errors.append({"row": row_no, "error": "签发日期格式应为 YYYY-MM-DD"}); continue
        clue_data = dict(clue.data or {}); serial_no = f"GZ{datetime.now():%Y%m%d%H%M%S}{row_no:04d}{uuid4().hex[:4].upper()}"
        item = BusinessRecord(module="notary", serial_no=serial_no, title=(row.get("公证标题") or row.get("title") or f"{clue.title}—公证审核").strip(), customer=clue.customer, status="等待材料", owner=(row.get("负责人") or row.get("owner") or clue.owner).strip(), department=clue.department, description=(row.get("说明") or row.get("description") or "批量导入公证记录").strip(), data={"clue_id": clue.id, "clue_no": clue.serial_no, "platform": clue_data.get("platform", ""), "product": clue_data.get("product", ""), "case_id": clue_data.get("converted_case_id"), "case_no": clue_data.get("converted_case_no", ""), "review_due_date": str(due), "certificate_no": certificate_no, "certificate_issued_date": issued_date, "certificate_storage_location": (row.get("存放位置") or row.get("storage_location") or "").strip(), "physical_received": (row.get("实物已收") or row.get("physical_received") or "").strip().casefold() in {"是", "true", "1", "yes"}, "imported_at": datetime.now().isoformat(timespec="seconds")})
        db.add(item); await db.flush(); previous = clue.status; clue.status = "已转案件" if clue_data.get("converted_case_id") else "待公证"; clue.data = {**clue_data, "notary": "等待公证书扫描件", "notary_record_id": item.id}
        db.add_all([WorkflowEvent(record_id=clue.id, action="批量导入公证信息", from_status=previous, to_status=clue.status, operator=identity["username"], comment=f"CSV 第 {row_no} 行，生成 {serial_no}"), WorkflowEvent(record_id=item.id, action="批量导入公证", to_status="等待材料", operator=identity["username"], comment=f"来源线索 {clue.serial_no}；等待公证书扫描件")])
        seen.add(clue.id); created_ids.append(item.id)
        if certificate_no: existing_certificate_nos.add(certificate_no.casefold())
    await db.commit()
    return {"created": len(created_ids), "created_ids": created_ids, "failed": len(errors), "errors": errors}


@app.post(f"{settings.api_prefix}/investigations/notaries/storage/import", status_code=status.HTTP_201_CREATED)
async def import_notary_storage(file: UploadFile = File(...), identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Update warehouse/certificate/invoice/case fields on existing notary records from UTF-8 CSV."""
    if not (file.filename or "").lower().endswith(".csv"):
        raise HTTPException(status_code=422, detail="仅支持 UTF-8 CSV 文件")
    raw = await file.read()
    if len(raw) > 5 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="导入文件不能超过 5MB")
    try:
        content = raw.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise HTTPException(status_code=422, detail="CSV 必须使用 UTF-8 编码") from exc
    records = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "notary", *(await _record_scope_conditions(identity, db))))).all()
    by_clue = {str((item.data or {}).get("clue_no", "")).strip().casefold(): item for item in records if (item.data or {}).get("clue_no")}
    by_certificate = {str((item.data or {}).get("certificate_no", "")).strip().casefold(): item for item in records if (item.data or {}).get("certificate_no")}
    updated = 0; errors: list[dict] = []; seen: set[int] = set(); imported_rows: list[dict] = []
    for row_no, row in enumerate(csv.DictReader(io.StringIO(content)), 2):
        clue_no = (row.get("线索号") or row.get("线索编号") or row.get("clue_no") or "").strip()
        certificate_no = (row.get("公证书号") or row.get("certificate_no") or "").strip()
        item = by_clue.get(clue_no.casefold()) if clue_no else None
        if item is None and certificate_no:
            item = by_certificate.get(certificate_no.casefold())
        if item is None:
            errors.append({"row": row_no, "error": "未找到对应公证记录（请提供有效线索号或公证书号）"}); continue
        if item.id in seen:
            errors.append({"row": row_no, "error": "同一公证记录在文件中重复"}); continue
        data = dict(item.data or {})
        values = {
            "certificate_no": certificate_no,
            "warehouse": (row.get("仓库") or row.get("仓库位置") or row.get("warehouse") or "").strip(),
            "invoice_no": (row.get("发票号") or row.get("invoice_no") or "").strip(),
            "case_no": (row.get("案号") or row.get("案件编号") or row.get("case_no") or "").strip(),
            "investigator": (row.get("调查员") or row.get("investigator") or "").strip(),
            "investigated_at": (row.get("调查时间") or row.get("investigated_at") or "").strip(),
            "infringement_method": (row.get("侵权方式") or row.get("infringement_method") or "").strip(),
            "shop_name": (row.get("店铺名称") or row.get("shop_name") or "").strip(),
            "address": (row.get("调查地址") or row.get("address") or "").strip(),
        }
        item.data = {**data, **{key: value for key, value in values.items() if value}}
        db.add(WorkflowEvent(record_id=item.id, action="导入公证仓库信息", from_status=item.status, to_status=item.status, operator=identity["username"], comment=f"CSV 第 {row_no} 行"))
        imported_rows.append({
            "id": item.id,
            "线索号": clue_no or str(data.get("clue_no", "")),
            "调查员": values["investigator"] or item.owner,
            "调查时间": values["investigated_at"] or str(data.get("investigated_at", "")),
            "侵权方式": values["infringement_method"] or str(data.get("infringement_method", "")),
            "店铺名称": values["shop_name"] or str(data.get("shop_name", "")),
            "调查地址": values["address"] or str(data.get("address", "")),
            "公证书号": values["certificate_no"] or str(data.get("certificate_no", "")),
            "仓库": values["warehouse"] or str(data.get("warehouse", "")),
            "发票号": values["invoice_no"] or str(data.get("invoice_no", "")),
            "案号": values["case_no"] or str(data.get("case_no", "")),
        })
        seen.add(item.id); updated += 1
    await db.commit()
    return {"created": updated, "updated": updated, "updated_ids": sorted(seen), "items": imported_rows, "failed": len(errors), "errors": errors}


async def _import_notary_named_file(file: UploadFile, match_field: str, reference_no: str, category: str, identity: dict, db: AsyncSession):
    filename = Path(file.filename or "").name
    if Path(filename).suffix.lower() != ".pdf":
        raise HTTPException(status_code=422, detail="仅支持 PDF 文件")
    lookup = reference_no.strip().casefold()
    if not lookup:
        label = "公证书号" if match_field == "certificate_no" else "发票号"
        raise HTTPException(status_code=422, detail=f"请填写{label}")
    records = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "notary", *(await _record_scope_conditions(identity, db))))).all()
    matches = [item for item in records if str((item.data or {}).get(match_field, "")).strip().casefold() == lookup]
    if not matches:
        label = "公证书号" if match_field == "certificate_no" else "发票号"
        raise HTTPException(status_code=422, detail=f"{label} {reference_no.strip()} 未匹配到任何公证记录")
    if len(matches) > 1:
        raise HTTPException(status_code=409, detail=f"{label}匹配到多条公证记录，请先清理重复编号")
    record = matches[0]
    duplicate = await db.scalar(select(FileAttachment).where(FileAttachment.record_id == record.id, FileAttachment.category == category, FileAttachment.original_name == filename))
    if duplicate:
        raise HTTPException(status_code=409, detail="该文件已经导入，请勿重复上传")
    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="单个文件不能超过 20MB")
    stored_name = f"{uuid4().hex}.pdf"; target = UPLOAD_ROOT / stored_name; target.write_bytes(content)
    item = FileAttachment(record_id=record.id, category=category, original_name=filename, stored_name=stored_name, content_type=file.content_type or "application/pdf", size=len(content), path=str(target), uploader=identity["username"], remark=f"按{match_field}显式编号匹配导入")
    db.add(item); db.add(WorkflowEvent(record_id=record.id, action=f"导入{category}", from_status=record.status, to_status=record.status, operator=identity["username"], comment=filename))
    await db.commit(); await db.refresh(item)
    return {"created": 1, "failed": 0, "errors": [], "record_id": record.id, "record_no": record.serial_no, "attachment": _attachment_dict(item, record)}


@app.post(f"{settings.api_prefix}/investigations/notaries/files/import", status_code=status.HTTP_201_CREATED)
async def import_notary_certificate_file(file: UploadFile = File(...), certificate_no: str = Form(...), identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    return await _import_notary_named_file(file, "certificate_no", certificate_no, "公证书扫描件", identity, db)


@app.post(f"{settings.api_prefix}/investigations/notaries/invoices/import", status_code=status.HTTP_201_CREATED)
async def import_notary_invoice_file(file: UploadFile = File(...), invoice_no: str = Form(...), identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    return await _import_notary_named_file(file, "invoice_no", invoice_no, "公证发票", identity, db)


@app.get(f"{settings.api_prefix}/investigations/notaries/files")
async def list_notary_files(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Return one query row per imported certificate or invoice attachment."""
    rows = (await db.execute(
        select(FileAttachment, BusinessRecord)
        .join(BusinessRecord, FileAttachment.record_id == BusinessRecord.id)
        .where(
            BusinessRecord.module == "notary",
            FileAttachment.category.in_({"公证书扫描件", "公证发票"}),
            *(await _record_scope_conditions(identity, db)),
        )
        .order_by(FileAttachment.created_at.desc(), FileAttachment.id.desc())
    )).all()
    items = []
    for attachment, record in rows:
        data = record.data or {}
        items.append({
            "id": attachment.id,
            "module": "notary_file",
            "serial_no": record.serial_no,
            "title": attachment.original_name,
            "customer": record.customer,
            "status": record.status,
            "owner": attachment.uploader or record.owner,
            "description": attachment.remark,
            "created_at": attachment.created_at.isoformat() if attachment.created_at else "",
            "updated_at": attachment.created_at.isoformat() if attachment.created_at else "",
            "data": {
                "attachment_id": attachment.id,
                "attachment_category": attachment.category,
                "invoice_no": data.get("invoice_no") or "",
                "certificate_no": data.get("certificate_no") or "",
                "collected_at": data.get("collected_at") or "",
                "notary_institution": data.get("notary_institution") or "",
                "clue_no": data.get("clue_no") or "",
                "case_no": data.get("case_no") or "",
                "document_type": attachment.category,
                "shop_name": data.get("shop_name") or record.title,
                "handler": attachment.uploader or record.owner,
                "imported_at": attachment.created_at.isoformat() if attachment.created_at else "",
            },
        })
    return {"items": items, "total": len(items)}


@app.post(f"{settings.api_prefix}/investigations/clues/{{clue_id}}/submit")
async def submit_investigation_clue(clue_id: int, body: TaskActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    clue = await _ensure_record_module(clue_id, "clue", identity, db); await _require_record_owner_or_manager(clue, identity, db)
    if clue.status not in {"草稿", "已驳回"}: raise HTTPException(status_code=409, detail="只有草稿或已驳回线索可以提交审批")
    data = clue.data or {}; missing = [name for name, value in {"客户": clue.customer, "调查平台": data.get("platform"), "侵权产品": data.get("product")}.items() if not value]
    if missing: raise HTTPException(status_code=422, detail="线索缺少：" + "、".join(missing))
    previous = clue.status; clue.status = "待审批"; clue.data = {**data, "submitted_at": datetime.now().isoformat(timespec="seconds"), "submitted_by": identity["username"], "review_comment": ""}
    db.add(WorkflowEvent(record_id=clue.id, action="提交线索审批", from_status=previous, to_status=clue.status, operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(clue); return _record_dict(clue)


@app.post(f"{settings.api_prefix}/investigations/clues/{{clue_id}}/review")
async def review_investigation_clue(clue_id: int, body: ClueReviewInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    reviewer = await db.scalar(select(User).where(User.username == identity["username"]))
    if not reviewer or not await _user_has_job_permission(reviewer, "线索审批", db):
        raise HTTPException(status_code=403, detail="当前账号没有线索审批岗位权限")
    clue = await _ensure_record_module(clue_id, "clue", identity, db)
    if clue.status != "待审批": raise HTTPException(status_code=409, detail="只有待审批线索可以审核")
    next_status = "待客户审核" if body.approved and bool((clue.data or {}).get("customer_review")) else "待取证" if body.approved else "已驳回"
    clue.status = next_status; clue.data = {**(clue.data or {}), "reviewer": identity["username"], "reviewed_at": datetime.now().isoformat(timespec="seconds"), "review_comment": body.comment, "rejection_reason": "" if body.approved else body.comment}
    db.add(WorkflowEvent(record_id=clue.id, action="线索内部审批通过，待客户审核" if next_status == "待客户审核" else "线索审批通过" if body.approved else "线索审批驳回", from_status="待审批", to_status=clue.status, operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(clue); return _record_dict(clue)


@app.post(f"{settings.api_prefix}/investigations/clues/{{clue_id}}/customer-review")
async def customer_review_investigation_clue(clue_id: int, body: ClueReviewInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    clue = await _ensure_record_module(clue_id, "clue", identity, db)
    if clue.status != "待客户审核":
        raise HTTPException(status_code=409, detail="只有待客户审核线索可以确认客户审核结果")
    if identity.get("role") != "admin":
        customer = await db.scalar(select(BusinessRecord).where(BusinessRecord.module == "customer", BusinessRecord.title == clue.customer))
        customer_managers = set((customer.data or {}).get("customer_managers") or ([customer.owner] if customer else []))
        if identity["username"] not in customer_managers:
            raise HTTPException(status_code=403, detail="仅该客户的客户管理人可以代录客户审核结果")
    clue.status = "待取证" if body.approved else "已驳回"
    clue.data = {**(clue.data or {}), "customer_reviewer": identity["username"], "customer_reviewed_at": datetime.now().isoformat(timespec="seconds"), "customer_review_comment": body.comment, "rejection_reason": "" if body.approved else body.comment}
    db.add(WorkflowEvent(record_id=clue.id, action="客户审核通过" if body.approved else "客户审核驳回", from_status="待客户审核", to_status=clue.status, operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(clue); return _record_dict(clue)


@app.post(f"{settings.api_prefix}/investigations/clues/{{clue_id}}/collect")
async def register_clue_collection(clue_id: int, body: ClueCollectionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    clue = await _ensure_record_module(clue_id, "clue", identity, db); await _require_record_owner_or_manager(clue, identity, db)
    if clue.status != "待取证": raise HTTPException(status_code=409, detail="只有审批通过的待取证线索可以登记取证")
    if body.collected_at > date.today(): raise HTTPException(status_code=422, detail="取证日期不能晚于今天")
    clue.status = "已取证"; clue.data = {**(clue.data or {}), "collected_at": str(body.collected_at), "notary_institution": body.notary_institution.strip(), "collected_by": identity["username"], "collection_registered_at": datetime.now().isoformat(timespec="seconds")}
    db.add(WorkflowEvent(record_id=clue.id, action="登记线索取证", from_status="待取证", to_status="已取证", operator=identity["username"], comment=f"取证日期 {body.collected_at}；公证机构 {body.notary_institution}。{body.comment}"))
    await db.commit(); await db.refresh(clue); return _record_dict(clue)


@app.post(f"{settings.api_prefix}/investigations/clues/{{clue_id}}/evidence", status_code=status.HTTP_201_CREATED)
async def create_evidence_from_clue(clue_id: int, body: EvidenceCreateInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    clue = await _ensure_record_module(clue_id, "clue", identity, db); await _require_record_owner_or_manager(clue, identity, db)
    if clue.status in {"草稿", "待审批", "已驳回"}: raise HTTPException(status_code=409, detail="线索审批通过后才能建立证据目录")
    user = await db.scalar(select(User).where(User.username == identity["username"])); owner = body.owner.strip() or identity["username"]
    if identity.get("role") == "user": owner = identity["username"]
    serial_no = f"ZJ{datetime.now():%Y%m%d%H%M%S%f}"; item = BusinessRecord(module="evidence", serial_no=serial_no, title=body.title.strip(), customer=clue.customer, status="待整理", owner=owner, department=user.department if user else clue.department, description=body.description, data={"source": body.source, "clue_id": clue.id, "clue_no": clue.serial_no, "platform": (clue.data or {}).get("platform", ""), "product": (clue.data or {}).get("product", "")})
    db.add(item); await db.flush(); evidence_ids = list((clue.data or {}).get("evidence_ids", [])); evidence_ids.append(item.id); clue.data = {**(clue.data or {}), "evidence_ids": list(dict.fromkeys(evidence_ids)), "evidence_count": len(set(evidence_ids))}
    db.add_all([WorkflowEvent(record_id=clue.id, action="建立证据目录", from_status=clue.status, to_status=clue.status, operator=identity["username"], comment=f"生成 {serial_no}"), WorkflowEvent(record_id=item.id, action="从线索建立证据", to_status="待整理", operator=identity["username"], comment=f"来源线索 {clue.serial_no}")])
    await db.commit(); await db.refresh(item); return _record_dict(item)


@app.post(f"{settings.api_prefix}/investigations/evidence/{{evidence_id}}/organize")
async def organize_evidence(evidence_id: int, body: TaskActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _ensure_record_module(evidence_id, "evidence", identity, db); await _require_record_owner_or_manager(item, identity, db)
    if item.status != "待整理": raise HTTPException(status_code=409, detail="只有待整理证据可以完成整理")
    item.status = "已整理"; item.data = {**(item.data or {}), "organized_at": datetime.now().isoformat(timespec="seconds"), "organized_by": identity["username"]}
    db.add(WorkflowEvent(record_id=item.id, action="完成证据整理", from_status="待整理", to_status=item.status, operator=identity["username"], comment=body.comment)); await db.commit(); await db.refresh(item); return _record_dict(item)


@app.post(f"{settings.api_prefix}/investigations/evidence/{{evidence_id}}/file")
async def file_evidence(evidence_id: int, body: TaskActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _ensure_record_module(evidence_id, "evidence", identity, db); await _require_record_owner_or_manager(item, identity, db)
    if item.status != "已整理": raise HTTPException(status_code=409, detail="证据整理完成后才能入卷")
    categories = set((await db.scalars(select(FileAttachment.category).where(FileAttachment.record_id == item.id))).all())
    if "证据目录" not in categories or not categories.intersection({"证据原件", "证据扫描件"}): raise HTTPException(status_code=422, detail="入卷前必须上传证据目录及证据原件或扫描件")
    item.status = "已入卷"; item.data = {**(item.data or {}), "filed_at": datetime.now().isoformat(timespec="seconds"), "filed_by": identity["username"], "file_categories": sorted(categories)}
    db.add(WorkflowEvent(record_id=item.id, action="证据入卷", from_status="已整理", to_status=item.status, operator=identity["username"], comment=body.comment)); await db.commit(); await db.refresh(item); return _record_dict(item)


@app.get(f"{settings.api_prefix}/notaries/lookup")
async def lookup_notary_by_certificate(certificate_no: str = Query(min_length=2, max_length=128), identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    lookup = certificate_no.strip().casefold()
    records = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "notary", *(await _record_scope_conditions(identity, db))))).all()
    matches = [item for item in records if str((item.data or {}).get("certificate_no", "")).strip().casefold() == lookup]
    if not matches:
        raise HTTPException(status_code=404, detail="未找到关联公证记录或当前账号无权查看")
    if len(matches) > 1:
        raise HTTPException(status_code=409, detail="公证书号匹配到多条记录，请先处理重复编号")
    return _record_dict(matches[0], await _allowed_field_keys(identity, db))


@app.post(f"{settings.api_prefix}/notaries/{{notary_id}}/certificate")
async def register_notary_certificate(notary_id: int, body: NotaryCertificateInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    notary = await _ensure_record_module(notary_id, "notary", identity, db); await _require_record_owner_or_manager(notary, identity, db)
    operator = await db.scalar(select(User).where(User.username == identity["username"]))
    if not operator or not await _user_has_job_permission(operator, "公证书号码登记", db):
        raise HTTPException(status_code=403, detail="当前账号没有公证书号码登记岗位权限")
    if notary.status not in {"等待材料", "待审核", "审核驳回", "审核通过"}: raise HTTPException(status_code=409, detail="当前公证记录不能登记公证书信息")
    duplicate = await db.scalar(select(BusinessRecord.id).where(BusinessRecord.module == "notary", BusinessRecord.id != notary.id, BusinessRecord.data["certificate_no"].as_string() == body.certificate_no.strip()))
    if duplicate: raise HTTPException(status_code=409, detail="公证书编号已经登记")
    notary.data = {**(notary.data or {}), "certificate_no": body.certificate_no.strip(), "certificate_issued_date": str(body.issued_date), "certificate_storage_location": body.storage_location.strip(), "physical_received": body.physical_received, "certificate_registered_at": datetime.now().isoformat(timespec="seconds"), "certificate_operator": identity["username"]}
    db.add(WorkflowEvent(record_id=notary.id, action="登记公证书", from_status=notary.status, to_status=notary.status, operator=identity["username"], comment=f"{body.certificate_no}；{body.storage_location}。{body.comment}")); await db.commit(); await db.refresh(notary); return _record_dict(notary)


@app.post(f"{settings.api_prefix}/investigations/{{record_id}}/assign")
async def assign_investigation_record(record_id: int, body: InvestigationAssignmentInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    record = await _ensure_record_visible(record_id, identity, db)
    if record.module not in {"investigation", "clue", "task"}:
        raise HTTPException(status_code=422, detail="仅调查授权、调查线索或调查任务可以分配调查员")
    if record.module == "task":
        _require_task_owner_or_initiator(record, identity, action="修改任务负责人")
    else:
        await _require_record_owner_or_manager(record, identity, db)
    if record.module == "clue" and record.status in {"已转案件"}:
        raise HTTPException(status_code=409, detail="已转案件线索不能更换调查员")
    if record.module == "task" and record.status in {"已完成", "待确认", "已验收", "已拒绝", "已撤回", "已停止", "已取消"}:
        raise HTTPException(status_code=409, detail="已结束任务不能更换调查员")
    if record.module == "investigation" and record.status in {"已完成", "已取消"}:
        raise HTTPException(status_code=409, detail="已结束调查授权不能更换调查员")
    previous_owner = record.owner
    record.owner = await _active_task_username(body.investigator, db, field_name="调查员")
    record.data = {**(record.data or {}), "investigator": record.owner, "assigner": identity["username"], "assigned_at": datetime.now().isoformat(timespec="seconds"), "assigned_by": identity["username"]}
    if record.module == "investigation" and record.status == "待分配": record.status = "进行中"
    assignment_event = WorkflowEvent(record_id=record.id, action="分配调查员", from_status=record.status, to_status=record.status, operator=identity["username"], comment=f"{previous_owner} → {record.owner}。{body.comment}")
    if record.module == "task":
        await _add_task_message_notifications(record, assignment_event, db, content="任务负责人已修改.")
    else:
        db.add(assignment_event)
    await db.commit(); await db.refresh(record)
    return _record_dict(record) if record.module != "task" else _task_dict(record)


@app.post(f"{settings.api_prefix}/investigations/clues/{{record_id}}/fee-application", status_code=status.HTTP_201_CREATED)
async def create_investigation_fee_application(record_id: int, body: InvestigationFeeInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    clue = await _ensure_record_module(record_id, "clue", identity, db); await _require_record_owner_or_manager(clue, identity, db)
    if clue.status not in {"已取证", "待公证", "已转案件"}:
        raise HTTPException(status_code=409, detail="线索完成取证后才能申请费用")
    if (clue.data or {}).get("fee_application_id"):
        raise HTTPException(status_code=409, detail="该线索已经提交过费用申请")
    serial_no = f"FY{datetime.now():%Y%m%d%H%M%S%f}"; fee = BusinessRecord(module="finance", serial_no=serial_no, title=f"调查费用—{clue.title}", customer=clue.customer, status="草稿", owner=identity["username"], department=clue.department, description=body.description, data={"fee_type": body.fee_type, "amount": body.amount, "clue_id": clue.id, "clue_no": clue.serial_no, "investigator": clue.owner, "source": "调查线索费用申请"})
    db.add(fee); await db.flush(); clue.data = {**(clue.data or {}), "fee_application_id": fee.id, "fee_no": serial_no, "fee_amount": body.amount, "fee_type": body.fee_type}
    db.add_all([WorkflowEvent(record_id=clue.id, action="申请调查费用", from_status=clue.status, to_status=clue.status, operator=identity["username"], comment=f"{serial_no}，{body.fee_type}，{body.amount:.2f}"), WorkflowEvent(record_id=fee.id, action="由调查线索生成费用申请", to_status="草稿", operator=identity["username"], comment=clue.serial_no)])
    await db.commit(); await db.refresh(fee); await db.refresh(clue)
    return {"fee": _record_dict(fee), "clue": _record_dict(clue)}


@app.post(f"{settings.api_prefix}/investigations/batch-delete")
async def batch_delete_investigation_records(body: InvestigationBatchDeleteInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    records = (await db.scalars(select(BusinessRecord).where(BusinessRecord.id.in_(set(body.record_ids)), BusinessRecord.module.in_(["investigation", "clue", "task"]), *(await _record_scope_conditions(identity, db))))).all()
    found = {item.id: item for item in records}; errors: list[dict] = []; deleted = 0; paths: list[Path] = []
    manager = identity.get("role") in {"admin", "manager"}
    for record_id in dict.fromkeys(body.record_ids):
        record = found.get(record_id)
        if not record:
            errors.append({"record_id": record_id, "error": "记录不存在或无权访问"}); continue
        if record.module == "investigation" and identity.get("role") != "admin":
            errors.append({"record_id": record_id, "record_no": record.serial_no, "error": "仅管理员可以删除调查任务"}); continue
        if record.module == "task" and identity.get("role") != "admin" and record.owner != identity["username"] and (record.data or {}).get("initiator") != identity["username"]:
            errors.append({"record_id": record_id, "record_no": record.serial_no, "error": "只能删除本人负责或发起的任务"}); continue
        if record.module not in {"task", "investigation"} and not manager and record.owner != identity["username"]:
            errors.append({"record_id": record_id, "record_no": record.serial_no, "error": "只能删除本人负责的记录"}); continue
        allowed = record.status in ({"草稿", "已驳回"} if record.module == "clue" else {"待接收", "未开始", "已驳回"} if record.module == "task" else {"待分配"})
        if not allowed:
            errors.append({"record_id": record_id, "record_no": record.serial_no, "error": f"当前状态“{record.status}”不允许删除"}); continue
        if record.module == "clue":
            related = await db.scalar(select(BusinessRecord.id).where(BusinessRecord.module.in_(["notary", "case"]), BusinessRecord.data["clue_id"].as_integer() == record.id))
            if related:
                errors.append({"record_id": record_id, "record_no": record.serial_no, "error": "已关联公证或案件，不能删除"}); continue
        elif record.module == "task":
            child = await db.scalar(select(BusinessRecord.id).where(BusinessRecord.module == "task", BusinessRecord.data["parent_task_id"].as_integer() == record.id))
            if child:
                errors.append({"record_id": record_id, "record_no": record.serial_no, "error": "任务存在子任务，不能删除"}); continue
        else:
            child = await db.scalar(select(BusinessRecord.id).where(BusinessRecord.module == "task", BusinessRecord.data["investigation_record_id"].as_integer() == record.id))
            if child:
                errors.append({"record_id": record_id, "record_no": record.serial_no, "error": "调查任务存在子任务，不能删除"}); continue
        attachments = (await db.scalars(select(FileAttachment).where(FileAttachment.record_id == record.id))).all()
        for attachment in attachments:
            paths.append(Path(attachment.path)); await db.delete(attachment)
        if record.module == "task":
            await _delete_task_notifications(record.id, db)
        await db.execute(delete(WorkflowEvent).where(WorkflowEvent.record_id == record.id)); await db.delete(record); deleted += 1
    await db.commit()
    for path in paths:
        if path.is_file() and UPLOAD_ROOT.resolve() in path.resolve().parents: path.unlink(missing_ok=True)
    return {"deleted": deleted, "failed": len(errors), "errors": errors}


@app.get(f"{settings.api_prefix}/investigations/{{record_id}}/tasks")
async def list_investigation_tasks(record_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    source = await _ensure_record_visible(record_id, identity, db)
    if source.module not in INVESTIGATION_MATERIAL_CATEGORIES: raise HTTPException(status_code=404, detail="调查业务记录不存在")
    tasks = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "task", BusinessRecord.data["investigation_record_id"].as_integer() == source.id, *(await _record_scope_conditions(identity, db))).order_by(BusinessRecord.created_at, BusinessRecord.id))).all()
    return {"record": _record_dict(source), "items": [_task_dict(item) for item in tasks], "total": len(tasks)}


@app.post(f"{settings.api_prefix}/investigations/{{record_id}}/close")
async def close_investigation(record_id: int, body: TaskActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    investigation = await _ensure_record_module(record_id, "investigation", identity, db)
    await _require_record_owner_or_manager(investigation, identity, db)
    if investigation.status in {"已完成", "已取消"}: raise HTTPException(status_code=409, detail="调查任务已经结束")
    tasks = list((await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "task", BusinessRecord.data["investigation_record_id"].as_integer() == investigation.id))).all())
    active_tasks = [item for item in tasks if item.status not in {"已完成", "已验收", "已拒绝", "已撤回", "已停止", "已取消"}]
    if active_tasks: raise HTTPException(status_code=409, detail=f"仍有 {len(active_tasks)} 个调查子任务未办结")
    task_ids = [item.id for item in tasks]
    clues = list((await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "clue", BusinessRecord.data["source_task_id"].as_integer().in_(task_ids)))).all()) if task_ids else []
    active_clues = [item for item in clues if item.status not in {"已转案件", "已驳回"}]
    if active_clues: raise HTTPException(status_code=409, detail=f"仍有 {len(active_clues)} 条调查线索未转案或未驳回")
    previous = investigation.status; investigation.status = "已完成"
    report_content = "\n".join([
        f"调查任务：{investigation.serial_no}｜{investigation.title}", f"客户：{investigation.customer}",
        f"负责人：{investigation.owner}", f"调查子任务：{len(tasks)} 个", f"调查线索：{len(clues)} 条",
        f"已转案件线索：{sum(1 for item in clues if item.status == '已转案件')} 条", f"驳回线索：{sum(1 for item in clues if item.status == '已驳回')} 条",
        f"关闭说明：{body.comment.strip() or '全部调查事项已经办结'}",
    ])
    content = _docx_bytes(f"调查任务报告-{investigation.serial_no}", report_content)
    stored_name = f"{uuid4().hex}.docx"; path = UPLOAD_ROOT / stored_name; path.write_bytes(content)
    attachment = FileAttachment(record_id=investigation.id, category="调查任务报告", original_name=f"调查任务报告-{investigation.serial_no}.docx", stored_name=stored_name, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", size=len(content), path=str(path), uploader=identity["username"], remark="调查任务关闭时由系统生成")
    db.add(attachment); await db.flush()
    investigation.data = {**(investigation.data or {}), "closed_at": datetime.now().isoformat(timespec="seconds"), "closed_by": identity["username"], "close_comment": body.comment.strip(), "report_attachment_id": attachment.id}
    db.add(WorkflowEvent(record_id=investigation.id, action="关闭调查任务并生成报告", from_status=previous, to_status="已完成", operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(investigation); await db.refresh(attachment)
    return {"record": _record_dict(investigation), "report": _attachment_dict(attachment, investigation)}


@app.post(f"{settings.api_prefix}/investigations/{{record_id}}/tasks", status_code=status.HTTP_201_CREATED)
async def create_investigation_task(record_id: int, body: InvestigationTaskInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    source = await _ensure_record_visible(record_id, identity, db)
    if source.module not in INVESTIGATION_MATERIAL_CATEGORIES: raise HTTPException(status_code=404, detail="调查业务记录不存在")
    await _require_record_owner_or_manager(source, identity, db)
    _validate_task_deadline(body.deadline)
    parent = None
    if body.parent_task_id:
        parent = await _ensure_record_module(body.parent_task_id, "task", identity, db)
        if int((parent.data or {}).get("investigation_record_id") or 0) != source.id: raise HTTPException(status_code=409, detail="父任务不属于当前调查事项")
    user = await db.scalar(select(User).where(User.username == identity["username"])); owner = body.owner.strip()
    if identity.get("role") == "user": owner = identity["username"]
    owner = await _active_task_username(owner, db, field_name="负责人")
    serial_no = f"RW{datetime.now():%Y%m%d%H%M%S%f}"; task = BusinessRecord(module="task", serial_no=serial_no, title=body.title.strip(), customer=source.customer, status="待接收", owner=owner, department=user.department if user else source.department, description=body.description, data={"deadline": str(body.deadline), "priority": body.priority, "source": "调查任务", "initiator": identity["username"], "collaborators": [], "case_no": "", "investigation_record_id": source.id, "investigation_no": source.serial_no, "investigation_module": source.module, "customer_review": bool((source.data or {}).get("customer_review")), "parent_task_id": parent.id if parent else None, "parent_task_no": parent.serial_no if parent else ""})
    db.add(task); await db.flush()
    db.add(WorkflowEvent(record_id=source.id, action="创建调查任务", from_status=source.status, to_status=source.status, operator=identity["username"], comment=f"{serial_no}：{body.title}"))
    await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action="创建调查任务", to_status="待接收", operator=identity["username"], comment=f"来源 {source.serial_no}" + (f"；父任务 {parent.serial_no}" if parent else "")), db, content="任务已分派.")
    await db.commit(); await db.refresh(task); return _task_dict(task)


@app.post(f"{settings.api_prefix}/investigations/clues/batch-cases", status_code=status.HTTP_201_CREATED)
async def batch_create_cases_from_clues(body: BatchClueCaseInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    contract = await _ensure_record_visible(body.contract_record_id, identity, db)
    if contract.module != "contract": raise HTTPException(status_code=422, detail="关联记录不是合同")
    if contract.status in {"草稿", "审批中", "已拒绝", "已撤回", "已作废"}: raise HTTPException(status_code=409, detail="合同审批通过后才能批量转案件")
    clues = {item.id: item for item in (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "clue", BusinessRecord.id.in_(set(body.clue_ids)), *(await _record_scope_conditions(identity, db))))).all()}
    contract_data = contract.data or {}; created_ids: list[int] = []; errors: list[dict] = []
    for clue_id in dict.fromkeys(body.clue_ids):
        clue = clues.get(clue_id)
        if not clue: errors.append({"clue_id": clue_id, "error": "线索不存在"}); continue
        clue_data = dict(clue.data or {})
        if clue_data.get("converted_case_id") or clue.status == "已转案件": errors.append({"clue_id": clue_id, "clue_no": clue.serial_no, "error": "线索已经转为案件"}); continue
        if clue.status not in {"已取证", "待公证"}: errors.append({"clue_id": clue_id, "clue_no": clue.serial_no, "error": "线索完成取证登记后才能转案件"}); continue
        if clue.customer.strip() != contract.customer.strip(): errors.append({"clue_id": clue_id, "clue_no": clue.serial_no, "error": "线索客户与合同客户不一致"}); continue
        serial_no = await _next_case_serial(body.case_type, db)
        case_record = BusinessRecord(module="case", serial_no=serial_no, title=clue.title, customer=contract.customer, status="等待公证书", owner=clue.owner, department=contract.department, description=f"由已取证线索 {clue.serial_no} 批量转案", data={"contract_id": contract.id, "contract_no": contract.serial_no, "external_contract_no": contract_data.get("external_contract_no", ""), "external_contract_numbers": contract_data.get("external_contract_numbers", []), "contract_title": contract.title, "clue_id": clue.id, "clue_no": clue.serial_no, "notary_id": clue_data.get("notary_record_id"), "case_type": body.case_type, "court": body.court, "opponent": clue_data.get("opponent", ""), "product": clue_data.get("product", ""), "batch_converted": True, "case_creation_step": "completed", "case_creation_approval_status": "自动通过", "case_creation_approved_by": "system"})
        db.add(case_record); await db.flush(); previous = clue.status; clue.status = "已转案件"; clue.data = {**clue_data, "converted_case_id": case_record.id, "converted_case_no": serial_no}
        notary = await db.get(BusinessRecord, int(clue_data.get("notary_record_id") or 0)) if clue_data.get("notary_record_id") else None
        if notary: notary.data = {**(notary.data or {}), "case_id": case_record.id, "case_no": serial_no}
        db.add_all([WorkflowEvent(record_id=clue.id, action="已取证线索批量转案件", from_status=previous, to_status="已转案件", operator=identity["username"], comment=f"关联合同 {contract.serial_no}，生成等待公证书案件 {serial_no}"), WorkflowEvent(record_id=case_record.id, action="线索批量转案", to_status="等待公证书", operator=identity["username"], comment=f"来源线索 {clue.serial_no} / 合同 {contract.serial_no}")])
        await _ensure_case_fixed_tasks(case_record, db, operator="system")
        created_ids.append(case_record.id)
    await db.commit()
    return {"created": len(created_ids), "created_ids": created_ids, "failed": len(errors), "errors": errors}


@app.get(f"{settings.api_prefix}/investigations/{{record_id}}/materials")
async def list_investigation_materials(record_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    record = await _ensure_record_visible(record_id, identity, db)
    if not record or record.module not in INVESTIGATION_MATERIAL_CATEGORIES:
        raise HTTPException(status_code=404, detail="调查、公证或证据记录不存在")
    items = (await db.scalars(select(FileAttachment).where(FileAttachment.record_id == record.id).order_by(FileAttachment.created_at.desc(), FileAttachment.id.desc()))).all()
    return {"record": _record_dict(record), "allowed_categories": INVESTIGATION_MATERIAL_CATEGORIES[record.module], "items": [_attachment_dict(x, record) for x in items], "total": len(items)}


@app.post(f"{settings.api_prefix}/investigations/{{clue_id}}/notary", status_code=status.HTTP_201_CREATED)
async def create_notary_from_clue(clue_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    clue = await _ensure_record_module(clue_id, "clue", identity, db)
    await _require_record_owner_or_manager(clue, identity, db)
    if clue.status not in {"已取证", "已转案件"}: raise HTTPException(status_code=409, detail="线索完成取证登记后才能建立公证记录")
    clue_data = clue.data or {}
    if clue_data.get("notary_record_id"):
        existing = await db.get(BusinessRecord, int(clue_data["notary_record_id"]))
        if existing:
            raise HTTPException(status_code=409, detail=f"该线索已生成公证记录 {existing.serial_no}")
    serial = f"GZ{datetime.now():%Y%m%d%H%M%S%f}"
    notary = BusinessRecord(
        module="notary", serial_no=serial, title=f"{clue.title}—公证审核",
        customer=clue.customer, status="等待材料", owner=clue.owner, department=clue.department,
        description="由调查线索自动生成",
        data={"clue_id": clue.id, "clue_no": clue.serial_no, "platform": clue_data.get("platform", ""), "product": clue_data.get("product", ""), "case_id": clue_data.get("converted_case_id"), "case_no": clue_data.get("converted_case_no", "")},
    )
    db.add(notary)
    await db.flush()
    previous_status = clue.status
    clue.status = "已转案件" if clue_data.get("converted_case_id") else "待公证"
    clue.data = {**clue_data, "notary": "等待公证书扫描件", "notary_record_id": notary.id}
    db.add_all([
        WorkflowEvent(record_id=clue.id, action="建立公证记录", from_status=previous_status, to_status=clue.status, operator=identity["username"], comment=f"生成公证记录 {serial}"),
        WorkflowEvent(record_id=notary.id, action="创建公证材料记录", to_status="等待材料", operator=identity["username"], comment=f"来源线索 {clue.serial_no}，等待上传公证书扫描件"),
    ])
    await db.commit()
    await db.refresh(notary)
    return _record_dict(notary)


async def _convert_notary_to_case(
    notary: BusinessRecord,
    clue: BusinessRecord,
    db: AsyncSession,
    *,
    operator: str,
    comment: str,
    case_type: str = "民事案件",
    court: str = "",
    automatic: bool = False,
) -> BusinessRecord:
    """将公证审核记录转为新案；人工审核和 30 日超期规则共用同一闭环。"""
    clue_data = clue.data or {}
    case_record = await db.get(BusinessRecord, int(clue_data.get("converted_case_id") or 0))
    if not case_record or case_record.module != "case":
        raise HTTPException(status_code=409, detail="请先从已取证线索批量生成“等待公证书”案件")
    if case_record.status not in {"等待公证书", "等待审核公证书"}:
        raise HTTPException(status_code=409, detail=f"案件 {case_record.serial_no} 当前阶段不允许公证审核")
    case_serial = case_record.serial_no
    previous_case_status = case_record.status
    case_record.status = "新案待分配"
    case_record.data = {**(case_record.data or {}), "notary_id": notary.id, "notary_no": notary.serial_no, "case_type": case_type or (case_record.data or {}).get("case_type", "民事案件"), "court": court or (case_record.data or {}).get("court", ""), "notary_review_automatic": automatic, "case_creation_step": "completed", "case_creation_approval_status": "自动通过", "case_creation_approved_by": "system"}
    action = "公证审核超期自动转案" if automatic else "公证审核通过"
    notary.status = "审核通过"
    notary.data = {
        **(notary.data or {}), "case_id": case_record.id, "case_no": case_serial,
        "auto_reviewed": automatic,
        **({"auto_reviewed_at": str(date.today())} if automatic else {}),
    }
    clue.status = "已转案件"
    clue.data = {
        **clue_data, "notary": "超期自动通过" if automatic else "审核通过",
        "converted_case_id": case_record.id, "converted_case_no": case_serial,
    }
    db.add_all([
        WorkflowEvent(record_id=notary.id, action=action, from_status="待审核", to_status="审核通过", operator=operator, comment=comment),
        WorkflowEvent(record_id=clue.id, action="自动转案件", from_status="待公证", to_status="已转案件", operator=operator, comment=f"生成案件 {case_serial}"),
        WorkflowEvent(record_id=case_record.id, action="公证审核完成", from_status=previous_case_status, to_status="新案待分配", operator=operator, comment=f"来源线索 {clue.serial_no} / 公证 {notary.serial_no}"),
    ])
    return case_record


async def _apply_notary_auto_conversion(db: AsyncSession) -> bool:
    """公证审核超过 30 日仍未处理时，自动进入“新案待分配”。"""
    notaries = (await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "notary", BusinessRecord.status == "待审核",
    ))).all()
    changed = False
    for notary in notaries:
        raw_due = (notary.data or {}).get("review_due_date")
        try:
            overdue = bool(raw_due) and date.fromisoformat(str(raw_due)) < date.today()
        except ValueError:
            overdue = False
        if not overdue:
            continue
        clue_id = int((notary.data or {}).get("clue_id", 0))
        clue = await db.get(BusinessRecord, clue_id)
        if not clue:
            continue
        await _convert_notary_to_case(
            notary, clue, db, operator="system", automatic=True,
            comment="公证书审核已超过 30 日，系统自动将案件进入“新案待分配”",
        )
        changed = True
    if changed:
        await db.commit()
    return changed


@app.post(f"{settings.api_prefix}/notaries/{{notary_id}}/review")
async def review_notary(notary_id: int, body: NotaryReviewInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    reviewer = await db.scalar(select(User).where(User.username == identity["username"]))
    if not reviewer or not await _user_has_job_permission(reviewer, "公证审核", db):
        raise HTTPException(status_code=403, detail="当前账号没有公证审核岗位权限")
    notary = await _ensure_record_module(notary_id, "notary", identity, db)
    if notary.status != "待审核":
        raise HTTPException(status_code=409, detail="该公证记录已完成审核")
    clue_id = int((notary.data or {}).get("clue_id", 0))
    clue = await db.get(BusinessRecord, clue_id)
    if not clue:
        raise HTTPException(status_code=409, detail="关联线索不存在")
    if not body.approved:
        notary.status = "审核驳回"
        clue.data = {**(clue.data or {}), "notary": "审核驳回，等待补正扫描件"}
        case_record = await db.get(BusinessRecord, int((clue.data or {}).get("converted_case_id") or 0))
        if case_record and case_record.status == "等待审核公证书": case_record.status = "等待公证书"
        db.add_all([
            WorkflowEvent(record_id=notary.id, action="公证审核驳回", from_status="待审核", to_status="审核驳回", operator=identity["username"], comment=body.comment),
            WorkflowEvent(record_id=clue.id, action="公证材料退回补正", from_status=clue.status, to_status=clue.status, operator=identity["username"], comment=body.comment),
        ])
        if case_record: db.add(WorkflowEvent(record_id=case_record.id, action="公证材料退回补正", from_status="等待审核公证书", to_status=case_record.status, operator=identity["username"], comment=body.comment))
        await db.commit()
        await db.refresh(notary)
        return {"notary": _record_dict(notary), "case": None}
    case_record = await _convert_notary_to_case(
        notary, clue, db, operator=identity["username"], comment=body.comment,
        case_type=body.case_type, court=body.court,
    )
    await db.commit()
    await db.refresh(notary)
    await db.refresh(case_record)
    return {"notary": _record_dict(notary), "case": _record_dict(case_record)}


async def _apply_task_auto_completion(db: AsyncSession) -> bool:
    """交接后未重新开始的任务，满 5 天自动完成。"""
    tasks = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "task"))).all()
    changed = False
    for task in tasks:
        data = task.data or {}
        confirm_at = data.get("completion_auto_confirm_at")
        if task.status in {"待确认", "已完成"} and confirm_at:
            try:
                should_confirm = date.fromisoformat(str(confirm_at)) <= date.today()
            except ValueError:
                should_confirm = False
            if should_confirm:
                previous = task.status
                task.status = "已验收"
                task.data = {**data, "auto_confirmed": True, "auto_confirmed_at": str(date.today())}
                await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action="任务完成自动验收", from_status=previous, to_status="已验收", operator="system", comment="负责人提交完成后满 5 日，发起人未重启，系统自动验收"), db, content="任务已确认完成.")
                await _advance_case_from_fixed_task(task, db, operator="system")
                changed = True
                continue
        auto_at = data.get("handoff_auto_complete_at")
        if not auto_at or data.get("handoff_restarted") or task.status != "待接收":
            continue
        try:
            should_complete = date.fromisoformat(str(auto_at)) <= date.today()
        except ValueError:
            should_complete = False
        if should_complete:
            previous = task.status
            task.status = "已完成"
            task.data = {**data, "auto_completed": True, "auto_completed_at": str(date.today())}
            await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action="交接任务自动完成", from_status=previous, to_status="已完成", operator="system", comment="交接满 5 天且未重新开始，系统自动完成"), db, content="任务已自动完成.")
            changed = True
    if changed:
        await db.commit()
    return changed


async def _apply_task_overdue_performance(db: AsyncSession) -> bool:
    """Persist overdue facts so performance reports do not depend on a page being open."""
    tasks = list((await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "task"))).all())
    changed = False
    terminal = {"已完成", "待确认", "已验收", "已拒绝", "已撤回", "已停止", "已取消"}
    for task in tasks:
        if task.status in terminal:
            continue
        data = task.data or {}
        try:
            overdue_days = (date.today() - date.fromisoformat(str(data.get("deadline") or ""))).days
        except ValueError:
            continue
        if overdue_days <= 0:
            continue
        previous = data.get("performance_impact") or {}
        impact = {
            "overdue": True, "overdue_days": overdue_days, "penalty_points": overdue_days,
            "recorded_for": str(date.today()), "responsible_user": task.owner,
        }
        if previous == impact:
            continue
        task.data = {**data, "performance_impact": impact}
        if not previous:
            db.add(WorkflowEvent(record_id=task.id, action="记录任务超期绩效", from_status=task.status, to_status=task.status, operator="system", comment=f"任务超期 {overdue_days} 天，记录绩效影响 {overdue_days} 分"))
        changed = True
    if changed:
        await db.commit()
    return changed


async def _apply_hearing_sms_reminders(db: AsyncSession) -> bool:
    """Create auditable 3-day/1-day hearing SMS records and send through a configured webhook."""
    today = date.today(); changed = False
    schedules = list((await db.scalars(select(HearingSchedule).where(HearingSchedule.status == "已排期", HearingSchedule.hearing_date.in_([today + timedelta(days=1), today + timedelta(days=3)])))).all())
    for hearing in schedules:
        days = (hearing.hearing_date - today).days
        duplicate = await db.scalar(select(BusinessRecord.id).where(BusinessRecord.module == "sms", BusinessRecord.data["hearing_id"].as_integer() == hearing.id, BusinessRecord.data["remind_days"].as_integer() == days))
        if duplicate: continue
        case_record = await db.get(BusinessRecord, hearing.case_record_id)
        if not case_record: continue
        names = list(dict.fromkeys(value for value in [hearing.hearing_lawyer, *((case_record.data or {}).get("handling_lawyers") or []), (case_record.data or {}).get("assistant", "")] if value))
        users = list((await db.scalars(select(User).where(User.is_active.is_(True), or_(User.username.in_(names), User.display_name.in_(names))))).all()) if names else []
        phones = list(dict.fromkeys(str((user.profile or {}).get("phone") or "").strip() for user in users if str((user.profile or {}).get("phone") or "").strip()))
        content = f"开庭提醒：案件 {case_record.serial_no} 将于 {hearing.hearing_date} {hearing.hearing_time} 在 {hearing.court}{(' ' + hearing.courtroom) if hearing.courtroom else ''} 开庭。"
        sms_status = "待配置短信通道" if not settings.sms_webhook_url else "待发送"
        if not phones: sms_status = "待补充手机号"
        response_excerpt = ""
        if phones and settings.sms_webhook_url:
            try:
                headers = {"Authorization": f"Bearer {settings.sms_webhook_token}"} if settings.sms_webhook_token else {}
                async with httpx.AsyncClient(timeout=10) as client:
                    response = await client.post(settings.sms_webhook_url, json={"phones": phones, "content": content, "case_no": case_record.serial_no, "hearing_id": hearing.id}, headers=headers)
                    response.raise_for_status(); response_excerpt = response.text[:500]
                sms_status = "已发送"
            except Exception as exc:
                sms_status = "发送失败"; response_excerpt = str(exc)[:500]
        sms = BusinessRecord(module="sms", serial_no=f"DX{datetime.now():%Y%m%d%H%M%S%f}", title=f"开庭短信提醒—{case_record.serial_no}", customer=case_record.customer, status=sms_status, owner="system", department=case_record.department, description=content, data={"hearing_id": hearing.id, "case_id": case_record.id, "case_no": case_record.serial_no, "remind_days": days, "phones": phones, "recipient_users": [user.username for user in users], "provider_response": response_excerpt})
        db.add(sms); await db.flush()
        db.add(WorkflowEvent(record_id=sms.id, action="生成开庭短信提醒", to_status=sms_status, operator="system", comment=f"开庭前 {days} 天；收件手机号 {len(phones)} 个"))
        for user in users:
            db.add(Notification(source_key=f"hearing-sms-{hearing.id}-{days}-{user.username}", source_type="case", source_id=case_record.id, sender="system", recipient=user.username, notification_type="系统通知", title=f"开庭短信：{sms_status}", content=content, level="info" if sms_status == "已发送" else "warning"))
        changed = True
    if changed: await db.commit()
    return changed


async def _business_rule_loop() -> None:
    """本地部署时持续执行不依赖用户打开页面的期限规则。"""
    while True:
        async with SessionLocal() as db:
            try:
                await _apply_notary_auto_conversion(db)
                await _apply_task_auto_completion(db)
                await _apply_task_overdue_performance(db)
                await _apply_hearing_sms_reminders(db)
            except Exception:
                await db.rollback()
        await asyncio.sleep(3600)


@app.get(f"{settings.api_prefix}/tasks")
async def list_tasks(
    keyword: str = "", status_filter: str = "", reminder_only: bool = False, scope: str = "default",
    relation: str = Query("", pattern="^(|initiated|owned|collaborating)$"), statuses: str = "",
    priority: str = "", serial_no: str = "", title: str = "", description: str = "",
    initiator: str = "", case_no: str = "", source: str = "", owner: str = "",
    plaintiff: str = "", defendant: str = "",
    created_from: date | None = None, created_to: date | None = None,
    deadline_from: date | None = None, deadline_to: date | None = None,
    sort_by: str = Query("created_at", pattern="^(created_at|deadline|days_remaining|updated_at)$"),
    sort_order: str = Query("desc", pattern="^(asc|desc)$"),
    page: int = Query(1, ge=1), page_size: int | None = Query(None, ge=1, le=200),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    await _apply_task_auto_completion(db)
    await _apply_task_overdue_performance(db)
    if scope not in {"default", "mine", "department", "company"}: raise HTTPException(status_code=422, detail="无效的任务范围")
    if created_from and created_to and created_from > created_to: raise HTTPException(status_code=422, detail="发起开始日期不能晚于结束日期")
    if deadline_from and deadline_to and deadline_from > deadline_to: raise HTTPException(status_code=422, detail="截止开始日期不能晚于结束日期")
    tasks = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "task").order_by(BusinessRecord.created_at.desc()))).all()
    if identity.get("role") != "admin" and scope in {"mine", "default"}:
        username = identity["username"]
        tasks = [task for task in tasks if task.owner == username or (task.data or {}).get("initiator") == username or username in (task.data or {}).get("collaborators", [])]
    elif scope == "department":
        if identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以查看部门任务")
        if identity.get("role") == "manager":
            user = await db.scalar(select(User).where(User.username == identity["username"]))
            if not user:
                raise HTTPException(status_code=401, detail="当前用户不存在")
            department_usernames = set((await db.scalars(select(User.username).where(
                User.department == user.department,
            ))).all())
            if relation == "owned":
                # “部门接受的任务”按当前负责人所属部门，而不是任务发起部门。
                tasks = [task for task in tasks if task.owner in department_usernames]
            elif relation == "collaborating":
                # “部门协作的任务”只要任一协作人属于本部门即可。
                tasks = [task for task in tasks if department_usernames.intersection((task.data or {}).get("collaborators", []))]
            else:
                # “部门发起的任务”以创建时写入的发起部门为准；不能再次缩成
                # 当前部门负责人的个人发起记录。
                tasks = [task for task in tasks if task.department == user.department]
    elif scope == "company" and identity.get("role") != "admin":
        raise HTTPException(status_code=403, detail="只有系统管理员可以查看全所任务")
    username = identity["username"]
    if scope != "department":
        # 管理员无论从“我的”“部门”或“公司”入口进入，都保留全所数据范围；
        # 不能因为选择“发起/接收/协作”再把最高权限缩回为本人参与的任务。
        is_admin_global_view = identity.get("role") == "admin"
        if relation == "initiated" and not is_admin_global_view:
            tasks = [task for task in tasks if (task.data or {}).get("initiator") == username]
        elif relation == "owned" and not is_admin_global_view:
            tasks = [task for task in tasks if task.owner == username]
        elif relation == "collaborating" and not is_admin_global_view:
            tasks = [task for task in tasks if username in (task.data or {}).get("collaborators", [])]
    items = [_task_dict(item) for item in tasks]

    def contains(value: object, needle: str) -> bool:
        return not needle.strip() or needle.strip().casefold() in str(value or "").casefold()

    if keyword:
        key = keyword.lower()
        items = [item for item in items if key in f"{item['serial_no']} {item['title']} {item['customer']} {item['owner']}".lower()]
    items = [item for item in items if
        contains(item.get("priority"), priority)
        and contains(item.get("serial_no"), serial_no)
        and contains(item.get("title"), title)
        and contains(item.get("description"), description)
        and contains(item.get("initiator"), initiator)
        and contains(item.get("case_no"), case_no)
        and contains(item.get("source"), source)
        and contains(item.get("owner"), owner)
        and contains(item.get("plaintiff"), plaintiff)
        and contains(item.get("defendant"), defendant)
    ]
    if created_from:
        items = [item for item in items if item.get("created_at") and item["created_at"].date() >= created_from]
    if created_to:
        items = [item for item in items if item.get("created_at") and item["created_at"].date() <= created_to]
    if deadline_from:
        items = [item for item in items if item.get("deadline") and item["deadline"] >= deadline_from]
    if deadline_to:
        items = [item for item in items if item.get("deadline") and item["deadline"] <= deadline_to]
    reverse_sort = sort_order == "desc"
    # 先按 ID 做稳定次排序，再按所选字段排序；空值始终位于末尾。
    items.sort(key=lambda item: item["id"], reverse=reverse_sort)
    populated = [item for item in items if item.get(sort_by) is not None and item.get(sort_by) != ""]
    missing = [item for item in items if item.get(sort_by) is None or item.get(sort_by) == ""]
    populated.sort(key=lambda item: item[sort_by], reverse=reverse_sort)
    items = populated + missing
    if reminder_only:
        items = [item for item in items if item["reminder_due"]]
    status_counts: dict[str, int] = {}
    for item in items:
        status_name = str(item.get("status") or "")
        status_counts[status_name] = status_counts.get(status_name, 0) + 1
    selected_statuses = {value.strip() for value in statuses.split(",") if value.strip()}
    if selected_statuses:
        items = [item for item in items if item["status"] in selected_statuses]
    if status_filter:
        items = [item for item in items if item["status"] == status_filter]
    all_items = [_task_dict(item) for item in tasks]
    if reminder_only:
        all_items = [item for item in all_items if item["reminder_due"]]
    total = len(items)
    effective_page_size = page_size or max(total, 1)
    if page_size is not None:
        start = (page - 1) * page_size
        items = items[start:start + page_size]
    return {
        "items": items, "total": total, "page": page, "page_size": effective_page_size,
        "status_counts": status_counts,
        "summary": {
            "total": len(all_items),
            "pending": sum(1 for item in all_items if item["status"] in {"待接收", "待处理"}),
            "processing": sum(1 for item in all_items if item["status"] == "处理中"),
            "awaiting_confirmation": sum(1 for item in all_items if item["status"] == "已完成"),
            "due_soon": sum(1 for item in all_items if item["days_remaining"] in {0, 1} and item["status"] not in {"已完成", "已撤回"}),
            "overdue": sum(1 for item in all_items if item["status"] == "已逾期"),
            "reminders": sum(1 for item in all_items if item["reminder_due"]),
        },
    }


@app.get(f"{settings.api_prefix}/tasks/unread-messages")
async def list_unread_task_messages(
    priority: str = "", serial_no: str = "", title: str = "", description: str = "",
    initiator: str = "", case_no: str = "", source: str = "", owner: str = "",
    plaintiff: str = "", defendant: str = "",
    created_from: date | None = None, created_to: date | None = None,
    deadline_from: date | None = None, deadline_to: date | None = None,
    sort_by: str = Query("", pattern="^(|created_at|deadline|days_remaining|updated_at)$"),
    sort_order: str = Query("desc", pattern="^(asc|desc)$"),
    page: int = Query(1, ge=1), page_size: int = Query(15, ge=1, le=200),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    """Aggregate the current recipient's unread task communication messages.

    Administrator visibility is intentionally not expanded here: unread state is
    personal, so even an administrator only sees notifications addressed to that
    administrator.  The task record itself still uses the normal participant
    guard for non-administrators.
    """
    if created_from and created_to and created_from > created_to:
        raise HTTPException(status_code=422, detail="发起开始日期不能晚于结束日期")
    if deadline_from and deadline_to and deadline_from > deadline_to:
        raise HTTPException(status_code=422, detail="截止开始日期不能晚于结束日期")

    notices = list((await db.scalars(
        select(Notification).where(
            Notification.recipient == identity["username"],
            Notification.recipient_deleted.is_(False),
            Notification.is_read.is_(False),
            Notification.source_type == "task",
            Notification.source_id.is_not(None),
            or_(
                Notification.source_key.like("task-message-%"),
                Notification.source_key.like("task-history-%"),
            ),
        ).order_by(Notification.created_at.desc(), Notification.id.desc())
    )).all())
    grouped: dict[int, list[Notification]] = {}
    for notice in notices:
        if notice.source_id is not None:
            grouped.setdefault(int(notice.source_id), []).append(notice)
    if not grouped:
        return {"items": [], "total": 0, "page": page, "page_size": page_size, "unread_messages": 0}

    tasks = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "task", BusinessRecord.id.in_(list(grouped)),
    ))).all())
    if identity.get("role") != "admin":
        tasks = [task for task in tasks if _is_task_participant(task, identity)]

    def contains(value: object, needle: str) -> bool:
        return not needle.strip() or needle.strip().casefold() in str(value or "").casefold()

    sender_usernames = {notice.sender for notice in notices if notice.sender and notice.sender != "system"}
    sender_names = {
        user.username: (user.display_name or user.username)
        for user in (await db.scalars(select(User).where(User.username.in_(sender_usernames)))).all()
    } if sender_usernames else {}
    items: list[dict] = []
    visible_notice_count = 0
    for task in tasks:
        row = _task_dict(task)
        if not (
            contains(row.get("priority"), priority)
            and contains(row.get("serial_no"), serial_no)
            and contains(row.get("title"), title)
            and contains(row.get("description"), description)
            and contains(row.get("initiator"), initiator)
            and contains(row.get("case_no"), case_no)
            and contains(row.get("source"), source)
            and contains(row.get("owner"), owner)
            and contains(row.get("plaintiff"), plaintiff)
            and contains(row.get("defendant"), defendant)
        ):
            continue
        if created_from and (not row.get("created_at") or row["created_at"].date() < created_from):
            continue
        if created_to and (not row.get("created_at") or row["created_at"].date() > created_to):
            continue
        if deadline_from and (not row.get("deadline") or row["deadline"] < deadline_from):
            continue
        if deadline_to and (not row.get("deadline") or row["deadline"] > deadline_to):
            continue
        task_notices = grouped[task.id]
        latest = task_notices[0]
        latest_content = latest.content or latest.title
        if latest.source_key.startswith("task-history-") and "｜" in latest_content:
            latest_content = latest_content.split("｜", 1)[1]
        visible_notice_count += len(task_notices)
        items.append({
            **row,
            "latest_unread_message": latest_content,
            "latest_unread_sender": "System" if latest.sender == "system" else sender_names.get(latest.sender, latest.sender or "System"),
            "latest_unread_at": latest.created_at,
            "latest_unread_notification_id": latest.id,
            "unread_count": len(task_notices),
        })
    if sort_by:
        reverse_sort = sort_order == "desc"
        items.sort(key=lambda item: item["id"], reverse=reverse_sort)
        populated = [item for item in items if item.get(sort_by) is not None and item.get(sort_by) != ""]
        missing = [item for item in items if item.get(sort_by) is None or item.get(sort_by) == ""]
        populated.sort(key=lambda item: item[sort_by], reverse=reverse_sort)
        items = populated + missing
    else:
        items.sort(key=lambda item: (item["latest_unread_at"], item["latest_unread_notification_id"]), reverse=True)
    total = len(items)
    start = (page - 1) * page_size
    return {
        "items": items[start:start + page_size], "total": total, "page": page, "page_size": page_size,
        "unread_messages": visible_notice_count,
    }


async def _active_task_username(value: str, db: AsyncSession, *, field_name: str) -> str:
    requested = value.strip()
    if not requested:
        raise HTTPException(status_code=422, detail=f"{field_name}不能为空")
    user = await db.scalar(select(User).where(User.username == requested))
    if not user:
        matches = list((await db.scalars(select(User).where(User.display_name == requested))).all())
        if len(matches) > 1:
            raise HTTPException(status_code=422, detail=f"{field_name}姓名不唯一，请填写账号")
        user = matches[0] if matches else None
    if not user or not user.is_active:
        raise HTTPException(status_code=422, detail=f"{field_name}不存在或已停用")
    return user.username


def _validate_task_deadline(deadline: date) -> None:
    duration = (deadline - date.today()).days
    if duration < 0:
        raise HTTPException(status_code=422, detail="任务截止日期不能早于今天")
    if duration > 30:
        raise HTTPException(status_code=422, detail="任务截止日期不能超过 30 天")


@app.post(f"{settings.api_prefix}/tasks", status_code=status.HTTP_201_CREATED)
async def create_task(body: TaskInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _validate_task_deadline(body.deadline)
    case_no = body.case_no.strip()
    source = body.source.strip() or "日常任务"
    if source == "案件任务" and not case_no:
        raise HTTPException(status_code=422, detail="案件任务必须关联有效案件")
    case_record = None
    if case_no:
        case_record = await db.scalar(select(BusinessRecord).where(BusinessRecord.module == "case", BusinessRecord.serial_no == case_no))
        if not case_record:
            raise HTTPException(status_code=404, detail="关联案件不存在")
        case_record = await _ensure_record_module(case_record.id, "case", identity, db)
        capabilities = await _case_detail_action_capabilities(case_record, identity, db)
        if not capabilities["can_create_case_task"]:
            raise HTTPException(status_code=403, detail="当前账号没有创建该案件任务的权限")
    serial = f"RW{datetime.now():%Y%m%d%H%M%S%f}"
    user = await db.scalar(select(User).where(User.username == identity["username"]))
    owner = await _active_task_username(body.owner, db, field_name="负责人")
    collaborators = []
    for value in body.collaborators:
        collaborator = await _active_task_username(value, db, field_name="协作人")
        if collaborator != owner and collaborator not in collaborators:
            collaborators.append(collaborator)
    task = BusinessRecord(module="task", serial_no=serial, title=body.title, customer=case_record.customer if case_record else body.customer, status="待接收", owner=owner, department=user.department if user else "上海分所", description=body.description, data={"deadline": str(body.deadline), "priority": body.priority, "source": source, "initiator": identity["username"], "collaborators": collaborators, "case_no": case_no, "case_record_id": case_record.id if case_record else None})
    db.add(task)
    await db.flush()
    await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action="发起任务", to_status="待接收", operator=identity["username"], comment=f"负责人：{owner}；截止日期：{body.deadline}"), db, content="任务已分派.")
    await db.commit()
    await db.refresh(task)
    return _task_dict(task)


@app.post(f"{settings.api_prefix}/tasks/batch-update")
async def batch_update_tasks(body: TaskBatchUpdateInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if body.owner is None and body.deadline is None and body.priority is None:
        raise HTTPException(status_code=422, detail="请至少选择一个需要修改的字段")
    if body.deadline is not None:
        duration = (body.deadline - date.today()).days
        if duration < 0: raise HTTPException(status_code=422, detail="任务截止日期不能早于今天")
        if duration > 30: raise HTTPException(status_code=422, detail="任务截止日期不能超过 30 天")
    if body.priority is not None and body.priority not in {"普通", "重要", "紧急"}:
        raise HTTPException(status_code=422, detail="任务优先级无效")
    tasks = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "task", BusinessRecord.id.in_(set(body.task_ids))))).all()
    if len(tasks) != len(set(body.task_ids)):
        raise HTTPException(status_code=404, detail="部分任务不存在")
    normalized_owner = None
    if body.owner is not None:
        normalized_owner = await _active_task_username(body.owner, db, field_name="负责人")
    for task in tasks:
        _require_task_owner_or_initiator(task, identity, action="批量修改任务")
        if task.status not in {"待接收", "待处理", "处理中"}:
            raise HTTPException(status_code=409, detail=f"任务 {task.serial_no} 已进入不可批量修改的状态")
        changes: list[str] = []
        if normalized_owner is not None:
            changes.append(f"负责人：{task.owner} → {normalized_owner}")
            task.owner = normalized_owner
        data = dict(task.data or {})
        if body.deadline is not None:
            changes.append(f"截止日期：{data.get('deadline', '')} → {body.deadline}")
            data["deadline"] = str(body.deadline)
        if body.priority is not None:
            changes.append(f"优先级：{data.get('priority', '')} → {body.priority}")
            data["priority"] = body.priority
        task.data = data
        await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action="批量修改任务", from_status=task.status, to_status=task.status, operator=identity["username"], comment="；".join(changes + ([body.comment] if body.comment else []))), db, content="任务已修改.")
    await db.commit()
    for task in tasks:
        await db.refresh(task)
    return {"updated": len(tasks), "items": [_task_dict(task) for task in tasks]}


@app.post(f"{settings.api_prefix}/tasks/batch-lifecycle")
async def batch_lifecycle_tasks(body: TaskBatchLifecycleInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Apply one dedicated task lifecycle action atomically to selected tasks.

    This deliberately does not reuse ``batch-update`` or the generic record
    transition endpoint: task state transitions have different participants,
    automatic deadlines and audit semantics.
    """
    task_ids = list(dict.fromkeys(body.task_ids))
    tasks = (await db.scalars(
        select(BusinessRecord).where(BusinessRecord.module == "task", BusinessRecord.id.in_(task_ids))
    )).all()
    if len(tasks) != len(task_ids):
        raise HTTPException(status_code=404, detail="部分任务不存在")
    tasks_by_id = {task.id: task for task in tasks}
    ordered_tasks = [tasks_by_id[task_id] for task_id in task_ids]
    comment = body.comment.strip()
    recipient = ""
    if body.action == "handoff":
        recipient = await _active_task_username(body.recipient, db, field_name="接收人")
    if body.action == "withdraw" and not comment:
        raise HTTPException(status_code=422, detail="批量撤回任务必须填写撤回原因")

    # Validate the full selection before changing any record, so a mixed
    # selection never creates a partial lifecycle result.
    for task in ordered_tasks:
        data = task.data or {}
        if body.action == "accept":
            if identity.get("role") != "admin" and task.owner != identity["username"]:
                raise HTTPException(status_code=403, detail=f"任务 {task.serial_no} 仅负责人可接收")
            if task.status not in {"待接收", "待处理"}:
                raise HTTPException(status_code=409, detail=f"任务 {task.serial_no} 当前状态不能接收")
        elif body.action == "complete":
            if identity.get("role") != "admin" and task.owner != identity["username"]:
                raise HTTPException(status_code=403, detail=f"任务 {task.serial_no} 仅负责人可提交完成")
            if task.status != "处理中":
                raise HTTPException(status_code=409, detail=f"任务 {task.serial_no} 仅处理中可提交完成")
        elif body.action == "handoff":
            if identity.get("role") != "admin" and task.owner != identity["username"]:
                raise HTTPException(status_code=403, detail=f"任务 {task.serial_no} 仅当前负责人可交接")
            if task.status in {"已完成", "已验收", "已撤回", "已停止", "已取消", "待确认", "已拒绝"}:
                raise HTTPException(status_code=409, detail=f"任务 {task.serial_no} 已结束，不能交接")
            if recipient == task.owner:
                raise HTTPException(status_code=422, detail=f"任务 {task.serial_no} 不能交接给当前负责人")
        else:  # withdraw
            if identity.get("role") != "admin" and data.get("initiator") != identity["username"]:
                raise HTTPException(status_code=403, detail=f"任务 {task.serial_no} 仅发起人可撤回")
            if task.status not in {"待接收", "待处理", "处理中"}:
                raise HTTPException(status_code=409, detail=f"任务 {task.serial_no} 当前状态不能撤回")

    auto_at = date.today() + timedelta(days=5)
    action_labels = {
        "accept": "批量接收任务",
        "complete": "批量提交任务完成",
        "handoff": "批量任务交接",
        "withdraw": "批量撤回任务",
    }
    content_labels = {
        "accept": "任务已批量接收.",
        "complete": "任务已批量提交完成，等待确认.",
        "handoff": "任务已批量交接.",
        "withdraw": "任务已批量撤回.",
    }
    for task in ordered_tasks:
        previous = task.status
        data = dict(task.data or {})
        if body.action == "accept":
            task.status = "处理中"
            task.data = {
                **data,
                "accepted_at": datetime.now().isoformat(timespec="seconds"),
                "handoff_restarted": True,
                "rejected_reason": "",
            }
        elif body.action == "complete":
            task.status = "已完成"
            task.data = {
                **data,
                "completion_submitted_at": datetime.now().isoformat(timespec="seconds"),
                "completion_auto_confirm_at": str(auto_at),
                "completion_comment": comment,
            }
        elif body.action == "handoff":
            previous_owner = task.owner
            task.owner = recipient
            task.status = "待接收"
            task.data = {
                **data,
                "handoff_from": previous_owner,
                "handoff_recipient": recipient,
                "handed_off_at": str(date.today()),
                "handoff_auto_complete_at": str(auto_at),
                "handoff_restarted": False,
            }
        else:
            task.status = "已撤回"
            task.data = {
                **data,
                "withdrawn_by": identity["username"],
                "withdrawn_at": datetime.now().isoformat(timespec="seconds"),
                "withdraw_comment": comment,
                "handoff_auto_complete_at": "",
                "completion_auto_confirm_at": "",
                "exception_request": {},
            }
        event_comment = comment
        if body.action == "handoff":
            event_comment = f"{previous_owner} 交接给 {recipient}；未重新开始将于 {auto_at} 自动完成。{comment}"
        elif body.action == "complete":
            event_comment = f"发起人应在 {auto_at} 前验收或退回重启。{comment}"
        await _add_task_message_notifications(
            task,
            WorkflowEvent(
                record_id=task.id,
                action=action_labels[body.action],
                from_status=previous,
                to_status=task.status,
                operator=identity["username"],
                comment=event_comment,
            ),
            db,
            content=content_labels[body.action],
        )
    await db.commit()
    for task in ordered_tasks:
        await db.refresh(task)
    return {"updated": len(ordered_tasks), "action": body.action, "items": [_task_dict(task) for task in ordered_tasks]}


async def _task_or_404(task_id: int, db: AsyncSession) -> BusinessRecord:
    task = await db.get(BusinessRecord, task_id)
    if not task or task.module != "task":
        raise HTTPException(status_code=404, detail="任务不存在")
    return task


def _is_task_participant(task: BusinessRecord, identity: dict) -> bool:
    data = task.data or {}
    username = identity["username"]
    return identity.get("role") == "admin" or username == task.owner or username == data.get("initiator") or username in data.get("collaborators", [])


@app.post(f"{settings.api_prefix}/tasks/{{task_id}}/accept")
async def accept_task(task_id: int, body: TaskActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    task = await _task_or_404(task_id, db)
    if identity.get("role") != "admin" and task.owner != identity["username"]:
        raise HTTPException(status_code=403, detail="只有任务负责人可以接收任务")
    if task.status not in {"待接收", "待处理"}:
        raise HTTPException(status_code=409, detail="当前状态不能接收任务")
    previous = task.status; task.status = "处理中"
    task.data = {**(task.data or {}), "accepted_at": datetime.now().isoformat(timespec="seconds"), "handoff_restarted": True, "rejected_reason": ""}
    await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action="接收任务", from_status=previous, to_status="处理中", operator=identity["username"], comment=body.comment), db, content="任务已接受.")
    await db.commit(); await db.refresh(task); return _task_dict(task)


@app.post(f"{settings.api_prefix}/tasks/{{task_id}}/reject")
async def reject_task(task_id: int, body: TaskActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    task = await _task_or_404(task_id, db)
    if identity.get("role") != "admin" and task.owner != identity["username"]:
        raise HTTPException(status_code=403, detail="只有任务负责人可以拒绝任务")
    if task.status not in {"待接收", "待处理"}:
        raise HTTPException(status_code=409, detail="当前状态不能拒绝任务")
    if not body.comment.strip(): raise HTTPException(status_code=422, detail="拒绝任务必须填写理由")
    previous = task.status; task.status = "已拒绝"
    task.data = {**(task.data or {}), "rejected_reason": body.comment.strip(), "rejected_at": datetime.now().isoformat(timespec="seconds"), "handoff_auto_complete_at": ""}
    await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action="拒绝任务", from_status=previous, to_status="已拒绝", operator=identity["username"], comment=body.comment), db, content="任务已拒绝.")
    await db.commit(); await db.refresh(task); return _task_dict(task)


@app.post(f"{settings.api_prefix}/tasks/{{task_id}}/withdraw")
async def withdraw_task(task_id: int, body: TaskActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Withdraw a live task through its dedicated lifecycle, never generic transition."""
    task = await _task_or_404(task_id, db)
    data = task.data or {}
    if identity.get("role") != "admin" and data.get("initiator") != identity["username"]:
        raise HTTPException(status_code=403, detail="只有任务发起人或系统管理员可以撤回任务")
    # ``待处理`` is retained for historical imported tasks and has the same pre-accept semantics.
    if task.status not in {"待接收", "待处理", "处理中"}:
        raise HTTPException(status_code=409, detail="只有待接收或处理中的任务可以撤回")
    comment = body.comment.strip()
    if not comment:
        raise HTTPException(status_code=422, detail="撤回任务必须填写撤回原因")
    previous = task.status
    task.status = "已撤回"
    task.data = {
        **data,
        "withdrawn_by": identity["username"],
        "withdrawn_at": datetime.now().isoformat(timespec="seconds"),
        "withdraw_comment": comment,
        "handoff_auto_complete_at": "",
        "completion_auto_confirm_at": "",
        "exception_request": {},
    }
    await _add_task_message_notifications(
        task,
        WorkflowEvent(
            record_id=task.id,
            action="撤回任务",
            from_status=previous,
            to_status="已撤回",
            operator=identity["username"],
            comment=comment,
        ),
        db,
        content="任务已撤回。",
    )
    await db.commit(); await db.refresh(task); return _task_dict(task)


@app.post(f"{settings.api_prefix}/tasks/{{task_id}}/resend")
async def resend_task(task_id: int, body: TaskHandoffInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    task = await _task_or_404(task_id, db); data = task.data or {}
    if identity.get("role") != "admin" and data.get("initiator") != identity["username"]:
        raise HTTPException(status_code=403, detail="只有任务发起人可以重新派发")
    if task.status != "已拒绝": raise HTTPException(status_code=409, detail="只有已拒绝任务可以重新派发")
    previous_owner = task.owner
    recipient = await _active_task_username(body.recipient, db, field_name="新负责人")
    auto_at = date.today() + timedelta(days=5)
    task.owner = recipient; task.status = "待接收"
    task.data = {**data, "rejected_reason": "", "resent_at": datetime.now().isoformat(timespec="seconds"), "handoff_from": previous_owner, "handoff_recipient": recipient, "handed_off_at": str(date.today()), "handoff_auto_complete_at": str(auto_at), "handoff_restarted": False}
    await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action="重新派发任务", from_status="已拒绝", to_status="待接收", operator=identity["username"], comment=f"{previous_owner} → {recipient}。{body.comment}"), db, content="任务已重新派发.")
    await db.commit(); await db.refresh(task); return _task_dict(task)


@app.post(f"{settings.api_prefix}/tasks/{{task_id}}/comments", status_code=status.HTTP_201_CREATED)
async def add_task_comment(task_id: int, body: TaskActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    task = await _task_or_404(task_id, db)
    if not _is_task_participant(task, identity): raise HTTPException(status_code=403, detail="只有任务参与人可以沟通")
    if not body.comment.strip(): raise HTTPException(status_code=422, detail="沟通内容不能为空")
    event = WorkflowEvent(record_id=task.id, action="任务沟通", from_status=task.status, to_status=task.status, operator=identity["username"], comment=body.comment.strip())
    await _add_task_message_notifications(task, event, db, content=body.comment.strip())
    await db.commit(); await db.refresh(event)
    return {"id": event.id, "operator": event.operator, "comment": event.comment, "created_at": event.created_at}


@app.post(f"{settings.api_prefix}/tasks/{{task_id}}/feedback", status_code=status.HTTP_201_CREATED)
async def create_task_feedback(
    task_id: int,
    comment: str = Form(...),
    files: list[UploadFile] = File(default=[]),
    identity: dict = Depends(current_identity),
    db: AsyncSession = Depends(get_db),
):
    """Create one task feedback message and all selected attachments atomically."""
    task = await _task_or_404(task_id, db)
    if not _is_task_participant(task, identity):
        raise HTTPException(status_code=403, detail="只有任务参与人可以提交反馈")
    normalized_comment = comment.strip()
    if not normalized_comment:
        raise HTTPException(status_code=422, detail="反馈内容不能为空")
    if len(normalized_comment) > 1000:
        raise HTTPException(status_code=422, detail="反馈内容不能超过 1000 个字符")
    if len(files) > 20:
        raise HTTPException(status_code=422, detail="一次最多上传 20 个反馈附件")

    allowed = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx", ".txt", ".png", ".jpg", ".jpeg", ".zip", ".rar"}
    prepared_files: list[tuple[str, str, bytes]] = []
    for file in files:
        suffix = Path(file.filename or "").suffix.lower()
        if suffix not in allowed:
            raise HTTPException(status_code=422, detail="不支持的文件格式")
        content = await file.read()
        if len(content) > 20 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="单个文件不能超过 20MB")
        prepared_files.append((Path(file.filename or f"task-feedback{suffix}").name, file.content_type or "application/octet-stream", content))

    written_paths: list[Path] = []
    attachments: list[FileAttachment] = []
    try:
        await _add_task_message_notifications(
            task,
            WorkflowEvent(
                record_id=task.id, action="任务沟通", from_status=task.status,
                to_status=task.status, operator=identity["username"], comment=normalized_comment,
            ),
            db,
            content=normalized_comment,
        )
        for original_name, content_type, content in prepared_files:
            suffix = Path(original_name).suffix.lower()
            target = UPLOAD_ROOT / f"{uuid4().hex}{suffix}"
            target.write_bytes(content)
            written_paths.append(target)
            attachment = FileAttachment(
                record_id=task.id, category="任务反馈附件", original_name=original_name,
                stored_name=target.name, content_type=content_type, size=len(content),
                path=str(target), uploader=identity["username"], remark=normalized_comment,
            )
            db.add(attachment)
            attachments.append(attachment)
            await _add_task_message_notifications(
                task,
                WorkflowEvent(
                    record_id=task.id, action="上传任务反馈附件", from_status=task.status,
                    to_status=task.status, operator=identity["username"],
                    comment=f"任务反馈附件：{original_name}",
                ),
                db,
                content=f"已上传任务反馈附件：{original_name}",
            )
        await db.commit()
        for attachment in attachments:
            await db.refresh(attachment)
    except Exception:
        await db.rollback()
        for path in written_paths:
            path.unlink(missing_ok=True)
        raise
    return {
        "comment": normalized_comment,
        "attachments": [_attachment_dict(attachment, task) for attachment in attachments],
    }


@app.post(f"{settings.api_prefix}/tasks/{{task_id}}/materials", status_code=status.HTTP_201_CREATED)
async def upload_task_materials(
    task_id: int,
    files: list[UploadFile] = File(...),
    remark: str = Form(""),
    identity: dict = Depends(current_identity),
    db: AsyncSession = Depends(get_db),
):
    """Upload task materials independently from feedback messages.

    Task materials are business files supplied at creation or during handling.  They
    must not be represented as feedback attachments because that would fabricate a
    communication entry and erase their distinct audit meaning.
    """
    task = await _task_or_404(task_id, db)
    if not _is_task_participant(task, identity):
        raise HTTPException(status_code=403, detail="只有任务参与人可以上传任务资料附件")
    if not files:
        raise HTTPException(status_code=422, detail="请至少选择一个任务资料附件")
    if len(files) > 20:
        raise HTTPException(status_code=422, detail="一次最多上传 20 个任务资料附件")

    allowed = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx", ".txt", ".png", ".jpg", ".jpeg", ".zip", ".rar"}
    prepared_files: list[tuple[str, str, bytes]] = []
    for file in files:
        suffix = Path(file.filename or "").suffix.lower()
        if suffix not in allowed:
            raise HTTPException(status_code=422, detail="不支持的文件格式")
        content = await file.read()
        if len(content) > 20 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="单个文件不能超过 20MB")
        prepared_files.append((Path(file.filename or f"task-material{suffix}").name, file.content_type or "application/octet-stream", content))

    written_paths: list[Path] = []
    attachments: list[FileAttachment] = []
    try:
        for original_name, content_type, content in prepared_files:
            suffix = Path(original_name).suffix.lower()
            target = UPLOAD_ROOT / f"{uuid4().hex}{suffix}"
            target.write_bytes(content)
            written_paths.append(target)
            attachment = FileAttachment(
                record_id=task.id, category="任务资料附件", original_name=original_name,
                stored_name=target.name, content_type=content_type, size=len(content),
                path=str(target), uploader=identity["username"], remark=remark.strip(),
            )
            db.add(attachment)
            attachments.append(attachment)
            await _add_task_message_notifications(
                task,
                WorkflowEvent(
                    record_id=task.id, action="上传任务资料附件", from_status=task.status,
                    to_status=task.status, operator=identity["username"],
                    comment=f"任务资料附件：{original_name}",
                ),
                db,
                content=f"已上传任务资料附件：{original_name}",
            )
        await db.commit()
        for attachment in attachments:
            await db.refresh(attachment)
    except Exception:
        await db.rollback()
        for path in written_paths:
            path.unlink(missing_ok=True)
        raise
    return {"attachments": [_attachment_dict(attachment, task) for attachment in attachments]}


@app.get(f"{settings.api_prefix}/tasks/{{task_id}}/history")
async def task_history(task_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    task = await _task_or_404(task_id, db)
    if not _is_task_participant(task, identity): raise HTTPException(status_code=403, detail="只有任务参与人可以查看沟通记录")
    events = (await db.scalars(select(WorkflowEvent).where(WorkflowEvent.record_id == task.id).order_by(WorkflowEvent.created_at.desc(), WorkflowEvent.id.desc()))).all()
    unread_keys = set((await db.scalars(select(Notification.source_key).where(
        Notification.recipient == identity["username"], Notification.recipient_deleted.is_(False),
        Notification.is_read.is_(False), Notification.source_type == "task", Notification.source_id == task.id,
        Notification.source_key.like(f"task-history-{task.id}-%-{identity['username']}"),
    ))).all())
    return {"items": [{
        "id": item.id, "action": item.action, "operator": item.operator, "comment": item.comment,
        "from_status": item.from_status, "to_status": item.to_status, "created_at": item.created_at,
        "unread": f"task-history-{task.id}-{item.id}-{identity['username']}" in unread_keys,
    } for item in events]}


@app.post(f"{settings.api_prefix}/tasks/{{task_id}}/history/{{event_id}}/mark-unread")
async def mark_task_history_unread(task_id: int, event_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    task = await _task_or_404(task_id, db)
    if not _is_task_participant(task, identity):
        raise HTTPException(status_code=403, detail="只有任务参与人可以标记沟通记录")
    event = await db.scalar(select(WorkflowEvent).where(WorkflowEvent.id == event_id, WorkflowEvent.record_id == task.id))
    if not event:
        raise HTTPException(status_code=404, detail="任务历史事件不存在")
    source_key = f"task-history-{task.id}-{event.id}-{identity['username']}"
    item = await db.scalar(select(Notification).where(Notification.source_key == source_key))
    if item:
        item.is_read = False
        item.read_at = None
        item.recipient_deleted = False
        if item.sender == identity["username"]:
            item.sender_deleted = False
    else:
        item = Notification(
            source_key=source_key, source_type="task", source_id=task.id,
            sender=event.operator or "system", recipient=identity["username"], notification_type="系统通知",
            title=f"任务历史待处理：{task.serial_no}",
            content=f"{event.action}｜{event.comment or '无备注'}", level="info", is_read=False,
        )
        db.add(item)
    await db.commit()
    await db.refresh(item)
    return _notification_dict(item)


@app.post(f"{settings.api_prefix}/tasks/{{task_id}}/messages/read")
async def read_task_messages(task_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    task = await _task_or_404(task_id, db)
    if not _is_task_participant(task, identity):
        raise HTTPException(status_code=403, detail="只有任务参与人可以读取任务消息")
    items = list((await db.scalars(select(Notification).where(
        Notification.recipient == identity["username"], Notification.recipient_deleted.is_(False),
        Notification.is_read.is_(False), Notification.source_type == "task", Notification.source_id == task.id,
        or_(
            Notification.source_key.like("task-message-%"),
            Notification.source_key.like("task-history-%"),
        ),
    ))).all())
    now = datetime.now()
    for item in items:
        item.is_read = True
        item.read_at = now
    await db.commit()
    return {"task_id": task.id, "updated": len(items), "is_read": True}


@app.post(f"{settings.api_prefix}/tasks/messages/batch-read")
async def batch_read_task_messages(body: TaskBatchReadInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Mark unread task messages read for the current recipient only.

    The original unread-task page has a selected-row "标记已读" action.  This
    command deliberately does not grant administrators access to somebody
    else's personal inbox: every notification is still filtered by recipient.
    All selected tasks are checked before any notification changes, so a mixed
    selection cannot partially succeed.
    """
    task_ids = list(dict.fromkeys(body.task_ids))
    tasks = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "task", BusinessRecord.id.in_(task_ids),
    ))).all())
    if len(tasks) != len(task_ids):
        raise HTTPException(status_code=404, detail="部分任务不存在")
    for task in tasks:
        if not _is_task_participant(task, identity):
            raise HTTPException(status_code=403, detail="只有任务参与人可以读取任务消息")
    items = list((await db.scalars(select(Notification).where(
        Notification.recipient == identity["username"], Notification.recipient_deleted.is_(False),
        Notification.is_read.is_(False), Notification.source_type == "task", Notification.source_id.in_(task_ids),
        or_(
            Notification.source_key.like("task-message-%"),
            Notification.source_key.like("task-history-%"),
        ),
    ))).all())
    now = datetime.now()
    for item in items:
        item.is_read = True
        item.read_at = now
    await db.commit()
    return {"task_ids": task_ids, "updated": len(items), "is_read": True}


@app.post(f"{settings.api_prefix}/tasks/{{task_id}}/handoff")
async def handoff_task(task_id: int, body: TaskHandoffInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    task = await _task_or_404(task_id, db)
    if identity.get("role") != "admin" and task.owner != identity["username"]:
        raise HTTPException(status_code=403, detail="只有当前负责人可以转交任务")
    if task.status in {"已完成", "已验收", "已撤回", "已停止", "已取消", "待确认", "已拒绝"}:
        raise HTTPException(status_code=409, detail="已结束任务不能交接")
    previous_owner, previous_status = task.owner, task.status
    recipient = await _active_task_username(body.recipient, db, field_name="接收人")
    if recipient == previous_owner:
        raise HTTPException(status_code=422, detail="任务不能转交给当前负责人")
    auto_at = date.today() + timedelta(days=5)
    task.owner = recipient
    task.status = "待接收"
    task.data = {**(task.data or {}), "handoff_from": previous_owner, "handoff_recipient": recipient, "handed_off_at": str(date.today()), "handoff_auto_complete_at": str(auto_at), "handoff_restarted": False}
    await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action="任务交接", from_status=previous_status, to_status="待接收", operator=identity["username"], comment=f"{previous_owner} 交接给 {recipient}；未重新开始将于 {auto_at} 自动完成。{body.comment}"), db, content="任务已交接.")
    await db.commit()
    await db.refresh(task)
    return _task_dict(task)


@app.post(f"{settings.api_prefix}/tasks/{{task_id}}/restart")
async def restart_task(task_id: int, body: TaskActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    task = await _task_or_404(task_id, db); data = task.data or {}
    if task.status == "已完成" and (task.data or {}).get("auto_completed"):
        raise HTTPException(status_code=409, detail="任务已自动完成，请新建后续任务")
    if task.status in {"待确认", "已完成"}:
        if identity.get("role") != "admin" and data.get("initiator") != identity["username"]: raise HTTPException(status_code=403, detail="只有发起人可以退回重启")
    elif identity.get("role") != "admin" and task.owner != identity["username"]:
        raise HTTPException(status_code=403, detail="只有任务负责人可以开始任务")
    if task.status == "已停止" and ((data.get("exception_request") or {}).get("action") != "挂起" or (data.get("exception_request") or {}).get("status") != "已通过"):
        raise HTTPException(status_code=409, detail="只有审批通过的挂起任务可以恢复")
    if task.status not in {"待接收", "待处理", "待确认", "已完成", "已停止"}: raise HTTPException(status_code=409, detail="当前状态不能重新开始")
    previous = task.status
    task.status = "处理中"
    task.data = {**data, "handoff_restarted": True, "restarted_at": str(date.today()), "completion_auto_confirm_at": "", "completion_comment": "", "exception_request": {**(data.get("exception_request") or {}), "resumed_at": datetime.now().isoformat(timespec="seconds"), "resumed_by": identity["username"]}}
    await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action="重新开始任务", from_status=previous, to_status="处理中", operator=identity["username"], comment=body.comment), db, content="任务已重新开始.")
    await db.commit()
    await db.refresh(task)
    return _task_dict(task)


@app.post(f"{settings.api_prefix}/tasks/{{task_id}}/exception-request")
async def request_task_exception(task_id: int, body: TaskExceptionRequestInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    task = await _task_or_404(task_id, db)
    if identity.get("role") != "admin" and task.owner != identity["username"]:
        raise HTTPException(status_code=403, detail="只有任务负责人可以申请挂起或取消")
    if task.status not in {"待接收", "待处理", "处理中"}:
        raise HTTPException(status_code=409, detail="当前任务状态不能申请挂起或取消")
    data = task.data or {}; pending = data.get("exception_request") or {}
    if pending.get("status") == "待审批":
        raise HTTPException(status_code=409, detail="已有任务特殊处理申请正在审批")
    request_data = {"action": body.action, "reason": body.reason.strip(), "status": "待审批", "requested_by": identity["username"], "requested_at": datetime.now().isoformat(timespec="seconds"), "status_before_request": task.status}
    task.data = {**data, "exception_request": request_data}
    await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action=f"申请任务{body.action}", from_status=task.status, to_status=task.status, operator=identity["username"], comment=body.reason.strip()), db, content=f"任务{body.action}申请待审批.")
    await db.commit(); await db.refresh(task)
    return _task_dict(task)


@app.post(f"{settings.api_prefix}/tasks/{{task_id}}/exception-review")
async def review_task_exception(task_id: int, body: TaskExceptionReviewInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    task = await _task_or_404(task_id, db); data = task.data or {}; pending = data.get("exception_request") or {}
    if pending.get("status") != "待审批":
        raise HTTPException(status_code=409, detail="该任务没有待审批的特殊处理申请")
    if identity.get("role") != "admin" and pending.get("requested_by") == identity["username"]:
        raise HTTPException(status_code=403, detail="申请人不能审批自己的任务挂起或取消申请")
    reviewer = await db.scalar(select(User).where(User.username == identity["username"]))
    is_same_department_manager = bool(
        identity.get("role") == "manager" and reviewer and reviewer.department == task.department
    )
    if identity.get("role") != "admin" and data.get("initiator") != identity["username"] and not is_same_department_manager:
        raise HTTPException(status_code=403, detail="只有任务发起人、部门负责人或管理员可以审批")
    if not body.approved and not body.comment.strip():
        raise HTTPException(status_code=422, detail="驳回时必须填写原因")
    previous = task.status; action_name = str(pending.get("action") or "")
    if body.approved:
        task.status = "已停止" if action_name == "挂起" else "已取消"
    reviewed = {**pending, "status": "已通过" if body.approved else "已驳回", "reviewed_by": identity["username"], "reviewed_at": datetime.now().isoformat(timespec="seconds"), "review_comment": body.comment.strip()}
    task.data = {**data, "exception_request": reviewed}
    await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action=f"任务{action_name}审批{'通过' if body.approved else '驳回'}", from_status=previous, to_status=task.status, operator=identity["username"], comment=body.comment), db, content=f"任务{action_name}申请已{'通过' if body.approved else '驳回'}.")
    await db.commit(); await db.refresh(task)
    return _task_dict(task)


@app.post(f"{settings.api_prefix}/tasks/{{task_id}}/complete")
async def complete_task(task_id: int, body: TaskActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    task = await _task_or_404(task_id, db)
    if identity.get("role") != "admin" and task.owner != identity["username"]: raise HTTPException(status_code=403, detail="只有任务负责人可以提交完成")
    if task.status != "处理中": raise HTTPException(status_code=409, detail="只有处理中的任务可以提交完成")
    previous = task.status
    task.status = "已完成"
    auto_at = date.today() + timedelta(days=5)
    task.data = {**(task.data or {}), "completion_submitted_at": datetime.now().isoformat(timespec="seconds"), "completion_auto_confirm_at": str(auto_at), "completion_comment": body.comment}
    await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action="提交任务完成", from_status=previous, to_status="已完成", operator=identity["username"], comment=f"发起人应在 {auto_at} 前验收或退回重启。{body.comment}"), db, content="任务已完成，等待确认.")
    await db.commit()
    await db.refresh(task)
    return _task_dict(task)


@app.post(f"{settings.api_prefix}/tasks/{{task_id}}/confirm")
async def confirm_task(task_id: int, body: TaskActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    task = await _task_or_404(task_id, db); data = task.data or {}
    if identity.get("role") != "admin" and data.get("initiator") != identity["username"]: raise HTTPException(status_code=403, detail="只有任务发起人可以确认完成")
    if task.status not in {"待确认", "已完成"}: raise HTTPException(status_code=409, detail="只有已完成待验收任务可以验收")
    previous = task.status
    task.status = "已验收"; task.data = {**data, "confirmed_at": datetime.now().isoformat(timespec="seconds"), "completion_auto_confirm_at": ""}
    await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action="验收任务", from_status=previous, to_status="已验收", operator=identity["username"], comment=body.comment), db, content="任务已确认完成.")
    await _advance_case_from_fixed_task(task, db, operator=identity["username"])
    await db.commit(); await db.refresh(task); return _task_dict(task)


async def _advance_case_from_fixed_task(task: BusinessRecord, db: AsyncSession, *, operator: str) -> None:
    data = task.data or {}
    if data.get("task_type") != "固定任务" or task.status != "已验收":
        return
    case_record = await db.get(BusinessRecord, int(data.get("case_id") or 0))
    if not case_record or case_record.module != "case" or case_record.status in {"待归档审核", "已归档"}:
        return
    targets = {"filing-registration": "一审立案受理", "service-tracking": "一审准备开庭"}
    target = targets.get(str(data.get("fixed_task_key") or ""))
    if not target:
        return
    ranks = {"新案待分配": 0, "文书准备": 1, "一审立案受理": 2, "一审准备开庭": 3, "待上诉": 4, "二审": 5, "执行": 6}
    if ranks.get(target, -1) <= ranks.get(case_record.status, -1):
        return
    previous = case_record.status; case_record.status = target
    case_record.data = {**(case_record.data or {}), "stage_advanced_by_task_id": task.id, "stage_advanced_at": datetime.now().isoformat(timespec="seconds"), "business_stage": "审理" if target == "一审准备开庭" else "立案"}
    db.add(WorkflowEvent(record_id=case_record.id, action="固定任务验收自动推进阶段", from_status=previous, to_status=target, operator=operator, comment=f"任务 {task.serial_no}：{task.title}"))


FINANCE_FEE_TYPES = {"官方费用", "代理费", "其他费用", "内部费用", "结算费用", "预损费用", "归档费用"}
EXPENSE_SUBTYPE_FEE_TYPE = {"官费": "官方费用", "第三方费用": "其他费用", "代理费": "代理费", "其他费用": "其他费用", "内部费用": "内部费用"}
EXPENSE_SCOPE_FEE_TYPES = {"律所": {"官方费用", "代理费", "其他费用"}, "平台": {"官方费用", "代理费", "其他费用"}, "内部": {"内部费用"}}
FINANCE_TRANSACTION_TYPES = {"付款", "开票", "回款", "退费"}
FINANCE_VOUCHER_CATEGORIES = {"付款凭证", "发票扫描件", "回款凭证", "退费凭证"}
FINANCE_DEFAULT_VOUCHER_CATEGORY = {"付款": "付款凭证", "开票": "发票扫描件", "回款": "回款凭证", "退费": "退费凭证"}


def _round_fee_amount(value: float) -> float:
    return float(Decimal(str(value)).quantize(Decimal("0.01"), rounding=ROUND_UP))


async def _finance_linked_case(case_no: str, identity: dict, db: AsyncSession) -> BusinessRecord | None:
    case_no = case_no.strip()
    if not case_no:
        return None
    item = await db.scalar(select(BusinessRecord).where(BusinessRecord.module == "case", BusinessRecord.serial_no == case_no, *(await _record_scope_conditions(identity, db))))
    if not item:
        raise HTTPException(status_code=404, detail="关联案件不存在或无权访问")
    return item


def _invoice_original_type(value: object) -> str:
    text = str(value or "").strip()
    return "专票" if "专用" in text or text == "专票" else "普票"


async def _invoice_list_rows(
    identity: dict,
    db: AsyncSession,
    *,
    scope: str,
    customer: str,
    application_no: str,
    invoice_type: str,
    invoice_title: str,
    invoice_no: str,
    invoice_status: str,
    invoiced_from: date | None,
    invoiced_to: date | None,
    case_no: str,
    applicant_filter: str = "",
    ids: set[int] | None = None,
) -> list[dict]:
    records = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "invoice",
        *(await _record_scope_conditions(identity, db)),
    ).order_by(BusinessRecord.updated_at.desc(), BusinessRecord.id.desc()))).all())
    if ids is not None:
        records = [item for item in records if item.id in ids]
    allowed_fields = await _allowed_field_keys(identity, db)
    personal_names = {str(identity.get("username", "")).strip(), str(identity.get("display_name", "")).strip()} - {""}
    user_display_names = {
        username: (display_name or username)
        for username, display_name in (await db.execute(select(User.username, User.display_name))).all()
    }

    def contains(value: object, needle: str) -> bool:
        return not needle.strip() or needle.strip().casefold() in str(value or "").casefold()

    rows: list[dict] = []
    for item in records:
        data = item.data or {}
        applicant_account = str(data.get("applicant") or item.owner or "").strip()
        applicant = user_display_names.get(applicant_account, applicant_account)
        if scope == "mine" and not ({applicant_account, applicant} & personal_names):
            continue
        if scope == "pending" and item.status != "待开票":
            continue
        if not contains(applicant, applicant_filter):
            continue
        if not contains(item.customer, customer) or not contains(item.serial_no, application_no):
            continue
        display_type = _invoice_original_type(data.get("invoice_type"))
        if invoice_type and display_type != invoice_type:
            continue
        if not contains(data.get("invoice_title"), invoice_title) or not contains(data.get("invoice_no"), invoice_no):
            continue
        if invoice_status and item.status != invoice_status:
            continue
        if not contains(data.get("case_no"), case_no):
            continue
        invoice_date_text = str(data.get("invoice_date") or "").strip()
        parsed_invoice_date = None
        if invoice_date_text:
            try:
                parsed_invoice_date = date.fromisoformat(invoice_date_text[:10])
            except ValueError:
                pass
        if invoiced_from and (not parsed_invoice_date or parsed_invoice_date < invoiced_from):
            continue
        if invoiced_to and (not parsed_invoice_date or parsed_invoice_date > invoiced_to):
            continue
        result = _record_dict(item, allowed_fields)
        result_data = dict(result.get("data") or {})
        result_data.update({
            "invoice_type_display": display_type,
            "applicant": applicant,
            "application_date": str(data.get("application_date") or item.created_at),
            "extra_amount": float(data.get("extra_amount", 0) or 0) if "finance.amount" in allowed_fields else None,
        })
        result["data"] = result_data
        rows.append(result)
    return rows


@app.get(f"{settings.api_prefix}/finance/invoices")
async def list_invoice_applications(
    scope: str = Query("company", pattern="^(mine|company|pending)$"),
    customer: str = "", application_no: str = "", invoice_type: str = Query("", pattern="^(|普票|专票)$"),
    invoice_title: str = "", invoice_no: str = "", invoice_status: str = "",
    invoiced_from: date | None = None, invoiced_to: date | None = None, case_no: str = "", applicant: str = "",
    page: int = Query(1, ge=1), page_size: int = Query(15, ge=1, le=200),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    if invoiced_from and invoiced_to and invoiced_from > invoiced_to:
        raise HTTPException(status_code=422, detail="开票开始日期不能晚于结束日期")
    if scope == "pending" and identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以处理开票")
    rows = await _invoice_list_rows(identity, db, scope=scope, customer=customer, application_no=application_no, invoice_type=invoice_type, invoice_title=invoice_title, invoice_no=invoice_no, invoice_status=invoice_status, invoiced_from=invoiced_from, invoiced_to=invoiced_to, case_no=case_no, applicant_filter=applicant)
    total_amount = round(sum(float((row.get("data") or {}).get("amount", 0) or 0) for row in rows if row.get("status") not in {"已撤回", "已作废"}), 2)
    total_extra_amount = round(sum(float((row.get("data") or {}).get("extra_amount", 0) or 0) for row in rows if row.get("status") not in {"已撤回", "已作废"}), 2)
    start = (page - 1) * page_size
    return {"items": rows[start:start + page_size], "total": len(rows), "total_amount": total_amount, "total_extra_amount": total_extra_amount, "page": page, "page_size": page_size}


@app.get(f"{settings.api_prefix}/finance/invoices/export")
async def export_invoice_applications(
    scope: str = Query("company", pattern="^(mine|company|pending)$"), ids: str = "",
    customer: str = "", application_no: str = "", invoice_type: str = Query("", pattern="^(|普票|专票)$"),
    invoice_title: str = "", invoice_no: str = "", invoice_status: str = "",
    invoiced_from: date | None = None, invoiced_to: date | None = None, case_no: str = "", applicant: str = "",
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    selected_ids = set(_export_ids(ids)) if ids.strip() else None
    if scope == "pending" and identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以处理开票")
    rows = await _invoice_list_rows(identity, db, scope=scope, customer=customer, application_no=application_no, invoice_type=invoice_type, invoice_title=invoice_title, invoice_no=invoice_no, invoice_status=invoice_status, invoiced_from=invoiced_from, invoiced_to=invoiced_to, case_no=case_no, applicant_filter=applicant, ids=selected_ids)
    if selected_ids is not None and not rows:
        raise HTTPException(status_code=422, detail="请选择需要导出的发票")
    headers = (
        ["请票单号", "申请人", "客户名称", "开票金额", "高开金额", "开票抬头", "备注"]
        if scope == "pending"
        else ["请票单号", "客户名称", "开票金额", "高开金额", "开票抬头", "发票号码", "申请人", "领票人", "开票日期", "状态"]
        if scope == "company"
        else ["请票单号", "客户名称", "开票金额", "高开金额", "发票编号", "领票人", "开票日期", "票据状态", "备注"]
    )
    def cell(value: object, *, number: bool = False) -> str:
        text_value = f"{float(value or 0):.2f}" if number else str(value or "")
        return f'<Cell><Data ss:Type="{"Number" if number else "String"}">{xml_escape(text_value)}</Data></Cell>'
    sheet_rows = ["<Row>" + "".join(cell(value) for value in headers) + "</Row>"]
    for row in rows:
        data = row.get("data") or {}
        values = (
            [row.get("serial_no"), data.get("applicant"), row.get("customer"), data.get("amount"), data.get("extra_amount"), data.get("invoice_title"), data.get("remark") or row.get("description")]
            if scope == "pending"
            else [row.get("serial_no"), row.get("customer"), data.get("amount"), data.get("extra_amount"), data.get("invoice_title"), data.get("invoice_no"), data.get("applicant"), data.get("recipient"), data.get("invoice_date"), row.get("status")]
            if scope == "company"
            else [row.get("serial_no"), row.get("customer"), data.get("amount"), data.get("extra_amount"), data.get("invoice_no"), data.get("recipient"), data.get("invoice_date"), row.get("status"), data.get("remark") or row.get("description")]
        )
        number_indexes = {3, 4} if scope == "pending" else {2, 3}
        sheet_rows.append("<Row>" + "".join(cell(value, number=index in number_indexes) for index, value in enumerate(values)) + "</Row>")
    sheet_name = "待处理开票" if scope == "pending" else "公司开票" if scope == "company" else "我的开票"
    workbook = '<?xml version="1.0" encoding="UTF-8"?><?mso-application progid="Excel.Sheet"?><Workbook xmlns="urn:schemas-microsoft-com:office:spreadsheet" xmlns:ss="urn:schemas-microsoft-com:office:spreadsheet"><Worksheet ss:Name="' + sheet_name + '"><Table>' + "".join(sheet_rows) + "</Table></Worksheet></Workbook>"
    filename = f"{sheet_name}-{date.today()}.xls"
    disposition = f"attachment; filename=my-invoices.xls; filename*=UTF-8''{quote(filename)}"
    return Response(content=workbook.encode("utf-8"), media_type="application/vnd.ms-excel", headers={"Content-Disposition": disposition})


def _case_fee_display_type(item: BusinessRecord) -> str:
    data = item.data or {}
    explicit = str(data.get("fee_type_name") or data.get("case_fee_type_name") or "").strip()
    if explicit:
        return explicit
    fee_type = str(data.get("fee_type") or "").strip()
    if fee_type == "代理费":
        return "律师代理费"
    if fee_type == "官方费用":
        return "官费"
    return fee_type or item.title


def _case_fee_date(value: object) -> date | None:
    text = str(value or "").strip()
    if not text:
        return None
    try:
        return date.fromisoformat(text[:10])
    except ValueError:
        return None


async def _invoice_case_fee_rows(
    identity: dict,
    db: AsyncSession,
    *,
    scope: str,
    case_no: str,
    court_case_no: str,
    notary_no: str,
    invoice_amount_from: float | None,
    invoice_amount_to: float | None,
    customer: str,
    paid_organization: str,
    invoice_status: str,
    invoice_from: date | None,
    invoice_to: date | None,
    hearing_lawyer: str,
    assistant: str,
    case_stages: str,
    paid_from: date | None,
    paid_to: date | None,
    fee_types: str,
    payer_name: str,
    cashed_from: date | None,
    cashed_to: date | None,
    ids: set[int] | None = None,
    include_all_fee_types: bool = False,
) -> list[dict]:
    scope_conditions = await _record_scope_conditions(identity, db)
    fee_conditions = list(scope_conditions)
    if scope == "mine":
        fee_conditions.append(BusinessRecord.owner == identity["username"])
    fees = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "finance", *fee_conditions
    ).order_by(BusinessRecord.updated_at.desc(), BusinessRecord.id.desc()))).all())
    if not include_all_fee_types:
        fees = [
            item for item in fees
            if str((item.data or {}).get("fee_type") or "")
            not in {"结算费用", "预损费用", "归档费用"}
        ]
    if ids is not None:
        fees = [item for item in fees if item.id in ids]
    if not fees:
        return []

    case_ids = {int((item.data or {}).get("case_id") or 0) for item in fees if (item.data or {}).get("case_id")}
    case_nos = {str((item.data or {}).get("case_no") or "") for item in fees if (item.data or {}).get("case_no")}
    case_conditions = []
    if case_ids:
        case_conditions.append(BusinessRecord.id.in_(case_ids))
    if case_nos:
        case_conditions.append(BusinessRecord.serial_no.in_(case_nos))
    cases = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "case",
        or_(*case_conditions) if case_conditions else false(),
        *scope_conditions,
    ))).all())
    cases_by_id = {item.id: item for item in cases}
    cases_by_no = {item.serial_no: item for item in cases}

    fees_by_case: dict[str, list[BusinessRecord]] = {}
    for item in fees:
        data = item.data or {}
        linked_case = cases_by_id.get(int(data.get("case_id") or 0)) or cases_by_no.get(str(data.get("case_no") or ""))
        linked_no = linked_case.serial_no if linked_case else str(data.get("case_no") or "")
        if linked_no:
            fees_by_case.setdefault(linked_no, []).append(item)

    invoices = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "invoice", BusinessRecord.status == "已开票", *scope_conditions
    ).order_by(BusinessRecord.updated_at.desc(), BusinessRecord.id.desc()))).all())
    invoice_by_fee: dict[int, list[tuple[BusinessRecord, float]]] = {}
    fee_ids = {item.id for item in fees}
    for invoice in invoices:
        data = invoice.data or {}
        explicit_ids = []
        raw_case_fee_ids = list(data.get("case_fee_ids") or [])
        has_explicit_link = bool(raw_case_fee_ids or data.get("case_fee_id"))
        for value in raw_case_fee_ids:
            try:
                fee_id = int(value)
            except (TypeError, ValueError):
                continue
            if fee_id in fee_ids and fee_id not in explicit_ids:
                explicit_ids.append(fee_id)
        if not explicit_ids and data.get("case_fee_id"):
            try:
                fee_id = int(data.get("case_fee_id"))
                if fee_id in fee_ids:
                    explicit_ids = [fee_id]
            except (TypeError, ValueError):
                pass
        if not explicit_ids and not has_explicit_link:
            candidates = fees_by_case.get(str(data.get("case_no") or ""), [])
            if len(candidates) == 1:
                explicit_ids = [candidates[0].id]
        if not explicit_ids:
            continue
        allocation_map: dict[int, float] = {}
        for allocation in data.get("case_fee_allocations") or []:
            try:
                allocation_map[int(allocation.get("fee_id"))] = float(allocation.get("amount") or 0)
            except (AttributeError, TypeError, ValueError):
                continue
        invoice_amount = float(data.get("amount") or 0)
        if not allocation_map:
            if len(explicit_ids) == 1:
                allocation_map[explicit_ids[0]] = invoice_amount
            else:
                weights = [abs(float((next(item for item in fees if item.id == fee_id).data or {}).get("amount") or 0)) for fee_id in explicit_ids]
                weight_total = sum(weights) or float(len(explicit_ids))
                allocated = 0.0
                for index, fee_id in enumerate(explicit_ids):
                    amount = invoice_amount - allocated if index == len(explicit_ids) - 1 else round(invoice_amount * (weights[index] or 1) / weight_total, 2)
                    allocation_map[fee_id] = amount
                    allocated += amount
        for fee_id in explicit_ids:
            invoice_by_fee.setdefault(fee_id, []).append((invoice, float(allocation_map.get(fee_id, 0))))

    transactions = list((await db.scalars(select(FinanceTransaction).where(
        FinanceTransaction.finance_record_id.in_(list(fee_ids))
    ).order_by(FinanceTransaction.transaction_date.desc(), FinanceTransaction.id.desc()))).all())
    payments_by_fee: dict[int, list[FinanceTransaction]] = {}
    for transaction in transactions:
        if transaction.transaction_type == "付款" and transaction.finance_record_id:
            payments_by_fee.setdefault(transaction.finance_record_id, []).append(transaction)

    incoming = list((await db.scalars(select(IncomingPayment).order_by(
        IncomingPayment.received_date.desc(), IncomingPayment.id.desc()
    ))).all())
    receipts_by_fee: dict[int, list[tuple[IncomingPayment, float]]] = {}
    for payment in incoming:
        for allocation in payment.allocations or []:
            if not isinstance(allocation, dict):
                continue
            linked_nested = False
            for settlement_item in allocation.get("settlement_items") or []:
                if not isinstance(settlement_item, dict):
                    continue
                try:
                    nested_fee_id = int(settlement_item.get("fee_record_id") or settlement_item.get("fee_id") or 0)
                except (TypeError, ValueError):
                    nested_fee_id = 0
                if nested_fee_id in fee_ids:
                    nested_amount = float(settlement_item.get("amount") or settlement_item.get("settlement_amount") or 0)
                    receipts_by_fee.setdefault(nested_fee_id, []).append((payment, nested_amount))
                    linked_nested = True
            if linked_nested:
                continue
            fee_id = 0
            try:
                fee_id = int(allocation.get("fee_id") or allocation.get("finance_record_id") or 0)
            except (AttributeError, TypeError, ValueError):
                pass
            if fee_id not in fee_ids:
                candidates = fees_by_case.get(str(allocation.get("case_no") or ""), []) if isinstance(allocation, dict) else []
                if len(candidates) == 1:
                    fee_id = candidates[0].id
            if fee_id in fee_ids:
                receipts_by_fee.setdefault(fee_id, []).append((payment, float(allocation.get("amount") or 0)))

    refunds = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "refund", *scope_conditions
    ).order_by(BusinessRecord.updated_at.desc(), BusinessRecord.id.desc()))).all())
    refunds_by_fee: dict[int, list[BusinessRecord]] = {}
    official_by_case: dict[str, list[BusinessRecord]] = {}
    fee_by_document_no: dict[str, BusinessRecord] = {}
    for fee in fees:
        fee_data = fee.data or {}
        if str(fee_data.get("fee_type") or "") == "官方费用":
            official_by_case.setdefault(str(fee_data.get("case_no") or ""), []).append(fee)
        document_no = str(fee_data.get("document_no") or "").strip()
        if document_no:
            fee_by_document_no[document_no] = fee
    for refund in refunds:
        refund_data = refund.data or {}
        linked_fee: BusinessRecord | None = None
        try:
            linked_fee_id = int(refund_data.get("fee_record_id") or 0)
        except (TypeError, ValueError):
            linked_fee_id = 0
        if linked_fee_id in fee_ids:
            linked_fee = next((fee for fee in fees if fee.id == linked_fee_id), None)
        if linked_fee is None:
            linked_fee = fee_by_document_no.get(str(refund_data.get("original_payment_no") or "").strip())
        if linked_fee is None:
            candidates = official_by_case.get(str(refund_data.get("case_no") or ""), [])
            if len(candidates) == 1:
                linked_fee = candidates[0]
        if linked_fee:
            refunds_by_fee.setdefault(linked_fee.id, []).append(refund)

    allowed_fields = await _allowed_field_keys(identity, db)
    show_amount = "finance.amount" in allowed_fields
    selected_stages = {value.strip() for value in case_stages.split(",") if value.strip()}
    selected_types = {value.strip() for value in fee_types.split(",") if value.strip()}
    normalized_types = {"代理费" if value == "律师代理费" else "官方费用" if value == "官费" else value for value in selected_types}

    def contains(value: object, needle: str) -> bool:
        return not needle.strip() or needle.strip().casefold() in str(value or "").casefold()

    rows: list[dict] = []
    for item in fees:
        data = item.data or {}
        linked_case = cases_by_id.get(int(data.get("case_id") or 0)) or cases_by_no.get(str(data.get("case_no") or ""))
        case_data = (linked_case.data or {}) if linked_case else {}
        linked_invoices = invoice_by_fee.get(item.id, [])
        invoiced_amount = round(sum(amount for _, amount in linked_invoices), 2)
        latest_invoice = linked_invoices[0][0] if linked_invoices else None
        latest_invoice_data = (latest_invoice.data or {}) if latest_invoice else {}
        invoice_date = _case_fee_date(latest_invoice_data.get("invoice_date"))
        payments = payments_by_fee.get(item.id, [])
        paid_amount = round(sum(float(tx.amount or 0) for tx in payments), 2)
        paid_date = payments[0].transaction_date if payments else _case_fee_date(data.get("paid_date") or data.get("payment_date"))
        receipts = receipts_by_fee.get(item.id, [])
        cashed_amount = round(sum(amount for _, amount in receipts), 2)
        latest_receipt = receipts[0][0] if receipts else None
        cashed_date = latest_receipt.received_date if latest_receipt else _case_fee_date(data.get("cashed_date") or data.get("receipt_date"))
        received_payer = latest_receipt.payer_name if latest_receipt else str(data.get("received_payer_name") or data.get("payer_name") or "")
        linked_refunds = refunds_by_fee.get(item.id, [])
        refund_requested_amount = round(sum(float((refund.data or {}).get("amount") or 0) for refund in linked_refunds if refund.status not in {"已驳回", "已作废"}), 2)
        refunded_amount = round(sum(float((refund.data or {}).get("amount") or 0) for refund in linked_refunds if refund.status == "已退款"), 2)
        fee_amount = float(data.get("amount") or 0)
        display_type = _case_fee_display_type(item)
        base_type = str(data.get("fee_type") or "")
        display_status = "已开票" if invoiced_amount > 0 else "未开票"
        case_stage = str(data.get("case_stage") or case_data.get("case_stage") or (linked_case.status if linked_case else ""))
        assistant_name = str(data.get("assistant") or data.get("lawyer_assistant") or case_data.get("assistant") or case_data.get("lawyer_assistant") or "")
        hearing_name = str(data.get("hearing_lawyer") or case_data.get("hearing_lawyer") or case_data.get("court_lawyer") or "")
        court_no = str(data.get("court_case_no") or case_data.get("court_case_no") or case_data.get("first_instance_case_no") or case_data.get("official_no") or "")
        certificate_no = str(data.get("certificate_no") or data.get("notary_no") or case_data.get("certificate_no") or case_data.get("notary_no") or "")
        court_name = str(data.get("court_name") or data.get("court") or case_data.get("court_name") or case_data.get("first_instance_court") or "")
        paid_org = str(data.get("paid_organization") or data.get("payee") or court_name)
        display_payment_status = str(data.get("payment_status") or "").strip()
        if not display_payment_status:
            if str(data.get("writeoff_status") or "") == "待核销":
                display_payment_status = "待核销"
            else:
                display_payment_status = {
                    "草稿": "创建待提交",
                    "待审批": "待审批",
                    "已审批": "待付款",
                    "部分付款": "待付款",
                    "已付款": "已付款",
                    "已退回": "已驳回",
                    "已驳回": "已驳回",
                    "已拒绝": "已驳回",
                    "已作废": "已作废",
                }.get(item.status, item.status)
        if not contains(data.get("case_no") or (linked_case.serial_no if linked_case else ""), case_no):
            continue
        if not contains(court_no, court_case_no) or not contains(certificate_no, notary_no):
            continue
        if invoice_amount_from is not None and invoiced_amount < invoice_amount_from:
            continue
        if invoice_amount_to is not None and invoiced_amount > invoice_amount_to:
            continue
        if not contains(item.customer or (linked_case.customer if linked_case else ""), customer) or not contains(paid_org, paid_organization):
            continue
        if invoice_status and display_status != invoice_status:
            continue
        if invoice_from and (not invoice_date or invoice_date < invoice_from):
            continue
        if invoice_to and (not invoice_date or invoice_date > invoice_to):
            continue
        if not contains(hearing_name, hearing_lawyer) or not contains(assistant_name, assistant):
            continue
        if selected_stages and case_stage not in selected_stages:
            continue
        if paid_from and (not paid_date or paid_date < paid_from):
            continue
        if paid_to and (not paid_date or paid_date > paid_to):
            continue
        if normalized_types and base_type not in normalized_types and display_type not in selected_types:
            continue
        if not contains(received_payer, payer_name):
            continue
        if cashed_from and (not cashed_date or cashed_date < cashed_from):
            continue
        if cashed_to and (not cashed_date or cashed_date > cashed_to):
            continue
        result = _record_dict(item, allowed_fields)
        result_data = dict(result.get("data") or {})
        result_data.update({
            "case_id": linked_case.id if linked_case else data.get("case_id"),
            "case_no": linked_case.serial_no if linked_case else data.get("case_no", ""),
            "case_stage": case_stage,
            "assistant": assistant_name,
            "hearing_lawyer": hearing_name,
            "court_case_no": court_no,
            "certificate_no": certificate_no,
            "fee_type": display_type,
            "base_fee_type": base_type,
            "amount": fee_amount if show_amount else None,
            "invoice_status": display_status,
            "invoice_date": invoice_date.isoformat() if invoice_date else "",
            "invoice_amount": invoiced_amount if show_amount else None,
            "invoice_no": latest_invoice_data.get("invoice_no", ""),
            "invoice_record_id": latest_invoice.id if latest_invoice else None,
            "cashed_date": cashed_date.isoformat() if cashed_date else "",
            "cashed_amount": cashed_amount if show_amount else None,
            "received_payer_name": received_payer,
            "paid_date": paid_date.isoformat() if paid_date else "",
            "paid_amount": paid_amount if show_amount else None,
            "refund_requested_amount": refund_requested_amount if show_amount else None,
            "refunded_amount": refunded_amount if show_amount else None,
            "court_name": court_name,
            "payment_status": display_payment_status,
            "paid_organization": paid_org,
            "contract_no": data.get("contract_no") or case_data.get("contract_no") or "",
        })
        result["data"] = result_data
        rows.append(result)
    return rows


@app.get(f"{settings.api_prefix}/finance/case-fees/invoice-status")
async def list_invoice_case_fees(
    scope: str = Query("mine", pattern="^(mine|company)$"),
    case_no: str = "", court_case_no: str = "", notary_no: str = "",
    invoice_amount_from: float | None = None, invoice_amount_to: float | None = None,
    customer: str = "", paid_organization: str = "",
    invoice_status: str = Query("未开票", pattern="^(未开票|已开票)$"),
    invoice_from: date | None = None, invoice_to: date | None = None,
    hearing_lawyer: str = "", assistant: str = "", case_stages: str = "",
    paid_from: date | None = None, paid_to: date | None = None,
    fee_types: str = "律师代理费", payer_name: str = "",
    cashed_from: date | None = None, cashed_to: date | None = None,
    page: int = Query(1, ge=1), page_size: int = Query(15, ge=1, le=200),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    if invoice_amount_from is not None and invoice_amount_to is not None and invoice_amount_from > invoice_amount_to:
        raise HTTPException(status_code=422, detail="开票最小金额不能大于最大金额")
    for start, end, label in [
        (invoice_from, invoice_to, "开票"), (paid_from, paid_to, "付款"), (cashed_from, cashed_to, "到账")
    ]:
        if start and end and start > end:
            raise HTTPException(status_code=422, detail=f"{label}开始日期不能晚于结束日期")
    rows = await _invoice_case_fee_rows(
        identity, db, scope=scope, case_no=case_no, court_case_no=court_case_no,
        notary_no=notary_no, invoice_amount_from=invoice_amount_from,
        invoice_amount_to=invoice_amount_to, customer=customer,
        paid_organization=paid_organization, invoice_status=invoice_status,
        invoice_from=invoice_from, invoice_to=invoice_to,
        hearing_lawyer=hearing_lawyer, assistant=assistant,
        case_stages=case_stages, paid_from=paid_from, paid_to=paid_to,
        fee_types=fee_types, payer_name=payer_name, cashed_from=cashed_from,
        cashed_to=cashed_to,
    )
    totals = {
        "amount": round(sum(float((row.get("data") or {}).get("amount") or 0) for row in rows), 2),
        "invoice_amount": round(sum(float((row.get("data") or {}).get("invoice_amount") or 0) for row in rows), 2),
        "cashed_amount": round(sum(float((row.get("data") or {}).get("cashed_amount") or 0) for row in rows), 2),
        "paid_amount": round(sum(float((row.get("data") or {}).get("paid_amount") or 0) for row in rows), 2),
    }
    start = (page - 1) * page_size
    return {"items": rows[start:start + page_size], "total": len(rows), "totals": totals, "page": page, "page_size": page_size}


@app.get(f"{settings.api_prefix}/finance/case-fees/invoice-status/export")
async def export_invoice_case_fees(
    scope: str = Query("mine", pattern="^(mine|company)$"), ids: str = "",
    case_no: str = "", court_case_no: str = "", notary_no: str = "",
    invoice_amount_from: float | None = None, invoice_amount_to: float | None = None,
    customer: str = "", paid_organization: str = "",
    invoice_status: str = Query("未开票", pattern="^(未开票|已开票)$"),
    invoice_from: date | None = None, invoice_to: date | None = None,
    hearing_lawyer: str = "", assistant: str = "", case_stages: str = "",
    paid_from: date | None = None, paid_to: date | None = None,
    fee_types: str = "律师代理费", payer_name: str = "",
    cashed_from: date | None = None, cashed_to: date | None = None,
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    selected_ids = set(_export_ids(ids)) if ids.strip() else None
    rows = await _invoice_case_fee_rows(
        identity, db, scope=scope, case_no=case_no, court_case_no=court_case_no,
        notary_no=notary_no, invoice_amount_from=invoice_amount_from,
        invoice_amount_to=invoice_amount_to, customer=customer,
        paid_organization=paid_organization, invoice_status=invoice_status,
        invoice_from=invoice_from, invoice_to=invoice_to,
        hearing_lawyer=hearing_lawyer, assistant=assistant,
        case_stages=case_stages, paid_from=paid_from, paid_to=paid_to,
        fee_types=fee_types, payer_name=payer_name, cashed_from=cashed_from,
        cashed_to=cashed_to, ids=selected_ids,
    )
    if selected_ids is not None and not rows:
        raise HTTPException(status_code=422, detail="请选择需要导出的费用.")
    headers = ["案号", "客户", "案件阶段", "助理", "开庭律师", "法院案号", "费用类型", "金额", "开票日期", "开票金额", "发票查看", "到账时间", "到账金额", "到账单位", "付款时间", "付款金额", "法院名称", "付款状态", "合同号"]
    keys = ["case_no", "customer", "case_stage", "assistant", "hearing_lawyer", "court_case_no", "fee_type", "amount", "invoice_date", "invoice_amount", "invoice_no", "cashed_date", "cashed_amount", "received_payer_name", "paid_date", "paid_amount", "court_name", "payment_status", "contract_no"]
    number_keys = {"amount", "invoice_amount", "cashed_amount", "paid_amount"}
    def cell(value: object, *, number: bool = False) -> str:
        text_value = f"{float(value or 0):.2f}" if number else str(value or "")
        return f'<Cell><Data ss:Type="{"Number" if number else "String"}">{xml_escape(text_value)}</Data></Cell>'
    sheet_rows = ["<Row>" + "".join(cell(value) for value in headers) + "</Row>"]
    for row in rows:
        data = row.get("data") or {}
        values = {**data, "customer": row.get("customer", "")}
        sheet_rows.append("<Row>" + "".join(cell(values.get(key), number=key in number_keys) for key in keys) + "</Row>")
    workbook = '<?xml version="1.0" encoding="UTF-8"?><?mso-application progid="Excel.Sheet"?><Workbook xmlns="urn:schemas-microsoft-com:office:spreadsheet" xmlns:ss="urn:schemas-microsoft-com:office:spreadsheet"><Worksheet ss:Name="未开票"><Table>' + "".join(sheet_rows) + "</Table></Worksheet></Workbook>"
    filename = f"{'公司未开票' if scope == 'company' else '未开票'}-{date.today()}.xls"
    disposition = f"attachment; filename=invoice-case-fees.xls; filename*=UTF-8''{quote(filename)}"
    return Response(content=workbook.encode("utf-8"), media_type="application/vnd.ms-excel", headers={"Content-Disposition": disposition})


async def _fee_query_rows(
    identity: dict, db: AsyncSession, *,
    case_no: str = "", court_case_no: str = "", notary_no: str = "",
    refund_amount_from: float | None = None, refund_amount_to: float | None = None,
    customer: str = "", paid_organization: str = "", payment_status: str = "",
    paid_from: date | None = None, paid_to: date | None = None,
    hearing_lawyer: str = "", assistant: str = "", case_stages: str = "",
    fee_types: str = "", ids: set[int] | None = None,
) -> list[dict]:
    rows = await _invoice_case_fee_rows(
        identity, db, scope="company", case_no=case_no,
        court_case_no=court_case_no, notary_no=notary_no,
        invoice_amount_from=None, invoice_amount_to=None, customer=customer,
        paid_organization=paid_organization, invoice_status="",
        invoice_from=None, invoice_to=None, hearing_lawyer=hearing_lawyer,
        assistant=assistant, case_stages=case_stages, paid_from=paid_from,
        paid_to=paid_to, fee_types=fee_types, payer_name="",
        cashed_from=None, cashed_to=None, ids=ids, include_all_fee_types=True,
    )
    filtered: list[dict] = []
    for row in rows:
        data = row.get("data") or {}
        refund_amount = data.get("refund_requested_amount")
        if refund_amount_from is not None and (refund_amount is None or float(refund_amount) < refund_amount_from):
            continue
        if refund_amount_to is not None and (refund_amount is None or float(refund_amount) > refund_amount_to):
            continue
        if payment_status and str(data.get("payment_status") or "") != payment_status:
            continue
        filtered.append(row)
    return filtered


@app.get(f"{settings.api_prefix}/finance/fees/query")
async def query_finance_fees(
    case_no: str = "", court_case_no: str = "", notary_no: str = "",
    refund_amount_from: float | None = None, refund_amount_to: float | None = None,
    customer: str = "", paid_organization: str = "",
    payment_status: str = Query("", pattern="^(|创建待提交|待审批|待付款|待核销|已付款|已驳回|已作废)$"),
    paid_from: date | None = None, paid_to: date | None = None,
    hearing_lawyer: str = "", assistant: str = "", case_stages: str = "", fee_types: str = "",
    page: int = Query(1, ge=1), page_size: int = Query(15, ge=1, le=200),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    if refund_amount_from is not None and refund_amount_to is not None and refund_amount_from > refund_amount_to:
        raise HTTPException(status_code=422, detail="退费最小金额不能大于最大金额")
    if paid_from and paid_to and paid_from > paid_to:
        raise HTTPException(status_code=422, detail="付款开始日期不能晚于结束日期")
    rows = await _fee_query_rows(
        identity, db, case_no=case_no, court_case_no=court_case_no,
        notary_no=notary_no, refund_amount_from=refund_amount_from,
        refund_amount_to=refund_amount_to, customer=customer,
        paid_organization=paid_organization, payment_status=payment_status,
        paid_from=paid_from, paid_to=paid_to, hearing_lawyer=hearing_lawyer,
        assistant=assistant, case_stages=case_stages, fee_types=fee_types,
    )
    amount_visible = "finance.amount" in await _allowed_field_keys(identity, db)
    totals = {
        key: round(sum(float((row.get("data") or {}).get(key) or 0) for row in rows), 2) if amount_visible else None
        for key in ("amount", "refund_requested_amount", "refunded_amount", "cashed_amount", "paid_amount")
    }
    start = (page - 1) * page_size
    return {"items": rows[start:start + page_size], "total": len(rows), "totals": totals, "page": page, "page_size": page_size}


@app.get(f"{settings.api_prefix}/finance/fees/query/export")
async def export_finance_fee_query(
    ids: str = "", selected_only: bool = False,
    case_no: str = "", court_case_no: str = "", notary_no: str = "",
    refund_amount_from: float | None = None, refund_amount_to: float | None = None,
    customer: str = "", paid_organization: str = "", payment_status: str = "",
    paid_from: date | None = None, paid_to: date | None = None,
    hearing_lawyer: str = "", assistant: str = "", case_stages: str = "", fee_types: str = "",
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    if refund_amount_from is not None and refund_amount_to is not None and refund_amount_from > refund_amount_to:
        raise HTTPException(status_code=422, detail="退费最小金额不能大于最大金额")
    if paid_from and paid_to and paid_from > paid_to:
        raise HTTPException(status_code=422, detail="付款开始日期不能晚于结束日期")
    if payment_status not in {"", "创建待提交", "待审批", "待付款", "待核销", "已付款", "已驳回", "已作废"}:
        raise HTTPException(status_code=422, detail="付款状态无效")
    selected_ids = set(_export_ids(ids)) if ids.strip() else None
    if selected_only and not selected_ids:
        raise HTTPException(status_code=422, detail="请选择需要导出的费用.")
    rows = await _fee_query_rows(
        identity, db, case_no=case_no, court_case_no=court_case_no,
        notary_no=notary_no, refund_amount_from=refund_amount_from,
        refund_amount_to=refund_amount_to, customer=customer,
        paid_organization=paid_organization, payment_status=payment_status,
        paid_from=paid_from, paid_to=paid_to, hearing_lawyer=hearing_lawyer,
        assistant=assistant, case_stages=case_stages, fee_types=fee_types,
        ids=selected_ids,
    )
    if selected_ids is not None and len(rows) != len(selected_ids):
        raise HTTPException(status_code=422, detail="部分费用不存在或无权导出")
    if not rows:
        raise HTTPException(status_code=422, detail="当前没有可导出的费用")
    headers = ["案号", "客户", "案件阶段", "助理", "开庭律师", "法院案号", "费用类型", "金额", "退费金额", "已退金额", "到账时间", "到账金额", "付款时间", "付款金额", "法院名称", "付款状态"]
    keys = ["case_no", "customer", "case_stage", "assistant", "hearing_lawyer", "court_case_no", "fee_type", "amount", "refund_requested_amount", "refunded_amount", "cashed_date", "cashed_amount", "paid_date", "paid_amount", "court_name", "payment_status"]
    number_keys = {"amount", "refund_requested_amount", "refunded_amount", "cashed_amount", "paid_amount"}
    def cell(value: object, *, number: bool = False) -> str:
        text_value = f"{float(value):.2f}" if number and value is not None else str(value or "")
        return f'<Cell><Data ss:Type="{"Number" if number and value is not None else "String"}">{xml_escape(text_value)}</Data></Cell>'
    sheet_rows = ["<Row>" + "".join(cell(value) for value in headers) + "</Row>"]
    for row in rows:
        data = row.get("data") or {}
        values = {**data, "customer": row.get("customer", "")}
        sheet_rows.append("<Row>" + "".join(cell(values.get(key), number=key in number_keys) for key in keys) + "</Row>")
    workbook = '<?xml version="1.0" encoding="UTF-8"?><?mso-application progid="Excel.Sheet"?><Workbook xmlns="urn:schemas-microsoft-com:office:spreadsheet" xmlns:ss="urn:schemas-microsoft-com:office:spreadsheet"><Worksheet ss:Name="费用查询"><Table>' + "".join(sheet_rows) + "</Table></Worksheet></Workbook>"
    filename = f"费用查询-{date.today()}.xls"
    disposition = f"attachment; filename=finance-fee-query.xls; filename*=UTF-8''{quote(filename)}"
    return Response(content=workbook.encode("utf-8"), media_type="application/vnd.ms-excel", headers={"Content-Disposition": disposition})


@app.post(f"{settings.api_prefix}/finance/invoices", status_code=status.HTTP_201_CREATED)
async def create_invoice_application(body: InvoiceApplicationInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _finance_linked_case(body.case_no, identity, db)
    if body.case_record_id:
        linked_case = await _ensure_record_visible(body.case_record_id, identity, db)
        if linked_case.module != "case": raise HTTPException(status_code=422, detail="关联记录不是案件")
        if case_record and case_record.id != linked_case.id: raise HTTPException(status_code=409, detail="案件编号与案件记录不一致")
        case_record = linked_case
    contract_record = None
    if body.contract_record_id:
        contract_record = await _ensure_record_visible(body.contract_record_id, identity, db)
        if contract_record.module != "contract": raise HTTPException(status_code=422, detail="关联记录不是合同")
    case_fee_ids = list(dict.fromkeys(body.case_fee_ids))
    case_fees: list[BusinessRecord] = []
    if case_fee_ids:
        case_fees = list((await db.scalars(select(BusinessRecord).where(
            BusinessRecord.id.in_(case_fee_ids), BusinessRecord.module == "finance",
            *(await _record_scope_conditions(identity, db)),
        ))).all())
        if len(case_fees) != len(case_fee_ids):
            raise HTTPException(status_code=404, detail="部分案件费用不存在或无权访问")
        linked_case_nos = {str((item.data or {}).get("case_no") or "").strip() for item in case_fees} - {""}
        if len(linked_case_nos) > 1:
            raise HTTPException(status_code=409, detail="同一张发票只能关联同一案件的费用")
        if case_record and linked_case_nos and case_record.serial_no not in linked_case_nos:
            raise HTTPException(status_code=409, detail="发票案件与所选案件费用不一致")
        if not case_record and linked_case_nos:
            case_record = await _finance_linked_case(next(iter(linked_case_nos)), identity, db)
    user = await db.scalar(select(User).where(User.username == identity["username"]))
    if not user: raise HTTPException(status_code=401, detail="当前用户不存在")
    serial = f"FP{datetime.now():%Y%m%d%H%M%S%f}"
    data = body.model_dump(); data["case_fee_ids"] = case_fee_ids; data["amount"] = _round_fee_amount(body.amount); data["extra_amount"] = _round_fee_amount(body.extra_amount); data["applicant"] = identity.get("display_name") or identity["username"]; data["case_id"] = case_record.id if case_record else None; data["contract_id"] = contract_record.id if contract_record else None
    if case_record: data["case_no"] = case_record.serial_no
    if contract_record: data["contract_no"] = contract_record.serial_no
    item = BusinessRecord(module="invoice", serial_no=serial, title=f"{body.customer}发票申请", customer=body.customer.strip(), status="草稿", owner=identity["username"], department=user.department, description=body.remark, data=data)
    db.add(item); await db.flush()
    db.add(WorkflowEvent(record_id=item.id, action="创建发票申请", to_status=item.status, operator=identity["username"], comment=f"{body.invoice_type}：{data['amount']:.2f} 元"))
    await db.commit(); await db.refresh(item)
    return await _record_dict_for_identity(item, identity, db)


@app.get(f"{settings.api_prefix}/customers")
async def list_customers(
    scope: str = Query("mine", pattern="^(mine|recycle|department|department_recycle|company|company_recycle|public|shared|recent_contact|recent_update)$"),
    customer_name: str = "",
    customer_type: str = Query("客户", pattern="^(客户|当事人)$"),
    manager: str = "",
    page: int = Query(1, ge=1),
    page_size: int = Query(15, ge=1, le=200),
    identity: dict = Depends(current_identity),
    db: AsyncSession = Depends(get_db),
):
    """Original customer-list scopes with authoritative server-side filtering and paging."""
    current_user = await db.scalar(
        select(User).where(User.username == identity["username"], User.is_active.is_(True))
    )
    if not current_user:
        raise HTTPException(status_code=401, detail="当前用户不存在或已停用")
    # Customer managers are stored as a JSON array.  Do not use a serialized
    # JSON ``contains`` predicate here: a user named ``ann`` would otherwise
    # match a manager named ``joann``.  Legacy rows may contain a display name
    # instead of a username, but it is safe to honor that alias only when the
    # active-user directory resolves it uniquely (the same rule used when
    # assigning customer managers).
    manager_tokens = {current_user.username}
    display_name = str(current_user.display_name or "").strip()
    if display_name:
        display_name_count = int(
            await db.scalar(
                select(func.count()).select_from(User).where(
                    User.is_active.is_(True), User.display_name == display_name
                )
            )
            or 0
        )
        if display_name_count == 1:
            manager_tokens.add(display_name)
    conditions = [BusinessRecord.module == "customer"]
    if scope in {"recycle", "department_recycle", "company_recycle"}:
        conditions.append(BusinessRecord.status == "已回收")
    elif scope == "public":
        if current_user.role not in {"admin", "manager", "user"}:
            raise HTTPException(status_code=403, detail="当前角色不能查看公海客户")
        conditions.append(BusinessRecord.status == "公海")
    elif scope == "shared":
        if current_user.role not in {"admin", "manager", "user"}:
            raise HTTPException(status_code=403, detail="当前角色不能查看共享客户")
        conditions.append(BusinessRecord.status.not_in(["已回收", "公海"]))
    elif scope == "recent_contact":
        # The original page is a projection of otherwise-visible active
        # customers, not a second customer owner register.  A non-empty real
        # contact timestamp is mandatory; recycled rows remain isolated in
        # their dedicated recycle-bin pages.
        conditions.append(BusinessRecord.status.not_in(["已回收", "公海"]))
    elif scope == "recent_update":
        # Original evidence shows this is an actor projection, not an all-firm
        # ``updated_at`` feed: rows last changed by other users on the same day
        # are absent.  Recycled rows remain eligible, while public-pool rows are
        # isolated on their dedicated page.
        if current_user.role not in {"admin", "manager", "user"}:
            raise HTTPException(status_code=403, detail="当前角色不能查看最近更新的客户")
        conditions.append(BusinessRecord.status != "公海")
    elif scope == "company":
        # The original company list is a full-firm register and includes rows
        # displayed as “已删除” (our persisted ``已回收`` status).  Public-pool
        # customers remain isolated on the separate 公海客户 page.
        conditions.append(BusinessRecord.status != "公海")
    else:
        conditions.append(BusinessRecord.status.not_in(["已回收", "公海"]))
    if scope in {"department", "department_recycle"}:
        if current_user.role not in {"admin", "manager"}:
            raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以查看部门客户")
        if current_user.role == "manager":
            conditions.append(BusinessRecord.department == current_user.department)
    elif scope in {"company", "company_recycle"} and current_user.role != "admin":
        detail = "只有系统管理员可以查看公司回收站" if scope == "company_recycle" else "只有系统管理员可以查看公司客户"
        raise HTTPException(status_code=403, detail=detail)
    normalized_name = customer_name.strip()
    if normalized_name:
        like = f"%{normalized_name}%"
        conditions.append(or_(BusinessRecord.serial_no.ilike(like), BusinessRecord.title.ilike(like)))
    conditions.append(
        func.coalesce(BusinessRecord.data["customer_type"].as_string(), "客户") == customer_type
    )
    normalized_manager = manager.strip()
    manager_search_tokens = {normalized_manager} if normalized_manager else set()
    if normalized_manager:
        # The original selector displays a person's name, while current rows
        # persist usernames.  Resolve an exact active display name only when it
        # identifies one account; duplicate display names must never broaden a
        # customer search to several unrelated managers.  Keep the literal
        # token as well for usernames and safely migrated legacy rows.
        display_name_usernames = list(
            (await db.scalars(
                select(User.username).where(
                    User.is_active.is_(True), User.display_name == normalized_manager
                )
            )).all()
        )
        if len(display_name_usernames) == 1:
            manager_search_tokens.add(display_name_usernames[0])
    candidate_rows = list(
        (await db.scalars(
        select(BusinessRecord)
        .where(*conditions)
        .order_by(BusinessRecord.updated_at.desc(), BusinessRecord.id.desc())
        )).all()
    )

    def exact_managers(item: BusinessRecord) -> set[str]:
        raw_managers = (item.data or {}).get("customer_managers", [])
        if not isinstance(raw_managers, list):
            return set()
        return {str(value).strip() for value in raw_managers if str(value).strip()}

    # Personal customer pages are still semantic subsets for normal users. Administrators
    # deliberately bypass that subset so their documented all-firm data access
    # cannot be reduced by this dedicated endpoint.  Recycle-bin and public-pool
    # rows remain on their own pages for every role.
    if scope in {"mine", "recycle"} and current_user.role != "admin":
        candidate_rows = [
            item
            for item in candidate_rows
            if str(item.owner or "").strip() in manager_tokens
            or bool(exact_managers(item) & manager_tokens)
        ]
    if scope == "shared":
        # “我的共享客户” contains only active customer rows with an explicit
        # recipient relationship.  Admin keeps full-firm audit visibility but
        # the semantic scope still excludes customers that were never shared.
        if current_user.role == "admin":
            candidate_rows = [
                item
                for item in candidate_rows
                if bool({str(value).strip() for value in (item.data or {}).get("shared_with", []) if str(value).strip()})
            ]
        else:
            candidate_rows = [
                item
                for item in candidate_rows
                if bool(
                    {
                        str(value).strip()
                        for value in (item.data or {}).get("shared_with", [])
                        if str(value).strip()
                    }
                    & manager_tokens
                )
            ]
    if scope in {"recent_contact", "recent_update"}:
        if current_user.role != "admin":
            visible_ids = await _visible_record_ids(identity, db)
            candidate_rows = [item for item in candidate_rows if item.id in visible_ids]
    if scope == "recent_contact":
        candidate_rows = [
            item for item in candidate_rows
            if _parse_customer_contact_at((item.data or {}).get("last_contact_at")) is not None
        ]
    if scope == "recent_update":
        latest_modifier_by_record: dict[int, str] = {}
        candidate_ids = [item.id for item in candidate_rows]
        if candidate_ids:
            modification_events = (await db.scalars(
                select(WorkflowEvent).where(
                    WorkflowEvent.record_id.in_(candidate_ids),
                    WorkflowEvent.action.in_(CUSTOMER_MODIFICATION_ACTIONS),
                ).order_by(WorkflowEvent.created_at.desc(), WorkflowEvent.id.desc())
            )).all()
            for event in modification_events:
                latest_modifier_by_record.setdefault(event.record_id, event.operator)
        candidate_rows = [
            item for item in candidate_rows
            if str(
                latest_modifier_by_record.get(item.id)
                or (item.data or {}).get("last_modified_by")
                or ""
            ).strip() == identity["username"]
        ]
    if normalized_manager:
        candidate_rows = [
            item
            for item in candidate_rows
            if str(item.owner or "").strip() in manager_search_tokens
            or bool(exact_managers(item) & manager_search_tokens)
        ]

    if scope == "recent_contact":
        candidate_rows.sort(
            key=lambda item: (
                _parse_customer_contact_at((item.data or {}).get("last_contact_at")) or datetime.min,
                item.id,
            ),
            reverse=True,
        )
    if scope == "recent_update":
        candidate_rows.sort(
            key=lambda item: (
                _parse_customer_contact_at(item.updated_at) or datetime.min,
                item.id,
            ),
            reverse=True,
        )

    total = len(candidate_rows)
    page_items = candidate_rows[(page - 1) * page_size : page * page_size]
    allowed_fields = await _allowed_field_keys(identity, db)
    return {
        "items": [_record_dict(item, allowed_fields) for item in page_items],
        "total": total,
        "page": page,
        "page_size": page_size,
        "summary": {
            "agency_fee_due": round(
                sum(float((item.data or {}).get("agency_fee_due") or 0) for item in candidate_rows), 2
            ),
            "official_fee_unreceived": round(
                sum(float((item.data or {}).get("official_fee_unreceived") or 0) for item in candidate_rows), 2
            ),
        },
    }


@app.post(f"{settings.api_prefix}/finance/invoices/{{invoice_id}}/submit")
async def submit_invoice_application(invoice_id: int, body: FinanceActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _ensure_record_module(invoice_id, "invoice", identity, db); await _require_record_owner_or_manager(item, identity, db)
    if item.status not in {"草稿", "已驳回"}: raise HTTPException(status_code=409, detail="当前发票申请不能提交")
    data = item.data or {}; missing = [name for name, value in {"客户名称": item.customer, "发票抬头": data.get("invoice_title"), "纳税人识别号": data.get("taxpayer_id"), "开票金额": data.get("amount")}.items() if not value]
    if data.get("delivery_method") == "电子发票" and not data.get("email"): missing.append("电子邮箱")
    if data.get("delivery_method") != "电子发票" and not data.get("delivery_address"): missing.append("邮寄地址")
    if missing: raise HTTPException(status_code=422, detail="发票申请缺少：" + "、".join(missing))
    previous = item.status; item.status = "待审批"
    db.add(WorkflowEvent(record_id=item.id, action="提交发票申请", from_status=previous, to_status=item.status, operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(item); return await _record_dict_for_identity(item, identity, db)


@app.post(f"{settings.api_prefix}/finance/invoices/{{invoice_id}}/withdraw")
async def withdraw_invoice_application(invoice_id: int, body: FinanceActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _ensure_record_module(invoice_id, "invoice", identity, db)
    await _require_record_owner_or_manager(item, identity, db)
    if item.status not in {"草稿", "待审批", "待开票", "已驳回"}:
        raise HTTPException(status_code=409, detail="当前发票申请不能撤回")
    previous = item.status
    item.status = "已撤回"
    item.data = {**(item.data or {}), "withdrawn_by": identity["username"], "withdrawn_at": datetime.now().isoformat(timespec="seconds"), "withdraw_comment": body.comment}
    db.add(WorkflowEvent(record_id=item.id, action="撤回发票申请", from_status=previous, to_status=item.status, operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(item)
    return await _record_dict_for_identity(item, identity, db)


@app.post(f"{settings.api_prefix}/finance/invoices/{{invoice_id}}/review")
async def review_invoice_application(invoice_id: int, body: FinanceReviewInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager", "auditor"}: raise HTTPException(status_code=403, detail="当前角色没有发票审批权限")
    item = await _ensure_record_module(invoice_id, "invoice", identity, db)
    if item.status != "待审批": raise HTTPException(status_code=409, detail="只有待审批发票申请可以审核")
    item.status = "待开票" if body.approved else "已驳回"
    item.data = {**(item.data or {}), "reviewer": identity["username"], "reviewed_at": datetime.now().isoformat(timespec="seconds"), "review_comment": body.comment}
    db.add(WorkflowEvent(record_id=item.id, action="发票审批通过" if body.approved else "发票审批驳回", from_status="待审批", to_status=item.status, operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(item); return await _record_dict_for_identity(item, identity, db)


@app.post(f"{settings.api_prefix}/finance/invoices/{{invoice_id}}/issue")
async def issue_invoice(invoice_id: int, body: InvoiceIssueInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以登记开票")
    item = await _ensure_record_module(invoice_id, "invoice", identity, db)
    if item.status != "待开票": raise HTTPException(status_code=409, detail="发票审批通过后才能登记开票")
    if await db.scalar(select(FinanceTransaction.id).where(FinanceTransaction.transaction_type == "开票", FinanceTransaction.voucher_no == body.invoice_no)): raise HTTPException(status_code=409, detail="发票号码已经登记")
    data = item.data or {}; tx = FinanceTransaction(finance_record_id=item.id, transaction_type="开票", amount=float(data.get("amount", 0)), transaction_date=body.invoice_date, voucher_no=body.invoice_no.strip(), counterparty=item.customer, operator=identity["username"], remark=f"发票申请 {item.serial_no}；{body.comment}")
    db.add(tx); await db.flush()
    item.status = "已开票"; item.data = {**data, "invoice_no": body.invoice_no.strip(), "invoice_date": str(body.invoice_date), "recipient": body.invoice_holder.strip(), "extra_amount": _round_fee_amount(body.extra_amount), "invoiced_opinion": body.comment.strip(), "invoice_transaction_id": tx.id, "issued_by": identity["username"], "issued_at": datetime.now().isoformat(timespec="seconds")}
    db.add(WorkflowEvent(record_id=item.id, action="登记开票", from_status="待开票", to_status=item.status, operator=identity["username"], comment=f"发票号：{body.invoice_no}。{body.comment}"))
    await db.commit(); await db.refresh(item); return await _record_dict_for_identity(item, identity, db)


@app.post(f"{settings.api_prefix}/finance/invoices/{{invoice_id}}/reject-issue")
async def reject_invoice_issue(invoice_id: int, body: FinanceActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以驳回开票")
    reason = body.comment.strip()
    if not reason:
        raise HTTPException(status_code=422, detail="请输入驳回原因")
    item = await _ensure_record_module(invoice_id, "invoice", identity, db)
    if item.status != "待开票":
        raise HTTPException(status_code=409, detail="只有待开票申请可以驳回")
    item.status = "已驳回"
    item.data = {**(item.data or {}), "invoiced_opinion": reason, "issue_rejected_by": identity["username"], "issue_rejected_at": datetime.now().isoformat(timespec="seconds")}
    db.add(WorkflowEvent(record_id=item.id, action="开票驳回", from_status="待开票", to_status=item.status, operator=identity["username"], comment=reason))
    await db.commit(); await db.refresh(item)
    return await _record_dict_for_identity(item, identity, db)


@app.post(f"{settings.api_prefix}/finance/invoices/{{invoice_id}}/void")
async def void_invoice(invoice_id: int, body: InvoiceVoidInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以作废发票")
    item = await _ensure_record_module(invoice_id, "invoice", identity, db)
    if item.status != "已开票": raise HTTPException(status_code=409, detail="只有已开票记录可以作废")
    data = item.data or {}; amount = float(data.get("amount", 0))
    db.add(FinanceTransaction(finance_record_id=item.id, transaction_type="开票", amount=-amount, transaction_date=date.today(), voucher_no=str(data.get("invoice_no", "")), counterparty=item.customer, operator=identity["username"], remark=f"作废冲销 {item.serial_no}：{body.reason}"))
    item.status = "已作废"; item.data = {**data, "void_reason": body.reason, "voided_by": identity["username"], "voided_at": datetime.now().isoformat(timespec="seconds")}
    db.add(WorkflowEvent(record_id=item.id, action="发票作废", from_status="已开票", to_status=item.status, operator=identity["username"], comment=body.reason))
    await db.commit(); await db.refresh(item); return await _record_dict_for_identity(item, identity, db)


@app.post(f"{settings.api_prefix}/finance/invoices/{{invoice_id}}/change-number")
async def change_invoice_number(invoice_id: int, body: InvoiceNumberChangeInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以修改发票号码")
    invoice_no = body.invoice_no.strip()
    if not invoice_no:
        raise HTTPException(status_code=422, detail="请输入新发票号码.")
    item = await _ensure_record_module(invoice_id, "invoice", identity, db)
    if item.status in {"已撤回", "已作废"}:
        raise HTTPException(status_code=409, detail="已撤回或已作废发票不能修改号码")
    duplicate = await db.scalar(select(FinanceTransaction.id).where(
        FinanceTransaction.transaction_type == "开票",
        FinanceTransaction.voucher_no == invoice_no,
        FinanceTransaction.finance_record_id != item.id,
    ))
    if duplicate:
        raise HTTPException(status_code=409, detail="发票号码已经登记")
    data = item.data or {}
    old_invoice_no = str(data.get("invoice_no") or "")
    transaction_id = data.get("invoice_transaction_id")
    if transaction_id:
        transaction = await db.scalar(select(FinanceTransaction).where(
            FinanceTransaction.id == int(transaction_id),
            FinanceTransaction.finance_record_id == item.id,
            FinanceTransaction.transaction_type == "开票",
        ))
        if transaction:
            transaction.voucher_no = invoice_no
    item.data = {**data, "invoice_no": invoice_no, "invoice_no_changed_by": identity["username"], "invoice_no_changed_at": datetime.now().isoformat(timespec="seconds")}
    db.add(WorkflowEvent(record_id=item.id, action="修改发票号", from_status=item.status, to_status=item.status, operator=identity["username"], comment=f"{old_invoice_no} → {invoice_no}"))
    await db.commit(); await db.refresh(item)
    return await _record_dict_for_identity(item, identity, db)


@app.post(f"{settings.api_prefix}/finance/invoices/{{invoice_id}}/change-date")
async def change_invoice_date(invoice_id: int, body: InvoiceDateChangeInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以修改发票日期")
    item = await _ensure_record_module(invoice_id, "invoice", identity, db)
    if item.status in {"已撤回", "已作废"}:
        raise HTTPException(status_code=409, detail="已撤回或已作废发票不能修改日期")
    data = item.data or {}
    old_application_date = str(data.get("application_date") or item.created_at)[:10]
    old_invoice_date = str(data.get("invoice_date") or "")[:10]
    transaction_id = data.get("invoice_transaction_id")
    if transaction_id:
        transaction = await db.scalar(select(FinanceTransaction).where(
            FinanceTransaction.id == int(transaction_id),
            FinanceTransaction.finance_record_id == item.id,
            FinanceTransaction.transaction_type == "开票",
        ))
        if transaction:
            transaction.transaction_date = body.invoice_date
    item.data = {
        **data,
        "application_date": str(body.application_date),
        "invoice_date": str(body.invoice_date),
        "invoice_date_changed_by": identity["username"],
        "invoice_date_changed_at": datetime.now().isoformat(timespec="seconds"),
    }
    db.add(WorkflowEvent(
        record_id=item.id,
        action="修改发票日期",
        from_status=item.status,
        to_status=item.status,
        operator=identity["username"],
        comment=f"申请日期 {old_application_date} → {body.application_date}；开票日期 {old_invoice_date} → {body.invoice_date}",
    ))
    await db.commit(); await db.refresh(item)
    return await _record_dict_for_identity(item, identity, db)


@app.post(f"{settings.api_prefix}/finance/refunds", status_code=status.HTTP_201_CREATED)
async def create_litigation_refund(body: LitigationRefundInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _finance_linked_case(body.case_no, identity, db)
    if not case_record: raise HTTPException(status_code=422, detail="诉讼费退款必须关联案件")
    if body.fee_record_id:
        fee_record = await _ensure_record_visible(body.fee_record_id, identity, db)
        if fee_record.module != "finance" or str((fee_record.data or {}).get("fee_type") or "") != "官方费用":
            raise HTTPException(status_code=422, detail="诉讼费退款只能关联官方费用")
        if str((fee_record.data or {}).get("case_no") or "") != case_record.serial_no:
            raise HTTPException(status_code=409, detail="退款费用与案件不一致")
    user = await db.scalar(select(User).where(User.username == identity["username"]))
    if not user: raise HTTPException(status_code=401, detail="当前用户不存在")
    serial = f"TF{datetime.now():%Y%m%d%H%M%S%f}"; data = body.model_dump(mode="json"); data["amount"] = _round_fee_amount(body.amount); data["case_id"] = case_record.id
    item = BusinessRecord(module="refund", serial_no=serial, title=f"{body.case_no}诉讼费退款", customer=body.customer.strip(), status="草稿", owner=identity["username"], department=user.department, description=body.remark, data=data)
    db.add(item); await db.flush(); db.add(WorkflowEvent(record_id=item.id, action="创建诉讼费退款申请", to_status=item.status, operator=identity["username"], comment=f"{body.court}：{data['amount']:.2f} 元"))
    await db.commit(); await db.refresh(item); return await _record_dict_for_identity(item, identity, db)


@app.post(f"{settings.api_prefix}/finance/refunds/{{refund_id}}/submit")
async def submit_litigation_refund(refund_id: int, body: FinanceActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _ensure_record_module(refund_id, "refund", identity, db); await _require_record_owner_or_manager(item, identity, db)
    if item.status not in {"草稿", "已驳回"}: raise HTTPException(status_code=409, detail="当前退款申请不能提交")
    data = item.data or {}; required = {"法院": data.get("court"), "原缴费票号": data.get("original_payment_no"), "申请人": data.get("applicant"), "退款账户名": data.get("refund_account_name"), "退款银行": data.get("refund_bank"), "退款账号": data.get("refund_account")}
    missing = [name for name, value in required.items() if not value]
    if missing: raise HTTPException(status_code=422, detail="退款申请缺少：" + "、".join(missing))
    previous = item.status; item.status = "待审批"; db.add(WorkflowEvent(record_id=item.id, action="提交诉讼费退款", from_status=previous, to_status=item.status, operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(item); return await _record_dict_for_identity(item, identity, db)


@app.post(f"{settings.api_prefix}/finance/refunds/{{refund_id}}/review")
async def review_litigation_refund(refund_id: int, body: FinanceReviewInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager", "auditor"}: raise HTTPException(status_code=403, detail="当前角色没有退款审批权限")
    item = await _ensure_record_module(refund_id, "refund", identity, db)
    if item.status != "待审批": raise HTTPException(status_code=409, detail="只有待审批退款申请可以审核")
    item.status = "退款办理中" if body.approved else "已驳回"; item.data = {**(item.data or {}), "reviewer": identity["username"], "reviewed_at": datetime.now().isoformat(timespec="seconds"), "review_comment": body.comment}
    db.add(WorkflowEvent(record_id=item.id, action="退款审批通过" if body.approved else "退款审批驳回", from_status="待审批", to_status=item.status, operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(item); return await _record_dict_for_identity(item, identity, db)


@app.post(f"{settings.api_prefix}/finance/refunds/{{refund_id}}/complete")
async def complete_litigation_refund(refund_id: int, body: RefundCompleteInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以登记退款到账")
    item = await _ensure_record_module(refund_id, "refund", identity, db)
    if item.status != "退款办理中": raise HTTPException(status_code=409, detail="退款审批通过后才能登记到账")
    data = item.data or {}; tx = FinanceTransaction(finance_record_id=item.id, transaction_type="退费", amount=float(data.get("amount", 0)), transaction_date=body.actual_date, voucher_no=body.voucher_no.strip(), counterparty=str(data.get("court", item.customer)), operator=identity["username"], remark=f"诉讼费退款 {item.serial_no}；{body.comment}")
    db.add(tx); await db.flush(); item.status = "已退款"; item.data = {**data, "actual_date": str(body.actual_date), "refund_voucher_no": body.voucher_no.strip(), "refund_transaction_id": tx.id, "completed_by": identity["username"], "completed_at": datetime.now().isoformat(timespec="seconds")}
    db.add(WorkflowEvent(record_id=item.id, action="登记退款到账", from_status="退款办理中", to_status=item.status, operator=identity["username"], comment=f"凭证号：{body.voucher_no}。{body.comment}"))
    await db.commit(); await db.refresh(item); return await _record_dict_for_identity(item, identity, db)


async def _finance_fee_readiness(item: BusinessRecord, identity: dict, db: AsyncSession) -> dict:
    data = item.data or {}
    case_no = str(data.get("case_no", "")).strip()
    case_record = await db.scalar(select(BusinessRecord).where(BusinessRecord.module == "case", BusinessRecord.serial_no == case_no, *(await _record_scope_conditions(identity, db)))) if case_no else None
    case_data = (case_record.data or {}) if case_record else {}
    personnel = {
        "客户管理人": bool(case_data.get("customer_manager")),
        "开庭律师": bool(case_data.get("hearing_lawyer")),
        "经办律师": bool(case_data.get("handling_lawyers")),
        "律师助理/文书": bool(case_data.get("assistant")),
    }
    court = {
        "关联案件": bool(case_record),
        "法院名称": bool(data.get("court") or case_data.get("court")),
        "缴费通知文号": bool(data.get("document_no")),
    }
    attachment_categories = set((await db.scalars(select(FileAttachment.category).where(FileAttachment.record_id == case_record.id))).all()) if case_record else set()
    document = {"法院缴费通知书": bool(data.get("document_no") or attachment_categories.intersection({"法院缴费通知书", "缴费通知书"}))}
    missing = [f"人员要素：{name}" for name, ready in personnel.items() if not ready]
    missing.extend(f"法院要素：{name}" for name, ready in court.items() if not ready)
    missing.extend(f"文档要素：{name}" for name, ready in document.items() if not ready)
    return {"case_id": case_record.id if case_record else None, "case_no": case_no, "personnel": personnel, "court": court, "document": document, "ready": not missing, "missing": missing}


def _internal_fee_payment_status(item: BusinessRecord, paid_amount: float) -> str:
    data = item.data or {}
    amount = float(data.get("amount", 0) or 0)
    explicit = str(data.get("payment_status", "")).strip()
    if explicit in {"已付", "已付款"} or item.status == "已付款" or (amount > 0 and paid_amount >= amount):
        return "已付"
    return "未付"


def _internal_fee_row(item: BusinessRecord, case_record: BusinessRecord | None, paid_amount: float, allowed_fields: set[str]) -> dict:
    data = item.data or {}
    case_data = (case_record.data or {}) if case_record else {}
    result = _record_dict(item, allowed_fields)
    visible_data = dict(result.get("data") or {})
    enriched = {
        **visible_data,
        "case_no": data.get("case_no") or (case_record.serial_no if case_record else ""),
        "case_stage": data.get("case_stage") or case_data.get("case_stage") or (case_record.status if case_record else ""),
        "plaintiff": data.get("plaintiff") or case_data.get("plaintiff") or case_data.get("appellant_names") or (case_record.customer if case_record else ""),
        "defendant": data.get("defendant") or case_data.get("defendant") or case_data.get("appellee_names") or case_data.get("opponent") or "",
        "handling_lawyer": data.get("handling_lawyer") or data.get("handler") or case_data.get("handling_lawyer") or case_data.get("case_lawyer") or ",".join(case_data.get("handling_lawyers") or []) or "",
        "lawyer_assistant": data.get("lawyer_assistant") or data.get("assistant") or case_data.get("lawyer_assistant") or case_data.get("assistant") or "",
        "case_source": data.get("case_source") or data.get("source_person") or case_data.get("case_source") or case_data.get("business_owner") or (case_record.owner if case_record else ""),
        "investigator": data.get("investigator") or case_data.get("investigator") or "",
        "archive_date": data.get("archive_date") or data.get("audited_time") or case_data.get("archive_date") or case_data.get("audited_time") or "",
        "application_date": data.get("application_date") or item.created_at,
        "internal_fee_type": data.get("commission_type") or data.get("fee_type_name") or item.title or data.get("fee_type") or "",
        "payee": data.get("payee") or "",
        "payment_status": _internal_fee_payment_status(item, paid_amount),
    }
    if "finance.amount" in allowed_fields:
        enriched["paid_amount"] = round(paid_amount, 2)
    result["data"] = enriched
    return result


async def _internal_fee_rows(
    identity: dict,
    db: AsyncSession,
    *,
    scope: str,
    case_no: str,
    handling_lawyer: str,
    assistant: str,
    source_person: str,
    customer: str,
    customer_manager: str,
    investigator: str,
    payment_status: str,
    paid_from: date | None,
    paid_to: date | None,
    payee: str,
    case_stages: str,
    fee_types: str,
    ids: set[int] | None = None,
) -> list[dict]:
    scope_conditions = await _record_scope_conditions(identity, db)
    fees = list((await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "finance", *scope_conditions).order_by(BusinessRecord.updated_at.desc(), BusinessRecord.id.desc()))).all())
    fees = [item for item in fees if (item.data or {}).get("fee_type") == "内部费用"]
    if ids is not None:
        fees = [item for item in fees if item.id in ids]
    case_records = list((await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "case", *scope_conditions))).all())
    cases_by_id = {item.id: item for item in case_records}
    cases_by_no = {item.serial_no: item for item in case_records}
    fee_ids = [item.id for item in fees]
    transactions = list((await db.scalars(select(FinanceTransaction).where(FinanceTransaction.finance_record_id.in_(fee_ids)))).all()) if fee_ids else []
    paid_by_fee: dict[int, float] = {}
    paid_dates_by_fee: dict[int, list[date]] = {}
    for transaction in transactions:
        if transaction.transaction_type != "付款" or not transaction.finance_record_id:
            continue
        paid_by_fee[transaction.finance_record_id] = paid_by_fee.get(transaction.finance_record_id, 0.0) + float(transaction.amount or 0)
        paid_dates_by_fee.setdefault(transaction.finance_record_id, []).append(transaction.transaction_date)
    allowed_fields = await _allowed_field_keys(identity, db)
    selected_stages = {value.strip() for value in case_stages.split(",") if value.strip()}
    selected_types = {value.strip() for value in fee_types.split(",") if value.strip()}
    personal_names = {str(identity.get("username", "")).strip(), str(identity.get("display_name", "")).strip()} - {""}

    def contains(value: object, needle: str) -> bool:
        return not needle.strip() or needle.strip().casefold() in str(value or "").casefold()

    rows: list[dict] = []
    for item in fees:
        data = item.data or {}
        linked_case = cases_by_id.get(int(data.get("case_id") or 0)) or cases_by_no.get(str(data.get("case_no") or ""))
        row = _internal_fee_row(item, linked_case, paid_by_fee.get(item.id, 0.0), allowed_fields)
        row_data = row["data"]
        paid_dates = list(paid_dates_by_fee.get(item.id, []))
        stored_paid_date = str(data.get("paid_date") or data.get("payment_date") or "").strip()
        if stored_paid_date:
            try:
                paid_dates.append(date.fromisoformat(stored_paid_date[:10]))
            except ValueError:
                pass
        row_data["paid_date"] = max(paid_dates).isoformat() if paid_dates else stored_paid_date
        payment_object = str(row_data.get("payee") or "").strip()
        if scope == "mine" and payment_object not in personal_names:
            continue
        if payee.strip() and not contains(payment_object, payee):
            continue
        if not contains(row_data.get("case_no"), case_no) or not contains(row_data.get("handling_lawyer"), handling_lawyer):
            continue
        if not contains(row_data.get("lawyer_assistant"), assistant) or not contains(row_data.get("case_source"), source_person):
            continue
        if not contains(item.customer, customer) or not contains((linked_case.data or {}).get("customer_manager") if linked_case else data.get("customer_manager"), customer_manager):
            continue
        if not contains(row_data.get("investigator"), investigator):
            continue
        if payment_status and row_data.get("payment_status") != payment_status:
            continue
        if paid_from and not any(value >= paid_from for value in paid_dates):
            continue
        if paid_to and not any(value <= paid_to for value in paid_dates):
            continue
        if selected_stages and str(row_data.get("case_stage") or "") not in selected_stages:
            continue
        if selected_types and str(row_data.get("internal_fee_type") or "") not in selected_types:
            continue
        rows.append(row)
    return rows


@app.get(f"{settings.api_prefix}/finance/internal-fees")
async def list_internal_fees(
    scope: str = Query("company", pattern="^(mine|company)$"),
    case_no: str = "", handling_lawyer: str = "", assistant: str = "", source_person: str = "",
    customer: str = "", customer_manager: str = "", investigator: str = "", payment_status: str = Query("", pattern="^(|已付|未付)$"),
    paid_from: date | None = None, paid_to: date | None = None, payee: str = "", case_stages: str = "", fee_types: str = "",
    page: int = Query(1, ge=1), page_size: int = Query(15, ge=1, le=200),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    if paid_from and paid_to and paid_from > paid_to:
        raise HTTPException(status_code=422, detail="付款开始日期不能晚于结束日期")
    rows = await _internal_fee_rows(identity, db, scope=scope, case_no=case_no, handling_lawyer=handling_lawyer, assistant=assistant, source_person=source_person, customer=customer, customer_manager=customer_manager, investigator=investigator, payment_status=payment_status, paid_from=paid_from, paid_to=paid_to, payee=payee, case_stages=case_stages, fee_types=fee_types)
    total = len(rows)
    total_amount = round(sum(float((row.get("data") or {}).get("amount", 0) or 0) for row in rows), 2)
    start = (page - 1) * page_size
    return {"items": rows[start:start + page_size], "total": total, "total_amount": total_amount, "page": page, "page_size": page_size}


@app.get(f"{settings.api_prefix}/finance/internal-fees/export")
async def export_internal_fees(
    scope: str = Query("company", pattern="^(mine|company)$"), ids: str = "",
    case_no: str = "", handling_lawyer: str = "", assistant: str = "", source_person: str = "",
    customer: str = "", customer_manager: str = "", investigator: str = "", payment_status: str = Query("", pattern="^(|已付|未付)$"),
    paid_from: date | None = None, paid_to: date | None = None, payee: str = "", case_stages: str = "", fee_types: str = "",
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    selected_ids = set(_export_ids(ids)) if ids.strip() else None
    rows = await _internal_fee_rows(identity, db, scope=scope, case_no=case_no, handling_lawyer=handling_lawyer, assistant=assistant, source_person=source_person, customer=customer, customer_manager=customer_manager, investigator=investigator, payment_status=payment_status, paid_from=paid_from, paid_to=paid_to, payee=payee, case_stages=case_stages, fee_types=fee_types, ids=selected_ids)
    if selected_ids is not None and not rows:
        raise HTTPException(status_code=422, detail="请选择需要导出的费用")
    headers = ["案号", "案件阶段", "原告", "被告", "经办律师", "律师助理", "案源人", "调查人", "归档时间", "申请时间", "内部费用类型", "金额", "收款人", "支付状态"]
    keys = ["case_no", "case_stage", "plaintiff", "defendant", "handling_lawyer", "lawyer_assistant", "case_source", "investigator", "archive_date", "application_date", "internal_fee_type", "amount", "payee", "payment_status"]
    def cell(value: object, *, number: bool = False) -> str:
        text_value = f"{float(value or 0):.2f}" if number else str(value or "")[:10] if isinstance(value, (date, datetime)) else str(value or "")
        data_type = "Number" if number else "String"
        return f'<Cell><Data ss:Type="{data_type}">{xml_escape(text_value)}</Data></Cell>'
    sheet_rows = ["<Row>" + "".join(cell(value) for value in headers) + "</Row>"]
    for row in rows:
        data = row.get("data") or {}
        sheet_rows.append("<Row>" + "".join(cell(data.get(key), number=key == "amount") for key in keys) + "</Row>")
    workbook = '<?xml version="1.0" encoding="UTF-8"?><?mso-application progid="Excel.Sheet"?><Workbook xmlns="urn:schemas-microsoft-com:office:spreadsheet" xmlns:ss="urn:schemas-microsoft-com:office:spreadsheet"><Worksheet ss:Name="内部费用明细"><Table>' + "".join(sheet_rows) + "</Table></Worksheet></Workbook>"
    filename = f"内部费用明细-{date.today()}.xls"
    disposition = f"attachment; filename=internal-fees.xls; filename*=UTF-8''{quote(filename)}"
    return Response(content=workbook.encode("utf-8"), media_type="application/vnd.ms-excel", headers={"Content-Disposition": disposition})


@app.get(f"{settings.api_prefix}/finance/summary")
async def finance_summary(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    finance_records = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module.in_(["finance", "invoice", "refund"]), *(await _record_scope_conditions(identity, db))))).all()
    fees = [item for item in finance_records if item.module == "finance"]
    invoices = [item for item in finance_records if item.module == "invoice"]
    refunds = [item for item in finance_records if item.module == "refund"]
    transactions = (await db.scalars(select(FinanceTransaction))).all()
    visible_record_ids = {item.id for item in finance_records}
    if identity.get("role") != "admin": transactions = [item for item in transactions if (item.finance_record_id and item.finance_record_id in visible_record_ids) or (not item.finance_record_id and item.operator == identity["username"])]
    amounts = {fee_type: sum(float((item.data or {}).get("amount", 0)) for item in fees if (item.data or {}).get("fee_type") == fee_type) for fee_type in FINANCE_FEE_TYPES}
    can_view_amount = "finance.amount" in await _allowed_field_keys(identity, db)
    incoming = (await db.scalars(select(IncomingPayment))).all()
    if identity.get("role") not in {"admin", "auditor"}:
        visible_customer_titles = set((await db.scalars(select(BusinessRecord.title).where(BusinessRecord.module == "customer", *(await _record_scope_conditions(identity, db))))).all())
        incoming = [item for item in incoming if item.operator == identity["username"] or item.claimant == identity["username"] or item.claimed_customer in visible_customer_titles]
    return {
        "fees": len(fees), "draft": sum(1 for item in fees if item.status == "草稿"),
        "pending": sum(1 for item in fees if item.status == "待审批"),
        "approved": sum(1 for item in fees if item.status in {"已审批", "部分付款"}),
        "paid": sum(1 for item in fees if item.status == "已付款"),
        "invoice_applications": len(invoices), "invoice_pending": sum(1 for item in invoices if item.status == "待审批"),
        "refund_applications": len(refunds), "refund_pending": sum(1 for item in refunds if item.status == "待审批"),
        "amount_visible": can_view_amount,
        "total_fee_amount": sum(amounts.values()) if can_view_amount else None, "amounts_by_type": amounts if can_view_amount else {},
        "paid_amount": sum(item.amount for item in transactions if item.transaction_type == "付款") if can_view_amount else None,
        "invoice_amount": sum(item.amount for item in transactions if item.transaction_type == "开票") if can_view_amount else None,
        "refund_amount": sum(item.amount for item in transactions if item.transaction_type == "退费") if can_view_amount else None,
        "incoming_payments": len(incoming), "incoming_unclaimed": sum(1 for item in incoming if item.status == "待认领"), "incoming_unallocated": sum(1 for item in incoming if item.status in {"待分配", "部分分配"}),
    }


@app.get(f"{settings.api_prefix}/finance/incoming-payments")
async def list_incoming_payments(payment_status: str = "", keyword: str = "", identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    items = (await db.scalars(select(IncomingPayment).order_by(IncomingPayment.received_date.desc(), IncomingPayment.id.desc()))).all()
    if identity.get("role") not in {"admin", "auditor"}:
        visible_customer_titles = set((await db.scalars(select(BusinessRecord.title).where(BusinessRecord.module == "customer", *(await _record_scope_conditions(identity, db))))).all())
        items = [item for item in items if item.operator == identity["username"] or item.claimant == identity["username"] or item.claimed_customer in visible_customer_titles]
    if payment_status: items = [item for item in items if item.status == payment_status]
    if keyword:
        key = keyword.casefold(); items = [item for item in items if key in f"{item.receipt_no} {item.payer_name} {item.bank_reference} {item.claimed_customer}".casefold()]
    can_view_amount = "finance.amount" in await _allowed_field_keys(identity, db)
    return {"items": [_incoming_payment_dict(item, show_amount=can_view_amount) for item in items], "total": len(items), "summary": {"total": len(items), "unclaimed": sum(1 for item in items if item.status == "待认领"), "unallocated": sum(1 for item in items if item.status in {"待分配", "部分分配"}), "completed": sum(1 for item in items if item.status == "已分配"), "amount": sum(item.amount for item in items) if can_view_amount else None, "remaining": sum(max(item.amount - item.allocated_amount, 0) for item in items) if can_view_amount else None}}


@app.post(f"{settings.api_prefix}/finance/incoming-payments", status_code=status.HTTP_201_CREATED)
async def create_incoming_payment(body: IncomingPaymentInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以登记银行到账")
    if await db.scalar(select(IncomingPayment.id).where(IncomingPayment.bank_reference == body.bank_reference.strip())): raise HTTPException(status_code=409, detail="银行流水号已经登记")
    item = IncomingPayment(receipt_no=f"HK{datetime.now():%Y%m%d%H%M%S%f}", received_date=body.received_date, amount=_round_fee_amount(body.amount), payer_name=body.payer_name.strip(), bank_reference=body.bank_reference.strip(), status="待认领", operator=identity["username"], remark=body.remark)
    db.add(item); await db.commit(); await db.refresh(item); return _incoming_payment_dict(item)


@app.post(f"{settings.api_prefix}/finance/incoming-payments/import")
async def import_incoming_payments(file: UploadFile = File(...), identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以导入银行到账")
    if not (file.filename or "").lower().endswith(".csv"):
        raise HTTPException(status_code=422, detail="仅支持 UTF-8 CSV 文件")
    raw = await file.read()
    if len(raw) > 5 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="CSV 文件不能超过 5MB")
    try:
        content = raw.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise HTTPException(status_code=422, detail="CSV 文件必须使用 UTF-8 编码") from exc
    existing = set((await db.scalars(select(IncomingPayment.bank_reference))).all())
    created = 0
    errors: list[dict] = []
    for row_no, row in enumerate(csv.DictReader(io.StringIO(content)), 2):
        try:
            payer = _csv_value(row, "对方户名", "付款人", "付款方", "payer_name")
            bank_reference = _csv_value(row, "银行流水号", "交易流水号", "业务编号", "bank_reference")
            received_text = _csv_value(row, "到账日期", "交易日期", "记账日期", "received_date")
            amount_text = _csv_value(row, "到账金额", "交易金额", "贷方发生额", "amount")
            if not payer or not bank_reference or not received_text or not amount_text:
                raise ValueError("缺少对方户名、银行流水号、到账日期或到账金额")
            if bank_reference in existing:
                raise ValueError("银行流水号已经登记")
            received_date = date.fromisoformat(received_text.replace("/", "-").strip())
            amount = _round_fee_amount(float(amount_text.replace(",", "").strip()))
            if amount <= 0:
                raise ValueError("到账金额必须大于 0")
            db.add(IncomingPayment(
                receipt_no=f"HK{datetime.now():%Y%m%d%H%M%S%f}{row_no}",
                received_date=received_date,
                amount=amount,
                payer_name=payer,
                bank_reference=bank_reference,
                status="待认领",
                operator=identity["username"],
                remark=_csv_value(row, "摘要", "备注", "remark"),
            ))
            existing.add(bank_reference)
            created += 1
        except (ValueError, TypeError) as exc:
            errors.append({"row": row_no, "error": str(exc) or "字段格式错误"})
    if created:
        await db.commit()
    return {"created": created, "errors": errors}


@app.post(f"{settings.api_prefix}/finance/incoming-payments/{{payment_id}}/claim")
async def claim_incoming_payment(payment_id: int, body: IncomingPaymentClaimInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await db.get(IncomingPayment, payment_id)
    if not item: raise HTTPException(status_code=404, detail="银行到账记录不存在")
    if item.status not in {"待认领", "待分配"} or item.allocated_amount > 0: raise HTTPException(status_code=409, detail="已发生分配的到账不能重新认领")
    customer = await db.scalar(select(BusinessRecord).where(BusinessRecord.module == "customer", BusinessRecord.title == body.customer.strip(), *(await _record_scope_conditions(identity, db))))
    if not customer: raise HTTPException(status_code=404, detail="客户不存在或无权认领")
    item.claimed_customer = customer.title; item.claimant = identity["username"]; item.status = "待分配"; item.remark = "；".join(part for part in [item.remark, body.comment] if part)
    db.add(WorkflowEvent(record_id=customer.id, action="认领银行到账", from_status=customer.status, to_status=customer.status, operator=identity["username"], comment=f"{item.receipt_no}｜{item.payer_name}｜{item.amount:.2f} 元。{body.comment}"))
    await db.commit(); await db.refresh(item); return _incoming_payment_dict(item, show_amount="finance.amount" in await _allowed_field_keys(identity, db))


@app.post(f"{settings.api_prefix}/finance/incoming-payments/{{payment_id}}/allocate")
async def allocate_incoming_payment(payment_id: int, body: IncomingPaymentAllocateInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await db.get(IncomingPayment, payment_id)
    if not item: raise HTTPException(status_code=404, detail="银行到账记录不存在")
    if item.status not in {"待分配", "部分分配"} or not item.claimed_customer: raise HTTPException(status_code=409, detail="到账认领到客户后才能分配")
    if len({entry.receivable_plan_id for entry in body.allocations}) != len(body.allocations): raise HTTPException(status_code=422, detail="同一次分配中应收计划不能重复")
    total = _round_fee_amount(sum(entry.amount for entry in body.allocations)); remaining_payment = _round_fee_amount(item.amount - item.allocated_amount)
    if total > remaining_payment + 0.001: raise HTTPException(status_code=409, detail=f"分配金额超过到账未分配余额 {remaining_payment:.2f} 元")
    prepared: list[tuple[IncomingPaymentAllocationItem, ReceivablePlan, BusinessRecord, BusinessRecord | None]] = []
    for entry in body.allocations:
        if entry.settlement_items:
            classified_total = _round_fee_amount(sum(item.amount for item in entry.settlement_items))
            if abs(classified_total - _round_fee_amount(entry.amount)) > 0.001:
                raise HTTPException(status_code=422, detail="结算费用明细金额之和必须等于本次分配金额")
            invalid_settlement = [item.fee_type for item in entry.settlement_items if item.settlement_amount + 0.001 < item.archive_fee]
            if invalid_settlement:
                raise HTTPException(status_code=422, detail="归档费不能大于结算金额：" + "、".join(invalid_settlement))
            excessive_settlement = [item.fee_type for item in entry.settlement_items if item.settlement_amount > item.amount + 0.001]
            if excessive_settlement:
                raise HTTPException(status_code=422, detail="结算金额不能大于分配金额：" + "、".join(excessive_settlement))
        plan = await db.get(ReceivablePlan, entry.receivable_plan_id)
        if not plan: raise HTTPException(status_code=404, detail=f"应收计划 {entry.receivable_plan_id} 不存在")
        contract = await _ensure_record_module(plan.contract_record_id, "contract", identity, db)
        if contract.customer.strip() != item.claimed_customer.strip(): raise HTTPException(status_code=409, detail=f"应收计划 {plan.phase} 的客户与到账认领客户不一致")
        remaining_plan = _round_fee_amount(plan.amount - plan.received_amount)
        if entry.amount > remaining_plan + 0.001: raise HTTPException(status_code=409, detail=f"{contract.serial_no}｜{plan.phase} 分配金额超过未收 {remaining_plan:.2f} 元")
        case_record = None
        if entry.case_no.strip():
            case_record = await db.scalar(select(BusinessRecord).where(BusinessRecord.module == "case", BusinessRecord.serial_no == entry.case_no.strip(), *(await _record_scope_conditions(identity, db))))
            if not case_record or int((case_record.data or {}).get("contract_id") or 0) != contract.id: raise HTTPException(status_code=409, detail=f"案件 {entry.case_no} 不属于所选合同")
        for settlement in entry.settlement_items:
            if settlement.fee_record_id is None:
                continue
            fee_record = await _ensure_record_module(settlement.fee_record_id, "finance", identity, db)
            fee_data = fee_record.data or {}
            if case_record and int(fee_data.get("case_id") or 0) not in {0, case_record.id} and str(fee_data.get("case_no") or "") != case_record.serial_no:
                raise HTTPException(status_code=409, detail=f"费用 {fee_record.serial_no} 不属于案件 {case_record.serial_no}")
        prepared.append((entry, plan, contract, case_record))
    allocation_rows = list(item.allocations or [])
    for entry, plan, contract, case_record in prepared:
        amount = _round_fee_amount(entry.amount); plan.received_amount = _round_fee_amount(plan.received_amount + amount); plan.status = "已收款" if plan.received_amount + 0.001 >= plan.amount else "部分收款"
        tx = FinanceTransaction(finance_record_id=contract.id, transaction_type="回款", amount=amount, transaction_date=item.received_date, voucher_no=item.bank_reference, counterparty=item.payer_name, operator=identity["username"], remark=f"银行到账 {item.receipt_no} 分配至 {contract.serial_no}｜{plan.phase}" + (f"｜案件 {case_record.serial_no}" if case_record else ""))
        db.add(tx); await db.flush(); row = {"receivable_plan_id": plan.id, "contract_id": contract.id, "contract_no": contract.serial_no, "phase": plan.phase, "case_id": case_record.id if case_record else None, "case_no": case_record.serial_no if case_record else "", "amount": amount, "payment_method": entry.payment_method.strip(), "settlement_items": [settlement.model_dump() for settlement in entry.settlement_items], "transaction_id": tx.id, "allocated_by": identity["username"], "allocated_at": datetime.now().isoformat(timespec="seconds")}; allocation_rows.append(row)
        db.add(WorkflowEvent(record_id=contract.id, action="分配银行回款", from_status=contract.status, to_status=contract.status, operator=identity["username"], comment=f"{item.receipt_no}｜{plan.phase}｜{amount:.2f} 元。{body.comment}"))
    item.allocated_amount = _round_fee_amount(item.allocated_amount + total); item.allocations = allocation_rows; item.status = "已分配" if item.allocated_amount + 0.001 >= item.amount else "部分分配"
    await db.commit(); await db.refresh(item); return _incoming_payment_dict(item, show_amount="finance.amount" in await _allowed_field_keys(identity, db))


@app.delete(f"{settings.api_prefix}/finance/incoming-payments/{{payment_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_incoming_payment(payment_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") != "admin": raise HTTPException(status_code=403, detail="仅管理员可删除银行到账")
    item = await db.get(IncomingPayment, payment_id)
    if not item: raise HTTPException(status_code=404, detail="银行到账记录不存在")
    for allocation in item.allocations or []:
        plan = await db.get(ReceivablePlan, int(allocation.get("receivable_plan_id") or 0)); amount = float(allocation.get("amount") or 0)
        if plan:
            plan.received_amount = max(_round_fee_amount(plan.received_amount - amount), 0); plan.status = "待收款" if plan.received_amount <= 0 else "部分收款"
        tx = await db.get(FinanceTransaction, int(allocation.get("transaction_id") or 0))
        if tx: await db.delete(tx)
    await db.delete(item); await db.commit(); return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.post(f"{settings.api_prefix}/finance/fees", status_code=status.HTTP_201_CREATED)
async def create_finance_fee(body: FinanceFeeInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if body.fee_type not in FINANCE_FEE_TYPES:
        raise HTTPException(status_code=422, detail="费用类型无效")
    if body.expense_scope and body.fee_type not in EXPENSE_SCOPE_FEE_TYPES[body.expense_scope]:
        raise HTTPException(status_code=422, detail="费用归属与费用类型不一致")
    if body.expense_subtype and EXPENSE_SUBTYPE_FEE_TYPE[body.expense_subtype] != body.fee_type:
        raise HTTPException(status_code=422, detail="费用子类型与费用类型不一致")
    if body.amount == 0: raise HTTPException(status_code=422, detail="费用金额不能为 0")
    if body.amount < 0 and body.fee_type != "内部费用": raise HTTPException(status_code=422, detail="只有内部费用可以使用负数冲销")
    case_record = await _finance_linked_case(body.case_no, identity, db)
    if body.case_record_id:
        linked_case = await _ensure_record_visible(body.case_record_id, identity, db)
        if linked_case.module != "case": raise HTTPException(status_code=422, detail="关联记录不是案件")
        if case_record and case_record.id != linked_case.id: raise HTTPException(status_code=409, detail="案件编号与案件记录不一致")
        case_record = linked_case
    contract_record = None
    if body.contract_record_id:
        contract_record = await _ensure_record_visible(body.contract_record_id, identity, db)
        if contract_record.module != "contract": raise HTTPException(status_code=422, detail="关联记录不是合同")
    user = await db.scalar(select(User).where(User.username == identity["username"]))
    if not user: raise HTTPException(status_code=401, detail="当前用户不存在")
    handler = identity["username"] if identity.get("role") == "user" else body.handler
    amount = _round_fee_amount(body.amount)
    serial = f"FY{datetime.now():%Y%m%d%H%M%S%f}"
    item = BusinessRecord(module="finance", serial_no=serial, title=body.title, customer=body.customer, status="草稿", owner=handler, department=user.department, description=body.description, data={"amount": amount, "fee_type": body.fee_type, "expense_scope": body.expense_scope or "", "expense_subtype": body.expense_subtype or "", "is_refund": body.fee_type == "内部费用" and amount < 0, "case_no": case_record.serial_no if case_record else body.case_no, "case_id": case_record.id if case_record else None, "contract_id": contract_record.id if contract_record else None, "contract_no": contract_record.serial_no if contract_record else "", "handler": handler, "court": body.court, "document_no": body.document_no, "payee": body.payee})
    db.add(item); await db.flush()
    db.add(WorkflowEvent(record_id=item.id, action="创建费用", to_status="草稿", operator=identity["username"], comment=f"{body.fee_type}：{amount:.2f} 元"))
    await db.commit(); await db.refresh(item)
    return await _record_dict_for_identity(item, identity, db)


def _settlement_fee_kind(fee_type: str) -> str:
    normalized = fee_type.strip()
    if "代理" in normalized:
        return "agency"
    if any(token in normalized for token in ("官费", "官方", "诉讼", "公证", "保全", "鉴定", "公告")):
        return "official"
    return "other"


async def _general_settlement_rows(
    identity: dict,
    db: AsyncSession,
    *,
    customer: str = "",
    case_no: str = "",
    received_from: date | None = None,
    received_to: date | None = None,
    payer: str = "",
    payment_method: str = "",
    case_customer: str = "",
    hearing_lawyer: str = "",
    assistant: str = "",
    customer_manager: str = "",
    source_person: str = "",
    receipt_ids: set[int] | None = None,
) -> list[dict]:
    """Build the original settlement candidates from bank receipts and their case allocations."""
    payments = list((await db.scalars(select(IncomingPayment).order_by(
        IncomingPayment.received_date.asc(), IncomingPayment.id.asc()
    ))).all())
    if receipt_ids is not None:
        payments = [item for item in payments if item.id in receipt_ids]

    allocations = [allocation for payment in payments for allocation in (payment.allocations or [])]
    case_ids = {int(row.get("case_id") or 0) for row in allocations if row.get("case_id")}
    case_nos = {str(row.get("case_no") or "").strip() for row in allocations if str(row.get("case_no") or "").strip()}
    cases = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "case",
        or_(BusinessRecord.id.in_(case_ids), BusinessRecord.serial_no.in_(case_nos)),
        *(await _record_scope_conditions(identity, db)),
    ))).all()) if case_ids or case_nos else []
    cases_by_id = {item.id: item for item in cases}
    cases_by_no = {item.serial_no: item for item in cases}

    visible_case_ids = set(cases_by_id)
    visible_case_nos = set(cases_by_no)
    fee_records = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "finance",
        *(await _record_scope_conditions(identity, db)),
    ).order_by(BusinessRecord.id.asc()))).all())
    fees_by_case: dict[str, list[BusinessRecord]] = {}
    for fee in fee_records:
        data = fee.data or {}
        fee_case_id = int(data.get("case_id") or 0)
        fee_case_no = str(data.get("case_no") or "").strip()
        if fee_case_id not in visible_case_ids and fee_case_no not in visible_case_nos:
            continue
        fee_type = _case_fee_display_type(fee)
        if not fee_type or fee_type in {"内部费用", "结算费用", "归档费用", "预损费用"}:
            continue
        key = str(fee_case_id or fee_case_no)
        fees_by_case.setdefault(key, []).append(fee)

    settlement_records = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "finance_settlement"
    ).order_by(BusinessRecord.created_at.desc(), BusinessRecord.id.desc()))).all())
    inactive_statuses = {"已拒绝", "已驳回", "已退回", "已撤回", "已作废"}
    active_receipt_ids = {
        int((record.data or {}).get("receipt_id") or 0)
        for record in settlement_records
        if record.status not in inactive_statuses
    }
    rejection_by_receipt: dict[int, str] = {}
    for record in settlement_records:
        receipt_id = int((record.data or {}).get("receipt_id") or 0)
        if receipt_id and record.status in inactive_statuses and receipt_id not in rejection_by_receipt:
            rejection_by_receipt[receipt_id] = str((record.data or {}).get("review_comment") or record.description or "")

    consumed_by_fee: dict[int, float] = {}
    rows: list[dict] = []
    can_view_amount = "finance.amount" in await _allowed_field_keys(identity, db)

    def contains(value: object, needle: str) -> bool:
        return not needle.strip() or needle.strip().casefold() in str(value or "").casefold()

    for payment in payments:
        payment_allocations = list(payment.allocations or [])
        if not payment_allocations or payment.id in active_receipt_ids:
            continue
        details: list[dict] = []
        for allocation in payment_allocations:
            linked_case = cases_by_id.get(int(allocation.get("case_id") or 0)) or cases_by_no.get(str(allocation.get("case_no") or ""))
            if not linked_case and identity.get("role") not in {"admin", "auditor"}:
                continue
            case_data = (linked_case.data or {}) if linked_case else {}
            remaining = _round_fee_amount(float(allocation.get("amount") or 0))
            explicit_items = list(allocation.get("settlement_items") or [])
            if explicit_items:
                for explicit in explicit_items:
                    fee_type = str(explicit.get("fee_type") or "其他费用")
                    current_amount = _round_fee_amount(float(explicit.get("amount") or 0))
                    details.append({
                        "fee_id": explicit.get("fee_record_id"),
                        "case_id": linked_case.id if linked_case else allocation.get("case_id"),
                        "case_no": linked_case.serial_no if linked_case else allocation.get("case_no", ""),
                        "case_stage": (case_data.get("case_stage") or linked_case.status) if linked_case else "",
                        "fee_type": fee_type,
                        "fee_total_amount": current_amount,
                        "fee_allocated_amount": current_amount,
                        "current_amount": current_amount,
                        "allocated_at": allocation.get("allocated_at", ""),
                        "settlement_amount": _round_fee_amount(float(explicit.get("settlement_amount") or 0)),
                        "archive_fee": _round_fee_amount(float(explicit.get("archive_fee") or 0)),
                        "customer": linked_case.customer if linked_case else payment.claimed_customer,
                        "handling_lawyer": case_data.get("handling_lawyers") or case_data.get("handling_lawyer") or linked_case.owner if linked_case else "",
                        "assistant": case_data.get("assistant") or case_data.get("lawyer_assistant", ""),
                        "contract_no": allocation.get("contract_no", ""),
                        "kind": _settlement_fee_kind(fee_type),
                    })
                continue
            case_key = str((linked_case.id if linked_case else 0) or allocation.get("case_no") or "")
            candidates = fees_by_case.get(case_key, [])
            for fee in candidates:
                if remaining <= 0.001:
                    break
                fee_data = fee.data or {}
                fee_total = abs(_round_fee_amount(float(fee_data.get("amount") or 0)))
                available = max(_round_fee_amount(fee_total - consumed_by_fee.get(fee.id, 0)), 0)
                current_amount = min(remaining, available)
                if current_amount <= 0.001:
                    continue
                fee_type = _case_fee_display_type(fee)
                kind = _settlement_fee_kind(fee_type)
                explicit_total = fee_data.get("settlement_amount")
                if explicit_total is not None and fee_total:
                    settlement_amount = _round_fee_amount(float(explicit_total) * current_amount / fee_total)
                elif kind == "agency":
                    settlement_amount = _round_fee_amount(current_amount * 0.8)
                else:
                    settlement_amount = current_amount
                explicit_archive = fee_data.get("archive_fee")
                if explicit_archive is not None and fee_total:
                    archive_fee = _round_fee_amount(float(explicit_archive) * current_amount / fee_total)
                elif kind == "agency" and "退费" not in fee_type:
                    archive_fee = _round_fee_amount(settlement_amount * 0.1)
                else:
                    archive_fee = 0.0
                details.append({
                    "fee_id": fee.id,
                    "case_id": linked_case.id if linked_case else allocation.get("case_id"),
                    "case_no": linked_case.serial_no if linked_case else allocation.get("case_no", ""),
                    "case_stage": (case_data.get("case_stage") or linked_case.status) if linked_case else "",
                    "fee_type": fee_type,
                    "fee_total_amount": fee_total,
                    "fee_allocated_amount": current_amount,
                    "current_amount": current_amount,
                    "allocated_at": allocation.get("allocated_at", ""),
                    "settlement_amount": settlement_amount,
                    "archive_fee": archive_fee,
                    "customer": linked_case.customer if linked_case else payment.claimed_customer,
                    "handling_lawyer": case_data.get("handling_lawyers") or case_data.get("handling_lawyer") or linked_case.owner if linked_case else "",
                    "assistant": case_data.get("assistant") or case_data.get("lawyer_assistant", ""),
                    "contract_no": allocation.get("contract_no", ""),
                    "kind": kind,
                })
                consumed_by_fee[fee.id] = _round_fee_amount(consumed_by_fee.get(fee.id, 0) + current_amount)
                remaining = _round_fee_amount(remaining - current_amount)
            if remaining > 0.001:
                details.append({
                    "fee_id": None,
                    "case_id": linked_case.id if linked_case else allocation.get("case_id"),
                    "case_no": linked_case.serial_no if linked_case else allocation.get("case_no", ""),
                    "case_stage": (case_data.get("case_stage") or linked_case.status) if linked_case else "",
                    "fee_type": "其他费用",
                    "fee_total_amount": remaining,
                    "fee_allocated_amount": remaining,
                    "current_amount": remaining,
                    "allocated_at": allocation.get("allocated_at", ""),
                    "settlement_amount": remaining,
                    "archive_fee": 0.0,
                    "customer": linked_case.customer if linked_case else payment.claimed_customer,
                    "handling_lawyer": case_data.get("handling_lawyers") or case_data.get("handling_lawyer") or linked_case.owner if linked_case else "",
                    "assistant": case_data.get("assistant") or case_data.get("lawyer_assistant", ""),
                    "contract_no": allocation.get("contract_no", ""),
                    "kind": "other",
                })
        if not details:
            continue
        for detail_index, detail in enumerate(details, start=1):
            detail["detail_id"] = f"{payment.id}-{detail_index}"
        detail_cases = [cases_by_id.get(int(item.get("case_id") or 0)) for item in details]
        detail_cases = [item for item in detail_cases if item]
        case_data_rows = [item.data or {} for item in detail_cases]
        row_customer = payment.claimed_customer or (details[0].get("customer") if details else "")
        row_case_customers = "、".join(dict.fromkeys(str(item.get("customer") or "") for item in details if item.get("customer")))
        row_case_nos = "、".join(dict.fromkeys(str(item.get("case_no") or "") for item in details if item.get("case_no")))
        row_hearing = "、".join(dict.fromkeys(str(data.get("hearing_lawyer") or "") for data in case_data_rows if data.get("hearing_lawyer")))
        row_assistant = "、".join(dict.fromkeys(str(data.get("assistant") or data.get("lawyer_assistant") or "") for data in case_data_rows if data.get("assistant") or data.get("lawyer_assistant")))
        row_manager = "、".join(dict.fromkeys(str(data.get("customer_manager") or "") for data in case_data_rows if data.get("customer_manager")))
        row_source = "、".join(dict.fromkeys(str(data.get("source_person") or data.get("case_source") or "") for data in case_data_rows if data.get("source_person") or data.get("case_source")))
        row_method = "、".join(dict.fromkeys(str(item.get("payment_method") or "") for item in payment_allocations if item.get("payment_method")))
        if not contains(row_customer, customer) or not contains(row_case_nos, case_no):
            continue
        if received_from and payment.received_date < received_from:
            continue
        if received_to and payment.received_date > received_to:
            continue
        if not contains(payment.payer_name, payer) or not contains(row_method, payment_method):
            continue
        if not contains(row_case_customers, case_customer) or not contains(row_hearing, hearing_lawyer):
            continue
        if not contains(row_assistant, assistant) or not contains(row_manager, customer_manager) or not contains(row_source, source_person):
            continue
        assigned = _round_fee_amount(sum(float(item["current_amount"]) for item in details))
        official = _round_fee_amount(sum(float(item["current_amount"]) for item in details if item["kind"] == "official"))
        agency = _round_fee_amount(sum(float(item["current_amount"]) for item in details if item["kind"] == "agency"))
        other = _round_fee_amount(sum(float(item["current_amount"]) for item in details if item["kind"] == "other"))
        agency_settlement = _round_fee_amount(sum(float(item["settlement_amount"]) for item in details if item["kind"] == "agency"))
        archive_fee = _round_fee_amount(sum(float(item["archive_fee"]) for item in details))
        actual = _round_fee_amount(sum(float(item["settlement_amount"]) for item in details) - archive_fee)
        rows.append({
            "id": payment.id,
            "serial_no": payment.receipt_no,
            "title": f"{row_customer}待结算回款",
            "customer": row_customer,
            "status": "待结算",
            "owner": payment.claimant or payment.operator,
            "data": {
                "receipt_id": payment.id,
                "receipt_no": payment.receipt_no,
                "customer_manager": row_manager,
                "payer_name": payment.payer_name,
                "received_date": str(payment.received_date),
                "receipt_amount": float(payment.amount) if can_view_amount else None,
                "allocated_amount": assigned if can_view_amount else None,
                "remaining_amount": max(_round_fee_amount(float(payment.amount) - assigned), 0) if can_view_amount else None,
                "assigned_official_fee": official if can_view_amount else None,
                "assigned_agency_fee": agency if can_view_amount else None,
                "assigned_other_fee": other if can_view_amount else None,
                "agency_settlement_amount": agency_settlement if can_view_amount else None,
                "archive_fee": archive_fee if can_view_amount else None,
                "actual_settlement_amount": actual if can_view_amount else None,
                "payment_method": row_method,
                "bank_remark": payment.remark or payment.bank_reference,
                "rejection_comment": rejection_by_receipt.get(payment.id, ""),
                "case_nos": row_case_nos,
                "case_customers": row_case_customers,
                "hearing_lawyer": row_hearing,
                "assistant": row_assistant,
                "source_person": row_source,
                "allocation_details": details if can_view_amount else [{**detail, "fee_total_amount": None, "fee_allocated_amount": None, "current_amount": None, "settlement_amount": None, "archive_fee": None} for detail in details],
            },
        })
    return sorted(rows, key=lambda row: (row["data"]["received_date"], row["id"]), reverse=True)


@app.get(f"{settings.api_prefix}/finance/general-settlements/pending")
async def list_general_settlement_candidates(
    customer: str = "", case_no: str = "",
    received_from: date | None = None, received_to: date | None = None,
    payer: str = "", payment_method: str = "", case_customer: str = "",
    hearing_lawyer: str = "", assistant: str = "", customer_manager: str = "", source_person: str = "",
    page: int = Query(1, ge=1), page_size: int = Query(10, ge=1, le=200),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    if received_from and received_to and received_from > received_to:
        raise HTTPException(status_code=422, detail="回款开始日期不能晚于结束日期")
    rows = await _general_settlement_rows(
        identity, db, customer=customer, case_no=case_no,
        received_from=received_from, received_to=received_to, payer=payer,
        payment_method=payment_method, case_customer=case_customer,
        hearing_lawyer=hearing_lawyer, assistant=assistant,
        customer_manager=customer_manager, source_person=source_person,
    )
    amount_keys = ["receipt_amount", "allocated_amount", "remaining_amount", "assigned_official_fee", "assigned_agency_fee", "assigned_other_fee", "agency_settlement_amount", "archive_fee", "actual_settlement_amount"]
    totals = {key: _round_fee_amount(sum(float((row.get("data") or {}).get(key) or 0) for row in rows)) for key in amount_keys}
    start = (page - 1) * page_size
    return {"items": rows[start:start + page_size], "total": len(rows), "totals": totals, "page": page, "page_size": page_size}


@app.post(f"{settings.api_prefix}/finance/general-settlements/apply", status_code=status.HTTP_201_CREATED)
async def apply_general_settlements(body: FinanceSettlementApplyInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    receipt_ids = list(dict.fromkeys(body.receipt_ids))
    rows = await _general_settlement_rows(identity, db, receipt_ids=set(receipt_ids))
    rows_by_id = {int(row["id"]): row for row in rows}
    missing = [str(receipt_id) for receipt_id in receipt_ids if receipt_id not in rows_by_id]
    if missing:
        raise HTTPException(status_code=409, detail="部分回款已申请结算、尚未分配或无权办理：" + "、".join(missing))
    created: list[BusinessRecord] = []
    now_key = datetime.now().strftime("%Y%m%d%H%M%S%f")
    for index, receipt_id in enumerate(receipt_ids):
        row = rows_by_id[receipt_id]
        row_data = dict(row.get("data") or {})
        application = BusinessRecord(
            module="finance_settlement",
            serial_no=f"JS{now_key}{index:02d}",
            title=f"{row.get('customer') or row_data.get('payer_name')}结算申请",
            customer=str(row.get("customer") or ""),
            status="待审批",
            owner=identity["username"],
            department=str(identity.get("department") or "上海分所"),
            description=body.comment.strip(),
            data={**row_data, "applied_by": identity["username"], "applied_at": datetime.now().isoformat(timespec="seconds")},
        )
        db.add(application)
        await db.flush()
        db.add(WorkflowEvent(record_id=application.id, action="申请结算", to_status="待审批", operator=identity["username"], comment=body.comment.strip()))
        created.append(application)
    await db.commit()
    return {"created": len(created), "application_ids": [item.id for item in created], "application_nos": [item.serial_no for item in created]}


def _settlement_application_scope(identity: dict) -> list:
    if identity.get("role") in {"admin", "auditor"}:
        return []
    if identity.get("role") == "manager":
        return [BusinessRecord.department == str(identity.get("department") or "")]
    return [BusinessRecord.owner == identity["username"]]


async def _pending_archive_settlement_rows(
    identity: dict,
    db: AsyncSession,
    *,
    case_type: str = "",
    case_stage: str = "",
    payer: str = "",
    received_from: date | None = None,
    received_to: date | None = None,
    hearing_lawyer: str = "",
    assistant: str = "",
    submitted_by: str = "",
    settled_from: date | None = None,
    settled_to: date | None = None,
    case_no: str = "",
    customer: str = "",
    reviewer: str = "",
    archive_from: date | None = None,
    archive_to: date | None = None,
    require_archived: bool = False,
    selected_ids: set[str] | None = None,
) -> list[dict]:
    """Flatten paid settlement details into the pending-archive or pending-payment list."""
    records = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "finance_settlement",
        BusinessRecord.status == "已付款",
        *_settlement_application_scope(identity),
    ).order_by(BusinessRecord.id.asc()))).all())
    details = [
        detail
        for record in records
        for detail in list((record.data or {}).get("allocation_details") or [])
        if float(detail.get("archive_fee") or 0) > 0.001
    ]
    case_ids = {int(detail.get("case_id") or 0) for detail in details if detail.get("case_id")}
    case_nos = {str(detail.get("case_no") or "").strip() for detail in details if str(detail.get("case_no") or "").strip()}
    cases = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "case",
        or_(BusinessRecord.id.in_(case_ids), BusinessRecord.serial_no.in_(case_nos)),
        *(await _record_scope_conditions(identity, db)),
    ))).all()) if case_ids or case_nos else []
    cases_by_id = {item.id: item for item in cases}
    cases_by_no = {item.serial_no: item for item in cases}
    archive_events: dict[int, list[WorkflowEvent]] = {}
    visible_case_ids = {item.id for item in cases}
    if require_archived and visible_case_ids:
        for event in (await db.scalars(select(WorkflowEvent).where(
            WorkflowEvent.record_id.in_(visible_case_ids),
            WorkflowEvent.action.in_({"提交归档审核", "归档审核通过"}),
        ).order_by(WorkflowEvent.created_at.desc(), WorkflowEvent.id.desc()))).all():
            archive_events.setdefault(event.record_id, []).append(event)
    decisions = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "finance_archive_settlement",
        BusinessRecord.status.in_({"已支付", "已拒绝"}),
        *_settlement_application_scope(identity),
    ))).all()) if require_archived else []
    decided_source_ids = {str((item.data or {}).get("source_row_id") or "") for item in decisions}
    can_view_amount = "finance.amount" in await _allowed_field_keys(identity, db)

    def contains(value: object, needle: str) -> bool:
        return not needle.strip() or needle.strip().casefold() in str(value or "").casefold()

    def date_in_range(value: object, start_date: date | None, end_date: date | None) -> bool:
        if not start_date and not end_date:
            return True
        try:
            current = date.fromisoformat(str(value or "")[:10])
        except ValueError:
            return False
        return (not start_date or current >= start_date) and (not end_date or current <= end_date)

    rows: list[dict] = []
    for record in records:
        data = record.data or {}
        for index, detail in enumerate(list(data.get("allocation_details") or []), start=1):
            archive_fee = _round_fee_amount(float(detail.get("archive_fee") or 0))
            if archive_fee <= 0.001:
                continue
            linked_case = cases_by_id.get(int(detail.get("case_id") or 0)) or cases_by_no.get(str(detail.get("case_no") or ""))
            if not linked_case:
                continue
            if require_archived and linked_case.status != "已归档":
                continue
            # Once a case is archived, the original workflow moves the fee to the pending-payment page.
            if not require_archived and linked_case.status == "已归档":
                continue
            case_data = linked_case.data or {}
            row_id = f"{record.id}:{detail.get('detail_id') or index}"
            if require_archived and row_id in decided_source_ids:
                continue
            if selected_ids is not None and row_id not in selected_ids:
                continue
            row_case_type = str(case_data.get("case_type") or "")
            if row_case_type in {"民事案件", "民事"}:
                row_case_type = "民事争议"
            row_case_stage = str(case_data.get("case_stage") or linked_case.status or "")
            row_hearing = case_data.get("hearing_lawyer") or detail.get("handling_lawyer") or ""
            row_assistant = case_data.get("assistant") or case_data.get("lawyer_assistant") or detail.get("assistant") or ""
            row_submitted_by = data.get("applied_by") or record.owner
            row_settled_at = data.get("paid_at") or record.updated_at
            case_events = archive_events.get(linked_case.id, [])
            archive_review_event = next((item for item in case_events if item.action == "归档审核通过"), None)
            archive_submit_event = next((item for item in case_events if item.action == "提交归档审核"), None)
            archive_date = case_data.get("archived_at") or (archive_review_event.created_at if archive_review_event else "")
            if not contains(row_case_type, case_type) or not contains(row_case_stage, case_stage):
                continue
            if not contains(data.get("payer_name"), payer) or not date_in_range(data.get("received_date"), received_from, received_to):
                continue
            if not contains(row_hearing, hearing_lawyer) or not contains(row_assistant, assistant):
                continue
            if not contains(row_submitted_by, submitted_by) or not date_in_range(row_settled_at, settled_from, settled_to):
                continue
            if not contains(linked_case.serial_no, case_no) or not contains(linked_case.customer, customer):
                continue
            if not contains(data.get("reviewer"), reviewer):
                continue
            if require_archived and not date_in_range(archive_date, archive_from, archive_to):
                continue
            rows.append({
                "id": row_id,
                "serial_no": row_id,
                "title": f"{linked_case.serial_no}归档费",
                "customer": linked_case.customer,
                "status": "待支付" if require_archived else "待归档",
                "owner": record.owner,
                "department": record.department,
                "created_at": record.created_at.isoformat() if record.created_at else "",
                "updated_at": record.updated_at.isoformat() if record.updated_at else "",
                "data": {
                    "application_id": record.id,
                    "receipt_id": data.get("receipt_id"),
                    "case_id": linked_case.id,
                    "case_no": linked_case.serial_no,
                    "case_type": row_case_type,
                    "case_stage": row_case_stage,
                    "assistant": row_assistant,
                    "hearing_lawyer": row_hearing,
                    "customer_manager": case_data.get("customer_manager") or "",
                    "fee_type": detail.get("fee_type") or "律师代理费",
                    "payment_method": data.get("payment_method") or "",
                    "payer_name": data.get("payer_name") or "",
                    "received_date": data.get("received_date") or "",
                    "receipt_amount": _round_fee_amount(float(detail.get("current_amount") or 0)) if can_view_amount else None,
                    "archive_fee_amount": archive_fee if can_view_amount else None,
                    "settlement_paid_at": str(row_settled_at or ""),
                    "submitted_by": row_submitted_by,
                    "reviewer": data.get("reviewer") or "",
                    "archive_reviewer": case_data.get("archive_reviewer") or (archive_review_event.operator if archive_review_event else ""),
                    "archive_reviewed_at": str(archive_review_event.created_at if archive_review_event else case_data.get("archived_at") or ""),
                    "archive_review_comment": case_data.get("archive_review_comment") or (archive_review_event.comment if archive_review_event else ""),
                    "archive_submitter": archive_submit_event.operator if archive_submit_event else linked_case.owner,
                    "archive_submitted_at": str(archive_submit_event.created_at if archive_submit_event else ""),
                    "archive_submit_comment": archive_submit_event.comment if archive_submit_event else "",
                    "archive_no": case_data.get("archive_no") or "",
                    "archive_date": str(archive_date or "")[:10],
                    "archive_status": "审核通过" if require_archived else "",
                },
            })
    return sorted(rows, key=lambda row: (str((row.get("data") or {}).get("case_no") or ""), str((row.get("data") or {}).get("received_date") or ""), row["id"]))


@app.get(f"{settings.api_prefix}/finance/archive-settlements/pending")
async def list_pending_archive_settlements(
    case_type: str = "", case_stage: str = "", payer: str = "",
    received_from: date | None = None, received_to: date | None = None,
    hearing_lawyer: str = "", assistant: str = "", submitted_by: str = "",
    settled_from: date | None = None, settled_to: date | None = None,
    case_no: str = "", customer: str = "", reviewer: str = "",
    page: int = Query(1, ge=1), page_size: int = Query(10, ge=1, le=200),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    for start_date, end_date, label in (
        (received_from, received_to, "回款"),
        (settled_from, settled_to, "结算支付"),
    ):
        if start_date and end_date and start_date > end_date:
            raise HTTPException(status_code=422, detail=f"{label}开始日期不能晚于结束日期")
    rows = await _pending_archive_settlement_rows(
        identity, db, case_type=case_type, case_stage=case_stage, payer=payer,
        received_from=received_from, received_to=received_to,
        hearing_lawyer=hearing_lawyer, assistant=assistant, submitted_by=submitted_by,
        settled_from=settled_from, settled_to=settled_to, case_no=case_no,
        customer=customer, reviewer=reviewer,
    )
    totals = {
        "receipt_amount": _round_fee_amount(sum(float((row.get("data") or {}).get("receipt_amount") or 0) for row in rows)),
        "archive_fee_amount": _round_fee_amount(sum(float((row.get("data") or {}).get("archive_fee_amount") or 0) for row in rows)),
    }
    start = (page - 1) * page_size
    return {"items": rows[start:start + page_size], "total": len(rows), "totals": totals, "page": page, "page_size": page_size}


@app.get(f"{settings.api_prefix}/finance/archive-settlements/export")
async def export_pending_archive_settlements(
    ids: str = "", identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    selected_ids = {item.strip() for item in ids.split(",") if item.strip()}
    if not selected_ids:
        raise HTTPException(status_code=422, detail="请选择需要导出的归档费.")
    rows = await _pending_archive_settlement_rows(identity, db, selected_ids=selected_ids)
    if len(rows) != len(selected_ids):
        raise HTTPException(status_code=409, detail="部分归档费不存在、已进入下一环节或无权导出")
    headers = ["案号", "客户", "案件阶段", "律师助理", "开庭律师", "客户管理人", "费用类型", "回款方式", "回款时间", "回款金额", "归档费金额", "结算时间"]
    numeric = {9, 10}
    values = [[
        row["data"].get("case_no"), row.get("customer"), row["data"].get("case_stage"),
        row["data"].get("assistant"), row["data"].get("hearing_lawyer"), row["data"].get("customer_manager"),
        row["data"].get("fee_type"), row["data"].get("payment_method"), row["data"].get("received_date"),
        row["data"].get("receipt_amount"), row["data"].get("archive_fee_amount"), row["data"].get("settlement_paid_at"),
    ] for row in rows]
    def cell(value: object, *, number: bool = False) -> str:
        value_text = f"{float(value or 0):.2f}" if number else str(value or "")
        return f'<Cell><Data ss:Type="{"Number" if number else "String"}">{xml_escape(value_text)}</Data></Cell>'
    sheet_rows = ["<Row>" + "".join(cell(value) for value in headers) + "</Row>"]
    sheet_rows.extend("<Row>" + "".join(cell(value, number=index in numeric and value is not None) for index, value in enumerate(row)) + "</Row>" for row in values)
    workbook = '<?xml version="1.0" encoding="UTF-8"?><?mso-application progid="Excel.Sheet"?><Workbook xmlns="urn:schemas-microsoft-com:office:spreadsheet" xmlns:ss="urn:schemas-microsoft-com:office:spreadsheet"><Worksheet ss:Name="待归档"><Table>' + "".join(sheet_rows) + "</Table></Worksheet></Workbook>"
    filename = f"待归档-{date.today()}.xls"
    disposition = f"attachment; filename=archive-settlement-pending.xls; filename*=UTF-8''{quote(filename)}"
    return Response(content=workbook.encode("utf-8"), media_type="application/vnd.ms-excel", headers={"Content-Disposition": disposition})


@app.get(f"{settings.api_prefix}/finance/archive-settlements/payment")
async def list_archive_settlement_payments(
    case_type: str = "", case_stage: str = "", payer: str = "",
    received_from: date | None = None, received_to: date | None = None,
    hearing_lawyer: str = "", assistant: str = "", submitted_by: str = "",
    settled_from: date | None = None, settled_to: date | None = None,
    case_no: str = "", customer: str = "", reviewer: str = "",
    archive_from: date | None = None, archive_to: date | None = None,
    page: int = Query(1, ge=1), page_size: int = Query(10, ge=1, le=200),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    for start_date, end_date, label in (
        (received_from, received_to, "回款"),
        (settled_from, settled_to, "结算支付"),
        (archive_from, archive_to, "归档"),
    ):
        if start_date and end_date and start_date > end_date:
            raise HTTPException(status_code=422, detail=f"{label}开始日期不能晚于结束日期")
    rows = await _pending_archive_settlement_rows(
        identity, db, case_type=case_type, case_stage=case_stage, payer=payer,
        received_from=received_from, received_to=received_to,
        hearing_lawyer=hearing_lawyer, assistant=assistant, submitted_by=submitted_by,
        settled_from=settled_from, settled_to=settled_to, case_no=case_no,
        customer=customer, reviewer=reviewer, archive_from=archive_from,
        archive_to=archive_to, require_archived=True,
    )
    totals = {
        "receipt_amount": _round_fee_amount(sum(float((row.get("data") or {}).get("receipt_amount") or 0) for row in rows)),
        "archive_fee_amount": _round_fee_amount(sum(float((row.get("data") or {}).get("archive_fee_amount") or 0) for row in rows)),
    }
    start = (page - 1) * page_size
    return {"items": rows[start:start + page_size], "total": len(rows), "totals": totals, "page": page, "page_size": page_size}


@app.post(f"{settings.api_prefix}/finance/archive-settlements/payment/review")
async def review_archive_settlement_payments(
    body: ArchiveSettlementPaymentReviewInput,
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    if identity.get("role") not in {"admin", "manager", "auditor"}:
        raise HTTPException(status_code=403, detail="当前角色没有归档费支付审核权限")
    settlement_ids = list(dict.fromkeys(item.strip() for item in body.settlement_ids if item.strip()))
    if len(settlement_ids) != len(body.settlement_ids):
        raise HTTPException(status_code=422, detail="归档费记录不能为空或重复")
    if not body.approved and not body.comment.strip():
        raise HTTPException(status_code=422, detail="拒绝支付时请输入备注.")
    rows = await _pending_archive_settlement_rows(
        identity, db, require_archived=True, selected_ids=set(settlement_ids),
    )
    if len(rows) != len(settlement_ids):
        raise HTTPException(status_code=409, detail="部分归档费不存在、已处理或无权审核")
    row_map = {row["id"]: row for row in rows}
    decided_at = datetime.now().isoformat(timespec="seconds")
    reusable = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "finance_archive_settlement",
        BusinessRecord.status == "已回滚",
        *_settlement_application_scope(identity),
    ))).all())
    reusable_by_source = {str((item.data or {}).get("source_row_id") or ""): item for item in reusable}
    created: list[BusinessRecord] = []
    for source_id in settlement_ids:
        row = row_map[source_id]
        data = row.get("data") or {}
        target_status = "已支付" if body.approved else "已拒绝"
        decision = reusable_by_source.get(source_id)
        if decision:
            previous_status = decision.status
            decision.status = target_status
            decision.title = row.get("title") or "归档费支付"
            decision.customer = row.get("customer") or ""
            decision.owner = row.get("owner") or identity["username"]
            decision.department = row.get("department") or str(identity.get("department") or "")
            decision.data = {
                **data,
                "source_row_id": source_id,
                "source_application_id": data.get("application_id"),
                "archive_payment_submitted_by": data.get("submitted_by") or row.get("owner") or identity["username"],
                "archive_payment_submitted_at": data.get("settlement_paid_at") or decided_at,
                "archive_payment_reviewer": identity["username"],
                "archive_payment_reviewed_at": decided_at,
                "archive_payment_comment": body.comment.strip(),
            }
        else:
            previous_status = "待支付"
            decision = BusinessRecord(
                module="finance_archive_settlement",
                serial_no=f"ARCP-{source_id.replace(':', '-')}",
                title=row.get("title") or "归档费支付",
                customer=row.get("customer") or "",
                status=target_status,
                owner=row.get("owner") or identity["username"],
                department=row.get("department") or str(identity.get("department") or ""),
                data={
                    **data,
                    "source_row_id": source_id,
                    "source_application_id": data.get("application_id"),
                    "archive_payment_submitted_by": data.get("submitted_by") or row.get("owner") or identity["username"],
                    "archive_payment_submitted_at": data.get("settlement_paid_at") or decided_at,
                    "archive_payment_reviewer": identity["username"],
                    "archive_payment_reviewed_at": decided_at,
                    "archive_payment_comment": body.comment.strip(),
                },
            )
        db.add(decision)
        await db.flush()
        db.add(WorkflowEvent(
            record_id=decision.id,
            action="归档费同意支付" if body.approved else "归档费拒绝支付",
            from_status=previous_status, to_status=target_status,
            operator=identity["username"], comment=body.comment.strip(),
        ))
        created.append(decision)
    await db.commit()
    return {"reviewed": len(created), "status": "已支付" if body.approved else "已拒绝", "record_ids": [item.id for item in created]}


@app.get(f"{settings.api_prefix}/finance/archive-settlements/payment/export")
async def export_archive_settlement_payments(
    ids: str = "", identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    selected_ids = {item.strip() for item in ids.split(",") if item.strip()}
    if not selected_ids:
        raise HTTPException(status_code=422, detail="请选择需要导出的归档费.")
    rows = await _pending_archive_settlement_rows(identity, db, require_archived=True, selected_ids=selected_ids)
    if len(rows) != len(selected_ids):
        raise HTTPException(status_code=409, detail="部分归档费不存在、已处理或无权导出")
    headers = ["案号", "客户", "案件阶段", "律师助理", "开庭律师", "客户管理人", "费用类型", "回款方式", "回款时间", "回款金额", "归档费金额", "支付时间", "归档号", "归档日期"]
    numeric = {9, 10}
    values = [[
        row["data"].get("case_no"), row.get("customer"), row["data"].get("case_stage"),
        row["data"].get("assistant"), row["data"].get("hearing_lawyer"), row["data"].get("customer_manager"),
        row["data"].get("fee_type"), row["data"].get("payment_method"), row["data"].get("received_date"),
        row["data"].get("receipt_amount"), row["data"].get("archive_fee_amount"), row["data"].get("settlement_paid_at"),
        row["data"].get("archive_no"), row["data"].get("archive_date"),
    ] for row in rows]
    def cell(value: object, *, number: bool = False) -> str:
        value_text = f"{float(value or 0):.2f}" if number else str(value or "")
        return f'<Cell><Data ss:Type="{"Number" if number else "String"}">{xml_escape(value_text)}</Data></Cell>'
    sheet_rows = ["<Row>" + "".join(cell(value) for value in headers) + "</Row>"]
    sheet_rows.extend("<Row>" + "".join(cell(value, number=index in numeric) for index, value in enumerate(row)) + "</Row>" for row in values)
    workbook = '<?xml version="1.0" encoding="UTF-8"?><?mso-application progid="Excel.Sheet"?><Workbook xmlns="urn:schemas-microsoft-com:office:spreadsheet" xmlns:ss="urn:schemas-microsoft-com:office:spreadsheet"><Worksheet ss:Name="待支付"><Table>' + "".join(sheet_rows) + "</Table></Worksheet></Workbook>"
    filename = f"待支付归档费-{date.today()}.xls"
    disposition = f"attachment; filename=archive-settlement-payment.xls; filename*=UTF-8''{quote(filename)}"
    return Response(content=workbook.encode("utf-8"), media_type="application/vnd.ms-excel", headers={"Content-Disposition": disposition})


async def _archive_settlement_decision_rows(
    identity: dict, db: AsyncSession, *, statuses: set[str],
    case_type: str = "", case_stage: str = "", payer: str = "",
    received_from: date | None = None, received_to: date | None = None,
    hearing_lawyer: str = "", assistant: str = "", submitted_by: str = "",
    settled_from: date | None = None, settled_to: date | None = None,
    submitted_from: date | None = None, submitted_to: date | None = None,
    case_no: str = "", customer: str = "", reviewer: str = "",
    reviewed_from: date | None = None, reviewed_to: date | None = None,
    archive_from: date | None = None, archive_to: date | None = None,
    payment_from: date | None = None, payment_to: date | None = None,
    selected_ids: set[int] | None = None,
) -> list[dict]:
    records = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "finance_archive_settlement",
        BusinessRecord.status.in_(statuses),
        *_settlement_application_scope(identity),
    ).order_by(BusinessRecord.updated_at.desc(), BusinessRecord.id.desc()))).all())
    can_view_amount = "finance.amount" in await _allowed_field_keys(identity, db)

    def contains(value: object, needle: str) -> bool:
        return not needle.strip() or needle.strip().casefold() in str(value or "").casefold()

    def date_in_range(value: object, start_date: date | None, end_date: date | None) -> bool:
        if not start_date and not end_date:
            return True
        try:
            current = date.fromisoformat(str(value or "")[:10])
        except ValueError:
            return False
        return (not start_date or current >= start_date) and (not end_date or current <= end_date)

    rows: list[dict] = []
    for record in records:
        if selected_ids is not None and record.id not in selected_ids:
            continue
        data = dict(record.data or {})
        if not contains(data.get("case_type"), case_type) or not contains(data.get("case_stage"), case_stage):
            continue
        if not contains(data.get("payer_name"), payer) or not date_in_range(data.get("received_date"), received_from, received_to):
            continue
        if not contains(data.get("hearing_lawyer"), hearing_lawyer) or not contains(data.get("assistant"), assistant):
            continue
        if not contains(data.get("archive_payment_submitted_by") or data.get("submitted_by"), submitted_by):
            continue
        if not date_in_range(data.get("settlement_paid_at"), settled_from, settled_to):
            continue
        if not date_in_range(data.get("archive_payment_submitted_at") or data.get("settlement_paid_at"), submitted_from, submitted_to):
            continue
        if not contains(data.get("case_no"), case_no) or not contains(record.customer, customer):
            continue
        if not contains(data.get("archive_payment_reviewer") or data.get("reviewer"), reviewer):
            continue
        if not date_in_range(data.get("archive_payment_reviewed_at"), reviewed_from, reviewed_to):
            continue
        if not date_in_range(data.get("archive_date"), archive_from, archive_to):
            continue
        if not date_in_range(data.get("archive_payment_reviewed_at"), payment_from, payment_to):
            continue
        if not can_view_amount:
            data["receipt_amount"] = None
            data["archive_fee_amount"] = None
        rows.append({
            "id": record.id,
            "serial_no": record.serial_no,
            "title": record.title,
            "customer": record.customer,
            "status": record.status,
            "owner": record.owner,
            "department": record.department,
            "created_at": record.created_at.isoformat() if record.created_at else "",
            "updated_at": record.updated_at.isoformat() if record.updated_at else "",
            "data": data,
        })
    return rows


@app.get(f"{settings.api_prefix}/finance/archive-settlements/paid")
async def list_paid_archive_settlements(
    case_type: str = "", case_stage: str = "", payer: str = "",
    received_from: date | None = None, received_to: date | None = None,
    hearing_lawyer: str = "", assistant: str = "", submitted_by: str = "",
    settled_from: date | None = None, settled_to: date | None = None,
    case_no: str = "", customer: str = "", reviewer: str = "",
    archive_from: date | None = None, archive_to: date | None = None,
    payment_from: date | None = None, payment_to: date | None = None,
    page: int = Query(1, ge=1), page_size: int = Query(10, ge=1, le=200),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    for start_date, end_date, label in (
        (received_from, received_to, "回款"), (settled_from, settled_to, "结算支付"),
        (archive_from, archive_to, "归档"), (payment_from, payment_to, "归档费支付"),
    ):
        if start_date and end_date and start_date > end_date:
            raise HTTPException(status_code=422, detail=f"{label}开始日期不能晚于结束日期")
    rows = await _archive_settlement_decision_rows(
        identity, db, statuses={"已支付"}, case_type=case_type, case_stage=case_stage,
        payer=payer, received_from=received_from, received_to=received_to,
        hearing_lawyer=hearing_lawyer, assistant=assistant, submitted_by=submitted_by,
        settled_from=settled_from, settled_to=settled_to, case_no=case_no,
        customer=customer, reviewer=reviewer, archive_from=archive_from,
        archive_to=archive_to, payment_from=payment_from, payment_to=payment_to,
    )
    totals = {
        "receipt_amount": _round_fee_amount(sum(float((row.get("data") or {}).get("receipt_amount") or 0) for row in rows)),
        "archive_fee_amount": _round_fee_amount(sum(float((row.get("data") or {}).get("archive_fee_amount") or 0) for row in rows)),
    }
    start = (page - 1) * page_size
    return {"items": rows[start:start + page_size], "total": len(rows), "totals": totals, "page": page, "page_size": page_size}


@app.post(f"{settings.api_prefix}/finance/archive-settlements/paid/rollback")
async def rollback_paid_archive_settlements(
    body: ArchiveSettlementRollbackInput,
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    if identity.get("role") not in {"admin", "manager", "auditor"}:
        raise HTTPException(status_code=403, detail="当前角色没有归档费支付回滚权限")
    if not body.comment.strip():
        raise HTTPException(status_code=422, detail="请输入备注.")
    record_ids = list(dict.fromkeys(body.record_ids))
    if len(record_ids) != len(body.record_ids):
        raise HTTPException(status_code=422, detail="归档费记录不能重复")
    records = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.id.in_(record_ids),
        BusinessRecord.module == "finance_archive_settlement",
        BusinessRecord.status == "已支付",
        *_settlement_application_scope(identity),
    ))).all())
    if len(records) != len(record_ids):
        raise HTTPException(status_code=409, detail="部分归档费不存在、不是已支付状态或无权回滚")
    rolled_back_at = datetime.now().isoformat(timespec="seconds")
    for record in records:
        record.status = "已回滚"
        record.data = {
            **(record.data or {}),
            "archive_payment_rollback_by": identity["username"],
            "archive_payment_rollback_at": rolled_back_at,
            "archive_payment_rollback_comment": body.comment.strip(),
        }
        db.add(WorkflowEvent(
            record_id=record.id, action="回滚归档费支付",
            from_status="已支付", to_status="已回滚",
            operator=identity["username"], comment=body.comment.strip(),
        ))
    await db.commit()
    return {"rolled_back": len(records), "status": "已回滚"}


@app.get(f"{settings.api_prefix}/finance/archive-settlements/paid/export")
async def export_paid_archive_settlements(
    ids: str = "", identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    try:
        selected_ids = {int(item.strip()) for item in ids.split(",") if item.strip()}
    except ValueError:
        raise HTTPException(status_code=422, detail="归档费记录编号无效")
    if not selected_ids:
        raise HTTPException(status_code=422, detail="请选择需要导出的归档费.")
    rows = await _archive_settlement_decision_rows(identity, db, statuses={"已支付"}, selected_ids=selected_ids)
    if len(rows) != len(selected_ids):
        raise HTTPException(status_code=409, detail="部分归档费不存在、已回滚或无权导出")
    headers = ["案号", "客户", "案件阶段", "律师助理", "开庭律师", "客户管理人", "费用类型", "回款方式", "回款时间", "回款金额", "归档费金额", "结算时间", "归档费支付日期"]
    numeric = {9, 10}
    values = [[
        row["data"].get("case_no"), row.get("customer"), row["data"].get("case_stage"),
        row["data"].get("assistant"), row["data"].get("hearing_lawyer"), row["data"].get("customer_manager"),
        row["data"].get("fee_type"), row["data"].get("payment_method"), row["data"].get("received_date"),
        row["data"].get("receipt_amount"), row["data"].get("archive_fee_amount"), row["data"].get("settlement_paid_at"),
        row["data"].get("archive_payment_reviewed_at"),
    ] for row in rows]
    def cell(value: object, *, number: bool = False) -> str:
        value_text = f"{float(value or 0):.2f}" if number else str(value or "")
        return f'<Cell><Data ss:Type="{"Number" if number else "String"}">{xml_escape(value_text)}</Data></Cell>'
    sheet_rows = ["<Row>" + "".join(cell(value) for value in headers) + "</Row>"]
    sheet_rows.extend("<Row>" + "".join(cell(value, number=index in numeric and value is not None) for index, value in enumerate(row)) + "</Row>" for row in values)
    workbook = '<?xml version="1.0" encoding="UTF-8"?><?mso-application progid="Excel.Sheet"?><Workbook xmlns="urn:schemas-microsoft-com:office:spreadsheet" xmlns:ss="urn:schemas-microsoft-com:office:spreadsheet"><Worksheet ss:Name="已支付"><Table>' + "".join(sheet_rows) + "</Table></Worksheet></Workbook>"
    filename = f"已支付归档费-{date.today()}.xls"
    disposition = f"attachment; filename=archive-settlement-paid.xls; filename*=UTF-8''{quote(filename)}"
    return Response(content=workbook.encode("utf-8"), media_type="application/vnd.ms-excel", headers={"Content-Disposition": disposition})


@app.get(f"{settings.api_prefix}/finance/archive-settlements/rejected")
async def list_rejected_archive_settlements(
    case_type: str = "", case_stage: str = "", payer: str = "",
    received_from: date | None = None, received_to: date | None = None,
    hearing_lawyer: str = "", assistant: str = "", submitted_by: str = "",
    submitted_from: date | None = None, submitted_to: date | None = None,
    case_no: str = "", customer: str = "", reviewer: str = "",
    reviewed_from: date | None = None, reviewed_to: date | None = None,
    page: int = Query(1, ge=1), page_size: int = Query(10, ge=1, le=200),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    for start_date, end_date, label in (
        (received_from, received_to, "回款"),
        (submitted_from, submitted_to, "提交"),
        (reviewed_from, reviewed_to, "审核"),
    ):
        if start_date and end_date and start_date > end_date:
            raise HTTPException(status_code=422, detail=f"{label}开始日期不能晚于结束日期")
    rows = await _archive_settlement_decision_rows(
        identity, db, statuses={"已拒绝"}, case_type=case_type, case_stage=case_stage,
        payer=payer, received_from=received_from, received_to=received_to,
        hearing_lawyer=hearing_lawyer, assistant=assistant, submitted_by=submitted_by,
        submitted_from=submitted_from, submitted_to=submitted_to,
        case_no=case_no, customer=customer, reviewer=reviewer,
        reviewed_from=reviewed_from, reviewed_to=reviewed_to,
    )
    totals = {
        "receipt_amount": _round_fee_amount(sum(float((row.get("data") or {}).get("receipt_amount") or 0) for row in rows)),
        "archive_fee_amount": _round_fee_amount(sum(float((row.get("data") or {}).get("archive_fee_amount") or 0) for row in rows)),
    }
    start = (page - 1) * page_size
    return {"items": rows[start:start + page_size], "total": len(rows), "totals": totals, "page": page, "page_size": page_size}


async def _rejected_archive_settlement_records(
    record_ids: list[int], identity: dict, db: AsyncSession,
) -> list[BusinessRecord]:
    unique_ids = list(dict.fromkeys(record_ids))
    if len(unique_ids) != len(record_ids):
        raise HTTPException(status_code=422, detail="归档费记录不能重复")
    records = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.id.in_(unique_ids),
        BusinessRecord.module == "finance_archive_settlement",
        BusinessRecord.status == "已拒绝",
        *_settlement_application_scope(identity),
    ))).all())
    if len(records) != len(unique_ids):
        raise HTTPException(status_code=409, detail="部分归档费不存在、不是已拒绝状态或无权处理")
    return records


@app.post(f"{settings.api_prefix}/finance/archive-settlements/rejected/rollback")
async def rollback_rejected_archive_settlements(
    body: ArchiveSettlementRejectedActionInput,
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    if identity.get("role") not in {"admin", "manager", "auditor"}:
        raise HTTPException(status_code=403, detail="当前角色没有归档费拒绝回滚权限")
    if not body.comment.strip():
        raise HTTPException(status_code=422, detail="请输入审核备注.")
    records = await _rejected_archive_settlement_records(body.record_ids, identity, db)
    changed_at = datetime.now().isoformat(timespec="seconds")
    for record in records:
        record.status = "已支付"
        record.data = {
            **(record.data or {}),
            "archive_rejection_rollback_by": identity["username"],
            "archive_rejection_rollback_at": changed_at,
            "archive_rejection_rollback_comment": body.comment.strip(),
        }
        db.add(WorkflowEvent(
            record_id=record.id, action="回滚归档费拒绝",
            from_status="已拒绝", to_status="已支付",
            operator=identity["username"], comment=body.comment.strip(),
        ))
    await db.commit()
    return {"rolled_back": len(records), "status": "已支付"}


@app.post(f"{settings.api_prefix}/finance/archive-settlements/rejected/reapply")
async def reapply_rejected_archive_settlements(
    body: ArchiveSettlementRejectedActionInput,
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    if identity.get("role") not in {"admin", "manager", "auditor"}:
        raise HTTPException(status_code=403, detail="当前角色没有归档费重新申请权限")
    records = await _rejected_archive_settlement_records(body.record_ids, identity, db)
    changed_at = datetime.now().isoformat(timespec="seconds")
    for record in records:
        record.status = "已回滚"
        record.data = {
            **(record.data or {}),
            "archive_payment_reapplied_by": identity["username"],
            "archive_payment_reapplied_at": changed_at,
            "archive_payment_reapply_comment": body.comment.strip(),
        }
        db.add(WorkflowEvent(
            record_id=record.id, action="重新申请归档费",
            from_status="已拒绝", to_status="已回滚",
            operator=identity["username"], comment=body.comment.strip(),
        ))
    await db.commit()
    return {"reapplied": len(records), "status": "待支付"}


@app.get(f"{settings.api_prefix}/finance/archive-settlements/rejected/export")
async def export_rejected_archive_settlements(
    ids: str = "", identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    try:
        selected_ids = {int(item.strip()) for item in ids.split(",") if item.strip()}
    except ValueError:
        raise HTTPException(status_code=422, detail="归档费记录编号无效")
    if not selected_ids:
        raise HTTPException(status_code=422, detail="请选择案件.")
    rows = await _archive_settlement_decision_rows(identity, db, statuses={"已拒绝"}, selected_ids=selected_ids)
    if len(rows) != len(selected_ids):
        raise HTTPException(status_code=409, detail="部分归档费不存在、已重新申请或无权导出")
    headers = ["案号", "客户", "案件阶段", "律师助理", "开庭律师", "客户管理人", "费用类型", "回款方式", "回款时间", "回款金额", "归档费金额", "结算时间", "支付状态"]
    numeric = {9, 10}
    values = [[
        row["data"].get("case_no"), row.get("customer"), row["data"].get("case_stage"),
        row["data"].get("assistant"), row["data"].get("hearing_lawyer"), row["data"].get("customer_manager"),
        row["data"].get("fee_type"), row["data"].get("payment_method"), row["data"].get("received_date"),
        row["data"].get("receipt_amount"), row["data"].get("archive_fee_amount"), row["data"].get("settlement_paid_at"),
        row.get("status"),
    ] for row in rows]
    def cell(value: object, *, number: bool = False) -> str:
        value_text = f"{float(value or 0):.2f}" if number else str(value or "")
        return f'<Cell><Data ss:Type="{"Number" if number else "String"}">{xml_escape(value_text)}</Data></Cell>'
    sheet_rows = ["<Row>" + "".join(cell(value) for value in headers) + "</Row>"]
    sheet_rows.extend("<Row>" + "".join(cell(value, number=index in numeric and value is not None) for index, value in enumerate(row)) + "</Row>" for row in values)
    workbook = '<?xml version="1.0" encoding="UTF-8"?><?mso-application progid="Excel.Sheet"?><Workbook xmlns="urn:schemas-microsoft-com:office:spreadsheet" xmlns:ss="urn:schemas-microsoft-com:office:spreadsheet"><Worksheet ss:Name="已拒绝"><Table>' + "".join(sheet_rows) + "</Table></Worksheet></Workbook>"
    filename = f"已拒绝归档费-{date.today()}.xls"
    disposition = f"attachment; filename=archive-settlement-rejected.xls; filename*=UTF-8''{quote(filename)}"
    return Response(content=workbook.encode("utf-8"), media_type="application/vnd.ms-excel", headers={"Content-Disposition": disposition})


@app.get(f"{settings.api_prefix}/finance/general-settlements/applications")
async def list_general_settlement_applications(
    customer: str = "", case_no: str = "", customer_manager: str = "",
    received_from: date | None = None, received_to: date | None = None,
    payer: str = "", payment_method: str = "", applied_by: str = "",
    applied_from: date | None = None, applied_to: date | None = None,
    hearing_lawyer: str = "", assistant: str = "", reviewer: str = "",
    reviewed_from: date | None = None, reviewed_to: date | None = None,
    paid_from: date | None = None, paid_to: date | None = None,
    source_person: str = "", application_status: str = Query("待审批", alias="status"),
    page: int = Query(1, ge=1), page_size: int = Query(10, ge=1, le=200),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    for start_date, end_date, label in (
        (received_from, received_to, "回款"),
        (applied_from, applied_to, "提交"),
        (reviewed_from, reviewed_to, "审核"),
        (paid_from, paid_to, "付款"),
    ):
        if start_date and end_date and start_date > end_date:
            raise HTTPException(status_code=422, detail=f"{label}开始日期不能晚于结束日期")
    allowed_statuses = {"待审批", "待付款", "部分付款", "已付款", "已拒绝", "已驳回", "已退回"}
    application_statuses = {item.strip() for item in application_status.split(",") if item.strip()}
    if not application_statuses or not application_statuses.issubset(allowed_statuses):
        raise HTTPException(status_code=422, detail="结算申请状态无效")
    records = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "finance_settlement",
        BusinessRecord.status.in_(application_statuses),
        *_settlement_application_scope(identity),
    ).order_by(BusinessRecord.created_at.desc(), BusinessRecord.id.desc()))).all())

    def contains(value: object, needle: str) -> bool:
        return not needle.strip() or needle.strip().casefold() in str(value or "").casefold()

    def date_in_range(value: object, start_date: date | None, end_date: date | None) -> bool:
        if not start_date and not end_date:
            return True
        try:
            current = date.fromisoformat(str(value or "")[:10])
        except ValueError:
            return False
        return (not start_date or current >= start_date) and (not end_date or current <= end_date)

    filtered: list[BusinessRecord] = []
    for record in records:
        data = record.data or {}
        if not contains(record.customer, customer) or not contains(data.get("case_nos"), case_no):
            continue
        if not contains(data.get("customer_manager"), customer_manager):
            continue
        if not date_in_range(data.get("received_date"), received_from, received_to):
            continue
        if not contains(data.get("payer_name"), payer) or not contains(data.get("payment_method"), payment_method):
            continue
        if not contains(data.get("applied_by") or record.owner, applied_by):
            continue
        if not date_in_range(data.get("applied_at") or record.created_at, applied_from, applied_to):
            continue
        if not contains(data.get("hearing_lawyer"), hearing_lawyer) or not contains(data.get("assistant"), assistant):
            continue
        if not contains(data.get("reviewer"), reviewer):
            continue
        if not date_in_range(data.get("reviewed_at"), reviewed_from, reviewed_to):
            continue
        if not date_in_range(data.get("paid_at"), paid_from, paid_to):
            continue
        if not contains(data.get("source_person"), source_person):
            continue
        filtered.append(record)

    items = [await _record_dict_for_identity(record, identity, db) for record in filtered]
    amount_keys = ["receipt_amount", "allocated_amount", "remaining_amount", "assigned_official_fee", "assigned_agency_fee", "assigned_other_fee", "agency_settlement_amount", "archive_fee", "actual_settlement_amount"]
    totals = {key: _round_fee_amount(sum(float((item.get("data") or {}).get(key) or 0) for item in items)) for key in amount_keys}
    start = (page - 1) * page_size
    return {"items": items[start:start + page_size], "total": len(items), "totals": totals, "page": page, "page_size": page_size}


@app.post(f"{settings.api_prefix}/finance/general-settlements/applications/reapply")
async def reapply_general_settlement_applications(body: FinanceSettlementReapplyInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    comment = body.comment.strip()
    if not comment:
        raise HTTPException(status_code=422, detail="请输入备注.")
    application_ids = list(dict.fromkeys(body.application_ids))
    records = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.id.in_(application_ids),
        BusinessRecord.module == "finance_settlement",
        *_settlement_application_scope(identity),
    ))).all())
    if len(records) != len(application_ids):
        raise HTTPException(status_code=404, detail="部分结算申请不存在或无权访问")
    allowed_from = {"已拒绝", "已驳回", "已退回"}
    invalid = [record.serial_no for record in records if record.status not in allowed_from]
    if invalid:
        raise HTTPException(status_code=409, detail="仅已拒绝或已退回结算可以重新申请：" + "、".join(invalid))
    reapplied_at = datetime.now().isoformat(timespec="seconds")
    for record in records:
        previous_status = record.status
        record.status = "待审批"
        record.description = comment
        record.data = {
            **(record.data or {}),
            "applied_by": identity["username"],
            "applied_at": reapplied_at,
            "reapplied_by": identity["username"],
            "reapplied_at": reapplied_at,
            "reapply_comment": comment,
        }
        db.add(WorkflowEvent(
            record_id=record.id,
            action="重新申请结算",
            from_status=previous_status,
            to_status="待审批",
            operator=identity["username"],
            comment=comment,
        ))
    await db.commit()
    return {"reapplied": len(records), "application_ids": application_ids, "status": "待审批"}


@app.post(f"{settings.api_prefix}/finance/general-settlements/applications/review")
async def review_general_settlement_applications(body: FinanceSettlementReviewInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager", "auditor"}:
        raise HTTPException(status_code=403, detail="当前角色没有结算审批权限")
    application_ids = list(dict.fromkeys(body.application_ids))
    records = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.id.in_(application_ids),
        BusinessRecord.module == "finance_settlement",
        *_settlement_application_scope(identity),
    ))).all())
    if len(records) != len(application_ids):
        raise HTTPException(status_code=404, detail="部分结算申请不存在或无权访问")
    invalid = [record.serial_no for record in records if record.status != "待审批"]
    if invalid:
        raise HTTPException(status_code=409, detail="仅待审批结算申请可以审核：" + "、".join(invalid))
    target_status = "待付款" if body.approved else "已拒绝"
    action = "同意结算" if body.approved else "拒绝结算"
    reviewed_at = datetime.now().isoformat(timespec="seconds")
    comment = body.comment.strip()
    for record in records:
        record.status = target_status
        record.data = {
            **(record.data or {}),
            "reviewer": identity["username"],
            "reviewed_at": reviewed_at,
            "review_comment": comment,
        }
        db.add(WorkflowEvent(
            record_id=record.id, action=action, from_status="待审批",
            to_status=target_status, operator=identity["username"], comment=comment,
        ))
    await db.commit()
    return {"reviewed": len(records), "application_ids": application_ids, "status": target_status}


@app.post(f"{settings.api_prefix}/finance/general-settlements/applications/payment")
async def pay_or_rollback_general_settlement_applications(body: FinanceSettlementPaymentInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager", "auditor"}:
        raise HTTPException(status_code=403, detail="当前角色没有结算付款权限")
    application_ids = list(dict.fromkeys(body.application_ids))
    records = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.id.in_(application_ids),
        BusinessRecord.module == "finance_settlement",
        *_settlement_application_scope(identity),
    ))).all())
    if len(records) != len(application_ids):
        raise HTTPException(status_code=404, detail="部分结算申请不存在或无权访问")
    invalid = [record.serial_no for record in records if record.status not in {"待付款", "已付款"}]
    if invalid:
        raise HTTPException(status_code=409, detail="仅待付款或已付款结算申请可以处理付款：" + "、".join(invalid))
    paid_again = [record.serial_no for record in records if record.status == "已付款" and body.action == "paid"]
    if paid_again:
        raise HTTPException(status_code=409, detail="已付款结算申请只能回退：" + "、".join(paid_again))
    comment = body.comment.strip()
    if body.action == "rollback" and not comment:
        raise HTTPException(status_code=422, detail="请输入审核备注.")
    if body.action == "rollback":
        archive_decisions = list((await db.scalars(select(BusinessRecord).where(
            BusinessRecord.module == "finance_archive_settlement",
            BusinessRecord.status.in_({"已支付", "已拒绝"}),
            *_settlement_application_scope(identity),
        ))).all())
        blocked_application_ids = {
            int((decision.data or {}).get("source_application_id") or 0)
            for decision in archive_decisions
        } & set(application_ids)
        if blocked_application_ids:
            raise HTTPException(
                status_code=409,
                detail="请先回滚或重新申请关联归档费，再回退结算",
            )
    processed_at = datetime.now().isoformat(timespec="seconds")
    if body.action == "paid":
        target_status = "已付款"
        action = "标记已支付"
        data_updates = {
            "paid_by": identity["username"],
            "paid_at": processed_at,
            "paid_comment": comment,
        }
    else:
        target_status = "已退回"
        action = "回退结算"
        data_updates = {
            "rollback_by": identity["username"],
            "rollback_at": processed_at,
            "rollback_comment": comment,
            "rejection_comment": comment,
        }
    for record in records:
        previous_status = record.status
        record.status = target_status
        record.data = {**(record.data or {}), **data_updates}
        db.add(WorkflowEvent(
            record_id=record.id,
            action=action,
            from_status=previous_status,
            to_status=target_status,
            operator=identity["username"],
            comment=comment,
        ))
    await db.commit()
    return {
        "processed": len(records),
        "application_ids": application_ids,
        "status": target_status,
        "action": action,
    }


@app.get(f"{settings.api_prefix}/finance/general-settlements/export")
async def export_general_settlements(
    kind: str = Query("settlement", pattern="^(settlement|receipt|case)$"), ids: str = "", application_ids: str = "",
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    if application_ids.strip():
        selected_application_ids = set(_export_ids(application_ids))
        if not selected_application_ids:
            raise HTTPException(status_code=422, detail="请选择需要导出的结算申请.")
        records = list((await db.scalars(select(BusinessRecord).where(
            BusinessRecord.id.in_(selected_application_ids),
            BusinessRecord.module == "finance_settlement",
            *_settlement_application_scope(identity),
        ))).all())
        if len(records) != len(selected_application_ids):
            raise HTTPException(status_code=409, detail="部分结算申请不存在或无权导出")
        rows = [{
            "id": record.id,
            "serial_no": str((record.data or {}).get("receipt_no") or record.serial_no),
            "customer": record.customer,
            "data": record.data or {},
        } for record in records]
    else:
        selected_ids = set(_export_ids(ids)) if ids.strip() else None
        if selected_ids is not None and not selected_ids:
            raise HTTPException(status_code=422, detail="请选择需要导出的回款.")
        rows = await _general_settlement_rows(identity, db, receipt_ids=selected_ids)
        if selected_ids is not None and len({int(row["id"]) for row in rows}) != len(selected_ids):
            raise HTTPException(status_code=409, detail="部分回款已申请结算、尚未分配或无权导出")
    if not rows:
        raise HTTPException(status_code=422, detail="没有可导出的待结算记录")
    def cell(value: object, *, number: bool = False) -> str:
        value_text = f"{float(value or 0):.2f}" if number else str(value or "")
        return f'<Cell><Data ss:Type="{"Number" if number else "String"}">{xml_escape(value_text)}</Data></Cell>'
    if kind == "receipt":
        headers = ["回款编号", "客户名称", "回款单位", "回款日期", "回款金额", "已分金额", "未分金额", "回款方式", "银行备注"]
        values = [[row["serial_no"], row["customer"], row["data"].get("payer_name"), row["data"].get("received_date"), row["data"].get("receipt_amount"), row["data"].get("allocated_amount"), row["data"].get("remaining_amount"), row["data"].get("payment_method"), row["data"].get("bank_remark")] for row in rows]
        numeric = {4, 5, 6}
        sheet_name = "到账清单"
    elif kind == "case":
        headers = ["回款编号", "案号", "阶段", "费用类型", "本笔分配金额", "本笔结算金额", "本笔归档费", "客户", "经办律师", "律师助理", "合同号"]
        values = [[row["serial_no"], detail.get("case_no"), detail.get("case_stage"), detail.get("fee_type"), detail.get("current_amount"), detail.get("settlement_amount"), detail.get("archive_fee"), detail.get("customer"), detail.get("handling_lawyer"), detail.get("assistant"), detail.get("contract_no")] for row in rows for detail in row["data"].get("allocation_details", [])]
        numeric = {4, 5, 6}
        sheet_name = "案件清单"
    else:
        headers = ["回款编号", "客户名称", "客户管理人", "回款单位", "回款日期", "回款金额", "已分金额", "未分金额", "已分官费", "已分代理费", "已分其他费用", "代理费结算金额", "扣归档费", "实际结算金额"]
        values = [[row["serial_no"], row["customer"], row["data"].get("customer_manager"), row["data"].get("payer_name"), row["data"].get("received_date"), row["data"].get("receipt_amount"), row["data"].get("allocated_amount"), row["data"].get("remaining_amount"), row["data"].get("assigned_official_fee"), row["data"].get("assigned_agency_fee"), row["data"].get("assigned_other_fee"), row["data"].get("agency_settlement_amount"), row["data"].get("archive_fee"), row["data"].get("actual_settlement_amount")] for row in rows]
        numeric = set(range(5, 14))
        sheet_name = "结算清单"
    sheet_rows = ["<Row>" + "".join(cell(value) for value in headers) + "</Row>"]
    sheet_rows.extend("<Row>" + "".join(cell(value, number=index in numeric) for index, value in enumerate(row)) + "</Row>" for row in values)
    workbook = '<?xml version="1.0" encoding="UTF-8"?><?mso-application progid="Excel.Sheet"?><Workbook xmlns="urn:schemas-microsoft-com:office:spreadsheet" xmlns:ss="urn:schemas-microsoft-com:office:spreadsheet"><Worksheet ss:Name="' + sheet_name + '"><Table>' + "".join(sheet_rows) + "</Table></Worksheet></Workbook>"
    filename = f"{sheet_name}-{date.today()}.xls"
    disposition = f"attachment; filename=settlement-export.xls; filename*=UTF-8''{quote(filename)}"
    return Response(content=workbook.encode("utf-8"), media_type="application/vnd.ms-excel", headers={"Content-Disposition": disposition})


@app.delete(f"{settings.api_prefix}/finance/general-settlements/applications/{{application_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_general_settlement_application(application_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可以删除结算申请")
    item = await db.get(BusinessRecord, application_id)
    if not item or item.module != "finance_settlement":
        raise HTTPException(status_code=404, detail="结算申请不存在")
    archive_decisions = (await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "finance_archive_settlement",
    ))).all()
    for decision in archive_decisions:
        if int((decision.data or {}).get("source_application_id") or 0) == item.id:
            await db.execute(delete(WorkflowEvent).where(WorkflowEvent.record_id == decision.id))
            await db.delete(decision)
    await db.execute(delete(WorkflowEvent).where(WorkflowEvent.record_id == item.id))
    await db.delete(item)
    await db.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.get(f"{settings.api_prefix}/finance/settlements/pending")
async def list_pending_finance_settlements(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Return paid internal fees enriched with their case and bank-receipt context."""
    fees = (await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "finance",
        BusinessRecord.status == "已付款",
        *(await _record_scope_conditions(identity, db)),
    ).order_by(BusinessRecord.updated_at.desc(), BusinessRecord.id.desc()))).all()
    fees = [item for item in fees if (item.data or {}).get("fee_type") == "内部费用" and not (item.data or {}).get("commission_paid")]
    case_ids = {int((item.data or {}).get("case_id") or 0) for item in fees if (item.data or {}).get("case_id")}
    case_nos = {str((item.data or {}).get("case_no") or "") for item in fees if (item.data or {}).get("case_no")}
    cases = (await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "case",
        or_(BusinessRecord.id.in_(case_ids), BusinessRecord.serial_no.in_(case_nos)),
        *(await _record_scope_conditions(identity, db)),
    ))).all() if case_ids or case_nos else []
    cases_by_id = {item.id: item for item in cases}
    cases_by_no = {item.serial_no: item for item in cases}
    payments = (await db.scalars(select(IncomingPayment).order_by(IncomingPayment.received_date.desc(), IncomingPayment.id.desc()))).all()
    if identity.get("role") not in {"admin", "auditor"}:
        visible_customers = {item.customer for item in cases}
        payments = [item for item in payments if item.operator == identity["username"] or item.claimant == identity["username"] or item.claimed_customer in visible_customers]
    can_view_amount = "finance.amount" in await _allowed_field_keys(identity, db)
    rows = []
    for fee in fees:
        fee_data = fee.data or {}
        case = cases_by_id.get(int(fee_data.get("case_id") or 0)) or cases_by_no.get(str(fee_data.get("case_no") or ""))
        case_data = (case.data or {}) if case else {}
        receipt_matches: list[tuple[IncomingPayment, dict]] = []
        for payment in payments:
            for allocation in payment.allocations or []:
                if (case and int(allocation.get("case_id") or 0) == case.id) or (fee_data.get("case_no") and allocation.get("case_no") == fee_data.get("case_no")):
                    receipt_matches.append((payment, allocation))
        receipt_amount = sum(float(allocation.get("amount") or 0) for _, allocation in receipt_matches)
        latest_payment = receipt_matches[0][0] if receipt_matches else None
        record = await _record_dict_for_identity(fee, identity, db)
        record["data"] = {
            **record.get("data", {}),
            "case_id": case.id if case else fee_data.get("case_id"),
            "case_no": case.serial_no if case else fee_data.get("case_no", ""),
            "plaintiff": case_data.get("plaintiff", ""),
            "defendant": case_data.get("defendant") or case_data.get("opponent", ""),
            "court_case_no": case_data.get("court_case_no", ""),
            "certificate_no": case_data.get("certificate_no", ""),
            "case_stage": case_data.get("case_stage") or case.status if case else "",
            "case_source": case_data.get("case_source") or case_data.get("source_person", ""),
            "hearing_lawyer": case_data.get("hearing_lawyer", ""),
            "assistant": case_data.get("assistant") or case_data.get("lawyer_assistant", ""),
            "investigator": case_data.get("investigator", ""),
            "quality_manager": case_data.get("quality_manager") or case_data.get("quality_control", ""),
            "receipt_amount": receipt_amount if can_view_amount else None,
            "receipt_date": str(latest_payment.received_date) if latest_payment else "",
            "payee": latest_payment.payer_name if latest_payment else fee_data.get("payee", ""),
            "settlement_status": fee.status,
        }
        rows.append(record)
    return {"items": rows, "total": len(rows)}


@app.post(f"{settings.api_prefix}/finance/settlements/mark-commission-paid")
async def mark_finance_settlements_commission_paid(body: FinanceSettlementMarkInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager", "auditor"}:
        raise HTTPException(status_code=403, detail="当前角色没有标识提成发放权限")
    fee_ids = list(dict.fromkeys(body.fee_ids))
    fees = (await db.scalars(select(BusinessRecord).where(
        BusinessRecord.id.in_(fee_ids),
        BusinessRecord.module == "finance",
        *(await _record_scope_conditions(identity, db)),
    ))).all()
    if len(fees) != len(fee_ids):
        raise HTTPException(status_code=404, detail="部分案件费用不存在或无权访问")
    invalid = [item.serial_no for item in fees if item.status != "已付款" or (item.data or {}).get("fee_type") != "内部费用" or (item.data or {}).get("commission_paid")]
    if invalid:
        raise HTTPException(status_code=409, detail="仅可标识尚未发放提成的已付款内部费用：" + "、".join(invalid))
    marked_at = datetime.now().isoformat(timespec="seconds")
    for item in fees:
        item.data = {
            **(item.data or {}),
            "commission_paid": True,
            "commission_paid_by": identity["username"],
            "commission_paid_at": marked_at,
            "commission_paid_comment": body.comment.strip(),
        }
        db.add(WorkflowEvent(record_id=item.id, action="标识提成已发", from_status=item.status, to_status=item.status, operator=identity["username"], comment=body.comment.strip()))
    await db.commit()
    return {"marked": len(fees), "fee_ids": fee_ids, "marked_at": marked_at}


@app.get(f"{settings.api_prefix}/finance/fees/refund-review-candidates")
async def list_internal_refund_review_candidates(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """List internal negative-amount requests for the dedicated refund review page."""
    items = (await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "finance",
        *(await _record_scope_conditions(identity, db)),
    ).order_by(BusinessRecord.updated_at.desc(), BusinessRecord.id.desc()))).all()
    refund_items = [
        item for item in items
        if (item.data or {}).get("fee_type") == "内部费用"
        and (
            bool((item.data or {}).get("is_refund"))
            or float((item.data or {}).get("amount") or 0) < 0
        )
    ]
    return {
        "items": [await _record_dict_for_identity(item, identity, db) for item in refund_items],
        "total": len(refund_items),
    }


async def _prepare_internal_payment_package(
    fee_ids: list[int], identity: dict, db: AsyncSession
) -> tuple[list[BusinessRecord], list[dict], str, float]:
    if identity.get("role") not in {"admin", "manager", "auditor"}:
        raise HTTPException(status_code=403, detail="当前角色没有打包付款权限")
    unique_ids = list(dict.fromkeys(fee_ids))
    rows = (await db.scalars(select(BusinessRecord).where(
        BusinessRecord.id.in_(unique_ids),
        BusinessRecord.module == "finance",
        *(await _record_scope_conditions(identity, db)),
    ))).all()
    by_id = {item.id: item for item in rows}
    if len(by_id) != len(unique_ids):
        raise HTTPException(status_code=404, detail="部分提成不存在或无权访问")
    fees = [by_id[item_id] for item_id in unique_ids]
    invalid_type = [item.serial_no for item in fees if (item.data or {}).get("fee_type") != "内部费用"]
    if invalid_type:
        raise HTTPException(status_code=409, detail="仅内部费用提成可以打包付款：" + "、".join(invalid_type))
    invalid_status = [item.serial_no for item in fees if item.status != "已审批"]
    if invalid_status:
        raise HTTPException(status_code=409, detail="仅待付款提成可以打包付款：" + "、".join(invalid_status))
    payees = {
        str((item.data or {}).get("payee") or (item.data or {}).get("applicant") or item.owner or "").strip()
        for item in fees
    }
    if "" in payees:
        raise HTTPException(status_code=422, detail="提成收款人不能为空")
    if len(payees) != 1:
        raise HTTPException(status_code=409, detail="请选择同一收款人的提成进行打包付款")
    payee = next(iter(payees))
    case_ids = {int((item.data or {}).get("case_id") or 0) for item in fees if (item.data or {}).get("case_id")}
    cases = {
        item.id: item
        for item in (await db.scalars(select(BusinessRecord).where(
            BusinessRecord.id.in_(case_ids),
            BusinessRecord.module == "case",
            *(await _record_scope_conditions(identity, db)),
        ))).all()
    } if case_ids else {}
    details: list[dict] = []
    for item in fees:
        data = item.data or {}
        amount = _round_fee_amount(float(data.get("actual_commission") if data.get("actual_commission") is not None else data.get("amount") or 0))
        linked_case = cases.get(int(data.get("case_id") or 0))
        details.append({
            "fee_id": item.id,
            "request_no": item.serial_no,
            "case_no": data.get("case_no", ""),
            "case_name": linked_case.title if linked_case else data.get("case_name") or item.title,
            "amount": amount,
            "commission_type": data.get("commission_type") or item.title or data.get("fee_type", ""),
            "payee": payee,
            "remark": data.get("remark") or item.description or "",
        })
    total_amount = _round_fee_amount(sum(float(item["amount"]) for item in details))
    return fees, details, payee, total_amount


def _new_internal_payment_package_no() -> str:
    numeric_suffix = int(uuid4().hex[:12], 16) % 100_000_000
    return f"P{datetime.now():%y%m%d}-{numeric_suffix:08d}"


@app.get(f"{settings.api_prefix}/finance/payment-packages")
async def list_internal_payment_packages(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    items = (await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "finance_package",
        *(await _record_scope_conditions(identity, db)),
    ).order_by(BusinessRecord.created_at.desc(), BusinessRecord.id.desc()))).all()
    return {"items": [await _record_dict_for_identity(item, identity, db) for item in items], "total": len(items)}


@app.post(f"{settings.api_prefix}/finance/payment-packages/preview")
async def preview_internal_payment_package(body: FinancePaymentPackagePreviewInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _fees, details, payee, total_amount = await _prepare_internal_payment_package(body.fee_ids, identity, db)
    return {
        "package_no": _new_internal_payment_package_no(),
        "print_date": str(date.today()),
        "payee": payee,
        "total_amount": total_amount,
        "items": details,
    }


@app.post(f"{settings.api_prefix}/finance/payment-packages", status_code=status.HTTP_201_CREATED)
async def create_internal_payment_package(body: FinancePaymentPackageCreateInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    fees, details, payee, total_amount = await _prepare_internal_payment_package(body.fee_ids, identity, db)
    if await db.scalar(select(BusinessRecord.id).where(BusinessRecord.serial_no == body.package_no)):
        raise HTTPException(status_code=409, detail="付款包号码已经存在")
    user = await db.scalar(select(User).where(User.username == identity["username"]))
    paid_at = datetime.now().isoformat(timespec="seconds")
    payment_date = str(date.today())
    package = BusinessRecord(
        module="finance_package",
        serial_no=body.package_no,
        title=f"{payee}提成付款申请单",
        customer="",
        status="待核销",
        owner=identity["username"],
        department=user.department if user else "上海分所",
        description=body.comment.strip(),
        data={
            "fee_ids": [item.id for item in fees],
            "payee": payee,
            "amount": total_amount,
            "total_amount": total_amount,
            "payment_date": payment_date,
            "payment_status": "待核销",
            "fee_type": "内部提成",
            "items": details,
            "submitted_at": paid_at,
            "submitted_by": identity["username"],
            "comment": body.comment.strip(),
        },
    )
    db.add(package)
    await db.flush()
    db.add(WorkflowEvent(record_id=package.id, action="创建付款包", from_status="", to_status="待核销", operator=identity["username"], comment=body.comment.strip() or "同一收款人提成打包付款"))
    for fee in fees:
        previous = fee.status
        fee.status = "已付款"
        fee.data = {
            **(fee.data or {}),
            "payment_status": "已付款",
            "payment_date": payment_date,
            "payment_package_id": package.id,
            "payment_package_no": package.serial_no,
            "paid_at": paid_at,
            "paid_by": identity["username"],
        }
        db.add(WorkflowEvent(record_id=fee.id, action="打包付款", from_status=previous, to_status="已付款", operator=identity["username"], comment=f"付款包 {package.serial_no}；{body.comment.strip()}"))
    await db.commit()
    await db.refresh(package)
    return await _record_dict_for_identity(package, identity, db)


@app.post(f"{settings.api_prefix}/finance/payment-packages/{{package_id}}/writeoff")
async def writeoff_internal_payment_package(package_id: int, body: FinancePaymentPackageWriteoffInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager", "auditor"}:
        raise HTTPException(status_code=403, detail="当前角色没有付款核销权限")
    package = await _ensure_record_module(package_id, "finance_package", identity, db)
    if package.status != "待核销":
        raise HTTPException(status_code=409, detail="仅待核销付款包可以核销")
    if body.payment_method not in {"自动扣款", "银行卡", "现金"}:
        raise HTTPException(status_code=422, detail="付款方式无效")
    if not body.invoice_no.strip():
        raise HTTPException(status_code=422, detail="请输入付款单据号.")
    package_data = package.data or {}
    expected_amount = _round_fee_amount(float(package_data.get("total_amount") or package_data.get("amount") or 0))
    confirmed_amount = _round_fee_amount(body.amount)
    if abs(confirmed_amount - expected_amount) > 0.001:
        raise HTTPException(status_code=409, detail=f"确认付款金额必须等于付款包金额 {expected_amount:.2f}")
    written_off_at = datetime.now().isoformat(timespec="seconds")
    package.status = "已付款"
    package.data = {
        **package_data,
        "payment_status": "已付款",
        "paid_date": str(body.paid_date),
        "payment_method": body.payment_method,
        "invoice_no": body.invoice_no.strip(),
        "remark": body.remark.strip(),
        "writeoff_status": "已核销",
        "written_off_at": written_off_at,
        "written_off_by": identity["username"],
    }
    fee_ids = [int(item_id) for item_id in package_data.get("fee_ids", [])]
    fees = (await db.scalars(select(BusinessRecord).where(
        BusinessRecord.id.in_(fee_ids), BusinessRecord.module == "finance"
    ))).all() if fee_ids else []
    if len(fees) != len(set(fee_ids)):
        raise HTTPException(status_code=409, detail="付款包关联的费用记录不完整")
    for fee in fees:
        data = fee.data or {}
        if int(data.get("payment_package_id") or 0) != package.id:
            raise HTTPException(status_code=409, detail=f"费用 {fee.serial_no} 的付款包关联不一致")
        fee.data = {
            **data,
            "writeoff_status": "已核销",
            "writeoff_voucher_no": body.invoice_no.strip(),
            "payment_method": body.payment_method,
            "written_off_at": written_off_at,
            "written_off_by": identity["username"],
        }
        db.add(WorkflowEvent(record_id=fee.id, action="付款包核销", from_status=fee.status, to_status=fee.status, operator=identity["username"], comment=f"付款包 {package.serial_no}；单据号 {body.invoice_no.strip()}；{body.remark.strip()}"))
    db.add(WorkflowEvent(record_id=package.id, action="付款核销", from_status="待核销", to_status="已付款", operator=identity["username"], comment=f"{body.payment_method}；单据号 {body.invoice_no.strip()}；{body.remark.strip()}"))
    await db.commit()
    await db.refresh(package)
    return await _record_dict_for_identity(package, identity, db)


@app.delete(f"{settings.api_prefix}/finance/payment-packages/{{package_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def cancel_internal_payment_package(package_id: int, reverse_paid: bool = Query(False), identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可以撤销付款包")
    package = await _ensure_record_module(package_id, "finance_package", identity, db)
    if package.status == "已付款" and not reverse_paid:
        raise HTTPException(status_code=409, detail="已核销付款包必须显式冲正后才能撤销")
    if package.status not in {"待核销", "已付款"}:
        raise HTTPException(status_code=409, detail="当前付款包状态不能撤销")
    package_data = package.data or {}
    fee_ids = [int(item_id) for item_id in package_data.get("fee_ids", [])]
    fees = (await db.scalars(select(BusinessRecord).where(
        BusinessRecord.id.in_(fee_ids), BusinessRecord.module == "finance"
    ))).all() if fee_ids else []
    payment_keys = {
        "payment_status", "payment_date", "payment_package_id",
        "payment_package_no", "paid_at", "paid_by", "writeoff_status",
        "writeoff_voucher_no", "payment_method", "written_off_at",
        "written_off_by",
    }
    for fee in fees:
        data = fee.data or {}
        if int(data.get("payment_package_id") or 0) != package.id:
            continue
        previous = fee.status
        fee.status = "已审批"
        fee.data = {key: value for key, value in data.items() if key not in payment_keys}
        action = "冲正已核销付款包" if package.status == "已付款" else "撤销打包付款"
        db.add(WorkflowEvent(record_id=fee.id, action=action, from_status=previous, to_status="已审批", operator=identity["username"], comment=f"撤销付款包 {package.serial_no}"))
    await db.execute(delete(WorkflowEvent).where(WorkflowEvent.record_id == package.id))
    await db.delete(package)
    await db.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.get(f"{settings.api_prefix}/finance/fees/{{fee_id}}/readiness")
async def finance_fee_readiness(fee_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _ensure_record_module(fee_id, "finance", identity, db)
    return await _finance_fee_readiness(item, identity, db)


@app.post(f"{settings.api_prefix}/finance/fees/{{fee_id}}/submit")
async def submit_finance_fee(fee_id: int, body: FinanceActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _ensure_record_module(fee_id, "finance", identity, db)
    await _require_record_owner_or_manager(item, identity, db)
    if item.status not in {"草稿", "已退回"}: raise HTTPException(status_code=409, detail="当前状态不能提交审批")
    data = item.data or {}; missing = []
    if not data.get("handler"): missing.append("经办人员")
    if not data.get("case_no"): missing.append("关联案号")
    if data.get("fee_type") == "官方费用":
        if not data.get("court"): missing.append("缴费法院/机构")
        if not data.get("document_no"): missing.append("缴费通知文号")
    if missing: raise HTTPException(status_code=422, detail="缺少费用审批要素：" + "、".join(missing))
    if data.get("fee_type") == "官方费用":
        readiness = await _finance_fee_readiness(item, identity, db)
        if not readiness["ready"]: raise HTTPException(status_code=422, detail="案件付款三要素不完整：" + "；".join(readiness["missing"]))
    previous = item.status; item.status = "待审批"
    db.add(WorkflowEvent(record_id=item.id, action="提交费用审批", from_status=previous, to_status="待审批", operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(item); return await _record_dict_for_identity(item, identity, db)


@app.post(f"{settings.api_prefix}/finance/fees/{{fee_id}}/approve")
async def approve_finance_fee(fee_id: int, body: FinanceActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _ensure_record_module(fee_id, "finance", identity, db)
    if identity.get("role") not in {"admin", "manager", "auditor"}: raise HTTPException(status_code=403, detail="当前角色没有费用审批权限")
    if item.status != "待审批": raise HTTPException(status_code=409, detail="仅待审批费用可以通过")
    item.status = "已审批"
    db.add(WorkflowEvent(record_id=item.id, action="费用审批通过", from_status="待审批", to_status="已审批", operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(item); return await _record_dict_for_identity(item, identity, db)


async def _review_finance_fee_records(items: list[BusinessRecord], approved: bool, comment: str, identity: dict, db: AsyncSession) -> None:
    if identity.get("role") not in {"admin", "manager", "auditor"}:
        raise HTTPException(status_code=403, detail="当前角色没有费用审批权限")
    invalid = [item.serial_no for item in items if item.status != "待审批"]
    if invalid:
        raise HTTPException(status_code=409, detail="仅待审批费用可以审核：" + "、".join(invalid))
    target_status = "已审批" if approved else "已驳回"
    normalized_comment = comment.strip() or ("审批通过" if approved else "审批拒绝")
    for item in items:
        data = item.data or {}
        is_refund = data.get("fee_type") == "内部费用" and (
            bool(data.get("is_refund")) or float(data.get("amount") or 0) < 0
        )
        action = (
            "内部提成退费审批通过" if approved else "内部提成退费审批驳回"
        ) if is_refund else ("费用审批通过" if approved else "费用审批驳回")
        item.status = target_status
        if is_refund and not data.get("is_refund"):
            item.data = {**data, "is_refund": True}
        db.add(WorkflowEvent(record_id=item.id, action=action, from_status="待审批", to_status=target_status, operator=identity["username"], comment=normalized_comment))


@app.post(f"{settings.api_prefix}/finance/fees/batch-review")
async def batch_review_finance_fees(body: FinanceFeeBatchReviewInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    fee_ids = list(dict.fromkeys(body.fee_ids))
    items = (await db.scalars(select(BusinessRecord).where(
        BusinessRecord.id.in_(fee_ids),
        BusinessRecord.module == "finance",
        *(await _record_scope_conditions(identity, db)),
    ))).all()
    if len(items) != len(fee_ids):
        raise HTTPException(status_code=404, detail="部分费用不存在或无权访问")
    await _review_finance_fee_records(items, body.approved, body.comment, identity, db)
    await db.commit()
    return {"reviewed": len(items), "fee_ids": fee_ids, "status": "已审批" if body.approved else "已驳回"}


@app.post(f"{settings.api_prefix}/finance/fees/{{fee_id}}/review")
async def review_finance_fee(fee_id: int, body: FinanceFeeReviewInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _ensure_record_module(fee_id, "finance", identity, db)
    await _review_finance_fee_records([item], body.approved, body.comment, identity, db)
    await db.commit(); await db.refresh(item)
    return await _record_dict_for_identity(item, identity, db)


@app.post(f"{settings.api_prefix}/finance/fees/{{fee_id}}/void")
async def void_rejected_finance_fee(fee_id: int, body: FinanceActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager", "auditor"}:
        raise HTTPException(status_code=403, detail="当前角色没有请款单作废权限")
    item = await _ensure_record_module(fee_id, "finance", identity, db)
    data = item.data or {}
    if data.get("fee_type") != "内部费用":
        raise HTTPException(status_code=409, detail="该入口仅可作废内部费用请款单")
    if item.status not in {"已拒绝", "已退回", "已驳回"}:
        raise HTTPException(status_code=409, detail="仅已拒绝的内部费用请款单可以作废")
    previous = item.status
    voided_at = datetime.now().isoformat(timespec="seconds")
    item.status = "已作废"
    item.data = {
        **data,
        "payment_status": "已作废",
        "voided_by": identity["username"],
        "voided_at": voided_at,
        "void_comment": body.comment.strip(),
    }
    db.add(WorkflowEvent(
        record_id=item.id,
        action="请款单作废",
        from_status=previous,
        to_status="已作废",
        operator=identity["username"],
        comment=body.comment.strip() or "已拒绝请款单作废",
    ))
    await db.commit(); await db.refresh(item)
    return await _record_dict_for_identity(item, identity, db)


@app.post(f"{settings.api_prefix}/finance/fees/{{fee_id}}/writeoff")
async def writeoff_finance_fee(fee_id: int, body: FinanceWriteoffInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager", "auditor"}:
        raise HTTPException(status_code=403, detail="当前角色没有付款核销权限")
    item = await _ensure_record_module(fee_id, "finance", identity, db)
    if item.status != "已付款":
        raise HTTPException(status_code=409, detail="费用全部付款后才能核销")
    data = item.data or {}
    if data.get("writeoff_status") == "已核销":
        raise HTTPException(status_code=409, detail="付款已经核销")
    payment_total = float(await db.scalar(select(func.coalesce(func.sum(FinanceTransaction.amount), 0)).where(FinanceTransaction.finance_record_id == item.id, FinanceTransaction.transaction_type == "付款")) or 0)
    if payment_total + 0.001 < abs(float(data.get("amount", 0) or 0)):
        raise HTTPException(status_code=409, detail="付款流水合计未达到申请金额，不能核销")
    item.data = {
        **data,
        "payment_status": "已付款",
        "writeoff_status": "已核销",
        "writeoff_voucher_no": body.voucher_no.strip(),
        "writeoff_comment": body.comment.strip(),
        "written_off_by": identity["username"],
        "written_off_at": datetime.now().isoformat(timespec="seconds"),
    }
    db.add(WorkflowEvent(record_id=item.id, action="付款核销", from_status=item.status, to_status=item.status, operator=identity["username"], comment=f"核销凭证：{body.voucher_no.strip()}。{body.comment}"))
    await db.commit(); await db.refresh(item)
    return await _record_dict_for_identity(item, identity, db)


@app.get(f"{settings.api_prefix}/finance/transactions")
async def list_finance_transactions(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    items = (await db.scalars(select(FinanceTransaction).order_by(FinanceTransaction.transaction_date.desc(), FinanceTransaction.id.desc()))).all()
    if identity.get("role") != "admin":
        visible_record_ids = await _visible_record_ids(identity, db)
        items = [item for item in items if (item.finance_record_id and item.finance_record_id in visible_record_ids) or (not item.finance_record_id and item.operator == identity["username"])]
    ids = {item.finance_record_id for item in items if item.finance_record_id}
    records = {item.id: item for item in (await db.scalars(select(BusinessRecord).where(BusinessRecord.id.in_(ids)))).all()} if ids else {}
    transaction_ids = {item.id for item in items}
    voucher_rows = (await db.scalars(select(FileAttachment).where(FileAttachment.finance_transaction_id.in_(transaction_ids)).order_by(FileAttachment.created_at.desc()))).all() if transaction_ids else []
    vouchers: dict[int, list[FileAttachment]] = {}
    for voucher in voucher_rows:
        vouchers.setdefault(int(voucher.finance_transaction_id or 0), []).append(voucher)
    show_amount = "finance.amount" in await _allowed_field_keys(identity, db)
    return {"items": [_finance_transaction_dict(item, records.get(item.finance_record_id), vouchers.get(item.id, []), show_amount=show_amount) for item in items], "total": len(items)}


@app.post(f"{settings.api_prefix}/finance/transactions", status_code=status.HTTP_201_CREATED)
async def create_finance_transaction(body: FinanceTransactionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if body.transaction_type not in FINANCE_TRANSACTION_TYPES: raise HTTPException(status_code=422, detail="流水类型无效")
    if body.transaction_type == "回款": raise HTTPException(status_code=409, detail="银行回款必须先进入回款管理，完成客户认领和应收分配")
    if identity.get("role") not in {"admin", "manager", "auditor"}:
        raise HTTPException(status_code=403, detail="当前角色没有登记财务流水的权限")
    if not body.finance_record_id:
        raise HTTPException(status_code=409, detail="付款、开票和退款流水必须关联费用记录并通过专用财务流程办理")
    record = await _ensure_record_module(body.finance_record_id, "finance", identity, db)
    if body.transaction_type != "付款":
        raise HTTPException(status_code=409, detail="开票和退款流水必须由发票或退费专用流程生成")
    if record.status not in {"已审批", "部分付款"}: raise HTTPException(status_code=409, detail="费用审批通过后才能付款")
    paid = await db.scalar(select(func.coalesce(func.sum(FinanceTransaction.amount), 0)).where(FinanceTransaction.finance_record_id == record.id, FinanceTransaction.transaction_type == "付款"))
    if float(paid or 0) + body.amount > float((record.data or {}).get("amount", 0)) + 0.001: raise HTTPException(status_code=409, detail="付款金额不能超过费用金额")
    item = FinanceTransaction(**body.model_dump(), operator=identity["username"]); db.add(item); await db.flush()
    if record:
        previous = record.status
        if body.transaction_type == "付款":
            paid_total = float(paid or 0) + body.amount
            record.status = "已付款" if paid_total + 0.001 >= float((record.data or {}).get("amount", 0)) else "部分付款"
        db.add(WorkflowEvent(record_id=record.id, action=f"登记{body.transaction_type}", from_status=previous, to_status=record.status, operator=identity["username"], comment=f"{body.amount:.2f} 元；{body.remark}"))
    await db.commit(); await db.refresh(item); return _finance_transaction_dict(item, record)


@app.delete(f"{settings.api_prefix}/finance/transactions/{{transaction_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_finance_transaction(transaction_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity["role"] != "admin": raise HTTPException(status_code=403, detail="仅管理员可删除")
    item = await db.get(FinanceTransaction, transaction_id)
    if not item: raise HTTPException(status_code=404, detail="财务流水不存在")
    record = await db.get(BusinessRecord, item.finance_record_id) if item.finance_record_id else None
    attachments = (await db.scalars(select(FileAttachment).where(FileAttachment.finance_transaction_id == item.id))).all()
    paths = [Path(x.path) for x in attachments]
    for attachment in attachments:
        await db.delete(attachment)
    await db.delete(item); await db.flush()
    if record and item.transaction_type == "付款":
        paid = float(await db.scalar(select(func.coalesce(func.sum(FinanceTransaction.amount), 0)).where(FinanceTransaction.finance_record_id == record.id, FinanceTransaction.transaction_type == "付款")) or 0)
        fee_amount = float((record.data or {}).get("amount", 0))
        record.status = "已审批" if paid <= 0 else ("已付款" if paid + 0.001 >= fee_amount else "部分付款")
    await db.commit()
    for path in paths:
        if path.is_file() and UPLOAD_ROOT.resolve() in path.resolve().parents:
            path.unlink()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.get(f"{settings.api_prefix}/finance/reconciliations")
async def list_reconciliations(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    items = (await db.scalars(select(ReconciliationBatch).order_by(ReconciliationBatch.date_to.desc(), ReconciliationBatch.id.desc()))).all()
    if identity.get("role") not in {"admin", "auditor"}: items = [item for item in items if item.operator == identity["username"]]
    show_amount = "finance.amount" in await _allowed_field_keys(identity, db)
    return {"items": [_reconciliation_dict(item, show_amount=show_amount) for item in items], "total": len(items)}


@app.post(f"{settings.api_prefix}/finance/reconciliations", status_code=status.HTTP_201_CREATED)
async def create_reconciliation(body: ReconciliationInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager", "auditor"}: raise HTTPException(status_code=403, detail="当前角色没有对账权限")
    if body.period_type not in {"周对账", "月对账"}: raise HTTPException(status_code=422, detail="对账周期无效")
    if body.date_from > body.date_to: raise HTTPException(status_code=422, detail="开始日期不能晚于结束日期")
    duplicate = await db.scalar(select(ReconciliationBatch.id).where(ReconciliationBatch.period_type == body.period_type, ReconciliationBatch.date_from == body.date_from, ReconciliationBatch.date_to == body.date_to))
    if duplicate: raise HTTPException(status_code=409, detail="该周期已经生成对账单")
    txs = (await db.scalars(select(FinanceTransaction).where(FinanceTransaction.transaction_date >= body.date_from, FinanceTransaction.transaction_date <= body.date_to))).all()
    if identity.get("role") != "admin":
        visible_record_ids = await _visible_record_ids(identity, db)
        txs = [item for item in txs if (item.finance_record_id and item.finance_record_id in visible_record_ids) or (not item.finance_record_id and item.operator == identity["username"])]
    item = ReconciliationBatch(**body.model_dump(), transaction_count=len(txs), total_amount=sum(tx.amount for tx in txs), status="待确认", operator=identity["username"])
    db.add(item); await db.commit(); await db.refresh(item); return _reconciliation_dict(item, show_amount="finance.amount" in await _allowed_field_keys(identity, db))


@app.post(f"{settings.api_prefix}/finance/reconciliations/{{batch_id}}/confirm")
async def confirm_reconciliation(batch_id: int, body: FinanceActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager", "auditor"}: raise HTTPException(status_code=403, detail="当前角色没有对账权限")
    item = await db.get(ReconciliationBatch, batch_id)
    if not item: raise HTTPException(status_code=404, detail="对账单不存在")
    if item.status == "已确认": raise HTTPException(status_code=409, detail="对账单已经确认")
    item.status = "已确认"; item.operator = identity["username"]; item.remark = body.comment or item.remark
    await db.commit(); await db.refresh(item); return _reconciliation_dict(item, show_amount="finance.amount" in await _allowed_field_keys(identity, db))


@app.delete(f"{settings.api_prefix}/finance/reconciliations/{{batch_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_reconciliation(batch_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity["role"] != "admin": raise HTTPException(status_code=403, detail="仅管理员可删除")
    item = await db.get(ReconciliationBatch, batch_id)
    if not item: raise HTTPException(status_code=404, detail="对账单不存在")
    await db.delete(item); await db.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.post(f"{settings.api_prefix}/cases/batch-update")
async def batch_update_cases(body: CaseBatchUpdateInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以批量修改案件")
    changes_requested = any(value is not None for value in (body.hearing_lawyer, body.handling_lawyers, body.assistant, body.case_stage))
    if not changes_requested:
        raise HTTPException(status_code=422, detail="请至少提供一个需要修改的案件字段")
    case_ids = list(dict.fromkeys(body.case_ids))
    case_nos = list(dict.fromkeys(value.strip() for value in body.case_nos if value.strip()))
    if not case_ids and not case_nos:
        raise HTTPException(status_code=422, detail="请至少选择一个案件 ID 或案号")
    if len(case_ids) + len(case_nos) > 100:
        raise HTTPException(status_code=422, detail="单次最多批量修改 100 个案件")
    handling_lawyers: list[str] | None = None
    handling_usernames: list[str] | None = None
    assistant_value: str | None = None
    assistant_username: str | None = None
    if body.handling_lawyers is not None:
        handling_lawyers, handling_usernames = await _resolve_active_case_people(body.handling_lawyers, db, field_name="经办律师")
        if not handling_lawyers:
            raise HTTPException(status_code=422, detail="请至少保留一名有效经办律师")
    if body.assistant is not None:
        assistant_values, assistant_usernames = await _resolve_active_case_people([body.assistant] if body.assistant.strip() else [], db, field_name="律师助理")
        assistant_value = assistant_values[0] if assistant_values else ""
        assistant_username = assistant_usernames[0] if assistant_usernames else ""
    requested = []
    if case_ids: requested.append(BusinessRecord.id.in_(case_ids))
    if case_nos: requested.append(BusinessRecord.serial_no.in_(case_nos))
    visible_cases = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "case", or_(*requested), *(await _record_scope_conditions(identity, db)),
    ))).all())
    by_id = {case.id: case for case in visible_cases}; by_no = {case.serial_no: case for case in visible_cases}
    missing_ids = [case_id for case_id in case_ids if case_id not in by_id]
    missing_nos = [case_no for case_no in case_nos if case_no not in by_no]
    if missing_ids or missing_nos:
        parts = []
        if missing_ids: parts.append("ID：" + "、".join(str(value) for value in missing_ids))
        if missing_nos: parts.append("案号：" + "、".join(missing_nos))
        raise HTTPException(status_code=404, detail="案件不存在或无权访问（" + "；".join(parts) + "）")
    cases = []
    for case in [*(by_id[value] for value in case_ids), *(by_no[value] for value in case_nos)]:
        if all(existing.id != case.id for existing in cases): cases.append(case)
    for case in cases:
        _require_case_creation_completed(case)
        if case.status in {"待归档审核", "已归档"}:
            raise HTTPException(status_code=409, detail=f"案件 {case.serial_no} 已进入归档流程，不能批量修改")
        data = dict(case.data or {})
        changes = []
        if body.hearing_lawyer is not None:
            changes.append(f"开庭律师：{data.get('hearing_lawyer', '')} → {body.hearing_lawyer.strip()}")
            data["hearing_lawyer"] = body.hearing_lawyer.strip()
        if body.handling_lawyers is not None:
            changes.append(f"经办律师：{','.join(data.get('handling_lawyers') or [])} → {','.join(handling_lawyers or [])}")
            data = _case_team_payload(data, handling_lawyers or [], handling_usernames or [], data.get("assistant", ""), str(data.get("assistant_username") or ""))
        if body.assistant is not None:
            changes.append(f"律师助理：{data.get('assistant', '')} → {assistant_value or ''}")
            data = _case_team_payload(data, list(data.get("handling_lawyers") or []), list(data.get("handling_lawyer_usernames") or []), assistant_value or "", assistant_username or "")
        if body.case_stage is not None:
            changes.append(f"案件阶段：{data.get('case_stage') or case.status} → {body.case_stage.strip()}")
            data["case_stage"] = body.case_stage.strip()
        case.data = data
        db.add(WorkflowEvent(record_id=case.id, action="批量修改案件", from_status=case.status, to_status=case.status, operator=identity["username"], comment="；".join(changes + ([body.comment.strip()] if body.comment.strip() else []))))
    await db.commit()
    for case in cases:
        await db.refresh(case)
    return {"updated": len(cases), "items": [await _record_dict_for_identity(case, identity, db) for case in cases]}


@app.get(f"{settings.api_prefix}/cases/{{case_id}}/reminders")
async def list_case_reminders(case_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    await _ensure_record_module(case_id, "case", identity, db)
    reminders = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "case_reminder",
    ).order_by(BusinessRecord.data["reminder_date"].as_string(), BusinessRecord.id))).all())
    items = [item for item in reminders if int((item.data or {}).get("case_id") or 0) == case_id]
    return {"items": [_record_dict(item) for item in items], "total": len(items)}


@app.post(f"{settings.api_prefix}/cases/{{case_id}}/reminders", status_code=status.HTTP_201_CREATED)
async def create_case_reminder(case_id: int, body: CaseReminderInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    await _require_case_detail_write_access(case_record, identity, db)
    if body.reminder_date > body.deadline:
        raise HTTPException(status_code=422, detail="提醒日期不能晚于截止日期")
    content = body.content.strip()
    if not content:
        raise HTTPException(status_code=422, detail="请输入提醒内容")
    item = BusinessRecord(
        module="case_reminder", serial_no=f"TX{datetime.now():%Y%m%d%H%M%S%f}",
        title=content[:255], customer=case_record.customer, status="有效",
        owner=identity["username"], department=case_record.department, description=content,
        data={"case_id": case_record.id, "case_no": case_record.serial_no,
              "reminder_date": str(body.reminder_date), "deadline": str(body.deadline)},
    )
    db.add(item)
    await db.flush()
    db.add(WorkflowEvent(
        record_id=case_record.id, action="新增案件提醒", from_status=case_record.status,
        to_status=case_record.status, operator=identity["username"],
        comment=f"提醒日期：{body.reminder_date}；截止日期：{body.deadline}；{content}",
    ))
    await db.commit()
    await db.refresh(item)
    return _record_dict(item)


@app.delete(f"{settings.api_prefix}/cases/{{case_id}}/reminders/{{reminder_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_case_reminder(case_id: int, reminder_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    await _require_case_detail_write_access(case_record, identity, db)
    item = await db.get(BusinessRecord, reminder_id)
    if not item or item.module != "case_reminder" or int((item.data or {}).get("case_id") or 0) != case_id:
        raise HTTPException(status_code=404, detail="案件提醒不存在")
    db.add(WorkflowEvent(
        record_id=case_record.id, action="删除案件提醒", from_status=case_record.status,
        to_status=case_record.status, operator=identity["username"],
        comment=f"提醒日期：{(item.data or {}).get('reminder_date', '')}；{item.description}",
    ))
    await db.delete(item)
    await db.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.get(f"{settings.api_prefix}/cases/{{case_id}}/logs")
async def list_case_logs(case_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    await _ensure_record_module(case_id, "case", identity, db)
    events = list((await db.scalars(select(WorkflowEvent).where(
        WorkflowEvent.record_id == case_id, WorkflowEvent.action == "新增案件日志",
    ).order_by(WorkflowEvent.created_at.desc(), WorkflowEvent.id.desc()))).all())
    return {"items": [{"id": item.id, "content": item.comment, "operator": item.operator, "created_at": item.created_at} for item in events], "total": len(events)}


@app.post(f"{settings.api_prefix}/cases/{{case_id}}/logs", status_code=status.HTTP_201_CREATED)
async def create_case_log(case_id: int, body: CaseLogInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    await _require_case_detail_write_access(case_record, identity, db)
    content = body.content.strip()
    if not content:
        raise HTTPException(status_code=422, detail="请输入日志内容")
    event = WorkflowEvent(
        record_id=case_record.id, action="新增案件日志", from_status=case_record.status,
        to_status=case_record.status, operator=identity["username"], comment=content,
    )
    db.add(event)
    await db.commit()
    await db.refresh(event)
    return {"id": event.id, "content": event.comment, "operator": event.operator, "created_at": event.created_at}


@app.post(f"{settings.api_prefix}/cases/batch-fees", status_code=status.HTTP_201_CREATED)
async def create_case_batch_fees(body: CaseBatchFeeInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_ids = list(dict.fromkeys(body.case_ids))
    cases = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "case", BusinessRecord.id.in_(case_ids),
        *(await _record_scope_conditions(identity, db)),
    ))).all())
    if len(cases) != len(case_ids):
        raise HTTPException(status_code=404, detail="存在无权访问或不存在的案件")
    expected_fee_type = EXPENSE_SUBTYPE_FEE_TYPE[body.expense_subtype]
    if expected_fee_type not in EXPENSE_SCOPE_FEE_TYPES[body.expense_scope]:
        raise HTTPException(status_code=422, detail="费用归属与费用子类型不一致")
    handler = body.handler.strip() or identity["username"]
    if identity.get("role") == "user":
        handler = identity["username"]
    handler_user = await db.scalar(select(User).where(User.username == handler, User.is_active.is_(True)))
    if not handler_user:
        raise HTTPException(status_code=422, detail="费用经办人不存在或已停用")
    ordered_cases = sorted(cases, key=lambda item: case_ids.index(item.id))
    for case_record in ordered_cases:
        await _require_record_owner_or_manager(case_record, identity, db)
        _require_case_creation_completed(case_record)
        if case_record.status in {"待归档审核", "已归档"}:
            raise HTTPException(status_code=409, detail=f"案件 {case_record.serial_no} 已进入归档流程，不能新增费用")
    created: list[BusinessRecord] = []
    amount = _round_fee_amount(body.amount)
    for case_record in ordered_cases:
        serial = f"FY{datetime.now():%Y%m%d%H%M%S%f}{uuid4().hex[:6]}"
        item = BusinessRecord(
            module="finance", serial_no=serial, title=f"{case_record.title}{body.expense_subtype}",
            customer=case_record.customer, status="草稿", owner=handler,
            department=case_record.department, description=body.description,
            data={"amount": amount, "fee_type": expected_fee_type,
                  "expense_scope": body.expense_scope, "expense_subtype": body.expense_subtype,
                  "is_refund": False, "case_no": case_record.serial_no, "case_id": case_record.id,
                  "contract_id": (case_record.data or {}).get("contract_record_id"),
                  "contract_no": (case_record.data or {}).get("contract_no", ""),
                  "handler": handler, "court": (case_record.data or {}).get("court", ""),
                  "document_no": "", "payee": (case_record.data or {}).get("court", "")},
        )
        db.add(item)
        await db.flush()
        created.append(item)
        db.add(WorkflowEvent(record_id=item.id, action="批量创建案件费用", to_status="草稿", operator=identity["username"], comment=f"{case_record.serial_no}｜{body.expense_scope}{body.expense_subtype}：{amount:.2f} 元"))
        db.add(WorkflowEvent(record_id=case_record.id, action="批量新增案件费用", from_status=case_record.status, to_status=case_record.status, operator=identity["username"], comment=f"{item.serial_no}｜{body.expense_scope}{body.expense_subtype}：{amount:.2f} 元"))
    await db.commit()
    for item in created:
        await db.refresh(item)
    return {"created": len(created), "items": [await _record_dict_for_identity(item, identity, db) for item in created]}


@app.get(f"{settings.api_prefix}/cases/summary")
async def case_summary(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    cases = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "case", *(await _record_scope_conditions(identity, db))))).all()
    return {
        "total": len(cases),
        "pending_assignment": sum(1 for item in cases if item.status == "新案待分配"),
        "in_progress": sum(1 for item in cases if item.status not in {"新案待分配", "已归档"}),
        "execution": sum(1 for item in cases if item.status == "执行"),
        "archived": sum(1 for item in cases if item.status == "已归档"),
    }


@app.post(f"{settings.api_prefix}/cases/invoice-files/import")
async def import_case_invoice_files(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """把当前用户上传的案件发票文件按文件名中的案件编号匹配到案件。"""
    pending = (await db.scalars(
        select(FileAttachment).where(
            FileAttachment.category == "案件发票文件",
            FileAttachment.uploader == identity["username"],
            FileAttachment.record_id.is_(None),
        ).order_by(FileAttachment.id)
    )).all()
    cases = (await db.scalars(
        select(BusinessRecord).where(
            BusinessRecord.module == "case",
            *(await _record_scope_conditions(identity, db)),
        )
    )).all()
    matched = 0
    unmatched = 0
    for attachment in pending:
        normalized = attachment.original_name.upper().replace(" ", "")
        case = next((item for item in cases if item.serial_no.upper() in normalized), None)
        if case:
            attachment.record_id = case.id
            attachment.remark = f"案件发票文件导入｜自动匹配 {case.serial_no}"
            db.add(WorkflowEvent(record_id=case.id, action="导入案件发票文件", from_status=case.status, to_status=case.status, operator=identity["username"], comment=attachment.original_name))
            matched += 1
        else:
            attachment.remark = "案件发票文件导入｜文件名未识别案件编号"
            unmatched += 1
    await db.commit()
    return {"processed": len(pending), "matched": matched, "unmatched": unmatched}


async def _query_counsel_cases(body: CounselCaseSearchInput, identity: dict, db: AsyncSession) -> list[BusinessRecord]:
    if body.scope not in {"mine", "department", "company"}:
        raise HTTPException(status_code=422, detail="法律顾问案件查询范围无效")
    if body.sort_order not in {"updated_desc", "case_no_asc", "case_no_desc"}:
        raise HTTPException(status_code=422, detail="法律顾问案件排序方式无效")
    records = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "case",
        *(await _record_scope_conditions(identity, db)),
    ))).all())
    records = [record for record in records if str((record.data or {}).get("case_type") or "") == "法律顾问"]
    if body.scope == "mine":
        identity_names = {str(identity.get("username") or ""), str(identity.get("display_name") or "")}
        records = [record for record in records if record.owner in identity_names]
    elif body.scope == "department":
        department = str(identity.get("department") or "").strip()
        records = [record for record in records if department and record.department == department]

    document_names: dict[int, str] = {}
    record_ids = [record.id for record in records]
    if record_ids:
        attachments = (await db.scalars(select(FileAttachment).where(FileAttachment.record_id.in_(record_ids)))).all()
        for attachment in attachments:
            if attachment.record_id is not None:
                document_names[attachment.record_id] = f"{document_names.get(attachment.record_id, '')} {attachment.original_name}".strip()

    def contains(value: object, expected: str) -> bool:
        return not expected.strip() or expected.strip().casefold() in str(value or "").casefold()

    filtered: list[BusinessRecord] = []
    for record in records:
        data = record.data or {}
        if not contains(record.customer, body.customer): continue
        if not contains(record.serial_no, body.serial_no): continue
        if body.keyword and not contains(f"{record.serial_no} {record.title} {record.customer}", body.keyword): continue
        if not contains(data.get("counsel_type"), body.counsel_type): continue
        if not contains(record.status, body.case_status): continue
        if not contains("、".join(data.get("handling_lawyers") or []), body.handling_lawyer): continue
        if not contains(data.get("assistant"), body.assistant): continue
        if not contains(document_names.get(record.id, ""), body.document_name): continue
        try:
            record_start = date.fromisoformat(str(data.get("counsel_start") or ""))
            record_end = date.fromisoformat(str(data.get("counsel_end") or ""))
        except ValueError:
            record_start = record_end = None
        if body.counsel_start and (not record_end or record_end < body.counsel_start): continue
        if body.counsel_end and (not record_start or record_start > body.counsel_end): continue
        filtered.append(record)
    if body.sort_order == "case_no_asc":
        filtered.sort(key=lambda item: (item.serial_no, item.id))
    elif body.sort_order == "case_no_desc":
        filtered.sort(key=lambda item: (item.serial_no, item.id), reverse=True)
    else:
        filtered.sort(key=lambda item: (item.updated_at or item.created_at, item.id), reverse=True)
    return filtered


@app.post(f"{settings.api_prefix}/cases/counsel/search")
async def search_counsel_cases(body: CounselCaseSearchInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    records = await _query_counsel_cases(body, identity, db)
    total = len(records)
    start = (body.page - 1) * body.page_size
    allowed_fields = await _allowed_field_keys(identity, db)
    return {
        "items": [_record_dict(record, allowed_fields) for record in records[start:start + body.page_size]],
        "total": total,
        "page": body.page,
        "page_size": body.page_size,
    }


@app.post(f"{settings.api_prefix}/cases/counsel/export")
async def export_counsel_cases(body: CounselCaseSearchInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    records = await _query_counsel_cases(body, identity, db)
    if body.selected_only:
        selected_ids = list(dict.fromkeys(body.selected_ids))
        if not selected_ids:
            raise HTTPException(status_code=422, detail="请选择需要导出的法律顾问案件")
        selected_set = set(selected_ids)
        records = [record for record in records if record.id in selected_set]
        if len(records) != len(selected_set):
            raise HTTPException(status_code=403, detail="选中的案件不存在、不可见或不符合当前查询条件")
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["案件编号", "案件名称", "顾问类型", "客户", "顾问开始日期", "顾问结束日期", "经办律师", "律师助理", "案源人", "案件阶段", "所属部门"])
    for record in records:
        data = record.data or {}
        writer.writerow([
            record.serial_no, record.title, data.get("counsel_type", ""), record.customer,
            data.get("counsel_start", ""), data.get("counsel_end", ""),
            "、".join(data.get("handling_lawyers") or []), data.get("assistant", ""),
            data.get("source_person") or record.owner, record.status, record.department,
        ])
    scope_label = "selected" if body.selected_only else "all"
    content = ("\ufeff" + output.getvalue()).encode("utf-8")
    return Response(
        content=content,
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="counsel-cases-{scope_label}-{date.today()}.csv"'},
    )


@app.get(f"{settings.api_prefix}/cases/eligible-contracts")
async def list_case_eligible_contracts(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Return every visible contract that can actually start the staged case flow."""
    conditions = [
        BusinessRecord.module == "contract",
        BusinessRecord.status.in_(CASE_SOURCE_CONTRACT_STATUSES),
        *(await _record_scope_conditions(identity, db)),
    ]
    contracts = (await db.scalars(
        select(BusinessRecord).where(*conditions).order_by(BusinessRecord.updated_at.desc(), BusinessRecord.id.desc())
    )).all()
    allowed_fields = await _allowed_field_keys(identity, db)
    return {"items": [_record_dict(item, allowed_fields) for item in contracts], "total": len(contracts)}


@app.get(f"{settings.api_prefix}/cases/reference-options")
async def list_case_reference_options(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Return active case dictionaries needed by the staged create form."""
    causes = (await db.scalars(select(SystemParameter).where(
        SystemParameter.category == "cause", SystemParameter.is_active.is_(True),
    ).order_by(SystemParameter.sort_order, SystemParameter.id))).all()
    return {
        "case_types": [
            {"value": "民事案件", "label": "民事争议"},
            {"value": "刑事案件", "label": "刑事案件"},
            {"value": "行政案件及国家赔偿", "label": "行政案件及国家赔偿"},
            {"value": "法律顾问", "label": "法律顾问"},
            {"value": "仲裁", "label": "仲裁"},
        ],
        "causes": [{"value": item.name, "label": item.name, "code": item.code} for item in causes],
        "right_types": ["商标权", "专利权", "著作权", "不正当竞争", "商业秘密", "其他"],
    }


async def _ensure_case_fixed_tasks(case_record: BusinessRecord, db: AsyncSession, *, operator: str) -> list[BusinessRecord]:
    """Create the mandatory standard task set once for each case."""
    existing = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.module == "task",
        BusinessRecord.data["case_id"].as_integer() == case_record.id,
        BusinessRecord.data["task_type"].as_string() == "固定任务",
    ))).all())
    existing_keys = {str((item.data or {}).get("fixed_task_key") or "") for item in existing}
    specs = [
        ("filing-registration", "立案登记", 7, "完成法院立案信息登记并上传受理材料"),
        ("service-tracking", "送达跟踪", 14, "跟踪法院送达情况并记录送达结果"),
    ]
    created: list[BusinessRecord] = []
    for key, title, days, description in specs:
        if key in existing_keys:
            continue
        task = BusinessRecord(
            module="task", serial_no=f"RW{datetime.now():%Y%m%d%H%M%S%f}{uuid4().hex[:4].upper()}",
            title=f"{title}—{case_record.serial_no}", customer=case_record.customer, status="待接收",
            owner=case_record.owner, department=case_record.department, description=description,
            data={
                "deadline": str(date.today() + timedelta(days=days)), "priority": "普通", "source": "案件任务",
                "task_type": "固定任务", "fixed_task_key": key, "initiator": operator,
                "collaborators": [], "case_no": case_record.serial_no, "case_id": case_record.id,
                "case_stage": "立案", "system_created_by": operator,
            },
        )
        db.add(task); await db.flush()
        await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action="生成案件固定任务", to_status="待接收", operator=operator, comment=f"案件 {case_record.serial_no} 创建时自动生成"), db, content="案件固定任务已生成.")
        created.append(task)
    if created:
        case_record.data = {**(case_record.data or {}), "fixed_tasks_generated": True, "fixed_task_ids": [*list((case_record.data or {}).get("fixed_task_ids", [])), *[item.id for item in created]]}
    return [*existing, *created]


async def _next_case_serial(case_type: str, db: AsyncSession) -> str:
    """Generate the compact legacy-style case number: SH + type + YY + 5 digits."""
    parameter_names = {case_type}
    if case_type == "民事案件":
        parameter_names.add("民事争议")
    parameter = await db.scalar(select(SystemParameter).where(
        SystemParameter.category == "case_type",
        SystemParameter.name.in_(parameter_names),
        SystemParameter.is_active.is_(True),
    ).order_by(SystemParameter.sort_order, SystemParameter.id))
    default_codes = {"民事案件": "MS", "刑事案件": "XS", "行政案件及国家赔偿": "XZ", "法律顾问": "GW", "仲裁": "ZC"}
    type_code = str((parameter.extra or {}).get("letter_code") if parameter else default_codes.get(case_type, "AJ")).strip().upper()
    prefix = f"SH{type_code}{datetime.now():%y}"
    existing = (await db.scalars(select(BusinessRecord.serial_no).where(
        BusinessRecord.module == "case",
        BusinessRecord.serial_no.like(f"{prefix}%"),
    ).order_by(BusinessRecord.serial_no.desc()))).all()
    sequence = max((int(match.group(1)) for value in existing if (match := re.fullmatch(rf"{re.escape(prefix)}(\d{{5}})", value))), default=0) + 1
    if sequence > 99999:
        raise HTTPException(status_code=409, detail=f"{datetime.now():%Y} 年{case_type}案件编号已用尽")
    return f"{prefix}{sequence:05d}"


@app.post(f"{settings.api_prefix}/cases", status_code=status.HTTP_201_CREATED)
async def create_case(body: CaseCreateInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """从有效合同建立案件；客户、部门和合同编号均以合同资料为准。"""
    title = body.title.strip()
    case_type = body.case_type.strip()
    cause_or_charge = body.cause_or_charge.strip()
    if not title:
        raise HTTPException(status_code=422, detail="案件名称不能为空")
    if case_type not in CASE_CREATABLE_TYPES:
        raise HTTPException(status_code=422, detail="案件类型不是原系统允许的新建类型")
    serial_no = body.serial_no.strip() or await _next_case_serial(case_type, db)
    if body.status != "新案待分配":
        raise HTTPException(status_code=422, detail="新建案件阶段必须为待分配")
    counsel_type = body.counsel_type.strip()
    if case_type != "法律顾问" and not cause_or_charge:
        raise HTTPException(status_code=422, detail="罪名或案由不能为空")
    if case_type == "法律顾问":
        if not counsel_type:
            raise HTTPException(status_code=422, detail="顾问类型不能为空")
        if len(counsel_type) > 128:
            raise HTTPException(status_code=422, detail="顾问类型过长")
        if not body.counsel_start or not body.counsel_end:
            raise HTTPException(status_code=422, detail="顾问期限不能为空")
        if body.counsel_start > body.counsel_end:
            raise HTTPException(status_code=422, detail="顾问结束日期不能早于开始日期")
        cause_or_charge = ""
    elif counsel_type or body.counsel_start or body.counsel_end:
        raise HTTPException(status_code=422, detail="仅法律顾问案件可以填写顾问类型和顾问期限")
    handling_lawyers = list(dict.fromkeys(str(item or "").strip() for item in body.handling_lawyers if str(item or "").strip()))
    if not handling_lawyers or any(len(item) > 128 for item in handling_lawyers):
        raise HTTPException(status_code=422, detail="请按顺序录入有效的经办律师")
    client_position = body.client_position.strip()
    if case_type == "法律顾问":
        client_position = ""
    allowed_client_positions = CASE_CLIENT_POSITIONS_BY_TYPE.get(case_type)
    if allowed_client_positions and client_position not in allowed_client_positions:
        raise HTTPException(status_code=422, detail=f"{case_type}客户诉讼地位无效")
    if case_type == "行政案件及国家赔偿" and client_position not in ADMINISTRATIVE_CLIENT_POSITIONS:
        raise HTTPException(status_code=422, detail="行政案件客户诉讼地位无效")
    right_type = body.right_type.strip()
    if case_type == "法律顾问":
        right_type = ""
    if len(right_type) > 128:
        raise HTTPException(status_code=422, detail="权利类型过长")
    assistant = body.assistant.strip()
    if len(assistant) > 128:
        raise HTTPException(status_code=422, detail="律师助理姓名过长")
    handling_lawyers, handling_usernames = await _resolve_active_case_people(handling_lawyers, db, field_name="经办律师")
    assistant_values, assistant_usernames = await _resolve_active_case_people([assistant] if assistant else [], db, field_name="律师助理")
    assistant = assistant_values[0] if assistant_values else ""
    assistant_username = assistant_usernames[0] if assistant_usernames else ""
    permission_key = CASE_CREATE_PERMISSION_BY_TYPE[case_type]
    if identity.get("role") != "admin":
        permission = await _permission_payload(identity.get("role", "user"), db)
        if permission_key not in set(permission.get("menu_keys", [])):
            raise HTTPException(status_code=403, detail="当前角色没有该案件类型的新建权限")
    contract = await _ensure_record_visible(body.contract_record_id, identity, db)
    if contract.module != "contract":
        raise HTTPException(status_code=422, detail="关联记录不是合同")
    if contract.status not in CASE_SOURCE_CONTRACT_STATUSES:
        raise HTTPException(status_code=409, detail="合同审批通过后才能新建案件")
    department = contract.department.strip()
    if not department or not await db.scalar(select(Department.id).where(Department.name == department, Department.is_active.is_(True))):
        raise HTTPException(status_code=409, detail="关联合同没有有效的所属部门")
    if await db.scalar(select(BusinessRecord.id).where(BusinessRecord.serial_no == serial_no)):
        raise HTTPException(status_code=409, detail="业务编号已存在")
    contract_data = contract.data or {}
    owner = body.owner.strip() or identity["username"]
    if identity.get("role") != "admin":
        user = await db.scalar(select(User).where(User.username == identity["username"]))
        if not user or not user.is_active:
            raise HTTPException(status_code=401, detail="当前用户不存在")
        owner = user.username
    owner_user = await db.scalar(select(User).where(User.username == owner, User.is_active.is_(True)))
    if not owner_user:
        raise HTTPException(status_code=422, detail="案件负责人必须是有效用户")
    record = BusinessRecord(
        module="case", serial_no=serial_no, title=title,
        customer=contract.customer, status="新案待分配", owner=owner,
        department=department, description="",
        data={
            "contract_id": contract.id,
            "contract_no": contract.serial_no,
            "external_contract_no": contract_data.get("external_contract_no", ""),
            "external_contract_numbers": contract_data.get("external_contract_numbers", []),
            "contract_title": contract.title,
            "case_type": case_type,
            "client_position": client_position,
            "cause_or_charge": cause_or_charge,
            "right_type": right_type,
            "source_person": contract_data.get("source_person") or contract.owner,
            "counsel_type": counsel_type,
            "counsel_start": str(body.counsel_start) if body.counsel_start else "",
            "counsel_end": str(body.counsel_end) if body.counsel_end else "",
            **_case_team_payload({}, handling_lawyers, handling_usernames, assistant, assistant_username),
            "case_creation_step": "basic",
            "case_creation_approval_status": "未提交",
            "business_stage": "立案",
        },
    )
    db.add(record)
    try:
        await db.flush()
    except IntegrityError as exc:
        await db.rollback()
        raise HTTPException(status_code=409, detail="业务编号已存在") from exc
    db.add(WorkflowEvent(
        record_id=record.id, action="从合同新建案件", to_status=record.status,
        operator=identity["username"], comment=f"关联合同：{contract.serial_no}｜{contract.title}",
    ))
    await db.commit()
    await db.refresh(record)
    return _record_dict(record)


@app.put(f"{settings.api_prefix}/cases/{{case_id}}/litigants")
async def update_case_litigants(case_id: int, body: CaseLitigantsInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    await _require_record_owner_or_manager(case_record, identity, db)
    if case_record.status in {"待归档审核", "已归档"}:
        raise HTTPException(status_code=409, detail="归档中的案件不能修改当事人")
    creation_step = str((case_record.data or {}).get("case_creation_step") or "")
    if creation_step not in {"basic", "litigants"}:
        raise HTTPException(status_code=409, detail="当前案件不处于新建当事人信息阶段")

    def clean_parties(values: list[str]) -> list[str]:
        result = list(dict.fromkeys(str(value or "").strip() for value in values if str(value or "").strip()))
        if any(len(value) > 256 for value in result):
            raise HTTPException(status_code=422, detail="当事人名称过长")
        return result

    plaintiffs = clean_parties(body.plaintiffs)
    plaintiff_agents = clean_parties(body.plaintiff_agents)
    defendants = clean_parties(body.defendants)
    defendant_agents = clean_parties(body.defendant_agents)
    third_parties = clean_parties(body.third_parties)
    third_party_agents = clean_parties(body.third_party_agents)
    case_type = str((case_record.data or {}).get("case_type") or "")
    permission_key = CASE_CREATE_PERMISSION_BY_TYPE.get(case_type)
    if permission_key and identity.get("role") != "admin":
        permission = await _permission_payload(identity.get("role", "user"), db)
        if permission_key not in set(permission.get("menu_keys", [])):
            raise HTTPException(status_code=403, detail="当前角色没有该案件类型的新建权限")
    if case_type in {"行政案件及国家赔偿", "仲裁"} and (not plaintiffs or not defendants):
        raise HTTPException(status_code=422, detail="请录入原告/申请人与被告/被申请人")
    case_record.data = {
        **(case_record.data or {}),
        "plaintiffs": plaintiffs,
        "plaintiff_agents": plaintiff_agents,
        "defendants": defendants,
        "defendant_agents": defendant_agents,
        "third_parties": third_parties,
        "third_party_agents": third_party_agents,
        "case_creation_step": "litigants",
    }
    db.add(WorkflowEvent(
        record_id=case_record.id,
        action="维护当事人信息",
        from_status=case_record.status,
        to_status=case_record.status,
        operator=identity["username"],
        comment=body.comment,
    ))
    await db.commit()
    await db.refresh(case_record)
    return _record_dict(case_record)


@app.put(f"{settings.api_prefix}/cases/{{case_id}}/complete-creation")
async def complete_case_creation(case_id: int, body: CaseCreationCompleteInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    await _require_record_owner_or_manager(case_record, identity, db)
    if case_record.status in {"待归档审核", "已归档"}:
        raise HTTPException(status_code=409, detail="归档中的案件不能完成新建")
    case_data = case_record.data or {}
    if str(case_data.get("case_type") or "") != "法律顾问":
        raise HTTPException(status_code=409, detail="当前案件类型必须通过司法机关步骤完成新建")
    if str(case_data.get("case_creation_step") or "") != "litigants":
        raise HTTPException(status_code=409, detail="请先完成当事人信息")
    permission_key = CASE_CREATE_PERMISSION_BY_TYPE["法律顾问"]
    if identity.get("role") != "admin":
        permission = await _permission_payload(identity.get("role", "user"), db)
        if permission_key not in set(permission.get("menu_keys", [])):
            raise HTTPException(status_code=403, detail="当前角色没有法律顾问案件新建权限")
    previous_status = case_record.status
    case_record.status = "待立案审批"
    case_record.data = {
        **case_data,
        "case_creation_step": "completed",
        "case_creation_completed_at": datetime.now().isoformat(timespec="seconds"),
        "case_creation_completed_by": identity["username"],
        "case_creation_approval_status": "待审批",
        "case_creation_submitted_at": datetime.now().isoformat(timespec="seconds"),
        "case_creation_submitted_by": identity["username"],
    }
    db.add(WorkflowEvent(
        record_id=case_record.id,
        action="完成法律顾问案件新建",
        from_status=previous_status,
        to_status=case_record.status,
        operator=identity["username"],
        comment=body.comment or "案件资料填写完成，提交案件主管审批",
    ))
    await db.commit()
    await db.refresh(case_record)
    return _record_dict(case_record)


@app.put(f"{settings.api_prefix}/cases/{{case_id}}/counsel-basic")
async def update_counsel_case_basic(case_id: int, body: CaseCounselBasicInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    await _require_record_owner_or_manager(case_record, identity, db)
    case_data = case_record.data or {}
    if str(case_data.get("case_type") or "") != "法律顾问":
        raise HTTPException(status_code=409, detail="该接口仅用于法律顾问案件")
    if case_record.status in {"待归档审核", "已归档"}:
        raise HTTPException(status_code=409, detail="归档中的法律顾问案件不能修改基本信息")
    if str(case_data.get("case_creation_step") or "") != "completed":
        raise HTTPException(status_code=409, detail="请先完成法律顾问案件新建流程")
    title = body.title.strip()
    counsel_type = body.counsel_type.strip()
    if not title or not counsel_type:
        raise HTTPException(status_code=422, detail="案件名称和顾问类型不能为空")
    if body.counsel_start > body.counsel_end:
        raise HTTPException(status_code=422, detail="顾问结束日期不能早于开始日期")
    handling_lawyers = list(dict.fromkeys(str(item or "").strip() for item in body.handling_lawyers if str(item or "").strip()))
    assistant = body.assistant.strip()
    handling_lawyers, handling_usernames = await _resolve_active_case_people(handling_lawyers, db, field_name="经办律师")
    assistant_values, assistant_usernames = await _resolve_active_case_people([assistant] if assistant else [], db, field_name="律师助理")
    assistant = assistant_values[0] if assistant_values else ""
    assistant_username = assistant_usernames[0] if assistant_usernames else ""
    if not handling_lawyers:
        raise HTTPException(status_code=422, detail="请至少保留一名有效经办律师")
    old_summary = f"{case_record.title}｜{case_data.get('counsel_type', '')}｜{case_data.get('counsel_start', '')}至{case_data.get('counsel_end', '')}"
    case_record.title = title
    case_record.data = _case_team_payload({
        **case_data,
        "counsel_type": counsel_type,
        "counsel_start": str(body.counsel_start),
        "counsel_end": str(body.counsel_end),
    }, handling_lawyers, handling_usernames, assistant, assistant_username)
    db.add(WorkflowEvent(
        record_id=case_record.id,
        action="修改法律顾问案件基本信息",
        from_status=case_record.status,
        to_status=case_record.status,
        operator=identity["username"],
        comment=f"修改前：{old_summary}" + (f"｜说明：{body.comment.strip()}" if body.comment.strip() else ""),
    ))
    await db.commit()
    await db.refresh(case_record)
    return _record_dict(case_record)


@app.put(f"{settings.api_prefix}/cases/{{case_id}}/normal-basic")
async def update_normal_case_basic(case_id: int, body: CaseNormalBasicInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Update the evidenced civil/criminal/administrative basic-information branch.

    The endpoint is intentionally separate from counsel-basic and from generic
    record PATCH so archived cases and the case lifecycle cannot be bypassed.
    """
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    await _require_record_owner_or_manager(case_record, identity, db)
    case_data = case_record.data or {}
    case_type = str(case_data.get("case_type") or "")
    if case_type not in NORMAL_CASE_BASIC_TYPES:
        raise HTTPException(status_code=409, detail="该接口仅用于民事、刑事、行政及国家赔偿案件")
    _require_case_creation_completed(case_record)
    if case_record.status in {"待归档审核", "已归档"}:
        raise HTTPException(status_code=409, detail="归档中的案件不能修改基本信息")
    title = body.title.strip()
    phase = body.case_phase.strip()
    cause_or_charge = body.cause_or_charge.strip()
    if phase not in CASE_BASIC_EDITABLE_PHASES:
        raise HTTPException(status_code=422, detail="案件阶段不是允许的办理阶段")
    if not title or not cause_or_charge:
        raise HTTPException(status_code=422, detail="案件名称、案由或罪名不能为空")
    customer = await _customer_or_404(body.customer_record_id, identity, db)
    if customer.status in {"公海", "已回收"}:
        raise HTTPException(status_code=409, detail="不能关联公海或回收站客户")
    handling_lawyers = list(dict.fromkeys(str(item or "").strip() for item in body.handling_lawyers if str(item or "").strip()))
    handling_lawyers, handling_usernames = await _resolve_active_case_people(handling_lawyers, db, field_name="经办律师")
    if not handling_lawyers:
        raise HTTPException(status_code=422, detail="请至少保留一名有效经办律师")
    assistant_values, assistant_usernames = await _resolve_active_case_people([body.assistant.strip()] if body.assistant.strip() else [], db, field_name="律师助理")
    investigator_values, _ = await _resolve_active_case_people([body.investigator.strip()] if body.investigator.strip() else [], db, field_name="调查员")
    business_owner_values, _ = await _resolve_active_case_people([body.business_owner.strip()] if body.business_owner.strip() else [], db, field_name="案源人")
    clue_ids = list(dict.fromkeys(body.investigation_clue_ids))
    clues = list((await db.scalars(select(BusinessRecord).where(
        BusinessRecord.id.in_(clue_ids), BusinessRecord.module == "clue", *(await _record_scope_conditions(identity, db)),
    ))).all()) if clue_ids else []
    clues_by_id = {item.id: item for item in clues}
    if len(clues_by_id) != len(clue_ids):
        raise HTTPException(status_code=404, detail="关联调查线索不存在或无权访问")
    ordered_clues = [clues_by_id[item_id] for item_id in clue_ids]
    right_type = body.right_type.strip()
    if case_type != "行政案件及国家赔偿" and right_type:
        raise HTTPException(status_code=422, detail="仅行政及国家赔偿案件可以修改权利类型")
    old_summary = f"{case_record.customer}｜{case_record.title}｜{case_record.status}｜{case_data.get('cause_or_charge', '')}"
    previous_status = case_record.status
    assistant = assistant_values[0] if assistant_values else ""
    assistant_username = assistant_usernames[0] if assistant_usernames else ""
    investigator = investigator_values[0] if investigator_values else ""
    business_owner = business_owner_values[0] if business_owner_values else ""
    clue_nos = [item.serial_no for item in ordered_clues]
    case_record.title = title
    case_record.customer = customer.title
    case_record.status = phase
    case_record.data = _case_team_payload({
        **case_data,
        "customer_record_id": customer.id,
        "customer_id": customer.id,
        "customer_no": customer.serial_no,
        "cause_or_charge": cause_or_charge,
        "right_type": right_type if case_type == "行政案件及国家赔偿" else str(case_data.get("right_type") or ""),
        "source_person": business_owner or str(case_data.get("source_person") or ""),
        "business_owner": business_owner or str(case_data.get("business_owner") or ""),
        "investigator": investigator,
        "investigation_clue_ids": clue_ids,
        "investigation_clue_nos": clue_nos,
        "investigation_clue_id": clue_ids[0] if clue_ids else None,
        "investigation_clue": "、".join(clue_nos),
        "clue_record_id": clue_ids[0] if clue_ids else None,
        "clue_no": clue_nos[0] if clue_nos else "",
    }, handling_lawyers, handling_usernames, assistant, assistant_username)
    db.add(WorkflowEvent(
        record_id=case_record.id, action="修改普通案件基本信息",
        from_status=previous_status, to_status=case_record.status, operator=identity["username"],
        comment=f"修改前：{old_summary}" + (f"｜说明：{body.comment.strip()}" if body.comment.strip() else ""),
    ))
    await db.commit()
    await db.refresh(case_record)
    return _record_dict(case_record)


@app.put(f"{settings.api_prefix}/cases/{{case_id}}/arbitration-basic")
async def update_arbitration_case_basic(case_id: int, body: CaseArbitrationBasicInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Keep the old arbitration edit branch isolated from normal/counsel cases."""
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    await _require_record_owner_or_manager(case_record, identity, db)
    case_data = case_record.data or {}
    if str(case_data.get("case_type") or "") != "仲裁":
        raise HTTPException(status_code=409, detail="该接口仅用于仲裁案件")
    _require_case_creation_completed(case_record)
    if case_record.status in {"待归档审核", "已归档"}:
        raise HTTPException(status_code=409, detail="归档中的仲裁案件不能修改基本信息")
    title, phase, cause_or_charge = body.title.strip(), body.case_phase.strip(), body.cause_or_charge.strip()
    if phase not in CASE_BASIC_EDITABLE_PHASES:
        raise HTTPException(status_code=422, detail="案件阶段不是允许的办理阶段")
    if not title or not cause_or_charge:
        raise HTTPException(status_code=422, detail="案件名称和案由不能为空")
    customer = await _customer_or_404(body.customer_record_id, identity, db)
    if customer.status in {"公海", "已回收"}:
        raise HTTPException(status_code=409, detail="不能关联公海或回收站客户")
    lawyers = list(dict.fromkeys(str(item or "").strip() for item in body.handling_lawyers if str(item or "").strip()))
    lawyers, lawyer_usernames = await _resolve_active_case_people(lawyers, db, field_name="经办律师")
    if not lawyers:
        raise HTTPException(status_code=422, detail="请至少保留一名有效经办律师")
    assistant_values, assistant_usernames = await _resolve_active_case_people([body.assistant.strip()] if body.assistant.strip() else [], db, field_name="律师助理")
    investigator_values, _ = await _resolve_active_case_people([body.investigator.strip()] if body.investigator.strip() else [], db, field_name="调查员")
    clue_ids = list(dict.fromkeys(body.investigation_clue_ids))
    clues = list((await db.scalars(select(BusinessRecord).where(BusinessRecord.id.in_(clue_ids), BusinessRecord.module == "clue", *(await _record_scope_conditions(identity, db))))).all()) if clue_ids else []
    if len({item.id for item in clues}) != len(clue_ids):
        raise HTTPException(status_code=404, detail="关联调查线索不存在或无权访问")
    by_id = {item.id: item for item in clues}; clue_nos = [by_id[item_id].serial_no for item_id in clue_ids]
    previous_status = case_record.status
    old_summary = f"{case_record.customer}｜{case_record.title}｜{case_record.status}｜{case_data.get('cause_or_charge', '')}"
    case_record.title, case_record.customer, case_record.status = title, customer.title, phase
    case_record.data = _case_team_payload({
        **case_data, "customer_record_id": customer.id, "customer_id": customer.id, "customer_no": customer.serial_no,
        "cause_or_charge": cause_or_charge, "investigator": investigator_values[0] if investigator_values else "",
        "investigation_clue_ids": clue_ids, "investigation_clue_nos": clue_nos,
        "investigation_clue_id": clue_ids[0] if clue_ids else None, "investigation_clue": "、".join(clue_nos),
        "clue_record_id": clue_ids[0] if clue_ids else None, "clue_no": clue_nos[0] if clue_nos else "",
    }, lawyers, lawyer_usernames, assistant_values[0] if assistant_values else "", assistant_usernames[0] if assistant_usernames else "")
    db.add(WorkflowEvent(record_id=case_record.id, action="修改仲裁案件基本信息", from_status=previous_status, to_status=case_record.status, operator=identity["username"], comment=f"修改前：{old_summary}" + (f"｜说明：{body.comment.strip()}" if body.comment.strip() else "")))
    await db.commit(); await db.refresh(case_record)
    return _record_dict(case_record)


async def _criminal_detail_maintenance_case(case_id: int, identity: dict, db: AsyncSession) -> BusinessRecord:
    record = await _ensure_record_module(case_id, "case", identity, db); await _require_record_owner_or_manager(record, identity, db)
    if str((record.data or {}).get("case_type") or "") != "刑事案件": raise HTTPException(status_code=409, detail="该接口仅用于刑事案件")
    _require_case_creation_completed(record)
    if record.status in {"待归档审核", "已归档"}: raise HTTPException(status_code=409, detail="归档中的刑事案件不能维护资料")
    return record


async def _save_criminal_detail(record: BusinessRecord, payload: dict, action: str, comment: str, identity: dict, db: AsyncSession):
    record.data = {**(record.data or {}), **payload}
    db.add(WorkflowEvent(record_id=record.id, action=action, from_status=record.status, to_status=record.status, operator=identity["username"], comment=comment.strip()))
    await db.commit(); await db.refresh(record); return _record_dict(record)


@app.put(f"{settings.api_prefix}/cases/{{case_id}}/criminal/litigants")
async def maintain_criminal_litigants(case_id: int, body: CaseLitigantsInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    record = await _criminal_detail_maintenance_case(case_id, identity, db)
    def clean(items: list[str]): return list(dict.fromkeys(str(x or "").strip() for x in items if str(x or "").strip()))
    payload = {key: clean(getattr(body, key)) for key in ("plaintiffs","plaintiff_agents","defendants","defendant_agents","third_parties","third_party_agents")}
    return await _save_criminal_detail(record, payload, "修改刑事案件当事人", body.comment, identity, db)


@app.put(f"{settings.api_prefix}/cases/{{case_id}}/criminal/public-security")
async def maintain_criminal_public_security(case_id: int, body: CriminalPublicSecurityMaintenanceInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    record = await _criminal_detail_maintenance_case(case_id, identity, db)
    return await _save_criminal_detail(record, {k: str(v or "").strip() for k,v in body.model_dump(exclude={"comment"}).items()}, "修改刑事案件公安机关信息", body.comment, identity, db)


@app.put(f"{settings.api_prefix}/cases/{{case_id}}/criminal/procuratorates")
async def maintain_criminal_procuratorates(case_id: int, body: CriminalProcuratorateMaintenanceInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    record = await _criminal_detail_maintenance_case(case_id, identity, db)
    return await _save_criminal_detail(record, {k: str(v or "").strip() for k,v in body.model_dump(exclude={"comment"}).items()}, "修改刑事案件检察院信息", body.comment, identity, db)


@app.put(f"{settings.api_prefix}/cases/{{case_id}}/criminal/courts")
async def maintain_criminal_courts(case_id: int, body: CriminalCourtMaintenanceInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    record = await _criminal_detail_maintenance_case(case_id, identity, db)
    return await _save_criminal_detail(record, {k: str(v or "").strip() for k,v in body.model_dump(exclude={"comment"}).items()}, "修改刑事案件审级法院信息", body.comment, identity, db)


@app.put(f"{settings.api_prefix}/cases/{{case_id}}/judicial")
async def update_case_judicial(case_id: int, body: CaseJudicialInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    await _require_record_owner_or_manager(case_record, identity, db)
    if case_record.status in {"待归档审核", "已归档"}:
        raise HTTPException(status_code=409, detail="归档中的案件不能修改司法机关信息")
    if str((case_record.data or {}).get("case_creation_step") or "") != "litigants":
        raise HTTPException(status_code=409, detail="请先完成当事人信息")
    case_type = str((case_record.data or {}).get("case_type") or "")
    if case_type == "法律顾问":
        raise HTTPException(status_code=409, detail="法律顾问案件不使用司法机关步骤")
    permission_key = CASE_CREATE_PERMISSION_BY_TYPE.get(case_type)
    if permission_key and identity.get("role") != "admin":
        permission = await _permission_payload(identity.get("role", "user"), db)
        if permission_key not in set(permission.get("menu_keys", [])):
            raise HTTPException(status_code=403, detail="当前角色没有该案件类型的新建权限")
    hearing_time = body.hearing_time.strip()
    if hearing_time and not re.fullmatch(r"(?:[01]\d|2[0-3]):[0-5]\d(?::[0-5]\d)?", hearing_time):
        raise HTTPException(status_code=422, detail="开庭时间格式必须为 HH:MM 或 HH:MM:SS")
    judicial_data: dict[str, object] = {}
    for key, value in body.model_dump().items():
        if isinstance(value, str):
            judicial_data[key] = value.strip()
        elif isinstance(value, date):
            judicial_data[key] = str(value)
        else:
            judicial_data[key] = value
    case_record.description = str(judicial_data.pop("description", ""))
    if case_type == "行政案件及国家赔偿":
        enabled_court_names = [
            str(judicial_data.get("first_court_name") or "").strip() if judicial_data.get("first_court_enabled") else "",
            str(judicial_data.get("second_court_name") or "").strip() if judicial_data.get("second_court_enabled") else "",
            str(judicial_data.get("retrial_court_name") or "").strip() if judicial_data.get("retrial_court_enabled") else "",
            str(judicial_data.get("court") or "").strip(),
        ]
        if not any(enabled_court_names):
            raise HTTPException(status_code=422, detail="行政案件请至少录入一个法院信息")
        for forbidden_key in (key for key in judicial_data if key.startswith(CRIMINAL_JUDICIAL_PREFIXES)):
            if str(judicial_data.get(forbidden_key) or "").strip():
                raise HTTPException(status_code=422, detail="行政案件不能填写公安或检察院信息")
    previous_status = case_record.status
    case_record.status = "待立案审批"
    case_record.data = {
        **(case_record.data or {}),
        **judicial_data,
        "case_creation_step": "completed",
        "case_creation_completed_at": datetime.now().isoformat(timespec="seconds"),
        "case_creation_completed_by": identity["username"],
        "case_creation_approval_status": "待审批",
        "case_creation_submitted_at": datetime.now().isoformat(timespec="seconds"),
        "case_creation_submitted_by": identity["username"],
    }
    nonempty_fields = [
        key for key, value in judicial_data.items()
        if value not in {"", None, False}
    ]
    db.add(WorkflowEvent(
        record_id=case_record.id,
        action="完成司法机关信息",
        from_status=previous_status,
        to_status=case_record.status,
        operator=identity["username"],
        comment=f"保存司法机关字段 {len(nonempty_fields)} 项并提交案件主管审批",
    ))
    await db.commit()
    await db.refresh(case_record)
    return _record_dict(case_record)


@app.post(f"{settings.api_prefix}/cases/{{case_id}}/creation/review")
async def review_case_creation(case_id: int, body: CaseCreationReviewInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或案件主管可以审批新建案件")
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    data = case_record.data or {}
    if data.get("batch_converted"):
        raise HTTPException(status_code=409, detail="线索自动生成案件无需人工立案审批")
    if data.get("case_creation_step") != "completed" or data.get("case_creation_approval_status") != "待审批" or case_record.status != "待立案审批":
        raise HTTPException(status_code=409, detail="该案件不在待立案审批状态")
    if not body.approved and not body.comment.strip():
        raise HTTPException(status_code=422, detail="驳回时必须填写原因")
    previous = case_record.status
    if body.approved:
        case_record.status = "新案待分配"
        approval_status = "已通过"
    else:
        case_record.status = "新案待分配"
        approval_status = "已驳回"
    case_record.data = {
        **data, "case_creation_approval_status": approval_status, "business_stage": "立案",
        "case_creation_reviewer": identity["username"], "case_creation_reviewed_at": datetime.now().isoformat(timespec="seconds"),
        "case_creation_review_comment": body.comment.strip(),
        **({"case_creation_step": "litigants"} if not body.approved else {}),
    }
    action = "案件创建审批通过" if body.approved else "案件创建审批驳回"
    db.add(WorkflowEvent(record_id=case_record.id, action=action, from_status=previous, to_status=case_record.status, operator=identity["username"], comment=body.comment))
    if body.approved:
        await _ensure_case_fixed_tasks(case_record, db, operator="system")
    await db.commit(); await db.refresh(case_record)
    return _record_dict(case_record)


def _require_case_creation_completed(case_record: BusinessRecord) -> None:
    """Prevent a newly created case from entering downstream workflows before all three steps are saved."""
    creation_step = str((case_record.data or {}).get("case_creation_step") or "")
    if creation_step and creation_step != "completed":
        raise HTTPException(status_code=409, detail="请先完成案件新建三步信息")
    approval_status = str((case_record.data or {}).get("case_creation_approval_status") or "")
    if approval_status and approval_status not in {"已通过", "自动通过"}:
        raise HTTPException(status_code=409, detail="案件创建尚未通过案件主管审批")


async def _require_case_detail_write_access(case_record: BusinessRecord, identity: dict, db: AsyncSession) -> None:
    """Apply one non-bypassable gate to every mutable case-detail feature."""
    if await _case_team_role(case_record, identity, db) == "none":
        raise HTTPException(status_code=403, detail="只有案件负责人、部门负责人、受派经办律师、律师助理或系统管理员可以办理案件详情")
    _require_case_creation_completed(case_record)
    if case_record.status in {"待归档审核", "已归档"}:
        raise HTTPException(status_code=409, detail="案件已进入归档流程，不能新增、删除或修改案件详情资料")


async def _require_case_progress_write_access(case_record: BusinessRecord, identity: dict, db: AsyncSession) -> None:
    """Only a responsible/manager user or assigned handling lawyer may advance a case."""
    role = await _case_team_role(case_record, identity, db)
    if role not in {"manager", "handling_lawyer"}:
        raise HTTPException(status_code=403, detail="只有案件负责人、部门负责人、受派经办律师或系统管理员可以维护案件进展和开庭排期")
    _require_case_creation_completed(case_record)
    if case_record.status in {"待归档审核", "已归档"}:
        raise HTTPException(status_code=409, detail="案件已进入归档流程，不能维护进展或开庭排期")


async def _case_detail_action_capabilities(case_record: BusinessRecord, identity: dict, db: AsyncSession) -> dict:
    role = await _case_team_role(case_record, identity, db)
    base = {
        "can_write": False, "can_upload_attachment": False,
        "can_delete_attachment": False, "can_create_reminder": False,
        "can_delete_reminder": False, "can_create_log": False,
        "can_update_progress": False, "can_manage_hearing": False,
        "can_create_case_task": False,
        "can_assign_team": role == "manager", "can_edit_basic": role == "manager",
        "can_close_case": role == "manager", "can_archive": role == "manager",
        "can_create_finance": role == "manager", "team_role": role, "reason": "",
    }
    try:
        await _require_case_detail_write_access(case_record, identity, db)
    except HTTPException as exc:
        return {**base, "reason": str(exc.detail)}
    can_progress = role in {"manager", "handling_lawyer"}
    return {
        **base,
        "can_write": True, "can_upload_attachment": True,
        "can_delete_attachment": True, "can_create_reminder": True,
        "can_delete_reminder": True, "can_create_log": True,
        "can_update_progress": can_progress, "can_manage_hearing": can_progress,
        "can_create_case_task": can_progress,
    }


@app.get(f"{settings.api_prefix}/cases/{{case_id}}/action-capabilities")
async def case_detail_action_capabilities(case_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    return {"case_id": case_record.id, **(await _case_detail_action_capabilities(case_record, identity, db))}


@app.post(f"{settings.api_prefix}/cases/{{case_id}}/assign")
async def assign_case(case_id: int, body: CaseAssignmentInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    await _require_record_owner_or_manager(case_record, identity, db)
    _require_case_creation_completed(case_record)
    previous = case_record.status
    handling_lawyers, handling_usernames = await _resolve_active_case_people(body.handling_lawyers, db, field_name="经办律师")
    if not handling_lawyers:
        raise HTTPException(status_code=422, detail="请至少分配一名有效经办律师")
    assistant_values, assistant_usernames = await _resolve_active_case_people([body.assistant] if body.assistant.strip() else [], db, field_name="律师助理")
    hearing_values, _ = await _resolve_active_case_people([body.hearing_lawyer], db, field_name="开庭律师")
    manager_values, _ = await _resolve_active_case_people([body.customer_manager] if body.customer_manager.strip() else [], db, field_name="客户管理人")
    assistant = assistant_values[0] if assistant_values else ""
    assistant_username = assistant_usernames[0] if assistant_usernames else ""
    case_data = _case_team_payload({
        **(case_record.data or {}), "customer_manager": manager_values[0] if manager_values else "",
        "hearing_lawyer": hearing_values[0],
    }, handling_lawyers, handling_usernames, assistant, assistant_username)
    case_record.data = case_data
    if case_record.status == "新案待分配":
        case_record.status = "文书准备"
    db.add(WorkflowEvent(record_id=case_record.id, action="案件人员分配", from_status=previous, to_status=case_record.status, operator=identity["username"], comment=f"开庭律师：{case_data['hearing_lawyer']}；经办律师：{','.join(handling_lawyers)}；助理：{assistant}。{body.comment}"))
    notary_id = int(case_data.get("notary_id") or 0)
    if notary_id and not case_data.get("notary_handoff_task_id"):
        notary = await db.get(BusinessRecord, notary_id)
        if notary:
            notary_data = notary.data or {}; scanner = str(notary_data.get("scan_uploaded_by") or notary.owner or identity["username"]).strip(); recipient = (assistant_username or next(iter(handling_usernames), "") or case_data["hearing_lawyer"])
            task = BusinessRecord(module="task", serial_no=f"RW{datetime.now():%Y%m%d%H%M%S%f}", title=f"公证书及公证费发票原件交接—{case_record.serial_no}", customer=case_record.customer, status="待接收", owner=scanner, department=case_record.department, description=f"扫描文员向案件文书人员 {recipient} 交接公证书及公证费发票原件", data={"deadline": str(date.today() + timedelta(days=5)), "priority": "紧急", "source": "自动任务", "initiator": recipient, "collaborators": [recipient] if recipient != scanner else [], "case_no": case_record.serial_no, "case_id": case_record.id, "notary_id": notary.id, "notary_no": notary.serial_no, "auto_task_type": "notary_original_handoff", "handoff_recipient": recipient, "system_created_by": identity["username"]})
            db.add(task); await db.flush(); case_record.data = {**case_record.data, "notary_handoff_task_id": task.id}; notary.data = {**notary_data, "handoff_task_id": task.id, "handoff_recipient": recipient}
            await _add_task_message_notifications(task, WorkflowEvent(record_id=task.id, action="系统生成原件交接任务", to_status="待接收", operator="system", comment=f"扫描文员 {scanner} 向 {recipient} 交接；来源案件 {case_record.serial_no}"), db, content="任务已分派.")
            db.add(WorkflowEvent(record_id=case_record.id, action="生成公证原件交接任务", from_status=case_record.status, to_status=case_record.status, operator="system", comment=f"任务 {task.serial_no}；负责人 {scanner}；接收人 {recipient}"))
    await db.commit()
    await db.refresh(case_record)
    return _record_dict(case_record)


@app.get(f"{settings.api_prefix}/cases/{{case_id}}/tasks")
async def list_case_tasks(case_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    task_rows = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "task").order_by(BusinessRecord.created_at.desc(), BusinessRecord.id.desc()))).all()
    items = [item for item in task_rows if _record_links_to_case(item, case_record)]
    return {"case": _record_dict(case_record), "items": [_task_dict(item) for item in items], "total": len(items)}


@app.post(f"{settings.api_prefix}/cases/{{case_id}}/progress")
async def update_case_progress(case_id: int, body: CaseProgressInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _ensure_record_module(case_id, "case", identity, db); await _require_case_progress_write_access(case_record, identity, db)
    if case_record.status in {"等待公证书", "等待审核公证书", "待归档审核", "已归档"}: raise HTTPException(status_code=409, detail="当前案件阶段不能登记诉讼进展")
    values = body.model_dump(); values["judgment_date"] = str(body.judgment_date) if body.judgment_date else ""
    if not any(values.get(key) for key in ["first_instance_court", "first_instance_case_no", "courtroom", "judge", "clerk", "judgment_date", "judgment_document_no", "second_instance_court", "second_instance_case_no"]): raise HTTPException(status_code=422, detail="请至少填写一项案件进展信息")
    previous = case_record.status; target = previous
    if body.second_instance_case_no.strip(): target = "二审"
    elif body.judgment_date or body.judgment_document_no.strip(): target = "待上诉"
    elif body.first_instance_case_no.strip(): target = "一审立案受理"
    stage_rank = {"新案待分配": 0, "文书准备": 1, "一审立案受理": 2, "一审准备开庭": 3, "待上诉": 4, "二审": 5, "执行": 6}
    if stage_rank.get(target, -1) < stage_rank.get(previous, -1): target = previous
    canonical_stage = "判决" if (body.judgment_date or body.judgment_document_no.strip()) else "审理" if body.second_instance_case_no.strip() else "立案"
    case_record.status = target; case_record.data = {**(case_record.data or {}), **{key: value.strip() if isinstance(value, str) else value for key, value in values.items() if key != "comment"}, "business_stage": canonical_stage}
    db.add(WorkflowEvent(record_id=case_record.id, action="登记案件诉讼进展", from_status=previous, to_status=target, operator=identity["username"], comment=body.comment or "根据法院案号、裁判日期等案件要素自动推进阶段"))
    await db.commit(); await db.refresh(case_record); return _record_dict(case_record)


@app.get(f"{settings.api_prefix}/hearings")
async def list_hearings(
    date_from: date | None = None, date_to: date | None = None,
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    conditions = []
    if date_from:
        conditions.append(HearingSchedule.hearing_date >= date_from)
    if date_to:
        conditions.append(HearingSchedule.hearing_date <= date_to)
    schedules = (await db.scalars(select(HearingSchedule).where(*conditions).order_by(HearingSchedule.hearing_date, HearingSchedule.hearing_time))).all()
    case_ids = {item.case_record_id for item in schedules}
    cases = {item.id: item for item in (await db.scalars(select(BusinessRecord).where(BusinessRecord.id.in_(case_ids), *(await _record_scope_conditions(identity, db))))).all()} if case_ids else {}
    items = [_hearing_dict(item, cases[item.case_record_id]) for item in schedules if item.case_record_id in cases]
    return {"items": items, "total": len(items)}


@app.post(f"{settings.api_prefix}/hearings", status_code=status.HTTP_201_CREATED)
async def create_hearing(body: HearingInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _ensure_record_module(body.case_record_id, "case", identity, db)
    await _require_case_progress_write_access(case_record, identity, db)
    item = HearingSchedule(**body.model_dump(), status="已排期")
    db.add(item)
    await db.flush()
    previous_status = case_record.status
    case_record.data = {**(case_record.data or {}), "court": body.court, "next_hearing_date": str(body.hearing_date), "next_hearing_time": body.hearing_time, "hearing_lawyer": body.hearing_lawyer, "business_stage": "审理"}
    if body.hearing_type.startswith("二审"):
        case_record.status = "二审"
    elif case_record.status in {"新案待分配", "文书准备", "一审立案受理"}:
        case_record.status = "一审准备开庭"
    db.add(WorkflowEvent(record_id=case_record.id, action="新增开庭排期并推进阶段", from_status=previous_status, to_status=case_record.status, operator=identity["username"], comment=f"{body.hearing_date} {body.hearing_time} {body.court} {body.courtroom}"))
    await db.commit()
    await db.refresh(item)
    return _hearing_dict(item, case_record)


@app.delete(f"{settings.api_prefix}/hearings/{{hearing_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_hearing(hearing_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity["role"] != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可删除排期")
    item = await db.get(HearingSchedule, hearing_id)
    if not item:
        raise HTTPException(status_code=404, detail="排期不存在")
    await db.delete(item)
    await db.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.get(f"{settings.api_prefix}/hearing-sms/outbox")
async def hearing_sms_outbox(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以查看短信发送记录")
    await _apply_hearing_sms_reminders(db)
    items = list((await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "sms", *(await _record_scope_conditions(identity, db))).order_by(BusinessRecord.created_at.desc(), BusinessRecord.id.desc()).limit(200))).all())
    return {"items": [_record_dict(item) for item in items], "total": len(items), "provider_configured": bool(settings.sms_webhook_url)}


@app.get(f"{settings.api_prefix}/cases/{{case_id}}/archive-readiness")
async def archive_readiness(case_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    data = case_record.data or {}
    checks = await _case_archive_checks(case_record, db)
    await db.commit()
    return {"case_id": case_id, "case_no": case_record.serial_no, "status": case_record.status, "checks": checks, "archive_no": data.get("archive_no", ""), "paper_archive_location": data.get("paper_archive_location", ""), "paper_volume_count": data.get("paper_volume_count", 1), "archive_reject_reason": data.get("archive_reject_reason", ""), "ready": all(checks.values())}


@app.post(f"{settings.api_prefix}/cases/{{case_id}}/close")
async def close_case_for_archive(case_id: int, body: TaskActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    await _require_record_owner_or_manager(case_record, identity, db)
    _require_case_creation_completed(case_record)
    if case_record.status in {"待归档审核", "已归档"}: raise HTTPException(status_code=409, detail="归档审核中或已归档案件不能重复办结")
    if (case_record.data or {}).get("case_closed_at"): raise HTTPException(status_code=409, detail="案件已经办理办结确认")
    tasks = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "task"))).all()
    active_tasks = [
        item for item in tasks
        if _record_links_to_case(item, case_record)
        and item.status not in {"已完成", "已验收", "已停止", "已撤回", "已拒绝"}
    ]
    if active_tasks: raise HTTPException(status_code=409, detail=f"仍有 {len(active_tasks)} 项案件任务未办结，不能确认案件办结")
    now = datetime.now(timezone.utc)
    case_record.data = {**(case_record.data or {}), "case_closed": True, "case_closed_at": now.isoformat(), "case_closed_by": identity["username"], "case_close_comment": body.comment.strip(), "business_stage": "结案"}
    db.add(WorkflowEvent(record_id=case_record.id, action="确认案件办结", from_status=case_record.status, to_status=case_record.status, operator=identity["username"], comment=body.comment.strip()))
    await db.commit(); await db.refresh(case_record)
    return await _record_dict_for_identity(case_record, identity, db)


@app.post(f"{settings.api_prefix}/cases/{{case_id}}/archive")
async def archive_case(case_id: int, body: ArchiveCheckInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    await _require_record_owner_or_manager(case_record, identity, db)
    _require_case_creation_completed(case_record)
    if case_record.status == "已归档": raise HTTPException(status_code=409, detail="案件已经归档")
    if case_record.status == "待归档审核" and body.submit: raise HTTPException(status_code=409, detail="案件已提交归档审核，请等待审核")
    checks = await _case_archive_checks(case_record, db)
    details = {"archive_no": body.archive_no.strip(), "paper_archive_location": body.paper_archive_location.strip(), "paper_volume_count": body.paper_volume_count}
    case_record.data = {**(case_record.data or {}), **checks, **details}
    if body.submit and not all(checks.values()):
        names = {"case_closed": "案件尚未办结", "fees_settled": "费用尚未结清", "documents_complete": "归档材料不完整", "finance_complete": "财务流程未完成"}
        missing = [names[key] for key, value in checks.items() if not value]
        raise HTTPException(status_code=409, detail="；".join(missing))
    if body.submit and not details["archive_no"]: raise HTTPException(status_code=422, detail="请填写案件归档号")
    if body.submit and not details["paper_archive_location"]: raise HTTPException(status_code=422, detail="请填写纸质卷宗存放位置")
    previous = case_record.status
    action = "保存归档检查"
    if body.submit:
        case_record.data = {**(case_record.data or {}), "status_before_archive": previous, "archive_submitted_at": datetime.now().isoformat(timespec="seconds"), "archive_submitter": identity["username"], "archive_reject_reason": ""}
        case_record.status = "待归档审核"
        action = "提交归档审核"
    db.add(WorkflowEvent(record_id=case_record.id, action=action, from_status=previous, to_status=case_record.status, operator=identity["username"], comment=body.comment or (f"归档号：{details['archive_no']}；纸质卷宗：{details['paper_archive_location']}，{details['paper_volume_count']} 卷" if body.submit else "更新归档检查项")))
    await db.commit()
    await db.refresh(case_record)
    return {"record": _record_dict(case_record), "checks": checks, "ready": all(checks.values())}


@app.post(f"{settings.api_prefix}/cases/{{case_id}}/archive/review")
async def review_case_archive(case_id: int, body: ArchiveReviewInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以审核归档")
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    if case_record.status != "待归档审核": raise HTTPException(status_code=409, detail="只有待归档审核案件可以审核")
    if body.approved:
        checks = await _case_archive_checks(case_record, db)
        if not all(checks.values()):
            raise HTTPException(status_code=409, detail="归档条件在审核前发生变化，请退回补齐后重新提交")
    data = case_record.data or {}; previous = case_record.status
    if body.approved:
        case_record.status = "已归档"
        case_record.data = {**data, "archived_at": datetime.now().isoformat(timespec="seconds"), "archive_reviewer": identity["username"], "archive_review_comment": body.comment, "archive_reject_reason": ""}
        action = "归档审核通过"
    else:
        restored_status = str(data.get("status_before_archive") or "执行")
        if restored_status in {"待归档审核", "已归档"}: restored_status = "执行"
        case_record.status = restored_status
        case_record.data = {**data, "archive_reviewer": identity["username"], "archive_reviewed_at": datetime.now().isoformat(timespec="seconds"), "archive_reject_reason": body.comment}
        action = "归档审核驳回"
    db.add(WorkflowEvent(record_id=case_record.id, action=action, from_status=previous, to_status=case_record.status, operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(case_record)
    return _record_dict(case_record)


@app.post(f"{settings.api_prefix}/cases/{{case_id}}/unarchive/request")
async def request_case_unarchive(case_id: int, body: CaseUnarchiveRequestInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    await _require_record_owner_or_manager(case_record, identity, db)
    if case_record.status != "已归档":
        raise HTTPException(status_code=409, detail="只有已归档案件可以申请解档")
    data = case_record.data or {}; pending = data.get("unarchive_request") or {}
    if pending.get("status") == "待审批":
        raise HTTPException(status_code=409, detail="该案件已有解档申请正在审批")
    request_data = {
        "status": "待审批", "reason": body.reason.strip(), "requested_by": identity["username"],
        "requested_at": datetime.now().isoformat(timespec="seconds"),
    }
    case_record.data = {**data, "unarchive_request": request_data}
    db.add(WorkflowEvent(record_id=case_record.id, action="提交解档申请", from_status="已归档", to_status="已归档", operator=identity["username"], comment=body.reason.strip()))
    await db.commit(); await db.refresh(case_record)
    return _record_dict(case_record)


@app.post(f"{settings.api_prefix}/cases/{{case_id}}/unarchive/review")
async def review_case_unarchive(case_id: int, body: CaseUnarchiveReviewInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以审批解档")
    case_record = await _ensure_record_module(case_id, "case", identity, db)
    data = case_record.data or {}; pending = data.get("unarchive_request") or {}
    if case_record.status != "已归档" or pending.get("status") != "待审批":
        raise HTTPException(status_code=409, detail="该案件没有待审批的解档申请")
    if pending.get("requested_by") == identity["username"] and identity.get("role") != "admin":
        raise HTTPException(status_code=403, detail="解档申请人不能审批自己的申请")
    if not body.approved and not body.comment.strip():
        raise HTTPException(status_code=422, detail="驳回时必须填写原因")
    previous = case_record.status
    reviewed = {
        **pending, "status": "已通过" if body.approved else "已驳回", "reviewed_by": identity["username"],
        "reviewed_at": datetime.now().isoformat(timespec="seconds"), "review_comment": body.comment.strip(),
    }
    if body.approved:
        restored_status = str(data.get("status_before_archive") or "执行")
        if restored_status in {"已归档", "待归档审核", "待立案审批"}: restored_status = "执行"
        case_record.status = restored_status
        case_record.data = {**data, "unarchive_request": reviewed, "unarchived_at": datetime.now().isoformat(timespec="seconds"), "unarchived_by": identity["username"], "archive_locked": False}
        action = "解档审批通过"
    else:
        case_record.data = {**data, "unarchive_request": reviewed}
        action = "解档审批驳回"
    db.add(WorkflowEvent(record_id=case_record.id, action=action, from_status=previous, to_status=case_record.status, operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(case_record)
    return _record_dict(case_record)


@app.get(f"{settings.api_prefix}/documents/summary")
async def document_summary(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    documents = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "document", *(await _record_scope_conditions(identity, db))))).all()
    attachments = (await db.scalars(select(FileAttachment))).all()
    attachments = await _filter_visible_attachments(attachments, identity, db)
    templates = (await db.scalars(select(DocumentTemplate))).all()
    return {
        "documents": len(documents),
        "pending_receipt": sum(1 for item in documents if item.status in {"待登记", "待签收"}),
        "received": sum(1 for item in documents if item.status == "已签收"),
        "attachments": len(attachments),
        "archive_materials": sum(1 for item in attachments if item.category in ARCHIVE_REQUIRED_CATEGORIES),
        "templates": sum(1 for item in templates if item.is_active),
    }


@app.post(f"{settings.api_prefix}/documents/{{document_id}}/transition")
async def transition_document(document_id: int, body: DocumentTransitionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _ensure_record_visible(document_id, identity, db)
    if item.module != "document": raise HTTPException(status_code=404, detail="收发文记录不存在")
    await _require_record_owner_or_manager(item, identity, db)
    allowed = WORKFLOW_TRANSITIONS["document"].get(item.status, [])
    if body.to_status not in allowed:
        raise HTTPException(status_code=409, detail=f"不能从“{item.status}”流转到“{body.to_status}”")
    if body.action_date > date.today(): raise HTTPException(status_code=422, detail="办理日期不能晚于今天")
    handler = body.handler.strip()
    archive_no = body.archive_no.strip()
    if body.to_status == "已签收" and not handler:
        raise HTTPException(status_code=422, detail="确认签收时必须填写签收人")
    if body.to_status == "已归档" and not archive_no:
        raise HTTPException(status_code=422, detail="归档时必须填写归档编号")
    previous = item.status
    data = dict(item.data or {})
    if body.to_status == "待签收":
        data.update({"registered_at": str(body.action_date), "register_operator": identity["username"]})
        action = "完成登记"
    elif body.to_status == "已签收":
        data.update({"signed_at": str(body.action_date), "signer": handler})
        action = "确认签收" if data.get("direction") == "收文" else "确认送达"
    else:
        data.update({"archived_at": str(body.action_date), "archive_no": archive_no, "archive_location": body.archive_location.strip()})
        action = "文档归档"
    item.data = data
    item.status = body.to_status
    db.add(WorkflowEvent(record_id=item.id, action=action, from_status=previous, to_status=item.status, operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(item)
    return _record_dict(item, await _allowed_field_keys(identity, db))


@app.post(f"{settings.api_prefix}/documents/official/process")
async def process_official_documents(body: OfficialDocumentProcessInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Mark selected official incoming documents as business processed/unprocessed.

    This status is intentionally separate from registration, signing and archive
    lifecycle.  It mirrors the legacy Patent Office batch actions without
    allowing a generic record update to bypass document lifecycle controls.
    """
    record_ids = list(dict.fromkeys(body.record_ids))
    if not record_ids:
        raise HTTPException(status_code=422, detail="请选择至少一条官文收文记录")
    target_status = "已处理" if body.processed else "未处理"
    action = "标记官文已处理" if body.processed else "标记官文未处理"
    changed: list[BusinessRecord] = []
    for record_id in record_ids:
        item = await _ensure_record_visible(record_id, identity, db)
        if item.module != "document" or (item.data or {}).get("direction", "收文") != "收文":
            raise HTTPException(status_code=422, detail="所选记录不是官文收文")
        await _require_record_owner_or_manager(item, identity, db)
        data = dict(item.data or {})
        previous = data.get("business_process_status", "未处理")
        if previous == target_status:
            continue
        data.update({
            "business_process_status": target_status,
            "business_processed_at": datetime.now().isoformat(timespec="seconds"),
            "business_processed_by": identity["username"],
        })
        item.data = data
        # Do not change item.status: document registration/sign/archive has its
        # own dedicated state machine and must remain independently auditable.
        db.add(WorkflowEvent(
            record_id=item.id,
            action=action,
            from_status=item.status,
            to_status=item.status,
            operator=identity["username"],
            comment=body.comment.strip(),
        ))
        changed.append(item)
    await db.commit()
    for item in changed:
        await db.refresh(item)
    return {"processed": len(changed), "business_process_status": target_status, "items": [_record_dict(item, await _allowed_field_keys(identity, db)) for item in changed]}


@app.post(f"{settings.api_prefix}/documents/official/receipt-date")
async def update_official_receipt_date(body: OfficialDocumentReceiptDateInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Correct the incoming-document date without using the generic record API.

    Legacy FIO supported this as a selected-row batch action.  It is metadata,
    not a document lifecycle transition, so registration/sign/archive status is
    deliberately left unchanged while every corrected document receives an
    independent audit event.
    """
    record_ids = list(dict.fromkeys(body.record_ids))
    if not record_ids:
        raise HTTPException(status_code=422, detail="请选择至少一条官文收文记录")
    changed: list[BusinessRecord] = []
    for record_id in record_ids:
        item = await _ensure_record_visible(record_id, identity, db)
        if item.module != "document" or (item.data or {}).get("direction", "收文") != "收文":
            raise HTTPException(status_code=422, detail="所选记录不是官文收文")
        await _require_record_owner_or_manager(item, identity, db)
        data = dict(item.data or {})
        previous_date = str(data.get("document_date") or data.get("received_at") or "")
        target_date = str(body.document_date)
        if previous_date == target_date:
            continue
        data.update({"document_date": target_date, "received_at": target_date})
        item.data = data
        db.add(WorkflowEvent(
            record_id=item.id,
            action="修改官文收文日期",
            from_status=item.status,
            to_status=item.status,
            operator=identity["username"],
            comment=body.comment.strip() or f"{previous_date or '未填写'} → {target_date}",
        ))
        changed.append(item)
    await db.commit()
    for item in changed:
        await db.refresh(item)
    return {"updated": len(changed), "document_date": str(body.document_date), "items": [_record_dict(item, await _allowed_field_keys(identity, db)) for item in changed]}


@app.get(f"{settings.api_prefix}/templates")
async def list_templates(category: str = "", _: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    query = select(DocumentTemplate).order_by(DocumentTemplate.category, DocumentTemplate.name)
    if category:
        query = query.where(DocumentTemplate.category == category)
    items = (await db.scalars(query)).all()
    return {"items": [_template_dict(item) for item in items], "total": len(items)}


@app.post(f"{settings.api_prefix}/templates", status_code=status.HTTP_201_CREATED)
async def create_template(body: TemplateInput, _: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if await db.scalar(select(DocumentTemplate.id).where(DocumentTemplate.name == body.name)):
        raise HTTPException(status_code=409, detail="模板名称已存在")
    item = DocumentTemplate(**body.model_dump())
    db.add(item)
    await db.commit()
    await db.refresh(item)
    return _template_dict(item)


@app.patch(f"{settings.api_prefix}/templates/{{template_id}}")
async def update_template(template_id: int, body: TemplateUpdate, _: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await db.get(DocumentTemplate, template_id)
    if not item:
        raise HTTPException(status_code=404, detail="模板不存在")
    changes = body.model_dump(exclude_unset=True)
    if not changes: return _template_dict(item)
    if "name" in changes:
        changes["name"] = str(changes["name"] or "").strip()
        if not changes["name"]: raise HTTPException(status_code=422, detail="模板名称不能为空")
        duplicate = await db.scalar(select(DocumentTemplate.id).where(DocumentTemplate.name == changes["name"], DocumentTemplate.id != template_id))
        if duplicate: raise HTTPException(status_code=409, detail="模板名称已存在")
    for key, value in changes.items():
        setattr(item, key, value)
    await db.commit()
    await db.refresh(item)
    return _template_dict(item)


@app.delete(f"{settings.api_prefix}/templates/{{template_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_template(template_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity["role"] != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可删除模板")
    item = await db.get(DocumentTemplate, template_id)
    if not item:
        raise HTTPException(status_code=404, detail="模板不存在")
    await db.delete(item)
    await db.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.get(f"{settings.api_prefix}/attachments")
async def list_attachments(record_id: int | None = None, finance_transaction_id: int | None = None, category: str = "", identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    record = None
    if record_id is not None:
        record = await _ensure_attachment_record_visible(record_id, identity, db)
    conditions = []
    if record_id is not None:
        conditions.append(FileAttachment.record_id == record_id)
    if finance_transaction_id is not None:
        conditions.append(FileAttachment.finance_transaction_id == finance_transaction_id)
    if category:
        conditions.append(FileAttachment.category == category)
    items = (await db.scalars(select(FileAttachment).where(*conditions).order_by(FileAttachment.created_at.desc(), FileAttachment.id.desc()))).all()
    if not (record and record.module == "task"):
        items = await _filter_visible_attachments(items, identity, db)
    record_ids = {item.record_id for item in items if item.record_id}
    records = {record.id: record for record in (await db.scalars(select(BusinessRecord).where(BusinessRecord.id.in_(record_ids)))).all()} if record_ids else {}
    return {"items": [_attachment_dict(item, records.get(item.record_id)) for item in items], "total": len(items), "required_archive_categories": sorted(ARCHIVE_REQUIRED_CATEGORIES)}


async def _sync_seal_document_names(record: BusinessRecord, db: AsyncSession) -> None:
    """Keep the legacy file-name projection derived from real seal attachments."""
    files = (await db.scalars(
        select(FileAttachment.original_name)
        .where(FileAttachment.record_id == record.id, FileAttachment.category == "用印文件")
        .order_by(FileAttachment.created_at, FileAttachment.id)
    )).all()
    record.data = {**(record.data or {}), "document_names": "、".join(files)}


@app.post(f"{settings.api_prefix}/cases/attachments/download")
async def download_case_attachments(body: AttachmentBatchInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    attachment_ids = list(dict.fromkeys(body.attachment_ids))
    attachments = list((await db.scalars(select(FileAttachment).where(FileAttachment.id.in_(attachment_ids)))).all())
    if len(attachments) != len(attachment_ids):
        raise HTTPException(status_code=404, detail="存在已删除或不存在的案件文件")
    ordered = sorted(attachments, key=lambda item: attachment_ids.index(item.id))
    paths: list[tuple[FileAttachment, Path]] = []
    for item in ordered:
        if not item.record_id:
            raise HTTPException(status_code=422, detail="所选文件不是案件文件")
        record = await _ensure_record_module(item.record_id, "case", identity, db)
        path = Path(item.path)
        if not path.is_file() or UPLOAD_ROOT.resolve() not in path.resolve().parents:
            raise HTTPException(status_code=404, detail=f"文件 {item.original_name} 的实体不存在")
        paths.append((item, path))
    output = io.BytesIO()
    used_names: set[str] = set()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for item, path in paths:
            filename = Path(item.original_name).name
            if filename in used_names:
                filename = f"{item.id}-{filename}"
            used_names.add(filename)
            archive.write(path, arcname=filename)
    output.seek(0)
    return StreamingResponse(
        output, media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="case-files-{date.today():%Y%m%d}.zip"'},
    )


@app.post(f"{settings.api_prefix}/cases/attachments/delete")
async def delete_case_attachments(body: AttachmentBatchInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    attachment_ids = list(dict.fromkeys(body.attachment_ids))
    attachments = list((await db.scalars(select(FileAttachment).where(FileAttachment.id.in_(attachment_ids)))).all())
    if len(attachments) != len(attachment_ids):
        raise HTTPException(status_code=404, detail="存在已删除或不存在的案件文件")
    prepared: list[tuple[FileAttachment, BusinessRecord, Path]] = []
    for item in attachments:
        if not item.record_id:
            raise HTTPException(status_code=422, detail="所选文件不是案件文件")
        record = await _ensure_record_module(item.record_id, "case", identity, db)
        await _require_case_detail_write_access(record, identity, db)
        prepared.append((item, record, Path(item.path)))
    affected_cases: dict[int, BusinessRecord] = {}
    for item, record, path in prepared:
        affected_cases[record.id] = record
        await db.delete(item)
        db.add(WorkflowEvent(
            record_id=record.id, action="批量删除案件文件", from_status=record.status,
            to_status=record.status, operator=identity["username"],
            comment=f"{item.category}：{item.original_name}",
        ))
    await db.flush()
    for record in affected_cases.values():
        await _sync_case_document_readiness(record, db)
    await db.commit()
    for _, _, path in prepared:
        if path.is_file() and UPLOAD_ROOT.resolve() in path.resolve().parents:
            path.unlink()
    return {"deleted": len(prepared)}


@app.put(f"{settings.api_prefix}/cases/attachments/{{attachment_id}}/rename")
async def rename_case_attachment(attachment_id: int, body: CaseAttachmentRenameInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Rename only the display/download name of one case attachment; never move its stored file."""
    item = await db.get(FileAttachment, attachment_id)
    if not item or not item.record_id:
        raise HTTPException(status_code=404, detail="案件文件不存在")
    case_record = await _ensure_record_module(item.record_id, "case", identity, db)
    await _require_case_detail_write_access(case_record, identity, db)
    requested_name = body.original_name.strip()
    if not requested_name or "/" in requested_name or "\\" in requested_name or Path(requested_name).name != requested_name or requested_name in {".", ".."}:
        raise HTTPException(status_code=422, detail="文件名不能为空，且不能包含路径")
    if Path(requested_name).suffix.lower() != Path(item.original_name).suffix.lower():
        raise HTTPException(status_code=422, detail="重命名不能修改文件扩展名")
    previous_name = item.original_name
    if requested_name == previous_name:
        return _attachment_dict(item, case_record)
    item.original_name = requested_name
    db.add(WorkflowEvent(
        record_id=case_record.id, action="重命名案件文件", from_status=case_record.status,
        to_status=case_record.status, operator=identity["username"],
        comment=f"{item.category}：{previous_name} → {requested_name}",
    ))
    await db.commit()
    await db.refresh(item)
    return _attachment_dict(item, case_record)


@app.post(f"{settings.api_prefix}/documents/official/upload", status_code=status.HTTP_201_CREATED)
async def upload_official_document(
    file: UploadFile = File(...), category: str = Form("收文附件"), remark: str = Form(""),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    suffix = Path(file.filename or "").suffix.lower()
    allowed = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx", ".txt", ".png", ".jpg", ".jpeg", ".zip", ".rar"}
    if suffix not in allowed:
        raise HTTPException(status_code=422, detail="不支持的文件格式")
    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="单个文件不能超过 20MB")
    user = await db.scalar(select(User).where(User.username == identity["username"]))
    if not user:
        raise HTTPException(status_code=401, detail="当前用户不存在")
    original_name = Path(file.filename or f"official{suffix}").name
    stored_name = f"{uuid4().hex}{suffix}"
    target = UPLOAD_ROOT / stored_name
    target.write_bytes(content)
    now = datetime.now()
    record = BusinessRecord(
        module="document", serial_no=f"SW{now:%Y%m%d%H%M%S%f}", title=original_name,
        customer="", status="待签收", owner=identity["username"], department=user.department,
        description=remark,
        data={"direction": "收文", "document_date": str(date.today()), "uploaded_at": str(date.today()), "uploader": identity["username"], "import_status": "已导入", "business_process_status": "未处理"},
    )
    try:
        db.add(record)
        await db.flush()
        attachment = FileAttachment(
            record_id=record.id, category=category or "收文附件", original_name=original_name,
            stored_name=stored_name, content_type=file.content_type or "application/octet-stream",
            size=len(content), path=str(target), uploader=identity["username"], remark=remark,
        )
        db.add(attachment)
        db.add(WorkflowEvent(record_id=record.id, action="上传官文收文", from_status="", to_status="待签收", operator=identity["username"], comment=original_name))
        await db.commit()
        await db.refresh(record)
        await db.refresh(attachment)
    except Exception:
        await db.rollback()
        target.unlink(missing_ok=True)
        raise
    return {"record": _record_dict(record), "attachment": _attachment_dict(attachment, record)}


@app.post(f"{settings.api_prefix}/attachments", status_code=status.HTTP_201_CREATED)
async def upload_attachment(
    file: UploadFile = File(...), record_id: int | None = Form(None),
    finance_transaction_id: int | None = Form(None),
    category: str = Form("普通附件"), remark: str = Form(""),
    identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db),
):
    record = None
    transaction = None
    if finance_transaction_id is not None:
        transaction = await db.get(FinanceTransaction, finance_transaction_id)
        if not transaction:
            raise HTTPException(status_code=404, detail="关联财务流水不存在")
        if record_id is not None and record_id != transaction.finance_record_id:
            raise HTTPException(status_code=409, detail="附件关联费用与财务流水不一致")
        record_id = transaction.finance_record_id
        await _ensure_record_visible(record_id, identity, db)
        expected_category = FINANCE_DEFAULT_VOUCHER_CATEGORY[transaction.transaction_type]
        if category == "普通附件":
            category = expected_category
        if category not in FINANCE_VOUCHER_CATEGORIES:
            raise HTTPException(status_code=422, detail="财务流水附件类型无效")
    if record_id is not None:
        record = await _ensure_attachment_record_visible(record_id, identity, db)
        if record.module == "case":
            await _require_case_detail_write_access(record, identity, db)
        if record.module == "task":
            if not _is_task_participant(record, identity):
                raise HTTPException(status_code=403, detail="只有任务参与人可以上传任务反馈附件")
            if category not in {"任务反馈附件", "任务资料附件"}:
                raise HTTPException(status_code=422, detail="任务附件类型无效")
        if record.module == "customer":
            await _require_record_owner_or_manager(record, identity, db)
        if record.module == "seal":
            await _require_record_owner_or_manager(record, identity, db)
            if record.status != "草稿":
                raise HTTPException(status_code=409, detail="只有草稿用印申请可以上传或替换用印文件")
            if category != "用印文件":
                raise HTTPException(status_code=422, detail="用印申请附件类型必须为用印文件")
        if record.module in INVESTIGATION_MATERIAL_CATEGORIES:
            await _require_record_owner_or_manager(record, identity, db)
            if category == "公证书扫描件":
                operator = await db.scalar(select(User).where(User.username == identity["username"]))
                if not operator or not await _user_has_job_permission(operator, "扫描上传", db):
                    raise HTTPException(status_code=403, detail="当前账号没有公证书扫描上传岗位权限")
        if record.module in INVESTIGATION_MATERIAL_CATEGORIES and category != "普通附件" and category not in INVESTIGATION_MATERIAL_CATEGORIES[record.module]:
            raise HTTPException(status_code=422, detail="材料类型与当前调查业务不匹配")
    suffix = Path(file.filename or "").suffix.lower()
    allowed = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx", ".txt", ".png", ".jpg", ".jpeg", ".zip", ".rar"}
    if suffix not in allowed:
        raise HTTPException(status_code=422, detail="不支持的文件格式")
    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="单个文件不能超过 20MB")
    stored_name = f"{uuid4().hex}{suffix}"
    target = UPLOAD_ROOT / stored_name
    target.write_bytes(content)
    item = FileAttachment(record_id=record_id, finance_transaction_id=finance_transaction_id, category=category, original_name=Path(file.filename or stored_name).name, stored_name=stored_name, content_type=file.content_type or "application/octet-stream", size=len(content), path=str(target), uploader=identity["username"], remark=remark)
    db.add(item)
    await db.flush()
    if record and record.module == "case":
        await _sync_case_document_readiness(record, db)
        db.add(WorkflowEvent(record_id=record.id, action="上传归档材料", from_status=record.status, to_status=record.status, operator=identity["username"], comment=f"{category}：{item.original_name}"))
    elif record and record.module in INVESTIGATION_MATERIAL_CATEGORIES:
        material_categories = await _sync_investigation_materials(record, db)
        db.add(WorkflowEvent(record_id=record.id, action="上传调查材料", from_status=record.status, to_status=record.status, operator=identity["username"], comment=f"{category}：{item.original_name}"))
        if record.module == "notary" and category == "公证书扫描件" and record.status in {"等待材料", "审核驳回"}:
            previous_notary_status = record.status; record.status = "待审核"; record.data = {**(record.data or {}), "review_due_date": str(date.today() + timedelta(days=30)), "scan_uploaded_at": datetime.now().isoformat(timespec="seconds"), "scan_uploaded_by": identity["username"]}
            clue = await db.get(BusinessRecord, int((record.data or {}).get("clue_id") or 0)); case_record = await db.get(BusinessRecord, int(((record.data or {}).get("case_id") or ((clue.data or {}).get("converted_case_id") if clue else 0)) or 0))
            if case_record and case_record.status == "等待公证书":
                case_record.status = "等待审核公证书"; db.add(WorkflowEvent(record_id=case_record.id, action="公证书扫描件已上传", from_status="等待公证书", to_status="等待审核公证书", operator=identity["username"], comment=f"公证记录 {record.serial_no}"))
            db.add(WorkflowEvent(record_id=record.id, action="提交公证书审核", from_status=previous_notary_status, to_status="待审核", operator=identity["username"], comment=f"扫描件 {item.original_name}；审核期限 30 日"))
    elif record and record.module == "customer":
        db.add(WorkflowEvent(record_id=record.id, action="上传客户文档", from_status=record.status, to_status=record.status, operator=identity["username"], comment=f"{category}：{item.original_name}"))
    elif record and record.module == "seal":
        await _sync_seal_document_names(record, db)
        db.add(WorkflowEvent(record_id=record.id, action="上传用印文件", from_status=record.status, to_status=record.status, operator=identity["username"], comment=f"{category}：{item.original_name}"))
    elif record and record.module == "task":
        await _add_task_message_notifications(
            record,
            WorkflowEvent(
                record_id=record.id, action=f"上传{category}",
                from_status=record.status, to_status=record.status,
                operator=identity["username"], comment=f"{category}：{item.original_name}",
            ),
            db,
            content=f"已上传{category}：{item.original_name}",
        )
    elif record and transaction:
        db.add(WorkflowEvent(record_id=record.id, action="上传财务凭证", from_status=record.status, to_status=record.status, operator=identity["username"], comment=f"{transaction.transaction_type}流水 #{transaction.id}｜{category}：{item.original_name}"))
    await db.commit()
    await db.refresh(item)
    return _attachment_dict(item, record)


@app.get(f"{settings.api_prefix}/attachments/{{attachment_id}}/download")
async def download_attachment(attachment_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await db.get(FileAttachment, attachment_id)
    if not item:
        raise HTTPException(status_code=404, detail="附件不存在")
    if item.record_id:
        record = await _ensure_attachment_record_visible(item.record_id, identity, db)
    elif identity.get("role") != "admin" and item.uploader != identity["username"]:
        raise HTTPException(status_code=404, detail="附件不存在或无权访问")
    path = Path(item.path)
    if not path.is_file() or UPLOAD_ROOT.resolve() not in path.resolve().parents:
        raise HTTPException(status_code=404, detail="附件文件不存在")
    return FileResponse(path, media_type=item.content_type, filename=item.original_name)


@app.delete(f"{settings.api_prefix}/attachments/{{attachment_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_attachment(attachment_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await db.get(FileAttachment, attachment_id)
    if not item:
        raise HTTPException(status_code=404, detail="附件不存在")
    record = await db.get(BusinessRecord, item.record_id) if item.record_id else None
    may_manage_customer_document = False
    may_manage_case_document = False
    may_manage_task_attachment = False
    may_manage_seal_attachment = False
    if record and record.module == "task":
        record = await _ensure_attachment_record_visible(record.id, identity, db)
        if item.category not in {"任务反馈附件", "任务资料附件"}:
            raise HTTPException(status_code=422, detail="任务附件类型无效")
        if identity.get("role") != "admin" and item.uploader != identity["username"]:
            raise HTTPException(status_code=403, detail="任务参与人只能删除自己上传的任务附件")
        may_manage_task_attachment = True
    if record and record.module == "case":
        record = await _ensure_record_module(record.id, "case", identity, db)
        await _require_case_detail_write_access(record, identity, db)
        may_manage_case_document = True
    if record and record.module == "customer":
        record = await _ensure_record_module(record.id, "customer", identity, db)
        await _require_record_owner_or_manager(record, identity, db)
        may_manage_customer_document = True
    if record and record.module == "seal":
        record = await _ensure_record_module(record.id, "seal", identity, db)
        await _require_record_owner_or_manager(record, identity, db)
        if record.status != "草稿":
            raise HTTPException(status_code=409, detail="只有草稿用印申请可以删除用印文件")
        if item.category != "用印文件":
            raise HTTPException(status_code=422, detail="用印申请附件类型无效")
        may_manage_seal_attachment = True
    may_manage_hr_document = identity.get("role") == "manager" and record and record.module == "hr" and item.category == "员工档案"
    if identity["role"] != "admin" and not may_manage_hr_document and not may_manage_customer_document and not may_manage_case_document and not may_manage_task_attachment and not may_manage_seal_attachment:
        raise HTTPException(status_code=403, detail="仅管理员可删除附件；客户负责人可删除客户文档，部门负责人可删除员工档案")
    path = Path(item.path)
    await db.delete(item)
    await db.flush()
    if record and record.module == "case":
        await _sync_case_document_readiness(record, db)
        db.add(WorkflowEvent(record_id=record.id, action="删除归档材料", from_status=record.status, to_status=record.status, operator=identity["username"], comment=f"{item.category}：{item.original_name}"))
    elif record and record.module in INVESTIGATION_MATERIAL_CATEGORIES:
        material_categories = await _sync_investigation_materials(record, db)
        db.add(WorkflowEvent(record_id=record.id, action="删除调查材料", from_status=record.status, to_status=record.status, operator=identity["username"], comment=f"{item.category}：{item.original_name}"))
        if record.module == "notary" and item.category == "公证书扫描件" and "公证书扫描件" not in material_categories and record.status == "待审核":
            record.status = "等待材料"; record.data = {**(record.data or {}), "review_due_date": "", "scan_uploaded_at": ""}
            clue = await db.get(BusinessRecord, int((record.data or {}).get("clue_id") or 0)); case_record = await db.get(BusinessRecord, int(((record.data or {}).get("case_id") or ((clue.data or {}).get("converted_case_id") if clue else 0)) or 0))
            if case_record and case_record.status == "等待审核公证书":
                case_record.status = "等待公证书"; db.add(WorkflowEvent(record_id=case_record.id, action="撤回公证书审核", from_status="等待审核公证书", to_status="等待公证书", operator=identity["username"], comment="公证书扫描件已删除"))
            db.add(WorkflowEvent(record_id=record.id, action="撤回公证书审核", from_status="待审核", to_status="等待材料", operator=identity["username"], comment="公证书扫描件已删除"))
    elif record and record.module == "customer":
        db.add(WorkflowEvent(record_id=record.id, action="删除客户文档", from_status=record.status, to_status=record.status, operator=identity["username"], comment=f"{item.category}：{item.original_name}"))
    elif record and record.module == "seal":
        await _sync_seal_document_names(record, db)
        db.add(WorkflowEvent(record_id=record.id, action="删除用印文件", from_status=record.status, to_status=record.status, operator=identity["username"], comment=f"{item.category}：{item.original_name}"))
    elif record and record.module == "task":
        await _add_task_message_notifications(
            record,
            WorkflowEvent(record_id=record.id, action=f"删除{item.category}", from_status=record.status, to_status=record.status, operator=identity["username"], comment=f"{item.category}：{item.original_name}"),
            db,
            content=f"已删除{item.category}：{item.original_name}",
        )
    elif record and item.finance_transaction_id:
        db.add(WorkflowEvent(record_id=record.id, action="删除财务凭证", from_status=record.status, to_status=record.status, operator=identity["username"], comment=f"流水 #{item.finance_transaction_id}｜{item.category}：{item.original_name}"))
    await db.commit()
    if path.is_file() and UPLOAD_ROOT.resolve() in path.resolve().parents:
        path.unlink()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.post(f"{settings.api_prefix}/receivables", status_code=status.HTTP_201_CREATED)
async def create_receivable(body: ReceivableInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    contract = await _ensure_record_module(body.contract_record_id, "contract", identity, db)
    await _require_record_owner_or_manager(contract, identity, db)
    plan = ReceivablePlan(**body.model_dump(), status="待收款")
    db.add(plan)
    await db.flush()
    db.add(WorkflowEvent(record_id=contract.id, action="新增应收计划", from_status=contract.status, to_status=contract.status, operator=identity["username"], comment=f"{body.phase}：{body.amount:.2f}元"))
    await db.commit()
    await db.refresh(plan)
    return _receivable_dict(plan, contract)


@app.post(f"{settings.api_prefix}/receivables/{{plan_id}}/receive")
async def receive_payment(plan_id: int, body: ReceivePaymentInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    plan = await db.get(ReceivablePlan, plan_id)
    if not plan:
        raise HTTPException(status_code=404, detail="应收计划不存在")
    contract = await _ensure_record_module(plan.contract_record_id, "contract", identity, db)
    await _require_record_owner_or_manager(contract, identity, db)
    remaining = max(plan.amount - plan.received_amount, 0)
    if body.amount > remaining + 0.001:
        raise HTTPException(status_code=409, detail=f"登记金额不能超过未收金额 {remaining:.2f} 元")
    plan.received_amount += body.amount
    plan.status = "已收款" if plan.received_amount + 0.001 >= plan.amount else "部分收款"
    db.add(WorkflowEvent(record_id=contract.id, action="登记回款", from_status=contract.status, to_status=contract.status, operator=identity["username"], comment=f"{plan.phase}回款 {body.amount:.2f} 元。{body.comment}"))
    await db.commit()
    await db.refresh(plan)
    return _receivable_dict(plan, contract)


@app.delete(f"{settings.api_prefix}/receivables/{{plan_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_receivable(plan_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity["role"] != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可删除应收计划")
    plan = await db.get(ReceivablePlan, plan_id)
    if not plan:
        raise HTTPException(status_code=404, detail="应收计划不存在")
    await db.delete(plan)
    await db.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


def _seal_asset_dict(item: SealAsset) -> dict:
    return {"id": item.id, "code": item.code, "name": item.name, "seal_type": item.seal_type, "custodian": item.custodian, "location": item.location, "status": item.status, "usage_count": item.usage_count, "last_used_at": item.last_used_at, "remark": item.remark, "created_at": item.created_at, "updated_at": item.updated_at}


async def _seal_record_dict(record: BusinessRecord, db: AsyncSession) -> dict:
    result = _record_dict(record)
    asset_id = int((record.data or {}).get("seal_asset_id") or 0)
    asset = await db.get(SealAsset, asset_id) if asset_id else None
    result["seal_asset"] = _seal_asset_dict(asset) if asset else None
    return result


async def _validated_seal_relations(body: SealApplicationInput, identity: dict, db: AsyncSession) -> tuple[str, str, str, str]:
    """Return canonical, visible seal references and prevent dangling business links."""
    case_no, contract_no, customer = body.case_no.strip(), body.contract_no.strip(), body.customer.strip()
    use_type = body.use_type.strip() or ("案件用印" if case_no else "合同用印" if contract_no else "行政用印")
    scope = await _record_scope_conditions(identity, db)

    async def visible(module: str, serial_no: str, label: str) -> BusinessRecord:
        row = await db.scalar(select(BusinessRecord).where(BusinessRecord.module == module, BusinessRecord.serial_no == serial_no, *scope))
        if not row:
            raise HTTPException(status_code=422, detail=f"关联{label}不存在或当前账号无权使用")
        return row

    if use_type == "案件用印" and not case_no:
        raise HTTPException(status_code=422, detail="案件用印必须选择关联案件")
    if use_type == "合同用印" and not contract_no:
        raise HTTPException(status_code=422, detail="合同用印必须选择关联合同")
    case = await visible("case", case_no, "案件") if case_no else None
    contract = await visible("contract", contract_no, "合同") if contract_no else None
    if customer:
        customer_row = await db.scalar(select(BusinessRecord).where(BusinessRecord.module == "customer", or_(BusinessRecord.title == customer, BusinessRecord.customer == customer, BusinessRecord.serial_no == customer), *scope))
        if not customer_row:
            raise HTTPException(status_code=422, detail="关联客户不存在或当前账号无权使用")
        customer = customer_row.title or customer_row.customer
    elif case:
        customer = case.customer
    elif contract:
        customer = contract.customer
    return case_no, contract_no, customer, use_type


@app.get(f"{settings.api_prefix}/seals/applications")
async def list_seal_applications(view: str = "my", keyword: str = "", record_status: str = "", serial_no: str = "", applicant: str = "", date_from: date | None = None, date_to: date | None = None, case_no: str = "", contract_no: str = "", customer: str = "", use_type: str = "", file_name: str = "", page: int = Query(1, ge=1), page_size: int = Query(20, ge=1, le=100), identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if view not in {"my", "audit", "all"}: raise HTTPException(status_code=422, detail="无效的用印视图")
    if view == "all" and identity.get("role") != "admin": raise HTTPException(status_code=403, detail="只有系统管理员可以查看全部用印申请")
    scope_conditions = await _record_scope_conditions(identity, db)
    conditions = [BusinessRecord.module == "seal", *scope_conditions]
    if view == "my": conditions.append(BusinessRecord.owner == identity["username"])
    if view == "audit": conditions.append(BusinessRecord.status.in_({"待审批", "待用印", "已拒绝"}))
    if record_status: conditions.append(BusinessRecord.status == record_status)
    def text_filter(column, value: str):
        if value.strip():
            conditions.append(column.ilike(f"%{value.strip()}%"))
    text_filter(BusinessRecord.serial_no, serial_no)
    text_filter(BusinessRecord.owner, applicant)
    text_filter(BusinessRecord.customer, customer)
    text_filter(BusinessRecord.data["case_no"].as_string(), case_no)
    text_filter(BusinessRecord.data["contract_no"].as_string(), contract_no)
    text_filter(BusinessRecord.data["document_names"].as_string(), file_name)
    if use_type.strip(): conditions.append(BusinessRecord.data["use_type"].as_string() == use_type.strip())
    if date_from: conditions.append(func.date(BusinessRecord.created_at) >= date_from)
    if date_to: conditions.append(func.date(BusinessRecord.created_at) <= date_to)
    if keyword:
        like = f"%{keyword.strip()}%"
        conditions.append(or_(BusinessRecord.serial_no.ilike(like), BusinessRecord.title.ilike(like), BusinessRecord.customer.ilike(like), BusinessRecord.owner.ilike(like), BusinessRecord.data["case_no"].as_string().ilike(like), BusinessRecord.data["contract_no"].as_string().ilike(like), BusinessRecord.data["document_names"].as_string().ilike(like)))
    total = int(await db.scalar(select(func.count()).select_from(BusinessRecord).where(*conditions)) or 0)
    rows = (await db.scalars(select(BusinessRecord).where(*conditions).order_by(BusinessRecord.updated_at.desc()).offset((page - 1) * page_size).limit(page_size))).all()
    all_seals = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "seal", *scope_conditions))).all()
    summary = {"total": len(all_seals), "pending": sum(x.status == "待审批" for x in all_seals), "waiting_stamp": sum(x.status == "待用印" for x in all_seals), "completed": sum(x.status in {"已用印", "已归档"} for x in all_seals)}
    return {"items": [await _seal_record_dict(x, db) for x in rows], "total": total, "page": page, "page_size": page_size, "summary": summary}


@app.post(f"{settings.api_prefix}/seals/applications/package-download")
async def package_download_seal_files(body: SealPackageDownloadInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    record_ids = list(dict.fromkeys(body.application_ids))
    records: dict[int, BusinessRecord] = {}
    for record_id in record_ids:
        records[record_id] = await _ensure_record_module(record_id, "seal", identity, db)
    attachments = (await db.scalars(
        select(FileAttachment)
        .where(FileAttachment.record_id.in_(record_ids))
        .order_by(FileAttachment.record_id, FileAttachment.created_at, FileAttachment.id)
    )).all()
    if not attachments:
        raise HTTPException(status_code=404, detail="所选用印申请暂无可下载附件")

    output = io.BytesIO()
    included = 0
    upload_root = UPLOAD_ROOT.resolve()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for attachment in attachments:
            path = Path(attachment.path)
            if not path.is_file() or upload_root not in path.resolve().parents:
                continue
            record = records[int(attachment.record_id)]
            safe_name = Path(attachment.original_name).name or attachment.stored_name
            archive.writestr(f"{record.serial_no}/{attachment.id}-{safe_name}", path.read_bytes())
            included += 1
    if not included:
        raise HTTPException(status_code=404, detail="所选附件文件不存在")
    output.seek(0)
    filename = f"seal-files-{date.today():%Y%m%d}.zip"
    return StreamingResponse(
        output,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post(f"{settings.api_prefix}/seals/applications", status_code=status.HTTP_201_CREATED)
async def create_seal_application(body: SealApplicationInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    asset = await db.get(SealAsset, body.seal_asset_id)
    if not asset: raise HTTPException(status_code=404, detail="印章不存在")
    if asset.status != "可用": raise HTTPException(status_code=409, detail=f"印章当前状态为“{asset.status}”，不能申请")
    case_no, contract_no, customer, use_type = await _validated_seal_relations(body, identity, db)
    serial = f"YY{datetime.now():%Y%m%d%H%M%S}{uuid4().hex[:3].upper()}"
    item = BusinessRecord(module="seal", serial_no=serial, title=body.title, customer=customer, status="草稿", owner=identity["username"], description=body.description, data={"case_no": case_no, "contract_no": contract_no, "use_type": use_type, "seal_asset_id": body.seal_asset_id, "seal_type": asset.seal_type, "seal_name": asset.name, "copies": body.copies, "purpose": body.purpose, "use_date": str(body.use_date), "delivery_method": body.delivery_method, "is_electronic_seal": body.is_electronic_seal, "is_offline_print": body.is_offline_print, "document_names": body.document_names})
    db.add(item); await db.flush()
    db.add(WorkflowEvent(record_id=item.id, action="创建用印申请", to_status="草稿", operator=identity["username"], comment=f"{asset.name}｜{body.copies}份｜{body.purpose}"))
    await db.commit(); await db.refresh(item)
    return await _seal_record_dict(item, db)


@app.patch(f"{settings.api_prefix}/seals/applications/{{record_id}}")
async def update_seal_application(record_id: int, body: SealApplicationInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _get_seal_application(record_id, identity, db)
    await _require_record_owner_or_manager(item, identity, db)
    if item.status != "草稿": raise HTTPException(status_code=409, detail="只有草稿用印申请可以修改")
    asset = await db.get(SealAsset, body.seal_asset_id)
    if not asset: raise HTTPException(status_code=404, detail="印章不存在")
    if asset.status != "可用": raise HTTPException(status_code=409, detail=f"印章当前状态为“{asset.status}”，不能申请")
    case_no, contract_no, customer, use_type = await _validated_seal_relations(body, identity, db)
    item.title = body.title.strip(); item.customer = customer; item.description = body.description.strip()
    existing_names = str((item.data or {}).get("document_names") or "")
    item.data = {"case_no": case_no, "contract_no": contract_no, "use_type": use_type, "seal_asset_id": body.seal_asset_id, "seal_type": asset.seal_type, "seal_name": asset.name, "copies": body.copies, "purpose": body.purpose, "use_date": str(body.use_date), "delivery_method": body.delivery_method, "is_electronic_seal": body.is_electronic_seal, "is_offline_print": body.is_offline_print, "document_names": existing_names or body.document_names}
    db.add(WorkflowEvent(record_id=item.id, action="修改用印草稿", from_status="草稿", to_status="草稿", operator=identity["username"], comment=f"{asset.name}｜{body.copies}份｜{body.purpose}"))
    await db.commit(); await db.refresh(item)
    return await _seal_record_dict(item, db)


@app.delete(f"{settings.api_prefix}/seals/applications/{{record_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_seal_application(record_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Remove only an untouched draft so page-created test data can be safely cleaned."""
    item = await _get_seal_application(record_id, identity, db)
    await _require_record_owner_or_manager(item, identity, db)
    if item.status != "草稿":
        raise HTTPException(status_code=409, detail="只有草稿用印申请可以删除；已提交申请请按流程撤回")
    attachment_count = int(await db.scalar(select(func.count()).select_from(FileAttachment).where(FileAttachment.record_id == item.id)) or 0)
    if attachment_count:
        raise HTTPException(status_code=409, detail="草稿已有关联附件，请先通过附件流程处理后再删除")
    await db.execute(delete(WorkflowEvent).where(WorkflowEvent.record_id == item.id))
    await db.delete(item)
    await db.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


async def _get_seal_application(record_id: int, identity: dict, db: AsyncSession) -> BusinessRecord:
    try:
        return await _ensure_record_module(record_id, "seal", identity, db)
    except HTTPException as exc:
        if exc.status_code == 404: raise HTTPException(status_code=404, detail="用印申请不存在或无权访问") from exc
        raise


@app.post(f"{settings.api_prefix}/seals/applications/{{record_id}}/submit")
async def submit_seal_application(record_id: int, body: TaskActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _get_seal_application(record_id, identity, db)
    await _require_record_owner_or_manager(item, identity, db)
    if item.status != "草稿": raise HTTPException(status_code=409, detail="只有草稿可以提交审批")
    attachment_count = int(await db.scalar(select(func.count()).select_from(FileAttachment).where(FileAttachment.record_id == item.id, FileAttachment.category == "用印文件")) or 0)
    if not attachment_count:
        raise HTTPException(status_code=409, detail="请先上传至少一个用印文件后再提交审批")
    old = item.status; item.status = "待审批"
    db.add(WorkflowEvent(record_id=item.id, action="提交用印审批", from_status=old, to_status=item.status, operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(item); return await _seal_record_dict(item, db)


@app.post(f"{settings.api_prefix}/seals/applications/{{record_id}}/withdraw")
async def withdraw_seal_application(record_id: int, body: TaskActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _get_seal_application(record_id, identity, db)
    if identity.get("role") != "admin" and item.owner != identity["username"]:
        raise HTTPException(status_code=403, detail="只有申请人或管理员可以撤回用印申请")
    if item.status not in {"待审批", "待用印"}:
        raise HTTPException(status_code=409, detail="只有待审批或已审待用印的申请可以撤回")
    previous = item.status
    item.status = "已撤回"
    db.add(WorkflowEvent(record_id=item.id, action="撤回用印申请", from_status=previous, to_status="已撤回", operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(item); return await _seal_record_dict(item, db)


@app.post(f"{settings.api_prefix}/seals/applications/{{record_id}}/approve")
async def approve_seal_application(record_id: int, body: SealApprovalInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _get_seal_application(record_id, identity, db)
    if identity.get("role") not in {"admin", "manager", "auditor"}: raise HTTPException(status_code=403, detail="当前角色没有用印审批权限")
    if item.status != "待审批": raise HTTPException(status_code=409, detail="申请不在待审批状态")
    old = item.status; item.status = "待用印" if body.approved else "已拒绝"
    item.data = {
        **(item.data or {}),
        "approver": identity["username"],
        "approved_at": datetime.now().isoformat(timespec="seconds"),
        "approval_comment": body.comment.strip(),
    }
    db.add(WorkflowEvent(record_id=item.id, action="用印审批通过" if body.approved else "用印审批拒绝", from_status=old, to_status=item.status, operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(item); return await _seal_record_dict(item, db)


@app.post(f"{settings.api_prefix}/seals/applications/{{record_id}}/stamp")
async def stamp_seal_application(record_id: int, body: SealStampInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _get_seal_application(record_id, identity, db)
    if identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以登记实际用印")
    if item.status != "待用印": raise HTTPException(status_code=409, detail="申请尚未审批通过或已经用印")
    requested = int((item.data or {}).get("copies") or 0)
    if body.actual_copies > requested: raise HTTPException(status_code=409, detail=f"实际用印份数不能超过申请份数 {requested}")
    asset = await db.get(SealAsset, int((item.data or {}).get("seal_asset_id") or 0))
    if not asset or asset.status != "可用": raise HTTPException(status_code=409, detail="关联印章不存在或当前不可用")
    old = item.status; item.status = "已用印"; data = dict(item.data or {})
    data.update({"actual_copies": body.actual_copies, "stamp_operator": body.operator or identity["username"], "stamped_at": datetime.now().isoformat(), "archive_no": body.archive_no}); item.data = data
    asset.usage_count += body.actual_copies; asset.last_used_at = datetime.now()
    db.add(WorkflowEvent(record_id=item.id, action="完成实际用印", from_status=old, to_status=item.status, operator=identity["username"], comment=f"实际 {body.actual_copies} 份；归档号：{body.archive_no}。{body.comment}"))
    await db.commit(); await db.refresh(item); await db.refresh(asset); return await _seal_record_dict(item, db)


@app.post(f"{settings.api_prefix}/seals/applications/{{record_id}}/archive")
async def archive_seal_application(record_id: int, body: TaskActionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _get_seal_application(record_id, identity, db)
    if identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以归档用印材料")
    if item.status != "已用印": raise HTTPException(status_code=409, detail="只有已用印申请可以归档")
    if not (item.data or {}).get("archive_no"): raise HTTPException(status_code=409, detail="请先在用印登记中填写归档号")
    item.status = "已归档"
    db.add(WorkflowEvent(record_id=item.id, action="用印材料归档", from_status="已用印", to_status="已归档", operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(item); return await _seal_record_dict(item, db)


@app.get(f"{settings.api_prefix}/seals/assets")
async def list_seal_assets(keyword: str = "", identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    conditions = []
    if keyword:
        like = f"%{keyword.strip()}%"; conditions.append(or_(SealAsset.code.ilike(like), SealAsset.name.ilike(like), SealAsset.seal_type.ilike(like), SealAsset.custodian.ilike(like)))
    items = (await db.scalars(select(SealAsset).where(*conditions).order_by(SealAsset.code))).all()
    return {"items": [_seal_asset_dict(x) for x in items], "total": len(items)}


@app.post(f"{settings.api_prefix}/seals/assets", status_code=status.HTTP_201_CREATED)
async def create_seal_asset(body: SealAssetInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity["role"] != "admin": raise HTTPException(status_code=403, detail="仅管理员可新增印章")
    if body.seal_type not in REQUIRED_SEAL_TYPES: raise HTTPException(status_code=422, detail="印章类型不在系统允许范围内")
    if await db.scalar(select(SealAsset.id).where(SealAsset.code == body.code)): raise HTTPException(status_code=409, detail="印章编号已存在")
    item = SealAsset(**body.model_dump()); db.add(item); await db.commit(); await db.refresh(item); return _seal_asset_dict(item)


@app.patch(f"{settings.api_prefix}/seals/assets/{{asset_id}}")
async def update_seal_asset(asset_id: int, body: SealAssetUpdate, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity["role"] != "admin": raise HTTPException(status_code=403, detail="仅管理员可维护印章")
    item = await db.get(SealAsset, asset_id)
    if not item: raise HTTPException(status_code=404, detail="印章不存在")
    changes = body.model_dump(exclude_unset=True)
    if changes.get("status") not in {None, "可用", "停用", "维修", "遗失"}: raise HTTPException(status_code=422, detail="无效的印章状态")
    if changes.get("seal_type") not in {None, *REQUIRED_SEAL_TYPES}: raise HTTPException(status_code=422, detail="印章类型不在系统允许范围内")
    for key, value in changes.items(): setattr(item, key, value)
    await db.commit(); await db.refresh(item); return _seal_asset_dict(item)


def _department_dict(item: Department) -> dict:
    return {"id": item.id, "code": item.code, "name": item.name, "manager": item.manager, "sort_order": item.sort_order, "is_active": item.is_active, "created_by": item.created_by, "updated_by": item.updated_by, "created_at": item.created_at, "updated_at": item.updated_at}


def _job_role_dict(item: JobRole) -> dict:
    return {"id": item.id, "code": item.code, "name": item.name, "permissions": item.permissions or [], "description": item.description, "sort_order": item.sort_order, "is_active": item.is_active, "created_by": item.created_by, "updated_by": item.updated_by, "created_at": item.created_at, "updated_at": item.updated_at}


@app.post(f"{settings.api_prefix}/hr/employees", status_code=status.HTTP_201_CREATED)
async def create_hr_employee(body: HrEmployeeCreateInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    account_type = (body.account_type or body.data.get("account_type") or "员工账号").strip()
    if account_type not in {"员工账号", "客户账号", "外部合作账号"}:
        raise HTTPException(status_code=422, detail="账号类型无效")
    username = body.username.strip().lower()
    employee_no = body.employee_no.strip()
    if await db.scalar(select(BusinessRecord.id).where(BusinessRecord.module == "hr", BusinessRecord.serial_no == employee_no)):
        raise HTTPException(status_code=409, detail="员工编号已存在")
    department = await db.scalar(select(Department).where(Department.name == body.department, Department.is_active.is_(True)))
    if not department:
        raise HTTPException(status_code=422, detail="所选部门不存在或已停用")
    position = await db.scalar(select(JobRole).where(JobRole.name == body.position, JobRole.is_active.is_(True)))
    if not position:
        raise HTTPException(status_code=422, detail="所选职务不存在或已停用")
    profile = {**body.data, "account_type": account_type, "employee_no": employee_no, "company": body.company.strip(), "position": body.position.strip()}
    user: User | None = None
    if account_type == "员工账号":
        if not username:
            raise HTTPException(status_code=422, detail="员工账号必须填写登录用户名")
        if username == "admin":
            raise HTTPException(status_code=409, detail="不能通过员工档案创建或覆盖管理员账号")
        if not re.fullmatch(r"[a-z0-9._-]+", username):
            raise HTTPException(status_code=422, detail="登录账号只能包含小写字母、数字、点、下划线或短横线")
        existing_employee = await db.scalar(select(BusinessRecord.id).where(BusinessRecord.module == "hr", or_(BusinessRecord.owner == username, BusinessRecord.data["username"].as_string() == username)))
        if existing_employee:
            raise HTTPException(status_code=409, detail="该登录账号已关联其他员工档案")
        user = await db.scalar(select(User).where(User.username == username))
        if user:
            if user.role == "admin":
                raise HTTPException(status_code=409, detail="不能通过员工档案覆盖管理员账号")
            # Existing, independently-created accounts are linked rather than
            # duplicated.  Do not reset their password or elevate their role.
            user.display_name = body.display_name.strip(); user.department = body.department.strip()
            user.is_active = body.is_active; user.profile = {**(user.profile or {}), **profile}
        else:
            policy = await _security_policy(db)
            if len(body.password) < policy.min_password_length:
                raise HTTPException(status_code=422, detail=f"员工账号密码至少需要 {policy.min_password_length} 位")
            user = User(
                username=username, display_name=body.display_name.strip(), department=body.department.strip(),
                # Job position controls investigation capability.  A new HR
                # account always starts as the least-privileged system user.
                role="user", profile=profile, password_hash=hash_password(body.password),
                is_active=body.is_active, password_changed_at=datetime.now(),
            )
    employee = BusinessRecord(
        module="hr", serial_no=employee_no, title=body.display_name.strip(), customer=body.company.strip(),
        status="在职" if body.is_active else "停用", owner=username if user else identity["username"], department=body.department.strip(), description="",
        data={**profile, "username": username if user else "", "role": user.role if user else "", "is_active": body.is_active},
    )
    db.add(employee)
    if user and not user.id:
        db.add(user)
    try:
        await db.flush()
        db.add(WorkflowEvent(record_id=employee.id, action="新建员工", from_status="", to_status=employee.status, operator=identity["username"], comment=f"账号类型：{account_type}；{'登录账号：' + username if user else '不创建系统登录账号'}；职务：{body.position}"))
        await db.commit()
        if user:
            await db.refresh(user)
        await db.refresh(employee)
    except Exception:
        await db.rollback()
        raise
    return {"employee": _record_dict(employee), "user": _system_user_dict(user) if user else None}


@app.patch(f"{settings.api_prefix}/hr/employees/{{employee_id}}")
async def update_hr_employee(employee_id: int, body: HrEmployeeUpdateInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity)
    employee = await db.get(BusinessRecord, employee_id)
    if not employee or employee.module != "hr": raise HTTPException(status_code=404, detail="员工档案不存在")
    if body.role not in {"admin", "manager", "auditor", "user"}: raise HTTPException(status_code=422, detail="角色值无效")
    if body.left_at and body.left_at < body.joined_at: raise HTTPException(status_code=422, detail="离职日期不能早于入职日期")
    department = await db.scalar(select(Department).where(Department.name == body.department, Department.is_active.is_(True)))
    if not department: raise HTTPException(status_code=422, detail="所选部门不存在或已停用")
    position = await db.scalar(select(JobRole).where(JobRole.name == body.position, JobRole.is_active.is_(True)))
    if not position: raise HTTPException(status_code=422, detail="所选职务不存在或已停用")
    account_type = str(body.data.get("account_type") or (employee.data or {}).get("account_type") or "员工账号").strip()
    if account_type not in {"员工账号", "客户账号", "外部合作账号"}:
        raise HTTPException(status_code=422, detail="账号类型无效")
    username = str((employee.data or {}).get("username") or employee.owner).strip().lower()
    user = await db.scalar(select(User).where(User.username == username))
    if account_type == "员工账号" and not user: raise HTTPException(status_code=409, detail="员工账号关联的登录用户不存在，不能只修改一侧资料")
    if user and user.username == "admin":
        raise HTTPException(status_code=409, detail="管理员账号不能通过员工档案修改、停用或改名")
    if account_type != "员工账号" and user:
        raise HTTPException(status_code=409, detail="请先在系统用户管理中解除登录账号关联，再变更为非员工账号")
    if not user:
        previous_status = employee.status
        profile = {**(employee.data or {}), **body.data, "account_type": account_type, "employee_no": employee.serial_no, "company": employee.customer, "position": body.position, "email": body.email.strip(), "mobile": body.mobile.strip(), "office_phone": body.office_phone.strip(), "joined_at": str(body.joined_at), "left_at": str(body.left_at) if body.left_at else ""}
        employee.title = body.display_name.strip(); employee.department = body.department.strip(); employee.data = profile
        db.add(WorkflowEvent(record_id=employee.id, action="修改员工资料", from_status=previous_status, to_status=employee.status, operator=identity["username"], comment=f"账号类型：{account_type}；未关联系统登录账号"))
        await db.commit(); await db.refresh(employee)
        return {"employee": _record_dict(employee), "user": None}
    username = await _rename_system_username(user, body.username, identity, db)
    if user.username == identity["username"] and not body.is_active: raise HTTPException(status_code=409, detail="不能停用当前登录账号")
    if user.username == identity["username"] and body.role != "admin": raise HTTPException(status_code=409, detail="不能取消当前登录账号的管理员角色")
    # Employment status is the source of truth for a linked employee login.
    # Account enable/disable must go through the dedicated HR lifecycle so the
    # reason, effective date and workflow event cannot be bypassed by a profile
    # edit.  Probationary accounts remain usable; only offboarded/disabled staff
    # are prevented from logging in.
    expected_active = employee.status not in {"离职", "停用"}
    if body.is_active != expected_active:
        raise HTTPException(status_code=409, detail="员工账号启停必须通过“办理状态”完成，以同步人事状态、原因和生效日期")
    previous_status = employee.status
    profile = {**(user.profile or {}), **body.data, "account_type": "员工账号", "employee_no": employee.serial_no, "company": employee.customer, "position": body.position, "email": body.email.strip(), "mobile": body.mobile.strip(), "office_phone": body.office_phone.strip(), "joined_at": str(body.joined_at), "left_at": str(body.left_at) if body.left_at else ""}
    user.display_name = body.display_name.strip(); user.department = body.department.strip(); user.role = body.role; user.is_active = body.is_active; user.profile = profile
    employee.title = body.display_name.strip(); employee.department = body.department.strip(); employee.data = {**(employee.data or {}), **profile, "username": username, "role": body.role, "is_active": body.is_active}
    db.add(WorkflowEvent(record_id=employee.id, action="修改员工资料", from_status=previous_status, to_status=employee.status, operator=identity["username"], comment=f"部门：{employee.department}；职务：{body.position}；账号：{'启用' if body.is_active else '停用'}"))
    await db.commit(); await db.refresh(employee); await db.refresh(user)
    return {"employee": _record_dict(employee), "user": _system_user_dict(user)}


@app.get(f"{settings.api_prefix}/hr/departments")
async def list_departments(keyword: str = "", active_only: bool = False, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    statement = select(Department)
    if keyword.strip():
        term = f"%{keyword.strip()}%"; statement = statement.where(or_(Department.code.ilike(term), Department.name.ilike(term), Department.manager.ilike(term)))
    if active_only: statement = statement.where(Department.is_active.is_(True))
    items = (await db.scalars(statement.order_by(Department.sort_order, Department.id))).all()
    return {"items": [_department_dict(item) for item in items], "total": len(items)}


@app.post(f"{settings.api_prefix}/hr/departments", status_code=status.HTTP_201_CREATED)
async def create_department(body: DepartmentInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity); code, name = body.code.strip().upper(), body.name.strip()
    if await db.scalar(select(Department.id).where(or_(Department.code == code, Department.name == name))): raise HTTPException(status_code=409, detail="部门代码或名称已存在")
    item = Department(**body.model_dump(exclude={"code", "name", "manager"}), code=code, name=name, manager=body.manager.strip(), created_by=identity["username"], updated_by=identity["username"])
    db.add(item); await db.commit(); await db.refresh(item); return _department_dict(item)


@app.patch(f"{settings.api_prefix}/hr/departments/{{department_id}}")
async def update_department(department_id: int, body: DepartmentUpdate, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity); item = await db.get(Department, department_id)
    if not item: raise HTTPException(status_code=404, detail="部门不存在")
    old_name = item.name; code = body.code.strip().upper() if body.code is not None else item.code; name = body.name.strip() if body.name is not None else item.name
    if await db.scalar(select(Department.id).where(Department.id != item.id, or_(Department.code == code, Department.name == name))): raise HTTPException(status_code=409, detail="部门代码或名称已存在")
    changes = body.model_dump(exclude_unset=True, exclude_none=True)
    for key, value in changes.items(): setattr(item, key, value.strip().upper() if key == "code" else value.strip() if key in {"name", "manager"} else value)
    item.updated_by = identity["username"]
    if item.name != old_name:
        await db.execute(update(User).where(User.department == old_name).values(department=item.name))
        await db.execute(update(BusinessRecord).where(BusinessRecord.department == old_name).values(department=item.name))
    await db.commit(); await db.refresh(item); return _department_dict(item)


@app.delete(f"{settings.api_prefix}/hr/departments/{{department_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_department(department_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity); item = await db.get(Department, department_id)
    if not item: raise HTTPException(status_code=404, detail="部门不存在")
    users = await db.scalar(select(func.count()).select_from(User).where(User.department == item.name))
    records = await db.scalar(select(func.count()).select_from(BusinessRecord).where(BusinessRecord.department == item.name))
    if (users or 0) + (records or 0) > 0: raise HTTPException(status_code=409, detail=f"部门仍被 {users or 0} 个账号和 {records or 0} 条业务记录使用，请先停用或迁移")
    await db.delete(item); await db.commit(); return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.get(f"{settings.api_prefix}/hr/job-roles")
async def list_job_roles(keyword: str = "", active_only: bool = False, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    statement = select(JobRole)
    if keyword.strip():
        term = f"%{keyword.strip()}%"; statement = statement.where(or_(JobRole.code.ilike(term), JobRole.name.ilike(term), JobRole.description.ilike(term)))
    if active_only: statement = statement.where(JobRole.is_active.is_(True))
    items = (await db.scalars(statement.order_by(JobRole.sort_order, JobRole.id))).all()
    return {"items": [_job_role_dict(item) for item in items], "total": len(items)}


@app.post(f"{settings.api_prefix}/hr/job-roles", status_code=status.HTTP_201_CREATED)
async def create_job_role(body: JobRoleInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity); code, name = body.code.strip().upper(), body.name.strip()
    if await db.scalar(select(JobRole.id).where(or_(JobRole.code == code, JobRole.name == name))): raise HTTPException(status_code=409, detail="岗位角色代码或名称已存在")
    permissions = list(dict.fromkeys(value.strip() for value in body.permissions if value.strip()))
    item = JobRole(**body.model_dump(exclude={"code", "name", "permissions", "description"}), code=code, name=name, permissions=permissions, description=body.description.strip(), created_by=identity["username"], updated_by=identity["username"])
    db.add(item); await db.commit(); await db.refresh(item); return _job_role_dict(item)


@app.patch(f"{settings.api_prefix}/hr/job-roles/{{role_id}}")
async def update_job_role(role_id: int, body: JobRoleUpdate, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity); item = await db.get(JobRole, role_id)
    if not item: raise HTTPException(status_code=404, detail="岗位角色不存在")
    old_name = item.name; code = body.code.strip().upper() if body.code is not None else item.code; name = body.name.strip() if body.name is not None else item.name
    if await db.scalar(select(JobRole.id).where(JobRole.id != item.id, or_(JobRole.code == code, JobRole.name == name))): raise HTTPException(status_code=409, detail="岗位角色代码或名称已存在")
    changes = body.model_dump(exclude_unset=True, exclude_none=True)
    for key, value in changes.items():
        if key == "permissions": value = list(dict.fromkeys(entry.strip() for entry in value if entry.strip()))
        elif key == "code": value = value.strip().upper()
        elif key in {"name", "description"}: value = value.strip()
        setattr(item, key, value)
    item.updated_by = identity["username"]
    if item.name != old_name:
        hr_records = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "hr"))).all()
        for record in hr_records:
            if (record.data or {}).get("position") == old_name: record.data = {**(record.data or {}), "position": item.name}
    await db.commit(); await db.refresh(item); return _job_role_dict(item)


@app.delete(f"{settings.api_prefix}/hr/job-roles/{{role_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_job_role(role_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    _require_admin(identity); item = await db.get(JobRole, role_id)
    if not item: raise HTTPException(status_code=404, detail="岗位角色不存在")
    hr_records = (await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "hr"))).all()
    used = sum(1 for record in hr_records if (record.data or {}).get("position") == item.name)
    if used: raise HTTPException(status_code=409, detail=f"岗位角色仍被 {used} 份员工档案使用，请先调整员工岗位")
    await db.delete(item); await db.commit(); return Response(status_code=status.HTTP_204_NO_CONTENT)


HR_SUBRECORD_KINDS = {"leave", "matter", "commission"}


def _hr_subrecord_dict(item: HrSubrecord) -> dict:
    return {
        "id": item.id, "employee_id": item.employee_id, "kind": item.kind,
        "data": item.data or {}, "created_by": item.created_by,
        "updated_by": item.updated_by, "created_at": item.created_at,
        "updated_at": item.updated_at,
    }


def _validate_hr_subrecord(kind: str, raw: dict) -> dict:
    if kind not in HR_SUBRECORD_KINDS:
        raise HTTPException(status_code=422, detail="不支持的员工附属记录类型")
    data = dict(raw or {})
    if kind == "leave":
        required = {"start_date": "请假开始", "end_date": "请假结束", "leave_type": "请假类型"}
        for key, label in required.items():
            if not str(data.get(key) or "").strip(): raise HTTPException(status_code=422, detail=f"请填写{label}")
        try: start, end = date.fromisoformat(str(data["start_date"])), date.fromisoformat(str(data["end_date"]))
        except ValueError as exc: raise HTTPException(status_code=422, detail="请假日期格式无效") from exc
        if end < start: raise HTTPException(status_code=422, detail="请假结束不能早于开始")
        hours = float(data.get("hours") or 0)
        if hours <= 0: raise HTTPException(status_code=422, detail="请假小时数必须大于 0")
        data.update({"start_date": str(start), "end_date": str(end), "hours": hours, "leave_type": str(data["leave_type"]).strip(), "remark": str(data.get("remark") or "").strip()})
    elif kind == "matter":
        content = str(data.get("content") or "").strip()
        if not content: raise HTTPException(status_code=422, detail="请填写事项内容")
        operation_date = str(data.get("operation_date") or date.today())
        try: date.fromisoformat(operation_date)
        except ValueError as exc: raise HTTPException(status_code=422, detail="事项日期格式无效") from exc
        data.update({"content": content, "operation_date": operation_date})
    else:
        start_date, end_date = str(data.get("start_date") or ""), str(data.get("end_date") or "")
        if not start_date: raise HTTPException(status_code=422, detail="请填写提成开始日期")
        try:
            start = date.fromisoformat(start_date); end = date.fromisoformat(end_date) if end_date else None
        except ValueError as exc: raise HTTPException(status_code=422, detail="提成日期格式无效") from exc
        if end and end < start: raise HTTPException(status_code=422, detail="提成结束不能早于开始")
        numeric = ["base_salary", "hearing_rate", "document_rate", "source_rate", "investigation_rate", "quality_rate"]
        for key in numeric:
            value = float(data.get(key) or 0)
            if value < 0: raise HTTPException(status_code=422, detail="提成或工资数值不能为负数")
            data[key] = value
        data.update({"start_date": str(start), "end_date": str(end) if end else ""})
    return data


@app.get(f"{settings.api_prefix}/hr/{{employee_id}}/subrecords")
async def list_hr_subrecords(employee_id: int, kind: str = "", identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    await _ensure_record_module(employee_id, "hr", identity, db)
    conditions = [HrSubrecord.employee_id == employee_id]
    if kind:
        if kind not in HR_SUBRECORD_KINDS: raise HTTPException(status_code=422, detail="不支持的员工附属记录类型")
        conditions.append(HrSubrecord.kind == kind)
    items = (await db.scalars(select(HrSubrecord).where(*conditions).order_by(HrSubrecord.created_at.desc(), HrSubrecord.id.desc()))).all()
    return {"items": [_hr_subrecord_dict(item) for item in items], "total": len(items)}


@app.post(f"{settings.api_prefix}/hr/{{employee_id}}/subrecords", status_code=status.HTTP_201_CREATED)
async def create_hr_subrecord(employee_id: int, body: HrSubrecordInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以维护员工记录")
    await _ensure_record_module(employee_id, "hr", identity, db)
    item = HrSubrecord(employee_id=employee_id, kind=body.kind, data=_validate_hr_subrecord(body.kind, body.data), created_by=identity["username"], updated_by=identity["username"])
    db.add(item); await db.commit(); await db.refresh(item); return _hr_subrecord_dict(item)


@app.patch(f"{settings.api_prefix}/hr/{{employee_id}}/subrecords/{{subrecord_id}}")
async def update_hr_subrecord(employee_id: int, subrecord_id: int, body: HrSubrecordUpdate, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以维护员工记录")
    await _ensure_record_module(employee_id, "hr", identity, db)
    item = await db.get(HrSubrecord, subrecord_id)
    if not item or item.employee_id != employee_id: raise HTTPException(status_code=404, detail="员工附属记录不存在")
    item.data = _validate_hr_subrecord(item.kind, body.data); item.updated_by = identity["username"]
    await db.commit(); await db.refresh(item); return _hr_subrecord_dict(item)


@app.delete(f"{settings.api_prefix}/hr/{{employee_id}}/subrecords/{{subrecord_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_hr_subrecord(employee_id: int, subrecord_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以维护员工记录")
    await _ensure_record_module(employee_id, "hr", identity, db)
    item = await db.get(HrSubrecord, subrecord_id)
    if not item or item.employee_id != employee_id: raise HTTPException(status_code=404, detail="员工附属记录不存在")
    await db.delete(item); await db.commit(); return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.post(f"{settings.api_prefix}/hr/{{employee_id}}/transition")
async def transition_employee(employee_id: int, body: HrTransitionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以办理人事状态")
    item = await _ensure_record_module(employee_id, "hr", identity, db)
    allowed = WORKFLOW_TRANSITIONS["hr"].get(item.status, [])
    if body.to_status not in allowed: raise HTTPException(status_code=409, detail=f"不能从“{item.status}”办理到“{body.to_status}”")
    reason = body.reason.strip()
    if body.to_status in {"离职", "停用"} and len(reason) < 2: raise HTTPException(status_code=422, detail="离职或停用必须填写原因")
    if body.effective_date < date(2000, 1, 1): raise HTTPException(status_code=422, detail="办理日期无效")
    previous = item.status; data = dict(item.data or {})
    if body.to_status == "在职":
        data.update({"regularized_at": str(body.effective_date), "regularized_by": identity["username"]}); action = "试用转正"
    elif body.to_status == "离职":
        data.update({"offboard_date": str(body.effective_date), "offboard_reason": reason, "handover_to": body.handover_to.strip(), "offboard_by": identity["username"]}); action = "办理离职"
    else:
        data.update({"disabled_at": str(body.effective_date), "disabled_reason": reason, "disabled_by": identity["username"]}); action = "停用员工"
    # The dedicated HR lifecycle is the authoritative employment switch for a
    # linked employee login.  Only ordinary employee users are synchronized:
    # an administrator (or a legacy elevated account) must never be altered by
    # an HR status transition.
    linked_username = str(data.get("username") or item.owner or "").strip().lower()
    linked_user = await db.scalar(select(User).where(User.username == linked_username)) if linked_username else None
    if linked_user and linked_user.role == "user":
        login_enabled = body.to_status == "在职"
        linked_user.is_active = login_enabled
        data["is_active"] = login_enabled
        data["login_account_synced_at"] = datetime.now().isoformat()
        data["login_account_sync"] = "已启用" if login_enabled else "已停用"
    item.status = body.to_status; item.data = data
    db.add(WorkflowEvent(record_id=item.id, action=action, from_status=previous, to_status=item.status, operator=identity["username"], comment=body.comment or reason))
    await db.commit(); await db.refresh(item); return _record_dict(item, await _allowed_field_keys(identity, db))


def _warehouse_evidence_status(item: BusinessRecord) -> str:
    data = item.data or {}
    explicit = str(data.get("evidence_status") or "").strip()
    if explicit:
        return explicit
    return {"在库": "已入库", "借出": "已出库", "归还中": "已出库", "报废": "已销毁"}.get(item.status, "未入库")


@app.post(f"{settings.api_prefix}/warehouse/evidence", status_code=status.HTTP_201_CREATED)
async def create_warehouse_evidence(body: WarehouseEvidenceInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以登记证物")
    if await db.scalar(select(BusinessRecord.id).where(BusinessRecord.serial_no == body.serial_no.strip())):
        raise HTTPException(status_code=409, detail="线索编号已存在")
    department = "上海分所"
    if identity.get("role") != "admin":
        user = await db.scalar(select(User).where(User.username == identity["username"]))
        if not user:
            raise HTTPException(status_code=401, detail="当前用户不存在")
        department = user.department
    evidence_data = body.model_dump(mode="json", exclude={"serial_no", "description"})
    evidence_data.update({
        "category": "证物",
        "quantity": 1,
        "unit": "件",
        "evidence_status": "未入库",
        "registered_at": str(date.today()),
        "registered_by": identity["username"],
    })
    item = BusinessRecord(
        module="warehouse",
        serial_no=body.serial_no.strip(),
        title=body.shop_name.strip(),
        customer=body.rights_holder.strip(),
        status="在库",
        owner=body.investigator.strip(),
        department=department,
        description=body.description.strip(),
        data=evidence_data,
    )
    db.add(item)
    await db.flush()
    db.add(WorkflowEvent(record_id=item.id, action="证物登记", to_status="未入库", operator=identity["username"], comment=body.description.strip()))
    await db.commit()
    await db.refresh(item)
    return _record_dict(item, await _allowed_field_keys(identity, db))


@app.patch(f"{settings.api_prefix}/warehouse/evidence/{{item_id}}")
async def update_warehouse_evidence(item_id: int, body: WarehouseEvidenceInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以修改证物资料")
    item = await _ensure_record_module(item_id, "warehouse", identity, db)
    evidence_status = _warehouse_evidence_status(item)
    if evidence_status in {"已出库", "已销毁"}:
        raise HTTPException(status_code=409, detail="已出库或已销毁的证物不能修改资料")
    duplicate = await db.scalar(select(BusinessRecord.id).where(BusinessRecord.serial_no == body.serial_no.strip(), BusinessRecord.id != item.id))
    if duplicate:
        raise HTTPException(status_code=409, detail="线索编号已存在")
    process_fields = {
        key: value for key, value in (item.data or {}).items()
        if key in {
            "category", "quantity", "unit", "evidence_status", "registered_at", "registered_by",
            "checked_in_at", "checked_in_by", "checked_out_at", "checked_out_by", "recipient",
            "checkout_purpose", "rechecked_in_at", "rechecked_in_by", "return_condition",
            "destroyed_at", "destroyed_by", "destroy_reason",
        }
    }
    evidence_data = body.model_dump(mode="json", exclude={"serial_no", "description"})
    evidence_data.update(process_fields)
    item.serial_no = body.serial_no.strip()
    item.title = body.shop_name.strip()
    item.customer = body.rights_holder.strip()
    item.owner = body.investigator.strip()
    item.description = body.description.strip()
    item.data = evidence_data
    db.add(WorkflowEvent(record_id=item.id, action="修改证物资料", from_status=evidence_status, to_status=evidence_status, operator=identity["username"], comment=body.description.strip()))
    await db.commit()
    await db.refresh(item)
    return _record_dict(item, await _allowed_field_keys(identity, db))


@app.post(f"{settings.api_prefix}/warehouse/evidence/{{item_id}}/check-in")
async def check_in_warehouse_evidence(item_id: int, body: WarehouseEvidenceCheckInInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以办理证物入库")
    item = await _ensure_record_module(item_id, "warehouse", identity, db)
    previous = _warehouse_evidence_status(item)
    if previous != "未入库":
        raise HTTPException(status_code=409, detail="只有未入库证物可以办理首次入库")
    item.status = "在库"
    item.data = {**(item.data or {}), "warehouse": body.warehouse.strip(), "location": body.location.strip(), "evidence_status": "已入库", "checked_in_at": str(date.today()), "checked_in_by": identity["username"]}
    db.add(WorkflowEvent(record_id=item.id, action="证物入库", from_status=previous, to_status="已入库", operator=identity["username"], comment=body.comment.strip()))
    await db.commit()
    await db.refresh(item)
    return _record_dict(item, await _allowed_field_keys(identity, db))


@app.post(f"{settings.api_prefix}/warehouse/evidence/{{item_id}}/check-out")
async def check_out_warehouse_evidence(item_id: int, body: WarehouseEvidenceCheckOutInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以办理证物出库")
    item = await _ensure_record_module(item_id, "warehouse", identity, db)
    previous = _warehouse_evidence_status(item)
    if previous not in {"已入库", "已重新入库"}:
        raise HTTPException(status_code=409, detail="只有在库证物可以办理出库")
    item.status = "借出"
    item.data = {**(item.data or {}), "evidence_status": "已出库", "checked_out_at": str(date.today()), "checked_out_by": identity["username"], "recipient": body.recipient.strip(), "checkout_purpose": body.purpose.strip()}
    db.add(WorkflowEvent(record_id=item.id, action="证物出库", from_status=previous, to_status="已出库", operator=identity["username"], comment=f"领取人：{body.recipient.strip()}；用途：{body.purpose.strip()}。{body.comment.strip()}"))
    await db.commit()
    await db.refresh(item)
    return _record_dict(item, await _allowed_field_keys(identity, db))


@app.post(f"{settings.api_prefix}/warehouse/evidence/{{item_id}}/recheck-in")
async def recheck_in_warehouse_evidence(item_id: int, body: WarehouseEvidenceRecheckInInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以办理证物重新入库")
    item = await _ensure_record_module(item_id, "warehouse", identity, db)
    previous = _warehouse_evidence_status(item)
    if previous != "已出库":
        raise HTTPException(status_code=409, detail="只有已出库证物可以办理重新入库")
    previous_data = item.data or {}
    item.status = "在库"
    item.data = {**previous_data, "warehouse": body.warehouse.strip(), "location": body.location.strip(), "evidence_status": "已重新入库", "rechecked_in_at": str(date.today()), "rechecked_in_by": identity["username"], "return_condition": body.condition.strip(), "last_recipient": previous_data.get("recipient", ""), "recipient": "", "checkout_purpose": ""}
    db.add(WorkflowEvent(record_id=item.id, action="证物重新入库", from_status=previous, to_status="已重新入库", operator=identity["username"], comment=f"物品状况：{body.condition.strip()}。{body.comment.strip()}"))
    await db.commit()
    await db.refresh(item)
    return _record_dict(item, await _allowed_field_keys(identity, db))


@app.post(f"{settings.api_prefix}/warehouse/evidence/{{item_id}}/destroy")
async def destroy_warehouse_evidence(item_id: int, body: WarehouseEvidenceDestroyInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") != "admin":
        raise HTTPException(status_code=403, detail="只有管理员可以销毁证物")
    item = await _ensure_record_module(item_id, "warehouse", identity, db)
    previous = _warehouse_evidence_status(item)
    if previous not in {"已入库", "已重新入库"}:
        raise HTTPException(status_code=409, detail="只有在库证物可以办理销毁")
    item.status = "报废"
    item.data = {**(item.data or {}), "evidence_status": "已销毁", "destroyed_at": str(date.today()), "destroyed_by": identity["username"], "destroy_reason": body.reason.strip()}
    db.add(WorkflowEvent(record_id=item.id, action="证物销毁", from_status=previous, to_status="已销毁", operator=identity["username"], comment=body.reason.strip()))
    await db.commit()
    await db.refresh(item)
    return _record_dict(item, await _allowed_field_keys(identity, db))


@app.post(f"{settings.api_prefix}/warehouse/{{item_id}}/borrow")
async def borrow_warehouse_item(item_id: int, body: WarehouseBorrowInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以办理借出")
    item = await _ensure_record_module(item_id, "warehouse", identity, db)
    if (item.data or {}).get("evidence_status"): raise HTTPException(status_code=409, detail="证物必须使用证物出库入口办理")
    if item.status != "在库": raise HTTPException(status_code=409, detail="只有在库物品可以借出")
    if body.due_date < date.today(): raise HTTPException(status_code=422, detail="预计归还日期不能早于今天")
    data = dict(item.data or {}); data.update({"borrower": body.borrower.strip(), "due_date": str(body.due_date), "borrow_purpose": body.purpose.strip(), "borrowed_at": str(date.today()), "borrowed_by": identity["username"], "return_requested_at": ""})
    item.status = "借出"; item.data = data
    db.add(WorkflowEvent(record_id=item.id, action="物品借出", from_status="在库", to_status="借出", operator=identity["username"], comment=f"借用人：{body.borrower}；预计归还：{body.due_date}。{body.comment}"))
    await db.commit(); await db.refresh(item); return _record_dict(item)


@app.post(f"{settings.api_prefix}/warehouse/{{item_id}}/return")
async def return_warehouse_item(item_id: int, body: WarehouseReturnInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item = await _ensure_record_module(item_id, "warehouse", identity, db)
    if (item.data or {}).get("evidence_status"): raise HTTPException(status_code=409, detail="证物必须使用证物重新入库入口办理")
    data = dict(item.data or {})
    if item.status != "借出": raise HTTPException(status_code=409, detail="只有已借出物品可以发起归还")
    if identity.get("role") not in {"admin", "manager"} and data.get("borrower") not in {identity["username"], identity.get("display_name", "")}: raise HTTPException(status_code=403, detail="只有借用人或管理人员可以发起归还")
    item.status = "归还中"; item.data = {**data, "return_requested_at": str(date.today()), "return_requested_by": identity["username"]}
    db.add(WorkflowEvent(record_id=item.id, action="发起归还", from_status="借出", to_status="归还中", operator=identity["username"], comment=body.comment))
    await db.commit(); await db.refresh(item); return _record_dict(item)


@app.post(f"{settings.api_prefix}/warehouse/{{item_id}}/return-confirm")
async def confirm_warehouse_return(item_id: int, body: WarehouseReturnConfirmInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") not in {"admin", "manager"}: raise HTTPException(status_code=403, detail="只有管理员或部门负责人可以验收归还")
    item = await _ensure_record_module(item_id, "warehouse", identity, db)
    if (item.data or {}).get("evidence_status"): raise HTTPException(status_code=409, detail="证物必须使用证物重新入库入口办理")
    if item.status != "归还中": raise HTTPException(status_code=409, detail="只有归还中的物品可以验收入库")
    previous_data = dict(item.data or {}); last_borrower = str(previous_data.get("borrower") or "")
    data = {**previous_data, "last_borrower": last_borrower, "last_due_date": previous_data.get("due_date", ""), "returned_at": str(date.today()), "return_condition": body.condition.strip(), "returned_by": identity["username"], "borrower": "", "due_date": "", "borrow_purpose": "", "return_requested_at": ""}
    if body.location.strip(): data["location"] = body.location.strip()
    item.status = "在库"; item.data = data
    db.add(WorkflowEvent(record_id=item.id, action="归还验收入库", from_status="归还中", to_status="在库", operator=identity["username"], comment=f"原借用人：{last_borrower}；物品状况：{body.condition}。{body.comment}"))
    await db.commit(); await db.refresh(item); return _record_dict(item)


@app.post(f"{settings.api_prefix}/warehouse/{{item_id}}/scrap")
async def scrap_warehouse_item(item_id: int, body: WarehouseScrapInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity.get("role") != "admin": raise HTTPException(status_code=403, detail="仅管理员可以报废物品")
    item = await _ensure_record_module(item_id, "warehouse", identity, db)
    if (item.data or {}).get("evidence_status"): raise HTTPException(status_code=409, detail="证物必须使用证物销毁入口办理")
    if item.status != "在库": raise HTTPException(status_code=409, detail="只有在库物品可以报废")
    item.status = "报废"; item.data = {**(item.data or {}), "scrapped_at": str(date.today()), "scrap_reason": body.reason.strip(), "scrapped_by": identity["username"]}
    db.add(WorkflowEvent(record_id=item.id, action="物品报废", from_status="在库", to_status="报废", operator=identity["username"], comment=body.reason))
    await db.commit(); await db.refresh(item); return _record_dict(item)


@app.post(f"{settings.api_prefix}/records", status_code=status.HTTP_201_CREATED)
async def create_record(body: RecordInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if body.module in INVESTIGATION_RECORD_MODULES:
        raise HTTPException(status_code=422, detail="调查、公证和证据记录必须使用调查中心专用入口创建")
    if body.module == "customer":
        raise HTTPException(status_code=422, detail="新建客户必须使用客户专用入口")
    if body.module == "contract":
        raise HTTPException(status_code=422, detail="新建合同必须使用合同专用入口")
    if body.module == "case":
        raise HTTPException(status_code=422, detail="新建案件必须选择已审批合同，请使用案件创建入口")
    if body.module == "task":
        raise HTTPException(status_code=422, detail="任务必须使用任务专用入口创建")
    if body.module == "finance_package":
        raise HTTPException(status_code=422, detail="付款包必须使用打包付款专用入口创建")
    if body.module == "finance_settlement":
        raise HTTPException(status_code=422, detail="结算申请必须使用结算管理专用入口创建")
    if body.module == "case_reminder":
        raise HTTPException(status_code=422, detail="案件提醒必须使用案件提醒专用入口创建")
    if body.module == "finance_archive_settlement":
        raise HTTPException(status_code=422, detail="归档费支付必须使用归档费结算专用入口创建")
    if body.module == "finance":
        raise HTTPException(status_code=422, detail="费用必须使用费用管理专用入口创建")
    if body.module in {"hr", "warehouse"} and identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="当前角色不能新建人事或仓库记录")
    if await db.scalar(select(BusinessRecord.id).where(BusinessRecord.serial_no == body.serial_no)):
        raise HTTPException(status_code=409, detail="业务编号已存在")
    payload = body.model_dump()
    if body.module == "clue":
        payload["status"] = "草稿"
    if body.module == "document":
        direction = str((body.data or {}).get("direction") or "").strip()
        if direction not in {"收文", "发文"}: raise HTTPException(status_code=422, detail="收发类型必须为收文或发文")
        case_no = str((body.data or {}).get("case_no") or "").strip()
        if case_no:
            case_record = await db.scalar(select(BusinessRecord).where(BusinessRecord.module == "case", BusinessRecord.serial_no == case_no))
            if not case_record: raise HTTPException(status_code=422, detail="关联案件不存在")
            await _ensure_record_visible(case_record.id, identity, db)
        payload["status"] = "待登记"
    if body.module == "hr":
        joined_at = str((body.data or {}).get("joined_at") or "").strip()
        if not body.title.strip() or not (body.data or {}).get("position") or not joined_at: raise HTTPException(status_code=422, detail="员工姓名、岗位和入职日期不能为空")
        payload["status"] = "在职" if body.status == "在职" else "试用"
    if body.module == "warehouse":
        data = dict(body.data or {}); quantity = int(data.get("quantity") or 0)
        if data.get("evidence_status"): raise HTTPException(status_code=422, detail="证物必须使用证物登记入口创建")
        if quantity < 1 or not str(data.get("category") or "").strip() or not str(data.get("location") or "").strip(): raise HTTPException(status_code=422, detail="物品类别、数量和存放位置不能为空")
        data.update({"quantity": quantity, "borrower": "", "due_date": "", "borrow_purpose": ""}); payload["data"] = data; payload["status"] = "在库"
    if identity.get("role") != "admin":
        user = await db.scalar(select(User).where(User.username == identity["username"]))
        if not user: raise HTTPException(status_code=401, detail="当前用户不存在")
        payload["department"] = user.department
        if identity.get("role") == "user": payload["owner"] = user.username
    if body.module == "customer":
        managers = await _resolve_active_customer_managers(list((payload.get("data") or {}).get("customer_managers") or [payload.get("owner")]), db)
        owner = (await _resolve_active_customer_managers([payload.get("owner")], db))[0]
        managers = [owner, *[manager for manager in managers if manager != owner]]
        payload["owner"] = owner
        payload["data"] = {**(payload.get("data") or {}), "customer_managers": managers}
    record = BusinessRecord(**payload)
    db.add(record)
    await db.flush()
    db.add(WorkflowEvent(record_id=record.id, action="创建", to_status=record.status, operator=identity["username"], comment="创建业务记录"))
    await db.commit()
    await db.refresh(record)
    return _record_dict(record)


@app.get(f"{settings.api_prefix}/records/{{record_id}}")
async def get_record(record_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    record = await _ensure_record_visible(record_id, identity, db)
    return _record_dict(record, await _allowed_field_keys(identity, db))


@app.get(f"{settings.api_prefix}/records/{{record_id}}/history")
async def record_history(record_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    record = await _ensure_record_visible(record_id, identity, db)
    events = await db.scalars(select(WorkflowEvent).where(WorkflowEvent.record_id == record_id).order_by(WorkflowEvent.created_at.desc(), WorkflowEvent.id.desc()))
    return {
        "transitions": WORKFLOW_TRANSITIONS.get(record.module, {}).get(record.status, []),
        "items": [{"id": e.id, "action": e.action, "from_status": e.from_status, "to_status": e.to_status, "operator": e.operator, "comment": e.comment, "created_at": e.created_at} for e in events],
    }


@app.patch(f"{settings.api_prefix}/records/{{record_id}}")
async def update_record(record_id: int, body: RecordUpdate, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    record = await _ensure_record_visible(record_id, identity, db)
    await _require_record_owner_or_manager(record, identity, db)
    changes = body.model_dump(exclude_unset=True)
    if record.module not in GENERIC_RECORD_EDITABLE_MODULES:
        raise HTTPException(status_code=409, detail="该业务必须使用专用入口办理")
    if record.module in {"hr", "warehouse"} and identity.get("role") not in {"admin", "manager"}:
        raise HTTPException(status_code=403, detail="当前角色不能修改人事或仓库资料")
    if "status" in changes and record.module in {"clue", "evidence", "invoice", "refund", "document", "hr", "warehouse"}:
        raise HTTPException(status_code=409, detail="该业务必须使用专用审批或办理入口变更状态")
    if record.module == "customer" and "status" in changes and changes["status"] != record.status:
        raise HTTPException(status_code=409, detail="客户生命周期状态必须通过领取、释放、回收或恢复专用入口变更")
    if record.module == "warehouse" and "data" in changes:
        if record.status != "在库": raise HTTPException(status_code=409, detail="借出或归还中的物品不能直接修改资料")
        protected = {"borrower", "due_date", "borrow_purpose", "borrowed_at", "borrowed_by", "return_requested_at", "returned_at", "return_condition", "scrapped_at", "scrap_reason", "evidence_status", "checked_in_at", "checked_in_by", "checked_out_at", "checked_out_by", "recipient", "checkout_purpose", "rechecked_in_at", "rechecked_in_by", "destroyed_at", "destroyed_by", "destroy_reason"}
        incoming_data = dict(changes.get("data") or {})
        for key in protected:
            if incoming_data.get(key) != (record.data or {}).get(key): raise HTTPException(status_code=409, detail="借还及报废信息必须通过专用办理入口修改")
    if identity.get("role") != "admin":
        user = await db.scalar(select(User).where(User.username == identity["username"]))
        if not user: raise HTTPException(status_code=401, detail="当前用户不存在")
        if "department" in changes: changes["department"] = user.department
        if identity.get("role") == "user" and "owner" in changes: changes["owner"] = user.username
    if record.module == "customer":
        requested_title = str(changes.get("title", record.title) or "").strip()
        await _ensure_unique_customer_name(requested_title, db, exclude_id=record.id)
        changes["title"] = requested_title
        if "title" in body.model_fields_set:
            changes["customer"] = requested_title
        owner = record.owner
        if "owner" in changes:
            requested_owner = (await _resolve_active_customer_managers([changes["owner"]], db))[0]
            if requested_owner != record.owner:
                raise HTTPException(status_code=409, detail="客户负责人必须通过客户管理人专用入口修改")
            changes["owner"] = record.owner
        if "data" in changes:
            customer_data = dict(changes.get("data") or {})
            existing_customer_data = dict(record.data or {})
            for protected_contact_field in CUSTOMER_SYSTEM_DATA_FIELDS:
                if (
                    protected_contact_field in customer_data
                    and customer_data.get(protected_contact_field) != existing_customer_data.get(protected_contact_field)
                ):
                    raise HTTPException(status_code=409, detail="客户系统维护字段必须通过对应专用入口修改")
                if protected_contact_field in existing_customer_data:
                    customer_data[protected_contact_field] = existing_customer_data[protected_contact_field]
                else:
                    customer_data.pop(protected_contact_field, None)
            existing_managers = [
                str(manager).strip()
                for manager in existing_customer_data.get("customer_managers", [])
                if str(manager).strip()
            ] or [record.owner]
            if "customer_managers" in customer_data:
                incoming_managers = [
                    str(manager).strip()
                    for manager in (customer_data.get("customer_managers") or [])
                    if str(manager).strip()
                ]
                if incoming_managers != existing_managers:
                    raise HTTPException(status_code=409, detail="客户管理人必须通过客户管理人专用入口修改")
            customer_data["customer_managers"] = existing_managers
            changes["data"] = customer_data
    old_status = record.status
    for field, value in changes.items():
        setattr(record, field, value)
    if record.module == "customer":
        _mark_customer_modified(record, identity)
    if "status" in changes and changes["status"] != old_status:
        db.add(WorkflowEvent(record_id=record.id, action="编辑变更", from_status=old_status, to_status=record.status, operator=identity["username"], comment="通过编辑表单变更状态"))
    else:
        db.add(WorkflowEvent(record_id=record.id, action="编辑", from_status=record.status, to_status=record.status, operator=identity["username"], comment="修改业务资料"))
    await db.commit()
    await db.refresh(record)
    return _record_dict(record)


@app.post(f"{settings.api_prefix}/records/{{record_id}}/transition")
async def transition_record(record_id: int, body: TransitionInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    record = await _ensure_record_visible(record_id, identity, db)
    if record.module not in GENERIC_RECORD_TRANSITION_MODULES:
        raise HTTPException(status_code=409, detail="该业务必须使用专用审批或办理入口变更状态")
    approval_access = identity.get("role") == "auditor" and record.module in {"contract", "finance", "invoice", "refund", "seal", "clue", "notary"} and record.status in {"待审批", "审批中", "待审核"}
    if not approval_access:
        await _require_record_owner_or_manager(record, identity, db)
    allowed = WORKFLOW_TRANSITIONS.get(record.module, {}).get(record.status, [])
    if body.to_status not in allowed:
        raise HTTPException(status_code=409, detail=f"不能从“{record.status}”流转到“{body.to_status}”")
    previous = record.status
    record.status = body.to_status
    action = "审批通过"
    if body.to_status in {"已拒绝", "已驳回", "已退回", "已撤回"}:
        action = "驳回/撤回"
    elif body.to_status in {"已完成", "已归档", "已用印", "已付款", "已对账", "已发布"}:
        action = "办结"
    db.add(WorkflowEvent(record_id=record.id, action=action, from_status=previous, to_status=body.to_status, operator=identity["username"], comment=body.comment))
    await db.commit()
    await db.refresh(record)
    return _record_dict(record)


@app.delete(f"{settings.api_prefix}/records/{{record_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_record(record_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    if identity["role"] != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可删除")
    record = await db.get(BusinessRecord, record_id)
    if not record:
        raise HTTPException(status_code=404, detail="记录不存在")
    if record.module not in GENERIC_RECORD_DELETABLE_MODULES:
        raise HTTPException(status_code=409, detail="该业务记录不能通过通用入口物理删除，请使用专用撤销、作废或冲正流程")
    if record.module == "hr":
        linked_username = str((record.data or {}).get("username") or "").strip().lower()
        if linked_username and await db.scalar(select(User.id).where(User.username == linked_username)):
            # A generic record delete must never leave an active login account
            # behind.  Employee exits are handled by the HR edit/disable flow.
            raise HTTPException(status_code=409, detail="该员工档案关联可登录账号，不能直接删除；请在员工资料中停用账号以保持同步")
    if record.module == "seal" and (record.data or {}).get("actual_copies"):
        asset = await db.get(SealAsset, int((record.data or {}).get("seal_asset_id") or 0))
        if asset:
            asset.usage_count = max(0, asset.usage_count - int((record.data or {}).get("actual_copies") or 0))
    attachments = (await db.scalars(select(FileAttachment).where(FileAttachment.record_id == record_id))).all()
    attachment_paths = [Path(item.path) for item in attachments]
    for attachment in attachments:
        await db.delete(attachment)
    await db.execute(delete(FinanceTransaction).where(FinanceTransaction.finance_record_id == record_id))
    await db.execute(delete(ContractApprovalStep).where(ContractApprovalStep.contract_record_id == record_id))
    await db.execute(delete(WorkflowEvent).where(WorkflowEvent.record_id == record_id))
    if record.module == "task":
        await _delete_task_notifications(record_id, db)
    await db.delete(record)
    await db.commit()
    for path in attachment_paths:
        if path.is_file() and UPLOAD_ROOT.resolve() in path.resolve().parents:
            path.unlink()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.delete(
    f"{settings.api_prefix}/testing/cases/{{case_id}}",
    status_code=status.HTTP_204_NO_CONTENT,
    include_in_schema=False,
)
async def delete_smoke_case(case_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Non-production cleanup for records created by the end-to-end smoke suite only."""
    if settings.app_env.strip().lower() in {"production", "prod"}:
        raise HTTPException(status_code=404, detail="接口不存在")
    if identity.get("role") != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可清理本地冒烟案件")
    record = await db.get(BusinessRecord, case_id)
    if not record:
        raise HTTPException(status_code=404, detail="案件不存在")
    if record.module != "case":
        raise HTTPException(status_code=422, detail="该测试清理入口仅支持案件")
    if not (record.serial_no.startswith("SMOKE-") or record.title.startswith("SMOKE")):
        raise HTTPException(status_code=403, detail="只能清理本地冒烟测试案件")
    attachments = list((await db.scalars(select(FileAttachment).where(FileAttachment.record_id == case_id))).all())
    attachment_paths = [Path(item.path) for item in attachments]
    for attachment in attachments:
        await db.delete(attachment)
    related_tasks = list((await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "task", BusinessRecord.data["case_id"].as_integer() == case_id))).all())
    for task in related_tasks:
        await _delete_task_notifications(task.id, db)
        await db.execute(delete(WorkflowEvent).where(WorkflowEvent.record_id == task.id))
        await db.delete(task)
    await db.execute(delete(HearingSchedule).where(HearingSchedule.case_record_id == case_id))
    await db.execute(delete(FinanceTransaction).where(FinanceTransaction.finance_record_id == case_id))
    await db.execute(delete(WorkflowEvent).where(WorkflowEvent.record_id == case_id))
    await db.delete(record)
    await db.commit()
    for path in attachment_paths:
        if path.is_file() and UPLOAD_ROOT.resolve() in path.resolve().parents:
            path.unlink()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.delete(
    f"{settings.api_prefix}/testing/records/{{record_id}}",
    status_code=status.HTTP_204_NO_CONTENT,
    include_in_schema=False,
)
async def delete_smoke_record(record_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Non-production cleanup that cannot be used for ordinary or historical records."""
    if settings.app_env.strip().lower() in {"production", "prod"}:
        raise HTTPException(status_code=404, detail="接口不存在")
    if identity.get("role") != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可清理本地冒烟记录")
    record = await db.get(BusinessRecord, record_id)
    if not record:
        raise HTTPException(status_code=404, detail="记录不存在")
    explicit_test_marker = (
        record.serial_no.startswith("SMOKE-")
        or "SMOKE" in record.title.upper()
        or "冒烟" in record.title
        # UI-driven cross-role acceptance records use this exact, deliberately
        # narrow marker.  It keeps cleanup available for terminal task flows
        # which have no business-page delete action, without admitting normal
        # records that merely contain a generic "验收" label.
        or "UI任务流转验收-" in record.title
        # Contact-edit page evidence uses this equally narrow, fixed UI marker.
        or "UI临时联系人验收-" in record.title
        or "SMOKE" in (record.customer or "").upper()
        or "冒烟" in (record.customer or "")
        or record.owner.lower().startswith("smoke_")
        or "smoke_" in json.dumps(record.data or {}, ensure_ascii=False).lower()
    )
    if not explicit_test_marker:
        raise HTTPException(status_code=403, detail="只能清理带明确测试标识的本地冒烟记录")
    attachments = list((await db.scalars(select(FileAttachment).where(FileAttachment.record_id == record_id))).all())
    attachment_paths = [Path(item.path) for item in attachments]
    for attachment in attachments:
        await db.delete(attachment)
    if record.module == "case":
        related_tasks = list((await db.scalars(select(BusinessRecord).where(BusinessRecord.module == "task", BusinessRecord.data["case_id"].as_integer() == record.id))).all())
        for task in related_tasks:
            await _delete_task_notifications(task.id, db)
            await db.execute(delete(WorkflowEvent).where(WorkflowEvent.record_id == task.id))
            await db.delete(task)
    await db.execute(delete(FinanceTransaction).where(FinanceTransaction.finance_record_id == record_id))
    await db.execute(delete(ContractApprovalStep).where(ContractApprovalStep.contract_record_id == record_id))
    await db.execute(delete(WorkflowEvent).where(WorkflowEvent.record_id == record_id))
    if record.module == "task":
        await _delete_task_notifications(record_id, db)
    await db.delete(record)
    await db.commit()
    for path in attachment_paths:
        if path.is_file() and UPLOAD_ROOT.resolve() in path.resolve().parents:
            path.unlink()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


def _agent_content_hash(content: str) -> str:
    return hashlib.sha256(content.encode("utf-8")).hexdigest()


def _agent_document_dict(item: AgentDocument, template: DocumentTemplate | None = None, record: BusinessRecord | None = None, capabilities: dict | None = None) -> dict:
    return {"id": item.id, "job_no": item.job_no, "template_id": item.template_id, "template_name": template.name if template else "", "record_id": item.record_id, "record_no": record.serial_no if record else "", "record_title": record.title if record else "", "title": item.title, "instruction": item.instruction, "content": item.content, "status": item.status, "content_version": item.content_version, "confirmed_by": item.confirmed_by, "confirmed_at": item.confirmed_at, "conversation_id": item.conversation_id, "dify_message_id": item.dify_message_id, "error": item.error, "creator": item.creator, "created_at": item.created_at, "updated_at": item.updated_at, "capabilities": capabilities or {}}


async def _ensure_agent_document_access(document_id: int, identity: dict, db: AsyncSession, *, write: bool = False) -> tuple[AgentDocument, BusinessRecord | None]:
    item = await db.get(AgentDocument, document_id)
    if not item:
        raise HTTPException(status_code=404, detail="智能文档任务不存在")
    linked_record = await db.get(BusinessRecord, item.record_id) if item.record_id else None
    if item.record_id:
        # A generated document contains a snapshot of the related business
        # record.  Its creator must never retain access after that record is
        # transferred, hidden, or otherwise revoked.  Current record scope is
        # therefore checked before creator status for every module.
        if not linked_record:
            raise HTTPException(status_code=404, detail="关联业务记录不存在或无权访问")
        record = await _ensure_record_visible(item.record_id, identity, db)
        # Customer documents may contain a complete historic customer snapshot.
        # A shared read-only recipient can view the customer record itself, but
        # must not gain access to the generated document or its prompt/content.
        if record.module == "customer":
            await _require_record_owner_or_manager(record, identity, db)
            return item, record
        if write:
            await _require_record_owner_or_manager(record, identity, db)
        return item, record
    if identity.get("role") == "admin" or item.creator == identity["username"]:
        return item, None
    raise HTTPException(status_code=404, detail="智能文档任务不存在或无权访问")


async def _agent_document_capabilities(item: AgentDocument, identity: dict, db: AsyncSession, record: BusinessRecord | None) -> dict:
    """Return UI capabilities after the same checks used by protected APIs."""
    is_creator_or_admin = identity.get("role") == "admin" or item.creator == identity["username"]
    can_write = is_creator_or_admin
    if record:
        try:
            await _require_record_owner_or_manager(record, identity, db)
        except HTTPException:
            can_write = False
        else:
            can_write = True
    can_writeback = bool(can_write and record and item.status == "已人工确认")
    if can_writeback and record and record.module == "case":
        try:
            await _require_case_detail_write_access(record, identity, db)
        except HTTPException:
            can_writeback = False
    has_written_attachment = False
    if record:
        has_written_attachment = bool(await db.scalar(select(FileAttachment.id).where(FileAttachment.record_id == record.id, FileAttachment.remark == f"Dify任务 {item.job_no}")))
    can_delete = bool(can_write and is_creator_or_admin and not item.confirmed_by and not item.confirmed_at and not has_written_attachment)
    return {"can_download": True, "can_edit": can_write, "can_retry": can_write, "can_confirm": can_write, "can_writeback": can_writeback, "can_delete": can_delete}


def _docx_bytes(title: str, content: str) -> bytes:
    document = Document(); document.add_heading(title, level=0)
    for raw in content.splitlines():
        line = raw.strip()
        if not line: continue
        if line.startswith("### "): document.add_heading(line[4:], level=3)
        elif line.startswith("## "): document.add_heading(line[3:], level=2)
        elif line.startswith("# "): document.add_heading(line[2:], level=1)
        elif line.startswith(("- ", "* ")): document.add_paragraph(line[2:], style="List Bullet")
        else: document.add_paragraph(line)
    output = io.BytesIO(); document.save(output); return output.getvalue()


async def _run_document_agent(item: AgentDocument) -> None:
    if not settings.dify_base_url or not settings.dify_api_key:
        item.status = "待配置"; item.error = "Dify 尚未配置；可先编辑系统生成的字段提纲，配置后点击重试。"; return
    item.status = "生成中"; item.error = ""
    payload = {"inputs": {"document_job_no": item.job_no}, "query": item.prompt, "response_mode": "blocking", "conversation_id": item.conversation_id, "user": item.creator}
    try:
        async with httpx.AsyncClient(timeout=120) as client:
            response = await client.post(f"{settings.dify_base_url.rstrip('/')}/v1/chat-messages", headers={"Authorization": f"Bearer {settings.dify_api_key}"}, json=payload)
            response.raise_for_status(); result = response.json()
        item.content = result.get("answer", "").strip(); item.conversation_id = result.get("conversation_id", ""); item.dify_message_id = result.get("message_id", "")
        if not item.content: raise ValueError("Dify 未返回文档内容")
        item.status = "已生成"
    except Exception as exc:
        item.status = "生成失败"; item.error = str(exc)[:1000]


@app.get(f"{settings.api_prefix}/agent/documents")
async def list_agent_documents(identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    items = (await db.scalars(select(AgentDocument).order_by(AgentDocument.created_at.desc()).limit(100))).all()
    accessible_items: list[tuple[AgentDocument, BusinessRecord | None]] = []
    for item in items:
        try:
            _, record = await _ensure_agent_document_access(item.id, identity, db)
        except HTTPException:
            continue
        accessible_items.append((item, record))
    template_ids = {item.template_id for item, _ in accessible_items}; record_ids = {item.record_id for item, _ in accessible_items if item.record_id}
    templates = {x.id: x for x in (await db.scalars(select(DocumentTemplate).where(DocumentTemplate.id.in_(template_ids)))).all()} if template_ids else {}
    records = {x.id: x for x in (await db.scalars(select(BusinessRecord).where(BusinessRecord.id.in_(record_ids)))).all()} if record_ids else {}
    result = []
    for item, visible_record in accessible_items:
        record = visible_record or records.get(item.record_id)
        capabilities = await _agent_document_capabilities(item, identity, db, record)
        result.append(_agent_document_dict(item, templates.get(item.template_id), record, capabilities))
    return {"items": result, "dify_configured": bool(settings.dify_base_url and settings.dify_api_key)}


@app.post(f"{settings.api_prefix}/agent/documents", status_code=status.HTTP_201_CREATED)
async def create_agent_document(body: AgentDocumentInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    template = await db.get(DocumentTemplate, body.template_id)
    if not template or not template.is_active: raise HTTPException(status_code=404, detail="文书模板不存在或已停用")
    record = await _ensure_record_visible(body.record_id, identity, db) if body.record_id else None
    if record and record.module == "customer":
        await _require_record_owner_or_manager(record, identity, db)
    context = {"模板": template.name, "模板分类": template.category, "要求字段": template.fields, "用户要求": body.instruction}
    if record:
        safe_record = _record_dict(record, await _allowed_field_keys(identity, db))
        context["业务数据"] = {"编号": safe_record["serial_no"], "标题": safe_record["title"], "客户": safe_record["customer"], "负责人": safe_record["owner"], "部门": safe_record["department"], "说明": safe_record["description"], "扩展字段": safe_record["data"]}
    prompt = "请根据以下结构化信息生成正式、严谨、可直接审核的中文法律文书。不得虚构未提供的事实，对缺失信息使用【待补充】标记。\n" + json.dumps(context, ensure_ascii=False, indent=2)
    outline = "\n".join([f"## {field}\n【待补充】" for field in template.fields]) or "## 正文\n【待补充】"
    item = AgentDocument(job_no=f"AI{datetime.now().strftime('%Y%m%d%H%M%S')}{uuid4().hex[:4].upper()}", template_id=template.id, record_id=record.id if record else None, title=body.title.strip(), instruction=body.instruction.strip(), prompt=prompt, content=outline, status="等待生成", creator=identity["username"])
    db.add(item); await db.flush(); await _run_document_agent(item); await db.commit(); await db.refresh(item)
    if record: db.add(WorkflowEvent(record_id=record.id, action="创建智能文档", from_status=record.status, to_status=record.status, operator=identity["username"], comment=f"{item.job_no}｜{template.name}")); await db.commit()
    return _agent_document_dict(item, template, record)


@app.post(f"{settings.api_prefix}/agent/documents/{{document_id}}/retry")
async def retry_agent_document(document_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item, _ = await _ensure_agent_document_access(document_id, identity, db, write=True)
    if item.status == "生成中": raise HTTPException(status_code=409, detail="文档正在生成中")
    item.confirmed_by = ""; item.confirmed_at = None; item.confirmed_content_hash = ""
    item.content_version = int(item.content_version or 1) + 1
    await _run_document_agent(item); await db.commit(); await db.refresh(item)
    return _agent_document_dict(item)


@app.patch(f"{settings.api_prefix}/agent/documents/{{document_id}}")
async def update_agent_document(document_id: int, body: AgentDocumentUpdate, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item, _ = await _ensure_agent_document_access(document_id, identity, db, write=True)
    if body.title is not None: item.title = body.title.strip()
    if body.content is not None:
        item.content = body.content; item.status = "已编辑"
        item.content_version = int(item.content_version or 1) + 1
        item.confirmed_by = ""; item.confirmed_at = None; item.confirmed_content_hash = ""
    await db.commit(); await db.refresh(item); return _agent_document_dict(item)


@app.post(f"{settings.api_prefix}/agent/documents/{{document_id}}/confirm")
async def confirm_agent_document(document_id: int, body: AgentDocumentConfirmInput, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item, record = await _ensure_agent_document_access(document_id, identity, db, write=True)
    if item.status not in {"已生成", "已编辑", "已人工确认"}:
        raise HTTPException(status_code=409, detail="文档生成完成并经人工检查后才能确认")
    if not item.content.strip():
        raise HTTPException(status_code=409, detail="空文档不能确认")
    item.status = "已人工确认"
    item.confirmed_by = identity["username"]
    item.confirmed_at = datetime.now(timezone.utc)
    item.confirmed_content_hash = _agent_content_hash(item.content)
    if record:
        db.add(WorkflowEvent(record_id=record.id, action="人工确认智能文档", from_status=record.status, to_status=record.status, operator=identity["username"], comment=f"{item.job_no}｜版本 {item.content_version}。{body.comment}"))
    await db.commit(); await db.refresh(item)
    return _agent_document_dict(item, record=record)


@app.get(f"{settings.api_prefix}/agent/documents/{{document_id}}/download")
async def download_agent_document(document_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item, _ = await _ensure_agent_document_access(document_id, identity, db)
    content = _docx_bytes(item.title, item.content)
    return StreamingResponse(io.BytesIO(content), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="{item.job_no}.docx"'})


@app.post(f"{settings.api_prefix}/agent/documents/{{document_id}}/writeback")
async def writeback_agent_document(document_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item, record = await _ensure_agent_document_access(document_id, identity, db, write=True)
    if not item.record_id: raise HTTPException(status_code=409, detail="文档未关联业务记录，不能回写")
    if not record: raise HTTPException(status_code=404, detail="关联业务记录已不存在")
    if record.module == "case":
        await _require_case_detail_write_access(record, identity, db)
    if item.status != "已人工确认" or not item.confirmed_by or not item.confirmed_at:
        raise HTTPException(status_code=409, detail="智能文档必须先由人工审核确认，才能回写业务附件")
    if item.confirmed_content_hash != _agent_content_hash(item.content):
        raise HTTPException(status_code=409, detail="文档内容在确认后已变化，请重新人工确认")
    existing = await db.scalar(select(FileAttachment).where(FileAttachment.record_id == record.id, FileAttachment.remark == f"Dify任务 {item.job_no}"))
    if existing:
        raise HTTPException(status_code=409, detail=f"该智能文档已经回写为附件 {existing.original_name}，请勿重复操作")
    content = _docx_bytes(item.title, item.content); stored_name = f"{uuid4().hex}.docx"; path = UPLOAD_ROOT / stored_name; path.write_bytes(content)
    attachment = FileAttachment(record_id=record.id, category="智能生成文书", original_name=f"{item.title}.docx", stored_name=stored_name, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", size=len(content), path=str(path), uploader=identity["username"], remark=f"Dify任务 {item.job_no}")
    db.add(attachment); db.add(WorkflowEvent(record_id=record.id, action="智能文档回写", from_status=record.status, to_status=record.status, operator=identity["username"], comment=f"{item.job_no}｜{item.title}")); await db.commit(); await db.refresh(attachment)
    return {"attachment_id": attachment.id, "record_id": record.id, "filename": attachment.original_name}


@app.delete(f"{settings.api_prefix}/agent/documents/{{document_id}}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_agent_document(document_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    item, record = await _ensure_agent_document_access(document_id, identity, db, write=True)
    if identity.get("role") != "admin" and item.creator != identity["username"]: raise HTTPException(status_code=403, detail="只能删除本人创建的智能文档任务")
    if item.confirmed_by or item.confirmed_at or item.status == "已人工确认":
        raise HTTPException(status_code=409, detail="已人工确认的智能文档不得删除，请保留审核与回写审计记录")
    if record:
        written_attachment = await db.scalar(select(FileAttachment).where(FileAttachment.record_id == record.id, FileAttachment.remark == f"Dify任务 {item.job_no}"))
        if written_attachment:
            raise HTTPException(status_code=409, detail="已回写业务附件的智能文档不得删除，请保留附件与审计记录")
        db.add(WorkflowEvent(record_id=record.id, action="删除智能文档任务", from_status=record.status, to_status=record.status, operator=identity["username"], comment=f"{item.job_no}｜{item.title}"))
    await db.delete(item); await db.commit(); return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.delete(f"{settings.api_prefix}/testing/agent-documents/{{document_id}}", status_code=status.HTTP_204_NO_CONTENT, include_in_schema=False)
async def delete_smoke_agent_document(document_id: int, identity: dict = Depends(current_identity), db: AsyncSession = Depends(get_db)):
    """Remove explicit smoke-only AI jobs when normal audit retention blocks deletion."""
    if settings.app_env.strip().lower() in {"production", "prod"}:
        raise HTTPException(status_code=404, detail="接口不存在")
    if identity.get("role") != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可清理本地冒烟智能文档")
    item = await db.get(AgentDocument, document_id)
    if not item:
        raise HTTPException(status_code=404, detail="智能文档任务不存在")
    if "SMOKE" not in item.title.upper() and "冒烟" not in item.title:
        raise HTTPException(status_code=403, detail="只能清理带明确测试标识的本地智能文档")
    await db.delete(item)
    await db.commit()
    return Response(status_code=status.HTTP_204_NO_CONTENT)


@app.post(f"{settings.api_prefix}/agent/chat")
async def agent_chat(body: DifyRequest, identity: dict = Depends(current_identity)):
    if not settings.dify_base_url or not settings.dify_api_key:
        raise HTTPException(status_code=503, detail="Dify 尚未配置")
    payload = {"inputs": {"operator": identity["username"]}, "query": body.query, "response_mode": "blocking", "user": identity["username"]}
    if body.conversation_id:
        payload["conversation_id"] = body.conversation_id
    async with httpx.AsyncClient(timeout=120) as client:
        response = await client.post(f"{settings.dify_base_url.rstrip('/')}/v1/chat-messages", headers={"Authorization": f"Bearer {settings.dify_api_key}"}, json=payload)
    if response.is_error:
        raise HTTPException(status_code=502, detail="Dify 调用失败")
    return response.json()
