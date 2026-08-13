from __future__ import annotations

import base64
import io
from dataclasses import dataclass
from functools import lru_cache
from pathlib import Path

import pypdfium2 as pdfium
from docx import Document


MAX_TEXT_CHARS = 36_000
MAX_PDF_PAGES = 40
MAX_VISUAL_PAGES = 4
IMAGE_SUFFIXES = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".webp": "image/webp"}
TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".csv", ".json", ".xml", ".html", ".htm"}


@dataclass(frozen=True)
class AttachmentReading:
    status: str
    text: str
    page_count: int
    images: tuple[dict[str, str], ...] = ()


def _clip(value: str, limit: int = MAX_TEXT_CHARS) -> str:
    normalized = "\n".join(line.rstrip() for line in value.replace("\r\n", "\n").replace("\r", "\n").split("\n")).strip()
    return normalized if len(normalized) <= limit else normalized[:limit] + "\n[正文过长，已截取前部内容]"


def _image_data_url(content: bytes, mime_type: str) -> str:
    return f"data:{mime_type};base64,{base64.b64encode(content).decode('ascii')}"


def _read_docx(path: Path) -> AttachmentReading:
    document = Document(path)
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    text = _clip("\n".join(lines))
    return AttachmentReading(status="parsed" if text else "empty", text=text, page_count=0)


def _read_text(path: Path) -> AttachmentReading:
    content = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            text = content.decode(encoding)
            return AttachmentReading(status="parsed", text=_clip(text), page_count=0)
        except UnicodeDecodeError:
            continue
    return AttachmentReading(status="unsupported_encoding", text="", page_count=0)


def _render_pdf_page(page: object, page_number: int) -> dict[str, str]:
    bitmap = page.render(scale=1.35)
    image = bitmap.to_pil().convert("RGB")
    buffer = io.BytesIO()
    image.save(buffer, format="JPEG", quality=82, optimize=True)
    return {"page": str(page_number), "mime_type": "image/jpeg", "data_url": _image_data_url(buffer.getvalue(), "image/jpeg")}


def _read_pdf(path: Path) -> AttachmentReading:
    document = pdfium.PdfDocument(path)
    page_count = len(document)
    text_parts: list[str] = []
    images: list[dict[str, str]] = []
    try:
        for index in range(min(page_count, MAX_PDF_PAGES)):
            page = document[index]
            try:
                text_page = page.get_textpage()
                try:
                    page_text = (text_page.get_text_range() or "").strip()
                finally:
                    text_page.close()
                if page_text:
                    text_parts.append(f"[第 {index + 1} 页]\n{page_text}")
                if len(page_text) < 40 and len(images) < MAX_VISUAL_PAGES:
                    images.append(_render_pdf_page(page, index + 1))
            finally:
                page.close()
            if sum(len(item) for item in text_parts) >= MAX_TEXT_CHARS:
                break
    finally:
        document.close()
    text = _clip("\n\n".join(text_parts))
    status = "parsed" if text else ("visual" if images else "empty")
    return AttachmentReading(status=status, text=text, page_count=page_count, images=tuple(images))


@lru_cache(maxsize=256)
def _read_cached(path_text: str, name: str, size: int, modified_ns: int) -> AttachmentReading:
    del size, modified_ns
    path = Path(path_text)
    suffix = Path(name or path.name).suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix in TEXT_SUFFIXES:
        return _read_text(path)
    if suffix in IMAGE_SUFFIXES:
        content = path.read_bytes()
        return AttachmentReading(
            status="visual",
            text="",
            page_count=1,
            images=({"page": "1", "mime_type": IMAGE_SUFFIXES[suffix], "data_url": _image_data_url(content, IMAGE_SUFFIXES[suffix])},),
        )
    return AttachmentReading(status="unsupported", text="", page_count=0)


def read_attachment(path: Path, name: str) -> AttachmentReading:
    stat = path.stat()
    return _read_cached(str(path.resolve()), name, stat.st_size, stat.st_mtime_ns)
