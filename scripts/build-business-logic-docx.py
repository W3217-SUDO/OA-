from __future__ import annotations

import csv
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Administrator\Desktop\系统重构")
ARCHIVE = Path(r"D:\OA系统源码归档_20260717")
OUT = ARCHIVE / "原OA系统业务逻辑分析报告_20260717.docx"
FLOW_IMAGE = ROOT / ".tmp-docx" / "oa-business-flow.png"
RULES_CSV = ROOT / "docs" / "原OA业务规则证据矩阵_20260717.csv"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "6B7280"
GOLD = "B58728"
RED = "9B1C1C"
GREEN = "276749"
WHITE = "FFFFFF"
BLACK = "111827"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=11, bold=None, color=BLACK, font="Microsoft YaHei", italic=None) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_end])
    set_font(run, size=9, color=MID_GRAY)


def add_text(doc, text: str, bold=False, color=BLACK, size=11, after=6, align=None, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullet(doc, text: str, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_font(r)
    return p


def add_number(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_font(r)
    return p


def add_callout(doc, title: str, body: str, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_row_cant_split(table.rows[0])
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=11, bold=True, color=color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body)
    set_font(r2, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_simple_table(doc, headers, rows, widths, header_fill=PALE_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, size=font_size, bold=True, color=NAVY)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if i == 0 and len(headers) <= 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_font(r, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def make_flow_image() -> None:
    FLOW_IMAGE.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1800, 1050), "white")
    draw = ImageDraw.Draw(image)
    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    bold_path = Path(r"C:\Windows\Fonts\msyhbd.ttc")
    font = ImageFont.truetype(str(font_path), 33)
    small = ImageFont.truetype(str(font_path), 27)
    bold = ImageFont.truetype(str(bold_path if bold_path.exists() else font_path), 38)

    def pc(color: str) -> str:
        return color if color.startswith("#") else f"#{color}"

    def box(x, y, w, h, text, fill, outline=BLUE, text_color=NAVY, f=font):
        draw.rounded_rectangle((x, y, x + w, y + h), radius=18, fill=pc(fill), outline=pc(outline), width=4)
        bbox = draw.multiline_textbbox((0, 0), text, font=f, spacing=4, align="center")
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        draw.multiline_text((x + (w - tw) / 2, y + (h - th) / 2), text, font=f, fill=pc(text_color), spacing=4, align="center")

    def arrow(x1, y1, x2, y2, color=BLUE, width=5):
        draw.line((x1, y1, x2, y2), fill=pc(color), width=width)
        import math
        angle = math.atan2(y2 - y1, x2 - x1)
        size = 16
        points = [(x2, y2), (x2 - size * math.cos(angle - 0.55), y2 - size * math.sin(angle - 0.55)), (x2 - size * math.cos(angle + 0.55), y2 - size * math.sin(angle + 0.55))]
        draw.polygon(points, fill=pc(color))

    draw.text((70, 45), "原 OA 系统业务主链", font=bold, fill=pc(NAVY))
    draw.text((70, 102), "客户、合同、案件是主轴；任务、调查、用印、财务和归档是受控生命周期", font=small, fill=pc(MID_GRAY))

    boxes = {
        "customer": (90, 220, 230, 105, "客户", "E8F1FA"),
        "contract": (390, 220, 230, 105, "合同", "E8F1FA"),
        "case": (745, 220, 250, 105, "案件", "DCEAF7"),
        "delivery": (1100, 220, 260, 105, "任务 / 文件\n排期 / 阶段", "E8F1FA"),
        "finance": (1450, 220, 260, 105, "费用 / 发票\n付款 / 结算", "FFF4DB"),
        "invest": (390, 470, 280, 120, "调查项目 / 任务", "EAF7F0"),
        "clue": (745, 470, 280, 120, "线索审核 / 取证", "EAF7F0"),
        "seal": (1090, 470, 290, 120, "用印申请 / 审批\n（独立模块）", "F3EAF8"),
        "archive": (1450, 470, 260, 120, "案件归档\n归档结算", "FFF4DB"),
        "govern": (390, 760, 990, 125, "员工 / 部门 / 角色 / 权限 / 审计日志", "F2F4F7"),
    }
    for value in boxes.values():
        box(*value)
    arrow(320, 272, 390, 272)
    arrow(620, 272, 745, 272)
    arrow(995, 272, 1100, 272)
    arrow(1360, 272, 1450, 272)
    arrow(505, 325, 505, 470)
    arrow(670, 530, 745, 530)
    arrow(885, 470, 885, 325)
    arrow(870, 325, 1190, 470, color="7B61A8")
    arrow(1580, 325, 1580, 470, color=GOLD)
    arrow(1450, 530, 1380, 530, color=GOLD)
    arrow(885, 760, 885, 590, color=MID_GRAY)
    draw.text((98, 935), "实线表示主要业务流；用印可关联合同或案件，但保持独立编号、审批、附件和结果。", font=small, fill=pc(MID_GRAY))
    image.save(FLOW_IMAGE)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.167

    if "Evidence Label" not in styles:
        style = styles.add_style("Evidence Label", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(9.5)
        style.font.bold = True
        style.paragraph_format.space_after = Pt(2)


def add_header_footer(section, first=False):
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = first
    if not first:
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        r = hp.add_run("原 OA 系统业务逻辑分析报告")
        set_font(r, size=9, color=MID_GRAY)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    r = fp.add_run("第 ")
    set_font(r, size=9, color=MID_GRAY)
    add_page_number(fp)
    r2 = fp.add_run(" 页")
    set_font(r2, size=9, color=MID_GRAY)


def build() -> None:
    make_flow_image()
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_header_footer(section, first=True)

    add_text(doc, "业务系统重构 · 原系统证据盘点", bold=True, color=GOLD, size=10.5, after=12)
    add_text(doc, "原 OA 系统业务逻辑分析报告", bold=True, color=NAVY, size=27, after=7)
    add_text(doc, "结合页面、源码、业务程序集与数据库状态的业务推断", color=DARK_BLUE, size=14, after=22)

    metadata = [
        ("报告日期", "2026年7月17日"),
        ("分析范围", "客户、合同、案件、任务、调查、用印、财务、人事、仓库与审计"),
        ("证据规模", "9,541个源码文件；253张表；1,349个MVC动作；20,289个程序集方法"),
        ("使用目的", "作为 FastAPI、React、微信小程序和后续智能体接口的业务边界依据"),
    ]
    add_simple_table(doc, ["项目", "内容"], metadata, [2100, 7260], header_fill=LIGHT_GRAY, font_size=10.2)
    add_callout(doc, "核心判断", "原系统不是菜单集合，而是由客户、合同、案件三类主对象驱动，由任务、调查、用印、财务、归档等独立状态机共同约束的法律服务交付系统。", color=NAVY)
    add_text(doc, "证据说明：A级=源码直接证明；B级=多项证据支持的强推断；C级=仍需原服务或真实样本继续验证。", color=MID_GRAY, size=9.5, after=0, italic=True)

    doc.add_page_break()
    doc.add_heading("执行摘要", level=1)
    for item in (
        "客户是业务主体根节点；合同定义商业授权与收费边界；案件是法律服务交付主对象。",
        "调查线索可以先于案件存在，经过内部审核、客户审核和取证后，再生成或并入案件。",
        "用印是独立模块，拥有独立编号、审批、附件、打印和用印结果，只能选择关联合同或案件。",
        "案件归档受财务到账约束；归档完成并不等同财务结算完成。",
        "任务完成与验收是两个状态；交接、拒绝、撤回、停止、重启都必须保留节点和审计记录。",
        "当前已整理33条关键规则：27条A级直接证据、4条B级强推断、2条C级待验证。",
    ):
        add_bullet(doc, item)

    doc.add_heading("整体业务流程", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FLOW_IMAGE), width=Inches(6.35))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run("图1  原 OA 系统业务主链")
    set_font(r, size=9.5, color=MID_GRAY, italic=True)

    doc.add_heading("一、客户与合同", level=1)
    doc.add_heading("1. 客户生命周期", level=2)
    add_text(doc, "客户经历建档、分配、联系、共享、签约、回收和公海拾回。删除是业务停用而非物理删除；共享增加协作范围，但不改变客户所有权。")
    add_simple_table(doc, ["直接证据", "业务含义"], [
        ("CustomerDelete / CustomerRestore", "客户删除后可以恢复，应采用回收站生命周期。"),
        ("CustomerOpen / CustomerClose", "公海客户可以被重新拾回并重新分配负责人。"),
        ("CustomerOwner / BusinessOwner / DepartmentId", "客户可见范围同时受负责人、业务负责人和组织范围影响。"),
        ("P_CRM_Customer_No_Create", "客户编号必须由服务端按公司和年度原子生成。"),
    ], [3000, 6360])
    doc.add_heading("2. 合同生命周期", level=2)
    add_text(doc, "合同主流程为：草稿 → 待提交 → 审批中 → 审批通过或驳回 → 履行 → 归档。")
    for item in (
        "合同必须关联有效客户。",
        "审批和归档必须使用专用动作，不能由通用状态修改接口绕过。",
        "归档合同不允许新建案件。",
        "开票要求合同存在并具有外部合同号。",
        "合同变更应保留审核轮次和历史版本。",
    ):
        add_bullet(doc, item)
    add_callout(doc, "硬规则（A级）", "原控制器明确返回“合同已归档，不允许新建案子”；发票控制器会拒绝不存在或没有外部合同号的合同。", color=RED)

    doc.add_heading("二、案件中心", level=1)
    add_text(doc, "案件有两条主要入口：从有效合同建案，或由审核后的调查线索生成/并入案件。建案后按案件类型逐步补录当事人、司法机关、承办律师、阶段、排期、文件和费用。")
    add_simple_table(doc, ["规则", "等级", "业务约束"], [
        ("未找到合同或客户时不能建案", "A", "建案命令必须重新读取并验证上游对象。"),
        ("线索已生成案件时不能重复生成", "A", "线索转案件必须幂等。"),
        ("合并目标案件必须存在", "A", "禁止输入任意案号直接合并。"),
        ("不同客户的案件禁止合并", "A", "合并前校验客户主体一致。"),
        ("案件费用未到账不能提交归档", "A", "归档动作必须读取实时财务状态。"),
        ("归档包含内部审核和正式审核", "A", "使用两级审批状态机。"),
    ], [4300, 900, 4160])
    doc.add_heading("案件信息为什么分阶段维护", level=2)
    for item in (
        "建案：客户、合同、案件类型、权利类型和负责人。",
        "当事人：根据刑事、行政、仲裁、顾问等案件类型显示不同字段。",
        "司法机关：公安、检察院、法院及一审、二审、再审、执行分支。",
        "交付：律师分配、任务、排期、文件、执行阶段和提醒日期。",
        "收尾：结案、归档申请、内部审核、正式审核和档案号。",
    ):
        add_number(doc, item)

    doc.add_heading("三、协同任务", level=1)
    add_text(doc, "任务由发起人、当前负责人、最初负责人和协作人共同构成；任务节点用于表达一次接收、分工或交接，不能在交接时覆盖历史负责人。")
    add_simple_table(doc, ["状态/动作", "业务含义"], [
        ("未处理 → 进行中", "负责人接受任务。"),
        ("进行中 → 已完成待验收", "负责人声明完成，但尚未最终验收。"),
        ("已完成待验收 → 已验收", "发起人或授权人确认完成。"),
        ("拒绝 / 停止 / 撤回", "不同操作者和原因形成的旁路终态。"),
        ("重新打开 / 重启", "合法终态回到进行中，并保留原节点历史。"),
    ], [2900, 6460])
    for item in (
        "创建任务必须指定负责人。",
        "结束时间必须晚于开始时间。",
        "完成与验收是两个不同状态。",
        "程序集和数据访问层存在 GetFinishedTaskListAfter5Days，SQL 查询完成状态且完成时间超过指定天数的任务。",
    ):
        add_bullet(doc, item)
    add_callout(doc, "仍需验证", "“交接满5天未开始自动完成”和“任务期限超过30天拒绝”是当前重构必须保留的规则，但原 Windows Service 主程序尚未归档，准确触发条件不能伪装成已由原源码完全证明。", color=RED)

    doc.add_heading("四、调查、线索、取证与证物", level=1)
    add_text(doc, "调查流程为：合同 → 调查项目 → 调查任务 → 调查线索 → 内部审核 → 客户审核（可配置）→ 取证/公证/证物 → 生成或并入案件。")
    for item in (
        "创建调查项目必须存在合同。",
        "线索同时保存内部审核人和客户审核人，证明两级审核相互独立。",
        "并案目标案号必须存在；线索不能重复生成案件。",
        "导入取证文件时必须匹配已有线索。",
        "证据具有入库、出库、重新入库、销毁和领取状态，不能当作普通附件删除。",
    ):
        add_bullet(doc, item)
    add_simple_table(doc, ["对象", "关键状态", "推断"], [
        ("线索", "待提交、待审核、待客户审核、通过、拒绝、撤回、终止、已取证", "审核是质量闸门，不是形式步骤。"),
        ("证据", "未入库、已入库、已出库、重新入库、销毁", "实物证物需要不可逆流转审计。"),
        ("公证", "公证机构、公证号、取得日期", "公证结果与证据关联但不是同一状态。"),
    ], [1800, 3900, 3660])

    doc.add_heading("五、用印管理（独立模块）", level=1)
    add_callout(doc, "确认结论", "用印拥有独立主表、审核表和文件表；拥有独立编号、审核轮次、印章类型、电子章、线下打印、打印份数、打印人和打印时间。它可以关联合同或案件，但不属于调查流程。", color=NAVY)
    add_simple_table(doc, ["阶段", "允许动作"], [
        ("待提交", "编辑、提交、撤回。"),
        ("待审核", "审核通过、审核拒绝、撤回。"),
        ("待用印", "打印/盖章、上传盖章后文件、终止。"),
        ("已用印", "查看、下载、审计追溯。"),
        ("拒绝/撤回/终止", "保留历史；是否允许重提需按状态规则处理。"),
    ], [2200, 7160])
    add_text(doc, "真实状态分布中已用印786条、撤回678条。由此强推断撤回是高频正常流程，不能用删除申请代替。", color=DARK_BLUE, bold=True)

    doc.add_heading("六、财务闭环", level=1)
    add_text(doc, "财务至少包含应收/回款、案件费用、内部费用/应付、发票和结算五类对象。它们状态不同、责任人不同、审核路径不同，不能合并成一张通用费用表。")
    add_simple_table(doc, ["流程", "关键业务规则"], [
        ("回款分配", "银行回款先进入总金额，再分配到案件和费用明细；多案件复杂分配必须由财务处理。"),
        ("内部请款", "创建、提交、审批、付款打包、核销、撤销和回滚是不同动作。"),
        ("发票", "实际开票金额为零不能提交；撤回申请与作废发票是两个状态。"),
        ("一般结算", "待结算、待审核、待付款、已付款、拒绝和回退。"),
        ("归档结算", "案件归档后另行处理归档金额，不等同一般回款结算。"),
    ], [2200, 7160])
    add_callout(doc, "跨模块阻断", "案件费用未到账会阻断案件归档。这说明财务状态是案件状态转换的前置条件，不能依赖前端提示。", color=RED)

    doc.add_heading("七、权限与审计", level=1)
    add_text(doc, "权限至少由账号状态、菜单权限、公司/部门范围、记录级关系和当前业务状态五层共同决定。按钮是否显示只是表现层，后端每个命令都必须重新校验。")
    for item in (
        "admin 是最高权限：全部菜单、全所数据、全部字段。",
        "普通用户的数据范围来自负责人、经办人、协作人、审核人和共享人关系。",
        "审批、撤回、拒绝、重开、付款、用印、归档等动作必须记录操作人、时间、意见和前后状态。",
        "系统日志约4.41亿行；重构应拆分业务审计、技术运行和敏感访问日志，并设置不同保留策略。",
    ):
        add_bullet(doc, item)

    doc.add_heading("八、重构落地原则", level=1)
    for item in (
        "以客户、合同、案件、任务、调查、用印、财务为聚合根，避免跨聚合任意更新。",
        "使用 submit/approve/reject/reopen/archive/mark-paid 等语义化命令，禁止任意状态字符串 PATCH。",
        "每个命令在同一事务中完成权限校验、状态校验、数据写入和审计记录。",
        "PostgreSQL 保存业务关系和元数据，MinIO 保存附件对象；附件必须保留校验和、业务编号、文件类型和版本关系。",
        "React 网页、微信小程序和 Dify 智能体共用同一 REST API；智能体只能调用授权业务工具，不能直接写数据库。",
    ):
        add_number(doc, item)

    doc.add_page_break()
    doc.add_heading("九、待继续验证的高风险点", level=1)
    risks = [
        ("任务自动规则", "交接5天与30天限制的准确状态条件、定时频率和操作者记录。"),
        ("案件类型分支", "各类案件的必填当事人和司法机关字段。"),
        ("合同与收费", "收费方式、合同金额与应收生成的精确公式。"),
        ("结算比例", "代理费、归档费和绩效金额的精确计算。"),
        ("权限矩阵", "各角色在页面、字段、导出、下载和审批上的精确权限。"),
        ("用印与证物", "用印终止/重提规则、证物销毁是否需要额外审批或双人确认。"),
    ]
    add_simple_table(doc, ["风险领域", "仍需取得的证据"], risks, [2500, 6860], font_size=8.7)

    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    landscape.orientation = WD_ORIENT.LANDSCAPE
    landscape.page_width = Inches(11)
    landscape.page_height = Inches(8.5)
    landscape.top_margin = Inches(0.7)
    landscape.bottom_margin = Inches(0.7)
    landscape.left_margin = Inches(0.5)
    landscape.right_margin = Inches(0.5)
    add_header_footer(landscape)
    doc.add_heading("附录：业务规则证据矩阵", level=1)
    add_text(doc, "A级为直接证明，B级为强推断，C级为待验证。该矩阵用于后续接口设计、状态机实现和逐页面验收。", size=10, color=MID_GRAY)
    with RULES_CSV.open("r", encoding="utf-8-sig", newline="") as handle:
        rules = list(csv.DictReader(handle))
    rows = [(r["模块"], r["业务规则"], r["证据等级"], r["直接依据"], r["重构约束"]) for r in rules]
    table = add_simple_table(doc, ["模块", "业务规则", "等级", "直接依据", "重构约束"], rows, [1100, 3300, 700, 3000, 6300], font_size=8.2)
    for row in table.rows[1:]:
        level = row.cells[2].text.strip()
        if level == "A":
            set_cell_shading(row.cells[2], "E6F4EA")
        elif level == "B":
            set_cell_shading(row.cells[2], "FFF4D6")
        else:
            set_cell_shading(row.cells[2], "FDE8E8")
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.core_properties.title = "原 OA 系统业务逻辑分析报告"
    doc.core_properties.subject = "页面、源码、程序集与数据库交叉分析"
    doc.core_properties.author = "系统重构项目组"
    doc.core_properties.keywords = "OA,业务逻辑,案件管理,合同,用印,财务,重构"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
