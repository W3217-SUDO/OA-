from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Administrator\Desktop\系统重构")
ARCHIVE = Path(r"D:\OA系统源码归档_20260717")
OUT = ARCHIVE / "律所管理系统业务蓝图确认稿_20260717.docx"
FLOW_IMAGE = ROOT / ".tmp-docx" / "partner-business-map.png"

# standard_business_brief preset + named Chinese-font override.
NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "6B7280"
GOLD = "B58728"
PALE_GOLD = "FFF4DB"
RED = "9B1C1C"
PALE_RED = "FDECEC"
GREEN = "276749"
PALE_GREEN = "EAF7F0"
WHITE = "FFFFFF"
BLACK = "111827"
FONT = "Microsoft YaHei"


def set_font(run, size=11, bold=None, color=BLACK, italic=None):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), FONT)
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def row_cant_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def repeat_header(row):
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(el)


def table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)
    for row in table.rows:
        row_cant_split(row)
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_font(run, size=9, color=MID_GRAY)


def add_para(doc, text, *, size=11, bold=False, color=BLACK, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    set_font(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    set_font(p.add_run(text))
    return p


def add_callout(doc, title, body, *, fill=PALE_BLUE, title_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(title), bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    set_font(p2.add_run(body), size=10.5)
    add_para(doc, "", after=1)


def add_table(doc, headers, rows, widths, *, font_size=9.4, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(header), size=font_size, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_font(p.add_run(str(value)), size=font_size)
    table_geometry(table, widths)
    add_para(doc, "", after=1)
    return table


def page_break(doc):
    doc.add_page_break()


def make_business_map():
    FLOW_IMAGE.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1800, 1050), "white")
    draw = ImageDraw.Draw(image)
    regular_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    bold_path = Path(r"C:\Windows\Fonts\msyhbd.ttc")
    regular = ImageFont.truetype(str(regular_path), 30)
    small = ImageFont.truetype(str(regular_path), 25)
    bold = ImageFont.truetype(str(bold_path if bold_path.exists() else regular_path), 42)

    def pc(c):
        return c if c.startswith("#") else f"#{c}"

    def box(x, y, w, h, text, fill, outline=BLUE, text_color=NAVY, font=regular):
        draw.rounded_rectangle((x, y, x+w, y+h), radius=20, fill=pc(fill), outline=pc(outline), width=4)
        b = draw.multiline_textbbox((0, 0), text, font=font, spacing=5, align="center")
        tw, th = b[2]-b[0], b[3]-b[1]
        draw.multiline_text((x+(w-tw)/2, y+(h-th)/2), text, font=font, fill=pc(text_color), spacing=5, align="center")

    def arrow(x1, y1, x2, y2, color=BLUE):
        draw.line((x1, y1, x2, y2), fill=pc(color), width=5)
        import math
        a = math.atan2(y2-y1, x2-x1)
        s = 18
        pts = [(x2, y2), (x2-s*math.cos(a-.55), y2-s*math.sin(a-.55)), (x2-s*math.cos(a+.55), y2-s*math.sin(a+.55))]
        draw.polygon(pts, fill=pc(color))

    draw.text((70, 45), "律所经营与案件交付主链", font=bold, fill=pc(NAVY))
    draw.text((70, 108), "每一笔业务都要能回答：客户是谁、依据什么合同、谁在承办、钱到哪里、材料是否完整。", font=small, fill=pc(MID_GRAY))
    main = [
        (70, 220, 230, 105, "客户与\n利益冲突", "E8F1FA"),
        (370, 220, 230, 105, "合同与\n审批", "E8F1FA"),
        (670, 220, 230, 105, "立案与\n律师分工", "DCEAF7"),
        (970, 220, 260, 105, "办理过程\n任务·文件·开庭", "E8F1FA"),
        (1300, 220, 260, 105, "收款·开票\n结算", "FFF4DB"),
    ]
    for b in main:
        box(*b)
    for i in range(len(main)-1):
        arrow(main[i][0]+main[i][2], 272, main[i+1][0], 272)
    box(1300, 470, 260, 105, "结案与归档", "FFF4DB")
    arrow(1430, 325, 1430, 470)
    box(370, 470, 530, 120, "调查与取证\n项目 → 任务 → 线索 → 审核 → 证据", "EAF7F0", GREEN)
    arrow(635, 325, 635, 470, GREEN)
    arrow(900, 530, 970, 325, GREEN)
    box(970, 470, 260, 120, "用印管理\n独立申请、审批、留痕", "F3EAF8", "7D4E9E")
    arrow(1100, 325, 1100, 470, "7D4E9E")
    box(370, 760, 860, 120, "共同底座：人员、部门、角色、审批权限、数据范围、操作记录", "F2F4F7")
    draw.line((800, 590, 800, 760), fill=pc(MID_GRAY), width=4)
    draw.text((70, 955), "说明：用印是单独模块，可关联合同或案件，但不能被普通案件操作替代。", font=small, fill=pc(RED))
    image.save(FLOW_IMAGE, quality=95)


def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run("律所管理系统｜业务蓝图确认稿"), size=9, bold=True, color=MID_GRAY)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    set_font(footer.add_run("仅用于业务范围与规则确认　｜　第 "), size=9, color=MID_GRAY)
    add_page_number(footer)
    set_font(footer.add_run(" 页"), size=9, color=MID_GRAY)
    return doc


def add_module(doc, name, purpose, flow, owner, controls, decision):
    doc.add_heading(name, level=2)
    add_table(doc, ["这项工作解决什么", "事情怎么流转"], [(purpose, flow)], [3300, 6060], font_size=9.6)
    add_table(doc, ["主要负责人", "必须守住的管理点", "需要合伙人确认"], [(owner, controls, decision)], [1900, 3730, 3730], font_size=9.2)


def build():
    ARCHIVE.mkdir(parents=True, exist_ok=True)
    make_business_map()
    doc = setup_document()

    # Customer-pack opening, adapted for a partner stakeholder packet.
    add_para(doc, "业务确认文件", size=10, bold=True, color=GOLD, after=2)
    add_para(doc, "律所管理系统业务蓝图确认稿", size=28, bold=True, color=NAVY, after=8)
    add_para(doc, "供律所合伙人确认：业务范围、责任分工、审批边界和关键规则", size=13, color=MID_GRAY, after=18)
    add_table(doc, ["文件用途", "确认对象"], [("确定系统应当怎样支持律所经营与案件办理", "管理合伙人、业务合伙人、行政、财务及档案负责人")], [2800, 6560], font_size=10)
    add_table(doc, ["本稿确认什么", "本稿暂不确认什么"], [("谁负责、事情如何流转、何时审批、哪些情况必须阻止", "页面颜色、按钮位置、技术产品选型和服务器方案")], [4680, 4680], font_size=10)
    add_callout(doc, "建议的确认方式", "先确认主流程和责任边界，再逐项勾选第 10 页的拍板清单。凡标注“待确认”的内容，在合伙人签字前都不作为最终制度。", fill=PALE_GOLD, title_color=GOLD)
    add_para(doc, "编制日期：2026年7月17日", size=10, color=MID_GRAY, after=3)
    add_para(doc, "版本：合伙人讨论稿 V1.0", size=10, color=MID_GRAY, after=3)
    add_para(doc, "依据：现有 OA 页面、历史程序结构及业务数据关系的只读梳理", size=9.5, color=MID_GRAY, italic=True, after=0)

    page_break(doc)
    doc.add_heading("一、一页读懂这套系统", level=1)
    add_callout(doc, "一句话目标", "让律所从接触客户到案件归档的全过程可分工、可审批、可追踪、可结算、可追责，同时避免利益冲突、越权操作和关键材料缺失。", fill=PALE_GREEN, title_color=GREEN)
    add_para(doc, "系统不是把纸质表格搬到电脑上，而是把律所的经营纪律落实到每一项工作中：", bold=True)
    for item in (
        "每个客户有明确的管理人，并能进行利益冲突检查。",
        "每个案件都能追溯到客户、合同或经批准的调查线索。",
        "每项任务都有负责人、期限、交接和验收记录。",
        "每次用印都单独申请、审批和留痕，不能混在案件备注里。",
        "每笔费用、开票、回款和结算都能对应到合同或案件。",
        "结案归档前，系统检查费用、文件和审批是否完整。",
    ):
        add_bullet(doc, item)
    doc.add_heading("合伙人最需要看见的五件事", level=2)
    add_table(doc, ["经营问题", "系统应给出的答案"], [
        ("客户归谁管理？", "客户管理人是谁，是否共享，是否长期未跟进。"),
        ("案件由谁负责？", "承办律师、协办律师、助理分别是谁，当前到哪一步。"),
        ("风险在哪里？", "利益冲突、超期、缺合同、缺材料、费用未到、审批未完成。"),
        ("钱是否清楚？", "应收、已收、开票、请款、结算及归档费是否相互对应。"),
        ("事情是否闭环？", "任务是否验收、用印是否回传、案件是否满足归档条件。"),
    ], [2700, 6660], font_size=9.7)

    page_break(doc)
    doc.add_heading("二、整体业务地图", level=1)
    add_para(doc, "下图是业主需要确认的总逻辑。客户、合同、案件是主轴；调查、任务、用印、财务和归档分别承担独立职责。", after=8)
    doc.add_picture(str(FLOW_IMAGE), width=Inches(6.5))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    add_callout(doc, "必须明确", "“关联”不等于“合并”。例如，用印申请可以选择某份合同或某个案件作为依据，但仍由用印模块独立审批、编号、留存文件并记录是否已经盖章。", fill=PALE_RED, title_color=RED)
    doc.add_heading("三条贯穿全所的原则", level=2)
    for item in (
        "一件事情只有一个当前负责人，但可以有协办人员；负责人变更必须留下交接记录。",
        "审批人不能只看结果，还应能看到申请依据、相关文件、历史意见和金额。",
        "系统里的完成不是删除记录，而是进入可查询、可审计的下一状态。",
    ):
        add_bullet(doc, item)

    page_break(doc)
    doc.add_heading("三、谁在系统里负责什么", level=1)
    add_para(doc, "角色名称可以调整，但责任不能空缺。同一人可以兼任多个角色，系统仍要按其当时承担的职责记录操作。")
    add_table(doc, ["角色", "主要责任", "不应越过的边界"], [
        ("管理合伙人", "确定制度、审批层级和重大事项；查看全所经营与风险。", "不以口头决定替代应留痕的审批。"),
        ("业务合伙人／部门负责人", "客户与业务把关、案件负责人安排、重大节点监督。", "只能在授权范围内查看或审批。"),
        ("承办律师", "对案件质量、进度、文件和对外工作负责。", "不能绕过合同、用印和财务流程。"),
        ("协办律师／助理", "完成分派任务、上传成果、反馈进度。", "不能擅自改变案件负责人或审批结果。"),
        ("行政／用印管理员", "核对申请依据、执行盖章、回传盖章文件。", "不能替代业务审批，也不能无申请直接用印。"),
        ("财务人员", "审核开票、回款、付款、费用分配和结算。", "不能由普通办案人员自行拆分多案件回款。"),
        ("档案管理员", "检查归档材料、接收档案、管理借阅。", "归档条件不满足时不能强行通过。"),
        ("系统管理员", "维护账号、部门、角色和基础设置。", "最高权限用于管理和排障，不代表可代替业务负责人作出专业审批。"),
    ], [1900, 4300, 3160], font_size=8.7)
    add_callout(doc, "关于最高权限", "管理员账号应拥有全部菜单、全所数据和全部字段的查看管理能力；但合同、用印、财务等专业决定仍应记录实际审批人的身份和意见。", fill=PALE_GOLD, title_color=GOLD)

    page_break(doc)
    doc.add_heading("四、客户、合同与案件主流程", level=1)
    add_module(doc, "1. 客户与利益冲突", "把潜在客户、正式客户、联系人、跟进和共享关系放在一起管理，并在承接业务前检查冲突。", "新客户登记 → 指定管理人 → 冲突查询 → 跟进／共享 → 签约；无效或长期无人跟进的客户进入回收或公海机制。", "客户管理人、业务合伙人", "客户不能无人负责；共享、转交和恢复都要留痕；冲突查询结果不能被普通人员删除。", "客户多久未跟进进入公海？谁可以认领？冲突命中后由谁决定是否承接？")
    add_module(doc, "2. 合同与审批", "确认委托范围、收费方式、责任部门和风险条款，为立案、开票、收款及用印提供正式依据。", "起草 → 提交 → 审批 → 通过／退回 → 履行 → 归档。审批退回后由申请人修改再提交。", "经办律师、业务合伙人、管理合伙人／指定审批人", "归档合同不能再直接创建新案件；开票前应有有效合同和对外合同编号；合同文件与审批意见应完整保留。", "合同按金额、类型还是部门分级审批？哪些合同必须管理合伙人审批？")
    page_break(doc)
    doc.add_heading("四、客户、合同与案件主流程（续）", level=1)
    add_module(doc, "3. 立案与案件分工", "将委托事项变成可办理、可分工、可统计的案件，并建立当事人、司法机关、律师团队和阶段信息。", "选择客户和可用合同／经批准线索 → 填写案件基本信息 → 登记当事人和司法机关 → 指定负责人和团队 → 开始办理。", "案件负责人、承办律师、部门负责人", "不同客户的案件不能随意合并；同一线索不能重复建案；刑事、行政、顾问、仲裁等案件类型使用各自需要的信息。", "谁有权新建立案？负责人变更由谁批准？无合同时哪些业务允许特批立案？")

    page_break(doc)
    doc.add_heading("五、案件如何被真正办完", level=1)
    add_module(doc, "4. 阶段、任务与交接", "把案件工作拆成有负责人、有期限、有成果、有验收的事项，避免“口头交办、无人跟踪”。", "发起任务 → 接收／拒绝 → 办理与协作 → 提交完成 → 发起人验收；需要换人时先交接，必要时可重新开始。", "任务发起人、任务负责人、案件负责人", "开始时间和截止时间要合理；完成与验收是两个动作；拒绝、撤回、停止、交接和重启都要说明原因。", "是否采用：超过30天未完成自动拒绝；交接后可重新开始；交接满5天未开始自动完成？自动处理前是否先提醒？")
    add_module(doc, "5. 文件、开庭与进度", "集中保存关键文件、法定期限、开庭安排和办理记录，让接手人员能迅速了解案件。", "上传／形成文件 → 归类 → 关联阶段或任务 → 记录开庭与期限 → 更新进度 → 形成结案材料。", "承办律师、协办人员、案件负责人", "重要文件不得只存在个人电脑或聊天记录；文件替换、作废和下载应留痕；关键日期应提醒。", "哪些文件属于必备归档材料？哪些材料只有案件团队可见？期限提醒提前几天？")
    page_break(doc)
    doc.add_heading("五、案件如何被真正办完（续）", level=1)
    add_module(doc, "6. 结案与归档", "在专业工作、费用和材料均达到要求后结束案件，并把档案交给专门人员管理。", "承办律师申请结案／归档 → 系统检查费用和材料 → 负责人／档案审核 → 接收归档 → 后续借阅或查询。", "案件负责人、财务、档案管理员", "案件费用未到账时不能直接提交归档；归档后原则上只读，补录和借阅需单独授权。", "哪些欠款情形允许特批归档？特批人是谁？纸质档案和电子档案是否必须同时齐全？")

    page_break(doc)
    doc.add_heading("六、调查取证与独立用印", level=1)
    add_module(doc, "7. 调查与取证", "在正式诉讼前或案件办理中，对调查项目、线索、审核、证据和公证进行闭环管理。", "关联合同建立调查项目 → 分派调查任务 → 提交线索 → 内部审核 → 必要时客户审核 → 取证／公证 → 证据入库 → 生成新案件或并入已有案件。", "调查负责人、调查人员、审核人、证物管理员", "调查项目应有合同依据；线索已经生成案件后不能重复生成；证物入库、领取、归还和销毁要逐次登记。", "哪些线索需要客户确认？谁可以将线索转成案件？实物证据销毁由几人审批？")
    add_callout(doc, "用印必须单列", "用印不是“案件里加一个按钮”，而是一条独立的风险控制流程。它可以引用案件或合同，但必须拥有自己的申请编号、审批意见、文件版本和盖章结果。", fill=PALE_RED, title_color=RED)
    add_module(doc, "8. 用印管理（独立模块）", "控制公章、合同章等印章的使用，确保每次盖章都有正当依据、明确审批和完整留档。", "提出申请并上传待盖文件 → 选择关联合同／案件（可选） → 业务审批 → 用印管理员核对 → 执行盖章 → 上传盖章版／标记完成；不符合要求则退回或终止。", "申请人、业务审批人、用印管理员", "审批通过前不得盖章；盖章内容与审批版本必须一致；线下紧急用印也应补齐申请和文件。", "不同印章由谁保管？几级审批？紧急用印如何授权？盖章后多久必须回传文件？")

    page_break(doc)
    doc.add_heading("七、财务如何与业务对得上", level=1)
    add_para(doc, "财务模块的目的不是单独记账，而是把合同、案件、费用、发票、回款和律师结算串成同一条可核对的链。")
    add_table(doc, ["业务环节", "发生什么", "管理要求"], [
        ("费用形成", "案件产生律师费、官费、差旅或内部费用。", "必须说明归属合同／案件、费用类型、金额和承担方。"),
        ("开票", "根据合同和收费约定提出开票申请。", "实际开票金额不能为零；合同应有效并有对外合同编号。"),
        ("回款", "银行到账后匹配客户、合同和案件。", "一笔回款涉及多个案件时，由财务统一分配，办案人员不能自行拆分。"),
        ("请款／付款", "对外支付官费、供应商费用或报销款。", "申请、审批、付款凭证和收款方必须一致。"),
        ("结算", "按所内规则计算律师、团队和律所应得金额。", "计算依据、扣减项目、审批和支付状态应可追溯。"),
        ("归档结算", "案件完成后核对费用和归档费。", "未达到财务归档条件的案件应被拦截或走特批。"),
    ], [1700, 3540, 4120], font_size=9.2)
    doc.add_heading("合伙人需要确定的财务口径", level=2)
    for item in (
        "收费方式：固定收费、计时收费、风险代理、阶段收费分别如何确认应收。",
        "审批权限：开票、费用、付款、退款和结算按什么金额分级。",
        "回款顺序：先冲哪一笔应收，能否跨合同或跨案件分配。",
        "结算公式：承办、案源、协办、部门及律所之间如何计算和调整。",
        "欠款归档：哪些情况一律禁止，哪些情况允许书面特批。",
    ):
        add_bullet(doc, item)
    add_callout(doc, "风险提醒", "财务数据一旦进入审批或结算，不应允许普通人员通过通用编辑功能直接改结果；需要调整时，应当走冲销、退回或重新审批。", fill=PALE_GOLD, title_color=GOLD)

    page_break(doc)
    doc.add_heading("八、三个典型场景", level=1)
    add_table(doc, ["场景", "从开始到结束", "合伙人关注点"], [
        ("新客户委托诉讼", "登记客户 → 冲突检查 → 合同审批 → 立案 → 分派律师 → 办案任务／文件／开庭 → 开票回款 → 结案归档。", "承接决定有依据；案件负责人明确；费用、文件和归档闭环。"),
        ("调查线索转案件", "合同下建立调查项目 → 分派任务 → 提交线索 → 内部／客户审核 → 取证 → 生成新案件或并入已有案件。", "线索不能重复建案；证据流转可追踪；合并案件客户一致。"),
        ("合同或诉讼文件用印", "上传待盖文件 → 选择关联事项 → 审批 → 管理员核对版本 → 盖章 → 回传盖章版。", "用印独立编号；审批版本和盖章版本一致；紧急用印可追溯。"),
    ], [1900, 4760, 2700], font_size=9.1)
    doc.add_heading("例外情况也必须有路可走", level=2)
    add_table(doc, ["例外", "建议处理方式"], [
        ("无正式合同但必须先采取保全等紧急措施", "由指定合伙人特批临时立案，并设置补合同期限和提醒。"),
        ("案件需要换负责人", "先发起交接，确认材料和未办事项，再由有权限的人完成变更。"),
        ("客户欠费但依法必须结案", "财务标记欠款，管理合伙人书面特批归档，保留追收责任。"),
        ("节假日紧急用印", "限定申请人和审批人，先电子审批，事后在规定时间内补齐原件和盖章版。"),
    ], [3200, 6160], font_size=9.5)
    add_callout(doc, "设计原则", "系统既要拦住违规操作，也要为真实的紧急业务提供“有权限、有原因、有期限、有追踪”的特批通道。", fill=PALE_GREEN, title_color=GREEN)

    page_break(doc)
    doc.add_heading("九、需要合伙人拍板的事项", level=1)
    add_para(doc, "请在每一项中勾选“同意建议”“需要修改”或“暂缓确定”，并在备注中写明最终口径。以下内容确认后，才能转化为系统规则。", color=DARK_BLUE, bold=True)
    decisions = [
        ("客户公海", "客户连续多久未跟进进入公海；谁可以认领；原管理人能否取回。"),
        ("利益冲突", "命中冲突后谁有权决定承接、拒绝或采取隔离措施。"),
        ("合同审批", "按合同类型、金额、部门设置几级审批及最终审批人。"),
        ("无合同立案", "允许特批的案件类型、批准人和补合同期限。"),
        ("案件负责人", "新建、变更、合并案件分别由谁批准。"),
        ("任务规则", "是否确认30天、交接重启、交接5天未开始自动处理三项规则；是否先提醒。"),
        ("用印规则", "印章种类、保管人、审批层级、紧急用印和盖章文件回传期限。"),
        ("欠款归档", "禁止归档的欠款条件、允许特批的场景和批准人。"),
        ("财务结算", "案源、承办、协办、团队、部门和律所的结算口径。"),
        ("数据权限", "普通律师、部门负责人、合伙人、财务和行政各能查看哪些数据。"),
        ("智能文书助手", "未来智能体可以读取哪些案件材料，生成文书后由谁审定，是否允许直接入卷。"),
    ]
    rows = [(str(i+1), topic, item, "☐ 同意建议\n☐ 需要修改\n☐ 暂缓确定", "") for i, (topic, item) in enumerate(decisions)]
    add_table(doc, ["序号", "事项", "需要确认的口径", "选择", "备注"], rows[:6], [580, 1450, 4010, 2000, 1320], font_size=8.2)
    page_break(doc)
    doc.add_heading("九、需要合伙人拍板的事项（续）", level=1)
    add_table(doc, ["序号", "事项", "需要确认的口径", "选择", "备注"], rows[6:], [580, 1450, 4010, 2000, 1320], font_size=8.2)

    page_break(doc)
    doc.add_heading("十、确认结论与签署", level=1)
    add_callout(doc, "确认范围", "本次签署只确认业务流程、角色责任、审批边界和关键规则。界面样式、字段细节和技术实现将在此基础上另行验收。", fill=PALE_BLUE, title_color=BLUE)
    add_para(doc, "总体意见", bold=True, color=NAVY, after=8)
    add_table(doc, ["请选择", "含义"], [
        ("☐ 原则同意", "同意本稿业务框架，按第九部分已确认的口径继续细化。"),
        ("☐ 修改后同意", "业务框架可用，但应先完成下列修改再进入详细设计。"),
        ("☐ 暂缓", "关键经营制度尚未确定，暂不把相关内容设为系统强制规则。"),
    ], [2500, 6860], font_size=10)
    add_para(doc, "需要修改或补充的事项", bold=True, color=NAVY, after=5)
    for _ in range(4):
        add_para(doc, "________________________________________________________________________________", size=10, color=MID_GRAY, after=7)
    add_table(doc, ["确认人", "职务", "签字", "日期"], [("", "", "", "")], [2300, 2300, 2380, 2380], font_size=10)
    add_para(doc, "后续建议", bold=True, color=NAVY, before=8)
    for item in (
        "由管理合伙人牵头完成第九部分的制度拍板。",
        "由业务、财务、行政和档案负责人分别确认本部门流程。",
        "确认后再制作逐页面、逐按钮的原型和验收清单。",
        "智能文书助手和微信小程序只复用已确认规则，不另起一套业务口径。",
    ):
        add_bullet(doc, item)
    add_para(doc, "备注：本稿根据现有 OA 的只读材料整理；不能从现有证据完全确认的内容均列为待拍板事项。", size=9, color=MID_GRAY, italic=True, before=10, after=0)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
